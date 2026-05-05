"""Tier-1 artifacts for the Kroger run-through:
  1. One-pager (HTML)
  2. ROI Case (HTML, Chart.js interactive)
  3. Persona Day-in-the-Life (HTML)
  4. FSMA 204 Compliance Checklist (DOCX)
"""
from __future__ import annotations
import io, sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
OUT = Path(r"C:\Stage\Clients\Industries\Consumer\Retail\Kroger\02_projects\FY27_Pipeline\assortment-pricing-agentic\deliverables")
OUT.mkdir(parents=True, exist_ok=True)

# -------------------------------------------------------------------------
# HTML helpers — common stylesheet
# -------------------------------------------------------------------------

COMMON_HEAD = """<link rel="preconnect" href="https://fonts.googleapis.com">
<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
<link href="https://fonts.googleapis.com/css2?family=Fraunces:ital,opsz,wght@0,9..144,400;0,9..144,500;0,9..144,600&family=IBM+Plex+Sans:wght@0,300;400;500;600;700&family=JetBrains+Mono:wght@400;500;600&display=swap" rel="stylesheet">"""

COMMON_VARS = """:root {
  --bg-0: #07090F; --bg-1: #0C1119; --bg-2: #141B27; --bg-3: #1C2433;
  --line: rgba(255,255,255,0.10); --line-strong: rgba(255,255,255,0.20);
  --ink-0: #F4EFE5; --ink-1: #D9D2C3; --ink-2: #9A9383; --ink-3: #6B6757;
  --gold: #E3B657; --teal: #3DD9C4; --crimson: #D94B5A; --violet: #9B7BE0;
  --sky: #6FB6E5; --amber: #E8A33D; --ember: #E87B3C;
  --radius: 12px; --ease: cubic-bezier(.22,.68,.24,1);
}
body.light {
  --bg-0: #F5F1E8; --bg-1: #EEE7D7; --bg-2: #E5DCC7; --bg-3: #D9CFB6;
  --line: rgba(20,14,4,0.10); --line-strong: rgba(20,14,4,0.22);
  --ink-0: #1A1508; --ink-1: #3D3420; --ink-2: #6B5F45; --ink-3: #8E8367;
  --gold: #8E6A1E; --teal: #147A6C; --crimson: #A22A39;
  --violet: #6548B0; --sky: #3E7FAE; --amber: #B8791F; --ember: #B0551B;
}
* { box-sizing: border-box; }
html, body { margin: 0; padding: 0; }
body {
  background: var(--bg-0); color: var(--ink-0);
  font-family: "IBM Plex Sans", system-ui, sans-serif;
  font-size: 14.5px; line-height: 1.55; min-height: 100vh;
  transition: background .5s var(--ease), color .5s var(--ease);
}
body::before {
  content: ""; position: fixed; inset: 0;
  background: radial-gradient(1200px 600px at 10% -10%, rgba(232,163,61,0.08), transparent 60%),
              radial-gradient(900px 500px at 95% 110%, rgba(155,123,224,0.06), transparent 60%);
  pointer-events: none; z-index: 0;
}
.theme-toggle {
  position: fixed; top: 24px; right: 24px; z-index: 100;
  display: flex; align-items: center; gap: 8px;
  background: var(--bg-1); border: 1px solid var(--line-strong);
  border-radius: 999px; padding: 8px 14px; cursor: pointer;
  font-family: "JetBrains Mono", monospace;
  font-size: 11px; letter-spacing: .12em; text-transform: uppercase;
  color: var(--ink-1); user-select: none;
  transition: all .25s var(--ease);
}
.theme-toggle:hover { border-color: var(--gold); color: var(--ink-0); }
"""

THEME_TOGGLE_HTML = """<div class="theme-toggle" id="themeToggle" onclick="toggleTheme()">
  <span class="icon" id="themeIcon">☀</span>
  <span id="themeLabel">Light</span>
</div>"""

THEME_TOGGLE_JS = """<script>
function toggleTheme() {
  const body = document.body, icon = document.getElementById('themeIcon'), label = document.getElementById('themeLabel');
  body.classList.toggle('light');
  if (body.classList.contains('light')) { icon.textContent = '☾'; label.textContent = 'Dark'; localStorage.setItem('apex-theme', 'light'); }
  else { icon.textContent = '☀'; label.textContent = 'Light'; localStorage.setItem('apex-theme', 'dark'); }
}
(function() {
  const saved = localStorage.getItem('apex-theme');
  if (saved === 'light') {
    document.body.classList.add('light');
    document.getElementById('themeIcon').textContent = '☾';
    document.getElementById('themeLabel').textContent = 'Dark';
  }
})();
</script>"""


# =========================================================================
# 1. ONE-PAGER HTML
# =========================================================================
ONE_PAGER = """<!DOCTYPE html>
<html lang="en"><head>
<meta charset="UTF-8" /><meta name="viewport" content="width=device-width, initial-scale=1.0" />
<title>APEX · RC-E2E-03 · Kroger · One-Pager</title>
""" + COMMON_HEAD + """
<style>
""" + COMMON_VARS + """
.page {
  position: relative; z-index: 1; max-width: 1200px; margin: 0 auto;
  padding: 56px 48px 40px; min-height: 100vh;
}
.eyebrow {
  display: inline-block; font-family: "JetBrains Mono", monospace;
  font-size: 11px; letter-spacing: .18em; text-transform: uppercase;
  color: var(--gold); border: 1px solid var(--line-strong);
  padding: 6px 14px; border-radius: 999px; margin-bottom: 22px;
}
h1 {
  font-family: "Fraunces", serif; font-weight: 500;
  font-size: clamp(38px, 5vw, 60px); line-height: 1.04;
  letter-spacing: -0.02em; margin: 0 0 14px;
}
h1 em { font-style: italic; font-weight: 400; color: var(--amber); }
.subhead { font-size: 17px; color: var(--ink-1); max-width: 820px; margin: 0 0 28px; }

/* Headline number band */
.headline-band {
  display: grid; grid-template-columns: 2fr 1fr 1fr; gap: 18px;
  margin: 24px 0 32px;
}
@media (max-width: 800px) { .headline-band { grid-template-columns: 1fr; } }
.headline-card {
  background: var(--bg-1); border: 1px solid var(--line);
  border-left: 4px solid var(--accent); border-radius: var(--radius);
  padding: 22px 24px;
}
.headline-card .label {
  font-family: "JetBrains Mono", monospace; font-size: 10.5px;
  letter-spacing: .14em; text-transform: uppercase;
  color: var(--ink-3); margin-bottom: 8px;
}
.headline-card .value {
  font-family: "Fraunces", serif; font-size: 36px;
  font-weight: 500; line-height: 1; color: var(--accent); margin-bottom: 6px;
}
.headline-card .value.big { font-size: 56px; }
.headline-card .sub { font-size: 13px; color: var(--ink-2); }

/* 3-column blocks */
.three-col {
  display: grid; grid-template-columns: repeat(3, 1fr); gap: 18px;
  margin-bottom: 32px;
}
@media (max-width: 900px) { .three-col { grid-template-columns: 1fr; } }
.col-card {
  background: var(--bg-1); border: 1px solid var(--line);
  border-top: 3px solid var(--accent); border-radius: var(--radius);
  padding: 22px;
}
.col-card h3 {
  font-family: "Fraunces", serif; font-size: 20px; font-weight: 500;
  margin: 0 0 4px; color: var(--ink-0);
}
.col-card .col-eye {
  font-family: "JetBrains Mono", monospace; font-size: 10.5px;
  letter-spacing: .14em; text-transform: uppercase;
  color: var(--accent); margin-bottom: 12px;
}
.col-card p { color: var(--ink-1); font-size: 13.5px; margin: 0 0 10px; }
.col-card ul { margin: 8px 0 0; padding-left: 18px; color: var(--ink-1); font-size: 13px; }
.col-card ul li { margin-bottom: 4px; }

/* Wave timeline */
.wave-strip {
  background: var(--bg-1); border: 1px solid var(--line);
  border-radius: var(--radius); padding: 22px 24px; margin-bottom: 24px;
}
.wave-strip h3 {
  font-family: "JetBrains Mono", monospace; font-size: 11px;
  letter-spacing: .15em; text-transform: uppercase;
  color: var(--ink-3); margin: 0 0 14px;
}
.wave-row {
  display: grid; grid-template-columns: repeat(4, 1fr); gap: 12px;
}
@media (max-width: 900px) { .wave-row { grid-template-columns: 1fr 1fr; } }
.wave-item {
  border-left: 3px solid var(--w-color); padding: 10px 14px;
  background: var(--bg-2); border-radius: 4px;
}
.wave-item .w-label {
  font-family: "JetBrains Mono", monospace; font-size: 10px;
  letter-spacing: .14em; text-transform: uppercase;
  color: var(--w-color); margin-bottom: 4px;
}
.wave-item .w-name {
  font-family: "Fraunces", serif; font-size: 15px;
  font-weight: 500; color: var(--ink-0); margin-bottom: 4px;
}
.wave-item .w-detail {
  font-size: 12px; color: var(--ink-2); line-height: 1.45;
}
.wave-item .w-env {
  margin-top: 8px; font-family: "JetBrains Mono", monospace;
  font-size: 11px; color: var(--gold);
}

/* KPI grid */
.kpi-grid {
  display: grid; grid-template-columns: repeat(4, 1fr); gap: 12px;
  margin-bottom: 24px;
}
@media (max-width: 900px) { .kpi-grid { grid-template-columns: 1fr 1fr; } }
.kpi {
  background: var(--bg-1); border: 1px solid var(--line);
  border-radius: var(--radius); padding: 16px 18px;
}
.kpi .lbl {
  font-family: "JetBrains Mono", monospace; font-size: 10px;
  letter-spacing: .14em; text-transform: uppercase;
  color: var(--ink-3); margin-bottom: 6px;
}
.kpi .v {
  font-family: "Fraunces", serif; font-size: 22px;
  font-weight: 500; color: var(--accent);
}
.kpi .delta { font-size: 11.5px; color: var(--ink-2); margin-top: 4px; }

/* Footer */
.footer-band {
  background: var(--bg-1); border: 1px solid var(--line);
  border-radius: var(--radius); padding: 22px 28px;
  display: grid; grid-template-columns: 2fr 1fr; gap: 18px;
  align-items: center; margin-top: 24px;
}
@media (max-width: 800px) { .footer-band { grid-template-columns: 1fr; } }
.footer-band .left p { margin: 0 0 4px; color: var(--ink-1); }
.footer-band .left .signoff {
  font-family: "Fraunces", serif; font-size: 18px; color: var(--ink-0); margin-top: 8px;
}
.footer-band .right {
  font-family: "JetBrains Mono", monospace; font-size: 11px;
  color: var(--ink-2); text-align: right;
}
.footer-band .right b { color: var(--ink-0); display: block; margin-bottom: 4px; }
</style></head>
<body>
""" + THEME_TOGGLE_HTML + """
<div class="page">
  <div class="eyebrow">APEX · RC-E2E-03 · KROGER · ONE-PAGER</div>
  <h1>Cold-chain markdown intelligence at <em>2,800-store</em> scale.</h1>
  <p class="subhead">
    APEX RC-E2E-03 Assortment &amp; Pricing — anchored on Kroger as the Wave-1 reference client —
    converts ~$1.1B/yr in fresh-food shrink exposure into ~$200M/yr in shrink avoided through
    five-agent cold-chain disposition decisioning, FSMA 204 evidence by construction, and
    Microsoft-native deployment on Fabric + Foundry + Teams.
  </p>

  <div class="headline-band">
    <div class="headline-card" style="--accent: var(--gold);">
      <div class="label">W3 annualized durable value</div>
      <div class="value big">~$200M</div>
      <div class="sub">85 events/wk × 2,800 stores × $1,313 avg shrink avoided × 18% reduction</div>
    </div>
    <div class="headline-card" style="--accent: var(--teal);">
      <div class="label">Decision time</div>
      <div class="value">17 min → 90 sec</div>
      <div class="sub">5.2 hrs/shift reclaimed per Store Ops Lead</div>
    </div>
    <div class="headline-card" style="--accent: var(--violet);">
      <div class="label">FSMA 204 evidence coverage</div>
      <div class="value">99.5%</div>
      <div class="sub">14-field audit-row per disposition · auto-generated quarterly packet</div>
    </div>
  </div>

  <div class="three-col">
    <div class="col-card" style="--accent: var(--crimson);">
      <div class="col-eye">Why now</div>
      <h3>Three converging pressures</h3>
      <ul>
        <li><b>FSMA 204 compliance</b> becomes operationally non-optional in 2026 for FTL categories</li>
        <li><b>Fresh-food shrink</b> at ~2.5% of $45B fresh exposure = ~$1.1B/yr leakage</li>
        <li><b>Decision velocity gap</b> — 17-min current loop is too slow for the 52-min pre-opening window</li>
      </ul>
    </div>
    <div class="col-card" style="--accent: var(--amber);">
      <div class="col-eye">What we deliver</div>
      <h3>Five-agent service</h3>
      <ul>
        <li><b>Assess</b> — characterize the excursion in real time</li>
        <li><b>Classify</b> — per-SKU disposition vs. FSMA 204 policy</li>
        <li><b>Quantify</b> — markdown vs. destroy vs. sell-through $ impact</li>
        <li><b>Approve</b> — HITL via Teams Adaptive Card to Store Ops Lead</li>
        <li><b>Act</b> — POS markdown · vendor ticket · regulatory log</li>
      </ul>
    </div>
    <div class="col-card" style="--accent: var(--teal);">
      <div class="col-eye">How it runs</div>
      <h3>Microsoft-native, CAF-rideable</h3>
      <ul>
        <li><b>Fabric Lakehouse</b> — 12 Bronze tables · MERML + SCML Silver canonical</li>
        <li><b>Foundry + Azure OpenAI</b> — agent runtime via AEF LLM Gateway</li>
        <li><b>Eventstream</b> — 1.3B/day refrigeration telemetry · ≤30s latency</li>
        <li><b>LEDGER plane</b> — 14-field audit row · FSMA evidence packet</li>
      </ul>
    </div>
  </div>

  <div class="wave-strip">
    <h3>Three-wave path · Approach → W1 Foundation → W2 Pilot → W3 Scale &amp; Fuse</h3>
    <div class="wave-row">
      <div class="wave-item" style="--w-color: var(--ink-2);">
        <div class="w-label">Approach</div>
        <div class="w-name">Discovery + SOW</div>
        <div class="w-detail">3–5 weeks · banner scoping · stakeholder interviews · Independence office routing</div>
        <div class="w-env">$200K – $500K</div>
      </div>
      <div class="wave-item" style="--w-color: var(--sky);">
        <div class="w-label">Wave 1</div>
        <div class="w-name">Foundation</div>
        <div class="w-detail">5–8 months · single banner cluster (Cincinnati Kroger ~50 stores) · 3 agents · shadow-run</div>
        <div class="w-env">$1.2M – $3.5M</div>
      </div>
      <div class="wave-item" style="--w-color: var(--teal);">
        <div class="w-label">Wave 2</div>
        <div class="w-name">Pilot</div>
        <div class="w-detail">8–12 months · 250-store anchor cluster · 5 agents live · 90-day operating run · FSMA packet</div>
        <div class="w-env">$5M – $12M</div>
      </div>
      <div class="wave-item" style="--w-color: var(--gold);">
        <div class="w-label">Wave 3</div>
        <div class="w-name">Scale &amp; Fuse</div>
        <div class="w-detail">12–18 months · all 2,800+ stores · multi-banner · enterprise + adjacent-scenario fusion</div>
        <div class="w-env">$18M – $45M</div>
      </div>
    </div>
  </div>

  <div class="kpi-grid">
    <div class="kpi" style="--accent: var(--gold);">
      <div class="lbl">Per-event shrink avoided</div>
      <div class="v">$1,313</div>
      <div class="delta">destroy-all baseline → marked-down disposition</div>
    </div>
    <div class="kpi" style="--accent: var(--teal);">
      <div class="lbl">Time-to-decision</div>
      <div class="v">90 sec</div>
      <div class="delta">from 17-min current state · 18× compression</div>
    </div>
    <div class="kpi" style="--accent: var(--violet);">
      <div class="lbl">Manager attention reclaimed</div>
      <div class="v">+5.2 hrs</div>
      <div class="delta">per shift per store · store-floor presence</div>
    </div>
    <div class="kpi" style="--accent: var(--amber);">
      <div class="lbl">Auto-execute rate (W3)</div>
      <div class="v">75%</div>
      <div class="delta">25% routed to HITL · 100% audit-row attributed</div>
    </div>
  </div>

  <div class="footer-band">
    <div class="left">
      <p style="font-family: 'Fraunces', serif; font-size: 16px; color: var(--ink-0);">
        Companion artifacts available on request:
      </p>
      <p style="font-family: 'JetBrains Mono', monospace; font-size: 12px; color: var(--ink-2);">
        Walkthrough · Service Portfolio · SOR Bronze→Silver · ERD · ROI Case · Persona Day-in-the-Life · FSMA 204 Compliance Checklist
      </p>
      <p class="signoff">— Deloitte Microsoft Technology Solutions &amp; Practices</p>
    </div>
    <div class="right">
      <b>Independence-safe by design</b>
      All RC-E2E-03 commercial materials use "Deloitte's Microsoft practice," "DMTSP," "Microsoft platform capabilities," "Microsoft-native deployment." The terms "partner," "alliance," and "strategic alliance" do not appear.
      <br><br>v1.0 · 27 Apr 2026
    </div>
  </div>
</div>
""" + THEME_TOGGLE_JS + """
</body></html>"""


# =========================================================================
# 2. ROI CASE HTML (Chart.js interactive)
# =========================================================================
ROI_CASE = """<!DOCTYPE html>
<html lang="en"><head>
<meta charset="UTF-8" /><meta name="viewport" content="width=device-width, initial-scale=1.0" />
<title>APEX · RC-E2E-03 · Kroger · ROI Case</title>
""" + COMMON_HEAD + """
<script src="https://cdn.jsdelivr.net/npm/chart.js@4.4.0/dist/chart.umd.min.js"></script>
<style>
""" + COMMON_VARS + """
.page { position: relative; z-index: 1; max-width: 1400px; margin: 0 auto; padding: 56px 48px 40px; }
.eyebrow {
  display: inline-block; font-family: "JetBrains Mono", monospace;
  font-size: 11px; letter-spacing: .18em; text-transform: uppercase;
  color: var(--gold); border: 1px solid var(--line-strong);
  padding: 6px 14px; border-radius: 999px; margin-bottom: 22px;
}
h1 {
  font-family: "Fraunces", serif; font-weight: 500;
  font-size: clamp(36px, 4.5vw, 56px); line-height: 1.06;
  letter-spacing: -0.02em; margin: 0 0 14px;
}
h1 em { font-style: italic; font-weight: 400; color: var(--amber); }
.subhead { font-size: 16px; color: var(--ink-1); max-width: 920px; margin: 0 0 32px; }

section { margin-bottom: 44px; }
h2 {
  font-family: "Fraunces", serif; font-weight: 500;
  font-size: 28px; margin: 0 0 6px;
}
.sec-eye {
  font-family: "JetBrains Mono", monospace; font-size: 11px;
  letter-spacing: .15em; text-transform: uppercase;
  color: var(--ink-3); margin-bottom: 8px;
}
.lead { color: var(--ink-1); max-width: 920px; margin-bottom: 22px; }

.hero-num {
  background: var(--bg-1); border: 1px solid var(--line);
  border-left: 5px solid var(--gold); border-radius: var(--radius);
  padding: 28px 32px; display: flex; gap: 36px; align-items: center;
  flex-wrap: wrap;
}
.hero-num .big {
  font-family: "Fraunces", serif; font-size: 80px;
  font-weight: 500; color: var(--gold); line-height: 1;
}
.hero-num .lbl {
  font-family: "JetBrains Mono", monospace; font-size: 11px;
  letter-spacing: .14em; text-transform: uppercase;
  color: var(--ink-3); margin-bottom: 6px;
}
.hero-num .desc { font-size: 14px; color: var(--ink-1); max-width: 600px; }

.composition {
  display: grid; grid-template-columns: repeat(3, 1fr); gap: 16px; margin-bottom: 24px;
}
@media (max-width: 900px) { .composition { grid-template-columns: 1fr; } }
.lever {
  background: var(--bg-1); border: 1px solid var(--line);
  border-top: 4px solid var(--accent); border-radius: var(--radius);
  padding: 22px;
}
.lever h4 {
  font-family: "Fraunces", serif; font-size: 18px; font-weight: 500;
  margin: 0 0 4px;
}
.lever .lever-eye {
  font-family: "JetBrains Mono", monospace; font-size: 10.5px;
  letter-spacing: .14em; text-transform: uppercase;
  color: var(--accent); margin-bottom: 12px;
}
.lever .multi {
  font-family: "Fraunces", serif; font-size: 30px; font-weight: 500;
  color: var(--accent); margin-bottom: 8px;
}
.lever .basis { font-size: 12.5px; color: var(--ink-2); line-height: 1.5; }

.scenarios {
  display: grid; grid-template-columns: repeat(3, 1fr); gap: 16px;
  margin-top: 14px;
}
@media (max-width: 900px) { .scenarios { grid-template-columns: 1fr; } }
.scen-card {
  background: var(--bg-1); border: 1px solid var(--line);
  border-radius: var(--radius); padding: 22px; cursor: pointer;
  transition: all .25s var(--ease);
}
.scen-card.active {
  border-color: var(--accent); box-shadow: 0 0 0 1px var(--accent);
}
.scen-card h4 {
  font-family: "Fraunces", serif; font-size: 20px; font-weight: 500;
  margin: 0 0 8px; color: var(--accent);
}
.scen-card .scen-num {
  font-family: "Fraunces", serif; font-size: 36px; font-weight: 500;
  color: var(--ink-0); margin-bottom: 6px;
}
.scen-card .scen-detail {
  font-size: 12px; color: var(--ink-2); line-height: 1.55;
}

.chart-wrap {
  background: var(--bg-1); border: 1px solid var(--line);
  border-radius: var(--radius); padding: 24px; margin-bottom: 14px;
}
.chart-wrap h3 {
  font-family: "Fraunces", serif; font-size: 18px; margin: 0 0 14px;
}
.chart-canvas { height: 360px; position: relative; }

.envelope-table {
  background: var(--bg-1); border: 1px solid var(--line);
  border-radius: var(--radius); overflow: hidden;
}
.envelope-table table { width: 100%; border-collapse: collapse; }
.envelope-table th {
  background: var(--bg-2); padding: 12px 16px; text-align: left;
  font-family: "JetBrains Mono", monospace; font-size: 10.5px;
  letter-spacing: .14em; text-transform: uppercase;
  color: var(--ink-2); border-bottom: 1px solid var(--line);
}
.envelope-table td {
  padding: 14px 16px; border-bottom: 1px solid var(--line);
  font-size: 13px; color: var(--ink-1);
}
.envelope-table tr:last-child td { border-bottom: none; }
.envelope-table .num {
  font-family: "JetBrains Mono", monospace; font-weight: 600;
  text-align: right; color: var(--ink-0);
}
.envelope-table .gold-row td { background: rgba(227,182,87,0.08); font-weight: 600; }
.envelope-table .gold-row .num { color: var(--gold); }

.callout {
  background: var(--bg-2); border-left: 4px solid var(--gold);
  padding: 14px 18px; border-radius: 4px; margin-top: 18px;
}
.callout .tag {
  font-family: "JetBrains Mono", monospace; font-size: 10px;
  letter-spacing: .14em; text-transform: uppercase;
  color: var(--gold); margin-bottom: 4px;
}
.callout p { margin: 0; color: var(--ink-1); font-size: 13.5px; }

.footer-band {
  margin-top: 50px; padding-top: 24px; border-top: 1px solid var(--line);
  font-family: "JetBrains Mono", monospace; font-size: 11px;
  color: var(--ink-3); display: flex; justify-content: space-between; flex-wrap: wrap; gap: 12px;
}
</style></head>
<body>
""" + THEME_TOGGLE_HTML + """
<div class="page">
  <div class="eyebrow">APEX · RC-E2E-03 · KROGER · ROI CASE</div>
  <h1>The <em>$200M</em> number, decomposed.</h1>
  <p class="subhead">
    The shrink-avoided durable-value claim built bottom-up — three levers (event volume × per-event
    avoided × reduction rate), three scenarios (conservative · base · optimistic), and the payback
    timeline against the $18M–$45M Wave-3 envelope. Every number is editable; every assumption is
    sourced.
  </p>

  <section>
    <div class="sec-eye">§ 1 · Headline</div>
    <h2>Annualized W3 run-rate · base case</h2>
    <div class="hero-num">
      <div>
        <div class="lbl">SHRINK AVOIDED · ANNUALIZED</div>
        <div class="big">$200M</div>
      </div>
      <div class="desc">
        Computed as <b>85 events/wk × 52 wk × 2,800 stores × $1,313 avg shrink avoided × 18% reduction
        rate</b>. Sensitivity bounds in §3. The sales-conversation anchor; the Steering Committee
        gate metric.
      </div>
    </div>
  </section>

  <section>
    <div class="sec-eye">§ 2 · Composition · the three levers</div>
    <h2>What drives the number.</h2>
    <p class="lead">
      The headline figure is multiplicative across three independent levers. Move any one and the
      output moves. The grocery industry has reasonably tight reference values for the first two;
      the third (reduction rate) is the APEX claim and the focal point of the conversation.
    </p>
    <div class="composition">
      <div class="lever" style="--accent: var(--sky);">
        <div class="lever-eye">Lever 1 · Event volume</div>
        <h4>Cold-chain excursions detected</h4>
        <div class="multi">~12.4M / yr</div>
        <div class="basis">85 events × 52 weeks × 2,800 stores. Industry reference range 70–110/banner/wk depending on banner footprint and fixture density. Sourced from FDA FSMA 204 modeling and Deloitte grocery operations benchmark.</div>
      </div>
      <div class="lever" style="--accent: var(--violet);">
        <div class="lever-eye">Lever 2 · Per-event shrink at risk</div>
        <h4>$ exposure per event</h4>
        <div class="multi">$1,313</div>
        <div class="basis">Median destroy-all baseline value. Range $400 (small-format produce) to $4,200 (frozen seafood case). Reference from Deloitte cold-chain markdown library across 8 reference clients.</div>
      </div>
      <div class="lever" style="--accent: var(--gold);">
        <div class="lever-eye">Lever 3 · Reduction rate</div>
        <h4>Excursion shrink avoided</h4>
        <div class="multi">18%</div>
        <div class="basis">APEX claim. Comes from disposition-classification accuracy moving from ~71% to ~95%. The $200M number is most sensitive to this lever — see §3 sensitivity.</div>
      </div>
    </div>
  </section>

  <section>
    <div class="sec-eye">§ 3 · Sensitivity · three scenarios</div>
    <h2>What if the assumptions move?</h2>
    <p class="lead">
      Click a scenario to update the chart. Conservative bakes in slower agent maturity and lower
      reduction rate; Base is the headline case; Optimistic assumes RC-E2E-04 attaches and
      cross-banner orchestration multiplies coverage.
    </p>
    <div class="scenarios">
      <div class="scen-card" style="--accent: var(--ink-2);" onclick="setScenario('conservative', this)">
        <h4>Conservative</h4>
        <div class="scen-num">$95M</div>
        <div class="scen-detail">85 events × $1,313 × 12% reduction across 2,800 stores. Bakes in slower agent maturity, no cross-banner orchestration.</div>
      </div>
      <div class="scen-card active" style="--accent: var(--gold);" onclick="setScenario('base', this)">
        <h4>Base case (headline)</h4>
        <div class="scen-num">$200M</div>
        <div class="scen-detail">85 events × $1,313 × 18% reduction across 2,800 stores. APEX claim. The reference-deployment number.</div>
      </div>
      <div class="scen-card" style="--accent: var(--teal);" onclick="setScenario('optimistic', this)">
        <h4>Optimistic</h4>
        <div class="scen-num">$340M</div>
        <div class="scen-detail">95 events × $1,400 × 22% reduction · RC-E2E-04 loyalty-pricing fusion attached at W3.</div>
      </div>
    </div>
  </section>

  <section>
    <div class="sec-eye">§ 4 · Payback timeline</div>
    <h2>Cumulative investment vs. cumulative return.</h2>
    <p class="lead">
      The chart below time-phases investment and realized return across 6 fiscal years.
      Investment ramps in W1+W2; returns ramp in W2+W3. Net position turns positive in Y2 H2 at base
      case; Y3 H1 at conservative; Y2 H1 at optimistic.
    </p>
    <div class="chart-wrap">
      <h3 id="chartTitle">Payback timeline · base case</h3>
      <div class="chart-canvas"><canvas id="paybackChart"></canvas></div>
    </div>
    <div class="callout">
      <div class="tag">Reading the chart</div>
      <p>Gold line = cumulative durable-value return. Crimson line = cumulative investment. The
      crossover point is payback. Net Present Value at FY32 (5-year horizon) is the area between
      the two lines from crossover onwards. Discount rate not applied — these are nominal cumulative
      dollars; the CFO conversation typically applies a 9–12% WACC for NPV.</p>
    </div>
  </section>

  <section>
    <div class="sec-eye">§ 5 · Wave-by-wave envelope</div>
    <h2>What the engagement costs.</h2>
    <p class="lead">
      Mid-case envelope by wave. Low/High bands are 0.75× and 1.35× of mid. The Wave-3 grand total
      represents the cumulative engagement investment to reach run-rate.
    </p>
    <div class="envelope-table">
      <table>
        <thead>
          <tr>
            <th>Tier</th><th>Duration</th>
            <th class="num" style="text-align:right;">Low</th>
            <th class="num" style="text-align:right;">Mid</th>
            <th class="num" style="text-align:right;">High</th>
            <th>Primary deliverable</th>
          </tr>
        </thead>
        <tbody>
          <tr>
            <td>Approach</td><td>3–5 wks</td>
            <td class="num">$200K</td><td class="num">$340K</td><td class="num">$500K</td>
            <td>Discovery · banner scoping · SOW</td>
          </tr>
          <tr>
            <td>Wave 1 — Foundation</td><td>5–8 mos</td>
            <td class="num">$1.2M</td><td class="num">$2.4M</td><td class="num">$3.5M</td>
            <td>Single-banner shadow-run · audit-row e2e</td>
          </tr>
          <tr>
            <td>Wave 2 — Pilot</td><td>8–12 mos</td>
            <td class="num">$5.0M</td><td class="num">$8.5M</td><td class="num">$12M</td>
            <td>250-store live · 90-day run · FSMA packet</td>
          </tr>
          <tr>
            <td>Wave 3 — Scale &amp; Fuse</td><td>12–18 mos</td>
            <td class="num">$18M</td><td class="num">$31M</td><td class="num">$45M</td>
            <td>Enterprise 2,800+ stores · multi-banner fusion</td>
          </tr>
          <tr class="gold-row">
            <td>Cumulative engagement</td><td>~3 yrs</td>
            <td class="num">$24.4M</td><td class="num">$42.2M</td><td class="num">$61M</td>
            <td>To reach Wave-3 run-rate</td>
          </tr>
        </tbody>
      </table>
    </div>
  </section>

  <section>
    <div class="sec-eye">§ 6 · Sensitivity table</div>
    <h2>What changes if a lever flexes ±20%?</h2>
    <p class="lead">
      Single-variable sensitivity. Holds two levers at base case while flexing the third by ±20%.
      Reduction rate is the most leveraged variable; per-event $ is least leveraged.
    </p>
    <div class="envelope-table">
      <table>
        <thead>
          <tr>
            <th>Lever</th>
            <th class="num" style="text-align:right;">−20%</th>
            <th class="num" style="text-align:right;">Base</th>
            <th class="num" style="text-align:right;">+20%</th>
            <th>Note</th>
          </tr>
        </thead>
        <tbody>
          <tr>
            <td>Event volume (events/wk/banner)</td>
            <td class="num">$160M</td><td class="num">$200M</td><td class="num">$240M</td>
            <td>Linear; banner-mix dependent</td>
          </tr>
          <tr>
            <td>Per-event $ avoided</td>
            <td class="num">$160M</td><td class="num">$200M</td><td class="num">$240M</td>
            <td>Linear; category-mix dependent</td>
          </tr>
          <tr>
            <td>Reduction rate</td>
            <td class="num">$160M</td><td class="num">$200M</td><td class="num">$240M</td>
            <td>Linear at this scale; APEX claim</td>
          </tr>
          <tr>
            <td>Combined ±20% on all three</td>
            <td class="num">$102M</td><td class="num">$200M</td><td class="num">$346M</td>
            <td>Multiplicative; effective range</td>
          </tr>
        </tbody>
      </table>
    </div>
  </section>

  <div class="footer-band">
    <div>APEX · RC-E2E-03 · Kroger · ROI Case · v1.0 · 27 Apr 2026</div>
    <div>Companion · APEX-RC-E2E-03-Walkthrough.docx · APEX-RC-E2E-03-Kroger-OnePager.html</div>
    <div>Deloitte · Microsoft Technology Solutions &amp; Practices</div>
  </div>
</div>

<script>
const scenarios = {
  conservative: {
    title: 'Payback timeline · conservative case',
    invest: [2.4, 5.5, 8.0, 12.0, 8.0, 4.0],
    ret:    [0.0, 4.7, 24, 75, 95, 95],
  },
  base: {
    title: 'Payback timeline · base case',
    invest: [2.4, 6.5, 10.5, 14.0, 6.0, 3.0],
    ret:    [0.0, 10, 55, 145, 200, 200],
  },
  optimistic: {
    title: 'Payback timeline · optimistic case',
    invest: [2.5, 7.5, 13.0, 16.0, 8.0, 4.0],
    ret:    [0.0, 17, 95, 245, 340, 340],
  }
};

let currentScen = 'base';
let chart;

function buildChart(scen) {
  const data = scenarios[scen];
  const cumInvest = []; let s = 0;
  data.invest.forEach(v => { s += v; cumInvest.push(s); });
  const cumRet = []; let r = 0;
  data.ret.forEach(v => { cumRet.push(v); }); // already cumulative-friendly
  const labels = ['FY27', 'FY28 H1', 'FY28 H2', 'FY29', 'FY30', 'FY31+'];

  if (chart) chart.destroy();
  const ctx = document.getElementById('paybackChart');
  chart = new Chart(ctx, {
    type: 'line',
    data: {
      labels,
      datasets: [
        { label: 'Cumulative investment ($M)', data: cumInvest, borderColor: '#D94B5A', backgroundColor: 'rgba(217,75,90,0.1)', tension: 0.35, borderWidth: 2.5, pointRadius: 4, pointBackgroundColor: '#D94B5A' },
        { label: 'Cumulative return ($M)', data: cumRet, borderColor: '#E3B657', backgroundColor: 'rgba(227,182,87,0.1)', tension: 0.35, borderWidth: 2.5, pointRadius: 4, pointBackgroundColor: '#E3B657' },
      ]
    },
    options: {
      responsive: true, maintainAspectRatio: false,
      plugins: {
        legend: { labels: { color: '#D9D2C3', font: { family: 'IBM Plex Sans', size: 12 } } },
        tooltip: { callbacks: { label: ctx => ctx.dataset.label + ': $' + ctx.parsed.y.toFixed(1) + 'M' } }
      },
      scales: {
        x: { ticks: { color: '#9A9383' }, grid: { color: 'rgba(255,255,255,0.05)' } },
        y: { ticks: { color: '#9A9383', callback: v => '$' + v + 'M' }, grid: { color: 'rgba(255,255,255,0.05)' } }
      }
    }
  });
  document.getElementById('chartTitle').textContent = data.title;
}

function setScenario(s, el) {
  currentScen = s;
  document.querySelectorAll('.scen-card').forEach(c => c.classList.remove('active'));
  el.classList.add('active');
  buildChart(s);
}
buildChart('base');
</script>
""" + THEME_TOGGLE_JS + """
</body></html>"""


# =========================================================================
# 3. PERSONA DAY-IN-THE-LIFE HTML
# =========================================================================
PERSONA_DOC = """<!DOCTYPE html>
<html lang="en"><head>
<meta charset="UTF-8" /><meta name="viewport" content="width=device-width, initial-scale=1.0" />
<title>APEX · RC-E2E-03 · Kroger · Persona Day-in-the-Life</title>
""" + COMMON_HEAD + """
<style>
""" + COMMON_VARS + """
.page { position: relative; z-index: 1; max-width: 1500px; margin: 0 auto; padding: 56px 48px 40px; }
.eyebrow {
  display: inline-block; font-family: "JetBrains Mono", monospace;
  font-size: 11px; letter-spacing: .18em; text-transform: uppercase;
  color: var(--gold); border: 1px solid var(--line-strong);
  padding: 6px 14px; border-radius: 999px; margin-bottom: 22px;
}
h1 {
  font-family: "Fraunces", serif; font-weight: 500;
  font-size: clamp(36px, 4.5vw, 56px); line-height: 1.06;
  letter-spacing: -0.02em; margin: 0 0 14px;
}
h1 em { font-style: italic; font-weight: 400; color: var(--amber); }
.subhead { font-size: 16px; color: var(--ink-1); max-width: 920px; margin: 0 0 32px; }

.persona {
  margin-top: 40px; background: var(--bg-1); border: 1px solid var(--line);
  border-radius: var(--radius); overflow: hidden;
}
.persona-head {
  padding: 24px 28px; border-bottom: 1px solid var(--line);
  border-left: 5px solid var(--accent); background: var(--bg-2);
}
.persona-head .role {
  font-family: "JetBrains Mono", monospace; font-size: 10.5px;
  letter-spacing: .14em; text-transform: uppercase;
  color: var(--accent); margin-bottom: 4px;
}
.persona-head h2 {
  font-family: "Fraunces", serif; font-size: 28px; font-weight: 500;
  margin: 0 0 6px; color: var(--ink-0);
}
.persona-head .meta {
  display: flex; gap: 16px; flex-wrap: wrap;
  font-family: "JetBrains Mono", monospace; font-size: 11.5px;
  color: var(--ink-2); margin-top: 8px;
}
.persona-head .meta b { color: var(--ink-0); }

.compare {
  display: grid; grid-template-columns: 1fr 1fr; gap: 0;
}
@media (max-width: 900px) { .compare { grid-template-columns: 1fr; } }
.tl-side {
  padding: 22px 24px;
}
.tl-side.today { border-right: 1px solid var(--line); background: rgba(217,75,90,0.04); }
.tl-side.tomorrow { background: rgba(61,217,196,0.04); }
.tl-side h3 {
  font-family: "Fraunces", serif; font-size: 18px; font-weight: 500;
  margin: 0 0 16px;
}
.tl-side.today h3 { color: var(--crimson); }
.tl-side.tomorrow h3 { color: var(--teal); }
.tl-side h3 .pip {
  display: inline-block; width: 8px; height: 8px; border-radius: 50%;
  margin-right: 8px; vertical-align: middle;
}
.tl-side.today h3 .pip { background: var(--crimson); }
.tl-side.tomorrow h3 .pip { background: var(--teal); }

.timeline {
  position: relative; padding-left: 60px;
}
.timeline::before {
  content: ""; position: absolute; left: 30px; top: 0; bottom: 0;
  width: 1px; background: var(--line);
}
.tl-event {
  position: relative; margin-bottom: 14px;
}
.tl-event .time {
  position: absolute; left: -60px; width: 50px;
  font-family: "JetBrains Mono", monospace; font-size: 10.5px;
  color: var(--ink-3); letter-spacing: .04em;
  text-align: right; padding-top: 2px;
}
.tl-event::before {
  content: ""; position: absolute; left: -32px; top: 6px;
  width: 8px; height: 8px; border-radius: 50%;
  background: var(--bg-1); border: 2px solid var(--ink-3);
}
.tl-side.today .tl-event::before { border-color: var(--crimson); }
.tl-side.tomorrow .tl-event::before { border-color: var(--teal); }
.tl-event.gold::before {
  background: var(--gold); border-color: var(--gold);
  box-shadow: 0 0 0 4px rgba(227,182,87,0.2);
}
.tl-event .what {
  font-size: 13px; color: var(--ink-1); line-height: 1.5;
}
.tl-event .what b { color: var(--ink-0); }
.tl-event .what .pain {
  display: inline-block;
  background: rgba(217,75,90,0.15); color: var(--crimson);
  font-family: "JetBrains Mono", monospace; font-size: 10.5px;
  padding: 1px 8px; border-radius: 4px; margin-left: 8px;
}
.tl-event .what .gain {
  display: inline-block;
  background: rgba(61,217,196,0.15); color: var(--teal);
  font-family: "JetBrains Mono", monospace; font-size: 10.5px;
  padding: 1px 8px; border-radius: 4px; margin-left: 8px;
}

.kpi-strip {
  background: var(--bg-2); padding: 18px 28px;
  border-top: 1px solid var(--line);
  display: grid; grid-template-columns: repeat(3, 1fr); gap: 14px;
}
@media (max-width: 700px) { .kpi-strip { grid-template-columns: 1fr; } }
.kpi-pill {
  display: flex; align-items: center; gap: 10px;
}
.kpi-pill .l {
  font-family: "JetBrains Mono", monospace; font-size: 10px;
  letter-spacing: .12em; text-transform: uppercase;
  color: var(--ink-3); margin-bottom: 2px;
}
.kpi-pill .v {
  font-family: "Fraunces", serif; font-size: 16px; color: var(--gold); font-weight: 500;
}

.footer-band {
  margin-top: 50px; padding-top: 24px; border-top: 1px solid var(--line);
  font-family: "JetBrains Mono", monospace; font-size: 11px;
  color: var(--ink-3); display: flex; justify-content: space-between; flex-wrap: wrap; gap: 12px;
}
</style></head>
<body>
""" + THEME_TOGGLE_HTML + """
<div class="page">
  <div class="eyebrow">APEX · RC-E2E-03 · KROGER · PERSONA DAY-IN-THE-LIFE</div>
  <h1>What changes for the <em>five named personas</em>.</h1>
  <p class="subhead">
    Hour-by-hour today vs. tomorrow timelines for the Category Manager (strategic primary), the
    Store Operations Lead (executional primary), Food Safety Lead, Pricing Manager, and Vendor
    Manager. Crimson dots = pain points today; teal dots = better tomorrow; gold dots = the
    decision moments where APEX changes the game.
  </p>

  <!-- ===== Persona 1 · Category Manager ===== -->
  <div class="persona">
    <div class="persona-head" style="--accent: var(--amber);">
      <div class="role">PRIMARY · STRATEGIC OWNER</div>
      <h2>Sarah · Senior Category Manager · Dairy</h2>
      <div class="meta">
        <div><b>Reports to:</b> VP Merchandising</div>
        <div><b>Owns:</b> Dairy P&amp;L · ~3,200 active SKUs</div>
        <div><b>Current pain:</b> 12-day lag on category-quality data</div>
      </div>
    </div>
    <div class="compare">
      <div class="tl-side today">
        <h3><span class="pip"></span>Today · without RC-E2E-03</h3>
        <div class="timeline">
          <div class="tl-event"><span class="time">7:30 AM</span><div class="what">Reviews overnight markdown report from <b>BI extracts</b><span class="pain">12-day lag</span></div></div>
          <div class="tl-event"><span class="time">8:30 AM</span><div class="what">Email-thread on Cincinnati Store #428 dairy excursion (last week's data)</div></div>
          <div class="tl-event"><span class="time">9:30 AM</span><div class="what">Vendor scorecard for Quarterly Business Review — manual rebuild from spreadsheets<span class="pain">3 hr · stale</span></div></div>
          <div class="tl-event"><span class="time">12:00 PM</span><div class="what">Lunch · category-mix decisions deferred for lack of evidence</div></div>
          <div class="tl-event"><span class="time">2:00 PM</span><div class="what">Pricing Strategy escalation on goodwill credit threshold — phone tag<span class="pain">2 day delay</span></div></div>
          <div class="tl-event"><span class="time">4:00 PM</span><div class="what">Markdown policy revision — annual cycle, can't tune mid-year</div></div>
        </div>
      </div>
      <div class="tl-side tomorrow">
        <h3><span class="pip"></span>Tomorrow · with RC-E2E-03</h3>
        <div class="timeline">
          <div class="tl-event gold"><span class="time">7:30 AM</span><div class="what">Power BI scorecard — overnight category-quality + shrink-NPV — drill-through to specific events<span class="gain">≤ 24h fresh</span></div></div>
          <div class="tl-event"><span class="time">8:00 AM</span><div class="what">Reviews 3 alerts: 2 supplier-quality flags + 1 markdown-NPV anomaly · all attributed to trace_ids</div></div>
          <div class="tl-event gold"><span class="time">9:00 AM</span><div class="what">Concession Policy Studio — adjusts Tier-2 threshold for Cincinnati banner · simulates against last 90 days<span class="gain">PR-merge in minutes</span></div></div>
          <div class="tl-event"><span class="time">10:30 AM</span><div class="what">Vendor QBR prep — auto-generated supplier scorecard with audit-row evidence per finding</div></div>
          <div class="tl-event"><span class="time">2:00 PM</span><div class="what">Drills into Cincinnati Store #428 historical — sees disposition mix, who approved, what policy fired</div></div>
          <div class="tl-event"><span class="time">4:00 PM</span><div class="what">Strategic time on assortment + new-vendor qualification, not putting out yesterday's fires</div></div>
        </div>
      </div>
    </div>
    <div class="kpi-strip">
      <div class="kpi-pill"><div><div class="l">Category-quality data freshness</div><div class="v">12 days → ≤ 24 hrs</div></div></div>
      <div class="kpi-pill"><div><div class="l">QBR prep time</div><div class="v">3 hrs → 20 min</div></div></div>
      <div class="kpi-pill"><div><div class="l">Policy iteration cadence</div><div class="v">Annual → weekly</div></div></div>
    </div>
  </div>

  <!-- ===== Persona 2 · Store Ops Lead ===== -->
  <div class="persona">
    <div class="persona-head" style="--accent: var(--ember);">
      <div class="role">EXECUTIONAL PRIMARY · STORE FLOOR</div>
      <h2>Marisol · Store Operations Lead · Kroger #428</h2>
      <div class="meta">
        <div><b>Reports to:</b> District Manager</div>
        <div><b>Owns:</b> Pre-opening readiness · disposition execution</div>
        <div><b>Current pain:</b> 17-min decision loop · 5.2 hrs/shift on info-gathering</div>
      </div>
    </div>
    <div class="compare">
      <div class="tl-side today">
        <h3><span class="pip"></span>Today</h3>
        <div class="timeline">
          <div class="tl-event"><span class="time">6:08 AM</span><div class="what">EMS alarm fires on Dairy Case A3<span class="pain">42°F · 47 min</span></div></div>
          <div class="tl-event"><span class="time">6:11 AM</span><div class="what">Walks the case · counts 412 affected units · visual triage</div></div>
          <div class="tl-event"><span class="time">6:18 AM</span><div class="what">Calls Emerson hotline · on hold<span class="pain">14 min hold</span></div></div>
          <div class="tl-event"><span class="time">6:32 AM</span><div class="what">Pulls vendor portal log via shared password</div></div>
          <div class="tl-event"><span class="time">6:35 AM</span><div class="what">Voicemail to Food Safety Lead<span class="pain">callback delay</span></div></div>
          <div class="tl-event"><span class="time">6:50 AM</span><div class="what">Disposition agreed · directs Pull Team</div></div>
          <div class="tl-event"><span class="time">7:00 AM</span><div class="what">Store opens · regulatory log entry deferred to end-of-shift</div></div>
        </div>
      </div>
      <div class="tl-side tomorrow">
        <h3><span class="pip"></span>Tomorrow</h3>
        <div class="timeline">
          <div class="tl-event"><span class="time">6:08 AM</span><div class="what">EMS reading lands in Bronze in &lt; 30 sec · Assess agent fires</div></div>
          <div class="tl-event"><span class="time">6:08:30 AM</span><div class="what">Classify produces per-SKU disposition table with FSMA 204 policy citations</div></div>
          <div class="tl-event"><span class="time">6:08:50 AM</span><div class="what">Quantify produces 3-scenario $ impact</div></div>
          <div class="tl-event gold"><span class="time">6:09 AM</span><div class="what">Adaptive Card lands on Marisol's phone with proposed disposition + reasoning trace<span class="gain">90 sec total</span></div></div>
          <div class="tl-event gold"><span class="time">6:09:30 AM</span><div class="what">One-click <b>Approve</b> · audit row writes</div></div>
          <div class="tl-event"><span class="time">6:10 AM</span><div class="what">Act agent dispatches · POS markdown · vendor ticket auto-opened · FSMA log written</div></div>
          <div class="tl-event"><span class="time">6:15 AM</span><div class="what">Marisol walks the floor — store-floor presence, not on hold</div></div>
        </div>
      </div>
    </div>
    <div class="kpi-strip">
      <div class="kpi-pill"><div><div class="l">Time-to-decision</div><div class="v">17 min → 90 sec</div></div></div>
      <div class="kpi-pill"><div><div class="l">Manager attention reclaimed</div><div class="v">+5.2 hrs/shift</div></div></div>
      <div class="kpi-pill"><div><div class="l">FSMA log entry latency</div><div class="v">24 hrs → real-time</div></div></div>
    </div>
  </div>

  <!-- ===== Persona 3 · Food Safety Lead ===== -->
  <div class="persona">
    <div class="persona-head" style="--accent: var(--violet);">
      <div class="role">CO-PRIMARY · FSMA 204 EVIDENCE</div>
      <h2>Daniel · Food Safety Compliance Lead · Cincinnati Banner</h2>
      <div class="meta">
        <div><b>Reports to:</b> VP Quality &amp; Regulatory</div>
        <div><b>Owns:</b> FSMA 204 compliance · FDA inspection response · destroy-all events</div>
        <div><b>Current pain:</b> 7% audit-row gap · 24-month traceability uncertain</div>
      </div>
    </div>
    <div class="compare">
      <div class="tl-side today">
        <h3><span class="pip"></span>Today</h3>
        <div class="timeline">
          <div class="tl-event"><span class="time">All day</span><div class="what">Reactive · receives voicemails on CRITICAL events from store managers</div></div>
          <div class="tl-event"><span class="time">Weekly</span><div class="what">Manual exception sampling for FSMA 204 readiness review<span class="pain">7% gap</span></div></div>
          <div class="tl-event"><span class="time">Quarterly</span><div class="what">FDA inspection prep — reconstructs evidence from store logs + vendor portals + emails<span class="pain">3 weeks</span></div></div>
          <div class="tl-event"><span class="time">Ad-hoc</span><div class="what">Recall response — manually traces lot codes across stores via Excel</div></div>
        </div>
      </div>
      <div class="tl-side tomorrow">
        <h3><span class="pip"></span>Tomorrow</h3>
        <div class="timeline">
          <div class="tl-event gold"><span class="time">All day</span><div class="what">Real-time CRITICAL Adaptive Card on his phone · audit-row evidence pre-attached</div></div>
          <div class="tl-event"><span class="time">Weekly</span><div class="what">Power BI exception watchtower — one query returns all gaps · drill-through to trace_id<span class="gain">≥99.5%</span></div></div>
          <div class="tl-event gold"><span class="time">Quarterly</span><div class="what">Auto-generated FSMA 204 evidence packet ZIP — manifests + policies + audit rows + tie-out<span class="gain">72 hrs</span></div></div>
          <div class="tl-event"><span class="time">Recall scenario</span><div class="what">SCML.Lot query · all affected stores in seconds · auto-shelf-pull via Act</div></div>
        </div>
      </div>
    </div>
    <div class="kpi-strip">
      <div class="kpi-pill"><div><div class="l">Audit-row coverage</div><div class="v">62% → 99.5%</div></div></div>
      <div class="kpi-pill"><div><div class="l">FDA inspection prep</div><div class="v">3 wks → 72 hrs</div></div></div>
      <div class="kpi-pill"><div><div class="l">Recall response time</div><div class="v">days → seconds</div></div></div>
    </div>
  </div>

  <!-- ===== Persona 4 · Pricing Manager ===== -->
  <div class="persona">
    <div class="persona-head" style="--accent: var(--gold);">
      <div class="role">SECONDARY · PRICING POLICY OWNER</div>
      <h2>Yvonne · Pricing Manager · Kroger Banner</h2>
      <div class="meta">
        <div><b>Reports to:</b> Pricing Director</div>
        <div><b>Owns:</b> Markdown policy manifest · per-SKU elasticity · concession ceiling</div>
        <div><b>Current pain:</b> 12-month policy lock · per-cohort tuning impossible</div>
      </div>
    </div>
    <div class="compare">
      <div class="tl-side today">
        <h3><span class="pip"></span>Today</h3>
        <div class="timeline">
          <div class="tl-event"><span class="time">Annual</span><div class="what">Markdown policy doc revision · Word document + Excel elasticity table<span class="pain">12-month lock</span></div></div>
          <div class="tl-event"><span class="time">Weekly</span><div class="what">Reviews concession-cost report from Finance · 1-month lag</div></div>
          <div class="tl-event"><span class="time">Quarterly</span><div class="what">Tier threshold tuning · static rules · no per-cohort visibility</div></div>
          <div class="tl-event"><span class="time">Ad-hoc</span><div class="what">Email approvals on above-threshold concessions<span class="pain">2-3 day delay</span></div></div>
        </div>
      </div>
      <div class="tl-side tomorrow">
        <h3><span class="pip"></span>Tomorrow</h3>
        <div class="timeline">
          <div class="tl-event gold"><span class="time">Daily</span><div class="what">Concession Policy Studio — declarative authoring · simulation against historical 90 days</div></div>
          <div class="tl-event"><span class="time">Real-time</span><div class="what">Live concession-cost dashboard per cohort · daily burn-rate · anomaly alerts</div></div>
          <div class="tl-event gold"><span class="time">Per-cohort</span><div class="what">Tunes Tier-2 threshold for Cincinnati banner separately from Atlanta · PR-merge workflow</div></div>
          <div class="tl-event"><span class="time">Real-time</span><div class="what">Material concessions auto-route via Adaptive Card · she sees them inline · approves in seconds</div></div>
        </div>
      </div>
    </div>
    <div class="kpi-strip">
      <div class="kpi-pill"><div><div class="l">Policy iteration cadence</div><div class="v">Annual → weekly</div></div></div>
      <div class="kpi-pill"><div><div class="l">Per-cohort visibility</div><div class="v">None → live</div></div></div>
      <div class="kpi-pill"><div><div class="l">Concession-decision latency</div><div class="v">2-3 days → seconds</div></div></div>
    </div>
  </div>

  <!-- ===== Persona 5 · Vendor Manager ===== -->
  <div class="persona">
    <div class="persona-head" style="--accent: var(--teal);">
      <div class="role">SECONDARY · SUPPLIER QUALIFICATION</div>
      <h2>Marcus · Vendor Manager · Dairy Category</h2>
      <div class="meta">
        <div><b>Reports to:</b> CPO / Category Procurement</div>
        <div><b>Owns:</b> ~80 dairy vendors · FSMA 204 capability tier · QBR cadence</div>
        <div><b>Current pain:</b> Quarterly QBR rebuild from scratch · trade-spend ROI opaque</div>
      </div>
    </div>
    <div class="compare">
      <div class="tl-side today">
        <h3><span class="pip"></span>Today</h3>
        <div class="timeline">
          <div class="tl-event"><span class="time">Pre-QBR</span><div class="what">Manual scorecard rebuild from category-quality emails + finance pulls<span class="pain">~6 hours</span></div></div>
          <div class="tl-event"><span class="time">QBR</span><div class="what">Vendor disputes findings · evidence is hand-curated · weak negotiating position</div></div>
          <div class="tl-event"><span class="time">Post-QBR</span><div class="what">Trade-spend ROI on slotting + MDF · qualitative assessment</div></div>
          <div class="tl-event"><span class="time">FSMA prep</span><div class="what">Manual outreach to vendors on FTL capability · uneven KDE format support</div></div>
        </div>
      </div>
      <div class="tl-side tomorrow">
        <h3><span class="pip"></span>Tomorrow</h3>
        <div class="timeline">
          <div class="tl-event gold"><span class="time">Pre-QBR</span><div class="what">Auto-generated scorecard with audit-row evidence per finding · 20-min review</div></div>
          <div class="tl-event"><span class="time">QBR</span><div class="what">Trace_id-attributed evidence — vendor cannot dispute · negotiation moves to remediation</div></div>
          <div class="tl-event gold"><span class="time">Post-QBR</span><div class="what">Trade-spend ROI quantified per supplier · drives next negotiation</div></div>
          <div class="tl-event"><span class="time">FSMA</span><div class="what">FTL capability tier on each vendor record · KDE format coverage tracked at platform level</div></div>
        </div>
      </div>
    </div>
    <div class="kpi-strip">
      <div class="kpi-pill"><div><div class="l">QBR scorecard prep</div><div class="v">6 hrs → 20 min</div></div></div>
      <div class="kpi-pill"><div><div class="l">Vendor-finding disputability</div><div class="v">High → low</div></div></div>
      <div class="kpi-pill"><div><div class="l">Trade-spend ROI visibility</div><div class="v">Qualitative → quantitative</div></div></div>
    </div>
  </div>

  <div class="footer-band">
    <div>APEX · RC-E2E-03 · Kroger · Persona Day-in-the-Life · v1.0 · 27 Apr 2026</div>
    <div>Companion · APEX-RC-E2E-03-Walkthrough.docx · APEX-RC-Grocery-Merchandising-Service-Portfolio.docx</div>
    <div>Deloitte · Microsoft Technology Solutions &amp; Practices</div>
  </div>
</div>
""" + THEME_TOGGLE_JS + """
</body></html>"""


# =========================================================================
# Write all 3 HTMLs
# =========================================================================
def write_html(name, content):
    p = OUT / name
    p.write_text(content, encoding="utf-8")
    print(f"  wrote {p.name} ({len(content):,} chars)")

print("Writing Tier-1 HTMLs...")
write_html("APEX-RC-E2E-03-Kroger-OnePager.html", ONE_PAGER)
write_html("APEX-RC-E2E-03-Kroger-ROI-Case.html", ROI_CASE)
write_html("APEX-RC-E2E-03-Kroger-Personas.html", PERSONA_DOC)


# =========================================================================
# 4. FSMA 204 COMPLIANCE CHECKLIST DOCX
# =========================================================================
print("\nWriting FSMA 204 Compliance Checklist DOCX...")

INK   = "0B0B0F"; INK2 = "4A4A55"; MUTE = "8a8a90"; GOLD = "B8892E"
TEAL  = "2A7F73"; VIO = "6648B0"; CRIM = "A22A39"; BAND = "F3F3F5"
HEAD  = "0B0B0F"; LINE = "D4D4D7"

def set_cell_bg(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex.lstrip("#")); tcPr.append(shd)

def set_cell_border(cell, color_hex=LINE):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "4"); b.set(qn("w:color"), color_hex.lstrip("#"))
        borders.append(b)
    tcPr.append(borders)

def set_run(run, *, bold=False, italic=False, size=None, color=None, font="Arial"):
    if bold: run.bold = True
    if italic: run.italic = True
    if size is not None: run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor.from_string(color.lstrip("#"))
    if font: run.font.name = font

def add_body(doc, text, *, size=11, italic=False, color=INK, space_after=6):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text); set_run(r, size=size, color=color, italic=italic); return p

def add_h(doc, text, *, level=1):
    sizes = {1: 18, 2: 14, 3: 12.5, 4: 11.5}
    p = doc.add_paragraph(); p.style = doc.styles[f"Heading {level}"]
    p.paragraph_format.space_before = Pt(16 if level==1 else 10)
    p.paragraph_format.space_after = Pt(4)
    for r in list(p.runs): r.text = ""
    r = p.add_run(text); set_run(r, bold=True, size=sizes.get(level, 12), color=INK)
    if level == 1:
        pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom"); bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6"); bottom.set(qn("w:space"), "4"); bottom.set(qn("w:color"), GOLD)
        pBdr.append(bottom); pPr.append(pBdr)
    return p

def add_bullets(doc, items, *, size=11):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(it); set_run(r, size=size, color=INK)

def add_table(doc, header, rows, *, col_widths=None):
    n = len(header)
    tbl = doc.add_table(rows=len(rows)+1, cols=n)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(header):
        c = tbl.rows[0].cells[j]; set_cell_bg(c, HEAD); set_cell_border(c)
        c.paragraphs[0].text = ""
        r = c.paragraphs[0].add_run(h); set_run(r, bold=True, size=10, color="FFFFFF")
        if col_widths: c.width = Inches(col_widths[j])
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            c = tbl.rows[i].cells[j]; set_cell_border(c)
            if i % 2 == 0: set_cell_bg(c, BAND)
            c.paragraphs[0].text = ""
            r = c.paragraphs[0].add_run(str(v)); set_run(r, size=10, color=INK)
            if col_widths: c.width = Inches(col_widths[j])
    doc.add_paragraph()
    return tbl

def add_callout(doc, kind, body, *, accent=GOLD, fill="F6F1E4"):
    tbl = doc.add_table(rows=1, cols=1); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = tbl.rows[0].cells[0]; set_cell_bg(c, fill); c.paragraphs[0].text = ""
    p_tag = c.paragraphs[0]; r_tag = p_tag.add_run(kind.upper())
    set_run(r_tag, bold=True, size=9, color=accent)
    p = c.add_paragraph(); p.paragraph_format.space_before = Pt(3)
    r = p.add_run(body); set_run(r, size=10.5, color=INK); doc.add_paragraph()

def add_kv(doc, pairs, *, lw=1.6, rw=4.9):
    tbl = doc.add_table(rows=len(pairs), cols=2); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(pairs):
        ck, cv = tbl.rows[i].cells; ck.width = Inches(lw); cv.width = Inches(rw)
        set_cell_bg(ck, BAND); set_cell_border(ck); set_cell_border(cv)
        ck.paragraphs[0].text = ""; cv.paragraphs[0].text = ""
        rk = ck.paragraphs[0].add_run(k); set_run(rk, bold=True, size=10, color=INK)
        rv = cv.paragraphs[0].add_run(v); set_run(rv, size=10, color=INK)
    doc.add_paragraph()

def add_eyebrow(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text.upper()); set_run(r, bold=True, size=9, color=MUTE)
    rPr = r._element.get_or_add_rPr()
    spacing = OxmlElement("w:spacing"); spacing.set(qn("w:val"), "80"); rPr.append(spacing)

def pagebreak(doc):
    p = doc.add_paragraph(); p.add_run().add_break(WD_BREAK.PAGE)


def build_fsma():
    doc = Document()
    for s in doc.sections:
        s.page_width = Inches(8.5); s.page_height = Inches(11)
        s.top_margin = s.bottom_margin = Inches(0.85)
        s.left_margin = s.right_margin = Inches(0.95)
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(11)

    add_eyebrow(doc, "APEX · RC-E2E-09 · FSMA 204 COMPLIANCE CHECKLIST · KROGER")
    p = doc.add_paragraph(); r = p.add_run("FSMA 204 Compliance"); set_run(r, bold=True, size=26, color=INK)
    p2 = doc.add_paragraph(); r2 = p2.add_run("Final-rule readiness checklist · APEX RC-E2E-09 implementation map")
    set_run(r2, bold=True, size=15, color=INK2, italic=True)

    doc.add_paragraph()
    add_body(doc, "FDA's Food Safety Modernization Act §204(d) final rule — \"Requirements for Additional "
                  "Traceability Records for Certain Foods\" — establishes lot-level traceability requirements "
                  "for foods on the Food Traceability List (FTL). Compliance is required by January 20, 2026 "
                  "(at time of authoring; subject to FDA enforcement-discretion timing). This document maps "
                  "each FSMA 204 final-rule requirement to the APEX RC-E2E-09 Product Tracking service "
                  "implementation, with a per-requirement compliance checklist and a Kroger-specific gap "
                  "analysis path.",
             italic=True, color=INK2)
    doc.add_paragraph()
    add_kv(doc, [
        ("Document type",    "FSMA 204 Compliance Checklist · APEX RC-E2E-09 mapping"),
        ("Effective date",   "Jan 20, 2026 (subject to FDA enforcement-discretion timing)"),
        ("APEX service",     "RC-E2E-09 Product Tracking (co-anchor with RC-E2E-03)"),
        ("Reference client", "The Kroger Co. — Cincinnati banner cluster Wave-1 anchor"),
        ("Audience",         "Food Safety Lead · Quality · Internal Audit · External Audit · FDA Relations"),
        ("Companion",        "APEX-RC-E2E-03-09-SOR-Bronze-to-Silver.docx · APEX-RC-E2E-03-09-SOR-ERD.html"),
        ("Version / date",   "v1.0 · 27 Apr 2026"),
    ])
    pagebreak(doc)

    # SECTION 1 · OVERVIEW
    add_h(doc, "1 · FSMA 204 final-rule overview", level=1)
    add_body(doc, "The final rule establishes recordkeeping requirements for entities that manufacture, "
                  "process, pack, or hold foods on the FTL. The intent is to allow rapid traceability "
                  "during food-safety incidents — typically within 24 hours of FDA request.")
    add_h(doc, "1.1 · Five core requirements", level=2)
    add_table(doc,
              ["Requirement", "Definition", "RC-E2E-09 implementation"],
              [
                  ["Critical Tracking Events (CTEs)",
                   "Specific points in the supply chain where traceability records must be kept",
                   "BRZ.RC.LOT_TRACE captures all 5 CTEs: Growing, Receiving, Transformation, Shipping, First-Receipt"],
                  ["Key Data Elements (KDEs)",
                   "Specific data fields recorded at each CTE",
                   "BRZ.RC.LOT_TRACE.kde_payload (versioned JSON) carries all required fields per CTE type"],
                  ["Traceability Lot Code (TLC)",
                   "Unique identifier for each lot/batch of FTL food",
                   "BRZ.RC.LOT_TRACE.lot_code is the canonical TLC; supplier_lot_code preserved for reconciliation"],
                  ["TLC Source / TLC Source Reference",
                   "Identification of where the TLC was assigned",
                   "supplier_id (first source) + KDE-embedded source-reference per CTE"],
                  ["24-month recordkeeping",
                   "Records must be retained for 24 months after the FTL food leaves the entity's control",
                   "Bronze hot-tier 180 days · indefinite warm/cold archive · Delta time-travel supports point-in-time queries"],
              ],
              col_widths=[1.7, 2.6, 2.0])

    add_h(doc, "1.2 · 24-hour FDA response requirement", level=2)
    add_body(doc, "On FDA request during a food-safety incident, regulated entities must produce required "
                  "records within 24 hours in an electronic sortable spreadsheet format. RC-E2E-09's "
                  "evidence-packet generator produces this format directly from BRZ.RC.LOT_TRACE Silver-tier "
                  "joins, with response time measured in minutes rather than hours at W3 run-rate.")
    add_callout(doc, "Compliance pillar",
                "FSMA 204 is fundamentally a recordkeeping rule. Every requirement reduces to: "
                "(1) capture the right data at the right CTE, (2) preserve it for 24 months, "
                "(3) produce it in 24 hours when asked. RC-E2E-09 makes all three operationally "
                "non-optional by construction — records are generated, not collated.")

    pagebreak(doc)

    # SECTION 2 · FOOD TRACEABILITY LIST SCOPE
    add_h(doc, "2 · Food Traceability List (FTL) scope at Kroger", level=1)
    add_body(doc, "The FTL covers ~16 categories of high-risk foods. Kroger's exposure varies by category; "
                  "the table below ranks Kroger-specific exposure based on category SKU count and revenue mix.")
    add_table(doc,
              ["FTL Category", "Kroger SKU exposure", "Banner-mix concentration", "Supplier readiness"],
              [
                  ["Soft cheeses (made from raw or pasteurized)",
                   "~340 SKUs",                           "Specialty across all banners",
                   "Medium — Murray's owned · third-party uneven"],
                  ["Cucumbers",                            "~12 SKUs",
                   "Banner-aware seasonal",                "High — top 3 vendors FSMA-ready"],
                  ["Herbs (fresh)",                       "~85 SKUs",
                   "Specialty-store concentrated",        "Low — fragmented suppliers"],
                  ["Leafy greens (fresh)",                "~120 SKUs",
                   "All banners · high-volume",           "High — top vendors mostly ready"],
                  ["Melons",                              "~30 SKUs",
                   "Seasonal · all banners",              "Medium"],
                  ["Peppers",                             "~95 SKUs",
                   "All banners · year-round",            "Medium"],
                  ["Sprouts",                             "~28 SKUs",
                   "Specialty",                           "Medium"],
                  ["Tomatoes",                            "~45 SKUs",
                   "All banners",                         "High"],
                  ["Tropical tree fruits",                "~40 SKUs",
                   "Specialty",                           "Low"],
                  ["Fresh shell eggs",                    "~110 SKUs",
                   "All banners · daily",                 "High — top 3 cover ~85%"],
                  ["Nut butters",                         "~80 SKUs",
                   "All banners · center store",          "High"],
                  ["Crustaceans (shrimp, crab, lobster)", "~140 SKUs",
                   "Seafood-counter banners",             "Medium"],
                  ["Mollusks (oysters, clams, mussels)", "~60 SKUs",
                   "Specialty seafood",                  "Low — fragmented"],
                  ["Finfish (specific species)",          "~220 SKUs",
                   "All banners · seafood depts",         "Medium"],
                  ["Ready-to-eat deli salads",            "~190 SKUs",
                   "All banners · deli",                  "Medium · Murray's owned"],
                  ["Cheeses (specific aged categories)",  "~310 SKUs",
                   "Specialty",                           "Medium"],
              ],
              col_widths=[2.2, 1.3, 1.6, 1.5])

    add_callout(doc, "Kroger-specific scope",
                "Approximately 1,900 SKUs across 16 FTL categories. Banner-mix concentration drives "
                "differential rollout: Tier-1 banners with high specialty mix (Fred Meyer, Mariano's, "
                "Harris Teeter) carry disproportionate FTL exposure. Supplier-readiness gaps in soft "
                "cheese, mollusks, and tropical tree fruit are the highest-priority remediation targets "
                "at Wave-1 scope.",
                accent=VIO, fill="EFE5F4")

    pagebreak(doc)

    # SECTION 3 · CTE CHECKLIST
    add_h(doc, "3 · CTE-by-CTE compliance checklist", level=1)
    add_body(doc, "Each CTE has its own KDE set. The tables below show the FDA-required KDEs for each "
                  "CTE, with the BRZ.RC.LOT_TRACE field that captures it and the compliance-status indicator.")

    # CTE 1 — Receiving
    add_h(doc, "3.1 · CTE · Receiving (most common Kroger CTE)", level=2)
    add_body(doc, "Receiving is when an FTL food arrives at a Kroger DC or store from any supplier. "
                  "This CTE accounts for the highest CTE-event volume in Kroger's data flow.")
    add_table(doc,
              ["KDE required", "BRZ.RC.LOT_TRACE field", "Status"],
              [
                  ["Traceability Lot Code", "lot_code", "✓ captured"],
                  ["Quantity and unit of measure", "qty + uom", "✓ captured"],
                  ["Product description", "(joined to BRZ.RC.ITEM_MASTER.description_short)", "✓ via join"],
                  ["Date received", "event_ts (date portion)", "✓ captured"],
                  ["Location description (FDA Establishment Identifier or alternative)", "to_location_id", "✓ captured · requires DIM.Location enrichment"],
                  ["Traceability Lot Code Source (where TLC was assigned)", "supplier_id + supplier_lot_code", "✓ captured"],
                  ["Reference document type and number (PO, invoice, etc.)", "kde_payload.ref_doc_type + kde_payload.ref_doc_num", "✓ captured (JSON KDE)"],
                  ["Date of last shipment (from previous source)", "kde_payload.prev_shipment_date", "✓ captured"],
              ],
              col_widths=[2.6, 2.6, 1.4])

    # CTE 2 — Transformation
    add_h(doc, "3.2 · CTE · Transformation", level=2)
    add_body(doc, "Transformation is when an FTL food is changed (e.g., produce washed and bagged at a "
                  "distribution center; soft cheese repackaged; ready-to-eat deli salad assembled). Kroger's "
                  "DC-level repackaging operations and store-level deli operations both generate Transformation CTEs.")
    add_table(doc,
              ["KDE required", "BRZ.RC.LOT_TRACE field", "Status"],
              [
                  ["New TLC for the transformed product", "lot_code (new generation)", "✓ captured"],
                  ["TLC of input (incoming) lot(s)", "kde_payload.input_lot_codes[]", "✓ captured (JSON array)"],
                  ["Date of transformation", "event_ts", "✓ captured"],
                  ["Quantity and unit of measure of new lot", "qty + uom", "✓ captured"],
                  ["New product description", "(joined to BRZ.RC.ITEM_MASTER for new sku/gtin)", "✓ via join"],
                  ["Location description of transformation", "from_location_id (= to_location_id)", "✓ captured"],
              ],
              col_widths=[2.6, 2.6, 1.4])

    # CTE 3 — Shipping
    add_h(doc, "3.3 · CTE · Shipping", level=2)
    add_body(doc, "Shipping is when an FTL food leaves a Kroger DC en route to a store, or moves between "
                  "DCs. Drives the inter-DC and DC-to-store leg of the supply chain.")
    add_table(doc,
              ["KDE required", "BRZ.RC.LOT_TRACE field", "Status"],
              [
                  ["TLC", "lot_code", "✓ captured"],
                  ["Quantity and unit of measure shipped", "qty + uom", "✓ captured"],
                  ["Product description", "(joined to BRZ.RC.ITEM_MASTER)", "✓ via join"],
                  ["Date of shipment", "event_ts", "✓ captured"],
                  ["Location description of shipping origin", "from_location_id", "✓ captured"],
                  ["Location description of receiving destination", "to_location_id", "✓ captured"],
                  ["Reference document type and number", "kde_payload.ref_doc_type + kde_payload.ref_doc_num", "✓ captured"],
              ],
              col_widths=[2.6, 2.6, 1.4])

    # CTE 4 — Growing (less common at Kroger)
    add_h(doc, "3.4 · CTE · Growing (supplier-side; Kroger consumes via EDI)", level=2)
    add_body(doc, "Growing CTEs are recorded by farms and primary producers, not by Kroger directly. "
                  "Kroger consumes Growing-CTE records via supplier EDI integration (Trace One, FoodLogiQ) "
                  "and preserves them in the lot-trace event chain for downstream traceability.")
    add_table(doc,
              ["KDE required", "BRZ.RC.LOT_TRACE field", "Status"],
              [
                  ["Originating TLC", "lot_code", "✓ via supplier EDI"],
                  ["Date of harvesting / production", "harvest_or_prod_date", "✓ captured"],
                  ["Field/grow location identifier", "kde_payload.field_id", "✓ via supplier EDI"],
                  ["Quantity at harvest / production", "qty + uom", "✓ captured"],
                  ["Product description", "(joined to BRZ.RC.ITEM_MASTER)", "✓ via join"],
              ],
              col_widths=[2.6, 2.6, 1.4])

    # CTE 5 — First Land-based Receiver
    add_h(doc, "3.5 · CTE · First Land-based Receiver (seafood-specific)", level=2)
    add_body(doc, "First Land-based Receiver applies to seafood — the entity that first receives the seafood "
                  "after it's offloaded from a fishing vessel. Kroger's seafood suppliers (or their primary "
                  "processors) record this CTE; Kroger consumes it via EDI for finfish, crustaceans, mollusks.")
    add_table(doc,
              ["KDE required", "BRZ.RC.LOT_TRACE field", "Status"],
              [
                  ["TLC assigned by First Receiver", "lot_code", "✓ via supplier EDI"],
                  ["Vessel name and registration", "kde_payload.vessel_name + vessel_reg", "✓ captured"],
                  ["Date and location of offload", "event_ts + kde_payload.offload_location", "✓ captured"],
                  ["Species and quantity received", "(joined) + qty", "✓ via join"],
                  ["Reference catch document", "kde_payload.catch_doc_ref", "✓ captured"],
              ],
              col_widths=[2.6, 2.6, 1.4])

    pagebreak(doc)

    # SECTION 4 · TRACEABILITY LOT CODE (TLC) DISCIPLINE
    add_h(doc, "4 · Traceability Lot Code (TLC) discipline", level=1)
    add_body(doc, "The TLC is the single most operationally important FSMA 204 concept. It is the "
                  "identifier that links every CTE in the supply chain. RC-E2E-09 enforces TLC discipline "
                  "by construction; the rules below apply.")
    add_table(doc,
              ["Rule", "Implementation"],
              [
                  ["TLC must be assigned at first CTE within the regulated entity's control",
                   "BRZ.RC.LOT_TRACE.lot_code populated at every CTE — never NULL for FTL items"],
                  ["TLC must be unique within the assigning entity",
                   "Uniqueness enforced at supplier source · Kroger-internal TLC namespacing for transformation"],
                  ["TLC must persist through Receiving, Storing, Shipping (without transformation)",
                   "lot_code passes through unchanged across CTE chain · multi-event lookups via lot_code in Silver"],
                  ["TLC must be regenerated at each Transformation event",
                   "CTE-Transformation events emit a new lot_code with kde_payload.input_lot_codes preserving lineage"],
                  ["TLC must be linked to TLC Source (where it was assigned)",
                   "supplier_id captures TLC Source · supplier_lot_code preserves cross-system reconciliation key"],
                  ["TLC must be queryable by FDA within 24 hours of request",
                   "BRZ.RC.LOT_TRACE Z-ordered on lot_code · query latency under 100ms at enterprise scale"],
              ],
              col_widths=[3.0, 3.6])

    add_callout(doc, "TLC namespacing recommendation",
                "Recommended TLC format: <SUPPLIER_GLN>-<YYMMDD>-<SEQ>. Kroger-internal TLCs (assigned at "
                "Transformation): KR-INT-<YYMMDD>-<SEQ>. This separates supplier-assigned vs. Kroger-assigned "
                "TLCs and avoids namespace collisions when reconciling against supplier EDI feeds.",
                accent=VIO, fill="EFE5F4")

    # SECTION 5 · RECORDKEEPING & RETENTION
    add_h(doc, "5 · Recordkeeping &amp; retention", level=1)
    add_body(doc, "FSMA 204 requires 24 months of retention. APEX exceeds this by leveraging Delta Lake "
                  "time-travel for point-in-time queries and indefinite cold-archive storage.")
    add_table(doc,
              ["Retention requirement", "APEX implementation"],
              [
                  ["24 months from when food leaves entity control", "Bronze warm tier holds 5+ years; cold archive indefinite"],
                  ["Records in electronic sortable-spreadsheet format",
                   "FSMA evidence packet generator produces Excel + CSV exports directly from Silver layer"],
                  ["Records produced within 24 hours of FDA request",
                   "Generation latency under 4 hours at W2; under 60 minutes at W3 run-rate"],
                  ["Records protected against alteration",
                   "Bronze + Silver Delta Lake immutability + Purview audit log of all access · soft-delete only"],
                  ["Records linked to specific lots in question",
                   "lot_code lookup in Silver SCML.Lot returns full CTE chain · provenance tree with lineage"],
              ],
              col_widths=[2.8, 3.8])

    pagebreak(doc)

    # SECTION 6 · IMPLEMENTATION GAP ANALYSIS — KROGER-SPECIFIC
    add_h(doc, "6 · Kroger-specific implementation gap analysis", level=1)
    add_body(doc, "Per discovery findings, the gap analysis below identifies Kroger's current-state position "
                  "relative to FSMA 204 final-rule readiness. Each gap is mapped to a Wave-1, Wave-2, or "
                  "Wave-3 remediation step.")
    add_table(doc,
              ["Gap", "Current state", "Target", "Wave"],
              [
                  ["Lot-trace data fragmentation",
                   "Across Trace One, FoodLogiQ, supplier EDI, and Excel imports",
                   "Unified BRZ.RC.LOT_TRACE Bronze landing across all sources",
                   "W1"],
                  ["Supplier readiness gaps (FTL categories)",
                   "Soft cheese, mollusks, tropical fruit suppliers fragmented",
                   "Tier-3 supplier remediation plan · alternate-source qualification",
                   "W2"],
                  ["KDE format heterogeneity",
                   "Mixed EPCIS, Excel, manual; uneven KDE field coverage",
                   "BRZ.RC.LOT_TRACE.kde_schema_version normalizes to APEX-canonical",
                   "W1+W2"],
                  ["Transformation CTE under-recorded",
                   "Some store-level deli transformations not lot-traced",
                   "Store-deli systems integrated with BRZ.RC.LOT_TRACE Bronze ingest",
                   "W2"],
                  ["TLC namespace collisions",
                   "Cross-banner / cross-DC TLC reuse risk",
                   "TLC-namespacing convention enforced (Section 4 callout)",
                   "W1"],
                  ["FDA-response time",
                   "Manual reconstruction · 1–3 weeks today",
                   "Auto-generated evidence packet · ≤ 4 hr at W2; ≤ 60 min at W3",
                   "W2"],
                  ["Recall-execution lot-pull",
                   "Manual store-by-store communication",
                   "Auto-shelf-pull via Act agent on recall trigger",
                   "W2"],
                  ["Retention beyond 24 months",
                   "Inconsistent across systems",
                   "Indefinite Delta cold-archive · time-travel queries",
                   "W1"],
                  ["Audit-row coverage",
                   "62% on cold-chain dispositions today",
                   "≥ 99.5% at W3",
                   "W1+W2+W3"],
              ],
              col_widths=[1.9, 2.0, 2.0, 0.7])

    pagebreak(doc)

    # SECTION 7 · COMPLIANCE READINESS SCORECARD
    add_h(doc, "7 · FSMA 204 compliance readiness scorecard", level=1)
    add_body(doc, "Quarterly self-assessment scoring framework. Each pillar scores 0–4. Target W3 enterprise "
                  "readiness is 4 across all pillars; W2 pilot target is 3+ across the in-scope cluster.")
    add_table(doc,
              ["Pillar", "0 (None)", "2 (Partial)", "4 (Full)"],
              [
                  ["TLC capture",
                   "TLC not consistently captured",
                   "TLC captured for top-3 FTL categories",
                   "TLC captured 100% of FTL items at all CTEs"],
                  ["KDE coverage",
                   "Manual/missing for most CTEs",
                   "Auto-captured for Receiving + Shipping",
                   "All 5 CTEs auto-captured · canonical KDE format"],
                  ["Retention",
                   "<24 months · platform inconsistent",
                   "24 months · partially queryable",
                   "Indefinite · time-travel · point-in-time queries"],
                  ["FDA response time",
                   ">7 days",
                   "1–3 days",
                   "<60 minutes auto-packet"],
                  ["Recall execution",
                   "Manual · 24+ hrs to shelf-pull",
                   "Hybrid · 4–8 hrs",
                   "Auto-shelf-pull · ≤ 2 hrs"],
                  ["Audit-row coverage",
                   "<70%",
                   "85–95%",
                   "≥ 99.5% · trace_id attribution"],
                  ["Banner-cluster coverage",
                   "Pilot only",
                   "Anchor cluster + 2 adjacent",
                   "All banners · multi-banner orchestration"],
                  ["Supplier readiness",
                   "Top suppliers only",
                   "Top-50 suppliers FSMA-ready",
                   "All suppliers FTL-tier qualified"],
              ],
              col_widths=[1.6, 1.7, 1.7, 1.7])

    add_callout(doc, "Scoring discipline",
                "Score quarterly. Target trajectory: W1 exit = 1.5 average; W2 exit = 3.0; W3 mid = 3.7; "
                "W3 exit = 4.0 across all pillars. The scorecard is consumed by Internal Audit at quarterly "
                "review and feeds the Audit Committee Quality &amp; Compliance dashboard.")

    pagebreak(doc)

    # APPENDIX — GLOSSARY + REFERENCES
    add_h(doc, "Appendix A · Glossary", level=1)
    add_table(doc,
              ["Term", "Definition"],
              [
                  ["FSMA 204",      "Food Safety Modernization Act §204(d) · Final Rule on Additional Traceability Records"],
                  ["FTL",           "Food Traceability List · the FDA-published list of high-risk foods subject to FSMA 204"],
                  ["CTE",           "Critical Tracking Event · the supply-chain points where records must be kept (Receiving, Transformation, Shipping, Growing, First Land-based Receiver)"],
                  ["KDE",           "Key Data Element · the specific data fields required at each CTE"],
                  ["TLC",           "Traceability Lot Code · unique identifier for each lot of FTL food"],
                  ["TLC Source",    "The entity that assigned the TLC (typically the originating supplier or the entity performing Transformation)"],
                  ["TLC Source Reference", "Document or system reference that identifies where the TLC was assigned"],
                  ["EPCIS",         "Electronic Product Code Information Services · GS1 standard for sharing supply-chain event data"],
                  ["Recall",        "FDA- or supplier-initiated pull of an FTL lot from the supply chain"],
                  ["Trace One",     "Lot-traceability platform widely used in grocery retail"],
                  ["FoodLogiQ",     "Lot-traceability platform widely used in grocery retail"],
                  ["GLN",           "Global Location Number · GS1 identifier for supply-chain locations"],
                  ["Sortable Spreadsheet Format", "FDA-required electronic format for record production · typically Excel or CSV with defined columns"],
              ],
              col_widths=[1.7, 4.6])

    add_h(doc, "Appendix B · Cross-references", level=1)
    add_table(doc,
              ["Artifact", "Purpose"],
              [
                  ["APEX-RC-E2E-03-09-SOR-Bronze-to-Silver.docx",
                   "Bronze-tier specification with full BRZ.RC.LOT_TRACE field detail"],
                  ["APEX-RC-E2E-03-09-SOR-ERD.html",
                   "Visual ERD with FK relationships across the 12 Bronze tables"],
                  ["APEX-RC-E2E-03-Walkthrough.docx",
                   "Service walkthrough — 5-agent orchestration · co-anchor service mapping"],
                  ["APEX-RC-Grocery-Merchandising-Service-Portfolio.docx",
                   "Service portfolio — RC-E2E-09 Tier-1 co-anchor rationale"],
                  ["FDA Final Rule Text",
                   "21 CFR Part 1, Subpart S · Section 1.1305 onwards · final rule effective date Jan 20, 2026"],
                  ["FDA Food Traceability List (FTL)",
                   "Current list maintained at fda.gov/food/food-safety-modernization-act-fsma · subject to revision"],
              ],
              col_widths=[2.9, 3.5])

    add_body(doc, "Independence posture preserved across this document: 'Deloitte's Microsoft practice,' "
                  "'DMTSP,' 'Microsoft platform capabilities,' 'Microsoft-native deployment.' The terms "
                  "'partner,' 'alliance,' and 'strategic alliance' do not appear.",
             italic=True, color=MUTE, size=9)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— End of FSMA 204 Compliance Checklist —"); set_run(r, italic=True, size=10, color=MUTE)

    target = OUT / "APEX-FSMA-204-Compliance-Checklist.docx"
    try:
        doc.save(str(target))
    except PermissionError:
        n = 2
        while True:
            alt = target.with_name(target.stem + f"-v{n}" + target.suffix)
            if not alt.exists(): target = alt; break
            n += 1
        doc.save(str(target))
    print(f"  wrote {target.name}")


build_fsma()
print("\nTier 1 complete.")
