"""Build the SOR · Bronze-to-Silver document for RC-E2E-03 + RC-E2E-09.

Documents the Systems of Record that feed the two Tier-1 RC services for
Grocery Merchandising, with concrete Bronze-tier table specifications and
the canonical mapping to the Silver tier (MERML and SCML schema families).

Output: Kroger 02_projects/FY27_Pipeline/assortment-pricing-agentic/deliverables/
        APEX-RC-E2E-03-09-SOR-Bronze-to-Silver.docx
"""

from __future__ import annotations

import io
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

OUT_DIR = Path(r"C:\Stage\Clients\Industries\Consumer\Retail\Kroger\02_projects\FY27_Pipeline\assortment-pricing-agentic\deliverables")
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT = OUT_DIR / "APEX-RC-E2E-03-09-SOR-Bronze-to-Silver.docx"

# ---- styling ------------------------------------------------------------

INK   = "0B0B0F"
INK2  = "4A4A55"
MUTE  = "8a8a90"
GOLD  = "B8892E"
TEAL  = "2A7F73"
VIO   = "6648B0"
CRIM  = "A22A39"
BAND  = "F3F3F5"
HEAD  = "0B0B0F"
LINE  = "D4D4D7"
BRONZE_BG = "F0E4D0"
SILVER_BG = "E0E5EA"


def set_cell_bg(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex.lstrip("#")); tcPr.append(shd)


def set_cell_border(cell, color_hex=LINE):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "4"); b.set(qn("w:color"), color_hex.lstrip("#"))
        borders.append(b)
    tcPr.append(borders)


def set_run(run, *, bold=False, italic=False, size=None, color=None, font="Arial"):
    if bold: run.bold = True
    if italic: run.italic = True
    if size is not None: run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor.from_string(color.lstrip("#"))
    if font: run.font.name = font


def add_body(doc, text, *, size=11, italic=False, color=INK, space_after=6, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None: p.alignment = align
    r = p.add_run(text)
    set_run(r, size=size, color=color, italic=italic)
    return p


def add_h(doc, text, *, level=1):
    sizes = {1: 18, 2: 14, 3: 12.5, 4: 11.5}
    p = doc.add_paragraph()
    p.style = doc.styles[f"Heading {level}"]
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(4)
    for r in list(p.runs): r.text = ""
    r = p.add_run(text)
    set_run(r, bold=True, size=sizes.get(level, 12), color=INK)
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6"); bottom.set(qn("w:space"), "4"); bottom.set(qn("w:color"), GOLD)
        pBdr.append(bottom); pPr.append(pBdr)
    return p


def add_bullets(doc, items, *, size=11):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(it); set_run(r, size=size, color=INK)


def add_table(doc, header, rows, *, col_widths=None,
              header_fill=HEAD, header_color="FFFFFF", zebra=True):
    n = len(header)
    tbl = doc.add_table(rows=len(rows) + 1, cols=n)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(header):
        c = tbl.rows[0].cells[j]
        set_cell_bg(c, header_fill); set_cell_border(c)
        c.paragraphs[0].text = ""
        r = c.paragraphs[0].add_run(h); set_run(r, bold=True, size=10, color=header_color)
        if col_widths: c.width = Inches(col_widths[j])
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            c = tbl.rows[i].cells[j]
            set_cell_border(c)
            if zebra and i % 2 == 0: set_cell_bg(c, BAND)
            c.paragraphs[0].text = ""
            r = c.paragraphs[0].add_run(str(v)); set_run(r, size=10, color=INK)
            if col_widths: c.width = Inches(col_widths[j])
    doc.add_paragraph()
    return tbl


def add_callout(doc, kind, body, *, accent=GOLD, fill="F6F1E4"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = tbl.rows[0].cells[0]
    set_cell_bg(c, fill); c.paragraphs[0].text = ""
    p_tag = c.paragraphs[0]; r_tag = p_tag.add_run(kind.upper())
    set_run(r_tag, bold=True, size=9, color=accent)
    p = c.add_paragraph(); p.paragraph_format.space_before = Pt(3)
    r = p.add_run(body); set_run(r, size=10.5, color=INK)
    doc.add_paragraph()


def add_kv(doc, pairs, *, lw=1.6, rw=4.9):
    tbl = doc.add_table(rows=len(pairs), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(pairs):
        ck, cv = tbl.rows[i].cells
        ck.width = Inches(lw); cv.width = Inches(rw)
        set_cell_bg(ck, BAND); set_cell_border(ck); set_cell_border(cv)
        ck.paragraphs[0].text = ""; cv.paragraphs[0].text = ""
        rk = ck.paragraphs[0].add_run(k); set_run(rk, bold=True, size=10, color=INK)
        rv = cv.paragraphs[0].add_run(v); set_run(rv, size=10, color=INK)
    doc.add_paragraph()


def add_eyebrow(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text.upper()); set_run(r, bold=True, size=9, color=MUTE)
    rPr = r._element.get_or_add_rPr()
    spacing = OxmlElement("w:spacing"); spacing.set(qn("w:val"), "80"); rPr.append(spacing)


def pagebreak(doc):
    p = doc.add_paragraph(); p.add_run().add_break(WD_BREAK.PAGE)


def add_field_table(doc, sor_id, sor_name, fields, *, source_system, refresh,
                     volume, latency_target, partition_by, zorder_by, notes):
    """Render a per-SOR Bronze-tier specification block."""
    # Header band — dark with bronze accent
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    c1 = tbl.rows[0].cells[0]; c2 = tbl.rows[0].cells[1]
    c1.width = Inches(2.4); c2.width = Inches(4.0)
    set_cell_bg(c1, BRONZE_BG); set_cell_bg(c2, BRONZE_BG)
    set_cell_border(c1); set_cell_border(c2)
    c1.paragraphs[0].text = ""; c2.paragraphs[0].text = ""
    rb = c1.paragraphs[0].add_run("BRONZE TABLE"); set_run(rb, bold=True, size=9, color=GOLD)
    rid = c1.add_paragraph().add_run(sor_id); set_run(rid, bold=True, size=11, color=INK, font="Consolas")
    rname = c2.paragraphs[0].add_run(sor_name); set_run(rname, bold=True, size=12, color=INK)
    doc.add_paragraph()

    # Meta KV
    add_kv(doc, [
        ("Source system",    source_system),
        ("Refresh cadence",  refresh),
        ("Approx. volume",   volume),
        ("Latency SLO",      latency_target),
        ("Partition by",     partition_by),
        ("Z-order by",       zorder_by),
        ("Notes",            notes),
    ], lw=1.7, rw=4.6)

    # Field table
    add_table(doc,
              ["Field name", "Type", "PK / FK", "Description", "Required"],
              fields,
              col_widths=[1.7, 1.0, 0.7, 2.5, 0.7])


# ============================================================================
# DATA — SOR catalog and field specs
# ============================================================================

SOR_INVENTORY = [
    # (Bronze table ID, Friendly name, Source system, Owner, Service primary)
    ("BRZ.RC.POS_TXN",          "POS transactions",                     "Toshiba/IBM/NCR POS · CDC",         "Store Systems / IT-Retail",  "RC-E2E-03"),
    ("BRZ.RC.ITEM_MASTER",      "Item master / PIM",                    "Item Master · PIM",                  "Merchandising Master Data",   "Both"),
    ("BRZ.RC.PRICING",          "Pricing & price-zone",                 "Revionics / 84.51° pricing engine",  "Pricing Operations",          "RC-E2E-03"),
    ("BRZ.RC.PROMO_MARKDOWN",   "Promotional & markdown events",         "Revionics Markdown · ad-system",     "Promotions / Pricing",        "RC-E2E-03"),
    ("BRZ.RC.INVENTORY",        "Store + DC inventory positions",        "SAP IM · Manhattan WMS",             "Supply Chain · Inventory",    "Both"),
    ("BRZ.RC.REFRIG_TELEMETRY", "Refrigeration telemetry (in-store)",    "Emerson E2E · Honeywell Niagara · OPC-UA over MQTT", "Facilities Engineering", "RC-E2E-03"),
    ("BRZ.RC.REFRIG_SERVICE",   "Refrigeration vendor service tickets",  "Emerson · Lennox · Hussmann portals (REST)", "Facilities Engineering",      "RC-E2E-03"),
    ("BRZ.RC.FOOD_SAFETY",      "Food-safety incident & disposition log", "Quality Mgmt System · Trace One",   "Food Safety / Quality",       "Both"),
    ("BRZ.RC.LOT_TRACE",        "FSMA 204 lot-provenance events",        "Trace One · FoodLogiQ · supplier EDI", "Food Safety / Vendor Mgmt", "RC-E2E-09"),
    ("BRZ.RC.VENDOR_MASTER",    "Vendor master + scorecards",            "Vendor Master · supplier portal",    "Category Procurement",        "RC-E2E-09"),
    ("BRZ.RC.RECEIPT",          "DC + store receipts (with lot codes)",  "Manhattan WMS · DC-receiving",       "Supply Chain · Receiving",    "RC-E2E-09"),
    ("BRZ.RC.LOYALTY_HH",       "Loyalty households (read-only feed)",   "84.51° loyalty platform",            "Customer / Loyalty",          "RC-E2E-03 (elasticity)"),
]

# Detailed field specs per Bronze table
# Format: (field_name, type, key_marker, description, required)

POS_TXN_FIELDS = [
    ("transaction_id",        "STRING",     "PK",  "Source POS transaction identifier",                              "Y"),
    ("transaction_line_no",   "INT",        "PK",  "Line number within transaction",                                  "Y"),
    ("transaction_ts",        "TIMESTAMP",  "",    "POS transaction timestamp (banner local)",                        "Y"),
    ("banner",                "STRING",     "FK",  "Banner code (KR, RAL, FM, KS, SMI, FRY, HT, MAR, QFC, etc.)",     "Y"),
    ("store_id",              "STRING",     "FK",  "Source store identifier",                                          "Y"),
    ("register_id",           "STRING",     "",    "Lane / register identifier",                                       "Y"),
    ("cashier_id",            "STRING",     "",    "Cashier / associate identifier",                                   "N"),
    ("gtin",                  "STRING",     "FK",  "Global Trade Item Number (preferred for traceability)",            "Y"),
    ("upc",                   "STRING",     "FK",  "UPC if GTIN unavailable",                                          "N"),
    ("sku",                   "STRING",     "FK",  "Internal SKU identifier (if used)",                                "N"),
    ("qty",                   "DECIMAL(10,3)", "", "Quantity (supports weighted items)",                               "Y"),
    ("unit_of_measure",       "STRING",     "",    "EA / LB / KG / OZ",                                                "Y"),
    ("unit_price",            "DECIMAL(10,4)", "", "Per-unit price at moment of transaction",                          "Y"),
    ("ext_price",             "DECIMAL(12,2)", "", "Extended price (qty × unit_price)",                                "Y"),
    ("discount_amount",       "DECIMAL(10,2)", "", "Discount applied at line",                                          "N"),
    ("promo_id",              "STRING",     "FK",  "Promotion identifier if any",                                      "N"),
    ("loyalty_household_id",  "STRING",     "FK",  "Hashed household ID if loyalty-linked",                            "N"),
    ("tender_type",           "STRING",     "",    "Cash / credit / debit / EBT / etc.",                               "Y"),
    ("void_flag",             "BOOLEAN",    "",    "Voided line indicator",                                            "Y"),
    ("source_system",         "STRING",     "",    "Source POS family identifier (Toshiba, IBM, NCR)",                 "Y"),
    ("source_record_id",      "STRING",     "",    "Source record GUID for idempotent dedup",                          "Y"),
    ("ingest_ts",             "TIMESTAMP",  "",    "Bronze landing timestamp (UTC)",                                   "Y"),
    ("batch_id",              "STRING",     "",    "Ingest batch / commit identifier",                                  "Y"),
    ("_is_deleted",           "BOOLEAN",    "",    "Soft-delete marker (CDC)",                                         "Y"),
]

ITEM_MASTER_FIELDS = [
    ("sku",                   "STRING",     "PK",  "Internal SKU identifier",                                          "Y"),
    ("gtin",                  "STRING",     "",    "Global Trade Item Number",                                          "Y"),
    ("upc",                   "STRING",     "",    "Universal Product Code",                                            "N"),
    ("description_short",     "STRING(80)", "",    "Short description for receipts and POS",                            "Y"),
    ("description_long",      "STRING(255)","",    "Full description for digital surfaces",                              "N"),
    ("brand",                 "STRING",     "",    "Brand name",                                                         "Y"),
    ("private_label_flag",    "BOOLEAN",    "",    "True for Our Brands / Simple Truth / Murray's Cheese / etc.",       "Y"),
    ("category_l1",           "STRING",     "FK",  "Top-level category (e.g., DAIRY)",                                  "Y"),
    ("category_l2",           "STRING",     "FK",  "Mid-level category (e.g., CHEESE)",                                  "Y"),
    ("category_l3",           "STRING",     "FK",  "Sub-category (e.g., SPECIALTY-CHEESE)",                              "N"),
    ("vendor_id",             "STRING",     "FK",  "Primary vendor identifier (FK to BRZ.RC.VENDOR_MASTER)",             "Y"),
    ("pack_size",             "STRING",     "",    "Package size descriptor",                                            "Y"),
    ("uom",                   "STRING",     "",    "Unit of measure for retail (EA, LB)",                               "Y"),
    ("temp_class",            "STRING",     "",    "AMBIENT / REFRIGERATED / FROZEN / DELI / FRESH",                     "Y"),
    ("ftl_flag",              "BOOLEAN",    "",    "FDA Food Traceability List indicator",                               "Y"),
    ("ftl_category",          "STRING",     "",    "FTL category code (cheese, leafy greens, etc.) if ftl_flag",         "N"),
    ("shelf_life_days",       "INT",        "",    "Default shelf life in days (override per lot)",                      "N"),
    ("hazmat_flag",           "BOOLEAN",    "",    "Hazmat indicator",                                                   "Y"),
    ("active_flag",           "BOOLEAN",    "",    "Active in catalog",                                                  "Y"),
    ("effective_date",        "DATE",       "",    "First-effective date for this row",                                  "Y"),
    ("end_date",              "DATE",       "",    "End-effective date (SCD Type 2 in Silver)",                          "N"),
    ("source_system",         "STRING",     "",    "Source PIM identifier",                                              "Y"),
    ("source_record_id",      "STRING",     "",    "Source record GUID",                                                  "Y"),
    ("ingest_ts",             "TIMESTAMP",  "",    "Bronze landing timestamp",                                            "Y"),
    ("batch_id",              "STRING",     "",    "Ingest batch identifier",                                             "Y"),
    ("_is_deleted",           "BOOLEAN",    "",    "Soft-delete marker (CDC)",                                            "Y"),
]

PRICING_FIELDS = [
    ("pricing_record_id",     "STRING",     "PK",  "Source pricing record identifier",                                   "Y"),
    ("sku",                   "STRING",     "FK",  "SKU identifier (FK to BRZ.RC.ITEM_MASTER)",                          "Y"),
    ("price_zone_id",         "STRING",     "FK",  "Pricing zone identifier",                                            "Y"),
    ("store_id",              "STRING",     "FK",  "Store identifier (NULL when zone-level)",                            "N"),
    ("list_price",            "DECIMAL(10,4)", "", "Everyday published price",                                            "Y"),
    ("promo_price",           "DECIMAL(10,4)", "", "Promotional price if active",                                          "N"),
    ("price_action",          "STRING",     "",    "REGULAR / PROMO / MARKDOWN / CLEARANCE / TPR",                       "Y"),
    ("promo_id",              "STRING",     "FK",  "Linked promotion (FK to BRZ.RC.PROMO_MARKDOWN)",                     "N"),
    ("elasticity_curve_id",   "STRING",     "FK",  "Linked per-SKU elasticity curve identifier",                          "N"),
    ("currency",              "STRING(3)",  "",    "ISO currency code (USD)",                                            "Y"),
    ("effective_date",        "TIMESTAMP",  "",    "First-effective timestamp",                                          "Y"),
    ("expiration_date",       "TIMESTAMP",  "",    "Expiration timestamp",                                               "Y"),
    ("authoring_user_id",     "STRING",     "",    "Pricing-decision authoring user identifier",                          "N"),
    ("source_system",         "STRING",     "",    "Source pricing engine identifier (Revionics / 84.51°)",              "Y"),
    ("source_record_id",      "STRING",     "",    "Source record GUID",                                                  "Y"),
    ("ingest_ts",             "TIMESTAMP",  "",    "Bronze landing timestamp",                                            "Y"),
    ("batch_id",              "STRING",     "",    "Ingest batch identifier",                                             "Y"),
    ("_is_deleted",           "BOOLEAN",    "",    "Soft-delete marker",                                                  "Y"),
]

PROMO_MARKDOWN_FIELDS = [
    ("promo_id",              "STRING",     "PK",  "Promotion identifier (source-system primary key)",                   "Y"),
    ("promo_type",            "STRING",     "",    "TPR / BOGO / AD_BLOCK / CLUB / MARKDOWN / CLEARANCE",                 "Y"),
    ("promo_subtype",         "STRING",     "",    "Sub-classification (e.g., 2-FOR-X, % off, $ off)",                    "N"),
    ("sku",                   "STRING",     "FK",  "SKU identifier (NULL if sku_group_id)",                              "N"),
    ("sku_group_id",          "STRING",     "FK",  "SKU group / category target if not single SKU",                       "N"),
    ("price_zone_id",         "STRING",     "FK",  "Zone scope (NULL if all zones)",                                      "N"),
    ("store_id",              "STRING",     "FK",  "Store scope (NULL if zone or banner)",                                "N"),
    ("discount_type",         "STRING",     "",    "PERCENT / DOLLAR / TARGET_PRICE",                                     "Y"),
    ("discount_value",        "DECIMAL(10,4)", "", "Discount value (percent as 0-1 or as $)",                             "Y"),
    ("min_qty",               "INT",        "",    "Minimum quantity (BOGO, multi-buy)",                                   "N"),
    ("vehicle",               "STRING",     "",    "DIGITAL / AD_BLOCK / IN_STORE_SIGN / LOYALTY / MAIL",                  "Y"),
    ("ad_block_id",           "STRING",     "FK",  "Linked ad-block identifier if vehicle = AD_BLOCK",                     "N"),
    ("loyalty_required_flag", "BOOLEAN",    "",    "True if loyalty membership required",                                  "Y"),
    ("start_ts",              "TIMESTAMP",  "",    "Promotion start (UTC)",                                                "Y"),
    ("end_ts",                "TIMESTAMP",  "",    "Promotion end (UTC)",                                                  "Y"),
    ("expected_lift_pct",     "DECIMAL(5,2)","",   "Authoring forecast lift",                                              "N"),
    ("source_system",         "STRING",     "",    "Source promo engine identifier",                                       "Y"),
    ("source_record_id",      "STRING",     "",    "Source record GUID",                                                   "Y"),
    ("ingest_ts",             "TIMESTAMP",  "",    "Bronze landing timestamp",                                             "Y"),
    ("batch_id",              "STRING",     "",    "Ingest batch identifier",                                              "Y"),
    ("_is_deleted",           "BOOLEAN",    "",    "Soft-delete marker",                                                   "Y"),
]

INVENTORY_FIELDS = [
    ("snapshot_id",           "STRING",     "PK",  "Source snapshot identifier",                                          "Y"),
    ("snapshot_ts",           "TIMESTAMP",  "PK",  "Snapshot timestamp (UTC)",                                            "Y"),
    ("location_type",         "STRING",     "",    "STORE / DC / IN_TRANSIT",                                              "Y"),
    ("location_id",           "STRING",     "FK",  "Store_id or DC_id",                                                    "Y"),
    ("sku",                   "STRING",     "FK",  "SKU identifier",                                                       "Y"),
    ("on_hand_qty",           "DECIMAL(12,3)", "", "Quantity available",                                                   "Y"),
    ("on_order_qty",          "DECIMAL(12,3)", "", "Quantity on order",                                                    "N"),
    ("in_transit_qty",        "DECIMAL(12,3)", "", "Quantity in transit",                                                  "N"),
    ("allocated_qty",         "DECIMAL(12,3)", "", "Quantity allocated",                                                   "N"),
    ("safety_stock_qty",      "DECIMAL(12,3)", "", "Safety stock target",                                                  "N"),
    ("last_receipt_ts",       "TIMESTAMP",  "",    "Most recent receipt timestamp",                                        "N"),
    ("next_expected_rcv_ts",  "TIMESTAMP",  "",    "Next expected receipt timestamp",                                      "N"),
    ("source_system",         "STRING",     "",    "Source IM/WMS identifier",                                             "Y"),
    ("source_record_id",      "STRING",     "",    "Source record GUID",                                                    "Y"),
    ("ingest_ts",             "TIMESTAMP",  "",    "Bronze landing timestamp",                                              "Y"),
    ("batch_id",              "STRING",     "",    "Ingest batch identifier",                                               "Y"),
    ("_is_deleted",           "BOOLEAN",    "",    "Soft-delete marker",                                                    "Y"),
]

REFRIG_TELEMETRY_FIELDS = [
    ("reading_id",            "STRING",     "PK",  "Source telemetry reading identifier",                                  "Y"),
    ("reading_ts",            "TIMESTAMP",  "PK",  "Sensor read timestamp (UTC, 30-second cadence target)",                "Y"),
    ("store_id",              "STRING",     "FK",  "Store identifier",                                                      "Y"),
    ("fixture_id",            "STRING",     "FK",  "Refrigerated fixture identifier",                                       "Y"),
    ("sensor_id",             "STRING",     "",    "Sensor probe identifier (for multi-probe fixtures)",                   "Y"),
    ("temperature_f",         "DECIMAL(6,2)", "",  "Temperature reading in Fahrenheit",                                     "Y"),
    ("temperature_setpoint_f","DECIMAL(6,2)", "",  "Setpoint at time of reading",                                           "Y"),
    ("humidity_pct",          "DECIMAL(5,2)", "",  "Relative humidity if instrumented",                                     "N"),
    ("door_open_flag",        "BOOLEAN",    "",    "Door-open indicator (when sensored)",                                   "N"),
    ("alarm_state",           "STRING",     "",    "NORMAL / WARN / CRITICAL / OFFLINE",                                    "Y"),
    ("alarm_severity",        "INT",        "",    "0-5 severity scale",                                                    "Y"),
    ("compressor_runtime_pct","DECIMAL(5,2)", "",  "Compressor duty cycle in last window",                                  "N"),
    ("offline_flag",          "BOOLEAN",    "",    "Sensor-offline indicator (drives data-quality)",                        "Y"),
    ("source_system",         "STRING",     "",    "EMS identifier (Emerson E2E, Honeywell Niagara, etc.)",                "Y"),
    ("source_record_id",      "STRING",     "",    "Source record GUID",                                                    "Y"),
    ("ingest_ts",             "TIMESTAMP",  "",    "Bronze landing timestamp (sub-second target)",                          "Y"),
    ("batch_id",              "STRING",     "",    "Eventstream micro-batch identifier",                                    "Y"),
]

REFRIG_SERVICE_FIELDS = [
    ("ticket_id",             "STRING",     "PK",  "Vendor ticket identifier",                                              "Y"),
    ("vendor_id",             "STRING",     "FK",  "Refrigeration vendor (Emerson, Lennox, Hussmann, Hill PHOENIX)",       "Y"),
    ("store_id",              "STRING",     "FK",  "Store identifier",                                                      "Y"),
    ("fixture_id",            "STRING",     "FK",  "Affected fixture identifier",                                           "Y"),
    ("open_ts",               "TIMESTAMP",  "",    "Ticket open timestamp",                                                 "Y"),
    ("close_ts",              "TIMESTAMP",  "",    "Ticket close timestamp",                                                "N"),
    ("severity",              "STRING",     "",    "P1 / P2 / P3 / P4",                                                     "Y"),
    ("ticket_status",         "STRING",     "",    "OPEN / DISPATCHED / IN_PROGRESS / CLOSED / CANCELLED",                  "Y"),
    ("root_cause_code",       "STRING",     "",    "Vendor-coded root cause",                                               "N"),
    ("repair_action",         "STRING",     "",    "Free-text repair description",                                          "N"),
    ("technician_id",         "STRING",     "",    "Vendor technician identifier",                                          "N"),
    ("parts_replaced",        "ARRAY<STRING>","",  "Parts list",                                                            "N"),
    ("labor_hours",           "DECIMAL(6,2)", "",  "Labor hours billed",                                                    "N"),
    ("ticket_cost_usd",       "DECIMAL(10,2)", "", "Total ticket cost",                                                     "N"),
    ("source_system",         "STRING",     "",    "Vendor portal identifier",                                              "Y"),
    ("source_record_id",      "STRING",     "",    "Source record GUID",                                                    "Y"),
    ("ingest_ts",             "TIMESTAMP",  "",    "Bronze landing timestamp",                                              "Y"),
    ("batch_id",              "STRING",     "",    "Ingest batch identifier",                                                "Y"),
    ("_is_deleted",           "BOOLEAN",    "",    "Soft-delete marker",                                                    "Y"),
]

FOOD_SAFETY_FIELDS = [
    ("incident_id",           "STRING",     "PK",  "Incident identifier",                                                   "Y"),
    ("incident_ts",           "TIMESTAMP",  "",    "Incident occurrence timestamp",                                         "Y"),
    ("incident_type",         "STRING",     "",    "EXCURSION / SPOILAGE / COMPLAINT / RECALL / FDA_INQUIRY / SUPPLIER",   "Y"),
    ("severity",              "STRING",     "",    "LOW / MEDIUM / HIGH / CRITICAL",                                        "Y"),
    ("store_id",              "STRING",     "FK",  "Store identifier (NULL if banner-wide)",                                "N"),
    ("fixture_id",            "STRING",     "FK",  "Fixture identifier (if cold-chain)",                                    "N"),
    ("sku",                   "STRING",     "FK",  "Affected SKU",                                                          "N"),
    ("lot_code",              "STRING",     "FK",  "Lot code (FK to BRZ.RC.LOT_TRACE)",                                     "N"),
    ("disposition",           "STRING",     "",    "SELL_THROUGH / MARKDOWN / DESTROY / HOLD / RETURN_TO_VENDOR",          "Y"),
    ("disposition_reason",    "STRING",     "",    "Free-text rationale",                                                   "N"),
    ("evidence_refs",         "ARRAY<STRING>","",  "URIs to supporting evidence (photos, lab reports, etc.)",              "N"),
    ("approver_id",           "STRING",     "",    "Identity that approved disposition (Entra ID)",                         "Y"),
    ("approver_role",         "STRING",     "",    "Role of approver (StoreOpsLead, FoodSafetyLead, CategoryMgr)",         "Y"),
    ("approval_ts",           "TIMESTAMP",  "",    "Approval timestamp",                                                    "Y"),
    ("hitl_kind",             "STRING",     "",    "ZERO_TOUCH / ACK_ONLY / HITL / ESCALATION",                            "Y"),
    ("trace_id",              "STRING",     "",    "APEX trace_id correlator",                                              "Y"),
    ("source_system",         "STRING",     "",    "Source QMS / FSMA system",                                              "Y"),
    ("source_record_id",      "STRING",     "",    "Source record GUID",                                                    "Y"),
    ("ingest_ts",             "TIMESTAMP",  "",    "Bronze landing timestamp",                                              "Y"),
    ("batch_id",              "STRING",     "",    "Ingest batch identifier",                                                "Y"),
    ("_is_deleted",           "BOOLEAN",    "",    "Soft-delete marker",                                                    "Y"),
]

LOT_TRACE_FIELDS = [
    ("lot_event_id",          "STRING",     "PK",  "FSMA 204 Critical Tracking Event (CTE) record ID",                     "Y"),
    ("event_ts",              "TIMESTAMP",  "",    "CTE event timestamp",                                                   "Y"),
    ("cte_type",              "STRING",     "",    "GROWING / RECEIVING / TRANSFORMING / SHIPPING / FIRST-RECEIPT",        "Y"),
    ("gtin",                  "STRING",     "FK",  "GTIN of the traced item",                                               "Y"),
    ("sku",                   "STRING",     "FK",  "Internal SKU mapping",                                                  "N"),
    ("lot_code",              "STRING",     "PK",  "Lot or batch identifier",                                                "Y"),
    ("supplier_id",           "STRING",     "FK",  "Originating supplier (FK to BRZ.RC.VENDOR_MASTER)",                    "Y"),
    ("supplier_lot_code",     "STRING",     "",    "Supplier-system lot code (often differs from internal)",               "N"),
    ("harvest_or_prod_date",  "DATE",       "",    "Original production / harvest date",                                    "Y"),
    ("expiration_date",       "DATE",       "",    "Expiration / use-by / sell-by date",                                    "N"),
    ("ftl_category",          "STRING",     "",    "FDA Food Traceability List category code",                              "Y"),
    ("from_location_id",      "STRING",     "FK",  "Originating location (DC, supplier, store)",                            "N"),
    ("to_location_id",        "STRING",     "FK",  "Destination location",                                                  "N"),
    ("qty",                   "DECIMAL(12,3)", "", "Quantity moved in this CTE",                                            "Y"),
    ("uom",                   "STRING",     "",    "Unit of measure",                                                       "Y"),
    ("kde_payload",           "STRING",     "",    "Key Data Elements as JSON (per FSMA 204 final rule)",                   "Y"),
    ("kde_schema_version",    "STRING",     "",    "FSMA 204 KDE schema version",                                           "Y"),
    ("source_system",         "STRING",     "",    "Trace One / FoodLogiQ / supplier-EDI source",                          "Y"),
    ("source_record_id",      "STRING",     "",    "Source record GUID",                                                    "Y"),
    ("ingest_ts",             "TIMESTAMP",  "",    "Bronze landing timestamp",                                              "Y"),
    ("batch_id",              "STRING",     "",    "Ingest batch identifier",                                                "Y"),
    ("_is_deleted",           "BOOLEAN",    "",    "Soft-delete marker",                                                    "Y"),
]

VENDOR_MASTER_FIELDS = [
    ("vendor_id",             "STRING",     "PK",  "Internal vendor identifier",                                            "Y"),
    ("vendor_name",           "STRING",     "",    "Vendor display name",                                                   "Y"),
    ("parent_company",        "STRING",     "",    "Parent company / corporate group",                                       "N"),
    ("vendor_type",           "STRING",     "",    "BRAND / PRIVATE_LABEL_CO_MAN / DSD / WHOLESALE",                        "Y"),
    ("ftl_capability_tier",   "STRING",     "",    "TIER_1 / TIER_2 / TIER_3 / NOT_FTL_QUALIFIED",                         "Y"),
    ("kde_format_supported",  "ARRAY<STRING>","",  "Supplier-supported KDE formats (EPCIS, Excel, manual)",                "Y"),
    ("scorecard_grade",       "STRING",     "",    "A / B / C / D · current category-supplier scorecard",                   "N"),
    ("last_audit_date",       "DATE",       "",    "Most recent supplier audit date",                                       "N"),
    ("category_l1_focus",     "ARRAY<STRING>","",  "Primary category coverage",                                             "N"),
    ("active_flag",           "BOOLEAN",    "",    "Active vendor status",                                                  "Y"),
    ("effective_date",        "DATE",       "",    "First-effective date for this row",                                     "Y"),
    ("end_date",              "DATE",       "",    "End-effective date (SCD Type 2 in Silver)",                             "N"),
    ("source_system",         "STRING",     "",    "Vendor master source",                                                  "Y"),
    ("source_record_id",      "STRING",     "",    "Source record GUID",                                                    "Y"),
    ("ingest_ts",             "TIMESTAMP",  "",    "Bronze landing timestamp",                                              "Y"),
    ("batch_id",              "STRING",     "",    "Ingest batch identifier",                                                "Y"),
    ("_is_deleted",           "BOOLEAN",    "",    "Soft-delete marker",                                                    "Y"),
]

RECEIPT_FIELDS = [
    ("receipt_id",            "STRING",     "PK",  "Receipt event identifier",                                              "Y"),
    ("receipt_ts",            "TIMESTAMP",  "",    "Receipt timestamp",                                                     "Y"),
    ("location_type",         "STRING",     "",    "DC / STORE",                                                            "Y"),
    ("location_id",           "STRING",     "FK",  "DC or store identifier",                                                "Y"),
    ("vendor_id",             "STRING",     "FK",  "Originating vendor",                                                    "Y"),
    ("po_number",             "STRING",     "",    "Purchase order number",                                                 "Y"),
    ("asn_number",            "STRING",     "",    "Advance ship notice number",                                            "N"),
    ("gtin",                  "STRING",     "FK",  "GTIN received",                                                         "Y"),
    ("sku",                   "STRING",     "FK",  "SKU received",                                                          "Y"),
    ("lot_code",              "STRING",     "FK",  "Lot code received (FK to BRZ.RC.LOT_TRACE)",                            "Y"),
    ("qty_received",          "DECIMAL(12,3)", "", "Quantity received",                                                     "Y"),
    ("qty_rejected",          "DECIMAL(12,3)", "", "Quantity rejected at dock",                                             "N"),
    ("temp_at_receipt_f",     "DECIMAL(6,2)", "",  "Temperature reading at receipt (cold chain)",                           "N"),
    ("source_system",         "STRING",     "",    "Source WMS / receiving system",                                          "Y"),
    ("source_record_id",      "STRING",     "",    "Source record GUID",                                                    "Y"),
    ("ingest_ts",             "TIMESTAMP",  "",    "Bronze landing timestamp",                                              "Y"),
    ("batch_id",              "STRING",     "",    "Ingest batch identifier",                                                "Y"),
    ("_is_deleted",           "BOOLEAN",    "",    "Soft-delete marker",                                                    "Y"),
]

LOYALTY_HH_FIELDS = [
    ("loyalty_household_id",  "STRING",     "PK",  "Hashed household identifier",                                            "Y"),
    ("registration_ts",       "TIMESTAMP",  "",    "Loyalty registration timestamp",                                         "Y"),
    ("primary_banner",        "STRING",     "",    "Primary shopping banner",                                                "N"),
    ("primary_store_id",      "STRING",     "FK",  "Primary store",                                                          "N"),
    ("segment",               "STRING",     "",    "84.51° segment label",                                                  "N"),
    ("ltv_estimate_usd",      "DECIMAL(12,2)", "", "84.51° lifetime-value estimate",                                         "N"),
    ("last_active_ts",        "TIMESTAMP",  "",    "Most recent transaction timestamp",                                       "N"),
    ("consent_state",         "STRING",     "",    "FULL / RESTRICTED / WITHDRAWN — for personalization gating",            "Y"),
    ("source_system",         "STRING",     "",    "84.51° loyalty platform identifier",                                    "Y"),
    ("source_record_id",      "STRING",     "",    "Source record GUID",                                                    "Y"),
    ("ingest_ts",             "TIMESTAMP",  "",    "Bronze landing timestamp",                                              "Y"),
    ("batch_id",              "STRING",     "",    "Ingest batch identifier",                                                "Y"),
    ("_is_deleted",           "BOOLEAN",    "",    "Soft-delete marker",                                                    "Y"),
]


# ============================================================================
# BUILD
# ============================================================================

def build():
    doc = Document()
    for section in doc.sections:
        section.page_width  = Inches(8.5); section.page_height = Inches(11)
        section.top_margin = section.bottom_margin = Inches(0.85)
        section.left_margin = section.right_margin = Inches(0.95)
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(11)

    # ============ COVER ============
    add_eyebrow(doc, "APEX · RC-E2E-03 + RC-E2E-09 · SOR · BRONZE-TO-SILVER")
    p = doc.add_paragraph()
    r = p.add_run("Systems of Record"); set_run(r, bold=True, size=26, color=INK)
    p1b = doc.add_paragraph()
    r1b = p1b.add_run("Bronze landing → Silver canonical"); set_run(r1b, bold=True, size=20, color=INK)
    p2 = doc.add_paragraph()
    r2 = p2.add_run("Microsoft Fabric · Assortment & Pricing + Product Tracking · Kroger reference")
    set_run(r2, size=13, color=INK2, italic=True)

    doc.add_paragraph()
    add_body(doc,
             "A Bronze-tier-focused systems-of-record (SOR) specification for the two Tier-1 RC "
             "services that anchor the Grocery Merchandising sub-category. Twelve Bronze tables "
             "with field-level definitions, ingest patterns, and partition / Z-order guidance, "
             "set up so the Silver canonical layer (MERML and SCML) lands cleanly on top.",
             italic=True, color=INK2)

    doc.add_paragraph()
    add_kv(doc, [
        ("Document type",     "SOR Specification · Bronze + Silver mapping"),
        ("Services",          "RC-E2E-03 Assortment & Pricing · RC-E2E-09 Product Tracking"),
        ("Data platform",     "Microsoft Fabric · Lakehouse + Eventstream + Real-Time Intelligence"),
        ("Bronze tables",     "12 (POS Txn, Item Master, Pricing, Promo/Markdown, Inventory, Refrig Telemetry, Refrig Service, Food Safety, Lot Trace, Vendor Master, Receipt, Loyalty HH)"),
        ("Silver families",   "MERML (Merchandising) · SCML (Supply Chain & Lot)"),
        ("Reference client",  "The Kroger Co."),
        ("Audience",          "Data Engineering · Architecture · Fabric admin · Source-system owners"),
        ("Companion",         "APEX-RC-E2E-03-Walkthrough.docx · APEX-RC-Grocery-Merchandising-Service-Portfolio.docx"),
        ("Version / date",    "v1.0 · 27 Apr 2026"),
    ])

    pagebreak(doc)

    # ============ SECTION 1 · ARCHITECTURE CONTEXT ============
    add_h(doc, "1 · Architecture context — Fabric medallion for APEX", level=1)
    add_body(doc,
             "APEX runs on Microsoft Fabric as the production analytics + AI substrate. Source "
             "data lands in a Bronze tier, gets canonicalized to a Silver tier (MERML and SCML "
             "schema families), and rolls up to a Gold tier of curated KPI marts. RC-E2E-03 and "
             "RC-E2E-09 agents primarily read from Silver; LEDGER (the APEX 14-field audit row "
             "store) sits alongside Gold and is consumed by Power BI and Privacy/Audit "
             "evidence-packet generation.")

    add_h(doc, "1.1 · Tier responsibilities", level=2)
    add_table(doc,
              ["Tier", "Responsibility", "Format / engine", "Schema discipline"],
              [
                  ["Bronze", "Raw, source-aligned, append-only landing. Preserves source fidelity. The contract surface with each upstream system.",
                   "Delta Lake on OneLake · partitioned by ingest_date",
                   "Source-system-shaped + ingest metadata · additive evolution only"],
                  ["Silver", "Canonical, conformed, deduplicated. The dimensional model APEX agents read against.",
                   "Delta Lake on OneLake · SCD Type 2 where appropriate",
                   "MERML + SCML schema families · controlled evolution via PR"],
                  ["Gold",   "Curated KPI marts and Power BI semantic models. Aggregations, derived metrics, BI surfaces.",
                   "Delta + Direct Lake semantic models",
                   "Subject-area marts · scenario-aligned · stable APIs"],
                  ["LEDGER", "APEX 14-field audit-row store · sits alongside Gold, queried directly by Power BI and Privacy/Audit packet generators.",
                   "Delta Lake + PostgreSQL index for trace_id lookup",
                   "14-field audit-row schema · append-only · immutable"],
              ],
              col_widths=[0.9, 2.7, 1.7, 2.0])

    add_h(doc, "1.2 · Bronze tier conventions for this engagement", level=2)
    add_bullets(doc, [
        "One Bronze table per logical source feed. Twelve tables documented in Section 4. Each lands in the lakehouse under a per-service schema (sch_rc_e2e_03, sch_rc_e2e_09).",
        "Append-only with soft-delete marker (_is_deleted). CDC drives soft-delete events; never hard-delete in Bronze.",
        "Ingest metadata is mandatory: source_system, source_record_id, ingest_ts (UTC), batch_id, _is_deleted. These five fields appear in every Bronze table.",
        "Partition by ingest_date (yyyy-MM-dd). Z-order on join keys per table (Section 4 specifies).",
        "Idempotency via dedup on (source_system, source_record_id) within a batch. Reprocess-safe.",
        "Schema evolution is additive only. Net-new fields land in Bronze without breaking downstream Silver.",
        "Real-time feeds (refrigeration telemetry, FSMA CTE events) ingest via Fabric Eventstream + Real-Time Intelligence; batch feeds (item master, vendor master) ingest via Data Factory pipelines.",
        "All timestamps in Bronze are UTC. Banner-local timestamps are preserved as separate fields where business-relevant.",
    ])

    add_callout(doc, "Why Bronze fidelity matters",
                "If Bronze drops fields or applies transformations beyond the universal ingest "
                "metadata, debugging Silver canonical issues becomes archeology. Bronze is the "
                "ground truth for every retroactive question — recall events, FSMA-204 "
                "inspections, audit reconciliations. The discipline is: land everything, "
                "transform later. Silver is where transformations live; Bronze just records.",
                accent=GOLD)

    add_h(doc, "1.3 · From Bronze to Silver — canonical mapping", level=2)
    add_body(doc,
             "Silver is the canonical layer that APEX agents read. It conforms to two schema "
             "families that ride across all RC scenarios:")
    add_table(doc,
              ["Schema family", "Owns", "Bronze inputs (primary)"],
              [
                  ["MERML.Price",        "Per-SKU per-zone pricing state with effective windows",
                   "BRZ.RC.PRICING · BRZ.RC.PROMO_MARKDOWN · BRZ.RC.ITEM_MASTER"],
                  ["MERML.Markdown",     "Markdown events with discount, vehicle, and outcome links",
                   "BRZ.RC.PROMO_MARKDOWN · BRZ.RC.PRICING · BRZ.RC.POS_TXN"],
                  ["MERML.Elasticity",   "Per-SKU price-elasticity curves",
                   "BRZ.RC.POS_TXN · BRZ.RC.PRICING (derived in Silver)"],
                  ["SCML.Lot",           "Per-lot provenance with FSMA 204 KDEs",
                   "BRZ.RC.LOT_TRACE · BRZ.RC.RECEIPT · BRZ.RC.VENDOR_MASTER"],
                  ["SCML.Inventory",     "Store and DC inventory positions over time",
                   "BRZ.RC.INVENTORY · BRZ.RC.RECEIPT · BRZ.RC.POS_TXN"],
                  ["SCML.Fixture",       "Refrigerated fixture state and excursion classification",
                   "BRZ.RC.REFRIG_TELEMETRY · BRZ.RC.REFRIG_SERVICE"],
                  ["SCML.Receipt",       "DC + store receipt events with lot and temp linkage",
                   "BRZ.RC.RECEIPT · BRZ.RC.LOT_TRACE"],
                  ["DIM.Item",           "Slowly-changing item master (SCD2)",
                   "BRZ.RC.ITEM_MASTER"],
                  ["DIM.Vendor",         "Slowly-changing vendor master (SCD2)",
                   "BRZ.RC.VENDOR_MASTER"],
                  ["DIM.Household",      "Hashed household dimension with consent state",
                   "BRZ.RC.LOYALTY_HH"],
              ],
              col_widths=[1.6, 2.6, 2.2])

    pagebreak(doc)

    # ============ SECTION 2 · SOR INVENTORY ============
    add_h(doc, "2 · SOR inventory — twelve Bronze tables", level=1)
    add_body(doc,
             "The twelve tables below are the complete Bronze surface for RC-E2E-03 and "
             "RC-E2E-09. Each table is fed by one or more upstream source systems. Detailed "
             "field specifications are in Section 4.")
    add_table(doc,
              ["Bronze table", "Friendly name", "Source system", "SOR owner (Kroger)", "Service primary"],
              [list(s) for s in SOR_INVENTORY],
              col_widths=[1.7, 1.7, 1.4, 1.4, 1.0])

    add_callout(doc, "What's in scope vs. out of scope",
                "All twelve tables ingest in Wave 1 (read-only foundation), so the canonical "
                "Silver layer is intact when Wave 2 turns on the Approve and Act agents. The "
                "BRZ.RC.LOYALTY_HH feed is read-only from 84.51° — Kroger Plus operates on its "
                "own platform; APEX consumes the household dimension for elasticity calibration. "
                "Out of scope for this document: HR, Finance, Legal, Real Estate, and customer-"
                "facing channels (digital, web, mobile) — those land in separate SOR documents "
                "for the Tier-2 and Tier-3 services.",
                accent=GOLD)

    pagebreak(doc)

    # ============ SECTION 3 · INGEST PATTERNS ============
    add_h(doc, "3 · Ingest patterns and latency targets", level=1)
    add_body(doc,
             "Different Bronze tables have different latency and freshness needs. The matrix "
             "below states the ingest pattern, refresh cadence, and target end-to-end "
             "latency from source-event to Bronze-landed.")
    add_table(doc,
              ["Bronze table", "Ingest pattern", "Refresh", "Bronze latency SLO"],
              [
                  ["BRZ.RC.POS_TXN",          "CDC via Eventstream from POS gateway",          "Streaming",      "≤ 60 sec from POS commit"],
                  ["BRZ.RC.ITEM_MASTER",      "Daily batch via Data Factory · CDC nightly",    "Daily 02:00 UTC", "T+24h"],
                  ["BRZ.RC.PRICING",          "CDC from pricing engine + nightly snapshot",    "Streaming + nightly", "≤ 5 min for changes"],
                  ["BRZ.RC.PROMO_MARKDOWN",   "Daily batch + on-publish trigger",              "Daily + event",  "≤ 15 min from publish"],
                  ["BRZ.RC.INVENTORY",        "Hourly snapshot + delta CDC",                   "Hourly",         "≤ 90 min"],
                  ["BRZ.RC.REFRIG_TELEMETRY", "Streaming via Eventstream (MQTT bridge)",       "Streaming",      "≤ 30 sec for sensor reads"],
                  ["BRZ.RC.REFRIG_SERVICE",   "REST polling from vendor portals · 15 min",     "15-min poll",    "≤ 30 min"],
                  ["BRZ.RC.FOOD_SAFETY",      "CDC from QMS · Eventstream",                    "Streaming",      "≤ 60 sec from incident logged"],
                  ["BRZ.RC.LOT_TRACE",        "Mixed — supplier EDI batch + Trace One stream", "Daily + event",  "≤ 4 hr (FSMA 204 SLA)"],
                  ["BRZ.RC.VENDOR_MASTER",    "Weekly batch via Data Factory",                  "Weekly Sun",     "T+7d acceptable"],
                  ["BRZ.RC.RECEIPT",          "CDC from WMS · Eventstream",                    "Streaming",      "≤ 5 min from dock-in"],
                  ["BRZ.RC.LOYALTY_HH",       "Daily snapshot from 84.51° (read-only)",        "Daily 04:00 UTC", "T+24h"],
              ],
              col_widths=[1.8, 2.4, 1.3, 1.5])

    add_h(doc, "3.1 · Why streaming for these specific feeds", level=2)
    add_bullets(doc, [
        "POS_TXN — drives elasticity calibration and within-day inventory; agents need fresh data for cold-chain markdown decisioning at the moment of decision.",
        "REFRIG_TELEMETRY — the trigger for the Assess agent; ≤ 30-second latency is the difference between catching an excursion at minute 5 versus minute 25.",
        "FOOD_SAFETY — every disposition decision becomes a closeout audit row; latency here is the lower bound on FSMA 204 evidence freshness.",
        "PRICING — when Pricing publishes a zone refresh, every downstream agent decision needs the new price within minutes.",
        "RECEIPT — lot-trace establishment at the receiving moment; SCML.Lot in Silver depends on this freshness for any later disposition.",
    ])

    add_h(doc, "3.2 · Idempotency and reprocessing", level=2)
    add_body(doc,
             "Every Bronze table is idempotent on (source_system, source_record_id). Reprocessing "
             "a batch produces no duplicates because the dedup operates on the natural key from "
             "source. The batch_id field allows targeted reprocess of a specific window without "
             "rebuilding the whole partition. This matters for FSMA 204 inspection windows where "
             "we may need to replay a quarter of lot-trace data on demand.")

    add_h(doc, "3.3 · Schema evolution discipline", level=2)
    add_body(doc,
             "Bronze schema evolution is additive only — net-new fields land without breaking "
             "anything downstream. Field renames happen in Silver, not Bronze. Field deletes "
             "happen by deprecation (mark not-required, default NULL, document end-of-life "
             "date). Source-system breaking changes (e.g., POS migration from Toshiba to NCR) "
             "land as a new source_system value with both legacy and new feeds running side-by-"
             "side until cutover.")

    pagebreak(doc)

    # ============ SECTION 4 · BRONZE TABLE SPECIFICATIONS ============
    add_h(doc, "4 · Bronze table specifications", level=1)
    add_body(doc,
             "Detailed field-level specifications for all twelve Bronze tables. Each table "
             "block names source system, refresh cadence, expected volume, partition + Z-order "
             "guidance, and the full field list with type and key-marker.")

    # ----- 4.1 POS TXN
    add_h(doc, "4.1 · BRZ.RC.POS_TXN — POS transactions", level=2)
    add_body(doc,
             "Highest-volume Bronze table. Drives elasticity calibration in MERML.Elasticity, "
             "post-event outcome measurement, and the link from sales to lot-traceability via "
             "the lot_code carried through SCML.Lot. PII handling: loyalty_household_id is a "
             "hashed identifier from 84.51° — clear-text PII never lands here.")
    add_field_table(doc, "BRZ.RC.POS_TXN", "POS transactions", POS_TXN_FIELDS,
        source_system="Toshiba 4690 · IBM TCx · NCR Counterpoint (CDC via Eventstream)",
        refresh="Streaming — sub-minute target",
        volume="~25M lines/day at full Kroger NA enterprise scale",
        latency_target="≤ 60 sec from POS commit to Bronze-landed",
        partition_by="ingest_date (yyyy-MM-dd)",
        zorder_by="store_id, transaction_ts",
        notes="PII guard: loyalty_household_id is hashed at source. Bronze never receives clear-text customer PII.")

    pagebreak(doc)

    # ----- 4.2 ITEM MASTER
    add_h(doc, "4.2 · BRZ.RC.ITEM_MASTER — Item master / PIM", level=2)
    add_body(doc,
             "Slowly-changing item dimension. SCD Type 2 in Silver (DIM.Item). "
             "ftl_flag is the boolean that tells RC-E2E-09 whether this SKU requires Critical "
             "Tracking Event ingestion under FSMA 204 final rule.")
    add_field_table(doc, "BRZ.RC.ITEM_MASTER", "Item master / PIM", ITEM_MASTER_FIELDS,
        source_system="Item Master / PIM · daily batch via Data Factory + nightly CDC",
        refresh="Daily 02:00 UTC + nightly CDC delta",
        volume="~1.2M active SKUs across banners; ~5K daily mutations",
        latency_target="T+24h",
        partition_by="ingest_date",
        zorder_by="sku, gtin, category_l1",
        notes="SCD Type 2 boundary lives in Silver DIM.Item — Bronze stores every revision as an append.")

    # ----- 4.3 PRICING
    add_h(doc, "4.3 · BRZ.RC.PRICING — Pricing & price-zone state", level=2)
    add_body(doc,
             "Per-SKU, per-zone (or per-store) pricing record with effective windows. The most "
             "frequently-mutated dimensional feed.")
    add_field_table(doc, "BRZ.RC.PRICING", "Pricing & price-zone state", PRICING_FIELDS,
        source_system="Revionics + 84.51° pricing engine · CDC from pricing-system · nightly full",
        refresh="Streaming + nightly snapshot",
        volume="~50M active price records · ~2M daily mutations",
        latency_target="≤ 5 min from pricing-system commit",
        partition_by="ingest_date",
        zorder_by="sku, price_zone_id, effective_date",
        notes="Joins to BRZ.RC.PROMO_MARKDOWN on promo_id; promo overlays are layered on top of list_price in Silver MERML.Price.")

    pagebreak(doc)

    # ----- 4.4 PROMO MARKDOWN
    add_h(doc, "4.4 · BRZ.RC.PROMO_MARKDOWN — Promotional & markdown events", level=2)
    add_body(doc,
             "Every promotion or markdown that affects published price for a window of time. "
             "Distinct from BRZ.RC.PRICING because a promotion is a temporal event with start "
             "and end timestamps, whereas pricing is a steady-state record. A SKU can have "
             "multiple active promotions overlaid (e.g., a TPR + a loyalty offer).")
    add_field_table(doc, "BRZ.RC.PROMO_MARKDOWN", "Promotional & markdown events", PROMO_MARKDOWN_FIELDS,
        source_system="Revionics Markdown · ad-system · loyalty platform feeds",
        refresh="Daily batch + on-publish trigger",
        volume="~150K active promotions at any time · ~10K weekly mutations",
        latency_target="≤ 15 min from publish",
        partition_by="ingest_date",
        zorder_by="promo_type, start_ts, sku_group_id",
        notes="MERML.Markdown derives from this feed joined to BRZ.RC.POS_TXN for outcome attribution.")

    # ----- 4.5 INVENTORY
    add_h(doc, "4.5 · BRZ.RC.INVENTORY — Store + DC inventory positions", level=2)
    add_body(doc,
             "Hourly snapshots plus delta CDC. Captures on-hand, on-order, in-transit, and "
             "allocated quantities at both store and DC level.")
    add_field_table(doc, "BRZ.RC.INVENTORY", "Inventory positions (store + DC)", INVENTORY_FIELDS,
        source_system="SAP IM (DC) + Manhattan WMS · hourly snapshot + delta CDC",
        refresh="Hourly snapshot + delta on movements",
        volume="~3.6B SKU-location-day records / quarter",
        latency_target="≤ 90 min for hourly cycle; deltas ≤ 5 min",
        partition_by="ingest_date",
        zorder_by="location_id, sku, snapshot_ts",
        notes="Silver SCML.Inventory layers in receipt and POS-depletion to give a derived 'true' on-hand for agent decisioning.")

    pagebreak(doc)

    # ----- 4.6 REFRIG TELEMETRY
    add_h(doc, "4.6 · BRZ.RC.REFRIG_TELEMETRY — Refrigeration telemetry", level=2)
    add_body(doc,
             "The single most latency-sensitive feed in this engagement. The Assess agent's "
             "trigger surface — every excursion decision starts with a temperature reading "
             "crossing setpoint. Sensor reads at 30-second cadence from in-store EMS gateways "
             "(typically Emerson E2E and Honeywell Niagara, OPC-UA over MQTT bridge to "
             "Eventstream).")
    add_field_table(doc, "BRZ.RC.REFRIG_TELEMETRY", "Refrigeration sensor telemetry", REFRIG_TELEMETRY_FIELDS,
        source_system="Emerson E2E · Honeywell Niagara · Hill PHOENIX EMS · OPC-UA over MQTT → Eventstream",
        refresh="Streaming · 30-second sensor cadence",
        volume="~1.8 fixtures × 30K sensors × 2,800 stores × 2,880 reads/day = ~1.3B reads/day enterprise",
        latency_target="≤ 30 sec from sensor read to Bronze-landed",
        partition_by="ingest_date · store_id (hash-bucket on hot stores if needed)",
        zorder_by="fixture_id, reading_ts",
        notes="No _is_deleted (telemetry never logically deletes). Highest-volume Bronze table; Z-order is critical for query performance. Consider hot-storage tiering for the most recent 7 days.")

    # ----- 4.7 REFRIG SERVICE
    add_h(doc, "4.7 · BRZ.RC.REFRIG_SERVICE — Refrigeration vendor service tickets", level=2)
    add_body(doc,
             "Service-ticket events from refrigeration vendor portals. Establishes the "
             "fixture-failure-history that informs root-cause analysis on excursion events. "
             "REST polling rather than streaming — vendor portals don't typically expose CDC.")
    add_field_table(doc, "BRZ.RC.REFRIG_SERVICE", "Refrigeration vendor service tickets", REFRIG_SERVICE_FIELDS,
        source_system="Emerson · Lennox · Hussmann · Hill PHOENIX vendor portals (REST)",
        refresh="15-minute polling cycle",
        volume="~2K ticket events/week enterprise",
        latency_target="≤ 30 min from vendor-portal commit",
        partition_by="ingest_date",
        zorder_by="store_id, fixture_id, open_ts",
        notes="Vendor-portal API rate limits constrain polling cadence; coordinate with vendor account managers to align SLAs.")

    pagebreak(doc)

    # ----- 4.8 FOOD SAFETY
    add_h(doc, "4.8 · BRZ.RC.FOOD_SAFETY — Food-safety incident & disposition log", level=2)
    add_body(doc,
             "Every disposition decision (sell-through, markdown, destroy, hold) and its "
             "approver, approval timestamp, and HITL kind. The closeout-side feed for "
             "RC-E2E-03; the audit-row evidence for RC-E2E-09 inspections. Includes the APEX "
             "trace_id correlator that links back to the agent decision audit row.")
    add_field_table(doc, "BRZ.RC.FOOD_SAFETY", "Food-safety incidents & dispositions", FOOD_SAFETY_FIELDS,
        source_system="Quality Mgmt System (Trace One QMS or in-house) · CDC via Eventstream",
        refresh="Streaming",
        volume="~24K incidents/week enterprise (~85/banner/wk × 8 banners + non-cold-chain)",
        latency_target="≤ 60 sec from incident logged",
        partition_by="ingest_date",
        zorder_by="store_id, incident_ts, incident_type",
        notes="trace_id field is the bridge between Bronze food-safety log and APEX LEDGER audit row — never null for excursions.")

    # ----- 4.9 LOT TRACE
    add_h(doc, "4.9 · BRZ.RC.LOT_TRACE — FSMA 204 lot-provenance events", level=2)
    add_body(doc,
             "The FSMA 204 final rule lot-traceability surface. Critical Tracking Events (CTEs) "
             "for SKUs on the FDA Food Traceability List, with Key Data Elements (KDEs) carried "
             "as a versioned JSON payload. Mixed ingest pattern: supplier-EDI batch for some "
             "categories, Trace One stream for others, and direct supplier-portal pulls for the "
             "long tail.")
    add_field_table(doc, "BRZ.RC.LOT_TRACE", "FSMA 204 lot-provenance CTE events", LOT_TRACE_FIELDS,
        source_system="Trace One · FoodLogiQ · supplier EDI · supplier-portal pulls",
        refresh="Mixed: streaming where available, daily batch otherwise",
        volume="~600K CTE events/week at FSMA 204 full-coverage",
        latency_target="≤ 4 hr (FSMA 204 SLA driver) · ≤ 15 min for streaming feeds",
        partition_by="ingest_date · ftl_category",
        zorder_by="lot_code, gtin, event_ts",
        notes="kde_payload schema versioning (kde_schema_version) is critical — FSMA 204 KDE format will evolve and Bronze must preserve every revision.")

    pagebreak(doc)

    # ----- 4.10 VENDOR MASTER
    add_h(doc, "4.10 · BRZ.RC.VENDOR_MASTER — Vendor master + scorecards", level=2)
    add_body(doc,
             "Slowly-changing vendor dimension. Carries the FTL capability tier — which "
             "suppliers can produce KDEs in which formats. The Category Manager and Vendor "
             "Manager use this to drive supplier qualification decisions for FSMA 204 readiness.")
    add_field_table(doc, "BRZ.RC.VENDOR_MASTER", "Vendor master + scorecards", VENDOR_MASTER_FIELDS,
        source_system="Vendor Master · supplier portal · 84.51° supplier scorecard feeds",
        refresh="Weekly batch (Sunday 03:00 UTC)",
        volume="~25K active vendor records",
        latency_target="T+7d acceptable",
        partition_by="ingest_date",
        zorder_by="vendor_id, vendor_type",
        notes="SCD Type 2 boundary in Silver DIM.Vendor — Bronze stores every revision append-only.")

    # ----- 4.11 RECEIPT
    add_h(doc, "4.11 · BRZ.RC.RECEIPT — DC + store receipt events", level=2)
    add_body(doc,
             "Every receipt event with linked lot_code. The bridge between vendor-side "
             "lot-trace data and the internal supply-chain. temp_at_receipt_f is the cold-chain "
             "compliance evidence at the dock.")
    add_field_table(doc, "BRZ.RC.RECEIPT", "DC + store receipts (with lot codes)", RECEIPT_FIELDS,
        source_system="Manhattan WMS · DC-receiving · store-receiving · CDC via Eventstream",
        refresh="Streaming",
        volume="~140K receipt events/day across DCs and stores",
        latency_target="≤ 5 min from dock-in",
        partition_by="ingest_date",
        zorder_by="location_id, lot_code, receipt_ts",
        notes="Joins to BRZ.RC.LOT_TRACE on lot_code; one-to-many — a single lot may have multiple receipt events across the supply chain.")

    pagebreak(doc)

    # ----- 4.12 LOYALTY HH
    add_h(doc, "4.12 · BRZ.RC.LOYALTY_HH — Loyalty households (read-only)", level=2)
    add_body(doc,
             "Hashed-household feed from 84.51°. Used by RC-E2E-03 to calibrate elasticity "
             "curves in MERML.Elasticity (basket-level signal vs. transaction-level). Read-"
             "only — Kroger's loyalty platform is the system of record. Consent state is "
             "respected at the canonical layer (decisions don't personalize for "
             "consent_state = 'WITHDRAWN').")
    add_field_table(doc, "BRZ.RC.LOYALTY_HH", "Loyalty household dimension", LOYALTY_HH_FIELDS,
        source_system="84.51° loyalty platform · daily snapshot",
        refresh="Daily 04:00 UTC",
        volume="~60M+ active hashed households",
        latency_target="T+24h",
        partition_by="ingest_date",
        zorder_by="loyalty_household_id, primary_banner",
        notes="No clear-text customer PII. consent_state drives downstream personalization-eligibility. Privacy/consent governance is co-owned with the 84.51° platform team.")

    pagebreak(doc)

    # ============ SECTION 5 · BRONZE → SILVER MAPPING ============
    add_h(doc, "5 · Bronze → Silver canonical mapping", level=1)
    add_body(doc,
             "How the twelve Bronze tables compose into the canonical Silver layer that APEX "
             "agents read. This is the contract: when Bronze is correct, the mapping below "
             "produces a Silver state that the agents can reason against without going back to "
             "source systems.")

    add_h(doc, "5.1 · MERML — Merchandising canonical family", level=2)
    add_table(doc,
              ["Silver shape", "Composition", "Primary key", "Update pattern"],
              [
                  ["MERML.Price",
                   "BRZ.RC.PRICING base + BRZ.RC.PROMO_MARKDOWN overlay (active windows) + BRZ.RC.ITEM_MASTER (current SCD2 row) + BRZ.RC.LOYALTY_HH (consent-aware personalization tier)",
                   "(sku, price_zone_id, valid_from)",
                   "SCD Type 2 on price changes; on-publish from Bronze CDC"],
                  ["MERML.Markdown",
                   "BRZ.RC.PROMO_MARKDOWN where promo_type ∈ (MARKDOWN, CLEARANCE, TPR) joined to BRZ.RC.POS_TXN for outcome attribution",
                   "(promo_id, sku)",
                   "Append on promo publish + outcome attribution at end_ts + 7 days"],
                  ["MERML.Elasticity",
                   "Per-SKU elasticity curves derived from BRZ.RC.POS_TXN (price × demand history) + BRZ.RC.PRICING (price points) + BRZ.RC.LOYALTY_HH (basket-level segmentation)",
                   "(sku, price_zone_id, model_version)",
                   "Recompute weekly + ad-hoc on category-mix change"],
                  ["MERML.Promotion",
                   "BRZ.RC.PROMO_MARKDOWN where promo_type ∈ (TPR, BOGO, AD_BLOCK, CLUB) joined to BRZ.RC.AD_BLOCK_CALENDAR (TBD — out of scope here)",
                   "(promo_id)",
                   "Append on publish + outcome at end_ts + 14 days"],
              ],
              col_widths=[1.6, 2.7, 1.3, 1.5])

    add_h(doc, "5.2 · SCML — Supply Chain & Lot canonical family", level=2)
    add_table(doc,
              ["Silver shape", "Composition", "Primary key", "Update pattern"],
              [
                  ["SCML.Lot",
                   "BRZ.RC.LOT_TRACE (CTE events) joined to BRZ.RC.RECEIPT (location lineage) + BRZ.RC.VENDOR_MASTER (supplier dimension) + BRZ.RC.ITEM_MASTER (SKU dimension)",
                   "(lot_code, gtin)",
                   "Append on every CTE; lot lifecycle window from production to disposition"],
                  ["SCML.Inventory",
                   "BRZ.RC.INVENTORY base + BRZ.RC.RECEIPT (additions) + BRZ.RC.POS_TXN (depletions) reconciled to a derived true-state per SKU-location-time",
                   "(location_id, sku, valid_from)",
                   "Reconcile hourly; agent-decision-time queries hit the latest valid state"],
                  ["SCML.Fixture",
                   "BRZ.RC.REFRIG_TELEMETRY rolled up to fixture-state windows + BRZ.RC.REFRIG_SERVICE (failure history) + DIM.Store",
                   "(fixture_id, valid_from)",
                   "Window-based aggregation every 5 min; excursion classification on the fly"],
                  ["SCML.Receipt",
                   "BRZ.RC.RECEIPT joined to BRZ.RC.LOT_TRACE (provenance lineage) + BRZ.RC.VENDOR_MASTER (supplier qualification at receipt time)",
                   "(receipt_id)",
                   "Append on every dock-in event; one-to-one with Bronze"],
                  ["SCML.Disposition",
                   "BRZ.RC.FOOD_SAFETY closeout records joined to BRZ.RC.LOT_TRACE (affected lots) + BRZ.RC.POS_TXN (post-disposition sales for sell-through outcome)",
                   "(incident_id)",
                   "Append on every disposition; outcome attribution at incident + 7 days"],
              ],
              col_widths=[1.6, 2.7, 1.3, 1.5])

    add_h(doc, "5.3 · Dimensional helpers", level=2)
    add_table(doc,
              ["Silver dimension", "Composition", "SCD"],
              [
                  ["DIM.Item",      "BRZ.RC.ITEM_MASTER · all rows preserved as SCD2 history",            "Type 2"],
                  ["DIM.Vendor",    "BRZ.RC.VENDOR_MASTER · all rows preserved as SCD2 history",          "Type 2"],
                  ["DIM.Store",    "Out of this document's scope — assumed pre-existing in Silver",       "Type 1"],
                  ["DIM.Banner",    "Static reference — out of this document's scope",                    "Type 0"],
                  ["DIM.Household", "BRZ.RC.LOYALTY_HH · current-state with consent flag",                "Type 1 (privacy-respecting)"],
                  ["DIM.PriceZone", "Out of this document's scope — assumed pre-existing in Silver",      "Type 1"],
                  ["DIM.Calendar",  "Static reference — out of this document's scope",                    "Type 0"],
              ],
              col_widths=[1.6, 4.0, 1.0])

    pagebreak(doc)

    # ============ SECTION 6 · IMPLEMENTATION CONSIDERATIONS ============
    add_h(doc, "6 · Implementation considerations", level=1)

    add_h(doc, "6.1 · Storage tiering and retention", level=2)
    add_table(doc,
              ["Bronze table", "Hot tier (hot storage)", "Warm tier", "Cold archive"],
              [
                  ["BRZ.RC.POS_TXN",          "30 days",  "13 months",  "7 years (SOX / audit)"],
                  ["BRZ.RC.ITEM_MASTER",      "90 days",  "Indefinite (all SCD2 history)",  "—"],
                  ["BRZ.RC.PRICING",          "90 days",  "Indefinite",  "—"],
                  ["BRZ.RC.PROMO_MARKDOWN",   "90 days",  "13 months",  "7 years"],
                  ["BRZ.RC.INVENTORY",        "30 days",  "13 months",  "—"],
                  ["BRZ.RC.REFRIG_TELEMETRY", "7 days",   "90 days",     "Aggregate-only after 90d"],
                  ["BRZ.RC.REFRIG_SERVICE",   "90 days",  "5 years",     "—"],
                  ["BRZ.RC.FOOD_SAFETY",      "180 days", "Indefinite",  "Indefinite (FDA evidence)"],
                  ["BRZ.RC.LOT_TRACE",        "180 days", "Indefinite",  "Indefinite (FSMA 204 24-month minimum)"],
                  ["BRZ.RC.VENDOR_MASTER",    "180 days", "Indefinite",  "—"],
                  ["BRZ.RC.RECEIPT",          "30 days",  "5 years",     "—"],
                  ["BRZ.RC.LOYALTY_HH",       "30 days",  "Indefinite (privacy-aware)", "—"],
              ],
              col_widths=[2.0, 1.4, 1.4, 1.7])

    add_h(doc, "6.2 · Data quality discipline", level=2)
    add_bullets(doc, [
        "Every Bronze table has a Great Expectations (or equivalent Fabric DQ pack) suite at land-time. Failed expectations route to a quarantine area, not to Silver.",
        "Schema-drift detection runs nightly; net-new fields are auto-accepted, type changes raise a PR review.",
        "Referential integrity checks run nightly between Bronze tables (e.g., POS_TXN.sku must exist in ITEM_MASTER active window).",
        "Late-arriving data discipline: Bronze accepts late events with the original event_ts preserved; Silver re-projects affected windows.",
        "Telemetry-quality SLOs: REFRIG_TELEMETRY ≥ 99.5% completeness per fixture per day; below threshold raises Audience-Lead-equivalent alert.",
    ])

    add_h(doc, "6.3 · Privacy, consent, and PII handling", level=2)
    add_bullets(doc, [
        "No clear-text customer PII lands in Bronze. Loyalty household_id is hashed at source.",
        "BRZ.RC.LOYALTY_HH carries consent_state; downstream Silver views filter on it for personalization-tier decisions.",
        "BRZ.RC.FOOD_SAFETY captures approver_id (Entra identity) — internal employees, not customers — and is governed by HR data-protection rules.",
        "FSMA 204 lot-trace data is supplier-confidential; access controls in Silver and Gold limit visibility to Food-Safety / Compliance / Vendor-Mgmt roles.",
        "All Bronze access is logged via Microsoft Purview; LEDGER audit-row generation references the Purview log for evidence-of-access proof.",
    ])

    add_h(doc, "6.4 · Operational runbook", level=2)
    add_table(doc,
              ["Runbook entry", "Trigger", "Action"],
              [
                  ["POS feed delay > 60 sec",            "Streaming-lag alert",       "Page on-call data engineer · escalate after 5 min"],
                  ["Refrig telemetry gap > 5 min",       "Telemetry-quality alert",   "Page Facilities Eng + on-call · check MQTT bridge"],
                  ["FSMA 204 KDE schema evolution",      "Source publishes new ver",  "Open PR to Bronze schema · run shadow ingest 7 days"],
                  ["Vendor-portal API timeout",          "Polling cycle fails",       "Auto-retry 3x · open ticket with vendor account mgr"],
                  ["Recall event in BRZ.RC.LOT_TRACE",   "FDA / supplier recall",     "Auto-trigger SCML.Lot recall-window query · alert Food Safety Lead"],
                  ["Schema drift detected (Bronze)",     "Nightly DQ run",            "Auto-PR with proposed addition · review by Data Eng"],
                  ["Late-arriving data > 24h late",      "Watermark check",           "Re-project Silver windows · log to LEDGER"],
                  ["Bronze partition reprocess",         "Source-system replay",      "Use batch_id targeting · validate via dedup"],
              ],
              col_widths=[2.4, 1.8, 2.3])

    pagebreak(doc)

    # ============ SECTION 7 · WAVE-BY-WAVE INGEST PRIORITIZATION ============
    add_h(doc, "7 · Wave-by-wave ingest prioritization", level=1)
    add_body(doc,
             "Not all twelve Bronze tables need to land in Wave 1. The sequence below maps "
             "Bronze landings to the Wave plan from APEX-RC-E2E-03-Walkthrough.docx Section 9.")
    add_table(doc,
              ["Wave", "Bronze tables landed", "Why this set is sufficient"],
              [
                  ["W1 Foundation (single-banner pilot)",
                   "ITEM_MASTER · PRICING · REFRIG_TELEMETRY · FOOD_SAFETY · LOYALTY_HH (read-only)",
                   "Read-only intelligence on cold-chain + markdown. Assess and Classify can run; Quantify needs MERML.Elasticity which derives from POS_TXN below."],
                  ["W1 Foundation extension",
                   "+ POS_TXN (single banner scope) · INVENTORY (snapshot only)",
                   "Enables Quantify and shadow-run elasticity calibration."],
                  ["W2 Pilot (250-store live)",
                   "+ PROMO_MARKDOWN · REFRIG_SERVICE · LOT_TRACE · VENDOR_MASTER · RECEIPT",
                   "Full five-agent orchestration. FSMA 204 evidence packet generation possible."],
                  ["W3 Scale & Fuse",
                   "All twelve at enterprise scale + cross-banner aggregation",
                   "Enterprise run-rate; FSMA 204 full coverage; cross-scenario fusion candidates."],
              ],
              col_widths=[2.0, 2.6, 2.0])

    add_callout(doc, "Sequencing rationale",
                "ITEM_MASTER and PRICING are foundational dimensional feeds — every other "
                "Bronze table joins to them in Silver. REFRIG_TELEMETRY is the trigger surface "
                "for the Assess agent. FOOD_SAFETY is the closeout-side feed. LOYALTY_HH is "
                "needed early for elasticity calibration but is read-only and doesn't impose "
                "platform-engineering effort. POS_TXN comes second-tier in W1 because the "
                "transaction volume is large and the elasticity work depends on it but doesn't "
                "block the read-only intelligence story. PROMO_MARKDOWN, REFRIG_SERVICE, "
                "LOT_TRACE, VENDOR_MASTER, RECEIPT all light up in W2 when the full agent fleet "
                "goes live.")

    pagebreak(doc)

    # ============ APPENDIX ============
    add_h(doc, "Appendix A · Glossary", level=1)
    add_table(doc,
              ["Term", "Definition"],
              [
                  ["Bronze (medallion)",     "Raw, source-aligned, append-only landing tier. Preserves source schema and adds ingest metadata."],
                  ["Silver (medallion)",     "Canonical, conformed, deduplicated tier. APEX agents read from here. MERML and SCML schema families."],
                  ["Gold (medallion)",       "Curated KPI marts and Power BI semantic models for executive consumption."],
                  ["LEDGER",                 "APEX 14-field audit-row store. Sits alongside Gold; fed by every agent decision."],
                  ["MERML",                  "APEX canonical schema family for Merchandising — Price, Markdown, Elasticity, Promotion."],
                  ["SCML",                   "APEX canonical schema family for Supply Chain & Lot — Lot, Inventory, Fixture, Receipt, Disposition."],
                  ["CTE (Critical Tracking Event)", "FSMA 204 final-rule term for a logged supply-chain event — receiving, transforming, shipping, etc."],
                  ["KDE (Key Data Element)", "FSMA 204 final-rule term for the data fields the FDA requires for each CTE."],
                  ["FTL (Food Traceability List)", "FDA list of foods subject to FSMA 204 lot-traceability requirements."],
                  ["SCD Type 2",             "Slowly-changing dimension pattern that preserves history with effective-date windows."],
                  ["Eventstream",            "Microsoft Fabric streaming ingest service. Used for low-latency Bronze feeds."],
                  ["Direct Lake",            "Microsoft Fabric semantic-model engine that reads Delta files directly without import."],
                  ["OneLake",                "Microsoft Fabric storage layer. All Bronze/Silver/Gold tables live in OneLake as Delta."],
                  ["MQTT",                   "Lightweight pub/sub protocol used by refrigeration EMS gateways to publish telemetry."],
                  ["OPC-UA",                 "Industrial-control standard for sensor data; bridged to MQTT for cloud ingest."],
                  ["Trace One / FoodLogiQ",  "Two of the major lot-traceability platforms used in grocery for supplier integration."],
                  ["84.51°",                 "Kroger's data-and-analytics subsidiary. Loyalty-platform source for BRZ.RC.LOYALTY_HH."],
                  ["Revionics",              "Kroger's pricing-engine vendor (one possibility). Source for BRZ.RC.PRICING and BRZ.RC.PROMO_MARKDOWN."],
                  ["Manhattan WMS",          "Likely Kroger DC warehouse-management system. Source for receipt and inventory feeds."],
                  ["Emerson E2E / Niagara",  "Two of the most common refrigeration EMS gateway platforms. Source for telemetry."],
              ],
              col_widths=[1.7, 4.6])

    add_body(doc,
             "Independence posture preserved: 'Deloitte's Microsoft practice,' 'DMTSP,' "
             "'Microsoft platform capabilities,' 'Microsoft-native deployment.' Source-system "
             "vendor names (Revionics, Trace One, Emerson, etc.) are illustrative — confirm "
             "actual Kroger source-system inventory in Approach-stage discovery.",
             italic=True, color=MUTE, size=9)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— End of SOR document —"); set_run(r, italic=True, size=10, color=MUTE)

    # save
    target = OUT
    try:
        doc.save(str(target))
    except PermissionError:
        n = 2
        while True:
            alt = target.with_name(target.stem + f"-v{n}" + target.suffix)
            if not alt.exists(): target = alt; break
            n += 1
        doc.save(str(target))
        print(f"[notice] {OUT.name} was locked — saved as {target.name} instead.")
    print(f"Wrote {target}")


if __name__ == "__main__":
    build()
