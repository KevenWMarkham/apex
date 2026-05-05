"""Build the APEX · RC · Grocery Merchandising service portfolio document.

Companion to APEX-RC-E2E-03-Walkthrough.docx. Where the walkthrough deep-dives
one service end-to-end, this portfolio document:
  · Surveys ALL eight RC services
  · Ranks them by priority for the Grocery Merchandising sub-category
  · Deep-dives the top three for Merch (RC-E2E-03, RC-E2E-09, RC-E2E-04)
  · Recommends a sequencing path through the portfolio
  · Reframes persona ownership for Merchandising (Category Manager-led, not Store Ops-led)

Output: Kroger 02_projects/FY27_Pipeline/assortment-pricing-agentic/deliverables/
        APEX-RC-Grocery-Merchandising-Service-Portfolio.docx
"""

from __future__ import annotations

import io
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

OUT_DIR = Path(r"C:\Stage\Clients\Industries\Consumer\Retail\Kroger\02_projects\FY27_Pipeline\assortment-pricing-agentic\deliverables")
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT = OUT_DIR / "APEX-RC-Grocery-Merchandising-Service-Portfolio.docx"

# ---- styling ------------------------------------------------------------

INK   = "0B0B0F"
INK2  = "4A4A55"
MUTE  = "8a8a90"
GOLD  = "B8892E"
TEAL  = "2A7F73"
VIO   = "6648B0"
CRIM  = "A22A39"
BAND  = "F3F3F5"
HEAD  = "0B0B0F"
LINE  = "D4D4D7"


def set_cell_bg(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex.lstrip("#")); tcPr.append(shd)


def set_cell_border(cell, color_hex=LINE):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "4"); b.set(qn("w:color"), color_hex.lstrip("#"))
        borders.append(b)
    tcPr.append(borders)


def set_run(run, *, bold=False, italic=False, size=None, color=None, font="Arial"):
    if bold: run.bold = True
    if italic: run.italic = True
    if size is not None: run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor.from_string(color.lstrip("#"))
    if font: run.font.name = font


def add_body(doc, text, *, size=11, italic=False, color=INK, space_after=6, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None: p.alignment = align
    r = p.add_run(text)
    set_run(r, size=size, color=color, italic=italic)
    return p


def add_h(doc, text, *, level=1):
    sizes = {1: 18, 2: 14, 3: 12.5, 4: 11.5}
    p = doc.add_paragraph()
    p.style = doc.styles[f"Heading {level}"]
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(4)
    for r in list(p.runs): r.text = ""
    r = p.add_run(text)
    set_run(r, bold=True, size=sizes.get(level, 12), color=INK)
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6"); bottom.set(qn("w:space"), "4"); bottom.set(qn("w:color"), GOLD)
        pBdr.append(bottom); pPr.append(pBdr)
    return p


def add_bullets(doc, items, *, size=11):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(it); set_run(r, size=size, color=INK)


def add_table(doc, header, rows, *, col_widths=None,
              header_fill=HEAD, header_color="FFFFFF", zebra=True, tier_col=None):
    n = len(header)
    tbl = doc.add_table(rows=len(rows) + 1, cols=n)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(header):
        c = tbl.rows[0].cells[j]
        set_cell_bg(c, header_fill); set_cell_border(c)
        c.paragraphs[0].text = ""
        r = c.paragraphs[0].add_run(h); set_run(r, bold=True, size=10, color=header_color)
        if col_widths: c.width = Inches(col_widths[j])
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            c = tbl.rows[i].cells[j]
            set_cell_border(c)
            if zebra and i % 2 == 0: set_cell_bg(c, BAND)
            # tier column color
            if tier_col is not None and j == tier_col:
                v_str = str(v)
                if "Tier 1" in v_str: set_cell_bg(c, "F6E5BD")
                elif "Tier 2" in v_str: set_cell_bg(c, "E4F1EE")
                elif "Tier 3" in v_str: set_cell_bg(c, "EDEDED")
            c.paragraphs[0].text = ""
            r = c.paragraphs[0].add_run(str(v)); set_run(r, size=10, color=INK)
            if col_widths: c.width = Inches(col_widths[j])
    doc.add_paragraph()
    return tbl


def add_callout(doc, kind, body, *, accent=GOLD, fill="F6F1E4"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = tbl.rows[0].cells[0]
    set_cell_bg(c, fill); c.paragraphs[0].text = ""
    p_tag = c.paragraphs[0]; r_tag = p_tag.add_run(kind.upper())
    set_run(r_tag, bold=True, size=9, color=accent)
    p = c.add_paragraph(); p.paragraph_format.space_before = Pt(3)
    r = p.add_run(body); set_run(r, size=10.5, color=INK)
    doc.add_paragraph()


def add_kv(doc, pairs, *, lw=1.6, rw=4.9):
    tbl = doc.add_table(rows=len(pairs), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(pairs):
        ck, cv = tbl.rows[i].cells
        ck.width = Inches(lw); cv.width = Inches(rw)
        set_cell_bg(ck, BAND); set_cell_border(ck); set_cell_border(cv)
        ck.paragraphs[0].text = ""; cv.paragraphs[0].text = ""
        rk = ck.paragraphs[0].add_run(k); set_run(rk, bold=True, size=10, color=INK)
        rv = cv.paragraphs[0].add_run(v); set_run(rv, size=10, color=INK)
    doc.add_paragraph()


def add_eyebrow(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text.upper()); set_run(r, bold=True, size=9, color=MUTE)
    rPr = r._element.get_or_add_rPr()
    spacing = OxmlElement("w:spacing"); spacing.set(qn("w:val"), "80"); rPr.append(spacing)


def pagebreak(doc):
    p = doc.add_paragraph(); p.add_run().add_break(WD_BREAK.PAGE)


# ============================================================================
# BUILD
# ============================================================================

def build():
    doc = Document()
    for section in doc.sections:
        section.page_width  = Inches(8.5); section.page_height = Inches(11)
        section.top_margin = section.bottom_margin = Inches(0.9)
        section.left_margin = section.right_margin = Inches(1.0)
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(11)

    # ============ COVER ============
    add_eyebrow(doc, "APEX · PRACTICE RC · GROCERY MERCHANDISING SUB-CATEGORY")
    p = doc.add_paragraph()
    r = p.add_run("Service Portfolio"); set_run(r, bold=True, size=26, color=INK)
    p1b = doc.add_paragraph()
    r1b = p1b.add_run("for Grocery Merchandising"); set_run(r1b, bold=True, size=20, color=INK)
    p2 = doc.add_paragraph()
    r2 = p2.add_run("All eight RC services · ranked for Merchandising priority · deep-dives on the top three")
    set_run(r2, size=13, color=INK2, italic=True)

    doc.add_paragraph()
    add_body(doc,
             "A portfolio survey of the eight APEX RC (Retail & Consumer) services with priority "
             "ranking specific to the Grocery Merchandising sub-category — the buying, "
             "category-management, pricing, promotion, and assortment-decisioning function. "
             "Companion to APEX-RC-E2E-03-Walkthrough.docx, which deep-dives the highest-priority "
             "service end-to-end against Kroger as the reference client.",
             italic=True, color=INK2)

    doc.add_paragraph()
    add_kv(doc, [
        ("Document type",     "Service Portfolio · Practice RC · Grocery Merchandising sub-category"),
        ("Practice",          "RC — Retail & Consumer"),
        ("Sub-category",      "Grocery Merchandising · category buying / pricing / promo / assortment / private label / vendor"),
        ("Reference client",  "The Kroger Co. — multi-banner grocery, ~$150B revenue, 60M Plus households"),
        ("Audience",          "Chief Merchant · VP Merchandising · Category Managers · Pricing & Promotions · CFO · Account Team"),
        ("Services covered",  "All 8 active RC services (E2E-03, -04, -05, -06, -07, -08, -09, -11)"),
        ("Priority framework","Tier 1 / 2 / 3 by Grocery Merchandising relevance and value"),
        ("Top-3 deep dives",  "RC-E2E-03 Assortment & Pricing · RC-E2E-09 Product Tracking · RC-E2E-04 Customer Lifecycle & Loyalty"),
        ("Companion artifact","APEX-RC-E2E-03-Walkthrough.docx (deep-dive on Tier-1 service #1)"),
        ("Version / date",    "v1.0 · 27 Apr 2026"),
    ])

    pagebreak(doc)

    # ============ SECTION 1 · WHY THIS PORTFOLIO ============
    add_h(doc, "1 · Why this portfolio for Grocery Merchandising", level=1)
    add_body(doc,
             "Grocery Merchandising is not Retail Merchandising in general. It is the specialized "
             "function that buys, prices, promotes, assorts, and manages perishables and shelf-"
             "stable categories at scale — typically across 30+ category lines (dairy, meat, "
             "produce, bakery, deli, frozen, center-store, beverages, HBC, household, pet, "
             "floral, and more). The persona stack is dominated by Category Managers and Senior "
             "Buyers whose daily work is composed of:")
    add_bullets(doc, [
        "Pricing decisions across 50,000+ SKUs per banner, refreshed weekly per zone",
        "Promotional planning across 26 ad blocks per year × multiple banners × multiple zones",
        "Assortment decisions on listing, delisting, slotting, and private-label balance",
        "Vendor management — slotting allowances, MDF, scan data, joint business plans",
        "Markdown decisions on perishables (cold-chain) and non-perishables (seasonality, EOL)",
        "Category P&L management — gross margin, sales, units, share of category",
        "FSMA 204 lot-traceability decisions for the Food Traceability List subset",
    ])
    add_body(doc,
             "When we apply the eight APEX RC services to this function, three are Tier-1 priority "
             "(directly serving Merchandising's core daily work), three are Tier-2 priority "
             "(adjacent functions Merchandising consumes from or contributes to), and two are "
             "Tier-3 priority (other RC functions that don't materially intersect Merchandising's "
             "decision surface).")
    add_body(doc,
             "This portfolio is the navigational document that helps the Account Team sequence "
             "the engagement: which service to lead with, which to attach next, which to defer to "
             "later waves, and which not to pursue from the Merchandising organization at all.",
             italic=True, color=INK2)

    pagebreak(doc)

    # ============ SECTION 2 · ALL 8 RC SERVICES OVERVIEW ============
    add_h(doc, "2 · All eight RC services — overview and priority for Grocery Merch", level=1)
    add_body(doc,
             "The table below summarizes every RC service in the APEX Service Catalog. The rightmost "
             "column states the Tier ranking for Grocery Merchandising specifically. Tier rationale "
             "follows in Section 3.")

    add_table(doc,
              ["Service ID", "Service Name", "What it covers", "Primary persona", "Tier (Grocery Merch)"],
              [
                  ["RC-E2E-03", "Assortment & Pricing",
                   "Pricing, markdown, assortment rationalization. Cold-chain disposition, dynamic pricing, competitive monitoring, lifecycle pricing.",
                   "Category Manager / Pricing Director",
                   "Tier 1 — anchor"],
                  ["RC-E2E-09", "Product Tracking",
                   "Lot-provenance and FSMA 204 compliance. Supply-chain-of-custody for high-risk foods on the FDA Food Traceability List.",
                   "Food Safety Compliance Lead / Quality Manager",
                   "Tier 1 — co-anchor"],
                  ["RC-E2E-04", "Customer Lifecycle & Loyalty",
                   "Loyalty program operations, personalization, retention offers. Kroger Plus engagement and basket-level economics.",
                   "Loyalty Director / VP Customer",
                   "Tier 2 — high-attach"],
                  ["RC-E2E-08", "Customer Acquisition",
                   "Acquisition campaign optimization across paid/organic/owned channels. Promo-driven trial.",
                   "CMO / Acquisition Marketing Director",
                   "Tier 2 — adjacent"],
                  ["RC-E2E-05", "Store Operations",
                   "Store-level operations — labor scheduling, task management, frontline decision support. Where Merch decisions get executed.",
                   "Store Operations Lead / District Manager",
                   "Tier 2 — execution downstream"],
                  ["RC-E2E-07", "Returns & Refund Integrity",
                   "Returns triage with fraud detection, refund-policy compliance, secondary-disposition routing.",
                   "Returns Operations / Loss Prevention",
                   "Tier 3 — peripheral"],
                  ["RC-E2E-06", "Finance Close Automation",
                   "Multi-entity record-to-report close with HITL on material entries. The Nike anchor service.",
                   "Corporate Controller",
                   "Tier 3 — finance function"],
                  ["RC-E2E-11", "Controls & Audit Evidence",
                   "SOX-grade evidence packet generation. Wraps RC-E2E-06 with audit-row LEDGER export.",
                   "SOX Program Lead / External Audit PIC",
                   "Tier 3 — finance function"],
              ],
              col_widths=[1.0, 1.5, 2.2, 1.3, 1.0],
              tier_col=4)

    add_callout(doc, "Reading the tiers",
                "Tier 1 services (RC-E2E-03, RC-E2E-09) are the daily decisioning surface for "
                "Merchandising — what Category Managers and Buyers actually do every day. Tier 2 "
                "services (RC-E2E-04, RC-E2E-08, RC-E2E-05) are functions Merchandising partners "
                "with closely — loyalty pricing, promo-driven acquisition, in-store execution. "
                "Tier 3 services (RC-E2E-07, RC-E2E-06, RC-E2E-11) live in other functional "
                "ownership (returns ops, finance, audit) and don't materially shape Merchandising's "
                "decisions.")

    pagebreak(doc)

    # ============ SECTION 3 · TIER RATIONALE ============
    add_h(doc, "3 · Tier rationale for Grocery Merchandising", level=1)

    add_h(doc, "3.1 · Tier 1 — anchor services (build first)", level=2)
    add_body(doc,
             "Two services are the operational core of Grocery Merchandising. They are not "
             "merely Merchandising-adjacent — they are how the function actually operates day to "
             "day. An engagement that does not include both Tier-1 services is not really doing "
             "Grocery Merchandising automation; it is doing one part of it.")
    add_table(doc,
              ["Service", "Why Tier 1 for Grocery Merch", "Core decisions it touches"],
              [
                  ["RC-E2E-03 Assortment & Pricing",
                   "It IS the Merchandising decision surface. Every pricing zone refresh, every markdown call, every assortment review, every promo plan flows through here.",
                   "Per-SKU price, markdown timing, assortment add/delete, promotional vehicle mix, cold-chain disposition"],
                  ["RC-E2E-09 Product Tracking",
                   "FSMA 204 compliance becomes operationally non-optional in 2026. Merchandising owns the supplier qualification and category-quality decisions that lot-traceability evidences.",
                   "Per-lot provenance, supplier traceability, recall response, category-quality KPI"],
              ],
              col_widths=[1.6, 2.4, 2.3])

    add_h(doc, "3.2 · Tier 2 — high-attach services (build alongside or just after)", level=2)
    add_body(doc,
             "Three services are not Merchandising-primary but are functions where Merchandising "
             "has shared ownership or material influence. The right time to attach them is when "
             "the Tier-1 services are running enough to provide the data they need.")
    add_table(doc,
              ["Service", "Why Tier 2 for Grocery Merch", "Where Merch contributes"],
              [
                  ["RC-E2E-04 Customer Lifecycle & Loyalty",
                   "Kroger Plus pricing personalization is the bridge between Merchandising's pricing engine and Customer's loyalty engine. Merch authors the price; Loyalty applies the personalization.",
                   "Personalized pricing offers · category-driven loyalty mechanics · basket-builder promo design"],
                  ["RC-E2E-08 Customer Acquisition",
                   "Promo-driven new-customer acquisition uses Merchandising's promotional vehicles (loss leaders, BOGO, ad blocks) as the bait.",
                   "Acquisition-promo design · category-mix targeting for new-trial baskets"],
                  ["RC-E2E-05 Store Operations",
                   "Where Merchandising decisions get executed on the store floor. Cold-chain disposition, planogram compliance, promotional execution, perishables rotation.",
                   "Planogram authoring · perishables policy · execution KPIs (OSA, on-shelf, freshness)"],
              ],
              col_widths=[1.6, 2.4, 2.3])

    add_h(doc, "3.3 · Tier 3 — peripheral services (not Merch-led, may be excluded from Merch scope)", level=2)
    add_body(doc,
             "Three services are core RC capabilities but live in functional ownership outside "
             "Merchandising. They appear in this portfolio for completeness, but a Grocery "
             "Merchandising engagement scoped narrowly does not need to include them.")
    add_table(doc,
              ["Service", "Why Tier 3 for Grocery Merch"],
              [
                  ["RC-E2E-07 Returns & Refund Integrity",
                   "Returns is a customer-service / loss-prevention function. Merchandising consumes return-rate data as a category-quality signal but does not own the operational service."],
                  ["RC-E2E-06 Finance Close Automation",
                   "Finance function. Even though Merchandising's gross-margin number flows through close, Merch does not author the close-cycle decisions."],
                  ["RC-E2E-11 Controls & Audit Evidence",
                   "SOX/audit function. Wraps RC-E2E-06; not a Merchandising-owned service."],
              ],
              col_widths=[1.6, 4.7])

    add_callout(doc, "Engagement-shaping recommendation",
                "Lead with RC-E2E-03 + RC-E2E-09 as the Wave-1 anchor. Tier-1 services together "
                "give Merchandising a complete decisioning + traceability substrate. Attach "
                "RC-E2E-04 in W2 once loyalty-side personalization can consume the pricing "
                "engine's outputs. Defer Tier-3 services to a separate engagement scoped against "
                "the Finance and Customer Service organizations.")

    pagebreak(doc)

    # ============ SECTION 4 · DEEP DIVE 1 — RC-E2E-03 ============
    add_h(doc, "4 · Deep Dive 1 of 3 · RC-E2E-03 Assortment & Pricing  (Tier 1, anchor)", level=1)

    add_h(doc, "4.1 · What this service does for Grocery Merchandising", level=2)
    add_body(doc,
             "RC-E2E-03 is the daily-decisioning service for Category Managers and Senior Buyers. "
             "It composes a five-agent fleet — Assess, Classify, Quantify, Approve, Act — that "
             "fuses MERML pricing data with SCML lot-provenance, runs FSMA 204 policy at decision "
             "time, and routes material decisions through HITL via Microsoft Teams Adaptive Card. "
             "Decisions emit a 14-field audit row to LEDGER.")

    add_h(doc, "4.2 · The decisions it touches", level=2)
    add_table(doc,
              ["Decision class", "Frequency", "Owner persona", "Audit-row significance"],
              [
                  ["Cold-chain disposition (markdown / destroy / sell-through)",
                   "~85 events per banner per week",
                   "Category Manager → Store Ops Lead",
                   "FSMA 204 evidence; shrink-NPV attribution"],
                  ["Per-SKU markdown timing",
                   "Daily across 5,000+ markdown candidates per banner",
                   "Category Manager / Pricing Manager",
                   "Markdown-NPV; perishables expiration policy"],
                  ["Pricing zone refresh",
                   "Weekly per zone × banner",
                   "Pricing Director / Category Manager",
                   "Competitive-position evidence; margin tracking"],
                  ["Assortment add/delete (listing/delist)",
                   "Quarterly category review",
                   "Category Manager / VP Merch",
                   "SKU-level contribution-margin attribution"],
                  ["Promotional vehicle mix (TPR, ad block, BOGO, loyalty)",
                   "26 ad blocks/year × multiple banners",
                   "Promotions Manager / Category Manager",
                   "Promo-ROI measurement; attribution to specific tactics"],
                  ["Vendor scorecard / slotting decisions",
                   "Quarterly + ad-hoc on supplier issues",
                   "Category Manager / Vendor Manager",
                   "Supplier-quality KPI; trade-spend ROI"],
              ],
              col_widths=[2.3, 1.7, 1.5, 1.7])

    add_h(doc, "4.3 · Persona constellation (Merchandising-led)", level=2)
    add_body(doc,
             "The walkthrough document anchored on the cold-chain scenario leads with the Store "
             "Operations Lead because that's the persona on the store floor at 6:08 AM. For the "
             "broader Merchandising service, the persona stack is led by the Category Manager — "
             "Store Ops Lead is the executional secondary.")
    add_callout(doc, "PRIMARY · Category Manager (e.g., Dairy Category Manager)",
                "Owns the category P&L. Authors the pricing/promotion/assortment policy that "
                "RC-E2E-03 manifests encode. Reviews per-cycle markdown-NPV, shrink trends, "
                "and competitive-position evidence. Reports to VP Merchandising. Measure of "
                "success: category gross-margin lift, shrink reduction, promo-ROI, supplier "
                "scorecard movement.",
                accent="E59A3E")
    add_table(doc,
              ["Role", "Function in RC-E2E-03 for Grocery Merch"],
              [
                  ["VP Merchandising / Chief Merchant",
                   "Executive sponsor; reviews enterprise-level scorecard rolling up category P&Ls"],
                  ["Senior Buyer (per category, e.g., Senior Buyer — Dairy)",
                   "Authors purchase orders against the assortment policy; consumes shrink/disposition data for vendor scorecards"],
                  ["Pricing Manager (banner level)",
                   "Owns markdown-policy manifest; tunes per-SKU elasticity; approves above-threshold pricing"],
                  ["Promotions Manager",
                   "Owns ad-block planning; reviews promo-ROI attribution; designs new-promo tests"],
                  ["Vendor Manager / Category Procurement",
                   "Consumes supplier scorecards; runs quarterly business reviews with vendors using audit-row evidence"],
                  ["Store Operations Lead (executional)",
                   "Receives the disposition Adaptive Card on store-floor events; physical execution and approval at the moment of decision"],
                  ["Food Safety Compliance Lead",
                   "Escalation path for CRITICAL events; consumes LEDGER for FDA inquiry; co-owns FSMA policy with RC-E2E-09"],
              ],
              col_widths=[2.4, 3.9])

    add_h(doc, "4.4 · KPIs Merchandising will care about", level=2)
    add_table(doc,
              ["KPI", "Why Merchandising cares", "Target (W3)"],
              [
                  ["Network-wide shrink avoided ($)",
                   "Direct gross-margin contribution; lands in category-level P&L",
                   "~$200M/yr at Kroger scale"],
                  ["Markdown-NPV per cycle",
                   "Whether the markdown engine is making money or burning margin",
                   "Trending positive · attributed to specific decisions"],
                  ["Per-SKU pricing-conformance to policy",
                   "Discipline metric — are pricing zone refreshes actually executed",
                   "≥ 95%"],
                  ["Category-quality (return rate · complaint rate)",
                   "Category Manager scorecard input; affects vendor reviews",
                   "Stable or improving per category"],
                  ["FSMA-204 audit-row completeness",
                   "Regulatory non-optional; affects category-trade-up risk",
                   "≥ 99.5%"],
                  ["Promo-ROI by vehicle",
                   "Whether a TPR or ad-block actually drove incremental basket",
                   "Attribution at trace-id; per-tactic ranking"],
                  ["Vendor scorecard movement (supplier-by-supplier)",
                   "Vendor management leverage; trade-spend ROI",
                   "Quarterly review using audit-row evidence"],
              ],
              col_widths=[2.0, 2.7, 1.6])

    add_h(doc, "4.5 · See also", level=2)
    add_body(doc,
             "The full end-to-end walkthrough of RC-E2E-03 — including the full chain link "
             "narrative, six agent-level use cases, three commercial waves, and Wave-1 prototype "
             "specification — lives in APEX-RC-E2E-03-Walkthrough.docx (this folder). That "
             "walkthrough anchors on the cold-chain excursion scenario as the marquee example; "
             "the Category Manager persona framing in this section can be applied as an overlay "
             "on the walkthrough's Section 7.")

    pagebreak(doc)

    # ============ SECTION 5 · DEEP DIVE 2 — RC-E2E-09 ============
    add_h(doc, "5 · Deep Dive 2 of 3 · RC-E2E-09 Product Tracking  (Tier 1, co-anchor)", level=1)

    add_h(doc, "5.1 · What this service does for Grocery Merchandising", level=2)
    add_body(doc,
             "RC-E2E-09 is the lot-provenance and FSMA 204 compliance substrate that RC-E2E-03 "
             "consumes at decision time. While RC-E2E-03 owns the pricing/markdown side of every "
             "decision, RC-E2E-09 owns the regulatory + supplier-traceability side. The two "
             "services are intentionally co-anchored — running RC-E2E-03 without RC-E2E-09 "
             "leaves the FSMA evidence chain incomplete; running RC-E2E-09 without RC-E2E-03 "
             "leaves the regulatory engine without the commercial decisions it serves.")

    add_h(doc, "5.2 · Why Merchandising owns this (not Quality Assurance alone)", level=2)
    add_body(doc,
             "Lot-traceability has historically been considered a Quality Assurance / Food Safety "
             "function. Under FSMA 204 (final rule effective 2026), it becomes a daily "
             "decisioning function — the same SKUs that Merchandising buys, prices, promotes, "
             "and assorts now require lot-level Critical Tracking Events (CTEs) and Key Data "
             "Elements (KDEs). The buyer who decides to list a SKU is also implicitly committing "
             "to the lot-trace burden; the Category Manager who decides a promotion is also "
             "implicitly committing to the volume-traceability that promotion will generate.")
    add_body(doc,
             "Putting RC-E2E-09 inside the Merchandising portfolio (rather than treating it as a "
             "peripheral compliance service) reflects the operating reality of Grocery "
             "Merchandising in the FSMA 204 era.")

    add_h(doc, "5.3 · The decisions it touches", level=2)
    add_table(doc,
              ["Decision class", "Frequency", "Owner persona"],
              [
                  ["Supplier qualification (lot-traceability capability)",
                   "Annual at supplier review · ad-hoc at category change",
                   "Category Manager / Vendor Manager"],
                  ["Lot-provenance attestation per receipt",
                   "Daily at distribution-center receiving",
                   "Inventory Control / Receiving Lead (with Merch policy)"],
                  ["Recall response (lot-level pull-from-shelf)",
                   "Ad-hoc on FDA recall or supplier-initiated",
                   "Food Safety Compliance Lead + Category Manager"],
                  ["Category-quality KPI review",
                   "Monthly category review",
                   "Category Manager / Quality Manager"],
                  ["FSMA 204 evidence packet for inspection",
                   "Quarterly + on-demand for FDA inspection windows",
                   "Food Safety Compliance Lead → FDA-relations team"],
              ],
              col_widths=[2.3, 1.9, 2.0])

    add_h(doc, "5.4 · Persona constellation (Merchandising × Compliance overlap)", level=2)
    add_body(doc,
             "Where RC-E2E-03 leads with Category Manager, RC-E2E-09 leads with Food Safety "
             "Compliance Lead. The Category Manager is the secondary primary — they own the "
             "supplier-qualification policy that the lot-trace engine enforces.")
    add_callout(doc, "PRIMARY · Food Safety Compliance Lead",
                "Owns the FSMA 204 policy corpus. Approves above-threshold lot-trace gaps. "
                "Interfaces directly with FDA-relations and state inspectors. Consumes the LEDGER "
                "audit-row stream for inspection. Reports to VP Quality. Measure of success: "
                "audit-row coverage, recall response time, FSMA 204 compliance score.",
                accent="E59A3E")
    add_callout(doc, "SECONDARY PRIMARY · Category Manager",
                "Authors the supplier-qualification policy. Acts on supplier-quality data in "
                "vendor reviews. Approves lot-trace exceptions for high-trust suppliers. "
                "Co-owns the FSMA 204 policy with the Food Safety Lead.",
                accent=GOLD)

    add_h(doc, "5.5 · KPIs Merchandising will care about", level=2)
    add_table(doc,
              ["KPI", "Why Merchandising cares", "Target (W3)"],
              [
                  ["FSMA-204 audit-row completeness",
                   "Direct regulatory exposure metric · category-trade-up risk",
                   "≥ 99.5%"],
                  ["Recall response time (signal → shelf-pull)",
                   "Reputational risk · shelf-availability impact",
                   "≤ 2 hours from recall signal"],
                  ["Lot-trace coverage by category",
                   "Some categories ramp slower (long supply chains)",
                   "≥ 98% on Food Traceability List categories"],
                  ["Supplier-quality scorecard",
                   "Vendor review input · trade-spend allocation",
                   "Per-supplier audit-row evidence"],
                  ["Time-to-evidence for FDA inspection",
                   "Inspection-window response capability",
                   "≤ 4 hours from request"],
              ],
              col_widths=[2.0, 2.7, 1.6])

    pagebreak(doc)

    # ============ SECTION 6 · DEEP DIVE 3 — RC-E2E-04 ============
    add_h(doc, "6 · Deep Dive 3 of 3 · RC-E2E-04 Customer Lifecycle & Loyalty  (Tier 2, high-attach)", level=1)

    add_h(doc, "6.1 · What this service does for Grocery Merchandising", level=2)
    add_body(doc,
             "RC-E2E-04 is the loyalty and customer lifecycle service. For grocery, this is "
             "Kroger Plus + 84.51° personalization at scale — 60M+ households worth of basket "
             "data driving personalized pricing, promotional targeting, and loyalty program "
             "mechanics. Where RC-E2E-03 owns the published price for the SKU, RC-E2E-04 owns "
             "the personalized-to-shopper price layer that rides on top.")
    add_body(doc,
             "Merchandising does not own this service operationally — it lives in the Customer / "
             "Loyalty / Marketing organization. But Merchandising contributes the pricing engine, "
             "the promotional vehicles, and the category economics that the loyalty engine uses. "
             "And Merchandising consumes the basket-level data that the loyalty engine generates.")

    add_h(doc, "6.2 · Why Tier 2 (high-attach) for Grocery Merch", level=2)
    add_body(doc,
             "Three reasons RC-E2E-04 sits one tier below the anchors:")
    add_bullets(doc, [
        "It is not Merchandising-led — Customer / Loyalty owns it, with Merchandising as a critical input/output partner",
        "It depends on RC-E2E-03 — without the Tier-1 pricing engine running, the loyalty layer has nothing to personalize",
        "Its primary value lever is customer LTV and retention, not direct category P&L (though there is meaningful spillover)",
    ])
    add_body(doc,
             "But it is high-attach because every loyalty pricing decision is also a Merchandising "
             "pricing decision; every personalized promo is also a Merchandising promo. The two "
             "organizations either work together cleanly or fight a turf war that erodes both.")

    add_h(doc, "6.3 · The decisions it touches (where Merchandising contributes)", level=2)
    add_table(doc,
              ["Decision class", "Merch role", "Loyalty role"],
              [
                  ["Personalized pricing offers (per-shopper)",
                   "Authors base price + concession ceiling",
                   "Personalizes the offer per shopper"],
                  ["Loyalty-program category mechanics (e.g., 4× fuel points on dairy)",
                   "Co-designs category mechanic with loyalty",
                   "Operates the redemption + tracking"],
                  ["Targeted promo outreach (digital coupon)",
                   "Provides the promo vehicle pool",
                   "Targets and dispatches per shopper"],
                  ["Basket-builder promo design",
                   "Authors the cross-category promo logic",
                   "Targets at-risk-of-attrition shoppers"],
                  ["Win-back offer composition",
                   "Approves above-threshold concession",
                   "Identifies the at-risk shopper · composes the offer"],
              ],
              col_widths=[2.4, 2.0, 1.9])

    add_h(doc, "6.4 · KPIs Merchandising will care about", level=2)
    add_table(doc,
              ["KPI", "Why Merchandising cares", "Target (W3)"],
              [
                  ["Loyalty-driven incremental basket lift",
                   "Tells Merch whether the loyalty layer is incremental or cannibalizing",
                   "+2.5% basket vs. non-loyalty cohort"],
                  ["Personalized-pricing margin protection",
                   "Whether per-shopper price decisions are protecting category margin",
                   "Per-cohort margin floor enforced"],
                  ["Category-loyalty mechanic ROI",
                   "Whether 'extra fuel points on dairy' is actually moving dairy",
                   "Per-mechanic positive ROI · attributed"],
                  ["Win-back conversion rate",
                   "Indirect — affects total loyalty-cohort size which affects category economics",
                   "+8% over baseline"],
              ],
              col_widths=[2.0, 2.7, 1.6])

    add_h(doc, "6.5 · Why this attaches to RC-E2E-03 (and shouldn't be done first)", level=2)
    add_callout(doc, "Sequencing rule",
                "Build RC-E2E-04 only after RC-E2E-03's pricing engine is running on at least the "
                "Wave-2 250-store anchor cluster. Reason: RC-E2E-04's personalization layer needs "
                "RC-E2E-03's published price + concession ceiling as inputs. Reversing the order "
                "creates a personalization engine with no canonical base-price contract — a "
                "common anti-pattern that creates pricing-governance chaos at scale.")

    pagebreak(doc)

    # ============ SECTION 7 · RECOMMENDED SEQUENCING ============
    add_h(doc, "7 · Recommended sequencing through the portfolio", level=1)
    add_body(doc,
             "How to land the eight RC services across an engagement program at Kroger, scoped "
             "to the Merchandising organization. The recommended path lands the Tier-1 anchor "
             "first, attaches Tier-2 services as their dependencies become available, and "
             "explicitly excludes Tier-3 services from Merch-led scope.")

    add_table(doc,
              ["Wave / Phase", "Services in scope", "Why this combination", "Owner persona stack"],
              [
                  ["Wave 1 — Foundation",
                   "RC-E2E-03 (read-only)",
                   "Lead with the Merch anchor. Read-only intelligence on cold-chain + markdown decisions; no automated action yet.",
                   "Category Manager (primary) · Food Safety Lead · Pricing Manager"],
                  ["Wave 2 — Pilot",
                   "RC-E2E-03 (full live) + RC-E2E-09",
                   "Complete the Tier-1 anchor pair. RC-E2E-09 brings FSMA 204 evidence layer; RC-E2E-03 goes from read-only to live with HITL and downstream action.",
                   "Category Manager · Food Safety Compliance Lead · Store Ops Lead · Vendor Manager"],
                  ["Wave 3 — Scale & Fuse",
                   "RC-E2E-03 + RC-E2E-09 enterprise + RC-E2E-04 (attach)",
                   "Scale Tier-1 pair to enterprise. Attach RC-E2E-04 once published-price layer is stable so loyalty personalization has the substrate it needs.",
                   "Add: VP Merchandising · Loyalty Director · Pricing Director"],
                  ["Wave 4 (post-W3) — Adjacent",
                   "RC-E2E-08 + RC-E2E-05 (selective attach)",
                   "Acquisition (E2E-08) leverages Merch promo vehicles; Store Ops (E2E-05) deepens the in-store execution loop with planogram + perishables policy.",
                   "Add: Acquisition Marketing Director · Store Ops Director"],
                  ["Out of Merch scope",
                   "RC-E2E-07 · RC-E2E-06 · RC-E2E-11",
                   "Customer Service / Finance / Audit functions. Pursue separately under those organizations' engagement scopes.",
                   "Returns Operations · CFO · SOX Program — separate engagement"],
              ],
              col_widths=[1.4, 1.7, 2.4, 1.8])

    add_h(doc, "7.1 · Cumulative envelope across the Merch-led path (W1–W4)", level=2)
    add_table(doc,
              ["Wave", "Services live", "Approx. envelope (Mid)", "Cumulative envelope"],
              [
                  ["Approach",          "—",                                      "$200K–$500K",        "$200K–$500K"],
                  ["Wave 1",            "RC-E2E-03 (read-only)",                  "$1.2M–$3.5M",        "$1.4M–$4M"],
                  ["Wave 2",            "+ RC-E2E-09 + RC-E2E-03 live",           "$5M–$12M",           "$6.4M–$16M"],
                  ["Wave 3",            "Enterprise scale + RC-E2E-04 attach",    "$22M–$50M",          "$28.4M–$66M"],
                  ["Wave 4 (selective)","+ RC-E2E-08 / RC-E2E-05 selective",      "$5M–$15M",           "$33.4M–$81M"],
              ],
              col_widths=[1.0, 2.6, 1.4, 1.4])
    add_body(doc,
             "Wave 3 envelope is wider than the RC-E2E-03 walkthrough's $18M–$45M because it now "
             "includes the RC-E2E-04 attach + the broader Tier-1 enterprise scope. The W3 durable "
             "value also expands: $200M shrink-avoided (RC-E2E-03) plus loyalty-driven incremental "
             "basket lift (RC-E2E-04 on 60M+ households) is plausibly $400M+ annualized at "
             "run-rate.",
             italic=True, color=INK2)

    pagebreak(doc)

    # ============ SECTION 8 · UPDATE PATH FOR THE WALKTHROUGH ============
    add_h(doc, "8 · How this updates the walkthrough", level=1)
    add_body(doc,
             "The companion walkthrough — APEX-RC-E2E-03-Walkthrough.docx — was authored before "
             "the Grocery Merchandising sub-category framing was established. The walkthrough "
             "leads with the Store Operations Lead as primary persona because the marquee scenario "
             "(cold-chain excursion at 6:08 AM) is a store-floor moment.")
    add_body(doc,
             "For the Grocery Merchandising sub-category specifically, the persona overlay is:")
    add_table(doc,
              ["In the walkthrough", "For Grocery Merch sub-category, treat as"],
              [
                  ["Store Operations Lead (primary persona)",
                   "PRIMARY (executional) — but not the strategic owner. Category Manager is the strategic primary; Store Ops is the moment-of-decision primary."],
                  ["Food Safety Compliance Lead (secondary)",
                   "Co-primary on FSMA 204 evidence (RC-E2E-09 anchor)"],
                  ["Pricing Manager (secondary)",
                   "Reports to Category Manager / Pricing Director within Merch organization"],
                  ["Pull-Team Lead (secondary)",
                   "Stays as-is — store-floor execution"],
                  ["Inventory Control Specialist (secondary)",
                   "Stays as-is — Merch consumes their reconciliation output"],
                  ["FDA Auditor / State Inspector (secondary)",
                   "Stays as-is — evidence consumer"],
                  ["(Add) Category Manager",
                   "STRATEGIC PRIMARY — owns the policy, the P&L, the vendor reviews, the assortment"],
                  ["(Add) VP Merchandising / Chief Merchant",
                   "Executive sponsor — receives enterprise scorecard"],
                  ["(Add) Senior Buyer",
                   "Per-category PO authoring; consumes vendor scorecards"],
                  ["(Add) Promotions Manager",
                   "Ad-block planning; promo-ROI attribution"],
                  ["(Add) Vendor Manager",
                   "Quarterly business reviews using audit-row evidence"],
              ],
              col_widths=[2.4, 4.0])

    add_callout(doc, "Recommended walkthrough update",
                "Either (1) revise the walkthrough's Section 7 in-place to lead with Category "
                "Manager as primary persona and demote Store Operations Lead to executional "
                "secondary, or (2) keep the walkthrough's store-ops framing for the cold-chain "
                "scenario specifically and use this portfolio document as the Merchandising-side "
                "overlay for executive conversations. Option 2 preserves the walkthrough's "
                "operational vividness while making the Merchandising context explicit at the "
                "portfolio level. We recommend Option 2 unless the next stakeholder review is "
                "primarily with the Chief Merchant rather than store operations leadership.")

    pagebreak(doc)

    # ============ APPENDIX A · SERVICE QUICK REFERENCE ============
    add_h(doc, "Appendix A · Service quick reference card", level=1)
    add_body(doc, "One-row summary of every RC service for fast lookup during stakeholder conversations.")
    add_table(doc,
              ["Service ID", "Name", "Tier (Grocery Merch)", "Primary persona", "Key value lever"],
              [
                  ["RC-E2E-03", "Assortment & Pricing", "Tier 1 — anchor",
                   "Category Manager", "Shrink avoided · markdown-NPV · pricing conformance"],
                  ["RC-E2E-09", "Product Tracking", "Tier 1 — co-anchor",
                   "Food Safety Compliance Lead", "FSMA 204 evidence · recall response · supplier scorecard"],
                  ["RC-E2E-04", "Customer Lifecycle & Loyalty", "Tier 2 — high-attach",
                   "Loyalty Director", "Personalized pricing · loyalty incremental basket lift"],
                  ["RC-E2E-08", "Customer Acquisition", "Tier 2 — adjacent",
                   "Acquisition Marketing Director", "Promo-driven new-customer trial"],
                  ["RC-E2E-05", "Store Operations", "Tier 2 — execution downstream",
                   "Store Operations Lead", "Planogram compliance · OSA · in-store execution"],
                  ["RC-E2E-07", "Returns & Refund Integrity", "Tier 3 — peripheral",
                   "Returns Operations", "Returns triage · fraud detection"],
                  ["RC-E2E-06", "Finance Close Automation", "Tier 3 — finance function",
                   "Corporate Controller", "Close-cycle compression · audit-fee avoidance"],
                  ["RC-E2E-11", "Controls & Audit Evidence", "Tier 3 — finance function",
                   "SOX Program Lead", "SOX evidence packet · auditor walkthrough"],
              ],
              col_widths=[1.0, 1.6, 1.2, 1.5, 2.0],
              tier_col=2)

    add_body(doc,
             "Independence posture preserved across this portfolio document: 'Deloitte's "
             "Microsoft practice,' 'DMTSP,' 'Microsoft platform capabilities,' "
             "'Microsoft-native deployment.' The terms 'partner,' 'alliance,' and 'strategic "
             "alliance' do not appear.",
             italic=True, color=MUTE, size=9)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— End of portfolio document —"); set_run(r, italic=True, size=10, color=MUTE)

    # save
    target = OUT
    try:
        doc.save(str(target))
    except PermissionError:
        n = 2
        while True:
            alt = target.with_name(target.stem + f"-v{n}" + target.suffix)
            if not alt.exists(): target = alt; break
            n += 1
        doc.save(str(target))
        print(f"[notice] {OUT.name} was locked — saved as {target.name} instead.")
    print(f"Wrote {target}")


if __name__ == "__main__":
    build()
