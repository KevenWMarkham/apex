"""Build the APEX · TMT-MED-01 service walkthrough (Disney+ reference client).

Models the same shape as Nike-R2R-NOAH-Walkthrough.docx:
  - 11 sections + 2 appendices
  - 30 tables, no figures
  - Reference-client framing — Disney+ anchors TMT-MED-01 the same way Nike
    anchors RC-E2E-06. No separate program codename — this is APEX.

Service:        TMT-MED-01 — Media Subscriber Lifecycle
Reference:      Disney+ (Walt Disney Streaming Services)
APEX scenario:  TMT-CX-02 — Streaming churn intervention (content recommendation)
Output:         Disney+ folder · 02_projects/FY27_Pipeline/subscriber-lifecycle-agentic/deliverables/
"""

from __future__ import annotations

import io
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

OUT_DIR = Path(r"C:\Stage\Clients\Industries\TMT\Global_Media_Entertainment\Walt_Disney\02_projects\FY27_Pipeline\subscriber-lifecycle-agentic\deliverables")
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT = OUT_DIR / "APEX-TMT-MED-01-Walkthrough.docx"

# ---- styling helpers ----------------------------------------------------

INK   = "0B0B0F"
INK2  = "4A4A55"
MUTE  = "8a8a90"
GOLD  = "B8892E"
TEAL  = "2A7F73"
VIO   = "6648B0"
BAND  = "F3F3F5"
HEAD  = "0B0B0F"
LINE  = "D4D4D7"


def set_cell_bg(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex.lstrip("#"))
    tcPr.append(shd)


def set_cell_border(cell, color_hex=LINE):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), color_hex.lstrip("#"))
        borders.append(b)
    tcPr.append(borders)


def set_run(run, *, bold=False, italic=False, size=None, color=None, font="Arial"):
    if bold: run.bold = True
    if italic: run.italic = True
    if size is not None: run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor.from_string(color.lstrip("#"))
    if font: run.font.name = font


def add_body(doc, text, *, size=11, italic=False, color=INK, space_after=6, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None: p.alignment = align
    r = p.add_run(text)
    set_run(r, size=size, color=color, italic=italic)
    return p


def add_h(doc, text, *, level=1):
    sizes = {1: 18, 2: 14, 3: 12.5}
    p = doc.add_paragraph()
    p.style = doc.styles[f"Heading {level}"]
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(4)
    # Override the Heading style's color/font
    p.runs.clear() if hasattr(p, "runs") and callable(getattr(p, "runs", None)) else None
    # Actually just clear children and add a fresh run
    for r in list(p.runs):
        r.text = ""
    r = p.add_run(text)
    set_run(r, bold=True, size=sizes.get(level, 12), color=INK)
    if level == 1:
        # gold border under H1
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), GOLD)
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p


def add_bullets(doc, items, *, size=11):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(it)
        set_run(r, size=size, color=INK)


def add_table(doc, header, rows, *, col_widths=None,
              header_fill=HEAD, header_color="FFFFFF", zebra=True):
    n = len(header)
    tbl = doc.add_table(rows=len(rows) + 1, cols=n)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(header):
        c = tbl.rows[0].cells[j]
        set_cell_bg(c, header_fill)
        set_cell_border(c)
        c.paragraphs[0].text = ""
        r = c.paragraphs[0].add_run(h)
        set_run(r, bold=True, size=10, color=header_color)
        if col_widths: c.width = Inches(col_widths[j])
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            c = tbl.rows[i].cells[j]
            set_cell_border(c)
            if zebra and i % 2 == 0:
                set_cell_bg(c, BAND)
            c.paragraphs[0].text = ""
            r = c.paragraphs[0].add_run(str(v))
            set_run(r, size=10, color=INK)
            if col_widths: c.width = Inches(col_widths[j])
    doc.add_paragraph()
    return tbl


def add_callout(doc, kind, body, *, accent=GOLD, fill="F6F1E4"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = tbl.rows[0].cells[0]
    set_cell_bg(c, fill)
    c.paragraphs[0].text = ""
    p_tag = c.paragraphs[0]
    r_tag = p_tag.add_run(kind.upper())
    set_run(r_tag, bold=True, size=9, color=accent)
    p = c.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    r = p.add_run(body)
    set_run(r, size=10.5, color=INK)
    doc.add_paragraph()


def add_kv(doc, pairs, *, lw=1.6, rw=4.9):
    tbl = doc.add_table(rows=len(pairs), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(pairs):
        ck, cv = tbl.rows[i].cells
        ck.width = Inches(lw)
        cv.width = Inches(rw)
        set_cell_bg(ck, BAND)
        set_cell_border(ck); set_cell_border(cv)
        ck.paragraphs[0].text = ""
        cv.paragraphs[0].text = ""
        rk = ck.paragraphs[0].add_run(k)
        set_run(rk, bold=True, size=10, color=INK)
        rv = cv.paragraphs[0].add_run(v)
        set_run(rv, size=10, color=INK)
    doc.add_paragraph()


def add_eyebrow(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text.upper())
    set_run(r, bold=True, size=9, color=MUTE)
    rPr = r._element.get_or_add_rPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:val"), "80")
    rPr.append(spacing)


def pagebreak(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=9.5, color=INK, font="Consolas")


# ============================================================================
# BUILD
# ============================================================================

def build():
    doc = Document()
    # US Letter, 1" margins
    for section in doc.sections:
        section.page_width  = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = section.bottom_margin = Inches(0.9)
        section.left_margin = section.right_margin = Inches(1.0)
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # ============ COVER ============
    add_eyebrow(doc, "APEX · TMT-MED-01 · REFERENCE WALKTHROUGH · DISNEY+")
    p = doc.add_paragraph()
    r = p.add_run("APEX · TMT-MED-01")
    set_run(r, bold=True, size=26, color=INK)
    p1b = doc.add_paragraph()
    r1b = p1b.add_run("Media Subscriber Lifecycle")
    set_run(r1b, bold=True, size=20, color=INK)
    p2 = doc.add_paragraph()
    r2 = p2.add_run("Audience targeting, retention, and churn intervention — Disney+ as reference client")
    set_run(r2, size=13, color=INK2, italic=True)

    doc.add_paragraph()
    add_body(doc,
             "A walkthrough of the APEX TMT-MED-01 Media Subscriber Lifecycle service across the "
             "APEX value-delivery chain — Scenario → Solution → Use Case → Service → Persona → "
             "KPI — anchored against Disney+ as the Wave-1 reference client. The same six-agent "
             "service pattern documented here ports across other media-and-entertainment clients "
             "(Warner Bros Discovery, Fox, Sinclair) and to telco subscriber-lifecycle and "
             "software customer-success scenarios with only the manifest layer changing.",
             italic=True, color=INK2)

    doc.add_paragraph()
    add_kv(doc, [
        ("Document type",      "APEX Service Walkthrough · Service-to-KPI"),
        ("Service",            "TMT-MED-01 · Media Subscriber Lifecycle"),
        ("APEX scenario",      "TMT-CX-02 · Streaming churn intervention (content recommendation)"),
        ("Practice",           "TMT — Technology, Media & Telecom · Media sub-Practice"),
        ("Reference client",   "The Walt Disney Company · Disney+ Streaming Services"),
        ("Audience",           "CFO · Streaming COO · VP Subscriber Retention · Audience Lead · External audit"),
        ("Chain links",        "Scenario · Solution · Use Case · Service · Persona · KPI"),
        ("Waves covered",      "Approach · W1 Foundation · W2 Pilot · W3 Scale & Fuse"),
        ("Companion artifacts","TMT-MED-01-Sprint-Plan.xlsx · TMT-MED-01-Resource-Plan.xlsx · TMT-MED-01-ROI-Case.html"),
        ("Version / date",     "v1.0 · 27 Apr 2026"),
    ])

    pagebreak(doc)

    # ============ CONTENTS ============
    add_h(doc, "Contents", level=1)
    add_body(doc,
             "This document walks Disney+'s subscriber lifecycle program through every link of the "
             "APEX value-delivery chain. Sections 1–2 set context. Sections 3–8 are the chain itself: "
             "one section per chain link. Section 9 projects the chain across the three delivery "
             "waves; Section 10 summarizes the commercial envelope; Section 11 describes the "
             "reference Wave 1 implementation. Appendices A and B carry the glossary and the "
             "artifact cross-reference.",
             italic=True, color=INK2)

    add_table(doc,
              ["#", "Section", "What it covers"],
              [
                  ["1",  "The Reference Cycle at Disney+",
                   "The subscriber-lifecycle decisions Disney+ makes every hour and the gap "
                   "between intervention timing and outcome value."],
                  ["2",  "Why Subscriber Lifecycle, Why Disney+, Why Now",
                   "The five forces converging on streaming retention; why Disney+ is the TMT "
                   "Media reference client; why FY27 is the build window."],
                  ["3",  "Chain Link 1 · Scenario",
                   "Disney+'s subscriber lifecycle today; current-state response; friction "
                   "quantified; scope boundaries."],
                  ["4",  "Chain Link 2 · Solution",
                   "TMT-MED-01 in one sentence; six-agent orchestration; commercial envelope tiering; "
                   "three design commitments."],
                  ["5",  "Chain Link 3 · Use Cases",
                   "Six agent-level use cases composing the streaming-retention scenario; mapping "
                   "to the sprint plan; incremental delivery path."],
                  ["6",  "Chain Link 4 · Services",
                   "TMT-MED-01 Media Subscriber Lifecycle + TMT-MED-03 Rights & Monetization; "
                   "service composition; tenant manifest; candidate derived services."],
                  ["7",  "Chain Link 5 · Personas",
                   "VP Subscriber Retention as primary; five secondary personas (Audience Lead, "
                   "Content Programming, Pricing Strategy, Marketing Ops, External Audit) with "
                   "plane bindings."],
                  ["8",  "Chain Link 6 · KPIs",
                   "Per-cohort KPIs, derived KPIs for Disney commercial cases, audit-row attribution."],
                  ["9",  "The Three Waves",
                   "W1 Foundation (Disney+ NA single market), W2 Pilot (NA + cross-cohort), "
                   "W3 Scale & Fuse (global + Hulu + ESPN+)."],
                  ["10", "Commercial Envelope",
                   "Approach / W1 / W2 / W3 cost envelopes and durable-value reference."],
                  ["11", "Reference Wave 1 Implementation",
                   "What a Disney+ on TMT-MED-01 W1 prototype looks like; demo paths; technology stack; "
                   "intentional scope boundaries."],
                  ["A",  "Glossary",
                   "Selected terms used in this document with APEX- or streaming-specific meaning."],
                  ["B",  "Artifact Cross-Reference",
                   "How this walkthrough fits with the sprint plan, resource plan, and ROI case."],
              ],
              col_widths=[0.4, 2.6, 3.5])

    pagebreak(doc)

    # ============ SECTION 1 · Reference Cycle ============
    add_h(doc, "Section 1 · The Reference Cycle at Disney+", level=1)
    add_body(doc,
             "At 8:14 AM Pacific on a Tuesday in February 2027, the VP Subscriber Retention at "
             "Disney+ opens her Power BI dashboard for the daily lifecycle review. Three numbers "
             "have moved overnight. Trial-conversion lift from yesterday's audience experiment is "
             "+1.4 points. At-risk-cohort growth is +12,000 subscribers — the largest single-day "
             "movement in the past 30 days. And the Hulu cross-attach rate has dipped 0.3 points "
             "for the third consecutive day. Three signals, three different agent fleets, three "
             "different decisions waiting on her approval. The clock on each is different too: "
             "the trial-conversion signal has a 48-hour intervention window before the population "
             "ages out; the at-risk cohort needs a 72-hour content + offer surface before churn "
             "probability climbs above 85%; the Hulu attach trend can wait a week.")
    add_body(doc,
             "This is the scenario. Everything downstream — the TMT-MED-01 agent architecture, the "
             "sprint use cases, the productized services, the persona cast, the KPI rollups — "
             "exists to answer three questions about the ~262M global subscribers and the daily "
             "intervention windows that the VP Retention has to make calls on:")
    add_bullets(doc, [
        "Which of the ~2.5M lifecycle interventions proposed each day can be auto-dispatched, and which must be reviewed by a human?",
        "Who at Disney has the authority to approve a price-concession or content-exclusivity offer above threshold, and how do they make it in seconds rather than the days it takes today?",
        "What evidence survives the decision — auditable, attributable to the specific subscriber, the specific signal, the specific intervention, and the specific outcome?",
    ])

    add_h(doc, "Today, without TMT-MED-01", level=3)
    add_body(doc,
             "Disney+'s ~85-person Subscriber Retention organization spends the first three hours "
             "of each day pulling cohort exports from the data warehouse, joining engagement "
             "telemetry against subscription state, prioritizing intervention candidates in Excel, "
             "and routing approvals via Slack and email. The end-of-day intervention cadence runs "
             "two business days behind the engagement signal. SOX-equivalent evidence for "
             "audit-grade attribution (which intervention drove which retention outcome at which "
             "subscriber on which device on which date) is reconstructed quarterly from BI exports, "
             "vendor logs, and CRM histories. When the audit committee asks for the top ten "
             "retention-NPV-driving interventions of the quarter, the answer takes three weeks.")

    add_h(doc, "Tomorrow, with TMT-MED-01", level=3)
    add_body(doc,
             "Six agents — Engagement-Sense, Risk-Score, Content-Match, Offer-Compose, "
             "Outreach-Dispatch, Outcome-Attribute — run in parallel where the work allows, "
             "sequentially where it requires. Material concession decisions route through the VP "
             "Subscriber Retention via Microsoft Teams Adaptive Card on her device. Every decision "
             "emits a 14-field audit row to LEDGER. The intervention-to-outcome attribution that "
             "today takes three weeks is generated continuously and queryable in seconds. "
             "Retention NPV is no longer a quarterly retrospective — it is a leading indicator the "
             "VP Retention checks before her morning leadership stand-up.")

    add_callout(doc, "Independence note",
                "TMT-MED-01's agent manifests, policy manifests, and LEDGER are Disney artifacts; "
                "Deloitte's role is advisory and implementation enablement. The SOW is structured "
                "so that the program meets Deloitte independence requirements regardless of "
                "Disney's external-audit firm relationship at the time of contracting. The "
                "underlying APEX pattern is independence-safe by design; the specific engagement "
                "structure is designed in concert with Deloitte's Independence office.")

    # ============ SECTION 2 · Why ============
    add_h(doc, "Section 2 · Why Subscriber Lifecycle, Why Disney+, Why Now", level=1)

    add_h(doc, "2.1 · The five forces converge on the streaming retention calendar", level=2)
    add_body(doc,
             "APEX exists because five forces converged across FY25–FY26. For streaming "
             "subscriber lifecycle specifically, all five land on the daily retention "
             "calendar: the decision-velocity gap that leaves intervention windows open while "
             "humans assemble cohort exports; the agentic-AI inflection that crossed enterprise "
             "viability for personalization at population scale; the post-peak-streaming "
             "competitive pressure that makes every retained subscriber more economically valuable "
             "than every newly-acquired one; the content-investment pressure that demands "
             "engagement-driven content programming feedback loops; and the privacy and "
             "consent-governance regime that demands per-decision evidence for every personalization "
             "and every concession.")

    add_h(doc, "2.2 · Why Disney+ is the APEX TMT Media reference client", level=2)
    add_body(doc,
             "Of the four global media-and-entertainment subjects in the DMTSP TMT account "
             "portfolio (Disney, Warner Bros Discovery, Fox, Sinclair), Disney carries the "
             "reference subscriber-lifecycle scenario because Disney's portfolio has all the "
             "hardness of a public-company streaming P&L plus distinctive operating realities "
             "that stress the agent architecture in ways pure-play streamers do not:")
    add_bullets(doc, [
        "Three-streaming-brand portfolio — Disney+, Hulu, ESPN+ each with distinct engagement, churn, and price elasticity profiles",
        "Cross-brand bundle dynamics — the Disney Bundle drives attach economics that pure-play models cannot represent",
        "Theatrical / streaming staggered release windows — content recommendations must reason about availability windows that change weekly",
        "Sports rights tier — ESPN+ subscriber behavior is dominated by live event seasonality (NFL, college football, soccer leagues)",
        "International market complexity — Disney+ EMEA, Greater China (via Hulu Japan analog), APAC, and LATAM each have distinct content rights, pricing tiers, and regulatory overlays",
        "Ad-supported tier coexistence — Disney+ Basic with ads requires audience-targeting infrastructure that ad-free SVOD does not",
    ])
    add_body(doc,
             "If TMT-MED-01's orchestration pattern handles Disney's retail-streaming subscriber "
             "lifecycle, it ports across the Deloitte DMTSP practices to telco subscriber lifecycle "
             "(churn intervention with bundle dynamics), software SaaS customer success (consumption "
             "and renewal), and travel loyalty programs (status-tier retention) with only the "
             "manifest layer changing.")

    add_h(doc, "2.3 · Why this document exists", level=2)
    add_body(doc,
             "Three conversations at Disney need the same material with different emphasis. The "
             "Streaming COO and CFO need to see a credible retention-NPV story with ROI that "
             "survives a rigorous scrub against the existing personalization stack. The VP "
             "Subscriber Retention and Audience Lead need to see an operating-model story that "
             "fits how their teams actually work today. The CTO and CDO need to see a "
             "Microsoft-native architecture that fits the existing data and identity estate "
             "without inventing net-new platform categories. This document is the executive "
             "walkthrough that anchors all three.")

    pagebreak(doc)

    # ============ SECTION 3 · SCENARIO ============
    add_h(doc, "Section 3 · Chain Link 1 · Scenario", level=1)
    add_body(doc,
             "The Scenario link of the APEX value-delivery chain is the customer's operational "
             "reality stated without abstraction. It is not yet a solution. It is the thing that "
             "has to be true before a solution can be designed against it.")

    add_h(doc, "3.1 · Disney's subscriber lifecycle today", level=2)
    add_body(doc,
             "Disney's three streaming brands carry approximately 262M global subscribers across "
             "Disney+, Hulu, and ESPN+. Each subscriber moves through a defined lifecycle — trial, "
             "active, engaged, at-risk, lapsed, churned, won-back — with state transitions driven "
             "by engagement signals (watch time, depth of viewing, content discovery, app usage), "
             "lifecycle events (subscription renewal, payment failure, plan change), and external "
             "shocks (competitor launch, content release, sports season transition).")
    add_body(doc,
             "On any given day, approximately 2.5 million lifecycle decisions are theoretically "
             "available to be made — content surfacing, retention offers, price concessions, "
             "cross-brand attach proposals, advertising-tier upsells, payment-recovery outreach. "
             "Today, perhaps 200,000 of those — eight percent — actually receive a personalized "
             "intervention. The other 92% receive either no intervention or a generic "
             "intervention that statistically underperforms a personalized one by a meaningful "
             "margin.")

    add_h(doc, "3.2 · The current-state response", level=2)
    add_body(doc,
             "Today, the response across the three Disney streaming brands is a blend of (1) "
             "static personalization rules built into the recommendation system, (2) batch "
             "intervention campaigns scheduled weekly by the lifecycle marketing team, (3) "
             "ad-hoc retention offers approved through email/Slack chains, and (4) reactive "
             "winback campaigns triggered after subscribers cancel. The typical sequence "
             "across a one-week intervention cycle:")
    add_bullets(doc, [
        "Day 1 (Monday) — Audience Lead pulls cohort segmentation reports from the data warehouse; lifecycle marketing builds intervention campaign list in marketing automation tooling",
        "Day 2 — Content Programming Director reviews proposed content surfaces against rights windows and flags conflicts (e.g., a recommendation candidate that has gone off the platform)",
        "Day 3 — Pricing Strategy reviews proposed price concessions and bundle attach offers against margin guardrails",
        "Day 4 — Approval routing through email/Slack with the VP Subscriber Retention for any concession above threshold",
        "Day 5 — Marketing Operations executes the campaign through the customer engagement platform",
        "Day 6–7 — Outcome measurement waits 7–14 days for engagement signal to mature; attribution to specific intervention tries to disentangle from organic signal",
    ])
    add_body(doc,
             "By the time the loop closes — typically 14–21 days after the original signal — the "
             "intervention window has long since closed. Subscribers who were going to churn have "
             "churned; subscribers who were going to engage have done so on whatever surface "
             "happened to be presented to them. The retention organization is sprinting, but the "
             "loop time is the limit.")

    add_h(doc, "3.3 · The friction, quantified", level=2)
    add_body(doc, "Measured over four representative quarters at Disney+ North America, the "
                  "current-state response produces the following friction profile:")
    add_table(doc,
              ["Dimension", "Typical value", "Implication"],
              [
                  ["Daily lifecycle decisions theoretically available",
                   "~2.5M / day across 3 brands NA",
                   "92% currently un-personalized; intervention coverage gap is the headline opportunity"],
                  ["Effective personalized intervention rate (today)",
                   "~8% of available daily decisions",
                   "Capacity-constrained: human routing limits scale"],
                  ["Intervention loop time — signal → action",
                   "~14–21 days median",
                   "Most signals expire before action; intervention windows close"],
                  ["At-risk cohort identification accuracy",
                   "~64% precision · ~58% recall (current model)",
                   "False negatives churn before intervention; false positives waste capacity"],
                  ["Retention concession review backlog",
                   "~3 days median for above-threshold offers",
                   "Best concession candidates often expire from the review queue"],
                  ["Outcome-attribution latency",
                   "Quarterly retrospective via BI exports",
                   "No leading indicators on intervention NPV; programmatic improvement is slow"],
                  ["Privacy/consent-decision evidence",
                   "Reconstructed quarterly from system logs",
                   "Audit-row evidence for personalization decisions is not native to the stack"],
                  ["Cross-brand attach forecasting",
                   "Aggregate-only, brand-by-brand",
                   "Per-subscriber bundle propensity not actionable in the intervention loop"],
              ],
              col_widths=[2.4, 1.9, 2.2])

    add_h(doc, "3.4 · Scope of the Disney scenario", level=2)
    add_body(doc, "The Disney+ on TMT-MED-01 scenario specifically includes:")
    add_bullets(doc, [
        "Engagement signal ingestion from the Disney streaming platforms (watch sessions, content interactions, app activity, profile activity)",
        "Subscriber lifecycle state computation and transition classification",
        "Per-subscriber churn-risk scoring with confidence and time-to-action",
        "Personalized content-surface recommendation respecting current rights window state",
        "Personalized retention offer composition (price freeze, bundle attach, exclusivity, ad-tier downgrade option)",
        "Outreach dispatch via approved channels (push, email, in-app) with HITL on material concessions",
        "Outcome attribution at trace-id granularity, feeding NPV measurement and policy refinement",
        "Privacy/consent decision evidence for every personalization touch — captured as APEX 14-field audit row",
    ])
    add_body(doc, "It specifically excludes:")
    add_bullets(doc, [
        "Content acquisition and programming decisions (separate scenario — content-rights monetization)",
        "Acquisition campaign optimization (separate scenario — performance marketing acquisition)",
        "Ad inventory pricing and yield management for the Disney+ Basic ad-supported tier (adjacent scenario)",
        "Live-event programming for ESPN+ (handled by sports rights orchestration scenario)",
    ])

    pagebreak(doc)

    # ============ SECTION 4 · SOLUTION ============
    add_h(doc, "Section 4 · Chain Link 2 · Solution", level=1)
    add_body(doc,
             "The Solution link is how APEX answers the Scenario. It is an architectural "
             "description — what runs, where it runs, what contracts bind the components, how "
             "humans stay in charge of the decisions that matter.")

    add_h(doc, "4.1 · TMT-MED-01 in one sentence", level=2)
    add_callout(doc, "Solution",
                "A six-agent subscriber-lifecycle orchestration — Engagement-Sense, Risk-Score, "
                "Content-Match, Offer-Compose, Outreach-Dispatch, Outcome-Attribute — composes "
                "with a HITL-gated commercial envelope, runs on Microsoft-native infrastructure "
                "(Fabric + Foundry + Azure OpenAI + Teams + Entra) and consumes CAF (Converge "
                "Agentic Framework) primitives via A2A, and writes every decision to a 14-field "
                "audit row so retention NPV and personalization evidence are generated, not "
                "reconstructed.",
                accent=TEAL, fill="E4F1EE")

    add_h(doc, "4.2 · The six agents", level=2)
    add_table(doc,
              ["Agent", "Responsibility", "Input", "Output", "HITL?"],
              [
                  ["Engagement-Sense",
                   "Ingest engagement telemetry; normalize to TMTML.Engagement",
                   "Streaming platform events (watch, click, search, profile)",
                   "TMTML.Engagement records",
                   "No"],
                  ["Risk-Score",
                   "Compute per-subscriber churn probability + LTV impact",
                   "TMTML.Subscriber + TMTML.Engagement",
                   "Risk score with confidence + time-to-action",
                   "If confidence < 0.85, escalate"],
                  ["Content-Match",
                   "Propose personalized content surfaces respecting rights windows",
                   "TMTML.Subscriber + TMTML.Content + TMTML.Rights",
                   "Ordered slate with reasoning trace",
                   "Rare (rights logic deterministic)"],
                  ["Offer-Compose",
                   "Compose retention offer with price/bundle/exclusivity dimensions",
                   "Risk score + LTV + per-cohort offer policy",
                   "Offer proposal with concession tier",
                   "If concession > threshold, route to VP Retention"],
                  ["Outreach-Dispatch",
                   "Dispatch the intervention via approved channel; respect consent",
                   "Offer + content slate + consent state",
                   "Channel call (push/email/in-app) + receipt",
                   "Always for above-threshold concessions"],
                  ["Outcome-Attribute",
                   "Observe outcome window; attribute to trace-id; emit audit row",
                   "Engagement post-intervention + subscription state",
                   "14-field audit row to LEDGER",
                   "No"],
              ],
              col_widths=[1.3, 1.5, 1.4, 1.3, 1.0])

    add_h(doc, "4.3 · Commercial envelope — Disney tier thresholds", level=2)
    add_body(doc,
             "The commercial envelope is the policy layer that decides 'do this automatically vs. "
             "require human approval.' For Disney+ on TMT-MED-01, the envelope is tiered by subscriber "
             "value × concession amount × confidence:")
    add_table(doc,
              ["Subscriber tier", "Concession", "Confidence", "Action"],
              [
                  ["Tier 1 (high LTV: long-tenure, multi-brand)",   "Any",          "Any",     "Always HITL — VP Retention or designate"],
                  ["Tier 2 (standard)",                              "> $40 / yr",   "Any",     "HITL — Director-level retention review"],
                  ["Tier 2 (standard)",                              "≤ $40 / yr",   "≥ 0.95",  "Auto-dispatch; notify"],
                  ["Tier 2 (standard)",                              "≤ $40 / yr",   "< 0.95",  "HITL — Manager-level"],
                  ["Tier 3 (low LTV / new trial)",                   "≤ $15 / yr",   "≥ 0.90",  "Auto-dispatch; log only"],
                  ["Any tier",                                        "Cross-brand bundle attach > $X", "Any",  "HITL — bundle pricing review"],
                  ["Any tier",                                        "Ad-tier conversion offer", "Any",  "Auto-dispatch (revenue accretive)"],
              ],
              col_widths=[2.4, 1.6, 1.0, 1.6])

    add_h(doc, "4.4 · Three design commitments", level=2)

    add_h(doc, "4.4.1 · Manifest-driven, not model-driven", level=3)
    add_body(doc,
             "Every choice that could change between Disney close cycles — which LLM is used for "
             "offer-narrative generation, which tools each agent can call, which policy corpus "
             "applies to concession thresholds, which Audience Lead has approval authority for "
             "which tier — is declared in a Git-versioned manifest. Updating any of these is a "
             "manifest PR with review and merge discipline, not an ops-ticket. Privacy/consent "
             "policy versioning is part of the manifest history; auditors and consent regulators "
             "see exactly what the system was configured to do at the time of any specific "
             "personalization decision.")

    add_h(doc, "4.4.2 · Audit row, not log line", level=3)
    add_body(doc,
             "Every decision — every risk score, every content-surface choice, every offer "
             "compose, every dispatch, every outcome attribution — emits a 14-field structured "
             "row carrying trace-id, three version stamps (manifest, policy, prompt), "
             "content-addressed input and output hashes, reasoning-trace reference, HITL status "
             "and approver, downstream-effect cross-reference, and KPI attribution. This is the "
             "artifact that turns 'Disney+ personalization' from a black box into an auditable "
             "substrate.")

    add_h(doc, "4.4.3 · HITL is first-class, not bolted on", level=3)
    add_body(doc,
             "The HITL gate between Offer-Compose and Outreach-Dispatch is not a retro-fitted "
             "human check. It is declared in the orchestration manifest as a first-class "
             "primitive. Its gate kind (ZERO_TOUCH, ACK_ONLY, HITL, or ESCALATION) is "
             "deterministically mapped from the commercial envelope. The VP Subscriber "
             "Retention sees only what she must see; everything else dispatches automatically "
             "with audit-row attribution. Material concessions arrive as Microsoft Teams "
             "Adaptive Cards on her device with one-click approve / modify / reject.")

    pagebreak(doc)

    # ============ SECTION 5 · USE CASES ============
    add_h(doc, "Section 5 · Chain Link 3 · Use Cases", level=1)
    add_body(doc,
             "The Use Case link is the set of discrete operational patterns the Solution "
             "activates. One Solution typically enables multiple Use Cases — each Use Case is a "
             "bounded, testable, individually-measurable pattern that can be delivered as a piece "
             "of a larger engagement or as a standalone increment.")
    add_body(doc,
             "TMT-MED-01's six primary use cases (at the agent orchestration level) compose with "
             "the 22 detailed sprint-level use cases documented in TMT-MED-01-Sprint-Plan.xlsx. "
             "The six agent-level use cases below are individually deliverable, individually "
             "measurable, and individually commercially-envelopable.")

    for title, payload in [
        ("5.1 · Use Case · Engagement Signal Ingest",
         [("Trigger", "Streaming platform events arrive (watch/click/search/profile)"),
          ("Agents",  "Engagement-Sense"),
          ("Tools",   "Streaming-platform event connector; TMTML.Engagement schema validator"),
          ("HITL",    "None — deterministic ingest and normalization"),
          ("Envelope","Included in W1 foundation")]),
        ("5.2 · Use Case · At-Risk Cohort Identification",
         [("Trigger", "Engagement window closes for a subscriber cohort"),
          ("Agents",  "Risk-Score"),
          ("Tools",   "Subscriber-state retrieval; ML model serving (CAF cap:churn-prediction)"),
          ("HITL",    "Confidence < 0.85 escalates to Audience Lead for cohort review"),
          ("Envelope","~15% of W2 scope")]),
        ("5.3 · Use Case · Content Surface Personalization",
         [("Trigger", "At-risk score crosses threshold OR scheduled cohort refresh"),
          ("Agents",  "Content-Match"),
          ("Tools",   "Rights-window service; recommendation model (CAF cap:content-recommendation)"),
          ("HITL",    "Rare — rights logic is deterministic; exceptions to Content Programming"),
          ("Envelope","~20% of W2 scope")]),
        ("5.4 · Use Case · Retention Offer Composition",
         [("Trigger", "Risk score + intervention-eligible state"),
          ("Agents",  "Offer-Compose"),
          ("Tools",   "Per-cohort offer policy; LTV computation; concession-budget service"),
          ("HITL",    "If concession > tier threshold, route via Adaptive Card"),
          ("Envelope","~30% of W2 scope (highest judgment surface area)")]),
        ("5.5 · Use Case · Multi-Channel Dispatch",
         [("Trigger", "Approved offer + content slate"),
          ("Agents",  "Outreach-Dispatch"),
          ("Tools",   "Push/email/in-app channel APIs; consent-state service"),
          ("HITL",    "Always for above-threshold concessions; receipt logged"),
          ("Envelope","~20% of W2 scope")]),
        ("5.6 · Use Case · Outcome Attribution & Close-Out",
         [("Trigger", "Outcome window closes (typically 7–14 days post-dispatch)"),
          ("Agents",  "Outcome-Attribute"),
          ("Tools",   "Subscription-state delta; engagement-delta; LEDGER write"),
          ("HITL",    "None — write-only; consumed by Audience Lead and VP Retention dashboards"),
          ("Envelope","~15% of W2 scope; foundation for NPV measurement")]),
    ]:
        add_h(doc, title, level=3)
        add_kv(doc, payload, lw=1.2, rw=5.0)

    add_h(doc, "5.7 · Mapping to the Disney sprint plan's 22 use cases", level=2)
    add_body(doc,
             "The six agent-level use cases above are each realized by multiple sprint-level use "
             "cases in TMT-MED-01-Sprint-Plan.xlsx. The sprint plan refines the agent-level surface "
             "with specific user stories, acceptance criteria, and Wave assignment.")
    add_table(doc,
              ["Agent-level use case", "Sprint-level use cases (TMT-MED-01-Sprint-Plan.xlsx)"],
              [
                  ["5.1 Engagement Signal Ingest",          "UC-01 · UC-02 · UC-03"],
                  ["5.2 At-Risk Cohort Identification",     "UC-04 · UC-05 · UC-06 · UC-07"],
                  ["5.3 Content Surface Personalization",   "UC-08 · UC-09 · UC-10 · UC-11"],
                  ["5.4 Retention Offer Composition",       "UC-12 · UC-13 · UC-14 · UC-15 · UC-16"],
                  ["5.5 Multi-Channel Dispatch",            "UC-17 · UC-18 · UC-19"],
                  ["5.6 Outcome Attribution & Close-Out",   "UC-20 · UC-21 · UC-22"],
              ],
              col_widths=[2.6, 3.7])

    add_h(doc, "5.8 · Incremental delivery path", level=2)
    add_body(doc,
             "Individually, each use case is a deliverable. Together they compose the full "
             "scenario. The composition is deliberate — Disney can start with the first two "
             "agent-level use cases (Engagement Signal Ingest + At-Risk Cohort Identification) "
             "as a 'Subscriber Risk Insight' increment, capture meaningful value at a smaller "
             "envelope, and expand from there. Adding Content-Match + Offer-Compose converts that "
             "into a full intervention pipeline. Adding Outreach-Dispatch + Outcome-Attribute "
             "closes the loop end-to-end.")

    pagebreak(doc)

    # ============ SECTION 6 · SERVICES ============
    add_h(doc, "Section 6 · Chain Link 4 · Services", level=1)
    add_body(doc,
             "The Service link is what gets productized. A Service is a named, subscribable APEX "
             "capability with a defined persona, SLO, commercial terms, and an anchor in the APEX "
             "Service Catalog.")

    add_h(doc, "6.1 · Anchor services", level=2)

    add_h(doc, "TMT-MED-01 · Media Subscriber Lifecycle", level=3)
    add_body(doc,
             "The foundational subscriber-lifecycle automation service. Owns the six-agent "
             "orchestration from Engagement-Sense through Outreach-Dispatch. Primary persona: "
             "VP Subscriber Retention (Disney+ NA). SLO: intervention-loop-time ≤ 24 hours from "
             "signal to action at W2 steady state; ≤ 4 hours at W3. Commercial envelope: "
             "tiered by streaming-brand subscriber count, geographic market scope, and "
             "intervention coverage depth.")

    add_h(doc, "TMT-MED-03 · Media Rights & Monetization", level=3)
    add_body(doc,
             "The rights-and-monetization service that wraps Content-Match's rights-window "
             "logic, broadcast-rights utilization tracking, and downstream royalty residual "
             "reconciliation. Primary persona: Rights Management Lead + Distribution Strategy. "
             "SLO: rights-window query response < 100ms; royalty residual close within 72 hours "
             "of period end. Commercial envelope: tiered by content portfolio scale (number of "
             "rights agreements, number of monetization windows, breadth of distribution "
             "partnerships).")

    add_h(doc, "6.2 · Service composition", level=2)
    add_table(doc,
              ["Commitment", "TMT-MED-01 Media Subscriber Lifecycle",
                            "TMT-MED-03 Media Rights & Monetization"],
              [
                  ["Canonical schemas",
                   "TMTML.Subscriber, .Engagement, .Content, .Rights",
                   "TMTML.Rights, .Content"],
                  ["Agents deployed",
                   "Engagement-Sense, Risk-Score, Content-Match, Offer-Compose, Outreach-Dispatch",
                   "Rights-Window-Query, Royalty-Reconcile"],
                  ["MCP tool surfaces",
                   "Streaming-platform event ingest, channel APIs (push/email/in-app), consent-state",
                   "Rights-management system query, royalty-engine compute"],
                  ["HITL policies",
                   "tier × concession × confidence",
                   "exceptions on residual variance > threshold"],
                  ["Power BI surfaces",
                   "Retention dashboard · NPV trending · cohort drilldown",
                   "Rights utilization dashboard · royalty close timeline"],
                  ["Teams surfaces",
                   "Above-threshold concession Adaptive Card",
                   "Rights-conflict notification card"],
                  ["SLO",
                   "Intervention loop time, attribution latency",
                   "Rights-query latency, residual close time"],
                  ["Primary persona",
                   "VP Subscriber Retention (Disney+ NA)",
                   "Rights Management Lead"],
              ],
              col_widths=[1.6, 2.5, 2.5])

    add_h(doc, "6.3 · Disney tenant manifest subscription", level=2)
    add_body(doc, "The Disney tenant manifest subscribes to the two services by declaration:")
    add_code(doc,
             '{ "tenant_id": "dis", "subscriptions": [\n'
             '  { "service_id": "TMT-MED-01", "service_version": "v1.0.0" },\n'
             '  { "service_id": "TMT-MED-03", "service_version": "v1.0.0" }\n'
             ']}')
    add_body(doc,
             "Version pinning is deliberate. When the APEX Service Catalog releases v1.1.0 of "
             "TMT-MED-01, Disney decides when to upgrade. PATCH bumps are automatic; MINOR and "
             "MAJOR bumps require an explicit accept-upgrade HITL gate. This is how Disney stays "
             "on stable versions while the catalog continues to evolve.")

    add_h(doc, "6.4 · Candidate derived services (post-W3)", level=2)
    add_body(doc,
             "TMT-MED-01 and TMT-MED-03 anchor the subscriber-lifecycle scenario today. Three "
             "adjacent service candidates share the same architectural substrate and become "
             "fusion candidates at the end of Wave 3:")
    add_bullets(doc, [
        "TMT-MED-04 · Audience-Tier Optimization — extends Risk-Score with ad-supported-tier conversion modeling and ad-yield optimization",
        "TMT-MED-05 · Cross-Streaming Brand Bundle Orchestration — extends Offer-Compose with cross-brand attach (Disney+ → Hulu → ESPN+) propensity modeling",
        "TMT-MED-06 · Theatrical-Streaming Window Strategy — fuses subscriber lifecycle with theatrical-release windowing decisions",
    ])

    pagebreak(doc)

    # ============ SECTION 7 · PERSONAS ============
    add_h(doc, "Section 7 · Chain Link 5 · Personas", level=1)
    add_body(doc,
             "The Persona link is the human-role cast that makes the scenario operational. "
             "Persona depth is what separates an architecturally correct solution from one that "
             "lands on the floor.")

    add_h(doc, "7.1 · Primary persona", level=2)
    add_callout(doc, "PRIMARY · VP Subscriber Retention (Disney+ NA)",
                "Owns the retention P&L target, the concession policy, and the HITL approval "
                "path for every above-threshold offer. Approves via Microsoft Teams Adaptive Card "
                "on her phone or laptop; her approvals, modifications, and rejections write to "
                "the audit row with her named identity resolved through Entra. She reports to the "
                "Streaming COO and interfaces directly with the Audience Lead and Pricing "
                "Strategy Lead. Her measure of success: retention NPV per cohort, intervention "
                "loop time, and audit-row coverage.",
                accent="E59A3E")
    add_body(doc,
             "She is named, not generic. 'VP Subscriber Retention' is a title; she is a person "
             "with named identity, named device, named approval authority. The hitl-authority "
             "policy registered in LEDGER binds to her named identity (resolved through Entra "
             "group-lookup at runtime). The Adaptive Card fires to her device. Her approval is "
             "what unblocks Outreach-Dispatch.")

    add_h(doc, "7.2 · Secondary personas", level=2)
    add_body(doc,
             "Five secondary personas each carry a distinct role in Disney+'s subscriber-lifecycle "
             "operating model. Roles are deliberate — The service does not invent roles, it maps to roles "
             "that already exist in Disney's retention organization.")
    add_table(doc,
              ["Role", "Function in the Disney+ on TMT-MED-01 scenario"],
              [
                  ["Audience Lead",
                   "Owns cohort definition and segmentation; reviews Risk-Score precision/recall; "
                   "escalation path on cohort-level model drift."],
                  ["Content Programming Director",
                   "Reviews Content-Match recommendations against rights windows; flags "
                   "programming-strategy conflicts; reviews W3 content-investment feedback loop."],
                  ["Pricing Strategy Lead",
                   "Owns concession-policy manifest; reviews W2 concession-cost trends; tunes "
                   "tier thresholds quarterly."],
                  ["Marketing Operations Director",
                   "Owns the dispatch channel infrastructure; manages consent-state integration; "
                   "first-line response for delivery exceptions."],
                  ["External Audit / Privacy Compliance",
                   "Consumes the LEDGER audit-row stream for personalization decision evidence; "
                   "interfaces on consent-regulation reviews (GDPR, CCPA, evolving regimes)."],
              ],
              col_widths=[2.0, 4.3])

    add_h(doc, "7.3 · Why persona depth matters", level=2)
    add_body(doc,
             "A seller who names only the primary persona says 'we built something for the VP "
             "Retention.' A seller who names the full cast shows Disney a complete operating "
             "model: who owns content surfacing, who owns concession policy, who owns the channel "
             "stack, who consumes the audit row, who resolves model drift. That completeness is "
             "what turns a demo into an engagement.")
    add_body(doc,
             "Persona depth also matters for change management. When the engagement rolls out, "
             "Disney's enablement team has to train all six named roles — not just the VP "
             "Retention. The six-role inventory informs training plan scope, org-readiness "
             "assessment, and the operational handoff checklist in the W2 → W3 gate review.")

    add_h(doc, "7.4 · Persona-to-plane mapping", level=2)
    add_table(doc,
              ["Persona", "Plane(s) touched", "Surface"],
              [
                  ["VP Subscriber Retention (primary)",  "Experience · Decision",
                   "Microsoft Teams Adaptive Card (mobile + desktop)"],
                  ["Audience Lead",                       "Experience · Decision",
                   "Power BI cohort dashboard · model-drift notification card"],
                  ["Content Programming Director",        "Experience · Integration",
                   "Power BI rights-utilization page · drill-through to trace-id"],
                  ["Pricing Strategy Lead",               "Experience · Runtime",
                   "Concession-policy manifest console · BI concession-cost dashboard"],
                  ["Marketing Operations Director",       "Integration · Runtime",
                   "Channel-ops console · dispatch-receipt dashboard"],
                  ["External Audit / Privacy Compliance", "Ledger",
                   "Quarterly evidence-packet export · consent-decision query"],
                  ["Streaming COO (executive sponsor)",   "Ledger · Experience",
                   "Power BI executive scorecard · per-period NPV rollup"],
              ],
              col_widths=[2.1, 1.8, 2.4])

    pagebreak(doc)

    # ============ SECTION 8 · KPIs ============
    add_h(doc, "Section 8 · Chain Link 6 · KPIs", level=1)
    add_body(doc,
             "The KPI link is where the chain closes. APEX's signature is that every KPI "
             "attributes back to the specific decision audit row that produced it — not just to "
             "the aggregate. The distinction matters because it is what makes the KPI durable "
             "under audit and actionable under operating review.")

    add_h(doc, "8.1 · Per-cohort KPIs", level=2)
    add_body(doc,
             "A closed intervention cohort emits a closeout audit row carrying the following "
             "measured KPI fields. Each is bound to the specific trace-id of that cohort. A "
             "Power BI dashboard rolls them up across cohorts; drill-through on any aggregated "
             "cell returns the individual audit rows.")
    add_table(doc,
              ["KPI", "Measurement", "Target (W2 Pilot)", "Target (W3 Scale)"],
              [
                  ["Intervention loop time",
                   "Hours from signal to dispatch", "≤ 24 hrs", "≤ 4 hrs"],
                  ["Personalized intervention rate",
                   "% of available daily decisions personalized", "35%", "75%"],
                  ["At-risk cohort identification — precision",
                   "P at risk-threshold", "≥ 0.78", "≥ 0.85"],
                  ["At-risk cohort identification — recall",
                   "R at risk-threshold", "≥ 0.72", "≥ 0.82"],
                  ["Retention NPV per intervention",
                   "Audit-row attributed NPV", "+$84", "+$143"],
                  ["Concession utilization efficiency",
                   "% of concession budget yielding positive NPV", "≥ 70%", "≥ 85%"],
                  ["Audit-row coverage",
                   "% of personalization decisions attributed", "≥ 98%", "≥ 99.5%"],
                  ["Cross-brand bundle attach rate",
                   "Per-cohort attach % above baseline", "+1.8 pp", "+4.2 pp"],
              ],
              col_widths=[2.1, 2.0, 1.2, 1.2])

    add_h(doc, "8.2 · Derived KPIs for Disney commercial cases", level=2)
    add_body(doc,
             "Beyond the per-cohort operational metrics, four derived KPIs anchor Disney-specific "
             "commercial conversations. Each is computed from the per-cohort metrics aggregated "
             "over time and attributes back to specific cohorts.")

    add_h(doc, "8.2.1 · Subscriber Retention NPV", level=3)
    add_body(doc,
             "The headline financial metric. Computed as SUM(retention_npv_per_intervention) × "
             "covered intervention volume, projected against Disney+ NA's ~50M-subscriber base. "
             "At the W3 target of +$143 NPV per intervention and ~75% intervention coverage of "
             "~2.5M daily decisions, the annualized retention-NPV uplift sits in the "
             "$200M–$400M range — depending on baseline churn assumptions and concession-cost "
             "absorption. This is the single number CFO and Streaming COO will gate the "
             "engagement on.")

    add_h(doc, "8.2.2 · Intervention Coverage Lift", level=3)
    add_body(doc,
             "The operational metric that matters to the VP Retention and Audience Lead. "
             "Computed as (personalized interventions today) ÷ (theoretically available daily "
             "decisions). Today: ~8%. W2 target: 35%. W3 target: 75%. The metric is leading — it "
             "predicts retention-NPV trajectory before retention-NPV itself moves, because "
             "coverage gaps are where retention NPV is being left on the table.")

    add_h(doc, "8.2.3 · Loop-Time Compression", level=3)
    add_body(doc,
             "The metric that matters to leadership stand-up cadence. Today's intervention-loop "
             "time is 14–21 days median; W2 target is ≤ 24 hours; W3 target is ≤ 4 hours. The "
             "value of compression isn't just operational — it's that intervention windows that "
             "today close before action are now actionable. A subscriber whose engagement signal "
             "ages out in 72 hours becomes an addressable subscriber.")

    add_h(doc, "8.2.4 · Personalization Audit-Row Coverage", level=3)
    add_body(doc,
             "The regulatory/governance metric. Computed as the percentage of personalization "
             "decisions whose audit rows carry all 14 required fields, every version stamp, and a "
             "verifiable signature. Target ≥ 99.5% at W3. Matters because it is the direct answer "
             "to the consent-regulator audit question and the audit-committee privacy-controls "
             "question.")

    add_h(doc, "8.3 · KPI attribution in the audit row", level=2)
    add_body(doc,
             "The closeout audit row — emitted once per resolved intervention cohort — carries "
             "the structure below. Every KPI attributes back to specific rows by trace-id.")
    add_code(doc,
             '{\n'
             '  "trace_id": "trc_2027-02-09_0814_dis_dtc_na_at_risk_001",\n'
             '  "scenario_id": "TMT-CX-02-streaming-churn-intervention",\n'
             '  "service_id": "TMT-MED-01",\n'
             '  "tenant_id": "dis",\n'
             '  "event_kind": "atlas_intervention_closeout",\n'
             '  "event_ts": "2027-02-23T08:14:32Z",\n'
             '  "manifest_version": "tmt-atlas-orchestration-v1.0.0",\n'
             '  "policy_version": "concession-corpus-v1.0.0",\n'
             '  "kpis": {\n'
             '    "intervention_loop_hours": 22.4,\n'
             '    "intervention_personalized": true,\n'
             '    "risk_score_at_decision": 0.91,\n'
             '    "concession_amount_usd": 18.00,\n'
             '    "outcome_retained": true,\n'
             '    "retention_npv_usd": 156.40,\n'
             '    "audit_row_complete": true\n'
             '  }\n'
             '}')

    pagebreak(doc)

    # ============ SECTION 9 · WAVES ============
    add_h(doc, "Section 9 · The Three Waves", level=1)
    add_body(doc,
             "The same chain — Scenario → Solution → Use Case → Service → Persona → KPI — "
             "projects across three delivery Waves. W1 Foundation is prerequisite work that "
             "enables the scenario. W2 Pilot is the scenario running live at proof-point scale. "
             "W3 Scale & Fuse is the scenario at enterprise run-rate, fused with adjacent "
             "scenarios.")

    add_h(doc, "9.1 · Wave 1 · Foundation — Disney+ NA single market", level=2)
    add_body(doc,
             "W1 is the one-time investment per tenant. It produces the rails that make the "
             "scenario possible and make every subsequent scenario cheap. For Disney, W1 scope "
             "is one streaming brand (Disney+ NA), three agents (Engagement-Sense, Risk-Score, "
             "Outcome-Attribute), and the full APEX platform substrate via CAF.")
    add_table(doc,
              ["W1 element", "What it looks like for Disney"],
              [
                  ["Scope",            "Disney+ NA only · single-brand · one cohort family"],
                  ["Agents live",      "3 of 6 (Engagement-Sense, Risk-Score, Outcome-Attribute)"],
                  ["Use cases",        "5.1 + 5.2 + partial 5.6"],
                  ["Personas onboarded","VP Retention + Audience Lead + Marketing Ops Director"],
                  ["Schemas",          "TMTML.Subscriber, TMTML.Engagement (NA scope only)"],
                  ["Audit row",        "Composed end-to-end on shadow data; replay against historical 4 quarters"],
                  ["Adaptive Card",    "VP Retention's device tested; Entra binding live"],
                  ["LEDGER plane",     "Stood up; pilot rows ingested; query path verified"],
                  ["Duration",         "5–8 months"],
                  ["Effort",           "~50 story points"],
                  ["Exit gate",        "VP Retention sign-off on shadow-run accuracy and audit-row coverage"],
              ],
              col_widths=[1.7, 4.6])

    add_h(doc, "9.2 · Wave 2 · Pilot — Disney+ NA live + cross-cohort intervention", level=2)
    add_body(doc,
             "W2 is where the scenario becomes real. Disney+ NA runs the live scenario "
             "end-to-end for 90 days across two cohort families plus the at-risk cross-cohort. "
             "The full six-agent chain fires for every intervention. Target: 90 consecutive days "
             "with no pipeline failure beyond documented exceptions, plus the first quarterly "
             "personalization-evidence packet delivered to Disney's Privacy office and Internal "
             "Audit.")
    add_table(doc,
              ["W2 element", "What it looks like for Disney"],
              [
                  ["Scope",            "Disney+ NA · 2 cohort families + at-risk cross-cohort"],
                  ["Agents live",      "All 6 (full orchestration)"],
                  ["Use cases",        "5.1 through 5.6 fully exercised"],
                  ["Personas onboarded","All 6 named roles operating live"],
                  ["Schemas",          "TMTML.Subscriber, .Engagement, .Content, .Rights"],
                  ["HITL volume",      "Tier-tuned to ~5–10% of interventions during steady state"],
                  ["Power BI",         "Retention dashboard · NPV trending · concession utilization"],
                  ["Privacy packet",   "First quarterly evidence packet exercised with Privacy office"],
                  ["Duration",         "8–12 months"],
                  ["Effort",           "~45 story points initial + ongoing enablement"],
                  ["Exit gate",        "Steering committee approval to proceed to W3 enterprise scale"],
              ],
              col_widths=[1.7, 4.6])

    add_h(doc, "9.3 · Wave 3 · Scale & Fuse — global · multi-brand · cross-streaming", level=2)
    add_body(doc,
             "W3 is where the scenario meets the Disney enterprise and where it stops running "
             "alone. Scope expands from Disney+ NA to Disney+ global (NA, EMEA, Greater China "
             "analogs, APAC, LATAM); brands expand from Disney+ to Disney+ + Hulu + ESPN+; "
             "cohort definitions expand from at-risk to the full lifecycle palette. And — most "
             "importantly — TMT-MED-01 fuses with adjacent scenarios (audience-tier optimization, "
             "cross-streaming-brand bundle, theatrical-streaming-window strategy) that share the "
             "architectural substrate.")
    add_table(doc,
              ["W3 element", "What it looks like for Disney"],
              [
                  ["Scope",            "Disney+ global + Hulu + ESPN+ · 5+ regional markets"],
                  ["Agents live",      "All 6 across all 3 brands; brand-aware policy manifests"],
                  ["Use cases",        "Full 22 sprint-level use cases plus fusion hooks"],
                  ["Schemas",          "All TMTML.* shapes; cross-brand aggregation views"],
                  ["Cross-brand fusion","Cross-streaming bundle (TMT-MED-05) and audience-tier (TMT-MED-04) candidates light up"],
                  ["Personas",         "Retention organization trained globally; Privacy/audit consuming evidence directly"],
                  ["KPI rollup",       "Enterprise-wide retention NPV in the $200M–$400M annualized range at run-rate"],
                  ["Loop time",        "≤ 4 hours signal-to-action across all brands and markets"],
                  ["Duration",         "12–18 months core; ongoing fusion"],
                  ["Effort",           "~70 story points initial + ongoing scenario-fusion"],
                  ["Exit gate",        "Steady-state operating review; expansion into adjacent scenarios"],
              ],
              col_widths=[1.7, 4.6])

    add_h(doc, "9.4 · Cross-wave load-bearing connections", level=2)
    add_body(doc,
             "Three structural dependencies bind the waves into one progression. They are not "
             "architectural coincidences — they are the load-bearing connections that make the "
             "wave model coherent.")
    add_callout(doc, "W1 → W2",
                "The HITL Adaptive Card prototyped in W1 (VP Retention's device with Entra "
                "identity resolution) powers the W2 HITL gate across all six agents. What is "
                "prototyped in W1 is consumed unchanged in W2.",
                accent=TEAL, fill="E4F1EE")
    add_callout(doc, "W1 → W2",
                "The LEDGER plane established in W1 accumulates the W2 audit rows. The "
                "personalization-evidence packet format is designed in W1, exercised with "
                "shadow audit rows, and scales to thousands of rows in W2.")
    add_callout(doc, "W2 → W3",
                "The W2 pilot KPIs unlock W3 scaling. The Steering Committee decision to commit "
                "to global + cross-streaming-brand rollout is specifically gated on W2 meeting "
                "intervention-loop-time and retention-NPV targets.",
                accent=TEAL, fill="E4F1EE")

    pagebreak(doc)

    # ============ SECTION 10 · COMMERCIAL ENVELOPE ============
    add_h(doc, "Section 10 · Commercial Envelope", level=1)
    add_body(doc,
             "The Disney engagement encompasses four commercial tiers corresponding to the "
             "Approach stage (Discovery and SOW build-out), Wave 1 Foundation, Wave 2 Pilot, and "
             "Wave 3 Scale & Fuse. Each carries a distinct cost-envelope range, duration, and "
             "deliverable signature.")
    add_table(doc,
              ["Tier", "Duration", "Cost envelope", "Effort", "Primary deliverable"],
              [
                  ["Approach",         "3–5 weeks",   "$200K–$500K",   "22 SP",   "Discovery report · scenario scoping · SOW"],
                  ["Wave 1 Foundation","5–8 months",  "$1M–$3M",       "50 SP",   "3-agent shadow-run · audit-row end-to-end"],
                  ["Wave 2 Pilot",     "8–12 months", "$4M–$10M",      "45 SP",   "Disney+ NA live · 90-day operating run · privacy packet"],
                  ["Wave 3 Scale&Fuse","12–18 months","$15M–$40M",     "70+ SP",  "Global · multi-brand · cross-streaming fusion"],
              ],
              col_widths=[1.4, 1.2, 1.3, 0.8, 2.5])

    add_h(doc, "10.1 · What the envelope buys", level=2)
    add_body(doc,
             "The commercial envelope is not a project-labor-hours calculation. It reflects the "
             "total value-delivery bundle: Deloitte-delivered solutioning and Practice expertise; "
             "Microsoft platform licensing and consumption (Fabric F-SKU, Foundry, Azure OpenAI, "
             "Teams, Entra, Power BI); CAF-substrate consumption and any Converge-tier costs that "
             "ride with it; ISV co-sell where applicable; client-side enablement and change "
             "management; and Privacy / Internal-Audit walkthrough coordination.")

    add_h(doc, "10.2 · Envelope vs. durable value at Disney scale", level=2)
    add_body(doc,
             "A global multi-brand W3 enterprise footprint produces retention-NPV uplift in the "
             "$200M–$400M annualized range, plus cross-brand bundle economics that show up "
             "primarily in the Disney Bundle ARPU line, plus operational and governance value "
             "(audit-grade personalization evidence, faster strategic experimentation, retained "
             "retention-organization talent). The $15M–$40M W3 envelope spans a rapid-payback "
             "range (under 12 months at the low end of returns; under 24 months at base case) "
             "to a strategic-investment range (3–4 years at the high end of investment, "
             "reflecting greater scenario-fusion scope in the upper tier).")

    add_h(doc, "10.3 · Independence posture in Disney commercial conversations", level=2)
    add_callout(doc, "Independence",
                "The seller and delivery team use Deloitte-appropriate language throughout: "
                "'Deloitte's Microsoft practice,' 'DMTSP,' 'Microsoft platform capabilities,' "
                "'Microsoft-native deployment.' The terms 'partner,' 'alliance,' and 'strategic "
                "alliance' do not appear in TMT-MED-01 commercial materials. The engagement is "
                "structured so that manifests, policies, and LEDGER are Disney artifacts; "
                "Deloitte's role is advisory and implementation enablement. Before any contract "
                "drafting, confirm Disney's external-audit firm relationship and route the SOW "
                "through Deloitte's Independence office.")

    pagebreak(doc)

    # ============ SECTION 11 · REFERENCE W1 IMPLEMENTATION ============
    add_h(doc, "Section 11 · Reference Wave 1 Implementation", level=1)
    add_body(doc,
             "The TMT-MED-01 Wave 1 prototype is the concrete proof point that anchors this walkthrough. "
             "Whereas the NOAH program at Nike has a runnable React + Node + Postgres prototype "
             "delivered under the FY27 Pipeline engagement, the Disney+ on TMT-MED-01 W1 prototype is the "
             "next analogous reference build. It will be a runnable web application that "
             "executes the three foundation use cases against synthetic Disney+ subscriber data "
             "with full audit-row capture.")

    add_h(doc, "11.1 · What the W1 prototype demonstrates", level=2)
    add_table(doc,
              ["Capability", "What it demonstrates"],
              [
                  ["Synthetic Disney+ NA subscriber population",
                   "100K synthetic subscribers across 8 cohort families with realistic engagement signals"],
                  ["Engagement-Sense agent (live)",
                   "Streaming-event ingest into TMTML.Engagement; per-subscriber state computation"],
                  ["Risk-Score agent (live)",
                   "Per-subscriber churn probability with confidence + time-to-action"],
                  ["Outcome-Attribute agent (live)",
                   "14-field audit row composition and LEDGER write"],
                  ["Adaptive Card flow",
                   "Mock VP Retention concession-approval flow with Entra identity resolution"],
                  ["LEDGER query surface",
                   "Drill-through from Power BI dashboard tile to specific trace_id"],
                  ["Audit-row export",
                   "Quarterly evidence packet ZIP with manifests, policies, and audit-row extract"],
                  ["Privacy/consent state model",
                   "Per-subscriber consent state with W2-ready integration shim"],
                  ["Manifest versioning",
                   "Git-backed manifests for agents, policy, schema; PR-merge discipline"],
                  ["Shadow-run replay",
                   "Replay against 4 historical quarters of synthetic engagement data"],
                  ["Demo console",
                   "Three canonical demo paths · 3–5 minutes each · narrated for Disney stakeholders"],
                  ["Self-host install",
                   "Local Docker + Ollama (offline-capable for sensitive demo environments)"],
              ],
              col_widths=[2.2, 4.1])

    add_h(doc, "11.2 · Three canonical demo paths", level=2)
    add_body(doc, "Three demo paths planned for Disney stakeholder reviews, each ~3–5 minutes:")
    add_body(doc,
             "Demo Path A · The intervention loop. Pick a synthetic subscriber from the "
             "at-risk cohort. Watch Engagement-Sense ingest the subscriber's last 30 days of "
             "events; watch Risk-Score classify; watch Content-Match propose a personalized "
             "slate; watch Offer-Compose produce a concession-tier proposal; watch the Adaptive "
             "Card fire to the VP Retention's device; click Approve; watch the audit row write.")
    add_body(doc,
             "Demo Path B · The retention-NPV trending. Open the Power BI retention dashboard. "
             "Show the per-cohort NPV trending over 90 days of shadow-run data. Drill through "
             "into a single high-NPV cell — get the contributing audit rows. Drill into one row — "
             "see the full reasoning trace, the version stamps, the HITL approver's identity.")
    add_body(doc,
             "Demo Path C · The audit trail. Pick a synthetic subscriber. Open the audit-row "
             "stream for that trace_id. Show the full chain of custody from engagement signal "
             "through risk score, through offer compose, through HITL approval, through "
             "dispatch, through outcome attribution. Show the manifest version and policy "
             "version stamps. Show the consent-state attestation captured at decision time.")

    add_h(doc, "11.3 · Prototype technology stack", level=2)
    add_table(doc,
              ["Layer", "Stack"],
              [
                  ["Frontend",            "React + Vite · Tailwind · Power BI Embedded"],
                  ["Backend",             "Node.js (Fastify) · pg-pool"],
                  ["Persistence",         "PostgreSQL (LEDGER + state)"],
                  ["Identity",            "Entra ID (mock binding for prototype)"],
                  ["Inference",           "Local Ollama (Llama-class model) for offline demos · Azure OpenAI for online demos"],
                  ["Orchestration",       "Lightweight DAG runner + manifest loader (CAF-shaped contract)"],
                  ["Observability",       "OpenTelemetry traces + LangFuse-equivalent local span store"],
                  ["Schema validation",   "Zod (TMTML.* shapes)"],
                  ["Demo packaging",      "Docker Compose · 'docker compose up' single-command boot"],
                  ["Source control",      "Git repo with manifest history; PR-template for policy edits"],
                  ["Synthetic data",      "100K synthetic subscriber population · 4 historical quarters of events"],
              ],
              col_widths=[1.6, 4.7])

    add_h(doc, "11.4 · Mapping prototype → APEX architecture", level=2)
    add_body(doc,
             "The prototype is deliberately a subset of the full APEX architecture — enough to "
             "make the design concrete for Disney stakeholders, not so much that it requires a "
             "full Azure/Fabric/Foundry deployment to demo:")
    add_table(doc,
              ["Prototype element", "APEX architecture element", "Disney production surface"],
              [
                  ["Local Postgres",          "LEDGER plane (audit row store)",
                   "Azure Fabric Lakehouse + Postgres index in production"],
                  ["Mock channel APIs",        "Integration plane",
                   "Disney's actual channel stack (push, email, in-app)"],
                  ["Mock Adaptive Card device","Experience plane",
                   "VP Retention's actual Microsoft Teams environment"],
                  ["Local Ollama model",       "Decision plane (LLM)",
                   "Azure OpenAI via AEF LLM Gateway (CAF substrate)"],
                  ["Lightweight DAG runner",   "Runtime plane",
                   "CAF Dynamic Orchestrator + A2A Swarm in production"],
                  ["Synthetic subscriber data","TMTML.Subscriber + TMTML.Engagement",
                   "Disney's actual streaming-event stream"],
                  ["Mock consent state",       "Privacy/consent decision evidence",
                   "Disney's actual consent platform"],
              ],
              col_widths=[1.9, 2.0, 2.4])

    add_h(doc, "11.5 · Prototype limitations and next steps", level=2)
    add_body(doc, "The prototype is deliberately scoped for Wave 1 demonstration. It does not:")
    add_bullets(doc, [
        "Connect to real Disney+ engagement event stream — synthetic data only; real integration is W2 scope",
        "Run on Azure — everything is local (Docker + Ollama on a laptop); Azure migration is W2",
        "Handle multi-cohort or cross-brand orchestration — single cohort family scope only",
        "Implement real Entra authentication — hardcoded VP Retention persona for demo; Entra integration is W2",
        "Produce a true Privacy evidence packet — audit timeline view demonstrates the substrate; packet assembly is W2",
        "Activate Content-Match, Offer-Compose, or Outreach-Dispatch agents — those land in W2",
    ])
    add_body(doc,
             "These are not gaps in the architecture — they are intentional scope boundaries. "
             "The W1 prototype demonstrates the agentic flow end-to-end with enough fidelity to "
             "make the Wave 1 exit review credible. Wave 2 closes each gap with the appropriate "
             "Disney production system.")

    pagebreak(doc)

    # ============ APPENDIX A · GLOSSARY ============
    add_h(doc, "Appendix A · Glossary", level=1)
    add_body(doc, "Selected terms used in this document with APEX- or streaming-specific meaning.")
    add_table(doc,
              ["Term", "Definition"],
              [
                  ["A2A Swarm",
                   "Cross-app agent collaboration primitive in CAF; the bridge by which APEX scenarios dispatch into CAF runtime."],
                  ["Adaptive Card",
                   "Microsoft Teams rich-card format used by TMT-MED-01 for HITL approval surfaces."],
                  ["Agent",
                   "A bounded unit of intent + tool-use that emits typed output and writes an audit row per decision."],
                  ["APEX",
                   "Agentic Platform for Enterprise eXecution — Deloitte's reference architecture for industry scenarios on Microsoft-native deployment."],
                  ["At-risk cohort",
                   "Subscriber group whose engagement signal indicates elevated churn probability within an actionable intervention window."],
                  ["Audit row",
                   "14-field structured record emitted by every agent decision; primary substrate of personalization-evidence and KPI attribution."],
                  ["CAF",
                   "Converge Agentic Framework — production platform substrate that TMT-MED-01 consumes via A2A."],
                  ["Concession",
                   "An economically-loaded retention offer (price freeze, discount, bundle attach, exclusivity) requiring policy gating."],
                  ["Content-Match",
                   "TMT-MED-01 agent that proposes personalized content surfaces respecting rights windows."],
                  ["Disney Bundle",
                   "Cross-brand bundle pairing Disney+, Hulu, and ESPN+; primary cross-streaming attach surface."],
                  ["Engagement-Sense",
                   "TMT-MED-01 agent that ingests streaming-platform events and normalizes to TMTML.Engagement."],
                  ["Entra",
                   "Microsoft identity platform; resolves named-persona HITL bindings at runtime."],
                  ["Fabric",
                   "Microsoft unified analytics platform; hosts Bronze/Silver/Gold data assets used by TMT-MED-01 agents in production."],
                  ["Foundry",
                   "Microsoft Azure AI Foundry; hosts model endpoints, tool registrations, and prompt management for APEX agents."],
                  ["HITL",
                   "Human-in-the-loop gate; first-class orchestration primitive declared in manifest."],
                  ["Intervention loop time",
                   "Elapsed time from engagement signal to dispatched intervention; the headline operational KPI for TMT-MED-01."],
                  ["LEDGER",
                   "The APEX audit-row persistence + retrieval plane; substrate of personalization evidence and KPI attribution."],
                  ["Manifest",
                   "Git-versioned declarative artifact describing agents, policies, schemas, and tenant subscriptions."],
                  ["NPV per intervention",
                   "Net present value of an intervention's incremental retention contribution, attributed to a specific trace_id."],
                  ["Offer-Compose",
                   "TMT-MED-01 agent that composes a retention offer (price/bundle/exclusivity) within the concession-policy envelope."],
                  ["Outcome-Attribute",
                   "TMT-MED-01 agent that observes the outcome window and writes the closeout 14-field audit row."],
                  ["Outreach-Dispatch",
                   "TMT-MED-01 agent that dispatches the intervention via push/email/in-app; gates on consent state."],
                  ["Persona",
                   "A named human role with declared authority, plane bindings, and surfaces. Not generic; named identity."],
                  ["Risk-Score",
                   "TMT-MED-01 agent that computes per-subscriber churn probability with confidence and time-to-action."],
                  ["Rights window",
                   "The contractual time/territory window during which Disney has rights to monetize a piece of content."],
                  ["TMT-MED-01",
                   "APEX Service Catalog ID for Media Subscriber Lifecycle. Anchored on this reference engagement."],
                  ["TMT-MED-03",
                   "APEX Service Catalog ID for Media Rights & Monetization. Secondary service in this engagement."],
                  ["TMTML",
                   "APEX canonical schema family for TMT objects (Subscriber, Engagement, Content, Rights, etc.)."],
                  ["Trace-id",
                   "Immutable correlation identifier that binds every audit row from one decision event."],
                  ["Wave",
                   "Delivery phase: W1 Foundation, W2 Pilot, W3 Scale & Fuse. Anchors commercial envelope and exit gates."],
              ],
              col_widths=[1.7, 4.6])

    pagebreak(doc)

    # ============ APPENDIX B · ARTIFACT CROSS-REFERENCE ============
    add_h(doc, "Appendix B · Artifact Cross-Reference", level=1)
    add_body(doc,
             "This document is one of several APEX TMT-MED-01 artifacts covering Disney's subscriber-"
             "lifecycle automation program. The cross-reference below shows which artifact "
             "serves which audience.")
    add_table(doc,
              ["Artifact", "Primary audience", "Purpose"],
              [
                  ["APEX-TMT-MED-01-Walkthrough.docx (this)",
                   "Streaming COO · CFO · VP Retention · CTO/CDO · External Audit",
                   "Executive walkthrough of the value-delivery chain"],
                  ["TMT-MED-01-Sprint-Plan.xlsx",
                   "Engagement Manager · Lead Engineer · Scrum lead",
                   "22 sprint use cases with story points, dependencies, exit criteria"],
                  ["TMT-MED-01-Resource-Plan.xlsx",
                   "DMTSP PMO · Engagement Manager",
                   "Resource loading by phase, role, sprint"],
                  ["TMT-MED-01-ROI-Case.html",
                   "CFO · Streaming COO · Steering Committee",
                   "Retention-NPV uplift case with sensitivity analysis"],
                  ["TMT-MED-01 service spec (APEX Service Catalog)",
                   "APEX Practice TMT lead · Engagement Architect",
                   "Service composition contract: schemas, agents, SLOs, persona"],
                  ["APEX-TMT-Reference-Architecture.html",
                   "All audiences",
                   "TMT Practice reference — six sub-Practices, planes, archetypes"],
                  ["APEX-CAF-Integration-Architecture.docx",
                   "AI & Eng leadership · CAF platform lead · APEX TMT Practice lead",
                   "Substrate integration: how TMT-MED-01 consumes CAF via A2A"],
                  ["TMTML schema spec",
                   "Engagement Architect · Data engineer",
                   "Subscriber, Engagement, Content, Rights canonical shapes"],
                  ["APEX_Design.md",
                   "Build team · Practice Lead",
                   "Reference architecture spec for agents, LEDGER, audit row"],
                  ["Professional-APEX-Sellers-Guide.html",
                   "Seller · Account Team",
                   "Runtime plane, audit-row extension, LEDGER feedback loop"],
                  ["Disney 01_account materials",
                   "Account Team · Pursuit Lead",
                   "Account context, Disney organizational priors, prior engagements"],
                  ["W1 prototype repo (planned)",
                   "Lead Engineer · Disney technical reviewers",
                   "Runnable Wave 1 demo · synthetic data · three demo paths"],
                  ["Privacy/Consent governance addendum (planned)",
                   "Disney Privacy Office · External Audit",
                   "Decision-evidence shape and quarterly packet format"],
              ],
              col_widths=[2.3, 2.0, 2.0])

    add_body(doc,
             "Independence posture preserved across every APEX TMT-MED-01 artifact: 'Deloitte's Microsoft "
             "practice,' 'DMTSP,' 'Microsoft platform capabilities,' 'Microsoft-native "
             "deployment.' The terms 'partner,' 'alliance,' and 'strategic alliance' do not "
             "appear in any artifact.",
             italic=True, color=MUTE, size=9)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— End of walkthrough —")
    set_run(r, italic=True, size=10, color=MUTE)

    # save
    target = OUT
    try:
        doc.save(str(target))
    except PermissionError:
        n = 2
        while True:
            alt = target.with_name(target.stem + f"-v{n}" + target.suffix)
            if not alt.exists():
                target = alt
                break
            n += 1
        doc.save(str(target))
        print(f"\n[notice] {OUT.name} was locked — saved as {target.name} instead.")
    print(f"Wrote {target}")


if __name__ == "__main__":
    build()
