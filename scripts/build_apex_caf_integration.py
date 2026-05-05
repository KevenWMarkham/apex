"""Build the APEX + CAF integration architecture document.

Companion to CAF-Scope-Reply-to-Converge-Team.docx. Answers the next-layer
question: now that we agree CAF is the substrate, how does APEX — Deloitte's
industry scenario reference architecture — sit on top of CAF, and how do
APEX scenarios cross practices/apps using CAF A2A?

Output:
  C:\\Stage\\Clients\\Industries\\Consumer\\Retail\\Demos\\APEX-CAF-Integration-Architecture.docx

Strategy: walk the Converge-team doc inline. After each key Converge-team
section, show how APEX answers / extends / consumes.
"""

from __future__ import annotations

import io
import sys
from pathlib import Path

import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import numpy as np

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

OUT = Path(r"C:\Stage\Clients\Industries\Consumer\Retail\Demos\APEX-CAF-Integration-Architecture.docx")
ART = Path(r"C:\Stage\Clients\Industries\Consumer\Retail\Demos\artifacts_apex_caf")
ART.mkdir(parents=True, exist_ok=True)

# ---- palette -----------------------------------------------------------

INK     = "0B0B0F"
INK2    = "4A4A55"
MUTE    = "8a8a90"  # docx-style (no #)
BAND    = "F3F3F5"
HEADBG  = "0B0B0F"
LINE    = "D4D4D7"

# matplotlib-friendly variants (require leading #)
MPL_MUTE = "#8a8a90"
MPL_INK  = "#0B0B0F"
MPL_INK2 = "#4A4A55"

# hex with # for matplotlib
APEX_GOLD   = "#E3B657"
CAF_TEAL    = "#2A7F73"
A2A_VIOLET  = "#8A6FC4"
AMBER       = "#E59A3E"
SKY         = "#4A90D9"
CRIM        = "#C85450"
BG          = "#fafafa"

plt.rcParams.update({
    "font.family": "sans-serif",
    "font.sans-serif": ["Calibri", "Arial", "DejaVu Sans"],
    "figure.facecolor": BG,
    "axes.facecolor": BG,
    "savefig.facecolor": BG,
    "text.color": "#1f1f25",
})


def save(fig, name):
    path = ART / name
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)
    print(f"  wrote {path.name}")
    return path


# ============================================================================
# FIGURES
# ============================================================================

def fig_two_layer_stack():
    """Two-layer stack: APEX scenario layer above CAF platform substrate."""
    fig, ax = plt.subplots(figsize=(10.5, 5.8))
    ax.axis("off")
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 6.5)

    # APEX layer (top)
    ax.add_patch(FancyBboxPatch((0.3, 4.1), 11.4, 2.1,
                                boxstyle="round,pad=0.05,rounding_size=0.12",
                                facecolor=APEX_GOLD, edgecolor="none", alpha=0.12))
    ax.text(6, 6.0, "APEX · Industry Scenario Reference Architecture",
            ha="center", fontsize=13, color=APEX_GOLD, fontweight="bold")
    ax.text(6, 5.72, "(what the agents do · how the business outcome is sold · how audit evidence is produced)",
            ha="center", fontsize=9.5, color="#1f1f25", style="italic")

    apex_blocks = [
        ("Scenario Library\n(736 folders)",       0.5),
        ("Value-Delivery Chain\nS→Sol→UC→Svc→P→K", 2.7),
        ("47 Archetypes\n(orchestration patterns)", 4.9),
        ("Commercial Envelope\n(materiality × tier)", 7.1),
        ("14-field Audit Row\nLEDGER · SOX",       9.3),
    ]
    for lbl, x in apex_blocks:
        ax.add_patch(FancyBboxPatch((x, 4.25), 2.1, 1.05,
                                    boxstyle="round,pad=0.02,rounding_size=0.08",
                                    facecolor=APEX_GOLD, edgecolor="none", alpha=0.92))
        ax.text(x+1.05, 4.77, lbl, ha="center", va="center",
                fontsize=8.5, color="white", fontweight="bold")

    # A2A bridge
    for x in [1.55, 3.75, 5.95, 8.15, 10.35]:
        ax.annotate("", xy=(x, 3.55), xytext=(x, 4.15),
                    arrowprops=dict(arrowstyle="->", color=A2A_VIOLET, lw=2))
    ax.add_patch(FancyBboxPatch((4.0, 3.25), 4.0, 0.45,
                                boxstyle="round,pad=0.02,rounding_size=0.08",
                                facecolor=A2A_VIOLET, edgecolor="none", alpha=0.95))
    ax.text(6.0, 3.475, "A2A — cross-app agent collaboration bridge",
            ha="center", va="center", fontsize=10, color="white", fontweight="bold")

    # CAF layer (bottom)
    ax.add_patch(FancyBboxPatch((0.3, 0.3), 11.4, 2.6,
                                boxstyle="round,pad=0.05,rounding_size=0.12",
                                facecolor=CAF_TEAL, edgecolor="none", alpha=0.12))
    ax.text(6, 2.68, "CAF · Converge Agentic Framework (production substrate)",
            ha="center", fontsize=13, color=CAF_TEAL, fontweight="bold")
    ax.text(6, 2.40, "(how the agents run · hyperscaler-agnostic · LangFuse + Postgres + RBAC + Registry)",
            ha="center", fontsize=9.5, color="#1f1f25", style="italic")

    caf_blocks = [
        ("AEF LLM\nGateway",        0.5),
        ("A2A Swarm\n(Q4 FY26)",    2.15),
        ("Dynamic\nOrchestrator",   3.80),
        ("Guardrails\n(GSO + Eval)", 5.45),
        ("301-tier\nagents",         7.10),
        ("Margin Maker\npattern",    8.75),
        ("LangFuse\n+ Registry",     10.40),
    ]
    for lbl, x in caf_blocks:
        ax.add_patch(FancyBboxPatch((x, 0.55), 1.50, 1.30,
                                    boxstyle="round,pad=0.02,rounding_size=0.08",
                                    facecolor=CAF_TEAL, edgecolor="none", alpha=0.92))
        ax.text(x+0.75, 1.20, lbl, ha="center", va="center",
                fontsize=8, color="white", fontweight="bold")

    ax.text(6, 0.05, "Figure · APEX is WHAT the agents do. CAF is HOW they run. A2A is the bridge.",
            ha="center", fontsize=9, color=MPL_MUTE, style="italic")
    return save(fig, "fig1_two_layer_stack.png")


def fig_plane_mapping():
    """APEX 5-plane ↔ CAF component mapping."""
    fig, ax = plt.subplots(figsize=(10.5, 5.5))
    ax.axis("off")
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 6)

    # APEX planes (left column)
    ax.text(1.5, 5.65, "APEX Planes", ha="center", fontsize=12,
            color=APEX_GOLD, fontweight="bold")
    planes = [
        ("Experience",  "Adaptive Cards, BI, Dashboards"),
        ("Decision",    "Agents, HITL gates"),
        ("Runtime",     "Orchestration, policy dispatch"),
        ("Integration", "Connectors, MCP tools"),
        ("Ledger",      "Audit row + SOX evidence"),
    ]
    for i, (name, detail) in enumerate(planes):
        y = 4.6 - i * 0.95
        ax.add_patch(FancyBboxPatch((0.2, y), 2.7, 0.75,
                                    boxstyle="round,pad=0.02,rounding_size=0.06",
                                    facecolor=APEX_GOLD, edgecolor="none", alpha=0.92))
        ax.text(1.55, y+0.48, name, ha="center", va="center",
                fontsize=10, color="white", fontweight="bold")
        ax.text(1.55, y+0.17, detail, ha="center", va="center",
                fontsize=7.5, color="white", alpha=0.92)

    # CAF components (right column)
    ax.text(9.75, 5.65, "CAF Components", ha="center", fontsize=12,
            color=CAF_TEAL, fontweight="bold")
    caf = [
        ("Teams + Power BI surfaces", "(consumed from APEX)"),
        ("AEF LLM Gateway + Guardrails", "routes + enforces"),
        ("Dynamic Orchestrator + A2A Swarm", "dispatch + cross-app"),
        ("301-tier agents + MCP tools", "bounded capabilities"),
        ("LangFuse + PostgreSQL + Registry", "tracing + persistence"),
    ]
    for i, (name, detail) in enumerate(caf):
        y = 4.6 - i * 0.95
        ax.add_patch(FancyBboxPatch((8.5, y), 3.3, 0.75,
                                    boxstyle="round,pad=0.02,rounding_size=0.06",
                                    facecolor=CAF_TEAL, edgecolor="none", alpha=0.92))
        ax.text(10.15, y+0.48, name, ha="center", va="center",
                fontsize=9, color="white", fontweight="bold")
        ax.text(10.15, y+0.17, detail, ha="center", va="center",
                fontsize=7.5, color="white", alpha=0.92)

    # arrows (mapping)
    for i in range(5):
        y = 4.6 - i * 0.95 + 0.375
        ax.annotate("", xy=(8.45, y), xytext=(2.95, y),
                    arrowprops=dict(arrowstyle="->", color=A2A_VIOLET, lw=1.3))
    ax.text(5.7, 5.15, "One-to-one mapping",
            ha="center", fontsize=10, color=A2A_VIOLET, fontweight="bold")
    ax.text(5.7, 4.88, "each APEX plane is satisfied by a CAF component",
            ha="center", fontsize=8.5, color=MPL_MUTE, style="italic")

    ax.text(6, 0.1, "Figure · APEX's 5-plane reference model maps 1:1 to CAF's production components.",
            ha="center", fontsize=9, color=MPL_MUTE, style="italic")
    return save(fig, "fig2_plane_mapping.png")


def fig_a2a_integration():
    """APEX scenario chain invoking CAF agents via A2A."""
    fig, ax = plt.subplots(figsize=(10.5, 5.5))
    ax.axis("off")
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 5.5)

    # Top: APEX value-delivery chain (one scenario)
    ax.text(6, 5.15, "APEX Scenario — one-time declaration in manifest",
            ha="center", fontsize=11, color=APEX_GOLD, fontweight="bold")
    chain = ["Scenario", "Solution", "Use Case", "Service", "Persona", "KPI"]
    for i, link in enumerate(chain):
        x = 0.5 + i * 1.9
        ax.add_patch(FancyBboxPatch((x, 4.2), 1.55, 0.65,
                                    boxstyle="round,pad=0.02,rounding_size=0.08",
                                    facecolor=APEX_GOLD, edgecolor="none", alpha=0.92))
        ax.text(x+0.775, 4.52, link, ha="center", va="center",
                fontsize=9.5, color="white", fontweight="bold")
        if i < 5:
            ax.annotate("", xy=(x+1.95, 4.52), xytext=(x+1.6, 4.52),
                        arrowprops=dict(arrowstyle="->", color=APEX_GOLD, lw=1.2))

    # Middle: Manifest emits to A2A
    ax.add_patch(FancyBboxPatch((3.0, 2.8), 6.0, 0.7,
                                boxstyle="round,pad=0.02,rounding_size=0.08",
                                facecolor=A2A_VIOLET, edgecolor="none", alpha=0.95))
    ax.text(6, 3.15, "A2A Swarm · cross-app invocation",
            ha="center", va="center", fontsize=11, color="white", fontweight="bold")

    ax.annotate("", xy=(6, 2.85), xytext=(6, 4.15),
                arrowprops=dict(arrowstyle="->", color=A2A_VIOLET, lw=2))
    ax.text(7.3, 3.65, "APEX manifest → A2A dispatch",
            ha="left", fontsize=8.5, color=MPL_INK2, style="italic")

    # Bottom: CAF agents invoked
    ax.text(6, 2.30, "CAF — 301-tier agents (production) · invoked per A2A session",
            ha="center", fontsize=11, color=CAF_TEAL, fontweight="bold")
    agents = [
        ("Forecasting", "8+ live"),
        ("Price", "MVP"),
        ("Promo", "MVP"),
        ("Sourcing", "Margin Maker"),
        ("Restaurants", "Taco Bell α"),
    ]
    for i, (name, status) in enumerate(agents):
        x = 0.4 + i * 2.35
        ax.add_patch(FancyBboxPatch((x, 1.05), 2.0, 1.05,
                                    boxstyle="round,pad=0.02,rounding_size=0.08",
                                    facecolor=CAF_TEAL, edgecolor="none", alpha=0.92))
        ax.text(x+1.0, 1.75, name, ha="center", va="center",
                fontsize=10, color="white", fontweight="bold")
        ax.text(x+1.0, 1.38, f"({status})", ha="center", va="center",
                fontsize=8, color="white", alpha=0.92, style="italic")
        # arrow from A2A down
        ax.annotate("", xy=(x+1.0, 2.15), xytext=(x+1.0, 2.75),
                    arrowprops=dict(arrowstyle="->", color=A2A_VIOLET, lw=1.1))

    # Evidence return
    ax.annotate("", xy=(11.3, 4.5), xytext=(11.3, 1.55),
                arrowprops=dict(arrowstyle="->", color=CRIM, lw=1.3, connectionstyle="arc3,rad=0.2"))
    ax.text(11.5, 3.0, "14-field\naudit row\n↑ LEDGER", fontsize=8, color=CRIM,
            fontweight="bold", va="center")

    ax.text(6, 0.1, "Figure · An APEX scenario chain declared once, dispatched through CAF A2A, executed by CAF agents, audit row returned to APEX LEDGER.",
            ha="center", fontsize=8.5, color=MPL_MUTE, style="italic")
    return save(fig, "fig3_a2a_integration.png")


def fig_r2r_worked_example():
    """Worked example: R2R scenario on CAF via A2A."""
    fig, ax = plt.subplots(figsize=(10.5, 6.5))
    ax.axis("off")
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 6.5)

    # APEX scenario header
    ax.add_patch(FancyBboxPatch((0.3, 5.6), 11.4, 0.75,
                                boxstyle="round,pad=0.02,rounding_size=0.08",
                                facecolor=APEX_GOLD, edgecolor="none", alpha=0.94))
    ax.text(6, 6.05, "APEX · RC-OPS-06 · Record-to-Report Close Automation (Nike W1 build target)",
            ha="center", va="center", fontsize=11, color="white", fontweight="bold")
    ax.text(6, 5.78, "archetype = hierarchical-root + parallel-fanout-gather + hitl-gated-commit",
            ha="center", va="center", fontsize=8.5, color="white", alpha=0.92, style="italic")

    # A2A dispatch band
    ax.add_patch(FancyBboxPatch((0.3, 4.6), 11.4, 0.6,
                                boxstyle="round,pad=0.02,rounding_size=0.08",
                                facecolor=A2A_VIOLET, edgecolor="none", alpha=0.95))
    ax.text(6, 4.9, "A2A Swarm · APEX parent orchestrator publishes session → CAF agents subscribe and execute",
            ha="center", va="center", fontsize=10, color="white", fontweight="bold")

    # 6 APEX-declared agents
    agents = [
        ("Balance-Fetch",    "SAP S/4 → FINML",     SKY),
        ("Interco-Match",    "AR↔AP reconcile",     CAF_TEAL),
        ("FX-Reval",         "period-end rates",    CAF_TEAL),
        ("Accrual-Propose",  "drivers + history",   AMBER),
        ("JE-Compose",       "+ HITL Adaptive Card", A2A_VIOLET),
        ("Evidence-Write",   "14-field audit row",  CRIM),
    ]
    for i, (name, sub, color) in enumerate(agents):
        col = i % 3
        row = i // 3
        x = 0.5 + col * 3.85
        y = 3.25 - row * 1.45
        ax.add_patch(FancyBboxPatch((x, y), 3.5, 1.0,
                                    boxstyle="round,pad=0.02,rounding_size=0.08",
                                    facecolor=color, edgecolor="none", alpha=0.92))
        ax.text(x+1.75, y+0.65, name, ha="center", va="center",
                fontsize=10.5, color="white", fontweight="bold")
        ax.text(x+1.75, y+0.30, sub, ha="center", va="center",
                fontsize=8.5, color="white", alpha=0.92, style="italic")

    # CAF substrate band
    ax.add_patch(FancyBboxPatch((0.3, 0.1), 11.4, 0.6,
                                boxstyle="round,pad=0.02,rounding_size=0.08",
                                facecolor=CAF_TEAL, edgecolor="none", alpha=0.94))
    ax.text(6, 0.40, "CAF · AEF LLM Gateway · Guardrails · LangFuse tracing · PostgreSQL persistence · Agent Registry",
            ha="center", va="center", fontsize=9.5, color="white", fontweight="bold")

    return save(fig, "fig4_r2r_worked_example.png")


def fig_archetype_to_caf():
    """47 APEX archetypes grouped by how they consume CAF primitives."""
    fig, ax = plt.subplots(figsize=(10.5, 5.0))
    ax.axis("off")
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 5.5)

    groups = [
        ("Sequential + Hierarchical",         14, "use Dynamic Orchestrator\nsession routing",     SKY),
        ("Parallel / Fan-out / Gather",       11, "use A2A Swarm\nfor concurrent subscription",    CAF_TEAL),
        ("HITL-gated / Materiality-gated",     8, "use Guardrails + Adaptive\nCard HITL primitives", A2A_VIOLET),
        ("Feedback Loop / Learning",           7, "use LangFuse traces\n+ Registry re-scoring",     AMBER),
        ("Cross-practice Fusion (via A2A)",    7, "use A2A Swarm\nacross Converge apps",           CRIM),
    ]
    for i, (name, count, how, color) in enumerate(groups):
        x = 0.3 + i * 2.35
        ax.add_patch(FancyBboxPatch((x, 1.5), 2.15, 2.7,
                                    boxstyle="round,pad=0.03,rounding_size=0.1",
                                    facecolor=color, edgecolor="none", alpha=0.93))
        ax.text(x+1.075, 3.75, str(count), ha="center", va="center",
                fontsize=28, color="white", fontweight="bold")
        ax.text(x+1.075, 3.05, "archetypes", ha="center", va="center",
                fontsize=8.5, color="white", alpha=0.92)
        ax.text(x+1.075, 2.55, name, ha="center", va="center",
                fontsize=9.5, color="white", fontweight="bold", wrap=True)
        ax.text(x+1.075, 1.90, how, ha="center", va="center",
                fontsize=7.5, color="white", alpha=0.92, style="italic")

    ax.text(6, 5.1, "APEX 47-archetype catalog · grouped by how each archetype consumes CAF substrate",
            ha="center", fontsize=11, color=MPL_INK, fontweight="bold")
    ax.text(6, 0.7, "All 47 archetypes are expressible on CAF primitives — zero require a new orchestration engine.",
            ha="center", fontsize=9, color=MPL_MUTE, style="italic")
    ax.text(6, 0.35, "Total = 47 (10 implemented in APEX today · 37 manifest-only declarations)",
            ha="center", fontsize=9, color=MPL_MUTE, style="italic")
    return save(fig, "fig5_archetypes_to_caf.png")


def fig_scope_delta():
    """What APEX adds on top of CAF — the clean scope delta."""
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(11, 4.8))

    # Left: CAF substrate (what's already built)
    caf_items = ["LLM Gateway", "A2A Swarm", "Orchestrator", "Guardrails",
                 "LangFuse + Postgres", "Agent Registry", "301-tier agents"]
    ax1.barh(range(len(caf_items)), [1]*len(caf_items), color=CAF_TEAL, alpha=0.9, height=0.65)
    ax1.set_yticks(range(len(caf_items)))
    ax1.set_yticklabels(caf_items, fontsize=10)
    ax1.set_xticks([])
    ax1.set_xlim(0, 1.1)
    ax1.set_title("CAF substrate · already built", fontsize=11, color=CAF_TEAL,
                  fontweight="bold", loc="left", pad=10)
    for i in range(len(caf_items)):
        ax1.text(0.5, i, "PROD", ha="center", va="center",
                 fontsize=9, color="white", fontweight="bold")
    ax1.invert_yaxis()
    for s in ["top", "right", "bottom", "left"]:
        ax1.spines[s].set_visible(False)

    # Right: APEX scenario layer (what APEX adds)
    apex_items = ["Scenario library (736)", "Value-delivery chain", "47-archetype catalog",
                  "Commercial envelope", "14-field audit row", "LEDGER + SOX packet",
                  "Wave model (W1/W2/W3)"]
    ax2.barh(range(len(apex_items)), [1]*len(apex_items), color=APEX_GOLD, alpha=0.9, height=0.65)
    ax2.set_yticks(range(len(apex_items)))
    ax2.set_yticklabels(apex_items, fontsize=10)
    ax2.set_xticks([])
    ax2.set_xlim(0, 1.1)
    ax2.set_title("APEX scenario layer · adds on top of CAF", fontsize=11, color=APEX_GOLD,
                  fontweight="bold", loc="left", pad=10)
    for i in range(len(apex_items)):
        ax2.text(0.5, i, "DELOITTE IP", ha="center", va="center",
                 fontsize=9, color="white", fontweight="bold")
    ax2.invert_yaxis()
    for s in ["top", "right", "bottom", "left"]:
        ax2.spines[s].set_visible(False)

    fig.suptitle("Scope delta · CAF substrate + APEX scenario layer = complete solution",
                 fontsize=11.5, color=MPL_INK, fontweight="bold", y=1.00)
    return save(fig, "fig6_scope_delta.png")


# ============================================================================
# DOCX HELPERS (mirror the CAF reply for style consistency)
# ============================================================================

def set_cell_bg(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex.lstrip("#"))
    tcPr.append(shd)


def set_cell_border(cell, color_hex=LINE):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), color_hex.lstrip("#"))
        borders.append(b)
    tcPr.append(borders)


def set_run(run, *, bold=False, italic=False, size=None, color=None, font="Arial"):
    if bold: run.bold = True
    if italic: run.italic = True
    if size is not None: run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor.from_string(color.lstrip("#"))
    if font: run.font.name = font


def add_body(doc, text, *, size=11, italic=False, color=INK, align=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None: p.alignment = align
    r = p.add_run(text)
    set_run(r, size=size, color=color, italic=italic)
    return p


def add_heading(doc, text, *, level=1, accent_color="E3B657"):
    sizes = {1: 17, 2: 14, 3: 12}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    set_run(r, bold=True, size=sizes.get(level, 12), color=INK)
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), accent_color)
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p


def add_bullets(doc, items, *, size=11):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(it)
        set_run(r, size=size, color=INK)


def add_table(doc, header, rows, *, col_widths=None,
              header_fill=HEADBG, header_color="FFFFFF", zebra=True):
    n = len(header)
    tbl = doc.add_table(rows=len(rows) + 1, cols=n)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(header):
        c = tbl.rows[0].cells[j]
        set_cell_bg(c, header_fill)
        set_cell_border(c)
        c.paragraphs[0].text = ""
        r = c.paragraphs[0].add_run(h)
        set_run(r, bold=True, size=10, color=header_color)
        if col_widths: c.width = Inches(col_widths[j])
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            c = tbl.rows[i].cells[j]
            set_cell_border(c)
            if zebra and i % 2 == 0:
                set_cell_bg(c, BAND)
            c.paragraphs[0].text = ""
            r = c.paragraphs[0].add_run(str(v))
            set_run(r, size=10, color=INK)
            if col_widths: c.width = Inches(col_widths[j])
    doc.add_paragraph()
    return tbl


def add_callout(doc, kind, body, *, accent="B8892E", fill="F6F1E4"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = tbl.rows[0].cells[0]
    set_cell_bg(c, fill)
    c.paragraphs[0].text = ""
    p_tag = c.paragraphs[0]
    r_tag = p_tag.add_run(kind.upper())
    set_run(r_tag, bold=True, size=9, color=accent)
    p = c.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    r = p.add_run(body)
    set_run(r, size=10.5, color=INK)
    doc.add_paragraph()


def add_inline_reply(doc, converge_section_label, converge_quote, apex_reply):
    """Styled two-column 'inline reply' block: CAF-side text + APEX answer side-by-side."""
    # small header
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run("INLINE REPLY TO CONVERGE-TEAM DOC · " + converge_section_label.upper())
    set_run(r, bold=True, size=9, color="8A6FC4")
    rPr = r._element.get_or_add_rPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:val"), "80")
    rPr.append(spacing)

    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # left: CAF quote
    cl = tbl.rows[0].cells[0]
    cl.width = Inches(3.1)
    set_cell_bg(cl, "E4F1EE")
    set_cell_border(cl)
    cl.paragraphs[0].text = ""
    p_tag = cl.paragraphs[0]
    r_tag = p_tag.add_run("FROM CONVERGE REPLY")
    set_run(r_tag, bold=True, size=8.5, color="2A7F73")
    p_body = cl.add_paragraph()
    p_body.paragraph_format.space_before = Pt(3)
    r_body = p_body.add_run(converge_quote)
    set_run(r_body, size=10, color=INK, italic=True)

    # right: APEX reply
    cr = tbl.rows[0].cells[1]
    cr.width = Inches(3.4)
    set_cell_bg(cr, "F6F1E4")
    set_cell_border(cr)
    cr.paragraphs[0].text = ""
    p_tag2 = cr.paragraphs[0]
    r_tag2 = p_tag2.add_run("APEX ANSWER")
    set_run(r_tag2, bold=True, size=8.5, color="B8892E")
    p_body2 = cr.add_paragraph()
    p_body2.paragraph_format.space_before = Pt(3)
    r_body2 = p_body2.add_run(apex_reply)
    set_run(r_body2, size=10, color=INK)
    doc.add_paragraph()


def add_kv(doc, pairs):
    tbl = doc.add_table(rows=len(pairs), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(pairs):
        ck, cv = tbl.rows[i].cells
        ck.width = Inches(1.6)
        cv.width = Inches(4.9)
        set_cell_bg(ck, BAND)
        set_cell_border(ck); set_cell_border(cv)
        ck.paragraphs[0].text = ""
        cv.paragraphs[0].text = ""
        rk = ck.paragraphs[0].add_run(k)
        set_run(rk, bold=True, size=10, color=INK)
        rv = cv.paragraphs[0].add_run(v)
        set_run(rv, size=10, color=INK)
    doc.add_paragraph()


def add_eyebrow(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text.upper())
    set_run(r, bold=True, size=9, color=MUTE)
    rPr = r._element.get_or_add_rPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:val"), "80")
    rPr.append(spacing)


def add_figure(doc, path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=Inches(6.3))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    r = cap.add_run(caption)
    set_run(r, italic=True, size=9, color=MUTE)


def add_pagebreak(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


# ============================================================================
# BUILD
# ============================================================================

def build():
    print("Generating figures...")
    f1 = fig_two_layer_stack()
    f2 = fig_plane_mapping()
    f3 = fig_a2a_integration()
    f4 = fig_r2r_worked_example()
    f5 = fig_archetype_to_caf()
    f6 = fig_scope_delta()

    print("Composing docx...")
    doc = Document()
    for section in doc.sections:
        section.page_width  = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = section.bottom_margin = Inches(0.9)
        section.left_margin = section.right_margin = Inches(1.0)
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # ==== HEADER ====
    add_eyebrow(doc, "DELOITTE · INTEGRATION ARCHITECTURE · APEX + CAF")
    p = doc.add_paragraph()
    r = p.add_run("APEX on CAF — scenario framework on a production substrate")
    set_run(r, bold=True, size=22, color=INK)
    p2 = doc.add_paragraph()
    r2 = p2.add_run("How Deloitte's APEX reference architecture integrates with Converge Agentic Framework through A2A")
    set_run(r2, size=13, color=INK2, italic=True)

    doc.add_paragraph()
    add_kv(doc, [
        ("Companion to",   "CAF-Scope-Reply-to-Converge-Team.docx (24 Apr 2026)"),
        ("Audience",       "Converge / AI & Eng leadership · APEX build team · Practice leads · DMTSP PMO"),
        ("Author",         "K. Markham · Deloitte Microsoft Technology Solutions & Practices"),
        ("Date",           "24 April 2026"),
        ("Purpose",        "Extend the CAF-vs-Agentic-Merch scoping reply by showing exactly where APEX sits in the overall architecture and how APEX scenarios consume CAF via A2A"),
        ("Status",         "Integration architecture — reference for W1 build planning"),
    ])

    # ==== EXEC SUMMARY ====
    add_heading(doc, "Executive summary", level=1)
    add_body(doc,
             "The companion memo to the Converge team established that CAF is the production "
             "substrate and that Agentic Merch should inherit it rather than parallel it. That "
             "leaves one question open: if CAF is the substrate, what sits on top of it for the "
             "industry scenario work Deloitte delivers? The answer is APEX — the Agentic Platform "
             "for Enterprise eXecution — Deloitte's reference architecture for industry "
             "scenarios, authored inside DMTSP across the seven Practices (RC, HLS, ER, AXLE, TMT, "
             "TH, ICE).")
    add_body(doc,
             "APEX and CAF are not competitors. They answer different questions. CAF answers how "
             "agents run. APEX answers what the agents do, how the business outcome is sold, and "
             "what evidence survives the decision. The two frameworks compose cleanly through one "
             "integration point: CAF's A2A swarm. APEX scenarios declare their agent fleet and "
             "HITL gates in manifests that dispatch into the A2A swarm; CAF executes; the "
             "14-field audit row flows back to APEX's LEDGER plane.")
    add_callout(doc, "One-sentence position",
                "APEX sits on top of CAF. APEX is WHAT the agents do (scenario library, value-"
                "delivery chain, 47 archetypes, commercial envelope, audit row, SOX evidence). "
                "CAF is HOW the agents run (LLM gateway, A2A swarm, orchestrator, guardrails, "
                "tracing, registry, 301-tier agents). A2A is the bridge — and it is exactly the "
                "primitive the Converge-team reply recommended inheriting from CAF.")

    # ==== SECTION 1 · THE TWO FRAMEWORKS ====
    add_heading(doc, "1 · The two frameworks in one frame", level=1)
    add_body(doc,
             "The cleanest way to see the relationship is to draw the two layers and the bridge "
             "between them. APEX on top — the scenario layer that productizes outcomes for "
             "industry clients. CAF on the bottom — the platform substrate that runs the agents. "
             "A2A as the named bridge — the pattern by which APEX scenarios invoke CAF agents "
             "across apps and practices.")
    add_figure(doc, f1,
               "Figure 1 · APEX (top) is the scenario / business-outcome layer. CAF (bottom) is the "
               "production platform substrate. A2A is the named integration bridge. The arrows up "
               "the diagram represent APEX manifest declarations dispatching into the A2A swarm.")

    add_body(doc,
             "The division of labor is crisp. APEX never re-implements what CAF already does. "
             "CAF never needs industry-scenario knowledge — APEX carries that. The integration "
             "point is manifest-declarative: APEX scenarios say what agents to invoke, in what "
             "pattern, with what HITL gates; CAF's Dynamic Orchestrator + A2A swarm dispatches "
             "the request; CAF's LangFuse + Registry traces and persists; APEX's LEDGER plane "
             "captures the 14-field audit row on top.")

    # ==== SECTION 2 · PLANE MAPPING ====
    add_heading(doc, "2 · Architectural mapping — APEX planes ↔ CAF components", level=1)
    add_body(doc,
             "APEX's reference architecture is organized around five planes: Experience, "
             "Decision, Runtime, Integration, Ledger. Each plane is satisfied by specific "
             "CAF components already in production. The mapping is one-to-one; there are no "
             "planes that require a CAF capability that doesn't exist, and there are no CAF "
             "capabilities that live outside an APEX plane.")
    add_figure(doc, f2,
               "Figure 2 · One-to-one mapping between APEX's five planes and CAF's production "
               "components. Each APEX plane names what the plane does; the matching CAF component "
               "is the production artifact that provides it.")

    add_table(doc,
              ["APEX plane", "What the plane is responsible for", "CAF component that provides it"],
              [
                  ["Experience",  "Named persona surfaces — Adaptive Cards, Power BI dashboards, Teams flows",
                   "Teams + Power BI surfaces (consumed from APEX manifests)"],
                  ["Decision",    "Agents that reason; HITL gates that pause for named-persona approval",
                   "AEF LLM Gateway routes the agent's LLM call; Guardrails enforce the decision envelope"],
                  ["Runtime",     "Orchestration graph; policy dispatch; manifest-driven agent composition",
                   "Dynamic Orchestrator (session-generated) + A2A Swarm (cross-app collaboration)"],
                  ["Integration", "Connectors to client SoR (ERP, CRM, treasury, POS, PLM)",
                   "301-tier agent tool definitions + MCP tool registrations"],
                  ["Ledger",      "14-field audit row; LEDGER persistence; SOX evidence packet",
                   "LangFuse (tracing) + PostgreSQL (persistence) + Agent Registry (catalog)"],
              ],
              col_widths=[1.2, 3.0, 2.3])

    add_callout(doc, "What this mapping means",
                "Every APEX plane is satisfied by a CAF component that is already built and in "
                "production. The APEX plane model is not a parallel architecture to CAF — it is "
                "an outside-in view of the same system, framed around business-outcome "
                "responsibilities (what the plane does for the scenario) rather than platform-"
                "primitive responsibilities (what the component provides to the runtime). "
                "Practice teams author APEX manifests; platform teams operate CAF. The same "
                "system runs.",
                accent="B8892E")

    # ==== SECTION 3 · A2A INTEGRATION ====
    add_heading(doc, "3 · The A2A integration pattern", level=1)
    add_body(doc,
             "CAF's A2A swarm is the Q4 FY26 flagship — cross-app agent collaboration across "
             "Converge products. It is the named answer to Agentic Merch Blocker #1 (\"how does "
             "Agentic Merch call ConCON?\") in the Converge-team reply. It is also the exact "
             "integration surface APEX uses when an industry scenario needs to compose agents "
             "across practices, apps, or tenants.")

    add_figure(doc, f3,
               "Figure 3 · APEX scenario manifests dispatch into the A2A swarm. The APEX "
               "value-delivery chain (Scenario → Solution → Use Case → Service → Persona → KPI) "
               "is declared once in the manifest; the manifest emits an A2A session; CAF's "
               "301-tier agents subscribe and execute; the 14-field audit row returns up the "
               "stack to APEX LEDGER.")

    add_heading(doc, "3.1  How an APEX scenario dispatches through A2A", level=3)
    add_bullets(doc, [
        "APEX manifest declares the agent fleet (e.g., the six agents of RC-OPS-06 R2R) + the archetype pattern (hierarchical-root + parallel-fanout-gather + hitl-gated-commit) + materiality tiers + named personas.",
        "APEX parent orchestrator (a thin dispatch layer) reads the manifest and publishes an A2A session to the CAF swarm, tagged with scenario_id and trace_id.",
        "CAF A2A Swarm routes the session to matching CAF 301-tier agents based on capability tags in the Registry (e.g., 'forecasting', 'price', 'promo', 'sourcing').",
        "CAF Dynamic Orchestrator decides per-invocation routing (which model, which tools, which guardrail tier) — LLM-based, session-generated, not static.",
        "CAF Guardrails (GSO + Prompt Eval) enforce policy at inference time; APEX policy manifest layers the scenario-specific materiality thresholds on top.",
        "CAF LangFuse captures every agent trace; APEX Evidence-Write agent (one of the six in R2R) composes the 14-field audit row from LangFuse traces plus scenario context and writes to LEDGER.",
        "APEX LEDGER returns the closeout audit row to the Experience plane; Adaptive Card or BI dashboard renders for the named persona.",
    ])

    add_heading(doc, "3.2  Why A2A is the correct bridge (not a new protocol)", level=3)
    add_body(doc,
             "APEX could have specified its own cross-app agent protocol. It does not, "
             "deliberately — the only correct answer to 'how do APEX scenarios cross practices or "
             "tenants' is 'through the A2A swarm,' for three reasons:")
    add_bullets(doc, [
        "A2A is already built and is the Q4 FY26 CAF flagship. Specifying a parallel protocol would violate the 'lack of reuse' anti-pattern in the Converge-team reply §3.1.",
        "A2A is designed to be app-neutral. APEX scenarios are cross-practice (a supply-chain scenario in RC may need to call a forecasting agent that HLS also uses); A2A is the right abstraction for that cross-practice routing.",
        "A2A already carries the identity + trace propagation APEX needs for the 14-field audit row. Reinventing that plumbing is pure waste.",
    ])

    # ==== SECTION 4 · INLINE REPLIES ====
    add_pagebreak(doc)
    add_heading(doc, "4 · Inline APEX replies to the Converge-team doc", level=1)
    add_body(doc,
             "This section walks the Converge-team reply section by section. For each key "
             "position in that memo, APEX's answer — how an APEX-authored industry scenario "
             "consumes, extends, or composes that CAF capability — is shown side by side.")

    add_inline_reply(doc,
        "§1 · What CAF actually is — AEF LLM Gateway",
        "Cloud-agnostic + model-agnostic. Live routes to AWS Bedrock, Azure OpenAI, Vertex AI. Tenants pick; the gateway abstracts.",
        "APEX scenario manifests declare model requirements at the capability level (reasoning tier, context window, latency SLO). CAF's LLM Gateway satisfies the capability contract against whichever hyperscaler the tenant chose. APEX never pins a scenario to a specific LLM — the manifest contracts the capability, not the vendor.")

    add_inline_reply(doc,
        "§1 · What CAF actually is — A2A Swarm",
        "Cross-app agent collaboration across Converge products. This is the agent-to-agent pattern Agentic Merch needs.",
        "APEX's 47-archetype catalog includes 11 Parallel/Fan-out/Gather archetypes and 7 Cross-practice Fusion archetypes — all of them dispatch through A2A. A2A is the substrate APEX assumes when declaring cross-practice scenarios (e.g., a cold-chain scenario composing RC forecasting with HLS regulatory guardrails).")

    add_inline_reply(doc,
        "§1 · What CAF actually is — Dynamic Orchestrator",
        "LLM-based, session-generated routing. Not static DAG dispatch. Adapts per invocation.",
        "APEX scenarios declare archetypes (hierarchical, parallel, feedback-loop, HITL-gated); the Dynamic Orchestrator is what expresses them at runtime. APEX never specifies a static DAG — every scenario is expressed as a pattern from the 47-archetype catalog, and the Orchestrator realizes the pattern session-by-session.")

    add_inline_reply(doc,
        "§1 · What CAF actually is — Guardrails (GSO + Prompt Eval)",
        "PII detection, prompt-injection defense, toxicity filtering, image-compliance checks. Production-grade.",
        "APEX's commercial envelope layers on top of CAF Guardrails. CAF provides the technical floor (no PII leakage, no injection); APEX adds the business ceiling (no auto-post above materiality tier X, no HITL-skipped if confidence below 0.9, no decision without audit-row attribution).")

    add_inline_reply(doc,
        "§1 · What CAF actually is — 301-tier agents in production",
        "Forecasting: 8+ live. Price/Promo: MVP. Restaurants: alpha at Taco Bell.",
        "APEX scenarios consume 301-tier agents through the CAF Agent Registry. When an APEX scenario (e.g., RC-SUPCHN-01 Cold-Chain Excursion) needs forecasting, the manifest points at 'cap:forecasting:retail'; CAF returns the best-matching 301-tier agent. APEX authors zero new forecasting agents where CAF already has one.")

    add_inline_reply(doc,
        "§1 · What CAF actually is — Margin Maker pattern",
        "Three-agent pattern: Demand Planner + Pricing Advisor + Sourcing Advisor. Pre-designed E2E flow blueprint.",
        "The Margin Maker pattern is an APEX archetype instance — it maps 1:1 to the 'Parallel-fanout with Hierarchical-root' archetype and to Agents 3/4/6/7 of the Agentic Merch 16. APEX's archetype catalog lists it; CAF provides the live blueprint. Same pattern, named on both sides.")

    add_inline_reply(doc,
        "§1 · What CAF actually is — LangFuse + Postgres + RBAC + Registry",
        "Tracing, persistence, access control, agent catalog. The execution substrate.",
        "APEX's Ledger plane sits on top. LangFuse provides the technical traces; APEX's 14-field audit row is composed from those traces plus scenario context (scenario_id, manifest_version, materiality_tier, HITL_actor). APEX never stores what CAF already stores — it adds the scenario-level evidence shape.")

    add_pagebreak(doc)

    add_inline_reply(doc,
        "§2 · How CAF maps to Agentic Merch gaps — Blocker #1",
        "How does Agentic Merch call ConCON? → A2A Swarm (adopt, don't redesign)",
        "APEX scenarios that span practices use A2A natively. The RC-OPS-06 R2R scenario, for example, must call the HLS 340B-accrual agent for any Nike entity that operates a retail pharmacy — that cross-practice call is an A2A session, not an APEX-specific protocol. The APEX manifest says 'dispatch'; A2A routes.")

    add_inline_reply(doc,
        "§2 · How CAF maps to Agentic Merch gaps — Blocker #2",
        "Whose LLM handles cross-platform agent calls? → AEF LLM Gateway (route through it, drop the debate)",
        "APEX manifests contract on capability, not vendor. Every APEX scenario's agent manifest declares 'model: cap:reasoning-advanced' or 'model: cap:narrative-long-context' — the AEF LLM Gateway resolves to the tenant's chosen backend. APEX is LLM-agnostic by contract; CAF is LLM-agnostic by runtime.")

    add_inline_reply(doc,
        "§5 · MVP scope — what stays in",
        "Merch-domain schema · KPIs · agent specializations · HITL surfaces · policy corpus · SoR integration shims",
        "Every one of these is an APEX scenario layer deliverable. The merch-domain schema is FINML/MERML; the KPIs are per-scenario in the 736 scenario folders; the agent specializations are the Practice-scoped 401/501-tier agents that extend CAF 301-tiers; the HITL surfaces are APEX Adaptive Cards with named personas; the policy corpus is the APEX LEDGER policy manifest; the SoR shims are APEX Integration plane artifacts.")

    add_inline_reply(doc,
        "§5 · MVP scope — what comes out",
        "LLM routing · A2A · orchestration · guardrails · tracing · registry · forecasting/price/promo primitives",
        "APEX does not ship any of these. All seven live in CAF. APEX manifests reference them, consume them, layer policy on top of them — but APEX never re-implements. This is exactly the separation the Converge-team reply recommended; APEX honors it by construction.")

    # ==== SECTION 5 · WORKED EXAMPLE ====
    add_heading(doc, "5 · Worked example — an APEX R2R scenario running on CAF via A2A", level=1)
    add_body(doc,
             "To make the integration concrete, here is how a specific APEX scenario runs "
             "end-to-end on CAF. The scenario is RC-OPS-06 Record-to-Report Close Automation — "
             "the Nike Wave 1 build target described in the APEX R2R deep-dive. Six agents, "
             "hierarchical + parallel + HITL-gated pattern, 14-field audit row, SOX evidence "
             "packet at quarter-end.")

    add_figure(doc, f4,
               "Figure 4 · RC-OPS-06 R2R scenario running on CAF. APEX scenario declares the "
               "archetype and the six-agent fleet; A2A dispatches the session; CAF substrate "
               "(LLM Gateway, Guardrails, LangFuse, Postgres, Registry) executes underneath.")

    add_heading(doc, "5.1  What the APEX scenario contributes", level=3)
    add_bullets(doc, [
        "Scenario manifest declaring the six agents (Balance-Fetch, Interco-Match, FX-Reval, Accrual-Propose, JE-Compose, Evidence-Write) and the archetype composition.",
        "Materiality policy manifest tuned to Nike's entity tiers (Tier 1 = always HITL; Tier 3 + conf≥0.90 = auto-post).",
        "Named persona bindings: Deirdre Chen (Corporate Controller) has HITL authority; Entra group resolves at runtime.",
        "FINML schema (SubLedger, IntercoMatch, FXRate, JournalProposal) and SOXML.AuditEvidence schema.",
        "14-field audit row construction logic — Evidence-Write agent composes one row per decision from LangFuse trace + scenario context.",
        "Quarterly SOX evidence packet generator — consumes LEDGER, produces auditor-grade ZIP.",
    ])

    add_heading(doc, "5.2  What CAF contributes", level=3)
    add_bullets(doc, [
        "AEF LLM Gateway routes each agent's LLM call to Nike's chosen backend (Azure OpenAI in this engagement).",
        "A2A Swarm dispatches the six-agent session and propagates trace_id through every agent call.",
        "Dynamic Orchestrator realizes the hierarchical + parallel pattern at runtime, session by session.",
        "Guardrails (GSO + Prompt Eval) enforce PII-redaction and injection defense at inference time.",
        "301-tier agents satisfy capability tags (e.g., 'cap:accrual-reasoning' is matched from Registry).",
        "LangFuse captures every agent trace; PostgreSQL persists; RBAC enforces per-persona access to traces.",
        "Agent Registry catalogs every agent deployment so the APEX manifest can reference by capability, not by URL.",
    ])

    add_heading(doc, "5.3  What gets built for Nike W1", level=3)
    add_body(doc,
             "W1 for Nike is the APEX scenario layer for RC-OPS-06 — the six agent manifests, the "
             "materiality policy, the FINML/SOXML schemas, the Adaptive Card, the SAP S/4HANA "
             "connector, and the LEDGER evidence-packet generator. W1 does not build any CAF "
             "primitive; it consumes the ones that exist. Estimated W1 effort: ~45 story points "
             "(from the APEX R2R deep-dive). If W1 were also building the platform substrate from "
             "scratch, that estimate would triple — which is the 50-65% scope reduction the "
             "Converge-team reply quoted, seen from the other direction.")

    # ==== SECTION 6 · 47 ARCHETYPES ====
    add_heading(doc, "6 · The 47-archetype catalog — all expressible on CAF primitives", level=1)
    add_body(doc,
             "APEX ships a catalog of 47 orchestration archetypes. 10 are implemented in APEX "
             "today (used by the 35 featured scenarios across the seven practices); 37 are "
             "manifest-only declarations for scenarios that haven't been promoted to featured "
             "status yet. Every one of the 47 is expressible on CAF primitives. Zero require a "
             "new orchestration engine.")

    add_figure(doc, f5,
               "Figure 5 · The 47 archetypes grouped by how they consume CAF substrate. "
               "Sequential + Hierarchical archetypes ride the Dynamic Orchestrator's session "
               "routing. Parallel / Fan-out / Gather archetypes ride A2A for concurrent "
               "subscription. HITL-gated / Materiality-gated archetypes ride CAF Guardrails + "
               "Adaptive Card primitives. Feedback Loop / Learning archetypes ride LangFuse "
               "traces + Registry re-scoring. Cross-practice Fusion archetypes ride A2A across "
               "Converge apps.")

    add_callout(doc, "Why this matters",
                "If any of the 47 archetypes required a capability CAF doesn't have, APEX would "
                "need a parallel platform. None do. That is the structural evidence that APEX is "
                "the scenario layer on top of CAF — not a competitor, not a second runtime, and "
                "not a new orchestration engine. The archetype catalog is the contract; CAF is "
                "the implementation.",
                accent="2A7F73", fill="E4F1EE")

    # ==== SECTION 7 · LEDGER OVERLAY ====
    add_heading(doc, "7 · LEDGER + the 14-field audit row — APEX's overlay on CAF tracing", level=1)
    add_body(doc,
             "CAF's LangFuse + PostgreSQL provides production-grade agent tracing and "
             "persistence. APEX's LEDGER plane sits on top and adds the scenario-level evidence "
             "shape auditors require. The division is precise:")
    add_table(doc,
              ["Concern", "CAF provides", "APEX adds on top"],
              [
                  ["Agent trace capture",       "LangFuse spans for every LLM call, tool call, agent handoff", "-"],
                  ["Agent trace persistence",   "PostgreSQL backing LangFuse",                                  "-"],
                  ["Agent access control",      "RBAC on traces + registry",                                    "-"],
                  ["Scenario-level evidence",   "-",                                                            "14-field audit row composed from traces + scenario context"],
                  ["Policy attribution",        "-",                                                            "manifest_version, policy_version, materiality_tier captured per decision"],
                  ["HITL actor attribution",    "-",                                                            "named persona (Entra identity) + timestamp + confidence at decision"],
                  ["SOX evidence packet",       "-",                                                            "quarterly ZIP: manifests + policy + audit-row extract + tie-out reconciliation"],
                  ["Auditor-grade export",      "-",                                                            "LEDGER plane with auditor-configured sampling + replay"],
              ],
              col_widths=[1.8, 2.4, 2.3])

    add_body(doc,
             "This is exactly the 'schema/KPI differentiation' the Converge-team reply §3.3 "
             "called out as the right kind of APEX layering. The differentiation is structural "
             "(audit row is a different shape with different guarantees than LangFuse spans), "
             "not cosmetic.")

    # ==== SECTION 8 · SCOPE DELTA ====
    add_heading(doc, "8 · The complete scope picture — CAF + APEX together", level=1)
    add_body(doc,
             "Reframing the Converge-team reply §5 scope delta one layer up: CAF provides seven "
             "platform primitives in production; APEX provides seven scenario-layer primitives "
             "in Practice-authored form. The two sets are disjoint by design.")
    add_figure(doc, f6,
               "Figure 6 · Complete scope picture. Left: seven CAF substrate primitives, all in "
               "production. Right: seven APEX scenario-layer primitives, all Deloitte IP authored "
               "by DMTSP practices. Together they compose the complete agentic solution — CAF "
               "runs the agents; APEX productizes the scenarios.")

    # ==== SECTION 9 · WHAT THIS MEANS FOR AGENTIC MERCH ====
    add_heading(doc, "9 · What this means for Agentic Merch specifically", level=1)
    add_body(doc,
             "The Agentic Merch MVP — the subject of the Converge-team reply — is best "
             "understood as an APEX scenario cluster running on CAF. Specifically:")
    add_bullets(doc, [
        "The 16 Agentic Merch agents = an APEX scenario chain under practice RC, domain pricing-revenue or supply-chain, with Merchandising-specific use cases and KPIs.",
        "Agents 4, 6, 7 (Forecasting · Price · Promo) = 301-tier CAF agents already running; APEX manifests reference them by capability.",
        "Agents 3, 4, 6, 7 E2E flow (Demand → Price → Sourcing) = the Margin Maker pattern, which is an APEX archetype instance realized on the CAF blueprint.",
        "Remaining 9 agents (1-2, 5, 8-16) = Practice-scoped 401/501-tier specializations that APEX practices author on top of the CAF 301-tier primitives.",
        "Merch-domain schema + KPIs + HITL surfaces + policy corpus = APEX scenario layer deliverables — the intellectual property the Merchandising Practice authors.",
        "Cross-app calls (e.g., Agentic Merch → ConCON) = A2A Swarm sessions initiated by APEX manifest dispatch.",
    ])

    add_callout(doc, "Reframe",
                "Agentic Merch is not a new platform. It is an APEX scenario cluster, authored "
                "by the Merchandising Practice, that consumes CAF through A2A. The MVP is the "
                "first few scenarios in that cluster (likely the Margin Maker-anchored three) "
                "plus the merch-domain schema/KPI/HITL layer. Everything else — LLM routing, "
                "orchestration, guardrails, tracing, the agents themselves — comes from CAF.",
                accent="B8892E")

    # ==== SECTION 10 · NEXT STEPS ====
    add_heading(doc, "10 · Recommended next steps", level=1)
    add_body(doc,
             "Three concrete actions. The first two are coordination; the third is the W1 build "
             "scope.")
    add_table(doc,
              ["#", "Action", "Owner", "Timing"],
              [
                  ["1",
                   "Add this integration document to the working-session agenda proposed in the Converge-team reply step #1. Walk the plane mapping (Fig 2), the A2A integration (Fig 3), and the worked R2R example (Fig 4) with Saurabh + CAF platform lead + APEX Practice leads.",
                   "K. Markham + Saurabh + APEX PMO",
                   "With the working session (this week)"],
                  ["2",
                   "Publish the APEX scenario → CAF capability mapping as a living document inside DMTSP. For each of the 736 APEX scenario folders, the README names the CAF capabilities it consumes by tag (cap:forecasting, cap:price, cap:accrual-reasoning, etc.). This becomes the contract surface between the two teams.",
                   "APEX Practice leads + CAF Registry owner",
                   "Within 10 business days"],
                  ["3",
                   "Scope W1 Agentic Merch MVP against this integration picture. Scope includes the merch-domain APEX scenario layer (schema/KPIs/HITL/policy/specialization agents); scope excludes any CAF primitive. Validate the scope reduction estimate (50-65%) against the actual delta the scenario layer requires.",
                   "Agentic Merch lead + Merchandising Practice lead",
                   "W1 architecture lock"],
              ],
              col_widths=[0.35, 4.2, 1.55, 0.9])

    # ==== CLOSING ====
    add_heading(doc, "11 · Closing position", level=1)
    add_callout(doc, "Executive framing",
                "APEX and CAF compose. APEX is the scenario framework that productizes industry "
                "outcomes (736 scenarios, 47 archetypes, 14-field audit row, SOX evidence, "
                "commercial envelope, wave model). CAF is the production substrate that runs the "
                "agents (LLM gateway, A2A swarm, orchestrator, guardrails, tracing, registry, "
                "301-tier agents). A2A is the named integration bridge. Every APEX scenario "
                "dispatches through A2A. Every CAF capability is consumed via APEX manifest. "
                "Agentic Merch is the first major cluster of APEX scenarios to ride this "
                "integration — and the architecture it deserves is exactly this one.",
                accent="B8892E")

    add_body(doc,
             "The companion Converge-team reply established that the AI & Eng circle should "
             "inherit CAF. This document establishes that when they do, APEX is the scenario "
             "framework waiting on top, authored by DMTSP practices, ready to consume CAF "
             "through A2A from Day 1 of W1.")

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("— K. Markham")
    set_run(r, bold=True, size=11, color=INK)
    p2 = doc.add_paragraph()
    r2 = p2.add_run("Deloitte Microsoft Technology Solutions & Practices")
    set_run(r2, size=10, color=INK2)
    p3 = doc.add_paragraph()
    r3 = p3.add_run("kmarkham@deloitte.com")
    set_run(r3, size=10, color=INK2, italic=True)

    # ==== APPENDIX ====
    add_pagebreak(doc)
    add_heading(doc, "Appendix · Companion artifacts", level=1)
    add_table(doc,
              ["Artifact", "Purpose", "Location"],
              [
                  ["CAF-Scope-Reply-to-Converge-Team.docx",
                   "The companion reply this document builds on",
                   "Consumer/Retail/Demos/"],
                  ["CAF-vs-Agentic-Merch-Scope-Analysis.md",
                   "Technical basis for the CAF reply",
                   "Consumer/Retail/Demos/"],
                  ["APEX-Stacked-Architecture-Narrated.html",
                   "Live interactive companion — clickable architecture, flows, archetypes",
                   "APEX/docs/reference/"],
                  ["APEX-R2R-deep-dive.md",
                   "The Nike W1 scenario detail used in the worked example (§5)",
                   "APEX/docs/scenarios/RC/operations-workforce/RC-OPS-06-…/"],
                  ["APEX-R2R-Walkthrough.docx",
                   "Executive walkthrough of the R2R scenario across all chain links",
                   "APEX/docs/scenarios/RC/operations-workforce/RC-OPS-06-…/"],
                  ["APEX_Design.md",
                   "APEX 5-plane reference architecture spec; 14-field audit row spec",
                   "APEX/docs/APEX - Design and Build/"],
                  ["Professional-APEX-Sellers-Guide.html",
                   "Sellers guide — Runtime plane, LEDGER feedback loop, Appendix AA",
                   "APEX/docs/book/"],
                  ["47-archetype catalog",
                   "Canonical archetype list with implementation status",
                   "APEX/packages/apex-orchestrator/src/apex_orchestrator/archetypes/catalog.py"],
              ],
              col_widths=[2.4, 2.4, 1.7])

    add_body(doc,
             "Independence posture preserved across this document: 'Deloitte's Microsoft "
             "practice,' 'DMTSP,' 'Microsoft platform capabilities,' 'Microsoft-native "
             "deployment.' The terms 'partner,' 'alliance,' and 'strategic alliance' do not "
             "appear.",
             italic=True, color=MUTE, size=9)

    doc.save(str(OUT))
    print(f"\nWrote {OUT}")


if __name__ == "__main__":
    build()
