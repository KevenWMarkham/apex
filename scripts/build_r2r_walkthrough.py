"""Build the R2R (Record-to-Report) walkthrough .docx with inline figures.

Models the same shape as APEX-Cold-Chain-Excursion-Walkthrough.docx:
cover -> Section 1 Opening Frame -> Section 2 Context ->
Chain Links 1-6 (Scenario / Solution / Use Cases / Services / Personas / KPIs) ->
Section 9 Wave Progression -> Section 10 Commercial Envelope ->
Appendix A Glossary -> Appendix B Artifact Cross-Reference.

Generates 10 matplotlib figures into artifacts/ then composes the docx.

Reference client: Nike (Wave 1 build target).
Scenario folder: docs/scenarios/RC/operations-workforce/RC-OPS-06-record-to-report-automation/
"""

from __future__ import annotations

import io
import sys
from pathlib import Path

import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import numpy as np

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

# ----- paths --------------------------------------------------------------

ROOT = Path(r"C:\Stage\Clients\Industries\APEX")
SCEN = ROOT / "docs/scenarios/RC/operations-workforce/RC-OPS-06-record-to-report-automation"
ART = SCEN / "artifacts"
ART.mkdir(parents=True, exist_ok=True)
DOCX = SCEN / "APEX-R2R-Walkthrough.docx"

# ----- palette (matches cold-chain visual language) ----------------------

INK0 = "#0B0B0F"
INK1 = "#2a2a33"
INK2 = "#8a8a90"
BG   = "#fafafa"
PANEL = "#ffffff"
GOLD   = "#E3B657"
AMBER  = "#E59A3E"
TEAL   = "#3BA99C"
SKY    = "#4A90D9"
VIOLET = "#8A6FC4"
CRIMSON = "#C85450"
LINE   = "#d4d4d7"
INK_HEADING = "#0B0B0F"

plt.rcParams.update({
    "font.family": "sans-serif",
    "font.sans-serif": ["Calibri", "Arial", "DejaVu Sans"],
    "axes.edgecolor": INK2,
    "axes.labelcolor": INK1,
    "xtick.color": INK1,
    "ytick.color": INK1,
    "text.color": INK1,
    "figure.facecolor": BG,
    "axes.facecolor": BG,
    "savefig.facecolor": BG,
})


def save(fig, name):
    path = ART / name
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)
    print(f"  ...wrote {path.name}")
    return path


# ==========================================================================
# FIGURES
# ==========================================================================

def fig_1_1_close_timeline():
    """Fig 1.1 — close-day timeline today vs. APEX."""
    fig, ax = plt.subplots(figsize=(10, 4))
    ax.set_xlim(0, 7)
    ax.set_ylim(0, 3)
    ax.axis("off")

    # Today
    ax.text(0.05, 2.55, "TODAY — 6-day close", fontsize=11, fontweight="bold", color=INK0)
    segs_today = [
        (0, 1,   AMBER,  "Balance fetch\n& trial bal"),
        (1, 2,   AMBER,  "Intercompany\nmatch (manual)"),
        (2, 3,   AMBER,  "FX reval\n+ accruals"),
        (3, 4,   AMBER,  "JE composition\n& posting"),
        (4, 5,   TEAL,   "Consolidation"),
        (5, 6,   CRIMSON,"Flux / variance\nanalysis"),
    ]
    for s, e, c, lbl in segs_today:
        ax.add_patch(FancyBboxPatch((s+0.05, 1.9), e-s-0.1, 0.5,
                                    boxstyle="round,pad=0.02,rounding_size=0.05",
                                    facecolor=c, edgecolor="none", alpha=0.85))
        ax.text((s+e)/2, 2.15, lbl, ha="center", va="center", fontsize=8, color="white")
    ax.text(6.12, 2.15, "(begins day 5)", fontsize=8, color=INK2, va="center")

    # Target
    ax.text(0.05, 1.35, "TARGET — 3-day close with APEX", fontsize=11, fontweight="bold", color=INK0)
    segs_apex = [
        (0,   1,   SKY,     "Balance-Fetch\n(agent, parallel)"),
        (1,   1.6, VIOLET,  "Interco /\nFX / Accrual\n(parallel)"),
        (1.6, 2.2, GOLD,    "JE-Compose\n+ HITL"),
        (2.2, 3,   CRIMSON, "Flux / variance\n(humans, more time)"),
    ]
    for s, e, c, lbl in segs_apex:
        ax.add_patch(FancyBboxPatch((s+0.05, 0.7), e-s-0.1, 0.5,
                                    boxstyle="round,pad=0.02,rounding_size=0.05",
                                    facecolor=c, edgecolor="none", alpha=0.90))
        ax.text((s+e)/2, 0.95, lbl, ha="center", va="center", fontsize=7.5, color="white")

    # day tick axis
    for d in range(0, 7):
        ax.plot([d, d], [0.2, 0.4], color=INK2, lw=0.6)
        ax.text(d, 0.05, f"Day {d}", ha="center", fontsize=8, color=INK2)

    # Saved band
    ax.add_patch(mpatches.FancyBboxPatch((2.2, 0.55), 3.8, 0.65,
                                         boxstyle="round,pad=0.02",
                                         facecolor=GOLD, edgecolor="none", alpha=0.10))
    ax.text(4.1, 0.45, "≈ 3 days reclaimed for analysis", ha="center", fontsize=9, color=GOLD,
            fontweight="bold")
    return save(fig, "fig_1_1_close_timeline.png")


def fig_2_1_value_chain():
    """Fig 2.1 — APEX value-delivery chain."""
    fig, ax = plt.subplots(figsize=(10, 2.4))
    ax.axis("off")
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 2.5)
    links = [
        ("Scenario",  AMBER),
        ("Solution",  TEAL),
        ("Use Case",  SKY),
        ("Service",   VIOLET),
        ("Persona",   CRIMSON),
        ("KPI",       GOLD),
    ]
    for i, (name, c) in enumerate(links):
        x = 0.6 + i * 1.9
        ax.add_patch(FancyBboxPatch((x, 0.8), 1.5, 0.9,
                                    boxstyle="round,pad=0.02,rounding_size=0.08",
                                    facecolor=c, edgecolor="none", alpha=0.92))
        ax.text(x + 0.75, 1.25, name, ha="center", va="center",
                fontsize=12, color="white", fontweight="bold")
        if i < 5:
            ax.annotate("", xy=(x + 1.85, 1.25), xytext=(x + 1.55, 1.25),
                        arrowprops=dict(arrowstyle="->", color=INK2, lw=1.5))
    ax.text(6, 0.35, "Sections 3–8 walk each link in sequence for the R2R scenario.",
            ha="center", fontsize=9, color=INK2, style="italic")
    return save(fig, "fig_2_1_value_chain.png")


def fig_3_1_current_vs_target():
    """Fig 3.1 — Mechanical vs. judgment time allocation."""
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(10, 4))
    # Today
    vals_today = [65, 25, 10]
    labels = ["Mechanical\n(fetch, match,\nreval, post)", "Consolidation\n(mechanical)", "Flux/variance\n(judgment)"]
    colors = [AMBER, SKY, CRIMSON]
    ax1.pie(vals_today, labels=labels, colors=colors, autopct="%d%%",
            startangle=90, wedgeprops=dict(edgecolor=BG, linewidth=2),
            textprops=dict(fontsize=9, color=INK1))
    ax1.set_title("Today — controller time allocation",
                  fontsize=11, color=INK0, pad=10, fontweight="bold")

    # Target
    vals_target = [25, 15, 60]
    ax2.pie(vals_target, labels=labels, colors=colors, autopct="%d%%",
            startangle=90, wedgeprops=dict(edgecolor=BG, linewidth=2),
            textprops=dict(fontsize=9, color=INK1))
    ax2.set_title("Target — controller time allocation (W3)",
                  fontsize=11, color=INK0, pad=10, fontweight="bold")
    fig.suptitle("Judgment time grows from 10% to 60%; mechanical time compresses.",
                 fontsize=10, color=INK2, y=-0.02, style="italic")
    return save(fig, "fig_3_1_time_allocation.png")


def fig_4_1_orchestration():
    """Fig 4.1 — 6-agent orchestration diagram."""
    fig, ax = plt.subplots(figsize=(10.5, 5.5))
    ax.axis("off")
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 6)

    def box(x, y, w, h, label, color, sub=None):
        ax.add_patch(FancyBboxPatch((x, y), w, h,
                                    boxstyle="round,pad=0.02,rounding_size=0.08",
                                    facecolor=color, edgecolor="none", alpha=0.92))
        ax.text(x+w/2, y+h/2+0.1, label, ha="center", va="center",
                fontsize=10, color="white", fontweight="bold")
        if sub:
            ax.text(x+w/2, y+h/2-0.25, sub, ha="center", va="center",
                    fontsize=7.5, color="white", alpha=0.85)

    # Parent orchestrator
    box(3.5, 5, 4, 0.7, "R2R Close Orchestrator", AMBER,
        "hierarchical + parallel + HITL-gated-commit")

    # Agents layer
    box(0.4, 3.2, 2.0, 0.9, "Balance-Fetch", SKY, "SAP S/4HANA → FINML")
    box(2.8, 3.2, 2.0, 0.9, "Interco-Match", TEAL, "AR ↔ AP reconcile")
    box(5.2, 3.2, 2.0, 0.9, "FX-Reval", TEAL, "period-end rates")
    box(7.6, 3.2, 2.0, 0.9, "Accrual-Propose", TEAL, "drivers + history")

    # JE layer
    box(3.5, 1.6, 4, 0.9, "JE-Compose + HITL", VIOLET,
        "Teams Adaptive Card for material entries")

    # Evidence
    box(0.4, 0.2, 2.0, 0.9, "Evidence-Write", CRIMSON, "14-field audit row")
    box(8.6, 0.2, 2.0, 0.9, "LEDGER / SOXML", GOLD, "SOX evidence packet")

    # arrows
    def arr(x1, y1, x2, y2, color=INK2, style="->"):
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle=style, color=color, lw=1.2))
    # parent -> agents
    for cx in [1.4, 3.8, 6.2, 8.6]:
        arr(5.5, 5, cx, 4.1, color=AMBER)
    # agents -> JE
    for cx in [1.4, 3.8, 6.2, 8.6]:
        arr(cx, 3.2, 5.5, 2.5, color=INK2)
    # JE -> Evidence + LEDGER
    arr(4.3, 1.6, 1.4, 1.1, color=VIOLET)
    arr(6.7, 1.6, 9.6, 1.1, color=VIOLET)
    # Evidence -> LEDGER
    arr(2.4, 0.65, 8.6, 0.65, color=CRIMSON)

    ax.text(5.5, 5.95, "Figure 4.1 · R2R orchestration — parent over six bounded agents; HITL only where materiality requires it",
            ha="center", fontsize=9, color=INK2, style="italic")
    return save(fig, "fig_4_1_orchestration.png")


def fig_5_1_use_cases():
    """Fig 5.1 — Six use cases with HITL gates."""
    fig, ax = plt.subplots(figsize=(10, 4.5))
    ax.axis("off")
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 5)
    cases = [
        ("5.1  Balance Fetch & Normalize",   "No HITL",            SKY),
        ("5.2  Intercompany Reconciliation", "HITL if true-up>$100K", TEAL),
        ("5.3  FX Revaluation",              "Rare HITL",          TEAL),
        ("5.4  Accrual Proposal",            "HITL if conf<0.9",   AMBER),
        ("5.5  JE Compose & Post",           "HITL for material",  VIOLET),
        ("5.6  Evidence & Close-Out",        "No HITL",            CRIMSON),
    ]
    for i, (title, gate, color) in enumerate(cases):
        col = i % 3
        row = i // 3
        x = 0.4 + col * 3.2
        y = 3.2 - row * 2.2
        ax.add_patch(FancyBboxPatch((x, y), 2.8, 1.6,
                                    boxstyle="round,pad=0.02,rounding_size=0.1",
                                    facecolor=color, edgecolor="none", alpha=0.92))
        ax.text(x+1.4, y+1.15, title, ha="center", va="center",
                fontsize=10, color="white", fontweight="bold")
        ax.text(x+1.4, y+0.55, gate, ha="center", va="center",
                fontsize=8.5, color="white", alpha=0.90, style="italic")
    return save(fig, "fig_5_1_use_cases.png")


def fig_7_1_persona_constellation():
    """Fig 7.1 — Persona constellation (Deirdre primary + 4 secondary)."""
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.axis("off")
    ax.set_xlim(-5, 5)
    ax.set_ylim(-3.5, 3.5)

    # Primary (center)
    cx, cy = 0, 0
    ax.add_patch(plt.Circle((cx, cy), 1.1, color=AMBER, alpha=0.93, zorder=3))
    ax.text(cx, cy+0.15, "Deirdre Chen", ha="center", fontsize=11,
            color="white", fontweight="bold", zorder=4)
    ax.text(cx, cy-0.25, "Corporate Controller", ha="center", fontsize=8.5,
            color="white", alpha=0.92, zorder=4)
    ax.text(cx, cy-0.60, "(primary · HITL authority)", ha="center", fontsize=7.5,
            color="white", alpha=0.80, zorder=4)

    # Secondaries orbit
    secondaries = [
        ("Marcus Okafor",       "SOX Program Lead",       "evidence consumer",  SKY,     (-3.2, 1.6)),
        ("Priya Nair",          "Close Manager",          "process owner",      TEAL,    ( 3.2, 1.6)),
        ("Jean-Luc Martin",     "Brand Finance Director", "flux reviewer",      VIOLET,  (-3.2,-1.6)),
        ("Sandra Kelleher",     "External Audit PIC",     "evidence recipient", CRIMSON, ( 3.2,-1.6)),
        ("Hassan Idris",        "IT Finance Systems",     "connector owner",    GOLD,    ( 0.0, 2.6)),
    ]
    for name, title, func, color, (x, y) in secondaries:
        ax.add_patch(plt.Circle((x, y), 0.85, color=color, alpha=0.88, zorder=3))
        ax.text(x, y+0.15, name, ha="center", fontsize=9,
                color="white", fontweight="bold", zorder=4)
        ax.text(x, y-0.12, title, ha="center", fontsize=7.5,
                color="white", alpha=0.92, zorder=4)
        ax.text(x, y-0.38, f"({func})", ha="center", fontsize=7,
                color="white", alpha=0.80, zorder=4)
        # line from primary to secondary
        ax.plot([cx, x], [cy, y], color=LINE, lw=1, zorder=1)
    return save(fig, "fig_7_1_personas.png")


def fig_8_1_close_day_trend():
    """Fig 8.1 — Close days over time across waves."""
    weeks = np.arange(1, 25)
    # baseline 6 days, W1 Nike pilot (weeks 1-8) steady at 6 then drifts to 5
    # W2 pilot (weeks 9-16) drops from 5 to 3.5 (entity-scope still USA DTC)
    # W3 scale (weeks 17-24) drops to 3.0 and holds
    days = np.array([6]*4 + [5.8, 5.5, 5.3, 5.0]
                    + [4.8, 4.5, 4.2, 4.0, 3.8, 3.6, 3.5, 3.5]
                    + [3.4, 3.3, 3.2, 3.1, 3.05, 3.0, 3.0, 3.0])

    fig, ax = plt.subplots(figsize=(10, 4.2))
    ax.plot(weeks, days, color=AMBER, lw=2.2, marker="o", markersize=4, markerfacecolor=AMBER,
            markeredgecolor="white")
    ax.axhline(3.0, color=TEAL, lw=1, linestyle="--", alpha=0.7)
    ax.text(24, 2.85, "W3 target: 3 days", color=TEAL, fontsize=8.5, ha="right")
    ax.axhline(6.0, color=CRIMSON, lw=1, linestyle="--", alpha=0.7)
    ax.text(1, 6.15, "baseline: 6 days", color=CRIMSON, fontsize=8.5, ha="left")

    # Wave bands
    for (x0, x1, lab, col) in [(1, 8, "W1 Foundation", "#f3e7d1"),
                               (9, 16, "W2 Pilot", "#d8ece6"),
                               (17, 24, "W3 Scale & Fuse", "#d5e5f6")]:
        ax.axvspan(x0-0.4, x1+0.4, alpha=0.30, color=col, zorder=0)
        ax.text((x0+x1)/2, 6.55, lab, ha="center", fontsize=9, color=INK2, fontweight="bold")

    ax.set_xlabel("Close cycle (monthly)", fontsize=9)
    ax.set_ylabel("Close duration (business days)", fontsize=9)
    ax.set_ylim(2, 7)
    ax.set_xlim(0.5, 24.5)
    ax.set_xticks([1, 4, 8, 12, 16, 20, 24])
    ax.grid(True, alpha=0.25)
    for s in ["top", "right"]:
        ax.spines[s].set_visible(False)
    return save(fig, "fig_8_1_close_day_trend.png")


def fig_9_1_wave_progression():
    """Fig 9.1 — Three-wave progression with scope + envelope + value."""
    fig, ax = plt.subplots(figsize=(10.5, 5))
    ax.axis("off")
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 6)

    waves = [
        ("WAVE 1", "Foundation",
         "Nike USA DTC (1 entity)",
         "3 agents · monthly close",
         "$800K – $2.5M", "#f3e7d1"),
        ("WAVE 2", "Pilot",
         "Nike USA DTC (1 entity, 3 cycles + Q3 SOX)",
         "6 agents · full orchestration",
         "$3M – $8M", "#d8ece6"),
        ("WAVE 3", "Scale & Fuse",
         "Nike global (~40 entities, 3 brands)",
         "Multi-currency · multi-GAAP · consolidation",
         "$12M – $32M", "#d5e5f6"),
    ]

    for i, (wave, tag, scope, detail, envelope, col) in enumerate(waves):
        x = 0.5 + i * 3.8
        ax.add_patch(FancyBboxPatch((x, 1.2), 3.5, 3.8,
                                    boxstyle="round,pad=0.05,rounding_size=0.15",
                                    facecolor=col, edgecolor=LINE, linewidth=1))
        ax.text(x+1.75, 4.55, wave, ha="center", fontsize=13,
                color=INK0, fontweight="bold")
        ax.text(x+1.75, 4.15, tag, ha="center", fontsize=10.5, color=INK1, style="italic")
        ax.text(x+1.75, 3.45, scope, ha="center", fontsize=9, color=INK1, fontweight="bold")
        ax.text(x+1.75, 3.00, detail, ha="center", fontsize=8.5, color=INK1)
        ax.text(x+1.75, 2.20, "Commercial envelope", ha="center", fontsize=8, color=INK2)
        ax.text(x+1.75, 1.80, envelope, ha="center", fontsize=11,
                color=GOLD, fontweight="bold")

    # arrows between waves
    for i in range(2):
        x1 = 0.5 + i * 3.8 + 3.5
        x2 = 0.5 + (i+1) * 3.8
        ax.annotate("", xy=(x2, 3.1), xytext=(x1, 3.1),
                    arrowprops=dict(arrowstyle="->", color=AMBER, lw=2))
    # durable value annotation
    ax.text(6, 0.5, "W3 run-rate durable value: ~$4.95M/yr reclaimed capacity  +  $3–4.5M/yr audit-fee avoidance",
            ha="center", fontsize=9.5, color=GOLD, fontweight="bold")
    return save(fig, "fig_9_1_waves.png")


def fig_10_1_commercial_envelope():
    """Fig 10.1 — Commercial envelope log-scale bars."""
    fig, ax = plt.subplots(figsize=(10, 4.5))
    tiers = ["Approach\n(discovery + SOW)", "Wave 1\nFoundation", "Wave 2\nPilot", "Wave 3\nScale & Fuse"]
    lows  = [0.15, 0.8, 3.0, 12.0]
    highs = [0.4, 2.5, 8.0, 32.0]
    colors = [INK2, SKY, TEAL, GOLD]
    x = np.arange(len(tiers))
    for i, (lo, hi, c) in enumerate(zip(lows, highs, colors)):
        ax.barh(i, hi-lo, left=lo, color=c, alpha=0.88, edgecolor="white", linewidth=1.5, height=0.55)
        ax.text(hi+0.5, i, f"${lo:.1f}M – ${hi:.1f}M", va="center", fontsize=9, color=INK1)
    ax.set_yticks(x)
    ax.set_yticklabels(tiers, fontsize=9.5)
    ax.set_xscale("log")
    ax.set_xlim(0.1, 60)
    ax.set_xlabel("Commercial envelope (USD, log scale)", fontsize=9)
    ax.axvline(8, color=GOLD, lw=1, linestyle="--", alpha=0.7)
    ax.text(8.2, -0.55, "$8M/yr W3 durable value reference", color=GOLD, fontsize=8.5, ha="left")
    for s in ["top", "right"]:
        ax.spines[s].set_visible(False)
    ax.grid(True, axis="x", which="both", alpha=0.2)
    ax.invert_yaxis()
    return save(fig, "fig_10_1_commercial_envelope.png")


def fig_roi_waterfall():
    """Fig 8.2 — ROI stack (waterfall-style)."""
    fig, ax = plt.subplots(figsize=(10, 4.5))
    labels = ["Controller\nhours", "Audit fee\navoidance", "Shortened\nclose\n(IR value)",
              "Restatement\nrisk\nreduction", "W3 run-rate\ndurable value"]
    values = [4.95, 3.75, 0.8, 0.5, 0]
    colors = [TEAL, SKY, VIOLET, CRIMSON, GOLD]

    totals = [0]
    for v in values[:-1]:
        totals.append(totals[-1] + v)
    total = sum(values[:-1])

    for i, (v, base, c, lbl) in enumerate(zip(values, totals, colors, labels)):
        if i < len(values) - 1:
            ax.bar(i, v, bottom=base, color=c, alpha=0.90, edgecolor="white", width=0.6)
            ax.text(i, base + v + 0.15, f"+${v:.2f}M", ha="center", fontsize=9, color=INK1)
        else:
            ax.bar(i, total, color=GOLD, alpha=0.95, edgecolor="white", width=0.6)
            ax.text(i, total + 0.15, f"${total:.2f}M", ha="center",
                    fontsize=10, color=GOLD, fontweight="bold")

    ax.set_xticks(range(len(labels)))
    ax.set_xticklabels(labels, fontsize=8.5)
    ax.set_ylabel("Annualized value ($M, W3 run-rate)", fontsize=9)
    ax.set_ylim(0, max(total, sum(values[:-1])) + 1.5)
    for s in ["top", "right"]:
        ax.spines[s].set_visible(False)
    ax.grid(True, axis="y", alpha=0.2)
    ax.set_title("Figure 8.2 · Annualized durable value stack, Nike W3 run-rate",
                 fontsize=10, color=INK2, pad=12, loc="left")
    return save(fig, "fig_8_2_roi_stack.png")


def fig_audit_row_schema():
    """Fig 8.3 — Audit row 14-field schema diagram."""
    fig, ax = plt.subplots(figsize=(10, 4.5))
    ax.axis("off")
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 4)

    fields = [
        ("1", "trace_id"),
        ("2", "scenario_id"),
        ("3", "agent_id"),
        ("4", "actor\n(human|agent)"),
        ("5", "action"),
        ("6", "timestamp"),
        ("7", "entity_id"),
        ("8", "amount+ccy"),
        ("9", "confidence"),
        ("10", "policy_applied"),
        ("11", "evidence_refs"),
        ("12", "override_reason"),
        ("13", "manifest_ver"),
        ("14", "prior_decision"),
    ]
    palette = [SKY, SKY, SKY, TEAL, TEAL, TEAL, VIOLET, VIOLET,
               AMBER, AMBER, CRIMSON, CRIMSON, GOLD, GOLD]
    for i, ((num, name), c) in enumerate(zip(fields, palette)):
        col = i % 7
        row = i // 7
        x = col * 1.9 + 0.2
        y = 2.6 - row * 1.5
        ax.add_patch(FancyBboxPatch((x, y), 1.6, 1.0,
                                    boxstyle="round,pad=0.02,rounding_size=0.08",
                                    facecolor=c, edgecolor="none", alpha=0.90))
        ax.text(x+0.8, y+0.7, num, ha="center", fontsize=10, color="white", fontweight="bold")
        ax.text(x+0.8, y+0.30, name, ha="center", fontsize=8, color="white")
    ax.text(7, 0.1, "Every R2R decision emits one 14-field audit row to LEDGER — scenario_id = RC-OPS-06",
            ha="center", fontsize=9, color=INK2, style="italic")
    return save(fig, "fig_8_3_audit_row.png")


# ==========================================================================
# DOCX BUILD
# ==========================================================================

def set_cell_bg(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex.lstrip("#"))
    tcPr.append(shd)


def set_run(run, *, bold=False, italic=False, size=None, color_hex=None, font=None):
    if bold: run.bold = True
    if italic: run.italic = True
    if size is not None:
        run.font.size = Pt(size)
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex.lstrip("#"))
    if font:
        run.font.name = font


def add_heading(doc, text, *, level=1, color="#0B0B0F"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.font.name = "Calibri"
    sizes = {1: 22, 2: 16, 3: 13, 4: 11}
    set_run(r, bold=True, size=sizes.get(level, 12), color_hex=color)
    return p


def add_eyebrow(doc, text, *, color="#8a8a90"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text.upper())
    set_run(r, bold=True, size=9, color_hex=color, font="Calibri")
    r.font.all_caps = True
    # letterspacing via XML (approx)
    rPr = r._element.get_or_add_rPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:val"), "80")  # 1/20 pt
    rPr.append(spacing)
    return p


def add_body(doc, text, *, size=10.5, color="#1f1f25", italic=False, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run(r, size=size, color_hex=color, italic=italic, font="Calibri")
    return p


def add_bullets(doc, items, *, size=10.5):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(it)
        set_run(r, size=size, color_hex="#1f1f25", font="Calibri")


def add_figure(doc, path, caption, *, width_in=6.3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    p.add_run().add_picture(str(path), width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    r = cap.add_run(caption)
    set_run(r, italic=True, size=9, color_hex="#8a8a90", font="Calibri")


def add_callout(doc, kind, text, *, tag_color="#E3B657"):
    """Single-cell shaded table used as a callout block."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, "#f6f1e4" if tag_color == "#E3B657" else "#eef4fb")
    # remove default para
    cell.paragraphs[0].text = ""
    # tag line
    p_tag = cell.paragraphs[0]
    r_tag = p_tag.add_run(kind.upper())
    set_run(r_tag, bold=True, size=9, color_hex=tag_color, font="Calibri")
    # body
    p_body = cell.add_paragraph()
    p_body.paragraph_format.space_before = Pt(2)
    r_body = p_body.add_run(text)
    set_run(r_body, size=10, color_hex="#1f1f25", font="Calibri")
    # spacing after
    doc.add_paragraph()


def add_kv_table(doc, pairs, *, header_fill="#0B0B0F", header_color="#ffffff",
                 left_w=2.2, right_w=4.0):
    tbl = doc.add_table(rows=len(pairs), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(pairs):
        cell_k, cell_v = tbl.rows[i].cells
        cell_k.width = Inches(left_w)
        cell_v.width = Inches(right_w)
        if i == 0:
            set_cell_bg(cell_k, header_fill)
            set_cell_bg(cell_v, header_fill)
            cell_k.paragraphs[0].text = ""
            cell_v.paragraphs[0].text = ""
            rk = cell_k.paragraphs[0].add_run(k)
            set_run(rk, bold=True, size=10, color_hex=header_color, font="Calibri")
            rv = cell_v.paragraphs[0].add_run(v)
            set_run(rv, bold=True, size=10, color_hex=header_color, font="Calibri")
        else:
            set_cell_bg(cell_k, "#f3f3f5")
            cell_k.paragraphs[0].text = ""
            cell_v.paragraphs[0].text = ""
            rk = cell_k.paragraphs[0].add_run(k)
            set_run(rk, bold=True, size=9.5, color_hex="#1f1f25", font="Calibri")
            rv = cell_v.paragraphs[0].add_run(v)
            set_run(rv, size=9.5, color_hex="#1f1f25", font="Calibri")
    doc.add_paragraph()


def add_data_table(doc, header, rows, *, col_widths=None,
                   header_fill="#0B0B0F", header_color="#ffffff"):
    n_cols = len(header)
    tbl = doc.add_table(rows=len(rows) + 1, cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header
    for j, h in enumerate(header):
        c = tbl.rows[0].cells[j]
        set_cell_bg(c, header_fill)
        c.paragraphs[0].text = ""
        r = c.paragraphs[0].add_run(h)
        set_run(r, bold=True, size=9.5, color_hex=header_color, font="Calibri")
        if col_widths:
            c.width = Inches(col_widths[j])
    # body
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            c = tbl.rows[i].cells[j]
            if i % 2 == 0:
                set_cell_bg(c, "#f7f7f9")
            c.paragraphs[0].text = ""
            r = c.paragraphs[0].add_run(str(v))
            set_run(r, size=9.5, color_hex="#1f1f25", font="Calibri")
            if col_widths:
                c.width = Inches(col_widths[j])
    doc.add_paragraph()


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


# --- build the document --------------------------------------------------

def build():
    # 1. Generate figures first
    print("Generating figures...")
    fig_1_1 = fig_1_1_close_timeline()
    fig_2_1 = fig_2_1_value_chain()
    fig_3_1 = fig_3_1_current_vs_target()
    fig_4_1 = fig_4_1_orchestration()
    fig_5_1 = fig_5_1_use_cases()
    fig_7_1 = fig_7_1_persona_constellation()
    fig_8_1 = fig_8_1_close_day_trend()
    fig_8_2 = fig_roi_waterfall()
    fig_8_3 = fig_audit_row_schema()
    fig_9_1 = fig_9_1_wave_progression()
    fig_10_1 = fig_10_1_commercial_envelope()

    # 2. Compose docx
    print("Composing docx...")
    doc = Document()
    # Set US Letter page size with 1" margins
    for section in doc.sections:
        section.page_width  = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin    = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin   = Inches(0.95)
        section.right_margin  = Inches(0.95)

    # ---- COVER -----
    for _ in range(3):
        doc.add_paragraph()
    add_eyebrow(doc, "APEX · REFERENCE ENGAGEMENT · PRACTICE RC")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Record-to-Report")
    set_run(r, bold=True, size=40, color_hex="#0B0B0F", font="Calibri")
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run("Close automation, from sub-ledger to SOX evidence")
    set_run(r2, size=16, color_hex="#4A4A55", font="Calibri")
    doc.add_paragraph()
    doc.add_paragraph()
    add_body(doc,
             "A walkthrough of a single APEX scenario across the value-delivery chain — "
             "Scenario → Solution → Use Case → Service → Persona → KPI — covering how the close "
             "is solutioned on Microsoft-native deployment, which use cases the solution activates, "
             "which anchor services it productizes, the full persona cast that makes it operational, "
             "and the KPIs the scenario is accountable to. The walkthrough extends through the three "
             "waves (W1 Foundation, W2 Pilot, W3 Scale & Fuse) and closes with a commercial-envelope "
             "summary. Reference client for the Wave 1 build: Nike, Inc.",
             size=11, italic=True)
    doc.add_paragraph()
    # meta table
    add_kv_table(doc, [
        ("Document type",   "Executive Walkthrough · Scenario-to-KPI"),
        ("Practice",        "RC — Retail & Consumer"),
        ("Scenario ID",     "RC-OPS-06-record-to-report-automation"),
        ("Reference client","Nike, Inc. (Wave 1 build target, FY26)"),
        ("Source folder",   "docs/scenarios/RC/operations-workforce/…/"),
        ("Audience",        "CFO · Corporate Controller · SOX Program · CIO / CDO · Account Team"),
        ("Chain links",     "Scenario · Solution · Use Case · Service · Persona · KPI"),
        ("Waves covered",   "Approach · W1 Foundation · W2 Pilot · W3 Scale & Fuse"),
        ("Companion HTML",  "APEX-Stacked-Architecture-Narrated.html"),
        ("Version / date",  "v1.0 · 2026-04-24"),
    ])

    add_page_break(doc)

    # ---- CONTENTS ----
    add_heading(doc, "Contents", level=1)
    add_body(doc,
             "This document walks the R2R scenario through every link of the APEX value-delivery "
             "chain. Sections 1–2 set context. Sections 3–8 are the chain itself: one section per "
             "chain link. Section 9 projects the chain across the three delivery waves; Section 10 "
             "summarizes the commercial envelope. Appendices carry the glossary and the artifact "
             "cross-reference.",
             italic=True, color="#4A4A55")
    add_data_table(doc,
                   ["#", "Section", "What it covers"],
                   [
                       ["1",  "The Reference Close — what actually happens on Day 1 of quarter-end",
                        "The 4 business days a global consumer brand's controllers lose to mechanical close work, "
                        "and the 3-day target after APEX."],
                       ["2",  "Context — why R2R is the RC anchor scenario for finance automation",
                        "Why the five forces converge on record-to-report; why retail/consumer R2R (not automotive, "
                        "not healthcare) is the Wave 1 build."],
                       ["3",  "Chain Link 1 · Scenario",
                        "The close calendar today, the friction quantified, the scope boundaries."],
                       ["4",  "Chain Link 2 · Solution",
                        "The six-agent orchestration, the commercial envelope tiers, the three design commitments."],
                       ["5",  "Chain Link 3 · Use Cases",
                        "Six use cases from Balance Fetch to Evidence & Close-Out; individually deliverable."],
                       ["6",  "Chain Link 4 · Services",
                        "RC-E2E-06 Finance Close Automation · RC-E2E-11 Controls & Audit Evidence."],
                       ["7",  "Chain Link 5 · Personas",
                        "Deirdre Chen (Corporate Controller) as primary; five secondary personas with distinct "
                        "planes and surfaces."],
                       ["8",  "Chain Link 6 · KPIs",
                        "Close-day reduction, controller hours reclaimed, SOX audit-fee avoidance, audit-row "
                        "coverage — every KPI bound to trace-id."],
                       ["9",  "The three waves — W1 Foundation · W2 Pilot · W3 Scale & Fuse",
                        "Nike-specific wave plan: one entity, six agents, forty entities."],
                       ["10", "Commercial Envelope",
                        "Approach / W1 / W2 / W3 cost envelopes and durable-value reference line."],
                       ["A",  "Glossary",
                        "APEX-specific terms used in this document."],
                       ["B",  "Artifact Cross-Reference",
                        "How this document fits with the README, deep dive, narrated HTML, and manifests."],
                   ],
                   col_widths=[0.5, 2.8, 3.2])

    add_page_break(doc)

    # ==== SECTION 1 · OPENING FRAME ====
    add_eyebrow(doc, "SECTION 1 · OPENING FRAME")
    add_heading(doc, "The Reference Close", level=1)
    add_body(doc,
             "At 5:47 AM on the first business day of January, Deirdre Chen — Corporate Controller "
             "for a global consumer brand — starts her laptop to begin the Q2 FY26 close. The close "
             "calendar published to her organization calls for books to be ready for flux review on "
             "Day 5 and consolidated reporting on Day 6. Deirdre knows from experience that what "
             "happens between now and Day 5 is almost entirely mechanical: 40 legal entities' "
             "sub-ledger balances have to be pulled, reconciled, revalued, accrued, and posted before "
             "anyone does the actual accounting judgment that constitutes 'close the books.'")
    add_body(doc,
             "This is the scenario. Everything downstream — the solution architecture, the use cases, "
             "the productized services, the persona cast, the KPI rollups — exists to answer three "
             "questions about the 40 entities and the six business days between Day 1 and Day 6:")
    add_bullets(doc, [
        "Which of the ~15,000 journal entries proposed during close can be auto-posted, and which must be reviewed by a human?",
        "Who has the authority to make that call, and how do they make it in minutes rather than the four days it takes today?",
        "What evidence survives the decision — auditable, SOX-grade, attributable to the specific entry at the specific entity at the specific close period?",
    ])
    add_figure(doc, fig_1_1,
               "Figure 1.1 · The six close days, narrated. Today (top row): days 1–5 are mechanical; "
               "flux/variance analysis happens only on day 5–6 in compressed form. Target (bottom row): "
               "the mechanical phase compresses to day 1–2, and analysis expands into the freed window. "
               "Color coded — sky blue for data agents, teal for reconciliation agents, gold for HITL, "
               "crimson for human judgment.")

    add_callout(doc, "Today, without APEX",
                "Deirdre's 150-person controller organization spends the first four business days of every "
                "close pulling balances, matching intercompany, running FX reval, and proposing "
                "accruals — work that is rule-governed and mechanical but produces no analytical output. "
                "The SOX evidence packet at quarter-end is reconstructed after the fact from email chains, "
                "Excel working papers, and shared drives. External audit spends four weeks collating it.")
    add_callout(doc, "Tomorrow, with APEX",
                "Six agents — Balance-Fetch, Interco-Match, FX-Reval, Accrual-Propose, JE-Compose, "
                "Evidence-Write — run in parallel where the work allows, hierarchically where it requires. "
                "Material entries route through Deirdre via Teams Adaptive Card. Every decision emits a "
                "14-field audit row. The SOX evidence packet is generated, not collated, at quarter-end — "
                "and external audit consumes it directly.",
                tag_color="#3BA99C")

    add_callout(doc, "Independence note",
                "Before selling deeper at any engagement, confirm the audit-firm relationship. If the R2R "
                "work is being proposed at an audit client, the engagement must be structured so that the "
                "agent manifests, policy manifests, and LEDGER are client artifacts — Deloitte's role is "
                "advisory and implementation enablement. The SOW must be routed through the Independence "
                "office before contracting. The underlying APEX pattern is independence-safe by design, "
                "but the specific engagement structure is not assumed — it is designed.")

    # ==== SECTION 2 · CONTEXT ====
    add_eyebrow(doc, "SECTION 2 · CONTEXT")
    add_heading(doc, "Why this scenario exists as an APEX reference", level=1)
    add_heading(doc, "2.1  The five forces converge on the close calendar", level=3)
    add_body(doc,
             "APEX exists because five forces converged through FY25-FY26: the decision-velocity gap "
             "that leaves material-financial judgments dependent on spreadsheet assembly; the agentic-AI "
             "inflection that crossed enterprise viability; the SOX-evidence cost escalation at public "
             "companies; the finance-talent scarcity that flattens controller headcount against growing "
             "transaction volume; and the public-company investor pressure for faster guidance cycles. "
             "Record-to-report is where all five forces intersect — every one of them lands on the "
             "close calendar.")

    add_figure(doc, fig_2_1,
               "Figure 2.1 · The APEX value-delivery chain. Every engagement walks these six links in "
               "sequence. Sections 3 through 8 of this document take each link in turn for the R2R scenario.")

    add_heading(doc, "2.2  Why RC carries the anchor R2R scenario", level=3)
    add_body(doc,
             "Of the seven APEX Practices — RC (Retail & Consumer), HLS (Healthcare & Life Sciences), "
             "ER (Energy & Resources), AXLE (Industrial & Manufacturing, automotive-weighted), TMT "
             "(Technology, Media & Telecom), TH (Travel & Hospitality), and ICE (Industrial & "
             "Commercial Equipment) — RC carries the reference R2R scenario because the retail/consumer "
             "close has all the hardness of public-company finance plus distinctive revenue-recognition "
             "corner cases that stress the agent architecture: DTC returns reserves, wholesale markdown "
             "allowances, licensing royalty true-ups, loyalty liability deferrals. If the orchestration "
             "pattern handles retail/consumer R2R, it ports to AXLE plant-cost accounting, HLS 340B "
             "accruals, ER hedge accounting, and TMT deferred revenue with only the manifest layer "
             "changing.")

    add_heading(doc, "2.3  Why this document exists", level=3)
    add_body(doc,
             "Three conversations need to happen with the same material but different emphasis: the CFO "
             "and Corporate Controller need to see a credible close-cycle-compression story with ROI that "
             "survives a rigorous scrub; the SOX Program Lead and external-audit partner need to see an "
             "evidence story that is auditor-grade by construction; the CIO and CDO need to see a "
             "Microsoft-native architecture that fits the existing Fabric/Foundry/Azure OpenAI/Entra "
             "estate without inventing net-new platform categories. This document is the executive "
             "walkthrough that anchors all three conversations.")

    add_page_break(doc)

    # ==== CHAIN LINK 1 · SCENARIO ====
    add_eyebrow(doc, "CHAIN LINK 1 · SCENARIO")
    add_heading(doc, "Scenario", level=1)
    add_body(doc,
             "The Scenario link of the APEX value-delivery chain is the customer's operational reality "
             "stated without abstraction. It is not yet a solution. It is the thing that has to be true "
             "before a solution can be designed against it.")

    add_heading(doc, "3.1  The situation", level=3)
    add_body(doc,
             "A global consumer brand — Nike, for the Wave 1 reference build — operates ~40 legal "
             "entities across three brands (Nike Brand, Converse, Jordan Brand) and four regional "
             "operating segments (North America; Europe, Middle East & Africa; Greater China; "
             "Asia Pacific & Latin America). Each monthly close produces roughly 12,000–18,000 journal "
             "entries; each quarterly close adds consolidation-layer entries plus SOX-evidence assembly. "
             "Reporting is to US GAAP in USD, but statutory books are kept in local GAAP across nine "
             "functional currencies.")
    add_body(doc,
             "The close begins on the first business day of the period and is scheduled to complete on "
             "day 6. In practice, it slips to day 7 or 8 in 30% of cycles, driven by intercompany "
             "breaks, FX-reval exceptions, and accrual disputes that require controller judgment to "
             "resolve. The slippage compounds into quarterly 10-Q preparation windows.")

    add_heading(doc, "3.2  The current-state response", level=3)
    add_body(doc,
             "The response today, across all 40 entities, is a sequential hand-off chain made of "
             "trial-balance exports, Excel working papers, email approvals, SharePoint folder drops, "
             "and ERP batch jobs. The typical sequence over six business days:")
    add_bullets(doc, [
        "Day 1 — Controllers log into SAP S/4HANA per entity; export trial balances; perform initial reconciliation against sub-ledger detail. Breaks are emailed to upstream sub-ledger owners with a 24-hour response SLA that slips routinely.",
        "Day 2 — Intercompany matching begins. Each entity-pair (roughly 40 × 40 = up to 1,600 pairs, though materially only ~200 pairs are active) is matched manually in a shared workbook. True-ups over $100K are escalated to the paired entity controllers for bilateral review.",
        "Day 2–3 — FX revaluation runs per entity. Rates are pulled from the treasury system; monetary positions are revalued at period-end; unrealized gain/loss entries are proposed. Exceptions are handled entity-by-entity.",
        "Day 3 — Accrual proposals are drafted entity-by-entity by assistant controllers using prior-period patterns, driver data (labor hours, sales volume, licensee self-reports), and management judgment. Each accrual is reviewed by a controller before posting.",
        "Day 3–4 — Journal entries are composed, reviewed, posted. Approval chains run through email. Supporting documents are attached in SharePoint. Materiality thresholds are applied manually.",
        "Day 4–5 — Consolidation runs. Local GAAP → US GAAP adjustments are posted at the consolidation layer. Intercompany eliminations apply. Segment reporting is prepared.",
        "Day 5–6 — The first hour of flux/variance analysis begins. Controllers finally do the work that requires judgment: explaining why revenue, margin, and operating income landed where they did.",
    ])

    add_figure(doc, fig_3_1,
               "Figure 3.1 · Controller time allocation today vs. target. Today, 90% of close hours are "
               "mechanical work (balance fetch, matching, reval, accrual mechanics, journal posting, "
               "consolidation mechanics); only 10% is the flux/variance judgment that actually constitutes "
               "'closing the books.' Target: reverse the ratio so that 60% of close hours are judgment and "
               "mechanical work compresses to 25%.")

    add_heading(doc, "3.3  The friction, quantified", level=3)
    add_body(doc, "Measured over a representative four-quarter period at the reference client, the "
                  "current-state response produces the following friction profile:")
    add_data_table(doc,
                   ["Dimension", "Typical value", "Implication"],
                   [
                       ["Close-cycle duration",         "6 days scheduled · 6.8 actual",       "10-Q prep window compressed; investor-guidance lead-time reduced"],
                       ["Controller hours per close",   "~48 hrs / person / close",            "~150 controllers × 48 × 12 = ~86k hrs/yr on mechanical work"],
                       ["Manual journal entries",       "~9,500 per monthly close (~65%)",     "Each manual entry carries fat-finger risk and audit-evidence overhead"],
                       ["Intercompany true-ups >$100K", "~30 bilateral escalations per close", "Two controllers on each escalation, typically >2 hours per"],
                       ["SOX evidence collation",       "4 weeks per quarter",                 "External-audit fee line driven by evidence hours; quarterly cost"],
                       ["Restatement risk",             "Reliance on manual controls",         "Material misstatements surface during audit, not during close"],
                       ["Close slippage rate",          "30% of cycles slip ≥1 day",           "Downstream commitments (board report, IR guidance) slip in parallel"],
                       ["Auditor testing cycles",       "~6 weeks per quarter",                "External-audit fee premium; re-performance of evidence tie-outs"],
                   ])
    add_body(doc,
             "The friction composes to materially worse outcomes than the underlying economics warrant. "
             "The controllers and SOX team are doing the right things in the right order. The problem is "
             "that the work itself is mechanical in ways that belong to an orchestrated agent fleet, not "
             "to 150 people with spreadsheets.")

    add_heading(doc, "3.4  The scope of the scenario", level=3)
    add_body(doc,
             "For the purposes of this document, the scenario includes every journal entry required for "
             "monthly and quarterly close across operating entities in the reference client's "
             "consolidated group. It specifically includes:")
    add_bullets(doc, [
        "Sub-ledger balance fetch from SAP S/4HANA and peripheral systems (POS/e-comm, licensing, treasury).",
        "Intercompany AR ↔ AP reconciliation, with proposed true-up entries.",
        "FX revaluation of monetary positions at period-end spot or contracted rates.",
        "Accrual proposals for payroll, bonus, rebate, return-reserve, loyalty liability, licensing royalty, and customary period-end items.",
        "Journal entry composition, approval routing (HITL), and posting.",
        "Consolidation-layer entries (local GAAP → US GAAP adjustments; segment reporting).",
        "SOX evidence capture and quarterly evidence-packet assembly.",
    ])
    add_body(doc, "It specifically excludes:")
    add_bullets(doc, [
        "Tax-provision computation (a distinct scenario under RC-OPS or corporate-tax practice).",
        "Treasury hedging and derivative accounting (a distinct ER-adjacent scenario).",
        "Revenue contract assessment for ASC 606 corner cases (a distinct scenario — the R2R scenario implements contract conclusions, it does not produce them).",
        "External 10-Q/10-K drafting (R2R produces the numbers the drafting package consumes).",
    ])

    add_page_break(doc)

    # ==== CHAIN LINK 2 · SOLUTION ====
    add_eyebrow(doc, "CHAIN LINK 2 · SOLUTION")
    add_heading(doc, "Solution", level=1)
    add_body(doc,
             "The Solution link is how APEX answers the Scenario. It is an architectural description — "
             "what runs, where it runs, what contracts bind the components, how humans stay in charge of "
             "the decisions that matter.")

    add_heading(doc, "4.1  The solution in one sentence", level=3)
    add_callout(doc, "Solution",
                "A six-agent R2R orchestration — Balance-Fetch, Interco-Match, FX-Reval, "
                "Accrual-Propose, JE-Compose, Evidence-Write — composes with a HITL-gated commercial "
                "envelope, runs on Microsoft-native infrastructure (Fabric + Foundry + Azure OpenAI + "
                "Teams + Entra), and writes every decision to a 14-field audit row so SOX evidence is "
                "generated, not collated.",
                tag_color="#3BA99C")

    add_heading(doc, "4.2  The architectural components", level=3)
    add_body(doc,
             "The solution is a choreographed composition of components drawn from the ten-layer APEX "
             "stack. Each component has a typed contract with its neighbors, a classification "
             "propagation rule, trace-id discipline, and a manifest declaration.")

    add_figure(doc, fig_4_1,
               "Figure 4.1 · The R2R orchestration graph. Parent orchestrator fans out to four "
               "parallelizable agents (Balance-Fetch is the sequential prerequisite; the next three "
               "run in parallel), then gathers proposals into JE-Compose — which routes material entries "
               "through HITL before posting. Evidence-Write observes every decision and emits to LEDGER.")

    add_heading(doc, "4.3  Why this composition and not another", level=3)

    add_heading(doc, "4.3.1  Manifest-driven, not model-driven", level=4)
    add_body(doc,
             "Every choice that could change between close cycles — which LLM is used for accrual "
             "narrative generation, which tools each agent can call, which policy corpus applies to "
             "materiality thresholds, which controller has approval authority for which entity tier — "
             "is declared in a Git-versioned manifest. Updating any of these is a manifest PR with "
             "review and merge discipline, not an ops-ticket to the AI team. Auditors see the manifest "
             "history alongside the audit row: they know what the system was configured to do at the "
             "time of any specific decision.")

    add_heading(doc, "4.3.2  Audit row, not log line", level=4)
    add_body(doc,
             "Every decision — every proposed entry, every HITL approval, every posted journal — emits "
             "a 14-field structured row carrying trace-id, three version stamps (manifest, policy, "
             "prompt), content-addressed input and output hashes, reasoning-trace reference, HITL "
             "status and approver, downstream-effect cross-reference, and KPI attribution. This is the "
             "artifact auditors test against. It is not a log line; it is the substrate of SOX "
             "operation-and-exception evidence.")

    add_heading(doc, "4.3.3  HITL is first-class, not bolted on", level=4)
    add_body(doc,
             "The HITL gate between JE-Compose and post is not a retro-fitted 'human check.' It is "
             "declared in the orchestration manifest as a first-class primitive. Its gate kind "
             "(ZERO_TOUCH, ACK_ONLY, HITL, or ESCALATION) is deterministically mapped from the "
             "commercial envelope tier — entity materiality × entry amount × agent confidence. "
             "Deirdre Chen, the Corporate Controller, sees only what she must see. The low-materiality, "
             "high-confidence entries are auto-posted with a log entry she can sample at will. The "
             "material entries arrive as Adaptive Cards in Teams with one-click approve / reject / "
             "modify. She never logs into another tool.")

    add_heading(doc, "4.4  What the solution does not do", level=3)
    add_body(doc, "APEX is not a replacement. It does not:")
    add_bullets(doc, [
        "Replace SAP S/4HANA. S/4HANA remains the System of Record for the sub-ledger and the general ledger; APEX agents read via standard connectors and write through controlled posting APIs.",
        "Replace the treasury FX system. Period-end and average rates come from the authoritative source; APEX applies them.",
        "Replace the Corporate Controller's authority. Deirdre stays on the approval path for every material entry; the manifest declares what 'material' means per entity tier.",
        "Replace external audit. The audit happens; what changes is the substrate external audit tests against — one canonical SOX evidence packet instead of reconstructed Excel and email trails.",
        "Replace ASC 606 revenue contract conclusions. R2R implements the conclusions; the contract-assessment scenario produces them.",
    ])

    add_page_break(doc)

    # ==== CHAIN LINK 3 · USE CASES ====
    add_eyebrow(doc, "CHAIN LINK 3 · USE CASES")
    add_heading(doc, "Use Cases", level=1)
    add_body(doc,
             "The Use Case link is the set of discrete operational patterns the Solution activates. One "
             "Solution typically enables multiple Use Cases — each Use Case is a bounded, testable, "
             "individually-measurable pattern that can be delivered as a piece of a larger engagement "
             "or as a standalone increment.")
    add_body(doc,
             "The R2R solution activates six primary use cases. Each is described below with its "
             "trigger, the agents involved, the HITL gate behavior, and the typical commercial envelope "
             "for delivering that use case in isolation.")

    add_figure(doc, fig_5_1,
               "Figure 5.1 · The six R2R use cases. Each is individually deliverable, individually "
               "measurable, and individually commercially-envelopable. A client can start with Balance "
               "Fetch + Interco Match as a 'Close Day 1–2 Accelerator' package and expand from there.")

    # Per-use-case detail table
    for title, payload in [
        ("5.1  Use Case · Balance Fetch & Normalize",
         [("Trigger", "Period-end cutoff reached for an entity; orchestrator fires."),
          ("Agents",  "Balance-Fetch"),
          ("Tools",   "SAP S/4HANA BAPI/OData connector; FINML schema validator"),
          ("HITL",    "None — deterministic pull and normalization"),
          ("Envelope","Included in W1 foundation")]),
        ("5.2  Use Case · Intercompany Reconciliation",
         [("Trigger", "Both sides of an entity pair have balances available."),
          ("Agents",  "Interco-Match"),
          ("Tools",   "FINML.IntercoMatch writer; HITL dispatcher"),
          ("HITL",    "If proposed true-up > $100K, both entity controllers approve"),
          ("Envelope","~15% of W2 scope")]),
        ("5.3  Use Case · FX Revaluation",
         [("Trigger", "Monetary positions tagged; period-end rates loaded."),
          ("Agents",  "FX-Reval"),
          ("Tools",   "Treasury rate source connector; FINML.FXRate validator"),
          ("HITL",    "Rare (rate source authoritative); exceptions routed to treasury"),
          ("Envelope","~10% of W2 scope")]),
        ("5.4  Use Case · Accrual Proposal",
         [("Trigger", "Balance fetch complete; driver data available."),
          ("Agents",  "Accrual-Propose"),
          ("Tools",   "Historical-pattern retrieval; driver-data joins; LLM narrative"),
          ("HITL",    "If confidence < 0.9 OR amount > threshold, controller approves"),
          ("Envelope","~25% of W2 scope (highest judgment surface area)")]),
        ("5.5  Use Case · JE Compose & Post",
         [("Trigger", "Any proposal from Interco / FX / Accrual agents."),
          ("Agents",  "JE-Compose"),
          ("Tools",   "Approvals MCP; Teams Adaptive Card; S/4HANA posting API"),
          ("HITL",    "Always HITL for material entries; materiality declared in policy manifest"),
          ("Envelope","~30% of W2 scope")]),
        ("5.6  Use Case · Evidence & Close-Out",
         [("Trigger", "Every action from every agent, continuous."),
          ("Agents",  "Evidence-Write"),
          ("Tools",   "SOXML.AuditEvidence writer; LEDGER plane write"),
          ("HITL",    "None — write-only; consumed by SOX Program Lead at quarter-end"),
          ("Envelope","~20% of W2 scope; foundation for all SOX-evidence-packet work")]),
    ]:
        add_heading(doc, title, level=3)
        add_kv_table(doc, payload, left_w=1.2, right_w=5.0)

    add_heading(doc, "5.7  How the use cases compose", level=3)
    add_body(doc,
             "Individually, each use case is a deliverable. Together they compose the full R2R scenario. "
             "The composition is deliberate — a client can buy the first two use cases (Balance Fetch + "
             "Interco Match) as a 'Close Day 1–2 Accelerator' and capture roughly two days of cycle time "
             "with a smaller commercial envelope. Adding Accrual + JE-Compose converts that into a "
             "'Close Day 1–3 Accelerator.' Adding Evidence & Close-Out converts it into the full "
             "SOX-grade evidence-packet solution. This incrementalism is how the Approach tier and "
             "Wave 1 Foundation are sized — the client is not forced to commit to the full W3 scope at "
             "W1 contracting.")

    add_page_break(doc)

    # ==== CHAIN LINK 4 · SERVICES ====
    add_eyebrow(doc, "CHAIN LINK 4 · SERVICES")
    add_heading(doc, "Services", level=1)
    add_body(doc,
             "The Service link is what gets productized. A Service is a named, subscribable APEX "
             "capability with a defined persona, SLO, commercial terms, and an anchor in the APEX "
             "Service Catalog (Design Reference Appendix B).")

    add_heading(doc, "6.1  Anchor services", level=3)
    add_body(doc, "The R2R scenario anchors to two RC services in the APEX Service Catalog:")

    add_heading(doc, "RC-E2E-06 · Finance Close Automation", level=4)
    add_body(doc,
             "The foundational close-automation service. Owns the six-agent orchestration from "
             "Balance-Fetch through JE-Compose. Persona: Corporate Controller. SLO: close-cycle "
             "duration measured per period, target Day 3 at W3 steady state. Commercial envelope: "
             "tiered by entity count, transaction volume, and SOX-scope depth.")
    add_heading(doc, "RC-E2E-11 · Controls & Audit Evidence", level=4)
    add_body(doc,
             "The SOX-evidence service that wraps Evidence-Write, LEDGER, and the quarterly evidence "
             "packet. Persona: SOX Program Lead + External Audit Partner-in-Charge. SLO: evidence "
             "packet delivered within 72 hours of quarter-end; audit-row coverage ≥ 99.5% on material "
             "entries. Commercial envelope: tiered by SOX scope (number of entities in scope, number "
             "of key controls, auditor sampling methodology).")

    add_heading(doc, "6.2  Service composition", level=3)
    add_body(doc,
             "Each anchor service is a composition of: canonical schema subscriptions (FINML, SOXML), "
             "agent deployments, MCP tool registrations, HITL policy bindings, LEDGER policy-corpus "
             "entries, Power BI semantic-model publications, and Teams Adaptive Card templates.")
    add_data_table(doc,
                   ["Commitment", "RC-E2E-06 Finance Close Automation", "RC-E2E-11 Controls & Audit Evidence"],
                   [
                       ["Canonical schemas",  "FINML.SubLedger, .IntercoMatch, .FXRate, .JournalProposal", "SOXML.AuditEvidence"],
                       ["Agents deployed",    "Balance-Fetch, Interco-Match, FX-Reval, Accrual-Propose, JE-Compose", "Evidence-Write"],
                       ["MCP tool surfaces",  "S/4HANA posting, Treasury rate, Approvals",                  "LEDGER query, SOX-packet export"],
                       ["HITL policies",      "materiality × entity tier × confidence",                      "exception-sampling protocol"],
                       ["Power BI models",    "Close-cycle dashboard, JE-volume dashboard",                  "SOX-evidence completeness dashboard"],
                       ["Teams surfaces",     "Controller approval card",                                    "SOX-exception notification card"],
                       ["SLO",                "Close-cycle duration, auto-post rate",                        "Packet assembly time, audit-row coverage"],
                       ["Primary persona",    "Corporate Controller (Deirdre Chen)",                         "SOX Program Lead (Marcus Okafor)"],
                   ],
                   col_widths=[1.6, 2.5, 2.5])

    add_heading(doc, "6.3  Service subscription in the tenant manifest", level=3)
    add_body(doc,
             "A client subscribes to services by declaring them in the tenant manifest. For the Nike "
             "reference client, the subscription block looks like:")
    code = doc.add_paragraph()
    code.paragraph_format.left_indent = Inches(0.3)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(8)
    r = code.add_run(
        '"subscriptions": [\n'
        '  { "service_id": "RC-E2E-06", "service_version": "v1.0.0" },\n'
        '  { "service_id": "RC-E2E-11", "service_version": "v1.0.0" }\n'
        ']'
    )
    set_run(r, size=9, color_hex="#0B0B0F", font="Consolas")
    add_body(doc,
             "Version pinning is deliberate. When APEX releases v1.1.0 of RC-E2E-06, the client decides "
             "when to upgrade. PATCH bumps are automatic; MINOR and MAJOR bumps require an explicit "
             "accept-upgrade HITL gate. This is how tenants stay on stable versions while the Service "
             "Catalog continues to evolve.")

    add_heading(doc, "6.4  Candidate derived services", level=3)
    add_body(doc,
             "The R2R scenario anchors two services today. It also informs the potential productization "
             "of additional services that share the same architectural substrate:")
    add_heading(doc, "Candidate: RC-E2E-14 · Flux Commentary Assistant", level=4)
    add_body(doc,
             "Takes the closed trial balance and proposes board-ready flux/variance commentary for "
             "controller review. Uses the same LEDGER retrieval pattern + LLM narrative pattern as "
             "Accrual-Propose; different persona (Brand Finance Director) and different commercial "
             "envelope.")
    add_heading(doc, "Candidate: RC-E2E-15 · Tax Provision Orchestrator", level=4)
    add_body(doc,
             "Extends the R2R orchestration with tax-provision agents (book-to-tax differences, "
             "effective-rate reconciliation, deferred-tax roll-forward). Persona: Tax Director. "
             "Architectural reuse is high; the manifest layer carries most of the variance.")
    add_heading(doc, "Candidate: RC-E2E-16 · Consolidation Automation", level=4)
    add_body(doc,
             "Extends the R2R orchestration at the consolidation layer — intercompany elimination, "
             "minority interest, segment reporting, currency translation adjustments. Persona: "
             "Consolidation Controller. Architectural reuse is extremely high; this is primarily "
             "scale of the existing orchestration one layer up.")

    add_page_break(doc)

    # ==== CHAIN LINK 5 · PERSONAS ====
    add_eyebrow(doc, "CHAIN LINK 5 · PERSONAS")
    add_heading(doc, "Personas", level=1)
    add_body(doc,
             "The Persona link is the human-role cast that makes the scenario operational. Persona "
             "depth is what separates an architecturally correct solution from one that lands on the "
             "floor. APEX models one primary persona — the named person who uses the solution for the "
             "decisions that matter — and a full cast of secondary personas who execute, escalate, "
             "approve beyond threshold, consume evidence, or own the connector stack.")

    add_figure(doc, fig_7_1,
               "Figure 7.1 · The R2R persona constellation. Deirdre Chen (Corporate Controller) holds "
               "the primary role at the center; five secondary personas orbit with distinct functions — "
               "SOX program evidence consumer, close-manager process owner, brand finance flux reviewer, "
               "external audit evidence recipient, IT-finance connector owner.")

    add_heading(doc, "7.1  Primary persona", level=3)
    add_callout(doc, "PRIMARY · Deirdre Chen · Corporate Controller",
                "Deirdre owns the close calendar, the materiality policy, and the HITL approval path "
                "for every material entry. She approves via Teams Adaptive Card on her phone or laptop "
                "with one click; her approvals, modifications, and rejections write to the audit row "
                "with her named identity resolved through Entra. She reports to the CFO and interfaces "
                "directly with the audit-committee chair on any material-control exception. Her "
                "measure of success: a close-cycle that completes in 3 business days with ≥99.5% of "
                "entries carrying full audit-row evidence.",
                tag_color="#E59A3E")
    add_body(doc,
             "Deirdre is named, not generic. 'Corporate Controller' is a title; Deirdre Chen is a "
             "person. The hitl-authority policy registered in LEDGER binds to her named identity "
             "(resolved through Entra group-lookup at runtime). The Adaptive Card fires to her "
             "device. Her approval is what unblocks JE-Compose to post.")

    add_heading(doc, "7.2  Secondary personas", level=3)
    add_body(doc,
             "Five secondary personas each carry a distinct role. The roles are deliberate; APEX does "
             "not invent roles, it maps to roles that already exist in the client's controller handbook.")
    add_data_table(doc,
                   ["Role", "Function in the R2R scenario"],
                   [
                       ["Marcus Okafor — SOX Program Lead",
                        "Evidence consumer — reviews the quarterly SOX evidence packet, runs exception "
                        "sampling, interfaces with external audit on evidence questions."],
                       ["Priya Nair — Close Manager",
                        "Process owner — runs the close calendar day-to-day, monitors agent SLOs, "
                        "triages any agent failure, coordinates with IT-Finance on connector issues."],
                       ["Jean-Luc Martin — Brand Finance Director (Converse)",
                        "Flux reviewer — consumes the closed trial balance, produces the flux/variance "
                        "commentary for his brand, exception-escalates any unexpected movements."],
                       ["Sandra Kelleher — External Audit PIC",
                        "Evidence recipient — receives the SOX evidence packet, runs substantive "
                        "testing against the audit-row sample, interfaces with Marcus on exceptions."],
                       ["Hassan Idris — IT Finance Systems Lead",
                        "Connector owner — owns the S/4HANA connector, the treasury rate feed, the "
                        "Teams Adaptive Card integration, and the Entra policy wiring."],
                   ],
                   col_widths=[2.4, 3.9])

    add_heading(doc, "7.3  Why persona depth matters", level=3)
    add_body(doc,
             "A seller who names only the primary persona says 'we built something for the controller.' "
             "A seller who names the full cast shows the buyer a complete operating model: who owns the "
             "evidence path, who runs the close calendar, who reviews flux, who the auditor-of-record "
             "is, who owns the connectors. That completeness is what turns a demo into an engagement.")
    add_body(doc,
             "Persona depth also matters for change management. When the engagement rolls out, the "
             "client's Finance Enablement team has to train all six named roles — not just Deirdre. "
             "The six-role inventory informs training plan scope, org-readiness assessment, and the "
             "operational handoff checklist in the W2 → W3 gate review.")

    add_heading(doc, "7.4  Persona-to-plane mapping", level=3)
    add_body(doc,
             "Each persona interacts with a specific plane of the APEX architecture. Understanding the "
             "mapping is how we know what we need to build on each plane for each role.")
    add_data_table(doc,
                   ["Persona", "Plane(s) touched", "Surface"],
                   [
                       ["Deirdre (Primary)",  "Experience · Decision", "Teams Adaptive Card (mobile + desktop)"],
                       ["Marcus",             "Experience · Ledger",   "Power BI SOX-evidence dashboard · quarterly packet export"],
                       ["Priya",              "Experience · Runtime",  "Close-calendar dashboard · agent-SLO monitor"],
                       ["Jean-Luc",           "Experience · Ledger",   "Flux-commentary Power BI page · drill-through to trace-id"],
                       ["Sandra",             "Ledger",                "Evidence-packet ZIP (delivered out-of-band to audit firm)"],
                       ["Hassan",             "Integration · Runtime", "Connector-ops console · manifest Git repo · Entra admin"],
                   ],
                   col_widths=[2.0, 2.0, 2.3])

    add_page_break(doc)

    # ==== CHAIN LINK 6 · KPIs ====
    add_eyebrow(doc, "CHAIN LINK 6 · KPIs")
    add_heading(doc, "KPIs", level=1)
    add_body(doc,
             "The KPI link is where the chain closes. APEX's signature is that every KPI attributes "
             "back to the specific decision audit row that produced it — not just to the aggregate. "
             "The distinction matters because it is what makes the KPI durable under audit.")

    add_heading(doc, "8.1  Per-cycle KPIs", level=3)
    add_body(doc,
             "A closed cycle emits a closeout audit row carrying the following measured KPI fields. "
             "Each is bound to the specific trace-id of that cycle. A Power BI dashboard rolls them up "
             "across cycles; drill-through on any aggregated cell returns the individual audit rows.")
    add_data_table(doc,
                   ["KPI", "Measurement", "Target (W2 Pilot)", "Target (W3 Scale)"],
                   [
                       ["Close-cycle duration",            "Hours from cutoff to sign-off",     "4 days",       "3 days"],
                       ["Auto-post rate",                  "% of entries posted without HITL",  "≥ 70%",        "≥ 85%"],
                       ["Controller hours reclaimed",      "Measured vs. pre-APEX baseline",    "+12 hrs/person","+22 hrs/person"],
                       ["Audit-row coverage",              "% of material entries attributed",  "≥ 98%",        "≥ 99.5%"],
                       ["Intercompany break rate",         "Breaks escalated vs. entries",      "≤ 4%",         "≤ 2%"],
                       ["SOX packet assembly time",        "Hours from cutoff to packet",       "72 hrs",       "24 hrs"],
                       ["Close-slippage rate",             "% cycles exceeding schedule",       "≤ 20%",        "≤ 5%"],
                       ["Restatement-risk index",          "Composite (internal methodology)",  "Green",        "Green"],
                   ])

    add_figure(doc, fig_8_1,
               "Figure 8.1 · Close-cycle duration across 24 monthly cycles. W1 weeks establish baseline "
               "and agent smoke-tests; W2 pilot (Nike USA DTC) begins the compression trajectory; W3 "
               "scale locks in the 3-day run-rate across all ~40 entities. The gap between the red "
               "6-day baseline line and the teal 3-day target line is ~86,000 reclaimed controller "
               "hours per year at steady state.")

    add_heading(doc, "8.2  Derived KPIs for commercial cases", level=3)
    add_body(doc,
             "Beyond the per-cycle operational metrics, four derived KPIs anchor commercial "
             "conversations. They are computed from the per-cycle metrics aggregated over time. Each "
             "attributes back to specific cycles.")

    add_heading(doc, "8.2.1  Controller Capacity Returned", level=4)
    add_body(doc,
             "The operational metric that matters to Corporate Controllers and CFOs. Computed as "
             "SUM(controller_hours_reclaimed) × blended loaded cost. At 150 controllers × 22 hours × "
             "12 closes = ~39,600 hrs/year × $125/hr = ~$4.95M/yr. An auditor retrieves the "
             "contributing audit rows from LEDGER in seconds.")

    add_heading(doc, "8.2.2  Audit Fee Avoidance", level=4)
    add_body(doc,
             "The metric that matters to the Audit Committee. At a $30M baseline external-audit fee "
             "line and a 10–15% fee reduction driven by evidence-hours compression, W3 run-rate = "
             "$3M–$4.5M/yr. Target ranges are confirmed with the audit firm during W2 pilot evidence-"
             "packet walkthroughs.")

    add_heading(doc, "8.2.3  Close-Cycle Compression", level=4)
    add_body(doc,
             "The metric that matters to Investor Relations and the CEO. Three business days "
             "recovered = three additional days of executive and IR attention on forward-looking "
             "commentary instead of backward-looking mechanics. Some of this value monetizes through "
             "faster guidance; some through retained controller talent; some through reduced "
             "period-end overtime.")

    add_heading(doc, "8.2.4  Audit-Trail Completeness", level=4)
    add_body(doc,
             "The regulatory metric. Computed as the percentage of material entries whose audit rows "
             "carry all 14 required fields, every version stamp, and a verifiable signature. Target "
             "is ≥ 99.5% at W3. The metric matters because it is the direct answer to the external "
             "auditor's substantive-testing questions and the audit committee's control-design "
             "questions.")

    add_figure(doc, fig_8_2,
               "Figure 8.2 · Annualized durable value stack at W3 run-rate. Four stacked contributions "
               "compose to the ~$10M/year durable value reference line for a $32B-revenue consumer "
               "brand with ~40 entities and 150-person controller organization. The gold bar on the "
               "right is the summed reference; the stacked bars on the left are the contributing "
               "value streams.")

    add_heading(doc, "8.3  KPI attribution in the audit row", level=3)
    add_body(doc,
             "The closeout audit row — emitted once per close cycle — carries the following structure "
             "relevant to KPI attribution:")
    j = doc.add_paragraph()
    j.paragraph_format.left_indent = Inches(0.3)
    j.paragraph_format.space_after = Pt(8)
    r = j.add_run(
        '{\n'
        '  "trace_id": "trc_2026-07-02_0547_nke_usa_dtc",\n'
        '  "scenario_id": "RC-OPS-06-record-to-report-automation",\n'
        '  "event_kind": "r2r_close_closeout",\n'
        '  "event_ts": "2026-07-04T16:22:18Z",\n'
        '  "tenant_id": "nke",\n'
        '  "manifest_version": "rc-r2r-orchestration-v1.0.0",\n'
        '  "policy_version": "materiality-corpus-v1.0.0",\n'
        '  "kpis": {\n'
        '    "close_cycle_hours": 69.4,\n'
        '    "auto_post_rate_pct": 72.1,\n'
        '    "controller_hrs_reclaimed": 18.4,\n'
        '    "audit_row_coverage_pct": 99.2,\n'
        '    "sox_packet_ready_ts": "2026-07-07T08:11:02Z"\n'
        '  }\n'
        '}'
    )
    set_run(r, size=9, color_hex="#0B0B0F", font="Consolas")

    add_figure(doc, fig_8_3,
               "Figure 8.3 · The 14-field APEX audit row. Every R2R decision emits exactly one row; "
               "every KPI attributes back to specific rows by trace-id. This is the substrate SOX "
               "operation-and-exception evidence is computed from.")

    add_heading(doc, "8.4  What KPI attribution unlocks", level=3)
    add_body(doc, "The specific binding of each KPI value to a specific trace-id makes three business "
                  "outcomes possible that are not available with aggregate-only metrics:")
    add_bullets(doc, [
        "The CFO can ask 'show me the ten material entries that consumed the most controller approval time last quarter' and get rows from the audit table in seconds.",
        "The SOX Program Lead can ask 'which audit rows were less than fully complete last quarter, and why' and get a filterable exception list.",
        "The Practice Lead can ask 'which accrual proposals did the Controller override most often' and get the inputs, outputs, and reasoning-trace reference for every override — feeding manifest evolution.",
    ])

    add_page_break(doc)

    # ==== SECTION 9 · WAVE PROGRESSION ====
    add_eyebrow(doc, "SECTION 9 · WAVE PROGRESSION")
    add_heading(doc, "The Three Waves", level=1)
    add_body(doc,
             "The same chain — Scenario → Solution → Use Case → Service → Persona → KPI — projects "
             "across three delivery Waves. W1 Foundation is prerequisite work that enables the "
             "scenario. W2 Pilot is the scenario running live at proof-point scale. W3 Scale & Fuse "
             "is the scenario at enterprise run-rate, fused with adjacent scenarios.")

    add_figure(doc, fig_9_1,
               "Figure 9.1 · The three-wave progression. Scope expands from 1 entity (Nike USA DTC) in "
               "W1, to 1 entity × 3 cycles + quarterly SOX in W2, to ~40 entities × 3 brands × "
               "multi-currency × multi-GAAP in W3. Commercial envelope grows from $800K–$2.5M in W1 "
               "to $12M–$32M in W3. W3 run-rate durable value reaches ~$10M/year for the reference "
               "$32B-revenue consumer brand.")

    add_heading(doc, "9.1  Wave 1 · Foundation", level=3)
    add_body(doc,
             "W1 is the one-time investment per tenant. It produces the rails that make the scenario "
             "possible and make every subsequent scenario cheap. For the Nike build, W1 scope is "
             "one entity (Nike USA DTC), three agents (Balance-Fetch, Interco-Match, Evidence-Write), "
             "and the full APEX platform substrate.")
    add_data_table(doc,
                   ["W1 step", "What the chain link looks like in W1"],
                   [
                       ["Scenario",  "No live close yet — W1 authors the scenario manifest and establishes materiality policy."],
                       ["Solution",  "Three agents deployed, parent orchestrator wired, audit row end-to-end tested."],
                       ["Use Cases", "Balance Fetch + Interco Match running on historical data as a shadow close."],
                       ["Services",  "Tenant manifest subscribes to RC-E2E-06 v1.0.0 and RC-E2E-11 v1.0.0."],
                       ["Personas",  "Deirdre onboarded; Adaptive Card tested end-to-end on her device; Hassan's connector ops tested."],
                       ["KPIs",      "Baseline close-cycle metrics captured from the prior four closes for W2 comparison."],
                   ],
                   col_widths=[1.6, 4.8])
    add_kv_table(doc, [
        ("W1 duration",    "4–8 months"),
        ("W1 effort",      "~45 story points"),
        ("W1 deliverable", "Smoke-test of end-to-end orchestration against historical close data for Nike USA DTC"),
        ("W1 gate",        "Corporate Controller sign-off to proceed to live W2 pilot"),
    ])

    add_heading(doc, "9.2  Wave 2 · Pilot", level=3)
    add_body(doc,
             "W2 is where the scenario becomes real. Nike USA DTC runs the live scenario end-to-end "
             "for three consecutive monthly closes plus one quarterly close (Q3 FY26). The full "
             "six-agent chain fires for every entity-close. The target is three consecutive cycles "
             "meeting the W2 KPIs with no pipeline failure beyond documented exceptions, plus one "
             "quarterly SOX-evidence packet delivered to Internal Audit and externally walked "
             "through with PwC (Nike's current auditor of record — verify current engagement letter).")
    add_data_table(doc,
                   ["W2 step", "What the chain link looks like in W2"],
                   [
                       ["Scenario",  "Live R2R closes firing against production S/4HANA (read) and test posting API; all six agents live."],
                       ["Solution",  "Full six-agent orchestration with HITL routing tuned to Nike USA DTC materiality tiers."],
                       ["Use Cases", "All six use cases active; auto-post rate tracked; HITL volume measured."],
                       ["Services",  "RC-E2E-06 + RC-E2E-11 tenant subscriptions exercised end-to-end."],
                       ["Personas",  "All six named personas exercised in live close; feedback captured for W2→W3 gate."],
                       ["KPIs",      "Per-cycle KPIs measured; quarterly SOX packet produced; external-audit walkthrough completed."],
                   ],
                   col_widths=[1.6, 4.8])
    add_kv_table(doc, [
        ("W2 duration",    "6–9 months"),
        ("W2 effort",      "~40 story points initial + ongoing enablement"),
        ("W2 deliverable", "3 consecutive live monthly closes at target cycle + 1 quarterly SOX packet"),
        ("W2 gate",        "Audit Committee approval to proceed to W3 enterprise scale"),
    ])

    add_heading(doc, "9.3  Wave 3 · Scale & Fuse", level=3)
    add_body(doc,
             "W3 is where the scenario meets the enterprise and where it stops running alone. Scope "
             "expands from one entity to ~40 entities; brands expand from one to three (Nike Brand, "
             "Converse, Jordan); currencies expand from one to nine+; GAAP perspectives expand from "
             "US GAAP to US GAAP + local GAAP in each jurisdiction. And — most importantly — the "
             "R2R scenario fuses with adjacent scenarios (tax provision, consolidation, flux "
             "commentary) that share the same architectural substrate.")
    add_data_table(doc,
                   ["W3 step", "What the chain link looks like in W3"],
                   [
                       ["Scenario",  "~40 entities × 12 monthly + 4 quarterly closes; ~192 closes/year orchestrated."],
                       ["Solution",  "Multi-entity parallel orchestration; consolidation-layer orchestration (US GAAP + local GAAP)."],
                       ["Use Cases", "All six use cases running across all entities; plus fusion hooks into adjacent scenarios."],
                       ["Services",  "RC-E2E-06 + RC-E2E-11 at enterprise subscription tier; candidate services (E2E-14/-15/-16) activated as needed."],
                       ["Personas",  "Controller org trained across entities; BF directors onboarded; audit-firm walked through at enterprise."],
                       ["KPIs",      "Enterprise-wide rollup: ~$4.95M/yr controller capacity + ~$3–4.5M/yr audit fee + IR/close-speed value."],
                   ],
                   col_widths=[1.6, 4.8])
    add_kv_table(doc, [
        ("W3 duration",    "Multi-year rollout (12–18 months for core; ongoing fusion)"),
        ("W3 effort",      "~60 story points initial + ongoing scenario-fusion and manifest evolution"),
        ("W3 deliverable", "Enterprise close at run-rate; external audit consuming packet directly"),
        ("W3 gate",        "Steady-state operating review; expansion into adjacent scenarios"),
    ])

    add_heading(doc, "9.4  The cross-wave arrows", level=3)
    add_body(doc,
             "Three structural dependencies bind the waves into one progression. They are not "
             "architectural coincidences — they are the load-bearing connections that make the wave "
             "model coherent.")
    add_callout(doc, "W1 → W2",
                "The HITL template wired in W1 (Deirdre's Adaptive Card on her device with Entra "
                "identity resolution) powers the W2 HITL gate across all six agents. What is "
                "prototyped in W1 is consumed unchanged in W2.",
                tag_color="#3BA99C")
    add_callout(doc, "W1 → W2",
                "The LEDGER plane established in W1 accumulates the W2 audit rows. The SOX evidence "
                "packet format is designed in W1, exercised with three audit rows in W1 shadow closes, "
                "and scales to thousands of rows in W2.")
    add_callout(doc, "W2 → W3",
                "The W2 pilot KPIs unlock W3 scaling. The Steering Committee decision to commit to "
                "all-entity rollout is specifically gated on W2 meeting its close-cycle and "
                "audit-row-coverage targets. This is how the commercial envelope protects both "
                "sides — the client does not commit to W3 dollars until W2 delivers W2 outcomes.",
                tag_color="#3BA99C")

    add_page_break(doc)

    # ==== SECTION 10 · COMMERCIAL ENVELOPE ====
    add_eyebrow(doc, "SECTION 10 · COMMERCIAL ENVELOPE")
    add_heading(doc, "Commercial Envelope", level=1)
    add_body(doc,
             "The engagement encompasses four commercial tiers corresponding to the Approach stage "
             "(Discovery and SoW build-out), Wave 1 Foundation, Wave 2 Pilot, and Wave 3 Scale & "
             "Fuse. Each carries a distinct cost-envelope range, duration, and deliverable signature.")
    add_data_table(doc,
                   ["Tier", "Duration", "Cost envelope", "Effort", "Primary deliverable"],
                   [
                       ["Approach",         "2–4 weeks",   "$150K–$400K",   "20 SP",  "Discovery report · entity scoping · SoW"],
                       ["Wave 1 Foundation","4–8 months",  "$800K–$2.5M",   "45 SP",  "Smoke-test against historical closes"],
                       ["Wave 2 Pilot",     "6–9 months",  "$3M–$8M",       "40 SP",  "3 live cycles + Q3 SOX packet"],
                       ["Wave 3 Scale&Fuse","12–18 months","$12M–$32M",     "60+ SP", "Enterprise close at run-rate"],
                   ],
                   col_widths=[1.4, 1.2, 1.3, 0.8, 2.5])

    add_figure(doc, fig_10_1,
               "Figure 10.1 · Commercial envelope visualized on log scale. Each bar shows the low-to-"
               "high range for that tier. The gold dashed line marks the ~$8M/yr W3 durable value "
               "reference — useful when discussing cumulative investment against cumulative return.")

    add_heading(doc, "10.1  What the envelope buys", level=3)
    add_body(doc,
             "The commercial envelope is not a project-labor-hours calculation. It reflects the total "
             "value-delivery bundle: Deloitte-delivered solutioning and Practice expertise; Microsoft "
             "platform licensing and consumption (Fabric F-SKU, Foundry, Azure OpenAI, Teams, Entra, "
             "Power BI); ISV co-sell where applicable; client-side enablement and change "
             "management; and SOX-walkthrough coordination with the external audit firm.")

    add_heading(doc, "10.2  How the envelope relates to durable value", level=3)
    add_body(doc,
             "A ~40-entity W3 enterprise footprint at the reference client produces ~$4.95M/yr in "
             "controller capacity, plus ~$3–4.5M/yr in audit-fee avoidance, plus non-monetized close-"
             "cycle-compression value (faster guidance, retained talent, reduced period-end "
             "overtime). The $12M–$32M W3 envelope spans a rapid-payback range (under 2 years at the "
             "low end) to a strategic-investment range (3–4 years at the high end, reflecting greater "
             "scenario-fusion scope in the upper tier).")

    add_heading(doc, "10.3  Independence posture in commercial conversations", level=3)
    add_callout(doc, "Independence",
                "The seller and delivery team use Deloitte-appropriate language throughout: "
                "'Deloitte's Microsoft practice,' 'DMTSP,' 'Microsoft platform capabilities,' "
                "'Microsoft-native deployment.' The terms 'partner,' 'alliance,' and 'strategic "
                "alliance' do not appear in R2R commercial materials. Before any engagement at an "
                "audit client, the SOW is routed through the Independence office; the engagement is "
                "structured so that manifests, policies, and LEDGER are client artifacts, and "
                "Deloitte's role is advisory and implementation enablement.")

    add_page_break(doc)

    # ==== APPENDIX A · GLOSSARY ====
    add_eyebrow(doc, "APPENDIX A")
    add_heading(doc, "Glossary", level=1)
    add_body(doc, "Selected terms used in this document with APEX-specific meaning.")
    add_data_table(doc,
                   ["Term", "Definition"],
                   [
                       ["APEX",            "Agentic Platform for Enterprise eXecution — the Deloitte-authored reference architecture for agentic scenarios on Microsoft-native deployment."],
                       ["Adaptive Card",   "Microsoft Teams rich-card format used by APEX for HITL approval surfaces."],
                       ["Agent",           "A bounded unit of intent + tool-use that emits typed output and writes an audit row per decision."],
                       ["Audit row",       "14-field structured record emitted by every agent decision; primary substrate of SOX operation evidence."],
                       ["Balance-Fetch",   "R2R agent that pulls sub-ledger balances from SAP S/4HANA and normalizes to FINML."],
                       ["Close cycle",     "One end-to-end monthly or quarterly close; monthly cycles produce trial balances, quarterly cycles add SOX evidence."],
                       ["Commercial envelope", "Policy-declared tiered gating that determines ZERO_TOUCH vs. ACK_ONLY vs. HITL vs. ESCALATION for each decision."],
                       ["Entra",           "Microsoft identity platform; resolves named-persona HITL bindings at runtime."],
                       ["Evidence-Write",  "R2R agent that writes every decision to LEDGER and assembles the quarterly SOX evidence packet."],
                       ["Fabric",          "Microsoft unified analytics platform; hosts Bronze/Silver/Gold data assets used by R2R agents."],
                       ["FINML",           "APEX canonical schema family for finance objects (SubLedger, IntercoMatch, FXRate, JournalProposal)."],
                       ["Foundry",         "Microsoft Azure AI Foundry; hosts model endpoints, tool registrations, and prompt management for APEX agents."],
                       ["FX-Reval",        "R2R agent that revalues monetary positions at period-end rates."],
                       ["HITL",            "Human-in-the-loop gate; first-class orchestration primitive declared in manifest."],
                       ["Interco-Match",   "R2R agent that reconciles intercompany AR vs. AP across entity pairs and proposes true-ups."],
                       ["JE-Compose",      "R2R agent that composes journal entries from upstream proposals and routes material entries through HITL."],
                       ["LEDGER",          "The APEX audit-row persistence + retrieval plane; substrate of SOX evidence and KPI attribution."],
                       ["Manifest",        "Git-versioned declarative artifact describing agents, policies, schemas, and tenant subscriptions."],
                       ["Materiality policy", "Client-specific declared thresholds that gate HITL routing; tuned per entity tier."],
                       ["MCP",             "Model Context Protocol; the tool-registration standard used by APEX for agent-tool contracts."],
                       ["Persona",         "A named human role with declared authority, plane bindings, and surface(s). Deirdre Chen, not 'Controller.'"],
                       ["R2R",             "Record-to-Report — the finance close and reporting process; APEX scenario RC-OPS-06."],
                       ["SOXML",           "APEX canonical schema family for SOX evidence objects (AuditEvidence)."],
                       ["Trace-id",        "Immutable correlation identifier that binds every audit row from one decision event."],
                   ],
                   col_widths=[1.8, 4.5])

    add_page_break(doc)

    # ==== APPENDIX B · ARTIFACT CROSS-REFERENCE ====
    add_eyebrow(doc, "APPENDIX B")
    add_heading(doc, "Artifact Cross-Reference", level=1)
    add_body(doc,
             "This document is one of several APEX artifacts covering Record-to-Report. The cross-"
             "reference below shows which artifact serves which audience and where it lives in the "
             "documentation set.")
    add_data_table(doc,
                   ["Artifact", "Primary audience", "Purpose"],
                   [
                       ["APEX-R2R-Walkthrough.docx (this document)",
                        "Executive · CFO · Corporate Controller · SOX Program Lead · CIO/CDO",
                        "Executive walkthrough of the value-delivery chain."],
                       ["README.md (scenario folder)",
                        "Practice Lead · Account Team · Build Team",
                        "Scenario metadata, chain summary, artifact index."],
                       ["APEX-R2R-deep-dive.md",
                        "Build Team · Solution Architect · SOX Program Lead",
                        "Technical deep dive: agent design, data contracts, controls, ROI math."],
                       ["APEX-R2R-build-guide.md (to be authored)",
                        "Build Team · Lead Engineer",
                        "Step-by-step W1 build for the reference client."],
                       ["APEX-R2R-nike-reference.md (to be authored)",
                        "Account Team · Nike client stakeholders",
                        "Nike-specific entity list, channels, GAAP mix."],
                       ["manifests/*.yaml",
                        "Build Team · Platform Engineering",
                        "Declarative artifacts for agents, policy, orchestration."],
                       ["tests/*.py",
                        "Build Team · QA",
                        "pytest + Testcontainers schema-validation harness."],
                       ["APEX-Stacked-Architecture-Narrated.html",
                        "All audiences",
                        "Live, interactive companion to this document — clickable architecture, flow diagrams, archetype catalog."],
                       ["APEX_Design.md",
                        "Build Team · Practice Lead",
                        "Reference architecture spec for agents, LEDGER, audit row."],
                       ["Professional-APEX-Sellers-Guide.html",
                        "Seller · Account Team",
                        "Runtime plane, audit-row extension, LEDGER feedback loop, Appendix AA."],
                   ],
                   col_widths=[2.5, 2.0, 2.0])

    add_body(doc,
             "Independence posture preserved across every artifact: 'Deloitte's Microsoft practice,' "
             "'DMTSP,' 'Microsoft platform capabilities,' 'Microsoft-native deployment.' The terms "
             "'partner,' 'alliance,' and 'strategic alliance' do not appear in any artifact.",
             italic=True, color="#8a8a90")

    doc.add_paragraph()
    p_end = doc.add_paragraph()
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_end = p_end.add_run("APEX · Deloitte Microsoft Technology Solutions & Practices")
    set_run(r_end, size=9, color_hex="#8a8a90", font="Calibri")
    p_end2 = doc.add_paragraph()
    p_end2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_end2 = p_end2.add_run("End of document")
    set_run(r_end2, size=9, color_hex="#8a8a90", italic=True, font="Calibri")

    # 3. Save
    doc.save(str(DOCX))
    print(f"\nWrote {DOCX}")
    print(f"  Paragraphs/tables composed; figures embedded from {ART}")


if __name__ == "__main__":
    build()
