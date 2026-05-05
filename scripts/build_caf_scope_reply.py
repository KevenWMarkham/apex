"""Build the full-scope reply memo to the Converge team as a Word document.

Output: C:\\Stage\\Clients\\Industries\\Consumer\\Retail\\Demos\\CAF-Scope-Reply-to-Converge-Team.docx

Content mirrors CAF-vs-Agentic-Merch-Scope-Analysis.md but is formatted as
a formal memo / reply suitable for forwarding to Saurabh + the Converge /
AI & Eng leadership circle.
"""

from __future__ import annotations

import io
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

OUT = Path(r"C:\Stage\Clients\Industries\Consumer\Retail\Demos\CAF-Scope-Reply-to-Converge-Team.docx")

# ---- styling helpers ---------------------------------------------------

INK  = "0B0B0F"
INK2 = "4A4A55"
MUTE = "8a8a90"
GOLD = "B8892E"
TEAL = "2A7F73"
CRIM = "A63F3C"
LINE = "D4D4D7"
BAND = "F3F3F5"
HEADBG = "0B0B0F"


def set_cell_bg(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex.lstrip("#"))
    tcPr.append(shd)


def set_cell_border(cell, color_hex=LINE):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), color_hex.lstrip("#"))
        borders.append(b)
    tcPr.append(borders)


def set_run(run, *, bold=False, italic=False, size=None, color=None, font="Arial"):
    if bold: run.bold = True
    if italic: run.italic = True
    if size is not None: run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor.from_string(color.lstrip("#"))
    if font: run.font.name = font


def add_body(doc, text, *, size=11, italic=False, color=INK, align=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None: p.alignment = align
    r = p.add_run(text)
    set_run(r, size=size, color=color, italic=italic)
    return p


def add_heading(doc, text, *, level=1):
    sizes = {1: 17, 2: 14, 3: 12}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    set_run(r, bold=True, size=sizes.get(level, 12), color=INK)
    # simple horizontal rule under H1
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), GOLD.lstrip("#"))
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p


def add_bullets(doc, items, *, size=11):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(it)
        set_run(r, size=size, color=INK)


def add_table(doc, header, rows, *, col_widths=None,
              header_fill=HEADBG, header_color="FFFFFF", zebra=True):
    n = len(header)
    tbl = doc.add_table(rows=len(rows) + 1, cols=n)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(header):
        c = tbl.rows[0].cells[j]
        set_cell_bg(c, header_fill)
        set_cell_border(c)
        c.paragraphs[0].text = ""
        r = c.paragraphs[0].add_run(h)
        set_run(r, bold=True, size=10, color=header_color)
        if col_widths: c.width = Inches(col_widths[j])
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            c = tbl.rows[i].cells[j]
            set_cell_border(c)
            if zebra and i % 2 == 0:
                set_cell_bg(c, BAND)
            c.paragraphs[0].text = ""
            r = c.paragraphs[0].add_run(str(v))
            set_run(r, size=10, color=INK)
            if col_widths: c.width = Inches(col_widths[j])
    doc.add_paragraph()
    return tbl


def add_callout(doc, kind, body, *, accent=GOLD, fill="F6F1E4"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = tbl.rows[0].cells[0]
    set_cell_bg(c, fill)
    c.paragraphs[0].text = ""
    # tag
    p_tag = c.paragraphs[0]
    r_tag = p_tag.add_run(kind.upper())
    set_run(r_tag, bold=True, size=9, color=accent)
    # body
    p = c.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    r = p.add_run(body)
    set_run(r, size=10.5, color=INK)
    doc.add_paragraph()


def add_kv(doc, pairs):
    tbl = doc.add_table(rows=len(pairs), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(pairs):
        ck, cv = tbl.rows[i].cells
        ck.width = Inches(1.6)
        cv.width = Inches(4.9)
        set_cell_bg(ck, BAND)
        set_cell_border(ck); set_cell_border(cv)
        ck.paragraphs[0].text = ""
        cv.paragraphs[0].text = ""
        rk = ck.paragraphs[0].add_run(k)
        set_run(rk, bold=True, size=10, color=INK)
        rv = cv.paragraphs[0].add_run(v)
        set_run(rv, size=10, color=INK)
    doc.add_paragraph()


def add_eyebrow(doc, text, *, color=MUTE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text.upper())
    set_run(r, bold=True, size=9, color=color)
    rPr = r._element.get_or_add_rPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:val"), "80")
    rPr.append(spacing)


def add_pagebreak(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


# ---- build ----------------------------------------------------------------

def build():
    doc = Document()
    # US Letter, 1" margins (per docx skill guidance)
    for section in doc.sections:
        section.page_width  = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = section.bottom_margin = Inches(0.9)
        section.left_margin = section.right_margin = Inches(1.0)
    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # ===== HEADER BLOCK =====
    add_eyebrow(doc, "DELOITTE · MEMO · FULL-SCOPE REPLY")
    p = doc.add_paragraph()
    r = p.add_run("Agentic Merch should inherit CAF — not parallel it")
    set_run(r, bold=True, size=22, color=INK)
    p2 = doc.add_paragraph()
    r2 = p2.add_run("Scope, mapping, risk, and recommended path forward")
    set_run(r2, size=13, color=INK2, italic=True)

    # meta table
    doc.add_paragraph()
    add_kv(doc, [
        ("To",       "Converge / AI & Eng leadership — attn: Saurabh and platform leads"),
        ("From",     "K. Markham · Deloitte Microsoft Technology Solutions & Practices"),
        ("Date",     "24 April 2026"),
        ("Subject",  "CAF vs. Agentic Merch — full-scope scoping reply"),
        ("Cc",       "Merchandising Practice lead · RC Practice lead · DMTSP PMO"),
        ("Status",   "Decision input ahead of MVP architecture lock"),
        ("References", "Agentic Merch Working Design · April 16 blocker set · AI & Eng CONTEXT.md April 17 Decisions Log · Anti-pattern wiki"),
    ])

    # ===== EXEC SUMMARY =====
    add_heading(doc, "Executive summary", level=1)
    add_body(doc,
             "This memo is the full-scope reply to the question the Converge / AI & Eng circle "
             "has been working in the run-up to the Agentic Merch MVP architecture lock: what "
             "does CAF actually give us today, what is the MVP about to rebuild on top of Fabric "
             "and MCP, and what is the cost of getting the substrate choice wrong.")
    add_body(doc,
             "Short answer: CAF — the Converge Agentic Framework — is not a slide deck and not a "
             "future roadmap. It is a production platform today. Its LLM gateway is live and "
             "cloud-agnostic. Its agent-to-agent swarm is the Q4 FY26 flagship. Its dynamic "
             "orchestrator is session-generated, not static. Its guardrails are shipping. Eight "
             "forecasting agents are in production; price/promo is at MVP; restaurants is in "
             "alpha at Taco Bell. Four of the sixteen Agentic Merch agents are already running "
             "as CAF 301-tier agents. Every platform primitive the April 16 blocker set asked "
             "for already exists.")
    add_callout(doc, "One-sentence position",
                "Agentic Merch should be architected as a merch-domain layer on top of CAF, not "
                "as a parallel platform. Inheriting CAF compresses the platform-build by an "
                "estimated 50–65%, honors the April 17 cloud-agnostic commitment in code rather "
                "than on a slide, and concentrates the MVP's effort on the schema and KPI layer "
                "where the merch-domain IP actually lives.",
                accent=GOLD)

    # ===== SECTION 1 =====
    add_heading(doc, "1 · What CAF actually is", level=1)
    add_body(doc,
             "Before the scope conversation goes any further, the team needs a shared view of "
             "what is already built. The table below is capability-by-capability, with current "
             "state of play. None of it is speculative; all of it is in production, in alpha, "
             "or documented as a design pattern already exercised.")
    add_table(doc,
              ["Component", "Already built", "State of play"],
              [
                  ["AEF LLM Gateway",
                   "Yes",
                   "Cloud-agnostic and model-agnostic. Live routes to AWS Bedrock, Azure OpenAI, and Vertex AI. Tenants pick; the gateway abstracts."],
                  ["A2A Swarm",
                   "Yes (Q4 FY26 flagship)",
                   "Cross-app agent collaboration across Converge products. This is the agent-to-agent pattern Agentic Merch needs."],
                  ["Dynamic Orchestrator",
                   "Yes",
                   "LLM-based, session-generated routing. Not static DAG dispatch. Adapts per invocation."],
                  ["Guardrails — GSO + Prompt Eval",
                   "Yes",
                   "PII detection, prompt-injection defense, toxicity filtering, image-compliance checks. Production-grade."],
                  ["301-tier agents in production",
                   "Yes",
                   "Forecasting: 8+ live agents. Price/Promo: MVP shipped. Restaurants: alpha running at Taco Bell."],
                  ["Margin Maker Agents — pattern",
                   "Yes (documented)",
                   "Three-agent pattern: Demand Planner + Pricing Advisor + Sourcing Advisor. Pre-designed E2E flow blueprint."],
                  ["LangFuse + PostgreSQL + RBAC + Registry",
                   "Yes",
                   "Tracing, persistence, access control, agent catalog. The execution substrate."],
              ],
              col_widths=[1.8, 1.5, 3.2])
    add_body(doc,
             "Bottom line: CAF is not something the AI & Eng team needs to build. It is "
             "something that either gets used or reinvented. There is no third option that is "
             "honest about the trade-off.", italic=True, color=INK2)

    # ===== SECTION 2 =====
    add_heading(doc, "2 · How CAF maps directly to the Agentic Merch gaps", level=1)
    add_body(doc,
             "The Agentic Merch workstream identified blockers on April 16. Every one of the "
             "items below has a direct answer inside CAF today. This is the mapping the MVP "
             "architecture lock has to reconcile before any greenfield primitive gets "
             "specified.")
    add_table(doc,
              ["Agentic Merch gap (April 16)", "CAF component that solves it", "Action"],
              [
                  ["Blocker #1: how does Agentic Merch call ConCON? (cross-app agent calls)",
                   "A2A Swarm — the cross-app agent collaboration pattern",
                   "Adopt, don't redesign"],
                  ["Blocker #2: whose LLM handles cross-platform agent calls?",
                   "AEF LLM Gateway — cloud + model agnostic, live",
                   "Route through gateway; drop the 'whose LLM' debate"],
                  ["Orchestration layer ambiguity — Fabric + MCP stack is about to rebuild routing",
                   "Dynamic Orchestrator — LLM-based, session-generated",
                   "Stop rebuilding; inherit"],
                  ["Merch Domain guardrail framework has no technical home",
                   "GSO + Prompt Eval — PII, injection, toxicity, image compliance",
                   "Merch domain policies ride on top; no new framework"],
                  ["Agents 4, 6, 7 of the 16 (Forecasting · Price · Promo)",
                   "Already running as CAF 301-tier agents today",
                   "Map the 16 to CAF Registry; surface the overlap"],
                  ["Agents 3, 4, 6, 7 E2E flow (Demand → Price → Sourcing)",
                   "Margin Maker pattern — pre-designed blueprint, 1:1 map",
                   "Adopt blueprint; do not redraft"],
                  ["Execution substrate the AI & Eng circle is about to scope from zero",
                   "LangFuse + Postgres + RBAC + Registry — operating today",
                   "Inherit. Do not scope greenfield."],
              ],
              col_widths=[2.6, 2.3, 1.6])

    add_callout(doc, "Summary of the overlap",
                "Of the 16 Agentic Merch agents, four are already running in CAF production "
                "(Agents 4, 6, 7 as part of forecasting + price/promo; overlap with Agent 3 "
                "sourcing via Margin Maker). Of the seven framework primitives (LLM routing, "
                "A2A, orchestration, guardrails, tracing, persistence + RBAC, registry), all "
                "seven are already built. What is genuinely new in Agentic Merch is the "
                "merch-domain schema and KPI layer — not the platform.",
                accent=TEAL, fill="E4F1EE")

    add_pagebreak(doc)

    # ===== SECTION 3 =====
    add_heading(doc, "3 · Anti-pattern risk — using our own published schema", level=1)
    add_body(doc,
             "Our own wiki defines the anti-patterns this platform work is supposed to avoid. "
             "The current AI & Eng direction — Fabric + medallion + MCP, greenfield "
             "orchestration — is on track to trip three of them. Each of the three is a direct "
             "quotation from the anti-pattern list and is annotated with the specific way the "
             "current direction violates it.")

    add_heading(doc, "3.1  Anti-pattern · 'Lack of reuse (reinventing primitives)'", level=3)
    add_bullets(doc, [
        "A new LLM-routing layer on MCP reinvents the AEF LLM Gateway.",
        "A new orchestrator on Fabric primitives reinvents the Dynamic Orchestrator.",
        "A new tracing + registry stack reinvents LangFuse + Agent Registry.",
    ])

    add_heading(doc, "3.2  Anti-pattern · 'Non-repeatable demos or fragile prototypes'", level=3)
    add_body(doc,
             "Without the CAF execution substrate, the MVP becomes a prototype that cannot move "
             "to a second client without a reset. CAF was designed for multi-tenant "
             "repeatability from day one; a Fabric-direct MVP is not. The cost of repeatability "
             "shows up at the second engagement, not the first.")

    add_heading(doc, "3.3  Anti-pattern · 'Cosmetic verticalization (no schema/KPI differentiation)'", level=3)
    add_body(doc,
             "If Agentic Merch is positioned primarily as 'the Fabric version' of generic "
             "agentic patterns, the differentiation is cosmetic. The actual differentiation "
             "lives in the merch-domain schema, the merch-specific KPIs, and the "
             "industry-specific agent specializations — and those are exactly the components "
             "that do not require a reinvented platform underneath.")

    # ===== SECTION 4 =====
    add_heading(doc, "4 · The April 17 commitment only holds if we inherit CAF", level=1)
    add_body(doc,
             "The April 17 Decisions Log entry in AI & Eng CONTEXT.md hedged the cloud "
             "commitment cleanly:")
    add_callout(doc, "Quoted · April 17 Decisions Log",
                "Microsoft Fabric is an accelerator, not a mandate; clients may choose AWS/GCP.",
                accent=INK2, fill="EDEDEF")
    add_body(doc,
             "The commitment is directionally right. The problem is that the MVP is still "
             "being architected as if Fabric is the default substrate — and if the MVP gets "
             "built directly on Fabric primitives (Delta Lake, Fabric notebooks, Eventstream, "
             "Fabric-native MCP), three things follow:")
    add_bullets(doc, [
        "The 'clients may choose AWS/GCP' commitment is not actually portable. It requires a second build to honor on a second cloud.",
        "CAF's AEF LLM Gateway already solved hyperscaler-agnostic at the LLM layer. Building past it requires redoing work that is already done.",
        "The claim that Agentic Merch is cloud-neutral becomes a paper claim — true on a slide, false in the code.",
    ])
    add_body(doc,
             "The commitment the team made on April 17 either rides on top of the LLM Gateway "
             "and holds, or gets rebuilt as a second cloud port and doesn't. Those are the two "
             "options. There is not a third.")

    add_pagebreak(doc)

    # ===== SECTION 5 =====
    add_heading(doc, "5 · MVP scope delta — what's in, what's out, why", level=1)
    add_body(doc,
             "The scope question at the architecture lock is cleaner than it looks. Below are "
             "the two lists — what stays in the Agentic Merch MVP because the intellectual "
             "property genuinely lives there, and what comes out because it is already built in "
             "CAF.")

    add_heading(doc, "5.1  In scope for the Agentic Merch MVP", level=2)
    add_table(doc,
              ["Scope item", "Why it stays in"],
              [
                  ["Merch-domain canonical schema",
                   "SKU, price, promo, forecast, inventory shapes — merch-specific canonicalization is IP"],
                  ["Merch-domain KPIs",
                   "Margin, sell-through, markdown effectiveness, basket lift — industry-specific metric layer"],
                  ["Merch-domain agent specializations",
                   "Category planner, allocator, markdown optimizer — merch flavors on top of CAF 301-tier primitives"],
                  ["Merch-domain HITL surfaces",
                   "Buyer / planner / merchant Adaptive Cards with merch-specific decision framing"],
                  ["Merch-domain policy corpus",
                   "Category-specific rules, brand-specific guardrails, pricing-governance policies"],
                  ["Integration shims to merch systems of record",
                   "ERP (SAP / Oracle), MDM, pricing engines, category-management platforms"],
              ],
              col_widths=[2.6, 3.9])

    add_heading(doc, "5.2  Out of scope — because it is already built in CAF", level=2)
    add_table(doc,
              ["Drop from MVP scope", "Inherit from CAF"],
              [
                  ["LLM routing layer",                      "AEF LLM Gateway"],
                  ["Agent-to-agent orchestration",           "A2A Swarm"],
                  ["Session-aware dynamic routing",          "Dynamic Orchestrator"],
                  ["PII / injection / toxicity / image guardrails", "GSO + Prompt Eval"],
                  ["Agent tracing + persistence + catalog",  "LangFuse + PostgreSQL + RBAC + Registry"],
                  ["Forecasting / Price / Promo agent primitives", "Existing CAF 301-tier agents"],
              ],
              col_widths=[3.3, 3.2])

    add_callout(doc, "Estimated scope reduction",
                "If the MVP correctly consumes CAF instead of rebuilding the platform under "
                "Agentic Merch, the platform-layer scope drops by an estimated 50–65%. The "
                "remaining scope concentrates on merch-domain differentiation — which is "
                "exactly where the intellectual property lives and exactly what is not portable "
                "from CAF without this work. The exact number requires a validation working "
                "session with the CAF platform lead; 50–65% is the reasonable a-priori range "
                "based on the capability overlap documented in Sections 1 and 2.",
                accent=GOLD)

    # ===== SECTION 6 =====
    add_heading(doc, "6 · Recommended next steps", level=1)
    add_body(doc,
             "Five steps, in order. The first three happen before the architecture lock; the "
             "last two happen inside the W1 build scope.")
    add_table(doc,
              ["#", "Step", "Owner", "Timing"],
              [
                  ["1", "Book a 60-minute working session: Saurabh + CAF platform lead + Agentic Merch AI lead. Agenda = walk the seven CAF components against the 16 Agentic Merch agents and the Apr 16 blocker set. No decisions yet — map the overlap on a shared whiteboard.",
                   "K. Markham + Saurabh", "This week"],
                  ["2", "Produce a scope-delta document from that session: what is in CAF, what is not, what the merch layer has to add on top. Validate the 50–65% estimate in Section 5.",
                   "Agentic Merch lead", "Within 5 business days of the working session"],
                  ["3", "Re-anchor the MVP architecture to 'CAF is the substrate; Fabric is one of several data-plane options on top.' Update the April 17 Decisions Log entry in AI & Eng CONTEXT.md to reflect this framing.",
                   "AI & Eng lead", "Before architecture lock"],
                  ["4", "Pick the first two 301-tier CAF agents to surface through the Agentic Merch experience as a pilot. Forecasting is the obvious one (8+ already running). Price or Promo is the second.",
                   "Merchandising + CAF leads jointly", "W1 kickoff"],
                  ["5", "Preserve the independence posture across all commercial materials. Nothing in this memo changes commercial language, audit-client structure, or the Microsoft-native default for clients that choose it. It changes only what we build ourselves versus what we reuse from CAF.",
                   "DMTSP PMO", "Ongoing"],
              ],
              col_widths=[0.35, 4.2, 1.4, 1.0])

    # ===== SECTION 7 =====
    add_heading(doc, "7 · Closing position", level=1)
    add_callout(doc, "Executive framing — one sentence",
                "Agentic Merch should be architected as a merch-domain layer on top of CAF, "
                "not as a parallel platform. The LLM gateway, the A2A swarm, the dynamic "
                "orchestrator, the guardrails, the tracing substrate, and four of the sixteen "
                "agents are already running in CAF production today. Re-scoping the MVP to "
                "inherit CAF compresses the platform-build by more than half, honors the "
                "April 17 cloud-agnostic commitment in code (not just in principle), and "
                "focuses Agentic Merch on the schema and KPI differentiation where the real "
                "intellectual property lives.",
                accent=GOLD)

    add_body(doc,
             "Happy to walk any of the sections above in more depth, to stand the working "
             "session up on short notice, or to produce the single-slide companion for "
             "Saurabh's next leadership review. The analysis does not require consensus on the "
             "full memo before the working session happens — it only requires that the seven "
             "CAF capabilities in Section 1 get an honest read against the Apr 16 blocker set "
             "in Section 2. If that mapping stays intact under scrutiny, the scope conclusion "
             "follows.")

    # ===== SIGN-OFF =====
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("— K. Markham")
    set_run(r, bold=True, size=11, color=INK)
    p2 = doc.add_paragraph()
    r2 = p2.add_run("Deloitte Microsoft Technology Solutions & Practices")
    set_run(r2, size=10, color=INK2)
    p3 = doc.add_paragraph()
    r3 = p3.add_run("kmarkham@deloitte.com")
    set_run(r3, size=10, color=INK2, italic=True)

    add_pagebreak(doc)

    # ===== APPENDIX =====
    add_heading(doc, "Appendix · Source references", level=1)
    add_table(doc,
              ["Artifact", "Location"],
              [
                  ["Agentic Merchandising Working Design",
                   "Industries/Consumer/Retail/Demos/Agentic_Merchandising_Working_Design.md"],
                  ["Agentic Merchandising Technical Whitepaper",
                   "Industries/Consumer/Retail/Demos/Agentic_Merchandising_Technical_Whitepaper.docx"],
                  ["Agentic Merch Architecture HTML",
                   "Industries/Consumer/Retail/Demos/agentic_merch_architecture.html"],
                  ["Agent Fleet Architecture",
                   "Industries/Consumer/Retail/Demos/agent-fleet-architecture.html"],
                  ["Narrated Briefing",
                   "Industries/Consumer/Retail/Demos/narrated-briefing.html"],
                  ["AI & Eng Decisions Log (April 17 entry)",
                   "AI & Eng CONTEXT.md"],
                  ["Anti-pattern schema (three items cited in §3)",
                   "Wiki — 'Lack of reuse' · 'Non-repeatable demos' · 'Cosmetic verticalization'"],
                  ["Companion analysis (this memo's technical basis)",
                   "Industries/Consumer/Retail/Demos/CAF-vs-Agentic-Merch-Scope-Analysis.md"],
              ],
              col_widths=[2.8, 3.7])

    add_body(doc,
             "Independence posture preserved across this memo: 'Deloitte's Microsoft practice,' "
             "'DMTSP,' 'Microsoft platform capabilities,' 'Microsoft-native deployment.' The "
             "terms 'partner,' 'alliance,' and 'strategic alliance' do not appear.",
             italic=True, color=MUTE, size=9)

    # save
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
