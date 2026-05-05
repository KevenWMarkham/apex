"""Build two coordinated artifacts for the RC-E2E-03 service:

1. MCP Server Deep Dive (DOCX) — comprehensive technical specification
   of the MCP server that hosts the 19 tool capabilities the 6 agents call.

2. Service Sequence Diagram (HTML) — interactive phase-by-phase sequence
   diagrams showing agent interactions via MCP, aligned with the ERD's
   Bronze tables and the LEDGER closeout row.

Both go into Services/RC-E2E-03_Assortment-and-Pricing/Tier2-Build/
"""
from __future__ import annotations
import io, sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
OUT_DIR = Path(r"C:\Stage\Clients\Industries\Consumer\Retail\Kroger\02_projects\FY27_Pipeline\assortment-pricing-agentic\deliverables\Services\RC-E2E-03_Assortment-and-Pricing\Tier2-Build")
OUT_DIR.mkdir(parents=True, exist_ok=True)

INK="0B0B0F"; INK2="4A4A55"; MUTE="8a8a90"; GOLD="B8892E"; TEAL="2A7F73"
VIO="6648B0"; CRIM="A22A39"; BAND="F3F3F5"; HEAD="0B0B0F"; LINE="D4D4D7"
SERIF="Fraunces"; SANS="IBM Plex Sans"; MONO="JetBrains Mono"

# ---- DOCX helpers ----
def cell_bg(cell, c):
    tcPr = cell._tc.get_or_add_tcPr()
    s = OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto")
    s.set(qn("w:fill"),c.lstrip("#")); tcPr.append(s)
def cell_border(cell, c=LINE):
    tcPr = cell._tc.get_or_add_tcPr()
    bs = OxmlElement("w:tcBorders")
    for e in ("top","left","bottom","right"):
        b = OxmlElement(f"w:{e}"); b.set(qn("w:val"),"single"); b.set(qn("w:sz"),"4"); b.set(qn("w:color"),c.lstrip("#"))
        bs.append(b)
    tcPr.append(bs)
def setrun(r, *, bold=False, italic=False, size=None, color=None, font=SANS):
    if bold: r.bold = True
    if italic: r.italic = True
    if size is not None: r.font.size = Pt(size)
    if color: r.font.color.rgb = RGBColor.from_string(color.lstrip("#"))
    if font:
        r.font.name = font
        rPr = r._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts"); rPr.insert(0, rFonts)
        for attr in ("w:ascii","w:hAnsi","w:eastAsia","w:cs"):
            rFonts.set(qn(attr), font)
def addbody(doc, t, *, size=11, italic=False, color=INK, sa=6, font=SANS):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(sa)
    r = p.add_run(t); setrun(r, size=size, color=color, italic=italic, font=font); return p
def addh(doc, t, *, level=1):
    sz = {1:18,2:14,3:12.5,4:11.5}
    p = doc.add_paragraph(); p.style = doc.styles[f"Heading {level}"]
    p.paragraph_format.space_before = Pt(16 if level==1 else 10); p.paragraph_format.space_after = Pt(4)
    for r in list(p.runs): r.text = ""
    r = p.add_run(t); setrun(r, bold=True, size=sz.get(level,12), color=INK, font=SERIF)
    if level == 1:
        pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom"); bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),"6"); bot.set(qn("w:space"),"4"); bot.set(qn("w:color"),GOLD)
        pBdr.append(bot); pPr.append(pBdr)
    return p
def addbullets(doc, items, *, size=11):
    for it in items:
        p = doc.add_paragraph(style="List Bullet"); r = p.add_run(it); setrun(r, size=size, color=INK, font=SANS)
def addtable(doc, header, rows, *, col_widths=None):
    n = len(header); tbl = doc.add_table(rows=len(rows)+1, cols=n); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j,h in enumerate(header):
        c = tbl.rows[0].cells[j]; cell_bg(c,HEAD); cell_border(c); c.paragraphs[0].text=""
        r = c.paragraphs[0].add_run(h); setrun(r,bold=True,size=10,color="FFFFFF", font=SANS)
        if col_widths: c.width = Inches(col_widths[j])
    for i,row in enumerate(rows, start=1):
        for j,v in enumerate(row):
            c = tbl.rows[i].cells[j]; cell_border(c)
            if i%2==0: cell_bg(c,BAND)
            c.paragraphs[0].text=""; r = c.paragraphs[0].add_run(str(v)); setrun(r,size=10,color=INK, font=SANS)
            if col_widths: c.width = Inches(col_widths[j])
    doc.add_paragraph(); return tbl
def addcallout(doc, kind, body, *, accent=GOLD, fill="F6F1E4"):
    tbl = doc.add_table(rows=1, cols=1); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = tbl.rows[0].cells[0]; cell_bg(c,fill); c.paragraphs[0].text=""
    rt = c.paragraphs[0].add_run(kind.upper()); setrun(rt,bold=True,size=9,color=accent, font=MONO)
    p = c.add_paragraph(); p.paragraph_format.space_before = Pt(3)
    r = p.add_run(body); setrun(r,size=10.5,color=INK, font=SANS); doc.add_paragraph()
def addkv(doc, pairs, *, lw=1.6, rw=4.9):
    tbl = doc.add_table(rows=len(pairs), cols=2); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,(k,v) in enumerate(pairs):
        ck,cv = tbl.rows[i].cells; ck.width=Inches(lw); cv.width=Inches(rw)
        cell_bg(ck,BAND); cell_border(ck); cell_border(cv)
        ck.paragraphs[0].text=""; cv.paragraphs[0].text=""
        rk = ck.paragraphs[0].add_run(k); setrun(rk,bold=True,size=10,color=INK, font=SANS)
        rv = cv.paragraphs[0].add_run(v); setrun(rv,size=10,color=INK, font=SANS)
    doc.add_paragraph()
def addeyebrow(doc, t):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(t.upper()); setrun(r,bold=True,size=9,color=MUTE, font=MONO)
    rPr = r._element.get_or_add_rPr(); sp = OxmlElement("w:spacing"); sp.set(qn("w:val"),"80"); rPr.append(sp)
def pb(doc):
    p = doc.add_paragraph(); p.add_run().add_break(WD_BREAK.PAGE)
def addcode(doc, t):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(t); setrun(r, size=9.5, color=INK, font=MONO)


# =========================================================================
# 1. MCP SERVER DEEP DIVE (DOCX)
# =========================================================================

def build_mcp_deep_dive():
    print("Building MCP Server Deep Dive docx...")
    doc = Document()
    for s in doc.sections:
        s.page_width = Inches(8.5); s.page_height = Inches(11)
        s.top_margin = s.bottom_margin = Inches(0.85); s.left_margin = s.right_margin = Inches(0.95)
    doc.styles["Normal"].font.name = SANS
    doc.styles["Normal"].font.size = Pt(11)

    # ===== COVER =====
    addeyebrow(doc, "APEX · RC-E2E-03 · MCP SERVER · DEEP DIVE")
    p = doc.add_paragraph(); r = p.add_run("MCP Server Deep Dive")
    setrun(r, bold=True, size=24, color=INK, font=SERIF)
    p2 = doc.add_paragraph(); r2 = p2.add_run("The 19 tool capabilities that the 6 agents call · architecture · security · maturation")
    setrun(r2, bold=True, size=14, color=INK2, italic=True, font=SERIF)
    doc.add_paragraph()
    addbody(doc,
            "Architectural deep dive on the Model Context Protocol (MCP) server that hosts the "
            "tool capabilities consumed by the six RC-E2E-03 agents. Defines the tool registry, "
            "JSON-RPC contract surface, security and authorization model, observability stack, "
            "rate-limiting and failure handling, and the wave-by-wave tool buildout. The MCP "
            "server is the single integration surface between APEX agents and the rest of "
            "Kroger's data and action estate — getting it right unlocks everything else.",
            italic=True, color=INK2)
    doc.add_paragraph()
    addkv(doc, [
        ("Document type",     "MCP Server Architectural Deep Dive"),
        ("Service",           "RC-E2E-03 Assortment & Pricing"),
        ("MCP tools cataloged","19 tool capabilities · 4 categories"),
        ("Reference client",  "The Kroger Co."),
        ("Audience",          "Lead Engineer · Solution Architect · Platform Engineering · CAF MCP team"),
        ("Companion",         "APEX-RC-E2E-03-Six-Agents-Deep-Dive · Service-Sequence-Diagram.html · manifests/*.yaml · APEX-RC-E2E-03-09-SOR-ERD.html"),
        ("Version / date",    "v1.0 · 27 Apr 2026"),
    ])
    pb(doc)

    # ===== SECTION 1 — WHAT MCP IS =====
    addh(doc, "1 · What MCP is, and why it matters here", level=1)
    addbody(doc,
            "Model Context Protocol (MCP) is the open JSON-RPC 2.0 standard for exposing tools to "
            "agents. The MCP server is a long-running process that publishes a catalog of tool "
            "capabilities; agents discover the catalog, invoke tools by name with typed arguments, "
            "and consume typed results. APEX standardizes on MCP because it provides three things "
            "that matter operationally:")
    addbullets(doc, [
        "Capability indirection — agents reference tools by capability tag (cap:scml-lot-query) rather than by URL or implementation detail; the server can rewire backends without manifest changes",
        "Strong typing — every tool publishes JSON Schema for arguments and results; the agent runtime validates both directions",
        "Auditability — every tool invocation logs to LEDGER with trace_id correlation; the audit row attribution chain stays intact across the agent-to-tool boundary",
    ])
    addh(doc, "1.1 · The architectural placement", level=2)
    addbody(doc,
            "The MCP server sits between the agent runtime (CAF Dynamic Orchestrator + AEF LLM "
            "Gateway) and the rest of Kroger's data + action estate. Three categories of consumers:")
    addtable(doc, ["Consumer", "How it calls MCP", "Trust boundary"],
             [["APEX agents (Assess, Classify, Quantify, Approve, Act, Evidence-Write)",
               "Via CAF agent runtime · stdio or streamable HTTP · auto-discovered capabilities",
               "Inside Kroger tenant boundary"],
              ["CAF orchestrator",
               "MCP discovery API for capability resolution at session start",
               "Inside CAF substrate"],
              ["Power BI / external dashboards",
               "Read-only MCP queries via service principal (limited tool subset)",
               "Tenant boundary · service-account authenticated"]],
             col_widths=[2.4, 2.8, 1.8])

    addcallout(doc, "Why MCP and not direct integration",
               "Without MCP, every agent embeds knowledge of POS APIs, vendor portals, FSMA "
               "systems, and so on. Each integration becomes a separate manifest dependency. "
               "With MCP, agents reference capability tags; the server hosts the integration. "
               "Adding a new banner's POS system means adding an MCP-server adapter, not "
               "modifying every agent.")

    pb(doc)

    # ===== SECTION 2 — TOOL CATALOG =====
    addh(doc, "2 · The 19-tool catalog", level=1)
    addbody(doc,
            "Nineteen tool capabilities organized in four categories. Each category has different "
            "operational characteristics — read tools are high-volume and latency-sensitive; "
            "ML inference tools are stateful and version-pinned; write tools are mutation-bearing "
            "and need idempotency; observability tools are write-heavy but tolerant of latency.")

    addh(doc, "2.1 · Category A · Read tools (data-access)", level=2)
    addbody(doc, "Nine read tools that pull from Bronze and Silver for agent decisioning. "
                 "These map directly to ERD entities (see SOR-ERD.html for the visual).")
    addtable(doc, ["Tool capability", "Bronze/Silver source", "Used by", "Latency SLO"],
             [["cap:telemetry-window-query",     "BRZ.RC.REFRIG_TELEMETRY (window)",     "Assess",                 "≤ 100ms"],
              ["cap:fixture-state-lookup",       "SCML.Fixture (Silver)",                "Assess",                 "≤ 50ms"],
              ["cap:fixture-failure-history",    "BRZ.RC.REFRIG_SERVICE",                "Assess",                 "≤ 200ms"],
              ["cap:scml-lot-query",             "SCML.Lot (Silver) + BRZ.RC.LOT_TRACE", "Classify · Quantify",    "≤ 100ms"],
              ["cap:dim-item-lookup",            "DIM.Item (Silver) ← BRZ.RC.ITEM_MASTER","Classify",              "≤ 50ms"],
              ["cap:fsma-policy-resolver",       "FSMA-204 policy corpus (versioned)",   "Classify",               "≤ 100ms"],
              ["cap:merml-price-lookup",         "MERML.Price (Silver)",                 "Quantify",               "≤ 100ms"],
              ["cap:merml-elasticity-curve",     "MERML.Elasticity (Silver)",            "Quantify",               "≤ 200ms"],
              ["cap:scml-lot-qty",               "SCML.Lot quantity-on-hand",            "Quantify",               "≤ 50ms"]],
             col_widths=[2.0, 2.4, 1.5, 1.0])

    addh(doc, "2.2 · Category B · ML inference tools", level=2)
    addbody(doc, "Two ML-inference tools backed by the production models in the AI/ML Model Spec.")
    addtable(doc, ["Tool capability", "Backing model", "Used by", "Latency SLO"],
             [["cap:risk-score",                 "Risk-Score (XGBoost ensemble)",        "Assess",                 "≤ 50ms (p99)"],
              ["cap:disposition-classification", "Disposition Classification (transformer + calibrated head)", "Classify", "≤ 200ms (p99)"]],
             col_widths=[2.4, 2.6, 1.5, 1.5])
    addh(doc, "2.3 · Category C · Action tools (mutation)", level=2)
    addbody(doc, "Five action tools that mutate external systems. Each must be idempotent and "
                 "must support compensating actions on partial failure.")
    addtable(doc, ["Tool capability", "Target system", "Used by", "Idempotency key"],
             [["cap:teams-adaptive-card",        "Microsoft Teams (Adaptive Card delivery)", "Approve",            "trace_id + decision_id"],
              ["cap:entra-identity-resolve",     "Microsoft Entra (group resolution)",       "Approve",            "approver_id + ts"],
              ["cap:pos-markdown-upload",        "POS systems (Toshiba/IBM/NCR · banner-aware)", "Act",            "decision_id + sku + store"],
              ["cap:vendor-portal-ticket",       "Refrigeration vendor portals (REST)",      "Act",                "decision_id + fixture"],
              ["cap:fsma-regulatory-log",        "Quality Mgmt System (Trace One/etc)",      "Act",                "decision_id + lot_code"]],
             col_widths=[2.0, 2.4, 1.5, 1.5])

    addh(doc, "2.4 · Category D · Observability + LEDGER tools", level=2)
    addbody(doc, "Three tools that write evidence and lineage. The LEDGER write is the architectural "
                 "spine — every agent decision flows through it.")
    addtable(doc, ["Tool capability", "Target", "Used by", "SLO"],
             [["cap:ledger-write",               "LEDGER (Delta + Postgres index)",      "All 6 agents",           "≤ 100ms · ≥ 99.5% success"],
              ["cap:fsma-packet-update",         "Rolling quarterly evidence packet",    "Evidence-Write",         "≤ 1s (async)"],
              ["cap:purview-lineage-emit",       "Microsoft Purview lineage graph",      "Evidence-Write",         "≤ 500ms (async)"]],
             col_widths=[2.0, 2.4, 1.5, 1.7])

    addcallout(doc, "Tool count summary",
               "9 read tools (Category A) + 2 ML tools (Category B) + 5 action tools (Category C) + "
               "3 observability tools (Category D) = 19 tool capabilities total. Each is a versioned "
               "capability tag; backend implementations are interchangeable behind the tag.")

    pb(doc)

    # ===== SECTION 3 — TOOL CONTRACT FORMAT =====
    addh(doc, "3 · Tool contract format", level=1)
    addh(doc, "3.1 · JSON-RPC 2.0 over MCP transport", level=2)
    addbody(doc,
            "MCP uses JSON-RPC 2.0 as the wire format. Three transport options supported by CAF:")
    addtable(doc, ["Transport", "Use case", "Notes"],
             [["stdio",                  "Local-process tools (e.g., dev-time tool prototyping)", "Lowest latency · single-process"],
              ["streamable HTTP (SSE)",  "Default production transport",                           "Server-Sent Events · multiplexed"],
              ["WebSocket",              "Bidirectional streaming (cap:teams-adaptive-card waits)", "Used for HITL await patterns"]],
             col_widths=[2.0, 3.0, 2.0])

    addh(doc, "3.2 · Capability declaration (tool registration)", level=2)
    addbody(doc, "Each tool declares its capability tag, JSON Schema for arguments, JSON Schema for "
                 "results, and operational metadata. Sample registration for cap:scml-lot-query:")
    addcode(doc,
            "{\n"
            '  \"capability_tag\": \"cap:scml-lot-query\",\n'
            '  \"version\": \"1.0.0\",\n'
            '  \"description\": \"Query SCML.Lot canonical for affected lots within an excursion window\",\n'
            '  \"args_schema\": {\n'
            '    \"type\": \"object\",\n'
            '    \"properties\": {\n'
            '      \"fixture_id\":   { \"type\": \"string\", \"description\": \"Fixture identifier\" },\n'
            '      \"event_window\": { \"type\": \"object\", \"properties\": {\n'
            '        \"start_ts\": { \"type\": \"string\", \"format\": \"date-time\" },\n'
            '        \"end_ts\":   { \"type\": \"string\", \"format\": \"date-time\" }\n'
            '      } },\n'
            '      \"banner\":       { \"type\": \"string\" }\n'
            '    },\n'
            '    \"required\": [\"fixture_id\", \"event_window\"]\n'
            '  },\n'
            '  \"result_schema\": {\n'
            '    \"type\": \"array\",\n'
            '    \"items\": { \"$ref\": \"#/schemas/SCML.Lot\" }\n'
            '  },\n'
            '  \"slo\": { \"latency_p99_ms\": 100, \"availability\": 99.9 },\n'
            '  \"backing_source\": \"sch_rc_e2e_09.slv_scml_lot\",\n'
            '  \"required_scopes\": [\"apex.scml.read\"],\n'
            '  \"observability\": { \"trace\": true, \"audit_row\": true },\n'
            '  \"rate_limit\": { \"per_session\": 100, \"per_minute\": 5000 }\n'
            "}")

    addh(doc, "3.3 · Invocation pattern", level=2)
    addbody(doc, "Agents invoke tools via JSON-RPC 2.0 with trace propagation:")
    addcode(doc,
            "// Request from Classify agent\n"
            "{\n"
            '  \"jsonrpc\": \"2.0\",\n'
            '  \"id\": \"req_2027-04-22_06:08:31_clf_001\",\n'
            '  \"method\": \"cap:scml-lot-query\",\n'
            '  \"params\": {\n'
            '    \"fixture_id\": \"KR-428-DAIRY-A3\",\n'
            '    \"event_window\": {\n'
            '      \"start_ts\": \"2027-04-22T05:38:00Z\",\n'
            '      \"end_ts\":   \"2027-04-22T07:00:00Z\"\n'
            '    },\n'
            '    \"banner\": \"KR\"\n'
            '  },\n'
            '  \"meta\": {\n'
            '    \"trace_id\":   \"trc_2027-04-22_0608_kr_428_a3\",\n'
            '    \"agent_id\":   \"rc-e2e-03-classify-agent\",\n'
            '    \"manifest_version\": \"v1.0.0\",\n'
            '    \"caller_identity\": \"sp:apex-classify-prod\"\n'
            '  }\n'
            "}")

    addh(doc, "3.4 · Versioning policy", level=2)
    addbullets(doc, [
        "Capability tags follow semver — cap:scml-lot-query · v1.0.0",
        "PATCH (e.g., 1.0.0 → 1.0.1) — bug fixes, no schema change · auto-deploy",
        "MINOR (1.0.0 → 1.1.0) — additive schema fields · backward-compatible · no manifest change required",
        "MAJOR (1.0.0 → 2.0.0) — breaking schema change · manifest version-up required · old version remains available 90 days for grace migration",
    ])

    pb(doc)

    # ===== SECTION 4 — SECURITY + AUTHORIZATION =====
    addh(doc, "4 · Security and authorization", level=1)
    addh(doc, "4.1 · Identity model", level=2)
    addbody(doc,
            "Every agent invocation carries an Entra service principal identity. The MCP server "
            "validates the identity, resolves scopes, and enforces tool-level permissions. Three "
            "principal categories:")
    addtable(doc, ["Principal type", "Issued to", "Scope examples"],
             [["Agent service principal",
               "Each named agent (assess-agent, classify-agent, etc.)",
               "apex.scml.read · apex.merml.read · apex.ledger.write"],
              ["Orchestrator service principal",
               "CAF Dynamic Orchestrator",
               "apex.discover · apex.session.start · apex.audit.read"],
              ["Human-via-Card principal",
               "Adaptive Card return path · Entra group-bound identity",
               "apex.hitl.approve · scope-narrowed to specific decision_id"]],
             col_widths=[2.0, 2.5, 2.5])

    addh(doc, "4.2 · Tool-scope enforcement", level=2)
    addbody(doc,
            "Each tool declares required_scopes. The MCP server checks the principal's scope "
            "claims at every invocation; mismatches return 403 with structured error.")
    addtable(doc, ["Scope", "Granted to", "What it permits"],
             [["apex.refrig.read",          "assess-agent",                        "cap:telemetry-window-query · cap:fixture-state-lookup · cap:fixture-failure-history"],
              ["apex.scml.read",            "classify-agent · quantify-agent",     "cap:scml-lot-query · cap:scml-lot-qty"],
              ["apex.dim.read",             "all agents",                          "cap:dim-item-lookup"],
              ["apex.fsma.read",            "classify-agent · evidence-write",     "cap:fsma-policy-resolver"],
              ["apex.merml.read",           "quantify-agent",                      "cap:merml-price-lookup · cap:merml-elasticity-curve"],
              ["apex.ml.invoke",            "assess · classify",                   "cap:risk-score · cap:disposition-classification"],
              ["apex.hitl.dispatch",        "approve-agent",                       "cap:teams-adaptive-card · cap:entra-identity-resolve"],
              ["apex.action.write",         "act-agent",                           "cap:pos-markdown-upload · cap:vendor-portal-ticket · cap:fsma-regulatory-log"],
              ["apex.ledger.write",         "all agents",                          "cap:ledger-write"],
              ["apex.evidence.compose",     "evidence-write-agent",                "cap:fsma-packet-update · cap:purview-lineage-emit"]],
             col_widths=[1.8, 2.5, 2.5])

    addh(doc, "4.3 · Privilege escalation prevention", level=2)
    addbullets(doc, [
        "Agent principals never receive apex.action.write except for act-agent — the architectural commitment that mutation is concentrated in one agent",
        "Read tools reject attempts to traverse to identifying customer PII fields — Purview classification rules enforce at the SQL surface",
        "HITL approvals capture the human's Entra identity into the audit row — not the agent's principal — preserving the bright line between machine-decision and human-approval",
        "Cross-tenant tool calls are blocked at the MCP server — tenant_id is derived from caller_identity and overrides any args payload",
    ])

    addh(doc, "4.4 · Audit trail of MCP invocations", level=2)
    addbody(doc,
            "Every MCP invocation produces a 'tool-call' audit row entry. The agent decision audit "
            "row references these tool-call rows by id, building the full reasoning chain. Auditors "
            "can drill from any agent decision down to the specific tool calls that produced it.")

    pb(doc)

    # ===== SECTION 5 — RATE LIMITING + FAILURE =====
    addh(doc, "5 · Rate limiting and failure handling", level=1)
    addh(doc, "5.1 · Rate-limit tiers", level=2)
    addtable(doc, ["Tool category", "Per-session limit", "Per-minute limit", "Notes"],
             [["A · Read tools",            "100/session", "5,000",  "Most-called category · scaled with banner volume"],
              ["B · ML inference",          "30/session",  "1,500",  "Model-server backed · GPU-bound at peak"],
              ["C · Action tools",          "10/session",  "200",    "Mutation-bearing · low-volume"],
              ["D · Observability",         "200/session", "10,000", "Highest-volume · async-tolerant"]],
             col_widths=[1.8, 1.5, 1.5, 2.5])

    addh(doc, "5.2 · Failure semantics", level=2)
    addtable(doc, ["Error code", "Meaning", "Agent response"],
             [["-32600 (invalid request)",  "Schema validation failed",                                       "Treat as bug · log + alert · do not retry"],
              ["-32601 (method not found)", "Capability tag not registered or version mismatch",              "Check manifest pinning · escalate to platform"],
              ["-32603 (internal error)",   "Backend integration failure",                                    "Retry per category policy · log to LEDGER"],
              ["-32000 (rate limited)",     "Per-tier throttle exceeded",                                     "Exponential backoff · respect Retry-After header"],
              ["-32001 (auth denied)",      "Scope mismatch or expired token",                                "Refresh token · do not retry on second 401"],
              ["-32002 (backend timeout)",  "Tool exceeded its SLO",                                          "Single retry · then circuit-break · degrade gracefully"]],
             col_widths=[2.0, 2.6, 2.4])

    addh(doc, "5.3 · Retry policy by tool category", level=2)
    addtable(doc, ["Category", "Retry on transient", "Backoff", "Circuit-break threshold"],
             [["A Read",          "Yes (3 retries)",     "Exponential 100ms-2s", "5 consecutive failures · 60s cool-down"],
              ["B ML inference",  "Yes (1 retry)",       "Linear 200ms",         "3 failures · 30s cool-down · fall back to heuristic"],
              ["C Action",        "Yes (3 retries)",     "Exponential 1s-30s",   "Per-target circuit · escalate to on-call"],
              ["D Observability", "Yes (5 retries)",     "Linear 100ms",         "Failure tolerance — buffer locally and replay"]],
             col_widths=[1.4, 1.7, 2.0, 2.4])

    addh(doc, "5.4 · Compensating actions for Category C tools", level=2)
    addbody(doc,
            "When an action tool succeeds and a parallel sibling fails, the system must compensate. "
            "Each Category C tool defines a compensating action that the act-agent invokes:")
    addtable(doc, ["Tool", "Compensating action on partial failure"],
             [["cap:pos-markdown-upload",        "If POS upload succeeded but vendor-ticket failed: schedule retry; do not roll back POS markdown"],
              ["cap:vendor-portal-ticket",       "If ticket created but POS or FSMA log failed: leave ticket open; retry siblings; alert Facilities Eng"],
              ["cap:fsma-regulatory-log",        "Most-critical · if FSMA log fails after POS+vendor succeed: alert Food Safety Lead immediately for manual entry · degrade to HITL"]],
             col_widths=[2.6, 4.0])

    pb(doc)

    # ===== SECTION 6 — OBSERVABILITY =====
    addh(doc, "6 · Observability", level=1)
    addh(doc, "6.1 · The four pillars", level=2)
    addtable(doc, ["Pillar", "Implementation", "Use case"],
             [["Tracing",      "OpenTelemetry traces · W3C traceparent propagation · LangFuse span store",
                                "Drill from any agent decision to the specific tool calls"],
              ["Metrics",      "Prometheus endpoint + Azure Monitor",
                                "Per-tool p50/p95/p99 latency · per-tool error rate · per-tier rate-limit utilization"],
              ["Logs",         "Application Insights + LEDGER tool-call audit rows",
                                "Operational debugging · audit reconstruction"],
              ["Lineage",      "Microsoft Purview lineage emission via cap:purview-lineage-emit",
                                "Bronze → Silver → Tool-call → Agent-decision → Outcome chain"]],
             col_widths=[1.4, 2.6, 2.6])

    addh(doc, "6.2 · Critical SLO dashboards", level=2)
    addbullets(doc, [
        "MCP server availability (target: 99.95%) — single SLO that gates the entire agent fleet",
        "Per-tool latency (target: p99 within stated SLO) — alerts on 1-hour rolling window violations",
        "LEDGER write success rate (target: 99.5%) — non-negotiable; audit-row coverage depends on it",
        "Tool-version drift (catalog vs. consumed) — alerts on stale-pin manifests",
        "Cross-trace integrity (sampled) — random audit of traceparent → LEDGER trace_id linkage",
    ])

    pb(doc)

    # ===== SECTION 7 — WAVE BUILDOUT =====
    addh(doc, "7 · Wave-by-wave tool buildout", level=1)
    addbody(doc,
            "Not all 19 tools land in Wave 1. The buildout sequences tools to match agent maturity "
            "and ingest readiness.")
    addtable(doc, ["Wave", "Tools landed", "Notes"],
             [["W1 Foundation",
               "Read: telemetry-window-query · fixture-state-lookup · scml-lot-query · dim-item-lookup · fsma-policy-resolver. ML: risk-score (heuristic baseline). Obs: ledger-write.",
               "7 tools · supports Assess + Classify in shadow-run · no mutation in W1"],
              ["W1 extension",
               "+ merml-price-lookup · merml-elasticity-curve · scml-lot-qty · fixture-failure-history",
               "Adds Quantify pre-requisites · still no mutation"],
              ["W2 Pilot",
               "+ disposition-classification (live) · teams-adaptive-card · entra-identity-resolve · pos-markdown-upload · vendor-portal-ticket · fsma-regulatory-log · fsma-packet-update · purview-lineage-emit",
               "All 19 tools live · full mutation surface · 90-day operating run"],
              ["W3 Scale & Fuse",
               "All 19 tools at enterprise scale · cross-banner adaptation · multi-tenant patterns",
               "Production-hardened · disaster recovery · multi-region replication"]],
             col_widths=[1.3, 3.5, 1.8])

    addh(doc, "7.1 · Tool-version pinning by wave", level=2)
    addbody(doc,
            "Tenant manifests pin to specific tool capability versions. Upgrade to MAJOR versions "
            "requires explicit manifest update; PATCH and MINOR are auto-served.")
    addcode(doc,
            "# tenant manifest snippet · kroger-w2-pilot.yaml\n"
            "tenant_id: kr\n"
            "wave: W2-pilot\n"
            "tool_pins:\n"
            "  - capability: cap:disposition-classification\n"
            "    pinned_version: \"1.0.x\"      # accept PATCH but not MINOR/MAJOR\n"
            "  - capability: cap:teams-adaptive-card\n"
            "    pinned_version: \"^2.0.0\"     # accept MINOR + PATCH\n"
            "  - capability: cap:fsma-policy-resolver\n"
            "    pinned_version: \"=1.0.0\"     # exact pin · regulatory traceability\n")

    pb(doc)

    # ===== SECTION 8 — TOOL MATURATION =====
    addh(doc, "8 · Tool maturation alignment", level=1)
    addbody(doc,
            "Just as agents mature (see Six-Agents Deep Dive § 3-4), tools mature alongside them. "
            "The MCP server's tool registry tracks per-tool maturity bands.")
    addtable(doc, ["Tool", "W1", "W2", "W3", "W4+ superagent fusion"],
             [["cap:telemetry-window-query",     "Reactive · 30-min window", "+ Multi-fixture cascade", "+ Predictive risk windowing", "Same tool · proactive callers"],
              ["cap:disposition-classification", "Single-FTL-category", "Cross-FTL-category", "Cross-banner pattern", "Cross-Practice (HLS 340B etc.)"],
              ["cap:merml-elasticity-curve",     "Per-SKU baseline", "Per-cohort overlay", "Cross-banner pooling", "Cross-Practice (TMT subscription)"],
              ["cap:teams-adaptive-card",        "Static template", "Tier-aware variants", "Approver-load-balanced", "Principle-based card generation"],
              ["cap:ledger-write",                "Fixed 14-field schema", "Schema additive evolution", "Cross-tenant aggregation patterns", "Cross-scenario evidence types"]],
             col_widths=[2.0, 1.4, 1.4, 1.4, 1.6])

    addh(doc, "8.1 · Tool registry as the authoritative catalog", level=2)
    addbody(doc,
            "The CAF MCP server's tool registry is the authoritative catalog. Per-capability "
            "metadata: version, maturity band, backing system, SLO, scope, deprecation timeline. "
            "Manifests resolve against this registry at session start; the orchestrator pins the "
            "correct version for the wave.")

    pb(doc)

    # ===== SECTION 9 — INDEPENDENCE + CAF =====
    addh(doc, "9 · Independence posture and CAF integration", level=1)
    addh(doc, "9.1 · Why MCP server is a Kroger artifact", level=2)
    addcallout(doc, "Independence by construction",
               "The MCP server runs in Kroger's tenant. Tool implementations are Kroger artifacts. "
               "Backing data sources are Kroger data. Deloitte's role is to advise on tool "
               "specifications and capability tagging — the integrations themselves are "
               "Kroger-owned. This preserves audit independence even if Kroger's external auditor "
               "relationship raises constraints.",
               accent=GOLD)

    addh(doc, "9.2 · CAF MCP shared-tool absorption", level=2)
    addbody(doc,
            "Several of the 19 tools have direct analogues in CAF's existing tool catalog. The "
            "MCP server can either implement Kroger-specific versions or consume CAF-shared "
            "implementations via the AEF LLM Gateway federation.")
    addtable(doc, ["Tool", "Kroger-specific or CAF-shared?", "Rationale"],
             [["cap:risk-score",                 "Kroger-specific",      "Trained on Kroger telemetry · banner-tuned"],
              ["cap:disposition-classification", "Kroger-specific",      "FSMA 204 policy corpus is Kroger's authority"],
              ["cap:merml-elasticity-curve",     "Kroger-specific",      "84.51° relationship · proprietary basket data"],
              ["cap:teams-adaptive-card",        "CAF-shared (federate)", "Identical pattern across all APEX scenarios"],
              ["cap:entra-identity-resolve",     "CAF-shared (federate)", "Microsoft Entra is the universal substrate"],
              ["cap:ledger-write",               "CAF-shared (federate)", "14-field audit row is APEX-canonical"],
              ["cap:purview-lineage-emit",       "CAF-shared (federate)", "Microsoft Purview is the universal substrate"]],
             col_widths=[2.4, 2.0, 2.6])

    pb(doc)

    # ===== APPENDICES =====
    addh(doc, "Appendix A · Glossary", level=1)
    addtable(doc, ["Term", "Definition"],
             [["MCP",                "Model Context Protocol · open JSON-RPC 2.0 standard for agent-tool integration"],
              ["Capability tag",     "cap:* identifier · agent-side reference to a tool · backend-implementation-agnostic"],
              ["Tool registry",      "CAF-managed authoritative catalog of capability tags + versions + metadata"],
              ["Service principal",  "Entra non-human identity used by agents and orchestrator"],
              ["Scope",              "Permission claim attached to a service principal · enforced at MCP invocation"],
              ["Compensating action","Cleanup logic invoked when a parallel action partially fails"],
              ["Idempotency key",    "Per-invocation unique identifier preventing duplicate side-effects on retry"],
              ["Rate-limit tier",    "Per-category throttle bucket (A/B/C/D) · per-session and per-minute limits"],
              ["Circuit-break",      "Fault-tolerance pattern · pauses calls to a failing dependency until recovery"],
              ["JSON-RPC 2.0",       "Wire-format standard used by MCP for request/response semantics"],
              ["W3C traceparent",    "OpenTelemetry trace propagation header · binds agent decisions to tool calls"]],
             col_widths=[1.8, 4.6])

    addh(doc, "Appendix B · Cross-references", level=1)
    addtable(doc, ["Artifact", "Where it complements"],
             [["APEX-RC-E2E-03-Six-Agents-Deep-Dive-and-Maturation.docx", "How the 6 agents call the 19 tools"],
              ["APEX-RC-E2E-03-Service-Sequence-Diagram.html",            "Visual sequence of agent ↔ MCP ↔ external system"],
              ["APEX-RC-E2E-03-09-SOR-ERD.html",                          "Bronze tables that back the 9 read tools"],
              ["APEX-RC-E2E-03-09-AI-ML-Model-Spec.docx",                 "Model details for the 2 ML inference tools"],
              ["manifests/parent-orchestrator.yaml",                       "Tool-version pinning at orchestrator session start"],
              ["manifests/{assess,classify,quantify,approve,act}-agent.yaml", "Per-agent tool capability declarations"],
              ["APEX-RC-E2E-03-09-Privacy-Data-Governance-Spec.docx",     "Scope-enforcement and consent-state-aware tool gating"]],
             col_widths=[3.4, 3.0])

    addbody(doc,
            "Independence posture preserved across this document: 'Deloitte's Microsoft practice,' "
            "'DMTSP,' 'Microsoft platform capabilities,' 'Microsoft-native deployment.' The terms "
            "'partner,' 'alliance,' and 'strategic alliance' do not appear.",
            italic=True, color=MUTE, size=9)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— End of MCP Server Deep Dive —"); setrun(r, italic=True, size=10, color=MUTE, font=SERIF)

    target = OUT_DIR / "APEX-RC-E2E-03-MCP-Server-Deep-Dive.docx"
    try: doc.save(str(target))
    except PermissionError:
        n=2
        while True:
            alt = target.with_name(target.stem + f"-v{n}" + target.suffix)
            if not alt.exists(): target = alt; break
            n+=1
        doc.save(str(target))
    print(f"  wrote {target.name}")


# =========================================================================
# 2. SERVICE SEQUENCE DIAGRAM (HTML)
# =========================================================================

SEQUENCE_HTML = r"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8" />
<meta name="viewport" content="width=device-width, initial-scale=1.0" />
<title>APEX · RC-E2E-03 · Service Sequence Diagram</title>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
<link href="https://fonts.googleapis.com/css2?family=Fraunces:ital,opsz,wght@0,9..144,400;0,9..144,500;0,9..144,600&family=IBM+Plex+Sans:wght@0,300;400;500;600;700&family=JetBrains+Mono:wght@400;500;600&display=swap" rel="stylesheet">
<style>
:root {
  --bg-0:#07090F;--bg-1:#0C1119;--bg-2:#141B27;--bg-3:#1C2433;
  --line:rgba(255,255,255,.10);--line-strong:rgba(255,255,255,.20);
  --ink-0:#F4EFE5;--ink-1:#D9D2C3;--ink-2:#9A9383;--ink-3:#6B6757;
  --gold:#E3B657;--teal:#3DD9C4;--crimson:#D94B5A;--violet:#9B7BE0;
  --sky:#6FB6E5;--amber:#E8A33D;--ember:#E87B3C;
  /* sequence-diagram colors */
  --c-trigger:#6FB6E5; --c-orch:#E8A33D; --c-assess:#6FB6E5;
  --c-classify:#9B7BE0; --c-quantify:#3DD9C4; --c-approve:#E3B657;
  --c-act:#E87B3C; --c-evidence:#D94B5A;
  --c-mcp:#9A9383; --c-external:#6B6757; --c-ledger:#E3B657;
  --diag-bg:#07090F; --diag-grid:rgba(255,255,255,.05);
}
body.light {
  --bg-0:#F5F1E8;--bg-1:#EEE7D7;--bg-2:#E5DCC7;--bg-3:#D9CFB6;
  --line:rgba(20,14,4,.10);--line-strong:rgba(20,14,4,.22);
  --ink-0:#1A1508;--ink-1:#3D3420;--ink-2:#6B5F45;--ink-3:#8E8367;
  --gold:#8E6A1E;--teal:#147A6C;--crimson:#A22A39;
  --violet:#6548B0;--sky:#3E7FAE;--amber:#B8791F;--ember:#B0551B;
  --c-trigger:#3E7FAE; --c-orch:#B8791F; --c-assess:#3E7FAE;
  --c-classify:#6548B0; --c-quantify:#147A6C; --c-approve:#8E6A1E;
  --c-act:#B0551B; --c-evidence:#A22A39;
  --c-mcp:#6B5F45; --c-external:#8E8367; --c-ledger:#8E6A1E;
  --diag-bg:#F0E9D6; --diag-grid:rgba(20,14,4,.06);
}
*{box-sizing:border-box}html,body{margin:0;padding:0}
body{
  background:var(--bg-0);color:var(--ink-0);
  font-family:"IBM Plex Sans",system-ui,sans-serif;
  font-size:14.5px;line-height:1.55;min-height:100vh;
  transition:background .5s,color .5s;
}
body::before{
  content:"";position:fixed;inset:0;
  background:radial-gradient(1200px 600px at 10% -10%,rgba(232,163,61,.08),transparent 60%),
             radial-gradient(900px 500px at 95% 110%,rgba(155,123,224,.06),transparent 60%);
  pointer-events:none;z-index:0;
}
.theme-toggle{
  position:fixed;top:24px;right:24px;z-index:100;
  display:flex;align-items:center;gap:8px;
  background:var(--bg-1);border:1px solid var(--line-strong);
  border-radius:999px;padding:8px 14px;cursor:pointer;
  font-family:"JetBrains Mono",monospace;font-size:11px;letter-spacing:.12em;text-transform:uppercase;
  color:var(--ink-1);user-select:none;
}
.theme-toggle:hover{border-color:var(--gold);color:var(--ink-0)}

.hero{position:relative;z-index:1;max-width:1500px;margin:0 auto;padding:56px 48px 32px;border-bottom:1px solid var(--line)}
.eyebrow{
  display:inline-block;font-family:"JetBrains Mono",monospace;
  font-size:11px;letter-spacing:.18em;text-transform:uppercase;color:var(--gold);
  border:1px solid var(--line-strong);padding:6px 14px;border-radius:999px;margin-bottom:22px;
}
h1{
  font-family:"Fraunces",serif;font-weight:500;
  font-size:clamp(36px,5vw,60px);line-height:1.06;letter-spacing:-0.02em;margin:0 0 14px;
}
h1 em{font-style:italic;font-weight:400;color:var(--amber)}
.subhead{font-size:16px;color:var(--ink-1);max-width:920px;margin:0 0 28px}
.meta-strip{
  display:flex;flex-wrap:wrap;gap:22px;
  font-family:"JetBrains Mono",monospace;font-size:11.5px;color:var(--ink-2);
  border-top:1px solid var(--line);padding-top:18px;
}
.meta-strip b{color:var(--ink-0);font-weight:500}

section{position:relative;z-index:1;max-width:1500px;margin:0 auto;padding:48px;border-bottom:1px solid var(--line)}
.section-head{margin-bottom:24px}
.section-eye{
  font-family:"JetBrains Mono",monospace;font-size:11px;letter-spacing:.15em;
  text-transform:uppercase;color:var(--ink-3);margin-bottom:8px;
}
h2{font-family:"Fraunces",serif;font-weight:500;font-size:clamp(26px,3vw,38px);line-height:1.12;margin:0 0 6px}
.lead{color:var(--ink-1);max-width:920px}

/* ---- LEGEND ---- */
.legend{
  display:flex;flex-wrap:wrap;gap:14px;margin-top:14px;
  font-family:"JetBrains Mono",monospace;font-size:11px;
}
.legend .item{
  display:flex;align-items:center;gap:8px;color:var(--ink-1);
  border:1px solid var(--line);padding:6px 12px;border-radius:999px;
}
.legend .swatch{width:12px;height:12px;border-radius:3px}

/* ---- DIAGRAM ---- */
.diag-wrap{
  background:var(--bg-1);border:1px solid var(--line);border-radius:12px;
  padding:16px;margin-top:24px;overflow-x:auto;
}
.diag-svg{width:100%;min-width:1300px;display:block;background:var(--diag-bg);border-radius:8px}

.phase-name{
  font-family:"JetBrains Mono",monospace;font-size:10.5px;letter-spacing:.16em;
  text-transform:uppercase;color:var(--accent);margin-bottom:6px;
}
h3.phase-h{font-family:"Fraunces",serif;font-size:22px;font-weight:500;margin:0 0 4px}
.phase-blurb{color:var(--ink-1);font-size:14px;margin:0 0 14px;max-width:920px}

/* ---- TOOL TABLE ---- */
.tool-table{
  background:var(--bg-1);border:1px solid var(--line);border-radius:12px;
  overflow:hidden;margin-top:18px;
}
.tool-table table{width:100%;border-collapse:collapse}
.tool-table th{
  background:var(--bg-2);padding:12px 16px;text-align:left;
  font-family:"JetBrains Mono",monospace;font-size:10.5px;letter-spacing:.14em;
  text-transform:uppercase;color:var(--ink-2);border-bottom:1px solid var(--line);
}
.tool-table td{padding:12px 16px;border-bottom:1px solid var(--line);font-size:13px;color:var(--ink-1)}
.tool-table tr:last-child td{border-bottom:none}
.tool-table .cap{font-family:"JetBrains Mono",monospace;color:var(--violet);font-weight:600}
.tool-table .agent{font-family:"JetBrains Mono",monospace;color:var(--gold);font-size:11.5px}

/* ---- FOOTER ---- */
footer{
  position:relative;z-index:1;max-width:1500px;margin:0 auto;
  padding:32px 48px 64px;font-family:"JetBrains Mono",monospace;font-size:11px;color:var(--ink-3);
  display:flex;justify-content:space-between;flex-wrap:wrap;gap:12px;
}
</style>
</head>
<body>

<div class="theme-toggle" onclick="toggleTheme()">
  <span id="themeIcon">☀</span>
  <span id="themeLabel">Light</span>
</div>

<header class="hero">
  <div class="eyebrow">APEX · RC-E2E-03 · SERVICE SEQUENCE DIAGRAM</div>
  <h1>Six agents · 19 MCP tools · <em>one</em> 90-second decision.</h1>
  <p class="subhead">
    Phase-by-phase sequence diagram showing how the six RC-E2E-03 agents call the MCP server
    to read Bronze/Silver tables (per the ERD), invoke ML inference, dispatch HITL, mutate
    external systems, and write the LEDGER closeout audit row. Each phase is its own diagram;
    the final phase shows the full end-to-end closeout. Aligned with both the SOR ERD and
    the MCP Server Deep Dive.
  </p>
  <div class="meta-strip">
    <div><b>Agents</b> · 6</div>
    <div><b>MCP tools</b> · 19</div>
    <div><b>Phases</b> · 6 + closeout</div>
    <div><b>End-to-end target</b> · 90 sec</div>
    <div><b>Trace propagation</b> · W3C traceparent → LEDGER trace_id</div>
    <div><b>Companion</b> · MCP-Server-Deep-Dive · SOR-ERD</div>
  </div>
</header>

<!-- ===== LEGEND ===== -->
<section>
  <div class="section-head">
    <div class="section-eye">§ Legend</div>
    <h2>How to read these diagrams.</h2>
    <p class="lead">
      Each diagram has vertical lifelines (one per actor) and time flowing top-to-bottom.
      Solid arrows are synchronous JSON-RPC calls (request → response); dashed arrows are
      async (fire-and-forget). Activation rectangles show when an actor is processing.
      Notes (gold-bordered boxes) annotate notable architectural moments.
    </p>
  </div>
  <div class="legend">
    <div class="item"><div class="swatch" style="background:var(--c-trigger)"></div>Trigger / Eventstream</div>
    <div class="item"><div class="swatch" style="background:var(--c-orch)"></div>Orchestrator</div>
    <div class="item"><div class="swatch" style="background:var(--c-assess)"></div>Assess</div>
    <div class="item"><div class="swatch" style="background:var(--c-classify)"></div>Classify</div>
    <div class="item"><div class="swatch" style="background:var(--c-quantify)"></div>Quantify</div>
    <div class="item"><div class="swatch" style="background:var(--c-approve)"></div>Approve</div>
    <div class="item"><div class="swatch" style="background:var(--c-act)"></div>Act</div>
    <div class="item"><div class="swatch" style="background:var(--c-evidence)"></div>Evidence-Write</div>
    <div class="item"><div class="swatch" style="background:var(--c-mcp)"></div>MCP Server</div>
    <div class="item"><div class="swatch" style="background:var(--c-external)"></div>External system</div>
    <div class="item"><div class="swatch" style="background:var(--c-ledger)"></div>LEDGER</div>
  </div>
</section>

<!-- ===== PHASE 1: TRIGGER + ASSESS ===== -->
<section>
  <div class="section-head">
    <span class="phase-name" style="--accent:var(--c-assess)">Phase 1 of 6 · Trigger + Assess</span>
    <h3 class="phase-h">Refrigeration sensor → Assess agent → ExcursionProfile</h3>
    <p class="phase-blurb">Eventstream watermark fires when a fixture breaches setpoint+5°F for ≥15 min.
    Orchestrator instantiates a trace_id and dispatches Assess, which calls three read tools to
    characterize the event.</p>
  </div>

  <div class="diag-wrap">
    <svg class="diag-svg" viewBox="0 0 1300 600" xmlns="http://www.w3.org/2000/svg">
      <defs>
        <marker id="arr-sync" viewBox="0 0 10 10" refX="9" refY="5" markerWidth="9" markerHeight="9" orient="auto">
          <path d="M0,0 L10,5 L0,10 z" class="ar-fill"/>
        </marker>
        <marker id="arr-async" viewBox="0 0 10 10" refX="9" refY="5" markerWidth="9" markerHeight="9" orient="auto">
          <path d="M0,0 L10,5 L0,10 z" class="ar-fill" opacity="0.5"/>
        </marker>
        <style>
          .l-life{stroke:var(--line);stroke-dasharray:2,4}
          .l-grid{stroke:var(--diag-grid);stroke-width:.5}
          .ar-fill{fill:var(--ink-1)}
          .lbl{font-family:"JetBrains Mono",monospace;fill:var(--ink-1);font-size:11px}
          .lbl-acc{font-family:"JetBrains Mono",monospace;font-size:10.5px}
          .actor-box{stroke-width:0;rx:6}
          .actor-text{font-family:"JetBrains Mono",monospace;font-size:11px;font-weight:600;fill:#0B0B0F;text-anchor:middle}
          .note{fill:var(--bg-2);stroke:var(--gold);stroke-width:1.5;rx:4}
          .note-text{font-family:"JetBrains Mono",monospace;font-size:10px;fill:var(--ink-1)}
        </style>
      </defs>

      <!-- horizontal time gridlines -->
      <line x1="60" y1="120" x2="1240" y2="120" class="l-grid"/>
      <line x1="60" y1="220" x2="1240" y2="220" class="l-grid"/>
      <line x1="60" y1="320" x2="1240" y2="320" class="l-grid"/>
      <line x1="60" y1="420" x2="1240" y2="420" class="l-grid"/>
      <line x1="60" y1="520" x2="1240" y2="520" class="l-grid"/>

      <!-- Lifelines -->
      <!-- Trigger (Eventstream) -->
      <rect x="30" y="40" width="180" height="36" class="actor-box" fill="var(--c-trigger)"/>
      <text x="120" y="62" class="actor-text">EVENTSTREAM</text>
      <line x1="120" y1="80" x2="120" y2="580" class="l-life"/>

      <!-- Orchestrator -->
      <rect x="240" y="40" width="170" height="36" class="actor-box" fill="var(--c-orch)"/>
      <text x="325" y="62" class="actor-text">ORCHESTRATOR</text>
      <line x1="325" y1="80" x2="325" y2="580" class="l-life"/>

      <!-- Assess -->
      <rect x="450" y="40" width="160" height="36" class="actor-box" fill="var(--c-assess)"/>
      <text x="530" y="62" class="actor-text">ASSESS AGENT</text>
      <line x1="530" y1="80" x2="530" y2="580" class="l-life"/>

      <!-- MCP Server -->
      <rect x="700" y="40" width="160" height="36" class="actor-box" fill="var(--c-mcp)"/>
      <text x="780" y="62" class="actor-text">MCP SERVER</text>
      <line x1="780" y1="80" x2="780" y2="580" class="l-life"/>

      <!-- Bronze (read source) -->
      <rect x="900" y="40" width="180" height="36" class="actor-box" fill="var(--c-external)"/>
      <text x="990" y="62" class="actor-text">BRZ.RC.REFRIG_*</text>
      <line x1="990" y1="80" x2="990" y2="580" class="l-life"/>

      <!-- LEDGER -->
      <rect x="1110" y="40" width="130" height="36" class="actor-box" fill="var(--c-ledger)"/>
      <text x="1175" y="62" class="actor-text">LEDGER</text>
      <line x1="1175" y1="80" x2="1175" y2="580" class="l-life"/>

      <!-- ===== Messages ===== -->
      <!-- 1. Watermark fires (Trigger → Orch) -->
      <line x1="120" y1="120" x2="320" y2="120" stroke="var(--ink-1)" stroke-width="1.5" marker-end="url(#arr-sync)"/>
      <text x="125" y="115" class="lbl-acc" fill="var(--c-trigger)">1. watermark trigger (temp&gt;setpoint+5° · ≥15min)</text>

      <!-- Orchestrator activation -->
      <rect x="318" y="120" width="14" height="430" fill="var(--c-orch)" opacity="0.18"/>

      <!-- 2. Orch → Assess: invoke -->
      <line x1="325" y1="170" x2="525" y2="170" stroke="var(--ink-1)" stroke-width="1.5" marker-end="url(#arr-sync)"/>
      <text x="330" y="165" class="lbl-acc" fill="var(--c-orch)">2. invoke(trace_id, fixture_id, window)</text>

      <!-- Assess activation -->
      <rect x="523" y="170" width="14" height="320" fill="var(--c-assess)" opacity="0.18"/>

      <!-- 3. Assess → MCP: cap:telemetry-window-query -->
      <line x1="530" y1="220" x2="775" y2="220" stroke="var(--ink-1)" stroke-width="1.5" marker-end="url(#arr-sync)"/>
      <text x="535" y="215" class="lbl-acc" fill="var(--c-assess)">3. cap:telemetry-window-query</text>

      <!-- MCP → Bronze -->
      <line x1="780" y1="240" x2="985" y2="240" stroke="var(--ink-1)" stroke-width="1.5" marker-end="url(#arr-sync)"/>
      <text x="785" y="235" class="lbl-acc" fill="var(--c-mcp)">    SQL: SELECT FROM brz_refrig_telemetry WHERE ...</text>

      <!-- Bronze → MCP (return) -->
      <line x1="985" y1="265" x2="785" y2="265" stroke="var(--ink-2)" stroke-width="1" stroke-dasharray="3,3" marker-end="url(#arr-sync)"/>

      <!-- MCP → Assess (return) -->
      <line x1="775" y1="285" x2="535" y2="285" stroke="var(--ink-2)" stroke-width="1" stroke-dasharray="3,3" marker-end="url(#arr-sync)"/>
      <text x="540" y="280" class="lbl-acc" fill="var(--ink-2)">    ← telemetry window (84 readings)</text>

      <!-- 4. Assess → MCP: cap:fixture-state-lookup -->
      <line x1="530" y1="325" x2="775" y2="325" stroke="var(--ink-1)" stroke-width="1.5" marker-end="url(#arr-sync)"/>
      <text x="535" y="320" class="lbl-acc" fill="var(--c-assess)">4. cap:fixture-state-lookup</text>
      <line x1="775" y1="345" x2="535" y2="345" stroke="var(--ink-2)" stroke-width="1" stroke-dasharray="3,3" marker-end="url(#arr-sync)"/>
      <text x="540" y="340" class="lbl-acc" fill="var(--ink-2)">    ← SCML.Fixture (status, prior excursions)</text>

      <!-- 5. Assess → MCP: cap:fixture-failure-history -->
      <line x1="530" y1="380" x2="775" y2="380" stroke="var(--ink-1)" stroke-width="1.5" marker-end="url(#arr-sync)"/>
      <text x="535" y="375" class="lbl-acc" fill="var(--c-assess)">5. cap:fixture-failure-history</text>
      <line x1="775" y1="400" x2="535" y2="400" stroke="var(--ink-2)" stroke-width="1" stroke-dasharray="3,3" marker-end="url(#arr-sync)"/>

      <!-- 6. Assess → MCP: cap:risk-score (ML) -->
      <line x1="530" y1="430" x2="775" y2="430" stroke="var(--c-classify)" stroke-width="1.5" marker-end="url(#arr-sync)"/>
      <text x="535" y="425" class="lbl-acc" fill="var(--c-classify)">6. cap:risk-score (ML inference)</text>
      <line x1="775" y1="450" x2="535" y2="450" stroke="var(--ink-2)" stroke-width="1" stroke-dasharray="3,3" marker-end="url(#arr-sync)"/>
      <text x="540" y="445" class="lbl-acc" fill="var(--ink-2)">    ← risk score 0.91 · severity CRITICAL</text>

      <!-- 7. Assess → LEDGER: capture audit row -->
      <line x1="530" y1="490" x2="1170" y2="490" stroke="var(--c-ledger)" stroke-width="1.5" stroke-dasharray="6,2" marker-end="url(#arr-async)"/>
      <text x="540" y="485" class="lbl-acc" fill="var(--c-ledger)">7. cap:ledger-write (Assess capture row)</text>

      <!-- 8. Assess → Orch: ExcursionProfile -->
      <line x1="525" y1="530" x2="335" y2="530" stroke="var(--ink-2)" stroke-width="1" stroke-dasharray="3,3" marker-end="url(#arr-sync)"/>
      <text x="335" y="525" class="lbl-acc" fill="var(--ink-2)">8. ← ExcursionProfile {peak:48.2°, dur:47min, severity:CRITICAL}</text>

      <!-- Phase note -->
      <rect x="900" y="510" width="280" height="60" class="note"/>
      <text x="910" y="528" class="note-text">Trace propagation</text>
      <text x="910" y="544" class="note-text" fill="var(--ink-2)">trace_id: trc_2027-04-22_</text>
      <text x="910" y="558" class="note-text" fill="var(--ink-2)">0608_kr_428_a3</text>
    </svg>
  </div>

  <div class="tool-table">
    <table>
      <thead><tr><th style="width:25%">Tool capability</th><th style="width:32%">Bronze/Silver source</th><th style="width:16%">Agent</th><th style="width:27%">Result type</th></tr></thead>
      <tbody>
        <tr><td class="cap">cap:telemetry-window-query</td><td>BRZ.RC.REFRIG_TELEMETRY (window)</td><td class="agent">Assess</td><td>Array&lt;TelemetryReading&gt;</td></tr>
        <tr><td class="cap">cap:fixture-state-lookup</td><td>SCML.Fixture (Silver)</td><td class="agent">Assess</td><td>FixtureState</td></tr>
        <tr><td class="cap">cap:fixture-failure-history</td><td>BRZ.RC.REFRIG_SERVICE</td><td class="agent">Assess</td><td>Array&lt;ServiceTicket&gt;</td></tr>
        <tr><td class="cap">cap:risk-score</td><td>(ML model serving)</td><td class="agent">Assess</td><td>RiskScore (0.0-1.0)</td></tr>
        <tr><td class="cap">cap:ledger-write</td><td>(LEDGER plane)</td><td class="agent">Assess</td><td>capture audit row</td></tr>
      </tbody>
    </table>
  </div>
</section>

<!-- ===== PHASE 2: CLASSIFY ===== -->
<section>
  <div class="section-head">
    <span class="phase-name" style="--accent:var(--c-classify)">Phase 2 of 6 · Classify</span>
    <h3 class="phase-h">ExcursionProfile → DispositionTable per affected lot</h3>
    <p class="phase-blurb">Classify reasons over the FSMA 204 policy corpus joined to lot-provenance to produce a per-SKU disposition. If any lot's confidence falls below 0.85, the table is escalated to Food Safety Lead before reaching Quantify.</p>
  </div>

  <div class="diag-wrap">
    <svg class="diag-svg" viewBox="0 0 1300 700" xmlns="http://www.w3.org/2000/svg">
      <defs>
        <marker id="arr2-sync" viewBox="0 0 10 10" refX="9" refY="5" markerWidth="9" markerHeight="9" orient="auto"><path d="M0,0 L10,5 L0,10 z" fill="var(--ink-1)"/></marker>
        <style>
          .l-life2{stroke:var(--line);stroke-dasharray:2,4}
          .l-grid2{stroke:var(--diag-grid);stroke-width:.5}
          .lbl-acc2{font-family:"JetBrains Mono",monospace;font-size:10.5px}
          .actor-box2{rx:6}
          .actor-text2{font-family:"JetBrains Mono",monospace;font-size:11px;font-weight:600;fill:#0B0B0F;text-anchor:middle}
          .note2{fill:var(--bg-2);stroke:var(--c-classify);stroke-width:1.5;rx:4}
        </style>
      </defs>

      <line x1="60" y1="130" x2="1240" y2="130" class="l-grid2"/>
      <line x1="60" y1="280" x2="1240" y2="280" class="l-grid2"/>
      <line x1="60" y1="430" x2="1240" y2="430" class="l-grid2"/>
      <line x1="60" y1="580" x2="1240" y2="580" class="l-grid2"/>

      <rect x="30" y="40" width="170" height="36" class="actor-box2" fill="var(--c-orch)"/><text x="115" y="62" class="actor-text2">ORCHESTRATOR</text>
      <line x1="115" y1="80" x2="115" y2="680" class="l-life2"/>

      <rect x="220" y="40" width="170" height="36" class="actor-box2" fill="var(--c-classify)"/><text x="305" y="62" class="actor-text2">CLASSIFY AGENT</text>
      <line x1="305" y1="80" x2="305" y2="680" class="l-life2"/>

      <rect x="430" y="40" width="160" height="36" class="actor-box2" fill="var(--c-mcp)"/><text x="510" y="62" class="actor-text2">MCP SERVER</text>
      <line x1="510" y1="80" x2="510" y2="680" class="l-life2"/>

      <rect x="630" y="40" width="160" height="36" class="actor-box2" fill="var(--c-external)"/><text x="710" y="62" class="actor-text2">SCML.LOT (Silver)</text>
      <line x1="710" y1="80" x2="710" y2="680" class="l-life2"/>

      <rect x="830" y="40" width="180" height="36" class="actor-box2" fill="var(--c-external)"/><text x="920" y="62" class="actor-text2">FSMA POLICY CORPUS</text>
      <line x1="920" y1="80" x2="920" y2="680" class="l-life2"/>

      <rect x="1050" y="40" width="180" height="36" class="actor-box2" fill="var(--c-external)"/><text x="1140" y="62" class="actor-text2">DISPOSITION ML</text>
      <line x1="1140" y1="80" x2="1140" y2="680" class="l-life2"/>

      <!-- 1. Orch → Classify -->
      <line x1="115" y1="130" x2="300" y2="130" stroke="var(--ink-1)" stroke-width="1.5" marker-end="url(#arr2-sync)"/>
      <text x="120" y="125" class="lbl-acc2" fill="var(--c-orch)">1. invoke(trace_id, ExcursionProfile)</text>

      <rect x="298" y="130" width="14" height="540" fill="var(--c-classify)" opacity="0.18"/>

      <!-- 2. Classify → MCP: scml-lot-query -->
      <line x1="305" y1="180" x2="505" y2="180" stroke="var(--ink-1)" stroke-width="1.5" marker-end="url(#arr2-sync)"/>
      <text x="310" y="175" class="lbl-acc2" fill="var(--c-classify)">2. cap:scml-lot-query (fixture_id)</text>
      <line x1="510" y1="200" x2="705" y2="200" stroke="var(--ink-1)" stroke-width="1.5" marker-end="url(#arr2-sync)"/>
      <line x1="705" y1="225" x2="515" y2="225" stroke="var(--ink-2)" stroke-width="1" stroke-dasharray="3,3" marker-end="url(#arr2-sync)"/>
      <line x1="505" y1="250" x2="315" y2="250" stroke="var(--ink-2)" stroke-width="1" stroke-dasharray="3,3" marker-end="url(#arr2-sync)"/>
      <text x="320" y="245" class="lbl-acc2" fill="var(--ink-2)">    ← 6 affected lots</text>

      <!-- 3. cap:dim-item-lookup (combined) -->
      <line x1="305" y1="295" x2="505" y2="295" stroke="var(--ink-1)" stroke-width="1.5" marker-end="url(#arr2-sync)"/>
      <text x="310" y="290" class="lbl-acc2" fill="var(--c-classify)">3. cap:dim-item-lookup (sku batch)</text>
      <line x1="505" y1="320" x2="315" y2="320" stroke="var(--ink-2)" stroke-width="1" stroke-dasharray="3,3" marker-end="url(#arr2-sync)"/>
      <text x="320" y="315" class="lbl-acc2" fill="var(--ink-2)">    ← FTL flags · shelf-life · temp_class</text>

      <!-- 4. cap:fsma-policy-resolver -->
      <line x1="305" y1="345" x2="505" y2="345" stroke="var(--ink-1)" stroke-width="1.5" marker-end="url(#arr2-sync)"/>
      <text x="310" y="340" class="lbl-acc2" fill="var(--c-classify)">4. cap:fsma-policy-resolver</text>
      <line x1="510" y1="365" x2="915" y2="365" stroke="var(--ink-1)" stroke-width="1.5" marker-end="url(#arr2-sync)"/>
      <line x1="915" y1="390" x2="515" y2="390" stroke="var(--ink-2)" stroke-width="1" stroke-dasharray="3,3" marker-end="url(#arr2-sync)"/>
      <line x1="505" y1="410" x2="315" y2="410" stroke="var(--ink-2)" stroke-width="1" stroke-dasharray="3,3" marker-end="url(#arr2-sync)"/>
      <text x="320" y="405" class="lbl-acc2" fill="var(--ink-2)">    ← applicable policy clauses (per FTL)</text>

      <!-- 5. cap:disposition-classification (ML) -->
      <line x1="305" y1="450" x2="505" y2="450" stroke="var(--c-classify)" stroke-width="1.5" marker-end="url(#arr2-sync)"/>
      <text x="310" y="445" class="lbl-acc2" fill="var(--c-classify)">5. cap:disposition-classification (ML)</text>
      <line x1="510" y1="470" x2="1135" y2="470" stroke="var(--c-classify)" stroke-width="1.5" marker-end="url(#arr2-sync)"/>
      <line x1="1135" y1="495" x2="515" y2="495" stroke="var(--ink-2)" stroke-width="1" stroke-dasharray="3,3" marker-end="url(#arr2-sync)"/>
      <line x1="505" y1="515" x2="315" y2="515" stroke="var(--ink-2)" stroke-width="1" stroke-dasharray="3,3" marker-end="url(#arr2-sync)"/>
      <text x="320" y="510" class="lbl-acc2" fill="var(--ink-2)">    ← per-lot disposition + confidence + policy_citation</text>

      <!-- 6. Classify → LEDGER: capture audit row (per lot) -->
      <line x1="305" y1="555" x2="1240" y2="555" stroke="var(--c-ledger)" stroke-width="1.5" stroke-dasharray="6,2"/>
      <text x="315" y="550" class="lbl-acc2" fill="var(--c-ledger)">6. cap:ledger-write × 6 (per-lot capture rows)</text>

      <!-- 7. Classify → Orch: DispositionTable -->
      <line x1="300" y1="615" x2="125" y2="615" stroke="var(--ink-2)" stroke-width="1" stroke-dasharray="3,3" marker-end="url(#arr2-sync)"/>
      <text x="135" y="610" class="lbl-acc2" fill="var(--ink-2)">7. ← DispositionTable[6 lots] · all conf ≥ 0.85 (no escalation)</text>

      <!-- HITL note -->
      <rect x="970" y="595" width="270" height="65" class="note2"/>
      <text x="980" y="613" class="lbl-acc2" fill="var(--c-classify)">If any conf &lt; 0.85:</text>
      <text x="980" y="629" class="lbl-acc2" fill="var(--ink-2)">escalation_required = true</text>
      <text x="980" y="645" class="lbl-acc2" fill="var(--ink-2)">→ Food Safety Lead via</text>
      <text x="980" y="660" class="lbl-acc2" fill="var(--ink-2)">cap:teams-adaptive-card (out-of-band)</text>
    </svg>
  </div>

  <div class="tool-table">
    <table>
      <thead><tr><th style="width:28%">Tool capability</th><th style="width:32%">Source</th><th style="width:14%">Agent</th><th style="width:26%">Result type</th></tr></thead>
      <tbody>
        <tr><td class="cap">cap:scml-lot-query</td><td>SCML.Lot ← BRZ.RC.LOT_TRACE + RECEIPT</td><td class="agent">Classify</td><td>Array&lt;Lot&gt;</td></tr>
        <tr><td class="cap">cap:dim-item-lookup</td><td>DIM.Item ← BRZ.RC.ITEM_MASTER</td><td class="agent">Classify</td><td>Array&lt;ItemAttrs&gt;</td></tr>
        <tr><td class="cap">cap:fsma-policy-resolver</td><td>FSMA-204 policy corpus (versioned)</td><td class="agent">Classify</td><td>Array&lt;PolicyClause&gt;</td></tr>
        <tr><td class="cap">cap:disposition-classification</td><td>(ML model · 4-class)</td><td class="agent">Classify</td><td>Array&lt;DispositionResult&gt;</td></tr>
        <tr><td class="cap">cap:ledger-write</td><td>(LEDGER plane)</td><td class="agent">Classify</td><td>per-lot capture rows</td></tr>
      </tbody>
    </table>
  </div>
</section>

<!-- ===== PHASE 3: QUANTIFY ===== -->
<section>
  <div class="section-head">
    <span class="phase-name" style="--accent:var(--c-quantify)">Phase 3 of 6 · Quantify</span>
    <h3 class="phase-h">DispositionTable → QuantifiedDisposition with $ impact</h3>
    <p class="phase-blurb">Quantify converts each per-lot disposition into three-scenario dollar impacts (sell-through, markdown, destroy) using MERML.Price + MERML.Elasticity + SCML.Lot quantity. The output is what arrives at the Approve agent's Adaptive Card.</p>
  </div>

  <div class="diag-wrap">
    <svg class="diag-svg" viewBox="0 0 1300 600" xmlns="http://www.w3.org/2000/svg">
      <defs>
        <marker id="arr3-sync" viewBox="0 0 10 10" refX="9" refY="5" markerWidth="9" markerHeight="9" orient="auto"><path d="M0,0 L10,5 L0,10 z" fill="var(--ink-1)"/></marker>
        <style>
          .l-life3{stroke:var(--line);stroke-dasharray:2,4}
          .l-grid3{stroke:var(--diag-grid);stroke-width:.5}
          .lbl-acc3{font-family:"JetBrains Mono",monospace;font-size:10.5px}
          .actor-text3{font-family:"JetBrains Mono",monospace;font-size:11px;font-weight:600;fill:#0B0B0F;text-anchor:middle}
        </style>
      </defs>

      <line x1="60" y1="130" x2="1240" y2="130" class="l-grid3"/>
      <line x1="60" y1="260" x2="1240" y2="260" class="l-grid3"/>
      <line x1="60" y1="390" x2="1240" y2="390" class="l-grid3"/>
      <line x1="60" y1="520" x2="1240" y2="520" class="l-grid3"/>

      <rect x="30" y="40" width="160" height="36" rx="6" fill="var(--c-orch)"/><text x="110" y="62" class="actor-text3">ORCHESTRATOR</text>
      <line x1="110" y1="80" x2="110" y2="580" class="l-life3"/>

      <rect x="220" y="40" width="170" height="36" rx="6" fill="var(--c-quantify)"/><text x="305" y="62" class="actor-text3">QUANTIFY AGENT</text>
      <line x1="305" y1="80" x2="305" y2="580" class="l-life3"/>

      <rect x="430" y="40" width="160" height="36" rx="6" fill="var(--c-mcp)"/><text x="510" y="62" class="actor-text3">MCP SERVER</text>
      <line x1="510" y1="80" x2="510" y2="580" class="l-life3"/>

      <rect x="630" y="40" width="170" height="36" rx="6" fill="var(--c-external)"/><text x="715" y="62" class="actor-text3">MERML.PRICE</text>
      <line x1="715" y1="80" x2="715" y2="580" class="l-life3"/>

      <rect x="840" y="40" width="180" height="36" rx="6" fill="var(--c-external)"/><text x="930" y="62" class="actor-text3">MERML.ELASTICITY</text>
      <line x1="930" y1="80" x2="930" y2="580" class="l-life3"/>

      <rect x="1060" y="40" width="170" height="36" rx="6" fill="var(--c-external)"/><text x="1145" y="62" class="actor-text3">SCML.LOT.QTY</text>
      <line x1="1145" y1="80" x2="1145" y2="580" class="l-life3"/>

      <line x1="110" y1="130" x2="300" y2="130" stroke="var(--ink-1)" stroke-width="1.5" marker-end="url(#arr3-sync)"/>
      <text x="115" y="125" class="lbl-acc3" fill="var(--c-orch)">1. invoke(DispositionTable)</text>

      <rect x="298" y="130" width="14" height="430" fill="var(--c-quantify)" opacity="0.18"/>

      <line x1="305" y1="180" x2="505" y2="180" stroke="var(--ink-1)" stroke-width="1.5" marker-end="url(#arr3-sync)"/>
      <text x="310" y="175" class="lbl-acc3" fill="var(--c-quantify)">2. cap:merml-price-lookup (sku × zone)</text>
      <line x1="510" y1="200" x2="710" y2="200" stroke="var(--ink-1)" stroke-width="1.5" marker-end="url(#arr3-sync)"/>
      <line x1="710" y1="225" x2="510" y2="225" stroke="var(--ink-2)" stroke-width="1" stroke-dasharray="3,3" marker-end="url(#arr3-sync)"/>

      <line x1="305" y1="290" x2="505" y2="290" stroke="var(--ink-1)" stroke-width="1.5" marker-end="url(#arr3-sync)"/>
      <text x="310" y="285" class="lbl-acc3" fill="var(--c-quantify)">3. cap:merml-elasticity-curve (sku batch)</text>
      <line x1="510" y1="310" x2="925" y2="310" stroke="var(--ink-1)" stroke-width="1.5" marker-end="url(#arr3-sync)"/>
      <line x1="925" y1="335" x2="510" y2="335" stroke="var(--ink-2)" stroke-width="1" stroke-dasharray="3,3" marker-end="url(#arr3-sync)"/>
      <line x1="505" y1="355" x2="315" y2="355" stroke="var(--ink-2)" stroke-width="1" stroke-dasharray="3,3" marker-end="url(#arr3-sync)"/>

      <line x1="305" y1="420" x2="505" y2="420" stroke="var(--ink-1)" stroke-width="1.5" marker-end="url(#arr3-sync)"/>
      <text x="310" y="415" class="lbl-acc3" fill="var(--c-quantify)">4. cap:scml-lot-qty</text>
      <line x1="510" y1="440" x2="1140" y2="440" stroke="var(--ink-1)" stroke-width="1.5" marker-end="url(#arr3-sync)"/>
      <line x1="1140" y1="460" x2="510" y2="460" stroke="var(--ink-2)" stroke-width="1" stroke-dasharray="3,3" marker-end="url(#arr3-sync)"/>

      <text x="320" y="490" class="lbl-acc3" fill="var(--c-quantify)">5. compute: sell-through_$ · markdown_$ · destroy_loss_$ per lot</text>
      <text x="320" y="505" class="lbl-acc3" fill="var(--c-quantify)">    + recommended markdown_pct (elasticity-optimal)</text>

      <line x1="305" y1="540" x2="1240" y2="540" stroke="var(--c-ledger)" stroke-width="1.5" stroke-dasharray="6,2"/>
      <text x="315" y="535" class="lbl-acc3" fill="var(--c-ledger)">6. cap:ledger-write (Quantify capture row)</text>

      <line x1="300" y1="565" x2="120" y2="565" stroke="var(--ink-2)" stroke-width="1" stroke-dasharray="3,3" marker-end="url(#arr3-sync)"/>
      <text x="125" y="560" class="lbl-acc3" fill="var(--ink-2)">7. ← QuantifiedDisposition[6 lots] · total $7,920 at risk</text>
    </svg>
  </div>
</section>

<!-- ===== PHASE 4: APPROVE (HITL) ===== -->
<section>
  <div class="section-head">
    <span class="phase-name" style="--accent:var(--c-approve)">Phase 4 of 6 · Approve · HITL gate</span>
    <h3 class="phase-h">QuantifiedDisposition → ApprovedDisposition (with human-in-the-loop)</h3>
    <p class="phase-blurb">Approve resolves the Store Operations Lead's Entra identity, dispatches a Microsoft Teams Adaptive Card, and awaits the human response (Approve / Modify / Reject). The 10-minute timeout escalates to Food Safety Lead. Decision writes a HITL log row to LEDGER with the human's named identity.</p>
  </div>

  <div class="diag-wrap">
    <svg class="diag-svg" viewBox="0 0 1300 700" xmlns="http://www.w3.org/2000/svg">
      <defs>
        <marker id="arr4-sync" viewBox="0 0 10 10" refX="9" refY="5" markerWidth="9" markerHeight="9" orient="auto"><path d="M0,0 L10,5 L0,10 z" fill="var(--ink-1)"/></marker>
        <style>
          .l-life4{stroke:var(--line);stroke-dasharray:2,4}
          .l-grid4{stroke:var(--diag-grid);stroke-width:.5}
          .lbl-acc4{font-family:"JetBrains Mono",monospace;font-size:10.5px}
          .actor-text4{font-family:"JetBrains Mono",monospace;font-size:11px;font-weight:600;fill:#0B0B0F;text-anchor:middle}
          .note4{fill:var(--bg-2);stroke:var(--c-approve);stroke-width:1.5;rx:4}
        </style>
      </defs>

      <line x1="60" y1="130" x2="1240" y2="130" class="l-grid4"/>
      <line x1="60" y1="280" x2="1240" y2="280" class="l-grid4"/>
      <line x1="60" y1="430" x2="1240" y2="430" class="l-grid4"/>
      <line x1="60" y1="580" x2="1240" y2="580" class="l-grid4"/>

      <rect x="30" y="40" width="160" height="36" rx="6" fill="var(--c-orch)"/><text x="110" y="62" class="actor-text4">ORCHESTRATOR</text>
      <line x1="110" y1="80" x2="110" y2="680" class="l-life4"/>

      <rect x="220" y="40" width="170" height="36" rx="6" fill="var(--c-approve)"/><text x="305" y="62" class="actor-text4">APPROVE AGENT</text>
      <line x1="305" y1="80" x2="305" y2="680" class="l-life4"/>

      <rect x="430" y="40" width="160" height="36" rx="6" fill="var(--c-mcp)"/><text x="510" y="62" class="actor-text4">MCP SERVER</text>
      <line x1="510" y1="80" x2="510" y2="680" class="l-life4"/>

      <rect x="630" y="40" width="170" height="36" rx="6" fill="var(--c-external)"/><text x="715" y="62" class="actor-text4">MS ENTRA</text>
      <line x1="715" y1="80" x2="715" y2="680" class="l-life4"/>

      <rect x="840" y="40" width="180" height="36" rx="6" fill="var(--c-external)"/><text x="930" y="62" class="actor-text4">MS TEAMS</text>
      <line x1="930" y1="80" x2="930" y2="680" class="l-life4"/>

      <rect x="1060" y="40" width="170" height="36" rx="6" fill="var(--c-trigger)"/><text x="1145" y="62" class="actor-text4">STORE OPS LEAD</text>
      <line x1="1145" y1="80" x2="1145" y2="680" class="l-life4"/>

      <line x1="110" y1="130" x2="300" y2="130" stroke="var(--ink-1)" stroke-width="1.5" marker-end="url(#arr4-sync)"/>
      <text x="115" y="125" class="lbl-acc4" fill="var(--c-orch)">1. invoke(QuantifiedDisposition · materiality_policy)</text>

      <rect x="298" y="130" width="14" height="540" fill="var(--c-approve)" opacity="0.18"/>

      <text x="320" y="170" class="lbl-acc4" fill="var(--c-approve)">2. evaluate materiality_policy → HITL required (tier 2 · $7,920 &gt; $2,500)</text>

      <line x1="305" y1="200" x2="505" y2="200" stroke="var(--ink-1)" stroke-width="1.5" marker-end="url(#arr4-sync)"/>
      <text x="310" y="195" class="lbl-acc4" fill="var(--c-approve)">3. cap:entra-identity-resolve (apex-rc-store-ops · 428)</text>
      <line x1="510" y1="220" x2="710" y2="220" stroke="var(--ink-1)" stroke-width="1.5" marker-end="url(#arr4-sync)"/>
      <line x1="710" y1="245" x2="510" y2="245" stroke="var(--ink-2)" stroke-width="1" stroke-dasharray="3,3" marker-end="url(#arr4-sync)"/>
      <line x1="505" y1="265" x2="315" y2="265" stroke="var(--ink-2)" stroke-width="1" stroke-dasharray="3,3" marker-end="url(#arr4-sync)"/>
      <text x="320" y="260" class="lbl-acc4" fill="var(--ink-2)">    ← marisol.reyes@kroger.com</text>

      <line x1="305" y1="320" x2="505" y2="320" stroke="var(--ink-1)" stroke-width="1.5" marker-end="url(#arr4-sync)"/>
      <text x="310" y="315" class="lbl-acc4" fill="var(--c-approve)">4. cap:teams-adaptive-card (compose &amp; send)</text>
      <line x1="510" y1="345" x2="925" y2="345" stroke="var(--ink-1)" stroke-width="1.5" marker-end="url(#arr4-sync)"/>
      <line x1="930" y1="370" x2="1140" y2="370" stroke="var(--ink-1)" stroke-width="1.5" marker-end="url(#arr4-sync)"/>
      <text x="935" y="365" class="lbl-acc4" fill="var(--c-mcp)">    push to device</text>

      <!-- Wait period -->
      <rect x="298" y="395" width="14" height="100" fill="var(--c-approve)" opacity="0.30" stroke="var(--c-approve)" stroke-dasharray="3,3"/>
      <text x="318" y="445" class="lbl-acc4" fill="var(--c-approve)">5. await human response · timeout 10 min</text>

      <line x1="1140" y1="495" x2="935" y2="495" stroke="var(--c-approve)" stroke-width="2" marker-end="url(#arr4-sync)"/>
      <text x="945" y="490" class="lbl-acc4" fill="var(--c-approve)">6. APPROVE clicked at 6:09:30</text>
      <line x1="925" y1="515" x2="510" y2="515" stroke="var(--ink-2)" stroke-width="1" stroke-dasharray="3,3" marker-end="url(#arr4-sync)"/>
      <line x1="505" y1="535" x2="315" y2="535" stroke="var(--ink-2)" stroke-width="1" stroke-dasharray="3,3" marker-end="url(#arr4-sync)"/>
      <text x="320" y="530" class="lbl-acc4" fill="var(--ink-2)">    ← {hitl:HITL, action:APPROVE, actor:marisol.reyes, ts:06:09:30Z}</text>

      <line x1="305" y1="600" x2="1240" y2="600" stroke="var(--c-ledger)" stroke-width="1.5" stroke-dasharray="6,2"/>
      <text x="315" y="595" class="lbl-acc4" fill="var(--c-ledger)">7. cap:ledger-write (Approve HITL log row · with named identity)</text>

      <line x1="300" y1="640" x2="125" y2="640" stroke="var(--ink-2)" stroke-width="1" stroke-dasharray="3,3" marker-end="url(#arr4-sync)"/>
      <text x="135" y="635" class="lbl-acc4" fill="var(--ink-2)">8. ← ApprovedDisposition</text>

      <rect x="850" y="610" width="380" height="65" class="note4"/>
      <text x="860" y="628" class="lbl-acc4" fill="var(--c-approve)">If REJECTED: terminate flow at this step.</text>
      <text x="860" y="645" class="lbl-acc4" fill="var(--ink-2)">If MODIFIED: human-edited disposition becomes</text>
      <text x="860" y="660" class="lbl-acc4" fill="var(--ink-2)">ground truth · feeds model retraining (UC-22)</text>
    </svg>
  </div>
</section>

<!-- ===== PHASE 5: ACT (parallel) ===== -->
<section>
  <div class="section-head">
    <span class="phase-name" style="--accent:var(--c-act)">Phase 5 of 6 · Act · parallel downstream actions</span>
    <h3 class="phase-h">ApprovedDisposition → 3 parallel mutations</h3>
    <p class="phase-blurb">Act fires three mutations in parallel: POS markdown upload (banner-aware), refrigeration vendor ticket (vendor-aware), and FSMA regulatory log entry. Each has independent retry policy and idempotency keys. Failures escalate without blocking siblings.</p>
  </div>

  <div class="diag-wrap">
    <svg class="diag-svg" viewBox="0 0 1300 600" xmlns="http://www.w3.org/2000/svg">
      <defs>
        <marker id="arr5-sync" viewBox="0 0 10 10" refX="9" refY="5" markerWidth="9" markerHeight="9" orient="auto"><path d="M0,0 L10,5 L0,10 z" fill="var(--ink-1)"/></marker>
        <style>
          .l-life5{stroke:var(--line);stroke-dasharray:2,4}
          .l-grid5{stroke:var(--diag-grid);stroke-width:.5}
          .lbl-acc5{font-family:"JetBrains Mono",monospace;font-size:10.5px}
          .actor-text5{font-family:"JetBrains Mono",monospace;font-size:11px;font-weight:600;fill:#0B0B0F;text-anchor:middle}
        </style>
      </defs>

      <line x1="60" y1="130" x2="1240" y2="130" class="l-grid5"/>
      <line x1="60" y1="240" x2="1240" y2="240" class="l-grid5"/>
      <line x1="60" y1="380" x2="1240" y2="380" class="l-grid5"/>
      <line x1="60" y1="500" x2="1240" y2="500" class="l-grid5"/>

      <rect x="30" y="40" width="170" height="36" rx="6" fill="var(--c-orch)"/><text x="115" y="62" class="actor-text5">ORCHESTRATOR</text>
      <line x1="115" y1="80" x2="115" y2="580" class="l-life5"/>

      <rect x="230" y="40" width="160" height="36" rx="6" fill="var(--c-act)"/><text x="310" y="62" class="actor-text5">ACT AGENT</text>
      <line x1="310" y1="80" x2="310" y2="580" class="l-life5"/>

      <rect x="430" y="40" width="160" height="36" rx="6" fill="var(--c-mcp)"/><text x="510" y="62" class="actor-text5">MCP SERVER</text>
      <line x1="510" y1="80" x2="510" y2="580" class="l-life5"/>

      <rect x="630" y="40" width="170" height="36" rx="6" fill="var(--c-external)"/><text x="715" y="62" class="actor-text5">POS (NCR/IBM)</text>
      <line x1="715" y1="80" x2="715" y2="580" class="l-life5"/>

      <rect x="840" y="40" width="180" height="36" rx="6" fill="var(--c-external)"/><text x="930" y="62" class="actor-text5">EMERSON PORTAL</text>
      <line x1="930" y1="80" x2="930" y2="580" class="l-life5"/>

      <rect x="1060" y="40" width="180" height="36" rx="6" fill="var(--c-external)"/><text x="1150" y="62" class="actor-text5">FSMA QMS</text>
      <line x1="1150" y1="80" x2="1150" y2="580" class="l-life5"/>

      <line x1="115" y1="130" x2="305" y2="130" stroke="var(--ink-1)" stroke-width="1.5" marker-end="url(#arr5-sync)"/>
      <text x="120" y="125" class="lbl-acc5" fill="var(--c-orch)">1. invoke(ApprovedDisposition)</text>

      <rect x="303" y="130" width="14" height="450" fill="var(--c-act)" opacity="0.18"/>

      <text x="325" y="170" class="lbl-acc5" fill="var(--c-act)">2. fire 3 actions in parallel · idempotency-key per action</text>

      <!-- Three parallel calls -->
      <line x1="310" y1="220" x2="505" y2="220" stroke="var(--ink-1)" stroke-width="1.5" marker-end="url(#arr5-sync)"/>
      <text x="315" y="215" class="lbl-acc5" fill="var(--c-act)">3a. cap:pos-markdown-upload</text>
      <line x1="510" y1="240" x2="710" y2="240" stroke="var(--ink-1)" stroke-width="1.5" marker-end="url(#arr5-sync)"/>

      <line x1="310" y1="270" x2="505" y2="270" stroke="var(--ink-1)" stroke-width="1.5" marker-end="url(#arr5-sync)"/>
      <text x="315" y="265" class="lbl-acc5" fill="var(--c-act)">3b. cap:vendor-portal-ticket</text>
      <line x1="510" y1="290" x2="925" y2="290" stroke="var(--ink-1)" stroke-width="1.5" marker-end="url(#arr5-sync)"/>

      <line x1="310" y1="320" x2="505" y2="320" stroke="var(--ink-1)" stroke-width="1.5" marker-end="url(#arr5-sync)"/>
      <text x="315" y="315" class="lbl-acc5" fill="var(--c-act)">3c. cap:fsma-regulatory-log</text>
      <line x1="510" y1="340" x2="1145" y2="340" stroke="var(--ink-1)" stroke-width="1.5" marker-end="url(#arr5-sync)"/>

      <!-- Returns -->
      <line x1="710" y1="395" x2="510" y2="395" stroke="var(--ink-2)" stroke-width="1" stroke-dasharray="3,3" marker-end="url(#arr5-sync)"/>
      <line x1="505" y1="415" x2="315" y2="415" stroke="var(--ink-2)" stroke-width="1" stroke-dasharray="3,3" marker-end="url(#arr5-sync)"/>
      <text x="320" y="410" class="lbl-acc5" fill="var(--ink-2)">    ← POS markdown active (rcv 06:10:08)</text>

      <line x1="925" y1="435" x2="510" y2="435" stroke="var(--ink-2)" stroke-width="1" stroke-dasharray="3,3" marker-end="url(#arr5-sync)"/>
      <line x1="505" y1="455" x2="315" y2="455" stroke="var(--ink-2)" stroke-width="1" stroke-dasharray="3,3" marker-end="url(#arr5-sync)"/>
      <text x="320" y="450" class="lbl-acc5" fill="var(--ink-2)">    ← vendor_ticket_id: EMR-2027-04-22-1841</text>

      <line x1="1145" y1="475" x2="510" y2="475" stroke="var(--ink-2)" stroke-width="1" stroke-dasharray="3,3" marker-end="url(#arr5-sync)"/>
      <line x1="505" y1="495" x2="315" y2="495" stroke="var(--ink-2)" stroke-width="1" stroke-dasharray="3,3" marker-end="url(#arr5-sync)"/>
      <text x="320" y="490" class="lbl-acc5" fill="var(--ink-2)">    ← fsma_log_id: FSMA-KR-428-2027-04-22-001</text>

      <line x1="310" y1="535" x2="1240" y2="535" stroke="var(--c-ledger)" stroke-width="1.5" stroke-dasharray="6,2"/>
      <text x="320" y="530" class="lbl-acc5" fill="var(--c-ledger)">4. cap:ledger-write (Act action-log row · 3 sub-actions)</text>

      <line x1="305" y1="565" x2="125" y2="565" stroke="var(--ink-2)" stroke-width="1" stroke-dasharray="3,3" marker-end="url(#arr5-sync)"/>
      <text x="135" y="560" class="lbl-acc5" fill="var(--ink-2)">5. ← ActOutcome {all_actions_completed:true}</text>
    </svg>
  </div>
</section>

<!-- ===== PHASE 6: EVIDENCE-WRITE (CLOSEOUT) ===== -->
<section>
  <div class="section-head">
    <span class="phase-name" style="--accent:var(--c-evidence)">Phase 6 of 6 · Evidence-Write · closeout</span>
    <h3 class="phase-h">All upstream traces → 14-field closeout audit row → LEDGER + Purview + FSMA packet</h3>
    <p class="phase-blurb">Evidence-Write composes the canonical 14-field closeout audit row from all upstream agent traces, writes it to LEDGER (Delta + Postgres index), emits Purview lineage, and updates the rolling quarterly FSMA evidence packet. The agent that makes "records are generated, not collated" structural.</p>
  </div>

  <div class="diag-wrap">
    <svg class="diag-svg" viewBox="0 0 1300 580" xmlns="http://www.w3.org/2000/svg">
      <defs>
        <marker id="arr6-sync" viewBox="0 0 10 10" refX="9" refY="5" markerWidth="9" markerHeight="9" orient="auto"><path d="M0,0 L10,5 L0,10 z" fill="var(--ink-1)"/></marker>
        <style>
          .l-life6{stroke:var(--line);stroke-dasharray:2,4}
          .l-grid6{stroke:var(--diag-grid);stroke-width:.5}
          .lbl-acc6{font-family:"JetBrains Mono",monospace;font-size:10.5px}
          .actor-text6{font-family:"JetBrains Mono",monospace;font-size:11px;font-weight:600;fill:#0B0B0F;text-anchor:middle}
          .note6{fill:var(--bg-2);stroke:var(--c-evidence);stroke-width:1.5;rx:4}
        </style>
      </defs>

      <line x1="60" y1="130" x2="1240" y2="130" class="l-grid6"/>
      <line x1="60" y1="260" x2="1240" y2="260" class="l-grid6"/>
      <line x1="60" y1="390" x2="1240" y2="390" class="l-grid6"/>
      <line x1="60" y1="500" x2="1240" y2="500" class="l-grid6"/>

      <rect x="30" y="40" width="170" height="36" rx="6" fill="var(--c-orch)"/><text x="115" y="62" class="actor-text6">ORCHESTRATOR</text>
      <line x1="115" y1="80" x2="115" y2="560" class="l-life6"/>

      <rect x="230" y="40" width="190" height="36" rx="6" fill="var(--c-evidence)"/><text x="325" y="62" class="actor-text6">EVIDENCE-WRITE</text>
      <line x1="325" y1="80" x2="325" y2="560" class="l-life6"/>

      <rect x="450" y="40" width="160" height="36" rx="6" fill="var(--c-mcp)"/><text x="530" y="62" class="actor-text6">MCP SERVER</text>
      <line x1="530" y1="80" x2="530" y2="560" class="l-life6"/>

      <rect x="640" y="40" width="170" height="36" rx="6" fill="var(--c-ledger)"/><text x="725" y="62" class="actor-text6">LEDGER</text>
      <line x1="725" y1="80" x2="725" y2="560" class="l-life6"/>

      <rect x="850" y="40" width="170" height="36" rx="6" fill="var(--c-external)"/><text x="935" y="62" class="actor-text6">FSMA PACKET</text>
      <line x1="935" y1="80" x2="935" y2="560" class="l-life6"/>

      <rect x="1060" y="40" width="170" height="36" rx="6" fill="var(--c-external)"/><text x="1145" y="62" class="actor-text6">MS PURVIEW</text>
      <line x1="1145" y1="80" x2="1145" y2="560" class="l-life6"/>

      <line x1="115" y1="130" x2="320" y2="130" stroke="var(--ink-1)" stroke-width="1.5" marker-end="url(#arr6-sync)"/>
      <text x="120" y="125" class="lbl-acc6" fill="var(--c-orch)">1. invoke(trace_id, all_upstream_audit_rows[5+])</text>

      <rect x="318" y="130" width="14" height="430" fill="var(--c-evidence)" opacity="0.18"/>

      <text x="340" y="170" class="lbl-acc6" fill="var(--c-evidence)">2. compose 14-field closeout audit row</text>
      <text x="340" y="185" class="lbl-acc6" fill="var(--c-evidence)">    + correlate LangFuse spans</text>
      <text x="340" y="200" class="lbl-acc6" fill="var(--c-evidence)">    + stamp manifest_version + policy_version + prompt_version</text>
      <text x="340" y="215" class="lbl-acc6" fill="var(--c-evidence)">    + bind hitl_actor (marisol.reyes)</text>
      <text x="340" y="230" class="lbl-acc6" fill="var(--c-evidence)">    + compute KPIs (time_to_decision_sec, shrink_avoided_usd)</text>

      <line x1="325" y1="295" x2="525" y2="295" stroke="var(--ink-1)" stroke-width="1.5" marker-end="url(#arr6-sync)"/>
      <text x="330" y="290" class="lbl-acc6" fill="var(--c-evidence)">3. cap:ledger-write (closeout · primary)</text>
      <line x1="530" y1="315" x2="720" y2="315" stroke="var(--ink-1)" stroke-width="1.5" marker-end="url(#arr6-sync)"/>
      <line x1="720" y1="340" x2="530" y2="340" stroke="var(--ink-2)" stroke-width="1" stroke-dasharray="3,3" marker-end="url(#arr6-sync)"/>
      <line x1="525" y1="360" x2="335" y2="360" stroke="var(--ink-2)" stroke-width="1" stroke-dasharray="3,3" marker-end="url(#arr6-sync)"/>
      <text x="340" y="355" class="lbl-acc6" fill="var(--ink-2)">    ← row_id: trc_2027-04-22_0608_kr_428_a3 · 99.7% complete</text>

      <line x1="325" y1="425" x2="525" y2="425" stroke="var(--ink-1)" stroke-width="1.5" stroke-dasharray="6,2" marker-end="url(#arr6-sync)"/>
      <text x="330" y="420" class="lbl-acc6" fill="var(--c-evidence)">4. cap:fsma-packet-update (async)</text>
      <line x1="530" y1="445" x2="930" y2="445" stroke="var(--ink-1)" stroke-width="1.5" stroke-dasharray="6,2" marker-end="url(#arr6-sync)"/>

      <line x1="325" y1="475" x2="525" y2="475" stroke="var(--ink-1)" stroke-width="1.5" stroke-dasharray="6,2" marker-end="url(#arr6-sync)"/>
      <text x="330" y="470" class="lbl-acc6" fill="var(--c-evidence)">5. cap:purview-lineage-emit (async)</text>
      <line x1="530" y1="495" x2="1140" y2="495" stroke="var(--ink-1)" stroke-width="1.5" stroke-dasharray="6,2" marker-end="url(#arr6-sync)"/>

      <line x1="320" y1="540" x2="125" y2="540" stroke="var(--ink-2)" stroke-width="1" stroke-dasharray="3,3" marker-end="url(#arr6-sync)"/>
      <text x="135" y="535" class="lbl-acc6" fill="var(--ink-2)">6. ← closeout complete · trace_id sealed · scenario complete</text>

      <rect x="850" y="510" width="380" height="50" class="note6"/>
      <text x="860" y="528" class="lbl-acc6" fill="var(--c-evidence)">FSMA packet + Purview emission are async.</text>
      <text x="860" y="544" class="lbl-acc6" fill="var(--ink-2)">Closeout audit row is the synchronous critical path.</text>
    </svg>
  </div>
</section>

<!-- ===== ERD/MCP ALIGNMENT ===== -->
<section>
  <div class="section-head">
    <div class="section-eye">§ ERD ↔ MCP alignment</div>
    <h2>Each MCP read tool maps to a specific Bronze (or Silver) entity.</h2>
    <p class="lead">
      The 19 tools tie into the 12 Bronze tables documented in the SOR ERD. The mapping below
      shows which Bronze table backs each tool. Tools that read Silver (canonical layer)
      transitively touch their underlying Bronze sources.
    </p>
  </div>

  <div class="tool-table">
    <table>
      <thead><tr>
        <th style="width:24%">MCP tool</th>
        <th style="width:22%">Direct backing</th>
        <th style="width:32%">Underlying Bronze (per ERD)</th>
        <th style="width:10%">Category</th>
        <th style="width:12%">Used by</th>
      </tr></thead>
      <tbody>
        <tr><td class="cap">cap:telemetry-window-query</td><td>BRZ.RC.REFRIG_TELEMETRY</td><td>BRZ.RC.REFRIG_TELEMETRY</td><td>A · Read</td><td class="agent">Assess</td></tr>
        <tr><td class="cap">cap:fixture-state-lookup</td><td>SCML.Fixture (Silver)</td><td>BRZ.RC.REFRIG_TELEMETRY · BRZ.RC.REFRIG_SERVICE</td><td>A · Read</td><td class="agent">Assess</td></tr>
        <tr><td class="cap">cap:fixture-failure-history</td><td>BRZ.RC.REFRIG_SERVICE</td><td>BRZ.RC.REFRIG_SERVICE</td><td>A · Read</td><td class="agent">Assess</td></tr>
        <tr><td class="cap">cap:scml-lot-query</td><td>SCML.Lot (Silver)</td><td>BRZ.RC.LOT_TRACE · RECEIPT · VENDOR_MASTER · ITEM_MASTER</td><td>A · Read</td><td class="agent">Classify · Quantify</td></tr>
        <tr><td class="cap">cap:dim-item-lookup</td><td>DIM.Item (Silver)</td><td>BRZ.RC.ITEM_MASTER (SCD2)</td><td>A · Read</td><td class="agent">Classify</td></tr>
        <tr><td class="cap">cap:fsma-policy-resolver</td><td>FSMA-204 policy corpus</td><td>(versioned policy artifact · not Bronze)</td><td>A · Read</td><td class="agent">Classify</td></tr>
        <tr><td class="cap">cap:merml-price-lookup</td><td>MERML.Price (Silver)</td><td>BRZ.RC.PRICING + BRZ.RC.PROMO_MARKDOWN overlay</td><td>A · Read</td><td class="agent">Quantify</td></tr>
        <tr><td class="cap">cap:merml-elasticity-curve</td><td>MERML.Elasticity (Silver)</td><td>BRZ.RC.POS_TXN + PRICING + LOYALTY_HH (derived)</td><td>A · Read</td><td class="agent">Quantify</td></tr>
        <tr><td class="cap">cap:scml-lot-qty</td><td>SCML.Lot.qty</td><td>BRZ.RC.LOT_TRACE · RECEIPT · POS_TXN (depletions)</td><td>A · Read</td><td class="agent">Quantify</td></tr>
        <tr><td class="cap">cap:risk-score</td><td>(ML model)</td><td>(trained on BRZ.RC.REFRIG_TELEMETRY + FOOD_SAFETY history)</td><td>B · ML</td><td class="agent">Assess</td></tr>
        <tr><td class="cap">cap:disposition-classification</td><td>(ML model)</td><td>(trained on BRZ.RC.FOOD_SAFETY closeouts + FSMA corpus)</td><td>B · ML</td><td class="agent">Classify</td></tr>
        <tr><td class="cap">cap:teams-adaptive-card</td><td>MS Teams</td><td>—</td><td>C · Action</td><td class="agent">Approve</td></tr>
        <tr><td class="cap">cap:entra-identity-resolve</td><td>MS Entra</td><td>—</td><td>C · Action</td><td class="agent">Approve</td></tr>
        <tr><td class="cap">cap:pos-markdown-upload</td><td>POS systems</td><td>(writes to source-of-truth that BRZ.RC.PRICING reads)</td><td>C · Action</td><td class="agent">Act</td></tr>
        <tr><td class="cap">cap:vendor-portal-ticket</td><td>Vendor portals</td><td>(writes to source-of-truth that BRZ.RC.REFRIG_SERVICE reads)</td><td>C · Action</td><td class="agent">Act</td></tr>
        <tr><td class="cap">cap:fsma-regulatory-log</td><td>QMS (Trace One/etc)</td><td>(writes to source-of-truth that BRZ.RC.FOOD_SAFETY reads)</td><td>C · Action</td><td class="agent">Act</td></tr>
        <tr><td class="cap">cap:ledger-write</td><td>LEDGER (Delta + PG)</td><td>(audit-row store · separate from Bronze/Silver)</td><td>D · Obs</td><td class="agent">All 6 agents</td></tr>
        <tr><td class="cap">cap:fsma-packet-update</td><td>FSMA packet generator</td><td>(rolls up LEDGER closeout rows · async)</td><td>D · Obs</td><td class="agent">Evidence-Write</td></tr>
        <tr><td class="cap">cap:purview-lineage-emit</td><td>MS Purview</td><td>(emits Bronze→Silver→tool→agent→outcome lineage)</td><td>D · Obs</td><td class="agent">Evidence-Write</td></tr>
      </tbody>
    </table>
  </div>
</section>

<!-- ===== END-TO-END TIMELINE ===== -->
<section>
  <div class="section-head">
    <div class="section-eye">§ End-to-end · 90-second decision</div>
    <h2>What "90 seconds" means in wall-clock terms.</h2>
    <p class="lead">
      The 6-phase sequence executes in ≤ 90 seconds at W3 steady state. Below is the wall-clock
      timing budget for the reference excursion (Cincinnati Kroger #428 dairy case at 6:08 AM).
    </p>
  </div>

  <div class="tool-table">
    <table>
      <thead><tr>
        <th style="width:8%">Time</th>
        <th style="width:18%">Phase</th>
        <th style="width:30%">Action</th>
        <th style="width:14%">Latency</th>
        <th style="width:30%">Note</th>
      </tr></thead>
      <tbody>
        <tr><td>06:08:00</td><td>Trigger</td><td>Watermark fires · trace_id allocated</td><td>≤ 30s from sensor read</td><td>Eventstream → Orchestrator</td></tr>
        <tr><td>06:08:30</td><td>Phase 1 · Assess</td><td>5 read + ML calls · ExcursionProfile composed</td><td>~3s</td><td>5 sequential MCP calls · LEDGER capture row</td></tr>
        <tr><td>06:08:33</td><td>Phase 2 · Classify</td><td>4 read + 1 ML calls · DispositionTable per lot</td><td>~5s</td><td>conf ≥ 0.85 · no escalation · 6 LEDGER rows</td></tr>
        <tr><td>06:08:38</td><td>Phase 3 · Quantify</td><td>3 read calls · QuantifiedDisposition + markdown_pct</td><td>~3s</td><td>$7,920 total at risk · LEDGER capture row</td></tr>
        <tr><td>06:08:41</td><td>Phase 4 · Approve</td><td>Identity resolve · Adaptive Card delivered</td><td>~2s</td><td>Card on Marisol's phone</td></tr>
        <tr><td>06:09:30</td><td>Phase 4 · HITL</td><td>Marisol clicks APPROVE</td><td>~49s human</td><td>Within timeout · LEDGER HITL log row</td></tr>
        <tr><td>06:09:31</td><td>Phase 5 · Act</td><td>POS + vendor + FSMA log · parallel</td><td>~7s (max of 3)</td><td>All 3 succeed · LEDGER action log</td></tr>
        <tr><td>06:09:38</td><td>Phase 6 · Evidence-Write</td><td>Compose + LEDGER closeout + async siblings</td><td>~3s sync · packet/lineage async</td><td>14-field row sealed · scenario complete</td></tr>
        <tr style="background: rgba(227,182,87,0.10)"><td><b>06:09:41</b></td><td colspan="3"><b>End-to-end: 1 min 41 sec total · 49s human + 52s machine</b></td><td>Of which only 17 min today is the equivalent</td></tr>
      </tbody>
    </table>
  </div>
</section>

<footer>
  <div>APEX · RC-E2E-03 · Service Sequence Diagram · v1.0 · 27 Apr 2026</div>
  <div>Companion: APEX-RC-E2E-03-MCP-Server-Deep-Dive.docx · APEX-RC-E2E-03-09-SOR-ERD.html · APEX-RC-E2E-03-Six-Agents-Deep-Dive-and-Maturation.docx</div>
  <div>Deloitte · Microsoft Technology Solutions &amp; Practices</div>
</footer>

<script>
function toggleTheme() {
  const body = document.body, icon = document.getElementById('themeIcon'), label = document.getElementById('themeLabel');
  body.classList.toggle('light');
  if (body.classList.contains('light')) { icon.textContent = '☾'; label.textContent = 'Dark'; localStorage.setItem('apex-theme', 'light'); }
  else { icon.textContent = '☀'; label.textContent = 'Light'; localStorage.setItem('apex-theme', 'dark'); }
}
(function() {
  const saved = localStorage.getItem('apex-theme');
  if (saved === 'light') {
    document.body.classList.add('light');
    document.getElementById('themeIcon').textContent = '☾';
    document.getElementById('themeLabel').textContent = 'Dark';
  }
})();
</script>
</body></html>
"""

print("Building artifacts...")
build_mcp_deep_dive()

print("\nWriting Service Sequence Diagram html...")
target_html = OUT_DIR / "APEX-RC-E2E-03-Service-Sequence-Diagram.html"
target_html.write_text(SEQUENCE_HTML, encoding="utf-8")
print(f"  wrote {target_html.name}")

print("\nDone.")
