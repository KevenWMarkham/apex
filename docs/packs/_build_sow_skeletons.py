"""Build the 3 SOW skeleton DOCX templates (Lite / Standard / Enterprise).

Each one is a redlinable starting point for a real client engagement.
Find-and-replace tokens: [CLIENT_NAME], [BANNER], [START_DATE], [SOW_NUMBER], [LCSP_NAME].
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = r"C:\Stage\Clients\Industries\APEX\docs\reference"

# ---------- Styles ----------
def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'), 'clear')
    tc_pr.append(shd)

def add_page_break(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    br = OxmlElement('w:br'); br.set(qn('w:type'), 'page')
    r._r.append(br)

def style_default(doc):
    s = doc.styles['Normal']
    s.font.name = 'Calibri'
    s.font.size = Pt(11)
    s.paragraph_format.space_after = Pt(6)
    for hname, hsize in [('Heading 1', 18), ('Heading 2', 14), ('Heading 3', 12)]:
        st = doc.styles[hname]
        st.font.name = 'Calibri'
        st.font.size = Pt(hsize)
        st.font.color.rgb = RGBColor(0x0E,0x14,0x23)
        st.font.bold = True

def h1(doc, t):
    p = doc.add_paragraph(t, style='Heading 1')
    p.paragraph_format.space_before = Pt(18)
    return p

def h2(doc, t):
    p = doc.add_paragraph(t, style='Heading 2')
    p.paragraph_format.space_before = Pt(12)
    return p

def h3(doc, t):
    p = doc.add_paragraph(t, style='Heading 3')
    p.paragraph_format.space_before = Pt(8)
    return p

def para(doc, t, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.font.size = Pt(11)
    r.bold = bold
    r.italic = italic
    return p

def bullet(doc, t):
    p = doc.add_paragraph(t, style='List Bullet')
    return p

def callout(doc, t):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0,0)
    set_cell_shading(cell, "E8F5E9")
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(t)
    r.font.size = Pt(10)
    r.italic = True
    r.font.color.rgb = RGBColor(0x0E,0x14,0x23)
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return table

# ---------- Common sections (shared across all 3 tiers) ----------

def cover(doc, tier_name, price_range, timeline):
    h1(doc, f"STATEMENT OF WORK — APEX {tier_name}")
    para(doc, f"Engagement: {tier_name} on Deloitte's APEX delivery framework", italic=True)
    para(doc, "")
    t = doc.add_table(rows=6, cols=2)
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    t.columns[0].width = Inches(2.0); t.columns[1].width = Inches(4.5)
    rows = [
        ("SOW Number",        "[SOW_NUMBER]"),
        ("Client",            "[CLIENT_NAME]"),
        ("Engagement variant","[Variant A — B2B2C  /  Variant B — B2C]"),
        ("Estimated value",   price_range),
        ("Estimated duration", timeline),
        ("Effective date",    "[START_DATE]"),
    ]
    for i,(k,v) in enumerate(rows):
        c0 = t.cell(i,0); c1 = t.cell(i,1)
        c0.text = k; c1.text = v
        c0.paragraphs[0].runs[0].bold = True
        set_cell_shading(c0, "F5F5F5")
    para(doc, "")
    callout(doc,
        "This Statement of Work is issued under the Master Services Agreement between [CLIENT_NAME] and "
        "Deloitte Consulting LLP dated [MSA_DATE]. All terms of the MSA apply. This SOW supersedes any prior "
        "estimates, proposals, or correspondence between the parties relating to the services described herein.")
    para(doc, "")
    para(doc, "Prepared by: [LCSP_NAME], Deloitte's Microsoft Practice (DMTSP)", italic=True)
    para(doc, "Internal · Deloitte's Microsoft Technology & Services Practice", italic=True)

def section_background(doc, tier_name):
    h2(doc, "1. Background and engagement context")
    para(doc, ("[CLIENT_NAME] (\"Client\") has engaged Deloitte Consulting LLP (\"Deloitte\") to deliver a scoped "
               f"agentic-AI engagement on the APEX framework. This SOW describes the {tier_name} engagement, "
               "the first of a potential multi-wave program of work that may grow into a full Pack rollout and "
               "ongoing Operate services."))
    para(doc, ("Deloitte's APEX framework is a cloud-neutral delivery framework comprising: a 14-interface "
               "Cloud Profile contract; a 6-Agent Fleet pattern; a 14-field WORM hash-chain audit ledger "
               "(\"LEDGER\"); Industry Solution Packs (10-asset bundles) including the Customer-Focused "
               "Merchandise Pack (\"CFMP\"); and standard Service Envelope tiers under which engagements are "
               "scoped and priced."))
    para(doc, "This engagement runs on the APEX-M (Microsoft) cloud profile.")

def section_independence(doc):
    h2(doc, "Independence and funding")
    callout(doc,
        "Deloitte serves as Microsoft Corporation's independent auditor. Per SEC and PCAOB rules, Deloitte "
        "does not accept Enterprise Customer Investment Funds (ECIF) or similar payments directly from "
        "Microsoft. Any Microsoft-sourced funding contemplated under this engagement flows via (i) Azure "
        "Marketplace burndown against [CLIENT_NAME]'s existing Microsoft commitments, or (ii) a "
        "non-competing Systems Integrator (SI) partner under separate Teaming Agreement. Deloitte's "
        "engagement fee is paid by [CLIENT_NAME].")
    para(doc, "")
    para(doc, ("[CLIENT_NAME] acknowledges that Deloitte's role under this SOW is independent of Microsoft. "
               "Deloitte makes no representations on behalf of Microsoft and shall not be characterized as "
               "Microsoft's partner, agent, alliance member, or joint-commercial party. Microsoft is referred "
               "to throughout this engagement solely as the 'cloud profile' or 'platform'."))

def section_assumptions(doc, extra_assumptions=None):
    h2(doc, "Assumptions and dependencies")
    base_assumptions = [
        "Client provides timely access to subject-matter experts, business stakeholders, and decision-makers as agreed in the engagement governance plan.",
        "Client provides access to systems, data, and credentials required for in-scope work; production-system access is requested at least two business days in advance.",
        "Client provides one or more Microsoft Azure subscriptions sufficient for the APEX-M cloud profile, including but not limited to: Foundry, Azure OpenAI, OneLake / Fabric, Entra ID, Purview, Key Vault, and Container Apps capacity.",
        "Client owns and operates all source systems referenced by the deployed Pack manifests; Deloitte develops adapters but does not assume operational responsibility for source systems.",
        "Client's existing identity infrastructure (Entra) is operational and capable of supporting the persona model deployed by the Pack.",
        "Client is responsible for legal review of the deployed Constitution.yaml (privacy hard rules), the Adaptive Card library, and customer-facing surfaces prior to production go-live.",
        "Client is responsible for any data-protection-impact assessments (DPIA), regulator notifications, or consumer-disclosure obligations triggered by the deployed services.",
        "Deloitte's services are advisory and engineering; Deloitte does not assume any role of 'data controller' under GDPR/CCPA or analogous regulations.",
        "Pricing in this SOW does not include third-party software, Azure consumption, Sonos hardware, edge devices, or other infrastructure costs.",
        "Schedule shown in §5 assumes no more than two-week aggregate delay from Client-side dependencies; greater delays require change control.",
    ]
    for a in base_assumptions:
        bullet(doc, a)
    if extra_assumptions:
        for a in extra_assumptions:
            bullet(doc, a)

def section_change_control(doc):
    h2(doc, "Change control")
    para(doc, ("Any change to the scope, schedule, deliverables, fees, or assumptions of this SOW shall be "
               "documented in a Change Request signed by both parties prior to commencement of the affected "
               "work. Verbal authorizations are not binding. Change Requests may be priced as either a fixed "
               "addition to this SOW or as a time-and-materials engagement extension."))

def section_ip(doc):
    h2(doc, "Intellectual property")
    para(doc, ("APEX, including the APEX framework code, the 14-interface Cloud Profile contract specification, "
               "the 6-Agent Fleet pattern, the LEDGER schema, the Industry Solution Pack catalog (including "
               "CFMP and any CHC overlays), the apex-replay tooling, and all related methodology and pre-existing "
               "Deloitte intellectual property (collectively the \"APEX IP\") remain the sole property of "
               "Deloitte and its affiliates. Client receives a non-exclusive, royalty-free, perpetual license "
               "to use the deployed Pack manifests, Constitution overlay, and configuration artifacts solely "
               "for Client's internal business purposes."))
    para(doc, ("Client overlays — including Client's CAD-derived Azure Maps Creator datasets, Client's planogram "
               "data, Client's identity bindings, Client's threshold tunings, and any client-specific scenario "
               "manifests authored under this SOW — are owned by Client."))
    para(doc, ("Pre-existing IP retains its prior ownership; nothing in this SOW transfers ownership of pre-existing "
               "IP between the parties. Newly developed work product is owned per the IP allocation above; ambiguity "
               "in any single artifact's classification is resolved in Deloitte's favor for general framework "
               "components and in Client's favor for Client-specific configuration."))

def section_acceptance(doc, criteria):
    h2(doc, "Acceptance criteria")
    para(doc, "The deliverables under this SOW are accepted when the following are demonstrated:")
    for c in criteria:
        bullet(doc, c)
    para(doc, ("Client has ten (10) business days from delivery of each acceptance demonstration to provide "
               "written acceptance or a written description of any deficiency. Absent written response within "
               "this window, deliverables are deemed accepted."))

def section_signatures(doc):
    h2(doc, "Signatures")
    t = doc.add_table(rows=4, cols=2)
    for i in range(4):
        for j in range(2):
            t.cell(i,j).width = Inches(3.0)
    t.cell(0,0).text = "For [CLIENT_NAME]:"
    t.cell(0,1).text = "For Deloitte Consulting LLP:"
    t.cell(0,0).paragraphs[0].runs[0].bold = True
    t.cell(0,1).paragraphs[0].runs[0].bold = True
    t.cell(1,0).text = "By: _________________________"
    t.cell(1,1).text = "By: _________________________"
    t.cell(2,0).text = "Name: [CLIENT_SIGNATORY]"
    t.cell(2,1).text = "Name: [LCSP_NAME]"
    t.cell(3,0).text = "Title: ______________________"
    t.cell(3,1).text = "Title: Vice President"

def appendix_envelope(doc, current_tier):
    add_page_break(doc)
    h1(doc, "Appendix A — Service Envelope mapping")
    para(doc, "This engagement sits within the APEX Service Envelope tier ladder:")
    t = doc.add_table(rows=6, cols=4)
    t.style = 'Light Grid Accent 1'
    hdr = t.rows[0].cells
    hdr[0].text = "Tier"; hdr[1].text = "Description"; hdr[2].text = "Price band"; hdr[3].text = "This SOW"
    tier_rows = [
        ("T1 Foundation",     "VV runtime · LEDGER · MCP · Entra/Purview wiring", "$400-700K", "Foundation"),
        ("T2 Pack Lite",      "3 scenarios live on client data",                  "$150-250K", "Lite"),
        ("T2 Pack Standard",  "Sub-domain in production · 3-5 scenarios",         "$500K-$1.5M", "Standard"),
        ("T2 Pack Enterprise","Full pack + multi-store · Operate-ready",          "$1.5-3.5M",  "Enterprise"),
        ("T5 Operate",        "24×7 monitor · tuning · pack version uptake",      "$25-110K/mo", "Operate"),
    ]
    for i,(a,b,c,d) in enumerate(tier_rows, start=1):
        row = t.rows[i].cells
        row[0].text = a; row[1].text = b; row[2].text = c
        row[3].text = "✓" if d == current_tier else ""
        if d == current_tier:
            for cc in row: set_cell_shading(cc, "E8F5E9")

def appendix_funding(doc):
    add_page_break(doc)
    h1(doc, "Appendix B — Funding paths and independence")
    para(doc, "Engagement fees may be sourced from any of the following independence-safe funding vehicles:")
    vehicles = [
        ("Client direct",            "Standard Deloitte-on-Deloitte-paper engagement; no special funding routing required."),
        ("BVA (client-funded)",      "Workshop fee absorbed into Client's marketing or innovation budget; produces the scenario shortlist and SOW skeleton that initiates Lite tier."),
        ("DCIF (Deloitte co-invest)","Deloitte commits its own balance sheet against a defined client commitment milestone; reduces Client's upfront cost."),
        ("Azure Marketplace burndown","Client's existing Microsoft Azure commitment is burned down against the engagement fee via an Azure Marketplace transaction; Microsoft pays the Marketplace, Marketplace pays Deloitte. Deloitte does not receive ECIF directly from Microsoft."),
        ("SI Teaming (POC funding)", "Non-competing SI partner under separate Teaming Agreement receives Microsoft POC funding; Deloitte delivers under SI sub-agreement. Used selectively for early-stage proof-of-concept engagements where Client's Azure commitment is insufficient."),
        ("Time & Materials (T&M)",   "For custom build (T4) tiers or scope extensions beyond Lite/Standard/Enterprise."),
    ]
    for k,v in vehicles:
        p = doc.add_paragraph()
        r = p.add_run(f"{k}.  "); r.bold = True
        p.add_run(v)

def appendix_pack_assets(doc, tier_name):
    add_page_break(doc)
    h1(doc, "Appendix C — Pack 10-asset bundle inventory")
    para(doc, ("Every APEX Industry Solution Pack ships the same 10-asset bundle. Asset coverage varies by "
               "engagement tier. Items checked ✓ are delivered under this SOW; items marked – are deferred to a "
               "higher tier."))
    coverage = {
        "Lite":       {1:True, 2:True, 3:True, 4:True, 5:True, 6:True, 7:True, 8:True, 9:"partial", 10:"partial"},
        "Standard":   {1:True, 2:True, 3:True, 4:True, 5:True, 6:True, 7:True, 8:True, 9:True, 10:True},
        "Enterprise": {1:True, 2:True, 3:True, 4:True, 5:True, 6:True, 7:True, 8:True, 9:True, 10:True},
    }
    assets = [
        (1,  "VV Manifests",         "Declarative federated views w/ thresholds + resolution routes."),
        (2,  "Scenario Manifests",   "Named business-outcome chains composed from VVs."),
        (3,  "Source Adapters",      "POS · Loyalty CRM · Planogram · Beacon vendor · Vision device · Phone app."),
        (4,  "Adaptive Cards",       "Customer phone + kiosk JSON; operator Teams cards."),
        (5,  "Persona Map",          "personas.yaml including the customer persona type."),
        (6,  "Demo Data",            "Synthetic catalog + planogram + simulated shoppers."),
        (7,  "BVA Worksheet",        "Pack-level ROI calculator (separate xlsx)."),
        (8,  "Sample SOW",           "Lite / Standard / Enterprise templates (this document)."),
        (9,  "Acceptance Tests",     "60-80+ tests proving Pack health in CI."),
        (10, "Runbook + Training",   "Operate procedures and analyst training materials."),
    ]
    t = doc.add_table(rows=len(assets)+1, cols=4)
    t.style = 'Light Grid Accent 1'
    hdr = t.rows[0].cells
    hdr[0].text = "#"; hdr[1].text = "Asset"; hdr[2].text = "Description"; hdr[3].text = "This SOW"
    cov = coverage[tier_name]
    for i,(n,name_,desc) in enumerate(assets, start=1):
        row = t.rows[i].cells
        row[0].text = str(n); row[1].text = name_; row[2].text = desc
        c = cov.get(n)
        row[3].text = "✓" if c is True else ("partial" if c == "partial" else "—")

def appendix_demo(doc):
    add_page_break(doc)
    h1(doc, "Appendix D — Working demo reference")
    para(doc, ("The APEX framework demonstration referenced in §1 is currently operational at the URL below. "
               "[CLIENT_NAME] personnel may access the demo as part of pre-engagement validation."))
    p = doc.add_paragraph()
    r = p.add_run("Demo URL: ")
    r.bold = True
    p.add_run("https://ca-visionkit-portal.gentlestone-9b49b099.eastus2.azurecontainerapps.io")
    para(doc, "")
    para(doc, "The deployed demo covers:")
    for b in [
        "800-product pgvector catalog with semantic search",
        "Microsoft Agent Framework SDK 1.6.0 GA orchestration with gpt-5-mini",
        "Azure Speech (Ava Multilingual) for text-to-speech and speech-to-text",
        "Azure Maps Creator (Indoor Maps + Wayfinding REST API) — local-fallback path until per-tenant Creator dataset provisioned",
        "Customer-facing CFMP journey: CHOOSE → SELECT → BUY → SERVICES",
        "Live LEDGER tail on the architecture page showing hash-chained AuditRow entries with HITL status",
        "APEX Tour flyout with framework guided walkthrough",
    ]:
        bullet(doc, b)

# ============================================================
# TIER 1: PACK LITE
# ============================================================

def build_lite():
    doc = Document()
    style_default(doc)

    cover(doc, "CFMP Pack Lite (T2A)", "$150,000 – $250,000 USD", "4 to 6 weeks")
    section_background(doc, "Pack Lite")

    h2(doc, "2. Scope of services")
    para(doc, ("Deloitte will deliver the APEX CFMP Pack Lite engagement, comprising three live scenarios on "
               "Client's stores or — if Variant B (CHC) is elected — on a pilot household cohort. The engagement "
               "produces a working production-grade scenario set on Client data, a Pack-Lite SOW for follow-on "
               "Standard tier, and an operate-readiness assessment."))

    h3(doc, "2.1 Scenarios in scope")
    para(doc, "Three scenarios from the CFMP / CHC catalog, to be selected during BVA workshop. Typical Lite shortlist:")
    for s in [
        "rc-sampling-table-engagement (CHOOSE phase, in-store)",
        "rc-cart-dwell-abandonment-rescue (BUY phase, in-store)",
        "cfmp-wayfinding-walk-to-product (SELECT phase, in-store, Azure Maps Creator)",
        "— or CHC equivalent triad: chc-pantry-inventory-inference, chc-low-stock-voice-prompt, chc-replenishment-cart-build (NEED phase, in-home)",
    ]:
        bullet(doc, s)
    para(doc, "Final selection is documented in the post-BVA addendum to this SOW.")

    h3(doc, "2.2 Deliverables")
    for d in [
        "Three scenario manifests (`scenarios/*.yaml`) deployed to the APEX-M runtime.",
        "Six Virtual View manifests (`views/*.yaml`) with thresholds and resolution routes.",
        "Three Adaptive Card templates (`cards/*.json`) — customer-facing and operator-facing variants.",
        "Persona map (`personas.yaml`) with Client's identity binding (Entra group OIDs for store associates; loyalty-ID binding for customer persona).",
        "Source adapter scaffolding for Client's POS, planogram, and loyalty systems (one production-grade adapter; two stubs for follow-on Standard tier).",
        "Constitution overlay (`constitution.yaml`) with six hard rules reviewed and signed off by Client's privacy team.",
        "60+ acceptance tests in the Pack CI suite, passing against Client data.",
        "Pack-level BVA worksheet (xlsx) with ROI calculator populated from Client baseline metrics.",
        "Pack-Lite Standard SOW skeleton (this document, refined post-engagement).",
        "Live demonstration of all three scenarios on Client data, signed off in the acceptance demonstration meeting.",
    ]:
        bullet(doc, d)

    section_acceptance(doc, [
        "All three scenarios fire end-to-end on Client data without manual intervention in a live demonstration witnessed by Client's CMO/CX-VP and DPO.",
        "The Pack CI suite (60+ tests) passes at ≥98% green against Client data.",
        "The LEDGER produces a hash-chained AuditRow for every customer-facing decision, with consent_hash captured at decision moment, verified via apex-replay reproduction.",
        "Pack-level BVA worksheet completed with Client CFO sign-off on baseline assumptions.",
        "Operate-readiness assessment delivered with go/no-go recommendation for Standard tier.",
    ])

    h2(doc, "3. Engagement approach")
    h3(doc, "3.1 Phases")
    phases_lite = [
        ("Phase 1 — BVA Workshop (Week 0, 4 hours on-site)", "Scenario shortlisting, baseline-metric capture, ROI calculator population, SOW addendum."),
        ("Phase 2 — Manifest Authoring (Weeks 1-2)",         "Deloitte authors VV, scenario, persona, and card manifests; Client provides data samples and source-system access."),
        ("Phase 3 — Adapter Build (Weeks 2-4)",              "Production-grade adapter for primary source system; stubs for secondaries."),
        ("Phase 4 — Constitution Review (Week 4)",           "Constitution.yaml drafted; Client privacy team reviews and signs."),
        ("Phase 5 — Integration + Test (Weeks 4-5)",         "End-to-end scenarios deployed to APEX-M runtime; CI suite turned green."),
        ("Phase 6 — Acceptance Demo + Handoff (Week 6)",     "Live demonstration; BVA worksheet finalized; Standard SOW skeleton handed to LCSP."),
    ]
    for k,v in phases_lite:
        p = doc.add_paragraph()
        r = p.add_run(f"{k}.  "); r.bold = True
        p.add_run(v)

    h2(doc, "4. Engagement team")
    para(doc, "Deloitte delivery team (typical Lite composition):")
    for r in [
        "Engagement Manager (50%) — overall delivery accountability and Client communications.",
        "APEX Solution Architect (75%) — framework deployment and integration design.",
        "Pack Manifest Engineer (100%) — VV and scenario manifest authoring.",
        "Adapter Engineer (75%) — source-system integration.",
        "Privacy Engineer (25%) — Constitution overlay and consent-flow review.",
        "DataOps Engineer (50%) — Pack CI suite and acceptance tests.",
    ]:
        bullet(doc, r)

    h2(doc, "5. Schedule")
    para(doc, "Total elapsed: 4-6 weeks from MSA effective date and SOW countersignature.")

    h2(doc, "6. Fees and payment")
    para(doc, ("Total engagement fee: USD 150,000 – 250,000, fixed fee, to be finalized in the post-BVA "
               "addendum based on final scenario selection and source-system complexity."))
    para(doc, "Payment schedule:")
    for s in [
        "30% at SOW countersignature.",
        "40% upon completion of Phase 4 (Constitution review).",
        "30% upon Acceptance per §2.2.",
    ]:
        bullet(doc, s)
    para(doc, ("Pricing assumes BVA workshop is bundled into this engagement. If BVA was previously delivered "
               "as a standalone engagement, its fee is credited against this SOW."))

    section_assumptions(doc, extra_assumptions=[
        "Lite engagement assumes Client provides one primary source-system adapter target (POS, planogram, OR loyalty); two additional source systems are stubbed for Standard-tier follow-on.",
        "Lite scope is limited to three scenarios; additional scenarios require Standard-tier scoping.",
    ])
    section_change_control(doc)
    section_ip(doc)
    section_independence(doc)
    section_signatures(doc)
    appendix_envelope(doc, "Lite")
    appendix_funding(doc)
    appendix_pack_assets(doc, "Lite")
    appendix_demo(doc)

    out = OUT_DIR + r"\SOW-APEX-CFMP-Pack-Lite.docx"
    doc.save(out)
    print("saved", out)

# ============================================================
# TIER 2: PACK STANDARD
# ============================================================

def build_standard():
    doc = Document()
    style_default(doc)

    cover(doc, "CFMP Pack Standard (T2B)", "$750,000 – $1,500,000 USD", "12 to 16 weeks")
    section_background(doc, "Pack Standard")

    h2(doc, "2. Scope of services")
    para(doc, ("Deloitte will deliver the APEX CFMP Pack Standard engagement, comprising one complete CFMP "
               "sub-domain in production plus 5-7 additional scenarios beyond the Lite-tier shortlist. The "
               "engagement targets production deployment at a single banner or business unit, with operate-"
               "readiness handoff at conclusion."))

    h3(doc, "2.1 Scenarios in scope (10 total)")
    para(doc, "Three carryover from Lite (if Standard follows Lite) plus seven from the catalog. Typical Standard shortlist:")
    for s in [
        "Three Lite-tier scenarios continuing in production",
        "rc-on-shelf-availability-oos-reduction (SELECT)",
        "rc-shelf-gap-realtime-restock-dispatch (SELECT)",
        "rc-customer-wait-time-prediction-at-checkout (BUY)",
        "rc-bopis-pickup-counter-load (BUY)",
        "rc-loyalty-churn-prediction-winback (SERVICES) — Featured chain",
        "rc-product-complaint-triage (SERVICES)",
        "chc-substitution-acceptance (SELECT) — CHC option",
    ]:
        bullet(doc, s)

    h3(doc, "2.2 Deliverables (incremental beyond Lite)")
    for d in [
        "Seven additional scenario manifests beyond the Lite-tier shortlist.",
        "Twelve to fifteen additional Virtual View manifests covering the chosen sub-domain.",
        "Adaptive Card library for all 10 scenarios — both customer-facing and operator-facing variants.",
        "Three production-grade source adapters (vs Lite's one); two additional stubs for Enterprise follow-on.",
        "Operate-tier runbook (`runbooks/`) covering monitoring, threshold tuning, escalation, and Pack version uptake.",
        "Training deck and analyst materials for Client's operations team (handover for Operate tier).",
        "Pack acceptance test suite expanded to 100+ tests with multi-scenario interaction coverage.",
        "Live demonstration of all 10 scenarios; in-store production rollout at one banner with phased cutover plan.",
        "Standard-to-Enterprise SOW skeleton with phased rollout plan to additional banners.",
    ]:
        bullet(doc, d)

    section_acceptance(doc, [
        "All 10 scenarios fire end-to-end on Client data in production at the pilot banner without manual intervention.",
        "Pack CI suite (100+ tests) at ≥98% green against Client data.",
        "LEDGER hash chain verifiable end-to-end via apex-replay across all scenarios.",
        "Client's operations team completes Operate-tier training and signs off on runbook ownership.",
        "DPO sign-off on Constitution.yaml (refreshed for expanded scenario set) and privacy panel UI.",
        "BVA worksheet refreshed with post-deployment actuals; Pack-Standard ROI realized at minimum 1.5× engagement cost in first 12 months.",
    ])

    h2(doc, "3. Engagement approach")
    h3(doc, "3.1 Phases")
    phases_std = [
        ("Phase 1 — Mobilization (Weeks 1-2)",              "Engagement kickoff; carryover from Lite; refined scope addendum."),
        ("Phase 2 — Scenario authoring (Weeks 2-6)",        "Seven additional scenario manifests + supporting VVs."),
        ("Phase 3 — Adapter build (Weeks 3-8)",             "Three production adapters; integration with Client SOR."),
        ("Phase 4 — Operate-runbook (Weeks 6-10)",          "Runbook drafted; operations team training delivered."),
        ("Phase 5 — Integration + test (Weeks 8-12)",       "End-to-end production at pilot banner; CI suite expanded."),
        ("Phase 6 — Production rollout (Weeks 12-14)",      "Phased cutover at pilot banner; LEDGER verification."),
        ("Phase 7 — Acceptance + handoff (Weeks 14-16)",    "Acceptance demonstration; Standard→Enterprise SOW skeleton; Operate handoff."),
    ]
    for k,v in phases_std:
        p = doc.add_paragraph()
        r = p.add_run(f"{k}.  "); r.bold = True
        p.add_run(v)

    h2(doc, "4. Engagement team")
    para(doc, "Deloitte delivery team (typical Standard composition):")
    for r in [
        "Engagement Manager (75%)",
        "APEX Solution Architect (100%)",
        "Pack Manifest Engineer × 2 (100% each)",
        "Adapter Engineer × 2 (100% each)",
        "Privacy Engineer (50%)",
        "DataOps Engineer (75%)",
        "Operate Lead (50%) — embedded for runbook authoring and handoff",
    ]:
        bullet(doc, r)

    h2(doc, "5. Schedule")
    para(doc, "Total elapsed: 12-16 weeks. Phasing per §3.1. Phase 6 production rollout may be sequenced over multiple stores within the pilot banner.")

    h2(doc, "6. Fees and payment")
    para(doc, ("Total engagement fee: USD 750,000 – 1,500,000, fixed fee, finalized at mobilization based on "
               "final scenario count, adapter complexity, and rollout footprint."))
    para(doc, "Payment schedule:")
    for s in [
        "20% at SOW countersignature",
        "25% upon completion of Phase 3 (Adapter build)",
        "25% upon completion of Phase 5 (Integration + test)",
        "20% upon completion of Phase 6 (Production rollout)",
        "10% upon Acceptance per §2.2",
    ]:
        bullet(doc, s)

    section_assumptions(doc, extra_assumptions=[
        "Standard tier assumes Client commits to a single pilot banner / business unit for production rollout; additional banners require Enterprise-tier scoping.",
        "Standard scope is limited to 10 scenarios drawn from the CFMP / CHC catalog; bespoke scenarios beyond the catalog require T4 Custom Build scoping.",
        "Operate-tier subscription begins no earlier than Week 14 of this engagement; Operate fees billed separately under a Master Services Agreement amendment.",
    ])
    section_change_control(doc)
    section_ip(doc)
    section_independence(doc)
    section_signatures(doc)
    appendix_envelope(doc, "Standard")
    appendix_funding(doc)
    appendix_pack_assets(doc, "Standard")
    appendix_demo(doc)

    out = OUT_DIR + r"\SOW-APEX-CFMP-Pack-Standard.docx"
    doc.save(out)
    print("saved", out)

# ============================================================
# TIER 3: PACK ENTERPRISE
# ============================================================

def build_enterprise():
    doc = Document()
    style_default(doc)

    cover(doc, "CFMP Pack Enterprise (T2C)", "$1,500,000 – $3,500,000 USD", "6 to 9 months")
    section_background(doc, "Pack Enterprise")

    h2(doc, "2. Scope of services")
    para(doc, ("Deloitte will deliver the APEX CFMP Pack Enterprise engagement, comprising the full CFMP "
               "catalog (all 18 scenarios from CFMP v0.2 §3, optionally extended with CHC overlay scenarios) "
               "in production across Client's multi-banner footprint, with full Operate-tier handoff and "
               "optional cross-pack scenario chains (Variant A only)."))

    h3(doc, "2.1 Scenarios in scope (all 18 CFMP + optional CHC overlay)")
    for s in [
        "All 5 CHOOSE scenarios: personalized offers · review summary · search relevance · in-store advertising · sampling table",
        "All 6 SELECT scenarios: wayfinding · OSA · shelf gap · aisle engagement · end-cap ROI · SPL audit",
        "All 4 BUY scenarios: queue prediction · cart-dwell · SCO shrink · BOPIS pickup",
        "All 3 SERVICES scenarios: loyalty churn winback (Featured) · tier migration · complaint triage",
        "CHC overlay (Variant B): 12 in-home scenarios across NEED → SERVICES",
        "Cross-pack scenario chain (Variant A only): chc-replenishment-cycle bridging TCP and CFMP per CHC v0.2 §8",
    ]:
        bullet(doc, s)

    h3(doc, "2.2 Deliverables (Enterprise full set)")
    for d in [
        "All 18 CFMP scenario manifests (plus optional 12 CHC overlay scenarios) in production.",
        "Full Virtual View catalog (25-30 manifests covering MERML + CXML + CFMP.StoreMap).",
        "Complete Adaptive Card library (60+ card variants across customer phone, kiosk, operator Teams).",
        "Five production-grade source adapters (POS · planogram · loyalty CRM · beacon · vision device).",
        "Multi-banner rollout: 5-50 banners with phased cutover plan, change-management, and per-banner threshold tuning.",
        "Cross-channel attribution scenario (chc-loyalty-fusion-attribution) lighting up in-store + in-home + online into one LEDGER.",
        "Operate-tier runbook with multi-tenant operations, FinOps metering, threshold-tuning loop, and DataOps for upstream sources.",
        "Pack acceptance test suite at 150+ tests including cross-scenario interaction, multi-banner correctness, and apex-replay forensic correctness.",
        "Training program for Client's operations team: classroom + hands-on + certification.",
        "Live demonstration of full Pack at three banners with cross-channel LEDGER trace.",
        "Operate-tier handoff with first 90 days of co-operation included.",
    ]:
        bullet(doc, d)

    section_acceptance(doc, [
        "All 18 CFMP scenarios in production at minimum 3 banners; LEDGER hash chain verifiable via apex-replay.",
        "Pack CI suite (150+ tests) at ≥99% green across all banners.",
        "Multi-banner threshold tuning loop in operation: Telemetry VVs observe band distribution and Pack maintainer reviews monthly.",
        "Operations team completes certification program and assumes Pack ownership in the Operate-tier subscription.",
        "Pack-Enterprise ROI realized at minimum 2.5× engagement cost in first 12 months at deployed banners.",
        "Cross-channel attribution (chc-loyalty-fusion-attribution) verifiable: a single household interaction is traceable across in-store + in-home + online in one LEDGER query.",
    ])

    h2(doc, "3. Engagement approach")
    h3(doc, "3.1 Phases")
    phases_ent = [
        ("Phase 1 — Mobilization (Month 1)",                       "Engagement kickoff; team assembly; rollout sequencing across banners."),
        ("Phase 2 — Full Pack authoring (Months 1-3)",             "All 18 scenarios + supporting VVs + complete card library."),
        ("Phase 3 — Adapter build (Months 1-4)",                    "Five production-grade adapters; integration with Client's full SOR footprint."),
        ("Phase 4 — Multi-banner pilot rollout (Months 4-5)",       "Three banners live; full Pack running; threshold tuning."),
        ("Phase 5 — Cross-channel attribution build (Months 4-6)",  "Loyalty-fusion scenario; in-store + in-home + online LEDGER unification."),
        ("Phase 6 — Production rollout to remaining banners (Months 6-8)","Phased cutover banner-by-banner; per-banner threshold tuning and acceptance."),
        ("Phase 7 — Operate handoff (Month 9)",                    "Operations team certification; Operate-tier subscription begins; 90-day co-operation."),
    ]
    for k,v in phases_ent:
        p = doc.add_paragraph()
        r = p.add_run(f"{k}.  "); r.bold = True
        p.add_run(v)

    h2(doc, "4. Engagement team")
    para(doc, "Deloitte delivery team (typical Enterprise composition):")
    for r in [
        "Engagement Partner (25%) — executive sponsorship and Client C-suite engagement",
        "Engagement Manager (100%)",
        "APEX Solution Architect (100%)",
        "Pack Manifest Engineer × 3 (100% each)",
        "Adapter Engineer × 3 (100% each)",
        "Privacy Engineer (75%)",
        "DataOps Engineer (100%)",
        "Operate Lead (75%) — embedded throughout; transitions to Operate tier",
        "Change Management Lead (75%) — Client-facing rollout and training",
        "Independence Compliance Reviewer (10%) — Deloitte-internal Independence checkpoint at each phase gate",
    ]:
        bullet(doc, r)

    h2(doc, "5. Schedule")
    para(doc, "Total elapsed: 6-9 months. Phasing per §3.1 with monthly steering committee reviews and quarterly Engagement Partner check-ins.")

    h2(doc, "6. Fees and payment")
    para(doc, ("Total engagement fee: USD 1,500,000 – 3,500,000, fixed fee, finalized at mobilization based on "
               "scenario count (CFMP only vs CFMP+CHC), adapter count, banner rollout footprint, and selected "
               "cross-channel attribution scope."))
    para(doc, "Payment schedule (milestone-driven):")
    for s in [
        "15% at SOW countersignature",
        "15% at end of Phase 2 (Full Pack authoring)",
        "15% at end of Phase 3 (Adapter build)",
        "20% at end of Phase 4 (Multi-banner pilot live)",
        "15% at end of Phase 5 (Cross-channel attribution live)",
        "15% at end of Phase 6 (Production rollout to all banners)",
        "5% upon Acceptance per §2.2 (Operate handoff complete)",
    ]:
        bullet(doc, s)
    para(doc, "")
    para(doc, ("Operate-tier subscription pricing is separate from this SOW and is governed by a follow-on "
               "MSA amendment to be signed at end of Phase 6. Typical Operate fee for Enterprise deployment: "
               "USD 60,000 – 110,000 per month."))

    section_assumptions(doc, extra_assumptions=[
        "Enterprise tier assumes Client commits to a full-banner rollout footprint with named banners specified in the mobilization addendum.",
        "Enterprise tier includes the full CFMP catalog (18 scenarios); CHC overlay scenarios are optional and priced as an addendum to this SOW.",
        "Cross-pack scenario chains (Variant A) require a separate engagement with the Telco partner; this SOW covers only the CFMP side.",
        "Multi-banner production rollout assumes that each banner has functional store-side IT infrastructure (network, edge compute, vision hardware where applicable) ready prior to that banner's Phase-6 cutover slot.",
        "Operate-tier handoff assumes Client's operations team is staffed and trained per §3.1 Phase 7; if Client elects to outsource Operate to Deloitte permanently, that arrangement is governed by a separate Operate Master Services Agreement.",
    ])
    section_change_control(doc)
    section_ip(doc)
    section_independence(doc)
    section_signatures(doc)
    appendix_envelope(doc, "Enterprise")
    appendix_funding(doc)
    appendix_pack_assets(doc, "Enterprise")
    appendix_demo(doc)

    # Extra appendix for Enterprise — multi-banner sequencing
    add_page_break(doc)
    h1(doc, "Appendix E — Multi-banner rollout sequencing")
    para(doc, ("Enterprise-tier production rollout follows a phased banner-by-banner sequence designed to "
               "validate Pack stability at small banner footprints before expanding to high-volume banners. "
               "The default sequence is documented below; a Client-specific sequence is finalized at "
               "mobilization."))
    t = doc.add_table(rows=6, cols=4)
    t.style = 'Light Grid Accent 1'
    hdr = t.rows[0].cells
    hdr[0].text = "Wave"; hdr[1].text = "Banners"; hdr[2].text = "Timing"; hdr[3].text = "Gate criteria"
    waves = [
        ("Pilot",     "1 banner (smallest footprint by store count)", "Phase 4",      "All 18 scenarios live · Pack CI green"),
        ("Wave 1",    "2 additional banners (small + medium)",         "Phase 6 — M1", "Pilot threshold tuning stable for 14 days"),
        ("Wave 2",    "Next 5-10 banners",                              "Phase 6 — M2-3","Wave 1 banners stable · Operate runbook validated"),
        ("Wave 3",    "Remaining banners",                              "Phase 6 — M4",  "Wave 2 banners stable · Cross-channel attribution live"),
        ("Operate",   "All deployed banners",                           "Phase 7+",      "Multi-banner threshold tuning loop active · DataOps mature"),
    ]
    for i,(a,b,c,d) in enumerate(waves, start=1):
        row = t.rows[i].cells
        row[0].text = a; row[1].text = b; row[2].text = c; row[3].text = d

    out = OUT_DIR + r"\SOW-APEX-CFMP-Pack-Enterprise.docx"
    doc.save(out)
    print("saved", out)

if __name__ == "__main__":
    build_lite()
    build_standard()
    build_enterprise()
