"""Build the APEX RC-E2E-03 Six-Agents Deep Dive + Maturation document.

Deep dive on the six agents that compose the service:
  1. Assess          - characterize the excursion
  2. Classify        - per-SKU FSMA disposition
  3. Quantify        - $ impact computation
  4. Approve         - HITL gate via Adaptive Card
  5. Act             - orchestrate downstream actions
  6. Evidence-Write  - composes 14-field audit row + FSMA evidence packet (NEW)

Plus a multi-section narrative on how these agents evolve into context-aware
superagents through Wave 1 → 4+, including cross-agent fusion patterns.

Output: Services/RC-E2E-03_Assortment-and-Pricing/Tier2-Build/
        APEX-RC-E2E-03-Six-Agents-Deep-Dive-and-Maturation.docx
"""
from __future__ import annotations
import io, sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
OUT_DIR = Path(r"C:\Stage\Clients\Industries\Consumer\Retail\Kroger\02_projects\FY27_Pipeline\assortment-pricing-agentic\deliverables\Services\RC-E2E-03_Assortment-and-Pricing\Tier2-Build")
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT = OUT_DIR / "APEX-RC-E2E-03-Six-Agents-Deep-Dive-and-Maturation.docx"

INK="0B0B0F"; INK2="4A4A55"; MUTE="8a8a90"; GOLD="B8892E"; TEAL="2A7F73"
VIO="6648B0"; CRIM="A22A39"; BAND="F3F3F5"; HEAD="0B0B0F"; LINE="D4D4D7"

# Google fonts (matching the established system)
SERIF = "Fraunces"
SANS  = "IBM Plex Sans"
MONO  = "JetBrains Mono"

def cell_bg(cell, c):
    tcPr = cell._tc.get_or_add_tcPr()
    s = OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto")
    s.set(qn("w:fill"),c.lstrip("#")); tcPr.append(s)
def cell_border(cell, c=LINE):
    tcPr = cell._tc.get_or_add_tcPr()
    bs = OxmlElement("w:tcBorders")
    for e in ("top","left","bottom","right"):
        b = OxmlElement(f"w:{e}"); b.set(qn("w:val"),"single"); b.set(qn("w:sz"),"4"); b.set(qn("w:color"),c.lstrip("#"))
        bs.append(b)
    tcPr.append(bs)
def setrun(r, *, bold=False, italic=False, size=None, color=None, font=SANS):
    if bold: r.bold = True
    if italic: r.italic = True
    if size is not None: r.font.size = Pt(size)
    if color: r.font.color.rgb = RGBColor.from_string(color.lstrip("#"))
    if font:
        r.font.name = font
        rPr = r._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts"); rPr.insert(0, rFonts)
        for attr in ("w:ascii","w:hAnsi","w:eastAsia","w:cs"):
            rFonts.set(qn(attr), font)

def addbody(doc, t, *, size=11, italic=False, color=INK, sa=6, font=SANS):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(sa)
    r = p.add_run(t); setrun(r, size=size, color=color, italic=italic, font=font); return p

def addh(doc, t, *, level=1):
    sz = {1:18,2:14,3:12.5,4:11.5}
    p = doc.add_paragraph(); p.style = doc.styles[f"Heading {level}"]
    p.paragraph_format.space_before = Pt(16 if level==1 else 10); p.paragraph_format.space_after = Pt(4)
    for r in list(p.runs): r.text = ""
    r = p.add_run(t); setrun(r, bold=True, size=sz.get(level,12), color=INK, font=SERIF)
    if level == 1:
        pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom"); bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),"6"); bot.set(qn("w:space"),"4"); bot.set(qn("w:color"),GOLD)
        pBdr.append(bot); pPr.append(pBdr)
    return p

def addbullets(doc, items, *, size=11):
    for it in items:
        p = doc.add_paragraph(style="List Bullet"); r = p.add_run(it); setrun(r, size=size, color=INK, font=SANS)

def addtable(doc, header, rows, *, col_widths=None):
    n = len(header); tbl = doc.add_table(rows=len(rows)+1, cols=n); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j,h in enumerate(header):
        c = tbl.rows[0].cells[j]; cell_bg(c,HEAD); cell_border(c); c.paragraphs[0].text=""
        r = c.paragraphs[0].add_run(h); setrun(r,bold=True,size=10,color="FFFFFF", font=SANS)
        if col_widths: c.width = Inches(col_widths[j])
    for i,row in enumerate(rows, start=1):
        for j,v in enumerate(row):
            c = tbl.rows[i].cells[j]; cell_border(c)
            if i%2==0: cell_bg(c,BAND)
            c.paragraphs[0].text=""; r = c.paragraphs[0].add_run(str(v)); setrun(r,size=10,color=INK, font=SANS)
            if col_widths: c.width = Inches(col_widths[j])
    doc.add_paragraph(); return tbl

def addcallout(doc, kind, body, *, accent=GOLD, fill="F6F1E4"):
    tbl = doc.add_table(rows=1, cols=1); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = tbl.rows[0].cells[0]; cell_bg(c,fill); c.paragraphs[0].text=""
    rt = c.paragraphs[0].add_run(kind.upper()); setrun(rt,bold=True,size=9,color=accent, font=MONO)
    p = c.add_paragraph(); p.paragraph_format.space_before = Pt(3)
    r = p.add_run(body); setrun(r,size=10.5,color=INK, font=SANS); doc.add_paragraph()

def addkv(doc, pairs, *, lw=1.6, rw=4.9):
    tbl = doc.add_table(rows=len(pairs), cols=2); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,(k,v) in enumerate(pairs):
        ck,cv = tbl.rows[i].cells; ck.width=Inches(lw); cv.width=Inches(rw)
        cell_bg(ck,BAND); cell_border(ck); cell_border(cv)
        ck.paragraphs[0].text=""; cv.paragraphs[0].text=""
        rk = ck.paragraphs[0].add_run(k); setrun(rk,bold=True,size=10,color=INK, font=SANS)
        rv = cv.paragraphs[0].add_run(v); setrun(rv,size=10,color=INK, font=SANS)
    doc.add_paragraph()

def addeyebrow(doc, t):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(t.upper()); setrun(r,bold=True,size=9,color=MUTE, font=MONO)
    rPr = r._element.get_or_add_rPr(); sp = OxmlElement("w:spacing"); sp.set(qn("w:val"),"80"); rPr.append(sp)

def pb(doc):
    p = doc.add_paragraph(); p.add_run().add_break(WD_BREAK.PAGE)

def addcode(doc, t):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(t); setrun(r, size=9.5, color=INK, font=MONO)


def build():
    doc = Document()
    for s in doc.sections:
        s.page_width = Inches(8.5); s.page_height = Inches(11)
        s.top_margin = s.bottom_margin = Inches(0.85); s.left_margin = s.right_margin = Inches(0.95)
    doc.styles["Normal"].font.name = SANS
    doc.styles["Normal"].font.size = Pt(11)

    # ============ COVER ============
    addeyebrow(doc, "APEX · RC-E2E-03 · SIX AGENTS · DEEP DIVE & MATURATION")
    p = doc.add_paragraph(); r = p.add_run("Six Agents → Six Superagents")
    setrun(r, bold=True, size=24, color=INK, font=SERIF)
    p2 = doc.add_paragraph(); r2 = p2.add_run("How RC-E2E-03's specialized agents mature into context-aware superagents — Wave 1 through Wave 4+")
    setrun(r2, bold=True, size=14, color=INK2, italic=True, font=SERIF)
    doc.add_paragraph()
    addbody(doc,
            "Architectural deep dive on the six agents that compose APEX RC-E2E-03 Assortment & "
            "Pricing — Assess, Classify, Quantify, Approve, Act, Evidence-Write — and the maturation "
            "model that takes them from narrow-task specialists in Wave 1 to context-aware superagents "
            "by Wave 4+. Per-agent specifications, evolution paths, cross-agent fusion patterns, and "
            "the substrate that enables maturation. Audience: Engagement Architect, Lead Engineer, "
            "ML Engineering, Practice leadership.",
            italic=True, color=INK2)
    doc.add_paragraph()
    addkv(doc, [
        ("Document type",     "Architectural Deep Dive · Agent Maturation Model"),
        ("Service",           "RC-E2E-03 Assortment & Pricing"),
        ("Agents covered",    "6 (Assess · Classify · Quantify · Approve · Act · Evidence-Write)"),
        ("Reference client",  "The Kroger Co."),
        ("Audience",          "Engagement Architect · Lead Engineer · ML Engineering · Practice leadership · CAF platform team"),
        ("Companion",         "APEX-RC-E2E-03-Walkthrough.docx · manifests/*.yaml · APEX-RC-E2E-03-09-AI-ML-Model-Spec.docx"),
        ("Version / date",    "v1.0 · 27 Apr 2026"),
    ])
    pb(doc)

    # ============ EXECUTIVE SUMMARY ============
    addh(doc, "Executive summary — from specialists to superagents", level=1)
    addbody(doc,
            "Six agents compose the RC-E2E-03 service today. Each is purpose-built for a narrow "
            "responsibility — characterize an excursion, classify a disposition, quantify a "
            "$ impact, route a HITL approval, execute downstream actions, write the audit row. "
            "This narrow scoping is deliberate: it makes Wave 1 implementable, debuggable, and "
            "auditable.")
    addbody(doc,
            "Over time, those same six agents accumulate context. They observe outcomes (via "
            "LEDGER feedback), learn from cross-event patterns (via cohort drift watchtower), "
            "and collaborate across scenarios (via CAF A2A). What starts as six narrow-task "
            "agents in Wave 1 becomes — by Wave 4+ — a smaller number of context-aware "
            "superagents that handle broader decision surfaces with less HITL and more "
            "proactive intelligence.")
    addcallout(doc, "The maturation arc",
               "Wave 1 = six specialists, high HITL, narrow scope. Wave 2 = same six, calibrated "
               "with real outcomes, tier-tuned HITL. Wave 3 = same six, cross-banner aware, "
               "fused with adjacent scenarios. Wave 4+ = three-to-four superagents that absorb "
               "specialized siblings as the substrate matures. The architectural commitment: "
               "narrow-by-design Wave 1 is what enables broad-by-evolution Wave 4+.")

    addh(doc, "Why this matters now", level=2)
    addbody(doc,
            "Three reasons to write the maturation model down at the start of the engagement, not "
            "the end:")
    addbullets(doc, [
        "Stakeholders expect to know what the agents become. If the answer is unclear, scope keeps changing.",
        "Manifest versioning policy depends on it. PATCH / MINOR / MAJOR semantics need maturation rules.",
        "Build investment depends on it. We over-invest in Wave 1 narrow-correctness because Wave 4+ depth depends on Wave 1 fidelity.",
    ])
    pb(doc)

    # ============ SECTION 1 — SIX AGENTS AT A GLANCE ============
    addh(doc, "1 · The six agents at a glance", level=1)
    addbody(doc,
            "Each agent has a single responsibility today. The sequence below shows how they "
            "compose into a 90-second cold-chain disposition decision.")
    addtable(doc, ["#", "Agent", "Responsibility (W1)", "HITL gate (W1)", "Audit row?"],
             [["1", "Assess",         "Characterize the excursion event",
                                       "ZERO_TOUCH", "Yes (capture)"],
              ["2", "Classify",       "Per-SKU FSMA 204 disposition",
                                       "ESCALATION (confidence < 0.85)", "Yes (capture)"],
              ["3", "Quantify",       "$ impact for each disposition path",
                                       "ZERO_TOUCH", "Yes (capture)"],
              ["4", "Approve",        "HITL gate · Adaptive Card to Store Ops Lead",
                                       "ALWAYS for material entries", "Yes (HITL log)"],
              ["5", "Act",            "POS markdown · vendor ticket · FSMA log dispatch",
                                       "ZERO_TOUCH (post-approval)", "Yes (action log)"],
              ["6", "Evidence-Write", "Compose 14-field audit row · LEDGER write · packet update",
                                       "ZERO_TOUCH", "Yes (closeout — primary)"]],
             col_widths=[0.4, 1.5, 2.6, 1.7, 1.0])

    addh(doc, "1.1 · Why six and not five", level=2)
    addbody(doc,
            "Earlier RC-E2E-03 documentation collapsed audit-row writing into the Act agent. "
            "We split Evidence-Write into its own agent for three reasons:")
    addbullets(doc, [
        "Separation of concerns — Act mutates external systems (POS, vendor portal, FSMA log); Evidence-Write mutates only LEDGER. Different blast-radius.",
        "Independent maturation — Evidence-Write's evolution path (toward auto-generated quarterly packets, FDA inquiry response, recall traceability) is distinct from Act's evolution (toward broader downstream channels and SLAs).",
        "FSMA 204 alignment — explicit Evidence-Write agent makes the 'records are generated, not collated' commitment legible to auditors and regulators.",
    ])

    addh(doc, "1.2 · The architectural pattern", level=2)
    addbody(doc,
            "All six agents follow the same pattern: bounded responsibility, typed I/O, manifest-"
            "declared HITL gate, audit-row emission. This is the APEX architectural commitment. "
            "What changes between agents is the responsibility scope; what stays constant is the "
            "contract surface.")
    addcallout(doc, "Architectural invariant",
               "Every agent decision emits a 14-field audit row. Every agent has a manifest. Every "
               "agent contracts on capability tags (cap:*) — not on specific LLMs or specific tools. "
               "This is what makes the agents portable across reference clients and what makes "
               "maturation possible without architecture rewrites.",
               accent=TEAL, fill="E4F1EE")

    pb(doc)

    # ============ SECTION 2 — PER-AGENT DEEP DIVES ============
    addh(doc, "2 · Per-agent deep dives", level=1)
    addbody(doc,
            "Each agent gets its own subsection: purpose, inputs, outputs, tools, HITL behavior, "
            "failure modes, audit-row contribution, and maturation roadmap.")

    # ---------- 2.1 ASSESS ----------
    addh(doc, "2.1 · Assess agent", level=2)
    addh(doc, "Purpose", level=3)
    addbody(doc,
            "Characterize the cold-chain excursion event. Reads the temperature telemetry window "
            "and fixture state; produces an ExcursionProfile that downstream agents reason against. "
            "First agent in the chain; the trigger that fires the rest.")
    addh(doc, "Inputs", level=3)
    addbullets(doc, [
        "BRZ.RC.REFRIG_TELEMETRY — window from (excursion_start − 30 min) to now",
        "SCML.Fixture (Silver) — current fixture state and recent excursion history",
        "BRZ.RC.REFRIG_SERVICE — fixture failure history (informs severity weighting)",
        "Watermark trigger — temperature > setpoint+5°F for ≥ 15 min consecutive",
    ])
    addh(doc, "Outputs", level=3)
    addbody(doc, "ExcursionProfile schema:")
    addcode(doc,
            "{\n"
            '  excursion_id: STRING (uuid)\n'
            '  fixture_id: STRING\n'
            '  store_id: STRING\n'
            '  peak_temp_f: DECIMAL\n'
            '  avg_temp_f: DECIMAL\n'
            '  duration_minutes: INT\n'
            '  door_open_event_count: INT\n'
            '  severity_class: ENUM[NORMAL, WATCH, CRITICAL]\n'
            '  confidence: DECIMAL (0.0-1.0)\n'
            '  trace_id: STRING (correlator)\n'
            "}")
    addh(doc, "Tools / capability tags", level=3)
    addbullets(doc, [
        "cap:telemetry-window-query — read REFRIG_TELEMETRY window",
        "cap:fixture-state-lookup — read SCML.Fixture",
        "cap:fixture-failure-history — read REFRIG_SERVICE",
    ])
    addh(doc, "Failure modes", level=3)
    addtable(doc, ["Mode", "Detection", "Mitigation"],
             [["Telemetry sensor offline", "offline_flag = true in window", "Watermark widens; severity defaults to WATCH; alerts Facilities Eng"],
              ["Window data incomplete (< 80% completeness)", "Coverage < 80% in 30-min window", "Returns confidence 0.5 with explicit data-quality flag"],
              ["Sensor reading anomaly (impossible value)", "DQ pack flag", "Discard reading; window-extends; logs to Bronze quarantine"],
              ["Multiple-fixture cascade event", "Co-correlated alarms across adjacent fixtures", "Compose multi-fixture profile; severity may upgrade"]],
             col_widths=[2.2, 2.0, 2.5])
    addh(doc, "Audit-row contribution", level=3)
    addbody(doc,
            "Emits a 'capture' audit row at output time with: trace_id, agent_id (assess), "
            "input window hash, ExcursionProfile output hash, manifest_version, prompt_version, "
            "decision_class (severity), confidence, latency_ms.")

    # ---------- 2.2 CLASSIFY ----------
    addh(doc, "2.2 · Classify agent", level=2)
    addh(doc, "Purpose", level=3)
    addbody(doc,
            "Per-SKU disposition classification — sell-through, markdown, destroy, or hold — "
            "based on the ExcursionProfile and FSMA 204 policy corpus. The most-consequential "
            "agent; its output drives the $ impact downstream.")
    addh(doc, "Inputs", level=3)
    addbullets(doc, [
        "Assess.output (ExcursionProfile)",
        "SCML.Lot — affected lots (joined on fixture_id)",
        "DIM.Item — temperature class, FTL flag, shelf-life",
        "FSMA 204 policy corpus (versioned, per-FTL-category)",
        "Historical disposition outcomes (BRZ.RC.FOOD_SAFETY, last 90 days similar events)",
    ])
    addh(doc, "Outputs", level=3)
    addbody(doc, "DispositionTable schema:")
    addcode(doc,
            "[{\n"
            '  lot_code: STRING\n'
            '  sku: STRING\n'
            '  disposition: ENUM[SELL_THROUGH, MARKDOWN, DESTROY, HOLD]\n'
            '  reason: STRING\n'
            '  policy_citation: STRING (FSMA-204 clause reference)\n'
            '  confidence: DECIMAL\n'
            '  escalation_required: BOOLEAN\n'
            "}]")
    addh(doc, "Tools / capability tags", level=3)
    addbullets(doc, [
        "cap:scml-lot-query — affected lots",
        "cap:fsma-policy-resolver — policy corpus lookup",
        "cap:dim-item-lookup — item attributes",
        "cap:disposition-classification — ML model (see ML Spec §4)",
    ])
    addh(doc, "HITL gate", level=3)
    addbody(doc,
            "If any lot has confidence < 0.85, the entire DispositionTable escalates to Food "
            "Safety Compliance Lead before reaching Quantify. This is the single most-tuned "
            "threshold in the system — UC-07 in the use-case catalog covers calibration.")
    addh(doc, "Failure modes", level=3)
    addtable(doc, ["Mode", "Detection", "Mitigation"],
             [["Policy corpus version mismatch", "kde_schema_version drift", "Quarantine event; escalate to Food Safety Lead"],
              ["Lot data missing for affected SKU", "SCML.Lot returns empty", "Default to HOLD disposition; require manual lot-trace"],
              ["Model confidence collapse (drift)", "Cohort Drift Watchtower alert", "Auto-trigger retrain workflow; escalate to ML + Audience Lead"],
              ["FTL category not in current corpus", "ftl_category not matched", "Default to most-conservative policy; escalate"]],
             col_widths=[2.2, 2.0, 2.5])
    addh(doc, "Audit-row contribution", level=3)
    addbody(doc, "Emits one 'capture' audit row per lot in the DispositionTable. policy_version "
                 "stamp is mandatory — auditors must see exactly which policy version fired.")

    # ---------- 2.3 QUANTIFY ----------
    addh(doc, "2.3 · Quantify agent", level=2)
    addh(doc, "Purpose", level=3)
    addbody(doc,
            "Convert disposition decisions into dollar impact. For each affected lot, compute "
            "three scenarios: sell-through value, markdown value (with optimal recommended "
            "discount %), destroy loss. The Quantify output is what arrives at the Approve agent's "
            "Adaptive Card.")
    addh(doc, "Inputs", level=3)
    addbullets(doc, [
        "Classify.output (DispositionTable)",
        "MERML.Price — current pricing",
        "MERML.Markdown — markdown-recommendation curves",
        "MERML.Elasticity — per-SKU elasticity (see ML Spec §3)",
        "SCML.Lot — quantity-on-hand per lot",
    ])
    addh(doc, "Outputs", level=3)
    addbody(doc, "QuantifiedDisposition schema:")
    addcode(doc,
            "[{\n"
            '  lot_code: STRING\n'
            '  sku: STRING\n'
            '  recommended_disposition: ENUM (echoes Classify)\n'
            '  sell_through_value_usd: DECIMAL\n'
            '  markdown_value_usd: DECIMAL\n'
            '  destroy_loss_usd: DECIMAL\n'
            '  markdown_recommendation_pct: DECIMAL (0.0-1.0)\n'
            '  expected_npv_usd: DECIMAL\n'
            '  confidence: DECIMAL\n'
            "}]")
    addh(doc, "Tools / capability tags", level=3)
    addbullets(doc, [
        "cap:merml-price-lookup",
        "cap:merml-elasticity-curve",
        "cap:scml-lot-qty",
    ])
    addh(doc, "Failure modes", level=3)
    addtable(doc, ["Mode", "Detection", "Mitigation"],
             [["Elasticity model unavailable for new SKU", "MERML.Elasticity returns no row", "Fall back to category-level prior; flag for cold-start retrain"],
              ["Price-zone mismatch (banner-aware)", "MERML.Price wrong zone returned", "Quarantine; escalate to Pricing Manager"],
              ["Quantity-on-hand zero (already-depleted lot)", "SCML.Lot.qty = 0", "Skip lot; log audit row 'no-action-needed'"]],
             col_widths=[2.7, 1.8, 2.2])

    pb(doc)

    # ---------- 2.4 APPROVE ----------
    addh(doc, "2.4 · Approve agent (HITL orchestrator)", level=2)
    addh(doc, "Purpose", level=3)
    addbody(doc,
            "The HITL gate. Routes material disposition decisions to the Store Operations Lead "
            "(or Food Safety Lead for destroy-all events) via Microsoft Teams Adaptive Card. "
            "Captures approve / modify / reject decision and writes a HITL-log audit row.")
    addh(doc, "Inputs", level=3)
    addbullets(doc, [
        "Quantify.output (QuantifiedDisposition)",
        "Materiality policy manifest (banner-aware tier thresholds)",
        "DIM.Store (for tier classification)",
        "Entra identity service (named-persona resolution)",
    ])
    addh(doc, "Outputs", level=3)
    addbody(doc, "ApprovedDisposition schema:")
    addcode(doc,
            "{\n"
            '  disposition_record: <QuantifiedDisposition>\n'
            '  approver_id: STRING (Entra identity)\n'
            '  approver_role: ENUM[StoreOpsLead, FoodSafetyLead, Designate]\n'
            '  approval_ts: TIMESTAMP\n'
            '  hitl_kind: ENUM[ZERO_TOUCH, ACK_ONLY, HITL, ESCALATION]\n'
            '  modifications: ARRAY<{lot_code, original, modified}>\n'
            '  rejection_reason: STRING (if rejected)\n'
            "}")
    addh(doc, "HITL gate logic", level=3)
    addtable(doc, ["Condition", "HITL kind", "Routed to"],
             [["Tier 1 store (any event)", "HITL", "Store Ops Lead"],
              ["Tier 2 + value > $2,500", "HITL", "Store Ops Lead"],
              ["Tier 2 + value $500-$2,500 + conf ≥ 0.95", "ZERO_TOUCH", "Auto + notify"],
              ["Tier 2 + value $500-$2,500 + conf < 0.95", "HITL", "Store Ops Lead"],
              ["Tier 3 + value ≤ $500 + conf ≥ 0.90", "ZERO_TOUCH", "Log only"],
              ["Any DESTROY-ALL event", "HITL (override)", "Food Safety Lead"],
              ["Any FSMA-evidence-gap event", "HITL (override)", "Food Safety Lead"]],
             col_widths=[2.8, 1.4, 2.2])
    addh(doc, "Adaptive Card composition", level=3)
    addbody(doc,
            "The card carries: excursion summary, affected lot count, proposed actions per lot, "
            "dollar impact (sell-through/markdown/destroy comparison), reasoning trace link, "
            "policy citations, three buttons (Approve / Modify / Reject), free-text comment field. "
            "Modifications write back per-lot deltas; rejection requires a reason field.")
    addh(doc, "Failure modes", level=3)
    addtable(doc, ["Mode", "Detection", "Mitigation"],
             [["Adaptive Card delivery failure", "Teams API timeout", "Retry; fallback to email; escalate to designate-approver"],
              ["Approver timeout (10 min default)", "No response in window", "Escalate to Food Safety Lead per policy"],
              ["Identity binding mismatch (Entra)", "approver_id resolution fails", "Block decision; alert IT-Security"],
              ["Approver rejects without reason", "rejection_reason empty", "Block; require reason"]],
             col_widths=[2.4, 1.7, 2.3])

    # ---------- 2.5 ACT ----------
    addh(doc, "2.5 · Act agent", level=2)
    addh(doc, "Purpose", level=3)
    addbody(doc,
            "Post-approval execution. Three downstream actions in parallel: POS markdown upload "
            "(Toshiba/IBM/NCR per banner), refrigeration vendor ticket creation (Emerson/Lennox/"
            "Hussmann/Hill PHOENIX), and FSMA regulatory log entry. The mutation surface — Act is "
            "the only agent that mutates external systems.")
    addh(doc, "Inputs", level=3)
    addbullets(doc, [
        "Approve.output (ApprovedDisposition) — only when hitl_kind != REJECTED",
    ])
    addh(doc, "Outputs", level=3)
    addbody(doc, "ActOutcome schema:")
    addcode(doc,
            "{\n"
            '  pos_markdown_status: ENUM[SUCCESS, RETRY, FAILED]\n'
            '  vendor_ticket_id: STRING\n'
            '  vendor_ticket_status: ENUM[OPEN, ACK, FAILED]\n'
            '  fsma_log_id: STRING\n'
            '  all_actions_completed: BOOLEAN\n'
            '  closeout_ts: TIMESTAMP\n'
            "}")
    addh(doc, "Tools / capability tags", level=3)
    addbullets(doc, [
        "cap:pos-markdown-upload (banner-aware)",
        "cap:vendor-portal-ticket (vendor-aware)",
        "cap:fsma-regulatory-log",
    ])
    addh(doc, "Parallelism + retry", level=3)
    addbody(doc,
            "Three actions execute in parallel. Each has independent retry policy (3 retries, "
            "exponential backoff). Failure of one does not block the others; failures escalate "
            "to on-call. Idempotency key per action prevents duplicates on retry.")
    addh(doc, "Failure modes", level=3)
    addtable(doc, ["Mode", "Detection", "Mitigation"],
             [["POS API rate limit", "429 from POS gateway", "Backoff + retry; queue for batch upload"],
              ["Vendor portal API down", "5xx or timeout", "Queue locally; alert Facilities Eng; resume on recovery"],
              ["FSMA log write failure", "DB write fails", "CRITICAL alert; manual log entry triggered; quarantine until resolved"],
              ["Partial-action completion", "1 of 3 succeeded", "Compensating-action workflow per action type"]],
             col_widths=[2.3, 1.6, 2.5])

    # ---------- 2.6 EVIDENCE-WRITE ----------
    addh(doc, "2.6 · Evidence-Write agent (NEW)", level=2)
    addh(doc, "Purpose", level=3)
    addbody(doc,
            "Composes the closeout 14-field audit row from the upstream agent traces, writes it "
            "to LEDGER, and updates the rolling FSMA 204 evidence packet. The agent that makes "
            "'records are generated, not collated' a structural commitment.")
    addh(doc, "Inputs", level=3)
    addbullets(doc, [
        "All upstream agent audit rows (Assess + Classify + Quantify + Approve + Act)",
        "Manifest + policy + prompt version stamps (current effective)",
        "Reasoning-trace references (LangFuse spans correlated by trace_id)",
        "Consent state (BRZ.RC.LOYALTY_HH) for any household-relevant decision",
    ])
    addh(doc, "Outputs", level=3)
    addbody(doc, "Closeout audit row — the 14-field structure that anchors LEDGER:")
    addcode(doc,
            "{\n"
            '  trace_id: STRING\n'
            '  scenario_id: \"RC-SUPCHN-01-cold-chain-excursion\"\n'
            '  service_id: \"RC-E2E-03\"\n'
            '  tenant_id: \"kr\"\n'
            '  banner: STRING\n'
            '  event_kind: \"rc_e2e_03_closeout\"\n'
            '  event_ts: TIMESTAMP\n'
            '  manifest_version: STRING\n'
            '  policy_version: STRING\n'
            '  prompt_version: STRING\n'
            '  reasoning_trace_ref: STRING (LangFuse span id)\n'
            '  hitl_actor: STRING (Entra identity if HITL)\n'
            '  hitl_kind: ENUM\n'
            '  kpis: { time_to_decision_sec, shrink_avoided_usd, policy_conformance,\n'
            '          fsma_audit_row_complete, manager_time_reclaimed_min }\n'
            "}")
    addh(doc, "Tools / capability tags", level=3)
    addbullets(doc, [
        "cap:ledger-closeout-write — Delta + Postgres index",
        "cap:fsma-packet-update — rolling quarterly packet contributor",
        "cap:purview-lineage-emit — lineage to all upstream Bronze/Silver inputs",
    ])
    addh(doc, "Why a separate agent", level=3)
    addbody(doc,
            "The Evidence-Write responsibility is structurally distinct from Act for three reasons:")
    addbullets(doc, [
        "Different blast radius — Act mutates external systems with potential customer/financial impact; Evidence-Write mutates only LEDGER (read-only to all downstream).",
        "Different SLO — Act has retry policies tuned for external-system failure; Evidence-Write must succeed near-100% (LEDGER coverage SLO is 99.5%).",
        "Different evolution — Evidence-Write matures toward auto-generated FDA inspection packets, recall traceability, and Audit Committee dashboards; Act matures toward broader downstream channels and SLAs.",
    ])
    addh(doc, "Failure modes", level=3)
    addtable(doc, ["Mode", "Detection", "Mitigation"],
             [["LEDGER write failure", "Delta or Postgres index failure", "CRITICAL alert; failover to backup write path"],
              ["Audit row schema validation fails", "DQ pack rejection", "Quarantine row; auto-PR for schema reconcile"],
              ["FSMA packet generation fails", "Quarterly job error", "Retry; manual generation backstop; FDA-relations notified"],
              ["LangFuse trace correlation gap", "trace_id not found upstream", "Mark row with degraded-evidence flag; investigate gap"]],
             col_widths=[2.4, 1.7, 2.5])

    pb(doc)

    # ============ SECTION 3 — THE MATURATION MODEL ============
    addh(doc, "3 · The maturation model", level=1)
    addbody(doc,
            "What does it mean for an agent to mature? In APEX, three concrete dimensions define "
            "agent maturity. Each dimension has measurable thresholds; each evolves on its own clock.")

    addh(doc, "3.1 · The three dimensions", level=2)
    addtable(doc, ["Dimension", "What it measures", "How we measure it"],
             [["Context depth", "How much history each decision considers",
                                "Average lookback-window size · cross-event correlation count · feature-engineering complexity"],
              ["Capability breadth", "Range of decisions in scope",
                                    "Number of disposition classes · number of supported scenarios · cross-domain reach"],
              ["Autonomy band", "% of decisions auto-executed without HITL",
                               "Auto-execute rate by tier · per-decision-class HITL rate · HITL-override rate"]],
             col_widths=[1.8, 2.6, 2.3])

    addh(doc, "3.2 · The four maturity bands", level=2)
    addtable(doc, ["Band", "Context depth", "Capability breadth", "Autonomy band"],
             [["Wave 1 · Specialist",
               "Shallow · single-event window",
               "Single-scenario · cold-chain only",
               "Low · 30-50% auto · high HITL"],
              ["Wave 2 · Calibrated",
               "Mid · 30-day rolling history",
               "Single-scenario · banner-aware",
               "Mid · 55% auto · tier-tuned HITL"],
              ["Wave 3 · Cross-banner",
               "Deep · multi-quarter outcome history",
               "Cross-banner · scenario-fused (E2E-11/12 candidates)",
               "High · 75% auto · materiality-only HITL"],
              ["Wave 4+ · Superagent",
               "Multi-scenario · cohort + temporal · proactive",
               "Multi-scenario · cross-Practice · proactive",
               "Selective · 85%+ auto · principle-based HITL"]],
             col_widths=[1.4, 2.0, 2.2, 1.9])

    addh(doc, "3.3 · What 'superagent' means precisely", level=2)
    addbody(doc,
            "Superagent is not a marketing term — it has structural meaning in APEX:")
    addbullets(doc, [
        "Multi-scenario reach — the same agent serves RC-E2E-03 + RC-E2E-11 + RC-E2E-12 from a unified codebase, with scenario-routing in the dispatch layer not the agent body.",
        "Proactive initiation — agent proposes interventions before a trigger fires (e.g., flagging fixture-failure-risk before excursion); shifts from reactive to predictive.",
        "Cross-agent context absorption — agent reasons over what other agents have decided historically; enables compound reasoning.",
        "Self-tuning thresholds — confidence calibration, materiality thresholds, and policy interpretation tune themselves based on LEDGER outcome attribution.",
        "Principle-based HITL — instead of static threshold rules, escalation routing reasons over case-by-case material novelty, regulatory exposure, and downstream-effect uncertainty.",
    ])
    addcallout(doc, "What superagent is NOT",
               "Superagent is not LLM-with-no-rules. The architectural commitments — manifest "
               "discipline, audit row per decision, HITL as first-class primitive, policy "
               "versioning, capability-tag indirection — all hold across maturation. The superagent "
               "of Wave 4+ is more capable, not less governed.",
               accent=GOLD)

    pb(doc)

    # ============ SECTION 4 — PER-AGENT EVOLUTION ============
    addh(doc, "4 · Per-agent evolution roadmaps", level=1)
    addbody(doc,
            "Per-agent specifically. Each agent has a distinct evolution surface; some mature "
            "faster than others; some are absorbed into siblings (Section 5).")

    addh(doc, "4.1 · Assess agent evolution", level=2)
    addtable(doc, ["Wave", "Context depth", "Capability breadth", "Autonomy"],
             [["W1", "30-min telemetry window", "Cold-chain excursions only", "ZERO_TOUCH (deterministic)"],
              ["W2", "+ 30-day fixture history", "+ Multi-fixture cascade detection", "ZERO_TOUCH"],
              ["W3", "+ Cross-banner pattern matching · weather context", "+ Predictive risk-scoring (pre-excursion)", "ZERO_TOUCH"],
              ["W4+", "+ Multi-scenario context (E2E-11 perishables integrity)", "Proactive — flags fixture-risk before excursion fires", "Selective HITL on proactive interventions"]],
             col_widths=[0.7, 2.4, 2.4, 1.8])
    addbody(doc,
            "Assess matures from a deterministic event characterizer (W1) to a proactive "
            "fixture-risk forecaster (W4+). The expansion is in capability breadth — same agent, "
            "broader trigger surface — and in temporal context (predictive vs. reactive).")

    addh(doc, "4.2 · Classify agent evolution", level=2)
    addtable(doc, ["Wave", "Context depth", "Capability breadth", "Autonomy"],
             [["W1", "FSMA 204 policy + lot-context", "Cold-chain disposition only", "Mid (escalate < 0.85)"],
              ["W2", "+ Historical disposition outcomes (90 days)", "+ Per-FTL-category specialization", "Mid+"],
              ["W3", "+ Cross-banner pattern · supplier-quality signal", "+ Adjacent perishables (RC-E2E-11)", "High"],
              ["W4+", "+ Cross-Practice (HLS 340B disposition logic, ER hazmat patterns)", "Multi-scenario disposition reasoning", "High · principle-based escalation"]],
             col_widths=[0.7, 2.4, 2.4, 1.8])
    addbody(doc,
            "Classify is the most-leveraged maturation surface — its accuracy directly determines "
            "shrink-avoided. The W1→W4+ arc is the journey from 'apply the FSMA 204 policy' to "
            "'apply the pattern of regulatory disposition decisions across food, pharma, and "
            "hazmat.' Cross-Practice generalization is the long-arc goal.")

    addh(doc, "4.3 · Quantify agent evolution", level=2)
    addtable(doc, ["Wave", "Context depth", "Capability breadth", "Autonomy"],
             [["W1", "MERML.Price + MERML.Markdown + lot qty", "Cold-chain markdown only", "ZERO_TOUCH"],
              ["W2", "+ MERML.Elasticity per-SKU curves", "+ Markdown optimization beyond cold-chain", "ZERO_TOUCH"],
              ["W3", "+ Cross-banner elasticity · loyalty-cohort overlay", "+ Dynamic pricing fusion (RC-E2E-04)", "ZERO_TOUCH"],
              ["W4+", "+ Cross-Practice elasticity (TMT subscription pricing)", "Multi-scenario pricing decisioning", "ZERO_TOUCH (principle-based)"]],
             col_widths=[0.7, 2.4, 2.4, 1.8])
    addbody(doc,
            "Quantify's evolution is quietly the most economically valuable. By W3, it absorbs "
            "RC-E2E-04 loyalty-pricing fusion logic; by W4+, the same agent that quantified "
            "cold-chain markdowns is reasoning about TMT streaming-subscription pricing, ER "
            "commodity-position revaluation, and other scenarios where elasticity-times-quantity "
            "is the universal pattern.")

    addh(doc, "4.4 · Approve agent evolution", level=2)
    addtable(doc, ["Wave", "Context depth", "Capability breadth", "Autonomy"],
             [["W1", "Static materiality policy", "Cold-chain disposition routing", "HITL-heavy"],
              ["W2", "+ Tier-tuned · banner-overrides · designate-approver workflow", "+ Multi-tier approver hierarchy", "HITL with tier discretion"],
              ["W3", "+ Per-cohort approval-pattern learning · approver-load balancing", "+ Cross-scenario approval routing (E2E-11/12)", "Tier-tuned + load-aware"],
              ["W4+", "+ Principle-based escalation · case-by-case novelty assessment", "Multi-scenario meta-routing across Practice", "Selective HITL on novel cases only"]],
             col_widths=[0.7, 2.4, 2.4, 1.8])
    addbody(doc,
            "Approve's W4+ form is a principle-based escalator — it reasons about whether a "
            "decision is genuinely novel or analogous to past approved decisions, and routes "
            "accordingly. Most decisions become ZERO_TOUCH because the system has confidence "
            "they fall within an already-approved precedent class. Only genuinely-novel cases "
            "escalate. This is the deepest behavioral shift across the arc.")

    addh(doc, "4.5 · Act agent evolution", level=2)
    addtable(doc, ["Wave", "Context depth", "Capability breadth", "Autonomy"],
             [["W1", "Three downstream actions in parallel", "POS markdown · vendor ticket · FSMA log", "ZERO_TOUCH (post-approval)"],
              ["W2", "+ Receipt-aware monitoring · auto-retry logic", "+ Multi-channel orchestration patterns", "ZERO_TOUCH"],
              ["W3", "+ Cross-banner action coordination · vendor-contract-aware", "+ Adjacent-scenario action surfaces (recall, supplier-pull)", "ZERO_TOUCH"],
              ["W4+", "+ Self-orchestrating compensating actions on partial failure", "Cross-Practice action library (any APEX scenario)", "ZERO_TOUCH (compensation-aware)"]],
             col_widths=[0.7, 2.4, 2.4, 1.8])

    addh(doc, "4.6 · Evidence-Write agent evolution", level=2)
    addtable(doc, ["Wave", "Context depth", "Capability breadth", "Autonomy"],
             [["W1", "14-field audit row · trace correlation", "RC-E2E-03 only · single-tenant", "ZERO_TOUCH"],
              ["W2", "+ Quarterly FSMA evidence packet · rolling stamps", "+ Multi-cohort packet stratification", "ZERO_TOUCH"],
              ["W3", "+ Cross-banner consolidated packet · auditor-walkthrough automation", "+ FDA inquiry response · recall traceability", "ZERO_TOUCH"],
              ["W4+", "+ Principle-based evidence interpretation for novel inquiries", "Cross-Practice evidence (SOX + FDA + state regulators)", "ZERO_TOUCH (multi-regulatory)"]],
             col_widths=[0.7, 2.4, 2.4, 1.8])

    pb(doc)

    # ============ SECTION 5 — CROSS-AGENT FUSION ============
    addh(doc, "5 · Cross-agent fusion patterns", level=1)
    addbody(doc,
            "Maturation isn't only per-agent depth — agents also fuse with each other. By Wave 4+, "
            "the original six agents may consolidate into three or four superagents that absorb "
            "specialized siblings. The patterns below show the most-likely fusion paths.")

    addh(doc, "5.1 · Pattern A · Assess + Classify → Risk-Classify superagent", level=2)
    addbody(doc,
            "Why fuse: Both agents reason about the same input domain (excursion + lot context) — "
            "Assess characterizes the event, Classify decides the disposition. Their reasoning "
            "patterns are increasingly correlated as the system matures. By W4+, the same model "
            "can produce both ExcursionProfile and DispositionTable in one inference, with "
            "tighter latency and higher accuracy from shared context.")
    addbullets(doc, [
        "Trigger for fusion: Classify's confidence becomes systematically dependent on Assess's severity_class — they're co-modeling the same latent",
        "Architectural change: Single agent emits both schemas with shared backbone",
        "Maturity gate: Pattern emerges naturally in W3; fusion implementation in W4+",
        "Risk: Loss of debuggability — must preserve trace separation in audit row",
    ])

    addh(doc, "5.2 · Pattern B · Quantify + Approve → Materiality-Decision superagent", level=2)
    addbody(doc,
            "Why fuse: Quantify's $ output is the input that determines Approve's HITL routing. "
            "By W4+, the routing logic and the quantification become a single judgment — 'this "
            "matters enough to require approval' is one decision, not two.")
    addbullets(doc, [
        "Trigger for fusion: Materiality policy starts incorporating Quantify confidence; Quantify starts respecting materiality thresholds proactively",
        "Architectural change: Quantify outputs proposed action with embedded HITL recommendation",
        "Maturity gate: Pattern emerges in W3 cohort-tuned environments; fusion implementation W4+",
        "Risk: Erodes the bright line between machine-quantification and human-approval — preserve audit-row distinction even after agent merge",
    ])

    addh(doc, "5.3 · Pattern C · Act + Evidence-Write → Closeout superagent", level=2)
    addbody(doc,
            "Why fuse: Both agents fire post-approval. Act mutates external systems; Evidence-"
            "Write mutates LEDGER. By W4+, both can be expressed as 'commit the decision to all "
            "downstream consumers' with consistent transactional semantics.")
    addbullets(doc, [
        "Trigger for fusion: Act + Evidence-Write begin sharing failure-correlation data; Evidence-Write needs visibility into Act's compensating actions",
        "Architectural change: Single agent commits decision to POS + vendor + FSMA + LEDGER + FSMA-packet in coordinated transaction",
        "Maturity gate: Strong fusion candidate — emerges naturally in W3; implementation W4+",
        "Risk: Coupling write-to-external and write-to-evidence may concentrate failure modes; maintain ability to write evidence even if external mutation fails",
    ])

    addh(doc, "5.4 · The W4+ shape · 6 → 3", level=2)
    addtable(doc, ["W1 / W2 / W3 (six agents)", "W4+ superagents (three agents)", "Capability scope"],
             [["Assess", "Risk-Classify", "Multi-scenario excursion + disposition reasoning"],
              ["Classify", "(absorbed)", ""],
              ["Quantify", "Materiality-Decision", "Multi-scenario quantification + HITL routing"],
              ["Approve", "(absorbed)", ""],
              ["Act", "Closeout", "Multi-scenario downstream commit + evidence emission"],
              ["Evidence-Write", "(absorbed)", ""]],
             col_widths=[2.6, 2.4, 2.4])
    addcallout(doc, "Note on parent orchestrator",
               "The parent-orchestrator manifest also evolves. In W1-W3 it sequences six agents; "
               "in W4+ it sequences three superagents with broader interaction patterns. Manifest "
               "MAJOR-version-bump is the appropriate signal at the W3→W4 transition.",
               accent=VIO, fill="EFE5F4")

    pb(doc)

    # ============ SECTION 6 — THE SUBSTRATE ============
    addh(doc, "6 · The substrate that enables maturation", level=1)
    addbody(doc,
            "Maturation doesn't happen automatically. It requires substrate — the platform-level "
            "patterns and feedback loops that let agents observe their own outcomes and improve. "
            "Five substrate elements:")

    addh(doc, "6.1 · LEDGER feedback loop", level=2)
    addbody(doc,
            "Every agent decision writes to LEDGER. Every outcome (was the markdown the right "
            "call? Did the disposition hold up? Did the FDA inspection pass?) writes back to "
            "LEDGER. Periodically (weekly to monthly), the system queries LEDGER for "
            "decision-vs-outcome correlation and feeds the result into model retraining + policy "
            "calibration. This is the closed loop that turns reactive agents into self-improving "
            "ones.")
    addh(doc, "6.2 · CAF agent registry maturity tagging", level=2)
    addbody(doc,
            "The CAF registry tracks each agent's capability tag plus a maturity level "
            "(W1-Specialist, W2-Calibrated, W3-Cross-banner, W4-Superagent). APEX manifests can "
            "specify minimum maturity required for production use. As agents mature, manifests "
            "auto-pin to the latest matching maturity level via the registry's resolution logic.")
    addh(doc, "6.3 · Manifest version-up policy", level=2)
    addtable(doc, ["Change type", "Version bump", "Auto-deploy?"],
             [["Prompt tweak (no behavior change)", "PATCH (1.0.0 → 1.0.1)", "Yes"],
              ["New capability with backward compat (e.g., new disposition class)", "MINOR (1.0.0 → 1.1.0)", "After accept-upgrade gate"],
              ["Schema/contract change · breaking", "MAJOR (1.0.0 → 2.0.0)", "Manual approval; new tenant-subscription"],
              ["Maturity-band promotion (W1→W2)", "MAJOR + maturity tag", "Steering Committee gate"],
              ["Cross-agent fusion (W3→W4 superagent)", "MAJOR + parent-orchestrator update", "Practice leadership gate"]],
             col_widths=[2.7, 2.0, 2.0])
    addh(doc, "6.4 · Cohort drift detection", level=2)
    addbody(doc,
            "Per UC-22 in the use-case catalog, the Cohort Drift Watchtower runs continuously, "
            "watching for performance gaps that suggest maturation friction. When detected, "
            "drift triggers retraining; when retraining doesn't restore performance, drift "
            "triggers maturity-band review (perhaps the agent has outgrown its current band and "
            "needs cross-scenario context).")
    addh(doc, "6.5 · Outcome attribution discipline", level=2)
    addbody(doc,
            "Maturation requires outcome attribution at trace-id granularity. Every audit row "
            "must be linkable to the specific decision and the specific outcome. Without that, "
            "we can't measure what mattered, and learning becomes anecdotal. The Evidence-Write "
            "agent's existence is what makes this discipline structural rather than aspirational.")

    pb(doc)

    # ============ SECTION 7 — GOVERNANCE ============
    addh(doc, "7 · Governance · how an agent matures with discipline", level=1)
    addbody(doc,
            "Agent maturation is a governed process. The patterns below define who approves "
            "promotion, what evidence is required, and how rollback works.")

    addh(doc, "7.1 · Promotion gates · maturity band changes", level=2)
    addtable(doc, ["Promotion", "Evidence required", "Approver"],
             [["W1 → W2 Calibrated",
               "30 consecutive operating days · per-tier confidence calibration · banner-aware policy tested",
               "Engagement Architect + Steering Committee"],
              ["W2 → W3 Cross-banner",
               "90 consecutive operating days · cross-banner correlation analysis · cross-scenario test runs",
               "Steering Committee + DMTSP_lead"],
              ["W3 → W4 Superagent",
               "180 days at W3 · cross-Practice test bench · principle-based-HITL pilot evidence",
               "Practice leadership + Independence office (re-review for audit clients)"]],
             col_widths=[1.8, 3.0, 1.9])

    addh(doc, "7.2 · Cross-agent fusion gates", level=2)
    addbullets(doc, [
        "Both source agents at W3 maturity for ≥ 90 days",
        "Shadow-fusion run for ≥ 30 days (fused agent runs alongside specialists, decisions compared)",
        "Independent audit-row preservation verified (single fused agent must still emit per-step evidence)",
        "Approver: Practice leadership + DMTSP_lead",
    ])

    addh(doc, "7.3 · Rollback policy", level=2)
    addbody(doc,
            "Agent maturity promotions are reversible. If post-promotion performance degrades, "
            "the agent can be rolled back to its prior maturity band via:")
    addbullets(doc, [
        "Manifest reversion to prior MAJOR version (CAF registry serves the prior model)",
        "Capability-tag re-pin (manifest references prior maturity level)",
        "Tenant-subscription override (per-tenant maturity floor)",
        "Steering Committee notification within 24 hours of rollback",
    ])
    addh(doc, "7.4 · Audit posture across maturation", level=2)
    addcallout(doc, "Audit invariant",
               "Maturation does not weaken audit. Every decision at every maturity band emits "
               "the same 14-field audit row. Every audit row references the manifest version, "
               "policy version, and maturity band of the agent that emitted it. Auditors at any "
               "point in time can determine exactly which agent (and which maturity band) made "
               "any specific decision. Independence-safe by construction across the full arc.",
               accent=GOLD)

    pb(doc)

    # ============ APPENDIX ============
    addh(doc, "Appendix A · Glossary", level=1)
    addtable(doc, ["Term", "Definition"],
             [["Agent",          "A bounded unit of intent + tool-use that emits typed output and writes an audit row per decision"],
              ["Audit row",      "14-field structured record emitted by every agent decision · primary substrate of FSMA evidence and KPI attribution"],
              ["Capability tag", "cap:* identifier in agent manifests · indirection that lets manifests reference capability without binding to specific model URL"],
              ["CAF",            "Converge Agentic Framework · platform substrate that hosts agents and registry"],
              ["Closeout",       "The 14-field audit row written at the end of an agent-orchestration cycle by Evidence-Write"],
              ["Cohort drift",   "Performance gap in agent decisioning detected by the Cohort Drift Watchtower · triggers retraining or maturity review"],
              ["Cross-agent fusion", "Maturation pattern where two or more agents merge into a single superagent at W4+"],
              ["Evidence-Write",  "Agent #6 in RC-E2E-03 · composes 14-field audit row · writes LEDGER · updates FSMA evidence packet"],
              ["HITL",           "Human-in-the-loop · first-class orchestration primitive declared in manifest"],
              ["LEDGER",         "APEX audit-row persistence + retrieval plane · substrate of FSMA evidence and KPI attribution"],
              ["Manifest",       "Git-versioned declarative artifact describing agents, policies, schemas, tenant subscriptions"],
              ["Maturation",     "The structured process by which an agent grows from W1-Specialist to W4+-Superagent"],
              ["Materiality-Decision", "W4+ superagent that absorbs Quantify + Approve · unifies $ quantification with HITL routing"],
              ["Outcome attribution", "Per-trace-id linkage from agent decision to measured outcome · the substrate of self-improvement"],
              ["Risk-Classify",  "W4+ superagent that absorbs Assess + Classify · unifies excursion characterization with disposition classification"],
              ["Specialist",     "W1 maturity band · single-scenario · narrow capability · high HITL"],
              ["Superagent",     "W4+ maturity band · multi-scenario · proactive · principle-based HITL"],
              ["Trace-id",       "Immutable correlation identifier that binds every audit row from one decision event"],
              ["W1 / W2 / W3 / W4+", "Wave bands · also map to agent maturity bands"]],
             col_widths=[1.8, 4.6])

    addh(doc, "Appendix B · Cross-references", level=1)
    addtable(doc, ["Artifact", "Where it complements"],
             [["APEX-RC-E2E-03-Walkthrough.docx", "5-agent storyline (this doc adds Evidence-Write as 6th agent)"],
              ["manifests/parent-orchestrator.yaml", "Will be updated to reference 6th agent in W2; updated again at W4+ fusion"],
              ["manifests/{assess,classify,quantify,approve,act}-agent.yaml", "5 existing manifests · 6th (evidence-write-agent.yaml) to be authored as part of W2"],
              ["APEX-RC-E2E-03-09-AI-ML-Model-Spec.docx", "ML model details for Risk-Score, Elasticity, Disposition Classification"],
              ["APEX-RC-E2E-03-Use-Case-Catalog.xlsx", "UC-19 (Closeout audit row) · UC-22 (Cohort Drift Watchtower)"],
              ["APEX-RC-E2E-03-Kroger-Risk-Register.xlsx", "R-06 (cohort drift) · R-07 (model bias) · R-12 (schema drift)"],
              ["APEX-RC-E2E-03-09-Privacy-Data-Governance-Spec.docx", "Consent state respect across maturation"]],
             col_widths=[3.4, 3.0])

    addbody(doc,
            "Independence posture preserved across this document: 'Deloitte's Microsoft practice,' "
            "'DMTSP,' 'Microsoft platform capabilities,' 'Microsoft-native deployment.' The terms "
            "'partner,' 'alliance,' and 'strategic alliance' do not appear.",
            italic=True, color=MUTE, size=9)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— End of Six-Agents Deep Dive & Maturation —"); setrun(r, italic=True, size=10, color=MUTE, font=SERIF)

    target = OUT
    try: doc.save(str(target))
    except PermissionError:
        n=2
        while True:
            alt = target.with_name(target.stem + f"-v{n}" + target.suffix)
            if not alt.exists(): target = alt; break
            n+=1
        doc.save(str(target))
    print(f"Wrote {target}")


build()
