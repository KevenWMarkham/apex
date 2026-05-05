"""DOCX adapter — Sprint 29 Task 29.9.2.

Uses ``python-docx`` (optional dep). When the dep is absent, falls back
to extracting ``word/document.xml`` from the .docx zip directly via
stdlib ``zipfile`` + ``xml.etree.ElementTree`` so the linter still runs.
"""

from __future__ import annotations

import xml.etree.ElementTree as ET
import zipfile
from pathlib import Path

from apex_compliance_lint.types import FileAdapter


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


class DOCXAdapter(FileAdapter):
    """Extract paragraph text from a .docx file."""

    file_extensions = (".docx",)

    def extract(self, path: Path) -> list[tuple[int, str]]:
        try:
            return self._extract_via_python_docx(path)
        except ImportError:
            return self._extract_via_zip(path)

    @staticmethod
    def _extract_via_python_docx(path: Path) -> list[tuple[int, str]]:
        from docx import Document  # type: ignore[import-not-found]

        doc = Document(str(path))
        out: list[tuple[int, str]] = []
        for i, para in enumerate(doc.paragraphs, 1):
            text = para.text.strip()
            if text:
                out.append((i, text))
        # Also walk tables — common in compliance artifacts
        line_no = len(out) + 1
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        out.append((line_no, text))
                        line_no += 1
        return out

    @staticmethod
    def _extract_via_zip(path: Path) -> list[tuple[int, str]]:
        """Stdlib-only fallback when python-docx isn't installed."""
        try:
            with zipfile.ZipFile(path) as zf:
                with zf.open("word/document.xml") as fh:
                    data = fh.read()
        except (zipfile.BadZipFile, KeyError):
            return []
        try:
            root = ET.fromstring(data)
        except ET.ParseError:
            return []

        out: list[tuple[int, str]] = []
        line_no = 1
        # Each <w:p> is a paragraph; collect its <w:t> text runs in order
        for p in root.iter(f"{{{W_NS}}}p"):
            texts: list[str] = []
            for t in p.iter(f"{{{W_NS}}}t"):
                if t.text:
                    texts.append(t.text)
            joined = "".join(texts).strip()
            if joined:
                out.append((line_no, joined))
                line_no += 1
        return out
