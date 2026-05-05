"""Tier A artifacts — close the end-to-end gaps in the deliverable set.

  1. Solution Architecture Document (DOCX)            — Shared/Tier0-Foundation
  2. Service Operations Playbook · SRE (DOCX)         — Shared/Tier2-Build
  3. Test Strategy & Test Plan (DOCX)                 — Shared/Tier2-Build
  4. Service Roadmap · W1-W4+ (HTML)                  — Shared/Tier4-Strategic
  5. Demo Script & Walkthrough Guide (DOCX)           — RC-E2E-03/Tier3-Governance
"""
from __future__ import annotations
import io, sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
ROOT = Path(r"C:\Stage\Clients\Industries\Consumer\Retail\Kroger\02_projects\FY27_Pipeline\assortment-pricing-agentic\deliverables")
SHARED_T0 = ROOT / "Services" / "Shared-Both-Services" / "Tier0-Foundation"
SHARED_T2 = ROOT / "Services" / "Shared-Both-Services" / "Tier2-Build"
SHARED_T4 = ROOT / "Services" / "Shared-Both-Services" / "Tier4-Strategic"
RC03_T3   = ROOT / "Services" / "RC-E2E-03_Assortment-and-Pricing" / "Tier3-Governance"
for d in (SHARED_T0, SHARED_T2, SHARED_T4, RC03_T3):
    d.mkdir(parents=True, exist_ok=True)

INK="0B0B0F"; INK2="4A4A55"; MUTE="8a8a90"; GOLD="B8892E"; TEAL="2A7F73"
VIO="6648B0"; CRIM="A22A39"; BAND="F3F3F5"; HEAD="0B0B0F"; LINE="D4D4D7"
SERIF="Fraunces"; SANS="IBM Plex Sans"; MONO="JetBrains Mono"


def cell_bg(cell, c):
    tcPr = cell._tc.get_or_add_tcPr()
    s = OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto")
    s.set(qn("w:fill"),c.lstrip("#")); tcPr.append(s)
def cell_border(cell, c=LINE):
    tcPr = cell._tc.get_or_add_tcPr()
    bs = OxmlElement("w:tcBorders")
    for e in ("top","left","bottom","right"):
        b = OxmlElement(f"w:{e}"); b.set(qn("w:val"),"single"); b.set(qn("w:sz"),"4"); b.set(qn("w:color"),c.lstrip("#"))
        bs.append(b)
    tcPr.append(bs)
def setrun(r, *, bold=False, italic=False, size=None, color=None, font=SANS):
    if bold: r.bold = True
    if italic: r.italic = True
    if size is not None: r.font.size = Pt(size)
    if color: r.font.color.rgb = RGBColor.from_string(color.lstrip("#"))
    if font:
        r.font.name = font
        rPr = r._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts"); rPr.insert(0, rFonts)
        for attr in ("w:ascii","w:hAnsi","w:eastAsia","w:cs"):
            rFonts.set(qn(attr), font)
def addbody(doc, t, *, size=11, italic=False, color=INK, sa=6, font=SANS):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(sa)
    r = p.add_run(t); setrun(r, size=size, color=color, italic=italic, font=font); return p
def addh(doc, t, *, level=1):
    sz = {1:18,2:14,3:12.5,4:11.5}
    p = doc.add_paragraph(); p.style = doc.styles[f"Heading {level}"]
    p.paragraph_format.space_before = Pt(16 if level==1 else 10); p.paragraph_format.space_after = Pt(4)
    for r in list(p.runs): r.text = ""
    r = p.add_run(t); setrun(r, bold=True, size=sz.get(level,12), color=INK, font=SERIF)
    if level == 1:
        pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom"); bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),"6"); bot.set(qn("w:space"),"4"); bot.set(qn("w:color"),GOLD)
        pBdr.append(bot); pPr.append(pBdr)
    return p
def addbullets(doc, items, *, size=11):
    for it in items:
        p = doc.add_paragraph(style="List Bullet"); r = p.add_run(it); setrun(r, size=size, color=INK, font=SANS)
def addtable(doc, header, rows, *, col_widths=None):
    n = len(header); tbl = doc.add_table(rows=len(rows)+1, cols=n); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j,h in enumerate(header):
        c = tbl.rows[0].cells[j]; cell_bg(c,HEAD); cell_border(c); c.paragraphs[0].text=""
        r = c.paragraphs[0].add_run(h); setrun(r,bold=True,size=10,color="FFFFFF", font=SANS)
        if col_widths: c.width = Inches(col_widths[j])
    for i,row in enumerate(rows, start=1):
        for j,v in enumerate(row):
            c = tbl.rows[i].cells[j]; cell_border(c)
            if i%2==0: cell_bg(c,BAND)
            c.paragraphs[0].text=""; r = c.paragraphs[0].add_run(str(v)); setrun(r,size=10,color=INK, font=SANS)
            if col_widths: c.width = Inches(col_widths[j])
    doc.add_paragraph(); return tbl
def addcallout(doc, kind, body, *, accent=GOLD, fill="F6F1E4"):
    tbl = doc.add_table(rows=1, cols=1); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = tbl.rows[0].cells[0]; cell_bg(c,fill); c.paragraphs[0].text=""
    rt = c.paragraphs[0].add_run(kind.upper()); setrun(rt,bold=True,size=9,color=accent, font=MONO)
    p = c.add_paragraph(); p.paragraph_format.space_before = Pt(3)
    r = p.add_run(body); setrun(r,size=10.5,color=INK, font=SANS); doc.add_paragraph()
def addkv(doc, pairs, *, lw=1.6, rw=4.9):
    tbl = doc.add_table(rows=len(pairs), cols=2); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,(k,v) in enumerate(pairs):
        ck,cv = tbl.rows[i].cells; ck.width=Inches(lw); cv.width=Inches(rw)
        cell_bg(ck,BAND); cell_border(ck); cell_border(cv)
        ck.paragraphs[0].text=""; cv.paragraphs[0].text=""
        rk = ck.paragraphs[0].add_run(k); setrun(rk,bold=True,size=10,color=INK, font=SANS)
        rv = cv.paragraphs[0].add_run(v); setrun(rv,size=10,color=INK, font=SANS)
    doc.add_paragraph()
def addeyebrow(doc, t):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(t.upper()); setrun(r,bold=True,size=9,color=MUTE, font=MONO)
    rPr = r._element.get_or_add_rPr(); sp = OxmlElement("w:spacing"); sp.set(qn("w:val"),"80"); rPr.append(sp)
def pb(doc):
    p = doc.add_paragraph(); p.add_run().add_break(WD_BREAK.PAGE)
def addcode(doc, t):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(t); setrun(r, size=9.5, color=INK, font=MONO)


def docinit(title, subtitle, meta_kv):
    doc = Document()
    for s in doc.sections:
        s.page_width = Inches(8.5); s.page_height = Inches(11)
        s.top_margin = s.bottom_margin = Inches(0.85); s.left_margin = s.right_margin = Inches(0.95)
    doc.styles["Normal"].font.name = SANS
    doc.styles["Normal"].font.size = Pt(11)
    addeyebrow(doc, title.split(":")[0])
    p = doc.add_paragraph(); r = p.add_run(title); setrun(r, bold=True, size=24, color=INK, font=SERIF)
    p2 = doc.add_paragraph(); r2 = p2.add_run(subtitle); setrun(r2, bold=True, size=14, color=INK2, italic=True, font=SERIF)
    doc.add_paragraph()
    return doc


def docsave(doc, path):
    target = path
    try: doc.save(str(target))
    except PermissionError:
        n=2
        while True:
            alt = target.with_name(target.stem + f"-v{n}" + target.suffix)
            if not alt.exists(): target = alt; break
            n+=1
        doc.save(str(target))
    print(f"  wrote {target.relative_to(ROOT)}")


# =========================================================================
# 1. SOLUTION ARCHITECTURE DOCUMENT (SAD)
# =========================================================================

def build_sad():
    print("Building Solution Architecture Document...")
    doc = docinit("APEX · RC-E2E-03 + RC-E2E-09 · Solution Architecture Document",
                  "The integrated end-to-end architecture · 5 planes · 12 Bronze tables · 6 agents · 19 MCP tools",
                  None)
    addbody(doc, "The integrative 'one ring' architecture document for the APEX RC-E2E-03 + RC-E2E-09 services. "
                 "Brings the existing 22 deliverables into a single coherent architectural narrative. Covers "
                 "the 5 planes, 12 Bronze tables, 6 agents, 19 MCP tools, deployment topology, network "
                 "architecture, security model, and cross-cutting concerns. The single document an Enterprise "
                 "Architect or CTO/CIO can read in 30 minutes to understand the whole.",
            italic=True, color=INK2)
    doc.add_paragraph()
    addkv(doc, [
        ("Document type",     "Solution Architecture Document (SAD)"),
        ("Services",          "RC-E2E-03 Assortment & Pricing · RC-E2E-09 Product Tracking"),
        ("Reference client",  "The Kroger Co."),
        ("Audience",          "Enterprise Architect · CTO · CIO · CISO · Lead Engineer · Practice leadership"),
        ("Companion artifacts","All 22 existing deliverables (this doc references them)"),
        ("Version / date",    "v1.0 · 27 Apr 2026"),
    ])
    pb(doc)

    addh(doc, "1 · Architectural principles", level=1)
    addbody(doc, "Six commitments that bind every architectural decision in this engagement:")
    addbullets(doc, [
        "Manifest-driven, not model-driven — agents, policies, schemas, and tenant subscriptions are Git-versioned manifests. Behavior changes are PRs, not ops tickets.",
        "Audit row, not log line — every agent decision emits a 14-field structured audit row to LEDGER. SOX/FSMA evidence is generated, not collated.",
        "HITL is first-class — human-in-the-loop is an architectural primitive declared in manifest, not bolted on. Material decisions route to named-persona Adaptive Cards.",
        "Capability indirection — agents reference tools via cap:* tags. Backend implementations swap without manifest changes.",
        "Microsoft-native, CAF-rideable — Fabric · Foundry · Teams · Entra · Purview as substrate. CAF Dynamic Orchestrator + A2A Swarm + AEF LLM Gateway as agent runtime.",
        "Independence-safe by construction — agent manifests, policy manifests, and LEDGER are Kroger artifacts. Deloitte's role is advisory and implementation enablement.",
    ])

    addh(doc, "2 · Logical architecture · the 5 planes", level=1)
    addbody(doc, "APEX organizes the system into 5 horizontal planes. Each plane has clear ownership boundaries and "
                 "talks to its neighbors via typed contracts.")
    addtable(doc, ["Plane", "Responsibility", "Key components", "Key artifacts (deliverable)"],
             [["Experience", "Named-persona surfaces · Adaptive Cards · BI dashboards",
               "Microsoft Teams · Power BI Embedded · Adaptive Card runtime",
               "Personas (Tier1) · Pitch Deck (Tier3) · One-pager"],
              ["Decision", "Agents that reason · HITL gates that pause for named-persona approval",
               "6 agents (Assess · Classify · Quantify · Approve · Act · Evidence-Write)",
               "Six-Agents Deep Dive · agent manifests"],
              ["Runtime", "Orchestration graph · policy dispatch · manifest-driven composition",
               "CAF Dynamic Orchestrator · A2A Swarm · parent-orchestrator manifest · materiality policy",
               "Walkthrough · manifests/parent-orchestrator.yaml"],
              ["Integration", "Connectors to client systems-of-record · MCP tools",
               "MCP Server (19 tools) · POS connectors · vendor portals · FSMA QMS",
               "MCP Server Deep Dive · Sequence Diagram"],
              ["Ledger", "14-field audit row store · LEDGER · SOX/FSMA evidence packets",
               "Delta Lake (audit row) · Postgres index · Purview lineage",
               "SOR Bronze-to-Silver · ERD · FSMA Compliance Checklist"]],
             col_widths=[1.2, 1.8, 2.1, 1.8])

    pb(doc)

    addh(doc, "3 · Component diagram (textual)", level=1)
    addbody(doc, "The component-level relationships — what calls what, what reads what.")
    addcode(doc,
            "[Refrigeration EMS]   [POS Systems]   [WMS]   [Vendor Portals]   [Trace One/FoodLogiQ]\n"
            "       │                    │           │            │                   │\n"
            "       └─── Eventstream ────┴─── CDC ───┴─── REST ───┴─── EDI/CDC ───────┘\n"
            "                                       │\n"
            "                                       ▼\n"
            "                           [Bronze · 12 tables · OneLake]\n"
            "                                       │\n"
            "                              (Notebooks · Data Factory)\n"
            "                                       │\n"
            "                                       ▼\n"
            "                       [Silver · MERML + SCML + DIM]\n"
            "                                       │\n"
            "                                       ▼\n"
            "                            [MCP Server · 19 tools]\n"
            "                                       ▲\n"
            "                                       │ (capability resolution)\n"
            "                                       │\n"
            "         ┌──────┬──────┬──────┬──────┬──────┬──────┐\n"
            "      ASSESS  CLASSIFY  QUANTIFY  APPROVE  ACT  EVIDENCE-WRITE\n"
            "         └──────┴──────┴──────┴──────┴──────┴──────┘\n"
            "                                       │\n"
            "                              [Parent Orchestrator]\n"
            "                                  (CAF Dynamic)\n"
            "                                       │\n"
            "                              ┌────────┴────────┐\n"
            "                              ▼                 ▼\n"
            "                     [LEDGER (Delta + PG)]  [Adaptive Card → Teams]\n"
            "                              │\n"
            "                              ▼\n"
            "                     [Power BI · FSMA Packet · Purview Lineage]\n")

    addh(doc, "4 · Data architecture", level=1)
    addbody(doc, "Three-tier medallion in Microsoft Fabric. See SOR Bronze-to-Silver doc and ERD HTML for the "
                 "exhaustive field-level specification.")
    addtable(doc, ["Tier", "Role", "Schema family"],
             [["Bronze", "Raw, source-aligned, append-only · 12 tables", "Source-shaped + ingest metadata (5 fields)"],
              ["Silver", "Conformed canonical · 12 shapes (4 MERML + 5 SCML + 3 DIM)", "MERML.* + SCML.* + DIM.*"],
              ["Gold",   "Curated KPI marts · semantic models · dashboards", "Subject-area marts (Direct Lake)"],
              ["LEDGER", "Audit-row store (parallel to Gold)", "14-field audit row (immutable append)"]],
             col_widths=[1.0, 3.0, 2.5])

    addh(doc, "5 · Deployment topology · Azure", level=1)
    addh(doc, "5.1 · Subscription and tenancy", level=2)
    addtable(doc, ["Layer", "Approach"],
             [["Tenant",            "Single Kroger Entra tenant · agent service principals scoped per Entra group"],
              ["Subscriptions",     "3 subscriptions: apex-rc-dev · apex-rc-uat · apex-rc-prod"],
              ["Resource groups",   "Per-workspace + per-shared-service (networking, IAM, monitoring)"],
              ["Region",            "Primary: East US 2 · Secondary (DR): West US 2"],
              ["Fabric capacity",   "Dev F32 · UAT F64 · Prod F128 (W2) → F256 (W3)"]],
             col_widths=[1.6, 5.0])

    addh(doc, "5.2 · Network architecture", level=2)
    addbullets(doc, [
        "Hub-spoke topology · hub holds shared services (Entra Connect, Purview, Defender, monitoring)",
        "Spoke per workspace · private endpoints for OneLake, Foundry, AOAI, PostgreSQL",
        "Express Route to Kroger on-prem (refrigeration EMS · POS · WMS) · 1 Gbps primary, 100 Mbps backup",
        "Private DNS zones for all PaaS endpoints · no public ingress to data plane",
        "Azure Front Door + WAF for any internet-facing surfaces (none planned for W1-W2)",
    ])

    addh(doc, "5.3 · DR and high availability", level=2)
    addtable(doc, ["Component", "RPO", "RTO", "Strategy"],
             [["Bronze (OneLake)",     "≤ 15 min",  "≤ 4 hr",  "Geo-redundant Delta replication · cross-region"],
              ["Silver (OneLake)",     "≤ 1 hr",    "≤ 8 hr",  "Rebuilt from Bronze on warm-failover"],
              ["LEDGER PostgreSQL",    "≤ 5 min",   "≤ 30 min","Active geo-replication · failover group"],
              ["Eventstream",          "≤ 2 min",   "≤ 15 min","Active-active across regions"],
              ["MCP Server",           "—",         "≤ 5 min", "Stateless · auto-scale · zone-redundant"]],
             col_widths=[1.8, 1.0, 1.0, 2.8])

    pb(doc)

    addh(doc, "6 · Security architecture", level=1)
    addh(doc, "6.1 · Identity model", level=2)
    addbullets(doc, [
        "Microsoft Entra ID for all identities · service principals per agent · group-based RBAC",
        "Privileged Identity Management (PIM) for production write access · time-bounded · approval workflow",
        "Conditional Access for human personas · MFA required for HITL approvers",
        "Workload Identity Federation between Fabric and Foundry · no shared secrets",
    ])

    addh(doc, "6.2 · Data classification + Purview", level=2)
    addtable(doc, ["Sensitivity", "Applied to", "Controls"],
             [["Public",             "Static reference data (DIM.Calendar, DIM.Banner)", "—"],
              ["Internal",           "Promo/markdown commercial data",                    "Internal RBAC"],
              ["Confidential",       "Pricing · POS transactions · MERML.*",              "Per-banner RBAC + per-role scope"],
              ["Restricted",         "Loyalty households · personalization decisions",   "Consent-aware filtering · Privacy Office audit"],
              ["FSMA-204",          "Lot-trace · food-safety records · LEDGER FSMA rows", "Auditor + FDA-relations only · indefinite retention"],
              ["Audit-Critical",     "LEDGER audit rows · evidence packets",              "Immutable · long-retention · auditor-accessible"]],
             col_widths=[1.4, 2.6, 2.6])

    addh(doc, "6.3 · Encryption + key management", level=2)
    addbullets(doc, [
        "All data encrypted at rest with Microsoft-managed keys (default) · Customer-managed keys via Key Vault for FSMA-204 + Audit-Critical",
        "TLS 1.3 for all in-flight · mTLS between MCP server and agents",
        "Secrets in Azure Key Vault · referenced by manifest via Key Vault secret URIs",
        "Per-tenant key scope · key rotation per Kroger key-management policy",
    ])

    addh(doc, "7 · Integration architecture", level=1)
    addbody(doc, "Source systems → Bronze → Silver → MCP tools → agents. All inbound integrations are read-only "
                 "in W1; mutation surfaces (POS, vendor portals, FSMA QMS) light up in W2 via the Act agent.")
    addtable(doc, ["Source system", "Integration pattern", "Wave"],
             [["Refrigeration EMS (Emerson/Honeywell/Hill PHOENIX)", "OPC-UA → MQTT → Eventstream", "W1"],
              ["POS (Toshiba/IBM/NCR)",                              "CDC → Eventstream",            "W1 (read) · W2 (write)"],
              ["Pricing engine (Revionics + 84.51°)",                "CDC + nightly snapshot",       "W1"],
              ["WMS (Manhattan)",                                    "CDC → Eventstream",            "W1"],
              ["Vendor portals (Emerson/Lennox/Hussmann/Hill PHOENIX)", "REST polling 15-min",      "W1 (read) · W2 (write)"],
              ["QMS / Trace One (food-safety)",                      "CDC → Eventstream",            "W2"],
              ["Lot-trace (Trace One/FoodLogiQ/EDI)",                "Mixed stream + batch",         "W2"],
              ["84.51° loyalty platform",                            "Daily snapshot",               "W1"],
              ["Microsoft Teams (Adaptive Card)",                    "Bi-directional · WebSocket",   "W1 mock · W2 live"]],
             col_widths=[2.6, 2.4, 1.6])

    pb(doc)

    addh(doc, "8 · Cross-cutting concerns", level=1)
    addh(doc, "8.1 · Observability", level=2)
    addbullets(doc, [
        "OpenTelemetry traces · W3C traceparent propagation across agents → MCP → backends",
        "LangFuse-equivalent span store for agent reasoning traces",
        "Application Insights + Log Analytics for application logs",
        "Microsoft Purview for data lineage (Bronze → Silver → tool → agent → outcome)",
        "Power BI semantic models for KPI rollups · drill-through to LEDGER trace_id",
    ])

    addh(doc, "8.2 · Compliance posture", level=2)
    addtable(doc, ["Regulatory regime", "Coverage", "Reference"],
             [["FSMA 204 (FDA)",        "Lot-traceability via RC-E2E-09 · CTE/KDE capture · 24-month retention",
                                        "FSMA 204 Compliance Checklist"],
              ["GDPR / CCPA (privacy)", "Consent state machine · WITHDRAWN-redaction · Purview audit",
                                        "Privacy & Data Governance Spec"],
              ["SOX-adjacent",          "14-field audit row · manifest version stamps · evidence packet",
                                        "AI/ML Model Spec · LEDGER architecture"],
              ["Independence (audit)",  "Manifests/policies/LEDGER are Kroger artifacts · Deloitte advisory",
                                        "All artifacts (consistent posture)"]],
             col_widths=[1.8, 3.2, 2.0])

    addh(doc, "8.3 · Versioning policy", level=2)
    addbullets(doc, [
        "Manifests semver versioned · PATCH auto-deploy · MINOR + MAJOR via PR-merge with approver gate",
        "Capability tags (cap:*) versioned independently · backend swaps require no manifest change",
        "Schema evolution additive only at Bronze · breaking changes happen in Silver canonical layer",
        "Maturity-band promotion (W1→W4+) · Steering Committee gate · evidence requirements per band",
    ])

    addh(doc, "9 · Architectural decisions log (key ADRs)", level=1)
    addtable(doc, ["ADR", "Decision", "Rationale"],
             [["ADR-001",  "Microsoft Fabric over Snowflake/Databricks",
                            "Native Microsoft estate alignment · Direct Lake · Purview lineage built-in"],
              ["ADR-002",  "MCP over direct API integration",
                            "Capability indirection · backend swappability · CAF substrate alignment"],
              ["ADR-003",  "6 agents (split Evidence-Write from Act)",
                            "Different blast-radius · different SLO · different evolution arc"],
              ["ADR-004",  "Manifests as Kroger artifacts (not Deloitte)",
                            "Independence preservation · audit-firm-agnostic engagement structure"],
              ["ADR-005",  "Eventstream over Logic Apps for streaming feeds",
                            "Low-latency requirement (≤ 30s for telemetry) · native Fabric integration"],
              ["ADR-006",  "Postgres index alongside Delta for LEDGER",
                            "Sub-100ms trace_id lookup · Delta-only insufficient for Power BI drill-through"],
              ["ADR-007",  "Per-cohort elasticity Bayesian model",
                            "Cold-start handling for new SKUs · category-level prior pooling"],
              ["ADR-008",  "Tier-based materiality policy (3 tiers × banner-aware)",
                            "Balances HITL volume against approval-quality · banner-aware reflects business reality"],
              ["ADR-009",  "Synthetic-data shadow-run before live W2",
                            "De-risks model accuracy assumptions · auditor-walkthrough rehearsal"],
              ["ADR-010",  "Parent-orchestrator pattern (vs. peer-to-peer agents)",
                            "Single coordinator simplifies trace propagation and HITL gating"]],
             col_widths=[1.0, 2.4, 3.0])

    pb(doc)

    addh(doc, "10 · How this SAD relates to the other 22 deliverables", level=1)
    addbody(doc, "This SAD is the integrative document. The table below maps SAD sections to the deliverables that "
                 "deep-dive each topic.")
    addtable(doc, ["SAD section", "Deep-dive deliverable(s)"],
             [["§2 5 planes",                  "Walkthrough · Service Portfolio"],
              ["§3 Component diagram",          "Sequence Diagram · Six-Agents Deep Dive"],
              ["§4 Data architecture",          "SOR Bronze-to-Silver · ERD HTML · AI/ML Model Spec"],
              ["§5 Deployment topology",        "Fabric Implementation Runbook · this SAD §5"],
              ["§6 Security architecture",      "Privacy & Data Governance Spec · Risk Register"],
              ["§7 Integration architecture",   "MCP Server Deep Dive · Sequence Diagram"],
              ["§8.1 Observability",            "Fabric Runbook §6 · MCP Deep Dive §6"],
              ["§8.2 Compliance",               "FSMA 204 Compliance Checklist · Privacy Spec"],
              ["§8.3 Versioning",               "Six-Agents Deep Dive §6 · MCP Deep Dive §3.4"],
              ["§9 ADRs",                       "(this SAD is the ADR home)"],
              ["Operations",                    "Service Operations Playbook (Tier A · companion)"],
              ["Testing",                       "Test Strategy & Test Plan (Tier A · companion)"],
              ["Roadmap",                       "Service Roadmap HTML (Tier A · companion)"],
              ["Demo",                          "Demo Script & Walkthrough Guide (Tier A · companion)"]],
             col_widths=[2.4, 4.0])

    addh(doc, "Appendix A · Reference architecture diagrams", level=1)
    addbody(doc, "Visual companions to this SAD live in adjacent HTML deliverables:")
    addbullets(doc, [
        "SOR ERD (data layer): APEX-RC-E2E-03-09-SOR-ERD.html",
        "Service Sequence (agent ↔ MCP ↔ data): APEX-RC-E2E-03-Service-Sequence-Diagram.html",
        "Service Roadmap (forward-looking): APEX-RC-E2E-03-Service-Roadmap.html (Tier A · this batch)",
    ])

    addbody(doc, "Independence posture preserved: 'Deloitte's Microsoft practice,' 'DMTSP,' 'Microsoft platform "
                 "capabilities,' 'Microsoft-native deployment.' The terms 'partner,' 'alliance,' and 'strategic "
                 "alliance' do not appear.",
            italic=True, color=MUTE, size=9)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— End of Solution Architecture Document —"); setrun(r, italic=True, size=10, color=MUTE, font=SERIF)

    docsave(doc, SHARED_T0 / "APEX-RC-E2E-03-09-Solution-Architecture-Document.docx")


# =========================================================================
# 2. SERVICE OPERATIONS PLAYBOOK (SRE)
# =========================================================================

def build_ops_playbook():
    print("\nBuilding Service Operations Playbook...")
    doc = docinit("APEX · RC-E2E-03 + RC-E2E-09 · Service Operations Playbook",
                  "Day-2 operations · SLOs · monitoring · on-call · incident response · 12 runbooks",
                  None)
    addbody(doc, "Operational playbook for the SRE / Operations team running the APEX RC-E2E-03 + RC-E2E-09 "
                 "services in production. Defines SLOs/SLIs, monitoring & alerting strategy, on-call rotation "
                 "and escalation paths, incident classification (P0-P4), 12 named runbooks for common scenarios, "
                 "and the postmortem process. Pairs with the Fabric Implementation Runbook.",
            italic=True, color=INK2)
    doc.add_paragraph()
    addkv(doc, [
        ("Document type",      "Service Operations Playbook · SRE"),
        ("Services",           "RC-E2E-03 + RC-E2E-09"),
        ("Reference client",   "The Kroger Co."),
        ("Audience",           "Operations · SRE · Lead Engineer · On-call rotation · Engineering Manager"),
        ("Companion",          "Fabric Implementation Runbook · Risk Register · MCP Server Deep Dive"),
        ("Version / date",     "v1.0 · 27 Apr 2026"),
    ])
    pb(doc)

    addh(doc, "1 · Service Level Objectives (SLOs)", level=1)
    addbody(doc, "Six headline SLOs gate the engagement. Each has an SLI definition, a target, and an alerting threshold "
                 "tied to the error budget.")
    addtable(doc, ["SLO", "SLI definition", "Target (W2)", "Target (W3)", "Alert at"],
             [["Decision-loop time",        "Sec from event to disposition", "≤ 4 min p95",  "≤ 90 sec p95",  "p95 > target for 30 min"],
              ["Audit-row coverage",        "% material decisions attributed", "≥ 98%",       "≥ 99.5%",       "< target on rolling 1 hr"],
              ["MCP Server availability",   "% successful tool invocations",   "≥ 99.9%",     "≥ 99.95%",      "< target on rolling 5 min"],
              ["Bronze-landing latency",    "Sec from source commit to Bronze","≤ 60 sec p99","≤ 30 sec p99",  "p99 > target for 10 min"],
              ["Refrig telemetry coverage", "% fixtures with ≥ 99% reads/day", "≥ 95% fixtures","≥ 99.5%",     "< target daily aggregate"],
              ["FSMA evidence packet time", "Hr from quarter-end to packet",   "≤ 72 hr",     "≤ 24 hr",       "Trend-based · weekly review"]],
             col_widths=[1.7, 2.0, 1.0, 1.0, 1.5])

    addh(doc, "1.1 · Error budget policy", level=2)
    addbullets(doc, [
        "Each SLO has a 30-day error budget · 0.05% for 99.95% SLOs · 0.5% for 99.5% SLOs · 0.1% for 99.9% SLOs",
        "When 50% of error budget consumed in <14 days: feature freeze for affected component",
        "When 100% consumed: incident-priority work only · postmortem required before resuming features",
        "Steering Committee notified at 75% consumption (rolling 30-day)",
    ])

    pb(doc)

    addh(doc, "2 · Monitoring & alerting", level=1)
    addh(doc, "2.1 · Tool stack", level=2)
    addtable(doc, ["Layer", "Tool", "What it monitors"],
             [["Infrastructure",  "Azure Monitor + Log Analytics", "VM/container/Fabric capacity · network · storage"],
              ["Application",     "Application Insights",          "Service availability · request rates · errors · dependencies"],
              ["Tracing",         "OpenTelemetry + LangFuse-equiv","Distributed traces · agent reasoning chains"],
              ["Lineage",         "Microsoft Purview",             "Bronze→Silver→tool→agent→outcome lineage"],
              ["Business",        "Power BI · LEDGER queries",     "KPI trajectory · audit-row coverage · shrink avoided"],
              ["Synthetic",       "Custom probes",                  "MCP tool round-trips · LEDGER write probes · end-to-end synthetic excursion"]],
             col_widths=[1.5, 2.0, 3.1])

    addh(doc, "2.2 · Alert classifications", level=2)
    addtable(doc, ["Severity", "Definition", "Response", "Notification"],
             [["P0 · Critical",  "Service down · audit-row coverage drops > 5% in 1 hr · safety-relevant",
                                  "Immediate · all hands",       "Page on-call · auto-page Engineering Lead + DMTSP_lead"],
              ["P1 · High",      "Major SLO breach · single-tool down · LEDGER write failures",
                                  "≤ 15 min response",            "Page on-call"],
              ["P2 · Medium",    "Performance degradation · single-banner impact",
                                  "≤ 1 hr response · within business hours OK", "Email + Slack alert"],
              ["P3 · Low",       "Single-event anomaly · drift detection · capacity warning",
                                  "Same business day",            "Slack alert · daily standup"],
              ["P4 · Info",      "FYI · trend warning",
                                  "Backlog item",                 "Dashboard only"]],
             col_widths=[1.4, 2.4, 1.6, 1.6])

    addh(doc, "2.3 · Critical alerts (P0/P1)", level=2)
    addtable(doc, ["Alert", "Trigger", "Severity"],
             [["LEDGER write failure",          "5xx from cap:ledger-write · success rate < 99% over 5 min", "P0"],
              ["MCP Server down",               "Health check failure · 3 consecutive failures",            "P0"],
              ["Refrigeration telemetry stalled","Bronze-landing rate drops > 50% from baseline · 5 min",   "P0"],
              ["FSMA-evidence-gap detected",    "Audit-row coverage < 95% · rolling 1 hr",                  "P0"],
              ["Eventstream consumer lag",      "Lag > 5 min · sustained 10 min",                            "P1"],
              ["POS markdown failure",          "cap:pos-markdown-upload error rate > 5% · 15 min",         "P1"],
              ["Adaptive Card delivery failure","cap:teams-adaptive-card delivery rate < 95% · 30 min",     "P1"],
              ["Cohort drift detected",         "Risk-Score precision drops > 5pp from baseline",           "P1"],
              ["Disposition confidence collapse","Avg confidence drops > 10pp · 1 hr",                       "P1"],
              ["FSMA packet generation failure","Quarterly job error · before deadline",                     "P1"]],
             col_widths=[2.4, 3.5, 0.7])

    pb(doc)

    addh(doc, "3 · On-call rotation & escalation", level=1)
    addh(doc, "3.1 · Rotation structure", level=2)
    addtable(doc, ["Role", "Coverage", "Escalation timer"],
             [["Primary on-call (engineer)",     "24/7 · weekly rotation · 5-engineer pool", "0–15 min"],
              ["Secondary on-call (engineer)",   "24/7 · weekly · escalation backup",        "15–30 min"],
              ["Engineering Manager",            "Business hours + on-call paged for P0",   "30–45 min"],
              ["Engagement Architect",           "Business hours + paged for P0",            "45 min"],
              ["DMTSP_lead",                     "Notified on P0",                           "P0 only"],
              ["Customer Engineering (Kroger IT)","Business hours · P0 cross-paged",         "P0 cross-org"]],
             col_widths=[2.4, 2.4, 1.8])

    addh(doc, "3.2 · Escalation paths", level=2)
    addbullets(doc, [
        "P0: primary → secondary at 15 min if no ack → Engineering Manager at 30 min → DMTSP_lead at 45 min · simultaneous Kroger Customer Eng page if cross-org",
        "P1: primary → secondary at 30 min if no ack → Engineering Manager at business hours",
        "P2: primary handles · escalates to Engineering Manager only if backlog > 24 hr",
        "All P0/P1: postmortem required within 5 business days",
    ])

    pb(doc)

    addh(doc, "4 · Twelve named runbooks", level=1)
    addbody(doc, "The 12 most-likely operational scenarios · each has a named runbook with detection, "
                 "diagnostic commands, recovery steps, and postmortem template.")

    runbooks = [
        ("RB-01 · LEDGER write failure",
         "Symptoms: cap:ledger-write 5xx · audit-row coverage SLO at risk",
         "1. Confirm via Application Insights query · 2. Check Postgres connection (failover group health) · "
         "3. Check Delta write logs · 4. If PG primary down: failover to secondary · 5. If Delta down: "
         "switch to PG-only mode (degraded) · 6. Backfill failed rows from local agent buffers · 7. Postmortem"),
        ("RB-02 · MCP Server unhealthy",
         "Symptoms: Health-check failures · agent invocations timing out",
         "1. Check pod/container logs · 2. Restart MCP server (auto-scale group) · 3. Validate registry health · "
         "4. If schema-validation errors: check recent capability deployments · 5. Roll back recent capability "
         "version if root cause · 6. Notify agents of capability availability via parent-orchestrator"),
        ("RB-03 · Refrigeration telemetry stalled",
         "Symptoms: Bronze-landing rate drops > 50% · per-fixture coverage SLO breached",
         "1. Check Eventstream consumer health · 2. Validate MQTT bridge connectivity · 3. Check EMS gateway "
         "connectivity at affected stores · 4. If multi-store: page Facilities Eng · 5. Quarantine affected "
         "fixture from agent decisioning until resolved · 6. Backfill via REST polling fallback"),
        ("RB-04 · Adaptive Card delivery failure",
         "Symptoms: cap:teams-adaptive-card delivery rate drops · approver complaints",
         "1. Check Microsoft Teams service health · 2. Check Entra identity resolution · 3. Validate device "
         "tokens for affected approvers · 4. If isolated: cross-page with Microsoft support · 5. Activate "
         "fallback channel (email + designate-approver) · 6. Resume primary channel post-recovery"),
        ("RB-05 · POS markdown upload failure",
         "Symptoms: cap:pos-markdown-upload error rate > 5%",
         "1. Identify affected banner (Toshiba/IBM/NCR cluster) · 2. Check POS gateway health · 3. "
         "Validate API rate-limit utilization · 4. Queue failed markdowns for batch retry · 5. Notify "
         "Store Operations Lead at affected stores · 6. Manual markdown fallback if persistent"),
        ("RB-06 · FSMA evidence-gap detected",
         "Symptoms: Audit-row coverage drops below 95% on FSMA-relevant decisions",
         "1. CRITICAL · page Food Safety Lead immediately · 2. Identify gap via LEDGER query · 3. "
         "Reconstruct from Bronze + agent traces · 4. Manual evidence backfill if reconstruction "
         "incomplete · 5. Notify Internal Audit · 6. Mandatory postmortem · 7. Privacy Office review"),
        ("RB-07 · Cohort drift detected (Watchtower alert)",
         "Symptoms: Risk-Score precision/recall drops > 5pp from baseline",
         "1. Validate alert via Drift Watchtower dashboard · 2. Check feature distribution (PSI > 0.1 confirms drift) · "
         "3. Trigger shadow-retrain pipeline · 4. Run champion-challenger A/B for 14 days · 5. Promote "
         "challenger if performance restored · 6. If not: escalate to ML Engineering for feature investigation"),
        ("RB-08 · Eventstream consumer lag",
         "Symptoms: Watermark delays > 5 min · streaming-feed lag warnings",
         "1. Check Eventstream throughput vs. capacity · 2. Scale capacity if at limit · 3. Check downstream "
         "consumer health (Notebooks, Data Factory) · 4. Validate partition skew (rebalance if needed) · "
         "5. Resume normal operation · 6. Capacity-planning review if recurring"),
        ("RB-09 · Bronze schema drift detected",
         "Symptoms: DQ pack rejections at land-time · auto-PR opened",
         "1. Review auto-PR · 2. Validate schema-additive (no breaking changes) · 3. Approve & merge if additive · "
         "4. If breaking: quarantine batch · escalate to source-system owner · 5. Coordinate Bronze MAJOR-version "
         "bump with consumers · 6. Resume ingest"),
        ("RB-10 · Vendor portal API timeout",
         "Symptoms: cap:vendor-portal-ticket failures · ticket creation backlog",
         "1. Identify affected vendor (Emerson/Lennox/Hussmann/Hill PHOENIX) · 2. Validate API status with "
         "vendor account manager · 3. Queue tickets for batch retry · 4. Activate vendor-side manual fallback "
         "if persistent · 5. Refrigeration vendor SLA review post-incident"),
        ("RB-11 · LEDGER PostgreSQL lag",
         "Symptoms: trace_id lookup latency > 200ms · drill-through slow",
         "1. Check PG replication lag · 2. Validate index health · 3. If lag > 30 sec: failover to read-replica "
         "for queries · 4. Reindex if fragmentation root cause · 5. Capacity review if recurring"),
        ("RB-12 · Recall event triggered",
         "Symptoms: FDA or supplier-initiated recall on FTL category",
         "1. CRITICAL · page Food Safety Lead + Engineering Architect · 2. SCML.Lot recall-window query · "
         "3. Auto-trigger Act agent for shelf-pull at affected stores · 4. Generate FSMA evidence packet · "
         "5. External communications coordination · 6. Mandatory after-action review"),
    ]
    for title, sym, steps in runbooks:
        addh(doc, title, level=2)
        addbody(doc, sym, italic=True, color=INK2, size=10.5)
        addbody(doc, steps, size=10.5)

    pb(doc)

    addh(doc, "5 · Postmortem process", level=1)
    addbullets(doc, [
        "Required for all P0 incidents · optional for P1 (Engineering Manager discretion)",
        "Blameless postmortem · focus on systems and process not individuals",
        "Timeline reconstruction · root-cause analysis (5-whys minimum) · contributing factors",
        "Action items · each with owner · target date · tracked in engagement risk register",
        "Distribution · DMTSP_lead · Steering Committee · cross-org if Kroger IT involved",
        "Public-blameless summary · anonymized · added to engagement runbook for future learning",
    ])

    addh(doc, "6 · Capacity management", level=1)
    addbody(doc, "Per the Fabric Implementation Runbook, capacity is managed at three layers:")
    addbullets(doc, [
        "Fabric F-SKU capacity · scaled per-wave (F32 dev · F64 UAT · F128 W2 · F256 W3) · monitored daily",
        "Eventstream throughput · partition-aware · scaled per banner volume",
        "Postgres LEDGER index · query-volume-aware · pre-scaled before each wave promotion",
        "Quarterly capacity review · Steering Committee gates W2→W3 capacity uplift",
    ])

    addh(doc, "7 · Change management", level=1)
    addbullets(doc, [
        "All production changes via PR · 2-reviewer approval (Engineering Lead + Engagement Architect)",
        "Manifest changes follow PATCH/MINOR/MAJOR policy · MAJOR requires Steering Committee gate",
        "Tool capability deployments via blue-green · 14-day champion-challenger for ML model promotions",
        "Rollback procedures documented per change · 5-min RTO for all production changes",
        "Change Advisory Board (CAB) reviews weekly · cross-org with Kroger Customer Eng",
    ])

    addbody(doc, "Independence posture preserved.", italic=True, color=MUTE, size=9)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— End of Service Operations Playbook —"); setrun(r, italic=True, size=10, color=MUTE, font=SERIF)

    docsave(doc, SHARED_T2 / "APEX-RC-E2E-03-09-Service-Operations-Playbook.docx")


# =========================================================================
# 3. TEST STRATEGY & TEST PLAN
# =========================================================================

def build_test_plan():
    print("\nBuilding Test Strategy & Test Plan...")
    doc = docinit("APEX · RC-E2E-03 + RC-E2E-09 · Test Strategy & Test Plan",
                  "Test pyramid · pytest + Testcontainers harness · UAT · performance · acceptance gates",
                  None)
    addbody(doc, "End-to-end test strategy and test plan for RC-E2E-03 + RC-E2E-09. Defines the test pyramid "
                 "(unit → integration → E2E → UAT → performance), the apex-test-harness pattern (pytest + "
                 "Testcontainers + synthetic data), per-UC acceptance criteria, release gates, and the "
                 "synthetic-vs-live data strategy across waves.",
            italic=True, color=INK2)
    doc.add_paragraph()
    addkv(doc, [
        ("Document type",      "Test Strategy & Test Plan"),
        ("Services",           "RC-E2E-03 + RC-E2E-09"),
        ("Audience",           "QA Lead · Lead Engineer · Engagement Architect · Customer UAT team"),
        ("Companion",          "Use Case Catalog · Service Operations Playbook · AI/ML Model Spec"),
        ("Version / date",     "v1.0 · 27 Apr 2026"),
    ])
    pb(doc)

    addh(doc, "1 · Test pyramid", level=1)
    addbody(doc, "Five layers · increasing scope · decreasing volume.")
    addtable(doc, ["Layer", "Scope", "Volume target", "Tools"],
             [["Unit",        "Single function · single class · pure logic", "~1,000 tests · runs < 60 sec",
               "pytest · Vitest (frontend)"],
              ["Integration", "Single agent · single MCP tool · single Bronze ingest path", "~200 tests · runs < 5 min",
               "pytest + Testcontainers (Postgres · Eventstream emulator)"],
              ["End-to-end",  "Full agent chain · Assess → Evidence-Write", "~30 scenarios · runs < 30 min",
               "pytest + synthetic-data lab + LEDGER probe"],
              ["UAT",         "Per-UC acceptance · stakeholder-driven", "24 UCs · ~3 acceptance tests each",
               "Manual + scripted · in-product"],
              ["Performance", "Throughput · latency · capacity", "Per-SLO load test · weekly run",
               "k6 · Locust · custom synthetic excursion generator"]],
             col_widths=[1.2, 2.0, 1.8, 2.0])

    addh(doc, "2 · The apex-test-harness", level=1)
    addbody(doc, "The reusable test harness that anchors integration + E2E layers. pytest + Testcontainers + "
                 "synthetic-data lab. Reusable across reference clients with manifest-layer adaptation.")
    addh(doc, "2.1 · Architecture", level=2)
    addbullets(doc, [
        "pytest as test runner · fixtures bring up Testcontainers (Postgres, MQTT broker, Eventstream emulator)",
        "Synthetic-data lab provides 100K subscriber population + 4 quarters of realistic events (per Fabric Runbook §11.1)",
        "Mock MCP server with deterministic stubs for tools that touch external systems",
        "Real LEDGER (Postgres + Delta-equivalent local) for audit-row validation",
        "Real model-serving for ML inference tools (cap:risk-score, cap:disposition-classification)",
        "Per-test trace_id isolation · no cross-test contamination",
    ])

    addh(doc, "2.2 · Sample harness scaffold", level=2)
    addcode(doc,
            "# tests/conftest.py\n"
            "import pytest\n"
            "from testcontainers.postgres import PostgresContainer\n"
            "from apex.mcp_server import MCPServer\n"
            "from apex.synthetic import SyntheticExcursionLab\n\n"
            "@pytest.fixture(scope='session')\n"
            "def ledger_pg():\n"
            "    with PostgresContainer('postgres:16') as pg:\n"
            "        yield pg.get_connection_url()\n\n"
            "@pytest.fixture(scope='session')\n"
            "def synthetic_lab():\n"
            "    return SyntheticExcursionLab(banner='KR', stores=50, fixtures_per_store=8)\n\n"
            "@pytest.fixture\n"
            "def mcp(ledger_pg, synthetic_lab):\n"
            "    server = MCPServer(\n"
            "        manifest_dir='manifests/',\n"
            "        ledger_url=ledger_pg,\n"
            "        synthetic=synthetic_lab,\n"
            "    )\n"
            "    yield server\n"
            "    server.shutdown()\n\n"
            "# tests/test_e2e_cold_chain.py\n"
            "def test_excursion_to_closeout_within_90s(mcp, synthetic_lab):\n"
            "    excursion = synthetic_lab.fire_excursion(\n"
            "        store='KR-428', fixture='DAIRY-A3',\n"
            "        peak_temp_f=48.2, duration_min=47\n"
            "    )\n"
            "    closeout = mcp.run_orchestrator(excursion, timeout_sec=90)\n"
            "    \n"
            "    assert closeout.audit_row.kpis['time_to_decision_sec'] < 90\n"
            "    assert closeout.audit_row.fsma_audit_row_complete\n"
            "    assert closeout.audit_row.shrink_avoided_usd > 0")

    pb(doc)

    addh(doc, "3 · Per-UC acceptance criteria mapping", level=1)
    addbody(doc, "Each of the 24 UCs in the Use Case Catalog has explicit acceptance criteria. The table below "
                 "maps UC IDs to test categories (UAT-eligible vs. integration-only).")
    addtable(doc, ["UC", "Title", "Test layer", "Acceptance pattern"],
             [["UC-01", "Refrigeration telemetry ingest", "Integration + Performance", "30s latency · 99.5% completeness"],
              ["UC-02", "Excursion-detection trigger",     "Integration", "Fires within 30s · false-positive < 5%"],
              ["UC-03", "Assess agent characterization",   "Integration + UAT", "Severity classification ≥ 95% match"],
              ["UC-04", "SCML.Lot canonical layer",        "Integration", "FTL items have non-null lot_code · 99.5%"],
              ["UC-05", "FSMA 204 policy corpus",          "UAT (Food Safety Lead)", "Policy v1.0 sign-off"],
              ["UC-06", "Classify agent disposition",      "Integration + UAT + ML eval", "Match ≥ 95% · escalation working"],
              ["UC-07", "Confidence threshold tuning",     "ML eval + UAT",  "Calibration plot · 5-10% escalation rate"],
              ["UC-08", "MERML.Price canonical layer",     "Integration",    "Lookup ≤ 100ms · zone resolution correct"],
              ["UC-09", "MERML.Markdown derived layer",    "Integration",    "Outcome attribution · NPV match"],
              ["UC-10", "MERML.Elasticity model fit",      "ML eval",        "Coefficients in range · drift detection"],
              ["UC-11", "Quantify agent",                  "Integration + UAT", "±5% match Pricing Strategy reference"],
              ["UC-12", "Markdown-NPV outcome",            "E2E",            "Trace_id roll-up to dashboard"],
              ["UC-13", "Materiality policy",              "UAT (Pricing + VP)", "Policy sign-off"],
              ["UC-14", "Approve agent routing",           "E2E + UAT",      "≥ 99% routing correctness"],
              ["UC-15", "Adaptive Card delivery",          "E2E + UAT",      "≤ 5s delivery · Entra binding · audit row"],
              ["UC-16", "POS markdown upload",             "Integration + E2E", "≤ 2 min · idempotent on retry"],
              ["UC-17", "Vendor-portal ticket",            "Integration + E2E", "Ticket within 5 min · ticket_id captured"],
              ["UC-18", "FSMA regulatory log",             "Integration + UAT", "100% on closeouts"],
              ["UC-19", "LEDGER 14-field audit row",       "Integration",    "≥ 99.5% · all fields populated"],
              ["UC-20", "Power BI dashboard",              "UAT",            "Drill-through to trace_id"],
              ["UC-21", "FSMA evidence packet generator",  "E2E + UAT (Privacy + Audit)", "≤ 4 hr · format approved"],
              ["UC-22", "Cohort segmentation engine",      "ML eval + Integration", "Drift watchtower alerting"],
              ["UC-23", "Multi-banner orchestration",      "E2E (W3)",       "Cross-banner aggregation"],
              ["UC-24", "Cross-scenario fusion",           "E2E (W3)",       "Adjacent-scenario hooks"]],
             col_widths=[0.7, 2.2, 1.6, 2.1])

    pb(doc)

    addh(doc, "4 · Synthetic vs. live data strategy", level=1)
    addtable(doc, ["Wave", "Data strategy", "Validation approach"],
             [["W1 Foundation",  "100% synthetic · 100K subscribers · 4 quarters · Cincinnati Kroger banner",
               "Shadow-run replay against historical 4 quarters · UAT walkthrough by VP Retention + Food Safety Lead"],
              ["W1 → W2 cutover","Synthetic + 30-day live shadow-run on Cincinnati cluster (read-only)",
               "Compare synthetic-derived disposition with real Food Safety Lead's manual decisions · target ≥ 90% agreement"],
              ["W2 Pilot",       "Live data · 250-store anchor cluster · synthetic data retained for regression tests",
               "90-day operating run · weekly KPI review · Food Safety + Privacy Office monthly review"],
              ["W3 Scale",       "All live · synthetic data only for new-banner cutover and disaster-recovery rehearsal",
               "Continuous monitoring · quarterly Steering Committee review"]],
             col_widths=[1.5, 2.5, 2.6])

    addh(doc, "5 · Performance test plan", level=1)
    addtable(doc, ["Test", "Target", "Frequency"],
             [["End-to-end decision latency", "p95 ≤ 90 sec · p99 ≤ 120 sec",     "Weekly automated"],
              ["MCP tool latency",            "Per-tool SLOs (see MCP Deep Dive)", "Daily synthetic probe"],
              ["LEDGER write throughput",     "1,000 writes/sec sustained",        "Pre-W2 cutover"],
              ["Eventstream throughput",      "Peak banner volume × 2",            "Pre-W2 + monthly"],
              ["Concurrent excursion handling","100 concurrent fires · all close within 90 sec",   "Pre-W2 + per major release"],
              ["Adaptive Card delivery",      "≤ 5 sec p95",                       "Daily"],
              ["FSMA packet generation",      "≤ 4 hr at 25K incident volume",     "Pre-W2 cutover · quarterly"]],
             col_widths=[2.4, 2.6, 1.6])

    addh(doc, "6 · Release gates", level=1)
    addtable(doc, ["Gate", "Required for",  "Approver"],
             [["All unit + integration tests pass",        "Every merge to main",   "PR reviewers"],
              ["E2E suite passes (full 30 scenarios)",     "Every release",          "Lead Engineer"],
              ["Performance regression < 10%",             "Every release",          "Engineering Manager"],
              ["UAT signoff (per UC)",                      "W1 exit",                "Stakeholder per UC"],
              ["Privacy Office signoff",                   "W1 exit + W2 entry",     "Privacy Office Lead"],
              ["Food Safety Lead signoff (FSMA UCs)",      "W1 exit + W2 entry",     "Food Safety Compliance Lead"],
              ["Internal Audit + External Audit signoff",   "W2 quarterly packet",    "VP Internal Audit + audit firm PIC"],
              ["Steering Committee signoff",                "W1 → W2 + W2 → W3",      "Steering Committee chair (COO)"],
              ["Independence-office review",               "Initial SOW + each MAJOR scope change", "Deloitte Independence office"]],
             col_widths=[3.0, 2.0, 1.6])

    addh(doc, "7 · Test environment topology", level=1)
    addtable(doc, ["Environment", "Purpose", "Data", "Refresh"],
             [["dev",   "Engineer-local + shared",         "Synthetic only",                   "Per-engineer · daily"],
              ["uat",   "Pre-production validation",        "Synthetic + sanitized historical", "Weekly · production-rehearsal cadence"],
              ["staging","Production-mirror for major releases", "Live shadow-feed (read-only)", "Real-time"],
              ["prod",  "Production",                       "Live",                              "—"],
              ["dr",    "Disaster recovery rehearsal",      "Synthetic + last-30-day prod replay", "Quarterly DR drill"]],
             col_widths=[1.4, 2.2, 2.2, 1.6])

    addbody(doc, "Independence posture preserved.", italic=True, color=MUTE, size=9)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— End of Test Strategy & Test Plan —"); setrun(r, italic=True, size=10, color=MUTE, font=SERIF)

    docsave(doc, SHARED_T2 / "APEX-RC-E2E-03-09-Test-Strategy-and-Test-Plan.docx")


# =========================================================================
# 4. SERVICE ROADMAP (HTML)
# =========================================================================

ROADMAP_HTML = r"""<!DOCTYPE html>
<html lang="en"><head>
<meta charset="UTF-8" /><meta name="viewport" content="width=device-width, initial-scale=1.0" />
<title>APEX · RC-E2E-03 + RC-E2E-09 · Service Roadmap · W1-W4+</title>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
<link href="https://fonts.googleapis.com/css2?family=Fraunces:ital,opsz,wght@0,9..144,400;0,9..144,500;0,9..144,600&family=IBM+Plex+Sans:wght@0,300;400;500;600;700&family=JetBrains+Mono:wght@400;500;600&display=swap" rel="stylesheet">
<style>
:root{--bg-0:#07090F;--bg-1:#0C1119;--bg-2:#141B27;--bg-3:#1C2433;
  --line:rgba(255,255,255,.10);--line-strong:rgba(255,255,255,.20);
  --ink-0:#F4EFE5;--ink-1:#D9D2C3;--ink-2:#9A9383;--ink-3:#6B6757;
  --gold:#E3B657;--teal:#3DD9C4;--crimson:#D94B5A;--violet:#9B7BE0;
  --sky:#6FB6E5;--amber:#E8A33D;--ember:#E87B3C;
  --w0:#9A9383;--w1:#6FB6E5;--w2:#3DD9C4;--w3:#E3B657;--w4:#9B7BE0;}
body.light{--bg-0:#F5F1E8;--bg-1:#EEE7D7;--bg-2:#E5DCC7;--bg-3:#D9CFB6;
  --line:rgba(20,14,4,.10);--line-strong:rgba(20,14,4,.22);
  --ink-0:#1A1508;--ink-1:#3D3420;--ink-2:#6B5F45;--ink-3:#8E8367;
  --gold:#8E6A1E;--teal:#147A6C;--crimson:#A22A39;--violet:#6548B0;
  --sky:#3E7FAE;--amber:#B8791F;--ember:#B0551B;
  --w0:#6B5F45;--w1:#3E7FAE;--w2:#147A6C;--w3:#8E6A1E;--w4:#6548B0;}
*{box-sizing:border-box}html,body{margin:0;padding:0}
body{background:var(--bg-0);color:var(--ink-0);font-family:"IBM Plex Sans",sans-serif;font-size:14.5px;line-height:1.55;min-height:100vh;transition:background .5s,color .5s}
body::before{content:"";position:fixed;inset:0;background:radial-gradient(1200px 600px at 10% -10%,rgba(232,163,61,.08),transparent 60%),radial-gradient(900px 500px at 95% 110%,rgba(155,123,224,.06),transparent 60%);pointer-events:none;z-index:0}
.theme-toggle{position:fixed;top:24px;right:24px;z-index:100;display:flex;align-items:center;gap:8px;background:var(--bg-1);border:1px solid var(--line-strong);border-radius:999px;padding:8px 14px;cursor:pointer;font-family:"JetBrains Mono",monospace;font-size:11px;letter-spacing:.12em;text-transform:uppercase;color:var(--ink-1)}
.hero{position:relative;z-index:1;max-width:1500px;margin:0 auto;padding:56px 48px 32px;border-bottom:1px solid var(--line)}
.eyebrow{display:inline-block;font-family:"JetBrains Mono",monospace;font-size:11px;letter-spacing:.18em;text-transform:uppercase;color:var(--gold);border:1px solid var(--line-strong);padding:6px 14px;border-radius:999px;margin-bottom:22px}
h1{font-family:"Fraunces",serif;font-weight:500;font-size:clamp(36px,5vw,60px);line-height:1.06;letter-spacing:-0.02em;margin:0 0 14px}
h1 em{font-style:italic;font-weight:400;color:var(--amber)}
.subhead{font-size:16px;color:var(--ink-1);max-width:920px;margin:0 0 28px}
.meta-strip{display:flex;flex-wrap:wrap;gap:22px;font-family:"JetBrains Mono",monospace;font-size:11.5px;color:var(--ink-2);border-top:1px solid var(--line);padding-top:18px}
.meta-strip b{color:var(--ink-0);font-weight:500}
section{position:relative;z-index:1;max-width:1500px;margin:0 auto;padding:48px;border-bottom:1px solid var(--line)}
.section-head{margin-bottom:24px}
.section-eye{font-family:"JetBrains Mono",monospace;font-size:11px;letter-spacing:.15em;text-transform:uppercase;color:var(--ink-3);margin-bottom:8px}
h2{font-family:"Fraunces",serif;font-weight:500;font-size:clamp(26px,3vw,38px);line-height:1.12;margin:0 0 6px}
.lead{color:var(--ink-1);max-width:920px}

/* TIMELINE */
.timeline-wrap{background:var(--bg-1);border:1px solid var(--line);border-radius:12px;padding:32px;margin-top:24px;overflow-x:auto}
.timeline{position:relative;min-width:1300px;height:80px;margin-bottom:24px}
.timeline::before{content:"";position:absolute;top:38px;left:0;right:0;height:4px;background:linear-gradient(to right,var(--w0),var(--w1),var(--w2),var(--w3),var(--w4));border-radius:2px;opacity:.5}
.tl-stop{position:absolute;top:0;width:200px;text-align:center}
.tl-stop .dot{width:16px;height:16px;border-radius:50%;background:var(--accent);margin:30px auto 0;position:relative;z-index:2;box-shadow:0 0 0 4px var(--bg-1)}
.tl-stop .name{font-family:"JetBrains Mono",monospace;font-size:11px;letter-spacing:.14em;color:var(--accent);margin-top:14px;font-weight:600}
.tl-stop .span{font-size:12px;color:var(--ink-2);margin-top:2px}
.tl-stop:nth-child(1){left:2%;--accent:var(--w0)}
.tl-stop:nth-child(2){left:18%;--accent:var(--w1)}
.tl-stop:nth-child(3){left:38%;--accent:var(--w2)}
.tl-stop:nth-child(4){left:60%;--accent:var(--w3)}
.tl-stop:nth-child(5){left:82%;--accent:var(--w4)}

.wave-card{background:var(--bg-1);border:1px solid var(--line);border-left:5px solid var(--accent);border-radius:12px;padding:28px;margin-top:18px}
.wave-card .wname{font-family:"JetBrains Mono",monospace;font-size:11px;letter-spacing:.14em;color:var(--accent);margin-bottom:6px}
.wave-card h3{font-family:"Fraunces",serif;font-size:24px;font-weight:500;margin:0 0 8px}
.wave-card .span-meta{font-family:"JetBrains Mono",monospace;font-size:11.5px;color:var(--ink-2);margin-bottom:18px}
.wave-card .span-meta b{color:var(--accent)}
.wave-grid{display:grid;grid-template-columns:repeat(3,1fr);gap:16px;margin-top:14px}
@media (max-width:900px){.wave-grid{grid-template-columns:1fr}}
.feat-block h4{font-family:"Fraunces",serif;font-size:15px;font-weight:500;margin:0 0 8px;color:var(--ink-0)}
.feat-block .feat-eye{font-family:"JetBrains Mono",monospace;font-size:10.5px;letter-spacing:.14em;color:var(--accent);margin-bottom:6px}
.feat-block ul{margin:0;padding-left:18px;font-size:13px;color:var(--ink-1);line-height:1.55}
.feat-block ul li{margin-bottom:4px}

.gates{background:var(--bg-2);border-radius:8px;padding:14px 18px;margin-top:18px;border-left:3px solid var(--gold)}
.gates .glabel{font-family:"JetBrains Mono",monospace;font-size:10.5px;letter-spacing:.14em;color:var(--gold);margin-bottom:6px}
.gates ul{margin:0;padding-left:18px;font-size:13px;color:var(--ink-1)}

footer{position:relative;z-index:1;max-width:1500px;margin:0 auto;padding:32px 48px 64px;font-family:"JetBrains Mono",monospace;font-size:11px;color:var(--ink-3);display:flex;justify-content:space-between;flex-wrap:wrap;gap:12px}
</style></head>
<body>
<div class="theme-toggle" onclick="toggleTheme()"><span id="themeIcon">☀</span><span id="themeLabel">Light</span></div>

<header class="hero">
  <div class="eyebrow">APEX · RC-E2E-03 + RC-E2E-09 · SERVICE ROADMAP · W1–W4+</div>
  <h1>From narrow specialists to <em>context-aware superagents</em>.</h1>
  <p class="subhead">Forward-looking product roadmap across the four delivery waves plus the Approach stage. What lands when, what gates the next wave, and what fusion candidates emerge in W3 and W4+. Pairs with the Six-Agents Maturation document for the agent-side narrative.</p>
  <div class="meta-strip">
    <div><b>Waves</b> · Approach · W1 · W2 · W3 · W4+</div>
    <div><b>Span</b> · ~3 years to W3 enterprise · open-ended W4+</div>
    <div><b>Cumulative envelope</b> · $24M – $61M to W3</div>
    <div><b>W3 run-rate value</b> · ~$200M annualized</div>
    <div><b>Reference</b> · Kroger</div>
  </div>
</header>

<section>
  <div class="section-head">
    <div class="section-eye">§ 1 · Timeline at a glance</div>
    <h2>Five stops · ~3 years to enterprise run-rate.</h2>
    <p class="lead">Approach 4 weeks · W1 5–8 months · W2 8–12 months · W3 12–18 months · W4+ open-ended.</p>
  </div>
  <div class="timeline-wrap">
    <div class="timeline">
      <div class="tl-stop"><div class="dot"></div><div class="name">APPROACH</div><div class="span">3–5 wks</div></div>
      <div class="tl-stop"><div class="dot"></div><div class="name">WAVE 1</div><div class="span">5–8 mo</div></div>
      <div class="tl-stop"><div class="dot"></div><div class="name">WAVE 2</div><div class="span">8–12 mo</div></div>
      <div class="tl-stop"><div class="dot"></div><div class="name">WAVE 3</div><div class="span">12–18 mo</div></div>
      <div class="tl-stop"><div class="dot"></div><div class="name">WAVE 4+</div><div class="span">Ongoing</div></div>
    </div>
  </div>
</section>

<section>
  <div class="section-head"><div class="section-eye">§ 2 · Approach</div><h2>Discovery, scoping, SOW.</h2></div>
  <div class="wave-card" style="--accent:var(--w0)">
    <div class="wname">Approach · 3–5 weeks · $200K–$500K</div>
    <h3>Stakeholder discovery + commercial sizing</h3>
    <div class="span-meta">Outcome: <b>Signed SOW + Independence-office approval</b> · Steering Committee charter · banner-cluster scope locked</div>
    <div class="wave-grid">
      <div class="feat-block">
        <div class="feat-eye">Discovery</div>
        <h4>Stakeholder + system map</h4>
        <ul><li>10+ persona interviews</li><li>Source-system inventory</li><li>FSMA 204 readiness baseline</li></ul>
      </div>
      <div class="feat-block">
        <div class="feat-eye">Sizing</div>
        <h4>Commercial envelope</h4>
        <ul><li>Risk register (20 risks)</li><li>Sensitivity bands</li><li>Banner-cluster recommendation</li></ul>
      </div>
      <div class="feat-block">
        <div class="feat-eye">Contracting</div>
        <h4>SOW + Independence</h4>
        <ul><li>Independence-office routing</li><li>SOW execution</li><li>Audit Committee notification</li></ul>
      </div>
    </div>
    <div class="gates"><div class="glabel">Exit gate</div><ul><li>VP Retention + Streaming COO sign-off · Independence office sign-off · SOW executed</li></ul></div>
  </div>
</section>

<section>
  <div class="section-head"><div class="section-eye">§ 3 · Wave 1 Foundation</div><h2>Single-banner anchor · 3 of 5 agents · shadow-run.</h2></div>
  <div class="wave-card" style="--accent:var(--w1)">
    <div class="wname">W1 · 5–8 months · $1.2M–$3.5M · 50 SP</div>
    <h3>Cincinnati Kroger · ~50 stores · 3 agents · synthetic-data lab</h3>
    <div class="span-meta">Outcome: <b>Runnable W1 prototype</b> · audit-row coverage E2E · 3 demo paths · VP Retention sign-off</div>
    <div class="wave-grid">
      <div class="feat-block">
        <div class="feat-eye">Data</div>
        <h4>Bronze landing</h4>
        <ul><li>5 streaming feeds (Eventstream)</li><li>7 batch feeds (Data Factory)</li><li>SCML.Lot · MERML.Price · DIM.Item Silver</li></ul>
      </div>
      <div class="feat-block">
        <div class="feat-eye">Agents</div>
        <h4>Three live</h4>
        <ul><li>Assess (live)</li><li>Classify (live, escalate &lt; 0.85)</li><li>Evidence-Write (live)</li></ul>
      </div>
      <div class="feat-block">
        <div class="feat-eye">Substrate</div>
        <h4>Foundation rails</h4>
        <ul><li>LEDGER scaffolding</li><li>Adaptive Card on VP Retention's device</li><li>14-field audit row E2E</li></ul>
      </div>
    </div>
    <div class="gates"><div class="glabel">Exit gate</div><ul><li>VP Retention + Food Safety Lead sign-off · shadow-run accuracy ≥ 95% on historical 4 quarters · 3 demo paths landed</li></ul></div>
  </div>
</section>

<section>
  <div class="section-head"><div class="section-eye">§ 4 · Wave 2 Pilot</div><h2>250-store anchor cluster · all 5 agents live · FSMA packet.</h2></div>
  <div class="wave-card" style="--accent:var(--w2)">
    <div class="wname">W2 · 8–12 months · $5M–$12M · 45 SP</div>
    <h3>Multi-banner anchor · 250 stores · all 6 agents · 90-day operating run</h3>
    <div class="span-meta">Outcome: <b>Live operating service</b> · first quarterly FSMA evidence packet · Steering Committee approval to scale</div>
    <div class="wave-grid">
      <div class="feat-block">
        <div class="feat-eye">Agents</div>
        <h4>All six live</h4>
        <ul><li>Quantify · Approve · Act live</li><li>Materiality policy tier-tuned</li><li>Designate-approver workflow</li></ul>
      </div>
      <div class="feat-block">
        <div class="feat-eye">Mutation</div>
        <h4>Production write surface</h4>
        <ul><li>POS markdown upload (banner-aware)</li><li>Vendor-portal ticket creation</li><li>FSMA regulatory log</li></ul>
      </div>
      <div class="feat-block">
        <div class="feat-eye">Evidence</div>
        <h4>FSMA + Privacy</h4>
        <ul><li>Quarterly FSMA evidence packet</li><li>Privacy/Audit walkthrough</li><li>Concession Policy Studio live</li></ul>
      </div>
    </div>
    <div class="gates"><div class="glabel">Exit gate</div><ul><li>Steering Committee approval · Privacy/Audit signoff on Q1 packet · 90-day KPI targets met · external-audit walkthrough complete</li></ul></div>
  </div>
</section>

<section>
  <div class="section-head"><div class="section-eye">§ 5 · Wave 3 Scale & Fuse</div><h2>Enterprise 2,800+ stores · cross-banner · adjacent-scenario fusion.</h2></div>
  <div class="wave-card" style="--accent:var(--w3)">
    <div class="wname">W3 · 12–18 months · $18M–$45M · 70+ SP</div>
    <h3>All Kroger banners · multi-banner orchestration · scenario fusion</h3>
    <div class="span-meta">Outcome: <b>Enterprise run-rate</b> · ~$200M annualized shrink avoided · adjacent-scenario fusion candidates active</div>
    <div class="wave-grid">
      <div class="feat-block">
        <div class="feat-eye">Scale</div>
        <h4>2,800+ stores</h4>
        <ul><li>All 10+ banners live</li><li>Cross-banner aggregation views</li><li>F-SKU scaled to F256</li></ul>
      </div>
      <div class="feat-block">
        <div class="feat-eye">Fusion</div>
        <h4>Adjacent scenarios</h4>
        <ul><li>RC-E2E-04 Loyalty attach</li><li>RC-E2E-11 Perishables Integrity</li><li>RC-E2E-12 Regulator Evidence</li></ul>
      </div>
      <div class="feat-block">
        <div class="feat-eye">Maturation</div>
        <h4>Agents at W3 band</h4>
        <ul><li>Cross-banner pattern aware</li><li>Cohort drift remediation</li><li>Auto-execute rate ≥ 75%</li></ul>
      </div>
    </div>
    <div class="gates"><div class="glabel">Exit gate</div><ul><li>Steady-state operating review · 90-day enterprise run · all-banner Steering Committee signoff · readiness for W4+ fusion</li></ul></div>
  </div>
</section>

<section>
  <div class="section-head"><div class="section-eye">§ 6 · Wave 4+ Superagent Era</div><h2>Six agents → three superagents · cross-Practice generalization.</h2></div>
  <div class="wave-card" style="--accent:var(--w4)">
    <div class="wname">W4+ · ongoing · TBD envelope</div>
    <h3>Cross-Practice intelligence · principle-based HITL · proactive interventions</h3>
    <div class="span-meta">Outcome: <b>Three superagents replace six specialists</b> · cross-Practice deployment (HLS, ER, AXLE, TMT) · principle-based escalation</div>
    <div class="wave-grid">
      <div class="feat-block">
        <div class="feat-eye">Fusion A</div>
        <h4>Risk-Classify</h4>
        <ul><li>Absorbs Assess + Classify</li><li>Multi-scenario excursion + disposition reasoning</li></ul>
      </div>
      <div class="feat-block">
        <div class="feat-eye">Fusion B</div>
        <h4>Materiality-Decision</h4>
        <ul><li>Absorbs Quantify + Approve</li><li>Multi-scenario quantification + HITL routing</li></ul>
      </div>
      <div class="feat-block">
        <div class="feat-eye">Fusion C</div>
        <h4>Closeout</h4>
        <ul><li>Absorbs Act + Evidence-Write</li><li>Multi-scenario downstream commit + evidence emission</li></ul>
      </div>
    </div>
    <div class="gates"><div class="glabel">Maturation gate</div><ul><li>180 days at W3 maturity · cross-Practice test bench · principle-based HITL pilot · Practice leadership + Independence office gate</li></ul></div>
  </div>
</section>

<section>
  <div class="section-head"><div class="section-eye">§ 7 · Cross-wave dependencies</div><h2>What gates each wave's promotion.</h2></div>
  <div class="wave-card" style="--accent:var(--gold)">
    <div class="wname">Cross-wave structural connections</div>
    <h3>The architectural commitments that bind the waves together</h3>
    <div class="wave-grid">
      <div class="feat-block">
        <div class="feat-eye">W1 → W2</div>
        <h4>HITL substrate</h4>
        <ul><li>Adaptive Card pattern</li><li>Entra identity binding</li><li>LEDGER plane established</li></ul>
      </div>
      <div class="feat-block">
        <div class="feat-eye">W1 → W2</div>
        <h4>FSMA packet</h4>
        <ul><li>Packet format designed in W1</li><li>Exercised with shadow audit rows</li><li>Scales to thousands in W2</li></ul>
      </div>
      <div class="feat-block">
        <div class="feat-eye">W2 → W3</div>
        <h4>Pilot KPIs unlock scale</h4>
        <ul><li>Decision-loop time target met</li><li>FSMA coverage ≥ 99%</li><li>Steering Committee gate</li></ul>
      </div>
    </div>
  </div>
</section>

<footer>
  <div>APEX · RC-E2E-03 + RC-E2E-09 · Service Roadmap · v1.0 · 27 Apr 2026</div>
  <div>Companion: Six-Agents Maturation · Walkthrough · SAD · Operations Playbook · Test Plan</div>
  <div>Deloitte · Microsoft Technology Solutions &amp; Practices</div>
</footer>

<script>
function toggleTheme(){const b=document.body,i=document.getElementById('themeIcon'),l=document.getElementById('themeLabel');b.classList.toggle('light');if(b.classList.contains('light')){i.textContent='☾';l.textContent='Dark';localStorage.setItem('apex-theme','light')}else{i.textContent='☀';l.textContent='Light';localStorage.setItem('apex-theme','dark')}}
(function(){const s=localStorage.getItem('apex-theme');if(s==='light'){document.body.classList.add('light');document.getElementById('themeIcon').textContent='☾';document.getElementById('themeLabel').textContent='Dark'}})();
</script></body></html>
"""

def build_roadmap():
    print("\nBuilding Service Roadmap html...")
    target = SHARED_T4 / "APEX-RC-E2E-03-09-Service-Roadmap.html"
    target.write_text(ROADMAP_HTML, encoding="utf-8")
    print(f"  wrote {target.relative_to(ROOT)}")


# =========================================================================
# 5. DEMO SCRIPT & WALKTHROUGH GUIDE
# =========================================================================

def build_demo_script():
    print("\nBuilding Demo Script & Walkthrough Guide...")
    doc = docinit("APEX · RC-E2E-03 · Demo Script & Walkthrough Guide",
                  "Three narrated demo paths · stakeholder Q&A prep · environment setup · 25-minute total runtime",
                  None)
    addbody(doc, "Step-by-step demo script for the W1 prototype. Three canonical paths (~5 min each), stakeholder Q&A "
                 "prep, demo environment setup, and the runbook for delivering the demo cleanly to a Kroger "
                 "stakeholder audience. Use this as the source-of-truth for live walkthroughs.",
            italic=True, color=INK2)
    doc.add_paragraph()
    addkv(doc, [
        ("Document type",      "Demo Script & Walkthrough Guide"),
        ("Service",            "RC-E2E-03 (with RC-E2E-09 evidence path)"),
        ("Reference client",   "The Kroger Co."),
        ("Audience",           "Account Team · Pursuit Lead · Engagement Architect · presenters"),
        ("Total runtime",      "~25 minutes (3 paths × 5 min + Q&A prep)"),
        ("Companion",          "Walkthrough · Sequence Diagram · Pitch Deck · ROI Case"),
        ("Version / date",     "v1.0 · 27 Apr 2026"),
    ])
    pb(doc)

    addh(doc, "1 · Demo prep checklist (T-24 hrs)", level=1)
    addbody(doc, "Run through this checklist 24 hours before any stakeholder demo. Most issues caught here are "
                 "much harder to recover from once the meeting starts.")
    addbullets(doc, [
        "Demo laptop charged + spare power adapter packed",
        "Docker daemon running · 'docker compose up' verified · all 6 services green",
        "Local Ollama model (offline path) verified · cap:disposition-classification responds",
        "Or: Azure OpenAI online path · backup plan if Wi-Fi flakes",
        "Synthetic data lab loaded · 100K subscribers · 4-quarter event stream",
        "Adaptive Card on demo phone (Marisol Reyes mock identity) · notification verified",
        "Power BI dashboard tiles loaded · drill-through to LEDGER trace_id verified",
        "Three demo paths walked end-to-end one time · timing within target",
        "Backup screenshots of each path saved (in case live demo fails)",
        "Stakeholder roster confirmed · personalize talk-track for who's in the room",
    ])

    addh(doc, "2 · Demo environment setup (T-2 hrs)", level=1)
    addbody(doc, "Two-hour pre-flight. Enough time to recover if something is broken.")
    addtable(doc, ["Step", "Action", "Verify"],
             [["1",  "Boot demo laptop · close all unrelated apps", "Activity Monitor: only demo processes"],
              ["2",  "Pull latest demo branch · git pull origin demo", "Last commit hash matches expected"],
              ["3",  "docker compose up -d · wait 60 sec",            "All services HEALTHY in compose ps"],
              ["4",  "Open http://localhost:3000 in primary browser", "Demo console loads · all tabs functional"],
              ["5",  "Open Power BI Desktop · load dashboard pbix",   "Dashboard tiles load · drill-through works"],
              ["6",  "Open Microsoft Teams on demo phone",            "Test card delivery from /api/test-card endpoint"],
              ["7",  "Run demo path A end-to-end as rehearsal",        "Closeout audit row appears in LEDGER · timing < 90 sec"],
              ["8",  "Take screenshots of each major step",            "Saved to demo-backup/ folder"],
              ["9",  "Test failure recovery: kill one service · restart", "System recovers gracefully"],
              ["10", "Final 'docker compose down && up -d' · clean slate", "Fresh state for live demo"]],
             col_widths=[0.5, 3.0, 3.1])

    pb(doc)

    addh(doc, "3 · Demo Path A · The intervention loop (5 min)", level=1)
    addbody(doc, "The marquee path. Shows the full agent chain end-to-end on a synthetic excursion. "
                 "Anchors the 90-second decision narrative.")
    addh(doc, "3.1 · Setup", level=2)
    addbullets(doc, [
        "Open the demo console · navigate to 'Excursion Simulator'",
        "Select: Banner = Kroger · Store = #428 · Fixture = DAIRY-A3 · Severity = CRITICAL",
        "Talk-track: 'This is the kind of event Kroger handles ~85 times per banner per week. Today it's a 17-minute decision loop. Watch what happens here in about 90 seconds.'",
    ])
    addh(doc, "3.2 · Step-by-step (with talk-track)", level=2)
    addtable(doc, ["Time", "What's on screen", "What you say"],
             [["0:00", "Click 'Fire Excursion' button · trace_id appears",
               "'Eventstream watermark fires. Trace_id is allocated. Watch the agent chain.'"],
              ["0:05", "Assess agent activates · ExcursionProfile renders",
               "'Assess just characterized the event. Peak 48.2°F, 47-min duration, 4 door-open events. Severity CRITICAL. This took 3 seconds and 4 MCP calls.'"],
              ["0:10", "Classify agent activates · DispositionTable populates",
               "'Classify is reasoning over the FSMA 204 policy corpus and the 6 affected lots. Per-SKU disposition with confidence and the policy clause that fired. All 6 lots above 0.85 confidence — no escalation needed.'"],
              ["0:20", "Quantify agent activates · 3-scenario $ comparison",
               "'Quantify converted dispositions to dollars. $7,920 if we destroy everything. $4,210 if we mark down. $0 risk if we sell-through. The system is recommending a 22% markdown.'"],
              ["0:30", "Adaptive Card lands on demo phone · show phone to camera",
               "'Card just lit up Marisol's phone. She's the Store Operations Lead at Store 428. She has 90 seconds of context she would have spent 12 minutes assembling.'"],
              ["0:50", "Tap APPROVE on phone · audit row writes",
               "'Marisol approves. The decision writes to LEDGER with her named identity. Now watch what fires in parallel.'"],
              ["1:00", "Three parallel actions visualized · all green",
               "'POS markdown went live. Refrigeration vendor ticket auto-opened with Emerson. FSMA regulatory log entry created. Three systems updated in parallel.'"],
              ["1:20", "Evidence-Write composes closeout · LEDGER tile shows new row",
               "'Evidence-Write is composing the 14-field closeout audit row right now. Manifest version, policy version, named approver, all KPIs measured. This is the FSMA evidence chain — generated, not collated.'"],
              ["1:30", "Closeout complete · scenario sealed",
               "'90 seconds. Today this is 17 minutes. Tomorrow it's 1.5. Compounded across 12.4 million events per year, this is roughly $200M in shrink avoided.'"]],
             col_widths=[0.6, 3.0, 3.0])

    pb(doc)

    addh(doc, "4 · Demo Path B · The retention-NPV trending (5 min)", level=1)
    addbody(doc, "Shows the executive-side view. Power BI dashboard with drill-through to LEDGER. Anchors the "
                 "'every KPI attributes to a specific decision' narrative.")
    addh(doc, "4.1 · Step-by-step", level=2)
    addtable(doc, ["Time", "What's on screen", "What you say"],
             [["0:00", "Open Power BI shrink dashboard · 12 weeks of trending",
               "'This is what the Chief Merchant sees every Monday morning. Per-banner shrink-avoided trending across 12 weeks of shadow-run data.'"],
              ["0:30", "Hover over a high-NPV cell in the heatmap",
               "'Cincinnati Kroger banner. Week of April 12. $48,200 in shrink avoided. Click and we drill through.'"],
              ["1:00", "Drill-through table loads · 37 contributing audit rows",
               "'37 specific decisions composed this number. Every single one is an audit row. We're not aggregating — we're retrieving.'"],
              ["1:30", "Click into a single audit row",
               "'This is one decision. Trace_id. Manifest version. Policy version. The specific FSMA 204 clause that fired. Marisol Reyes approved. Confidence 0.91.'"],
              ["2:30", "Show the reasoning trace link",
               "'Click here and we get the full LangFuse reasoning trace. Every MCP tool call. Every model inference. Every prompt. Reproducible to the second.'"],
              ["3:00", "Back to the dashboard · show category-quality cell",
               "'Same drill works on category quality, vendor scorecard, FSMA evidence completeness — anything sourced from LEDGER.'"],
              ["4:00", "Show the 'Export evidence packet' button",
               "'And this button right here generates the quarterly FSMA evidence packet. Auto-composed from the same audit rows. The artifact the FDA inspector consumes.'"],
              ["4:45", "Wrap",
               "'KPIs are durable because they reduce to specific decisions. Auditor can sample any cell. Every aggregate has a forensic path back to the original decision.'"]],
             col_widths=[0.6, 3.0, 3.0])

    pb(doc)

    addh(doc, "5 · Demo Path C · The audit trail (5 min)", level=1)
    addbody(doc, "Shows the regulatory/audit perspective. Pick a synthetic excursion and walk the chain of custody "
                 "end-to-end. Anchors the FSMA 204 'records are generated, not collated' commitment.")
    addh(doc, "5.1 · Step-by-step", level=2)
    addtable(doc, ["Time", "What's on screen", "What you say"],
             [["0:00", "Open LEDGER explorer · search by store + date",
               "'Imagine FDA inspector calls. They want all cold-chain dispositions for Store 428 in March. We type two filters.'"],
              ["0:30", "Results: 14 audit rows · sortable spreadsheet",
               "'14 audit rows. Sortable spreadsheet, exactly what the FSMA 204 final rule requires. Notice every row has a trace_id, every field is non-null on FSMA-relevant decisions.'"],
              ["1:00", "Click into one row · show the 14 fields",
               "'Trace_id. Scenario_id. Manifest version 1.0.0. Policy version FSMA-204-cold-chain-corpus-v1.0.0. Approver — named Entra identity. Approval timestamp. KPIs.'"],
              ["1:45", "Show the reasoning trace · LangFuse spans",
               "'Reasoning trace: every agent call, every MCP tool call, every prompt. The auditor sees not just what happened, but why.'"],
              ["2:30", "Show the lot-trace lineage",
               "'Click here and we follow the lot-trace lineage back through SCML.Lot. Receipt event, supplier, FSMA 204 KDEs. The full chain of custody.'"],
              ["3:15", "Show consent state for personalization decisions",
               "'And for any decision that touched personalization, here's the consent state attestation captured at decision time. Privacy Office sees what the system saw at the moment.'"],
              ["4:00", "Click 'Generate evidence packet for this period'",
               "'Auto-generated evidence packet. Manifests, policies, audit rows, tie-out reconciliation. Within 4 hours at W2; under 60 minutes at W3.'"],
              ["4:45", "Wrap",
               "'This is what FSMA 204 readiness looks like by construction. Not a binder of paper. Not a quarterly assembly project. A queryable audit substrate that's always current.'"]],
             col_widths=[0.6, 3.0, 3.0])

    pb(doc)

    addh(doc, "6 · Anticipated questions and prepared answers", level=1)
    addbody(doc, "The most-likely Kroger stakeholder questions. Pre-rehearsed answers below.")

    qa_pairs = [
        ("Q · How does this not become another shelf-ware AI initiative?",
         "Three architectural commitments prevent shelf-ware. (1) Manifest-driven — every decision is reproducible from a Git history. "
         "(2) Audit-row-by-construction — evidence is generated, not collated, so it's auditor-grade from day one. (3) HITL is "
         "first-class — humans stay on every material decision, so adoption isn't predicated on automation acceptance. The "
         "5-agent pattern lands in 5–8 months at single-banner scope; if it doesn't work in shadow-run, you don't go to W2."),
        ("Q · What about 84.51° — does this displace it?",
         "No. RC-E2E-03 reads from 84.51° loyalty data via cap:loyalty-hh consent-aware feed. The pricing engine remains 84.51° + "
         "Revionics. RC-E2E-03 adds the cold-chain disposition layer that 84.51° doesn't currently address. Complementary, not "
         "competitive."),
        ("Q · How do you handle the Independence question if Deloitte is our auditor?",
         "Two architectural commitments make this independence-safe: (1) manifests, policy, and LEDGER are Kroger artifacts, "
         "with Deloitte's role limited to advisory and implementation; (2) before any contract, the SOW is routed through "
         "Deloitte's Independence office. We've confirmed Kroger's external auditor is PwC, which simplifies but doesn't "
         "obviate the routing."),
        ("Q · What's the disaster-recovery posture for the LEDGER?",
         "Active geo-replication on the Postgres index between East US 2 and West US 2 with a failover group. Delta Lake "
         "in OneLake gets geo-redundant replication. RTO 30 min for LEDGER · RPO 5 min. Quarterly DR drills validate the path. "
         "Importantly: agents continue to write to local Bronze even if LEDGER is briefly unavailable, and reconcile on recovery."),
        ("Q · How do we trust the ML models?",
         "Three layers. (1) Per-decision confidence calibration — Classify confidence below 0.85 escalates to Food Safety Lead, "
         "so the human stays on uncertain cases. (2) Cohort drift watchtower — UC-22 monitors precision/recall weekly per cohort. "
         "(3) Champion-challenger A/B for promotions — new models shadow live for 14 days before deployment. Plus the Audit-Row "
         "trail means every model decision is reproducible and auditable."),
        ("Q · What's the path to extend beyond cold-chain?",
         "Two paths. (1) Wave-3 fusion — RC-E2E-11 Perishables Integrity Response and RC-E2E-12 Regulator Evidence Service "
         "share the architectural substrate; they activate post-Wave-2. (2) Wave-4+ superagent absorption — the same six "
         "agents fuse into three superagents that handle multi-scenario decisioning. Same audit-row contract preserved across both."),
        ("Q · Who owns the model when it goes wrong?",
         "Kroger ML Engineering owns the model. Audience Lead owns cohort drift triage. Food Safety Lead owns FSMA-related "
         "exceptions. Privacy Office owns consent-related exceptions. Deloitte advises and supports — but the model artifact "
         "and lifecycle decisions are Kroger's. Independence-safe by construction."),
        ("Q · What's the Wave-1 commitment if we say go today?",
         "5–8 months · $1.2M–$3.5M · 50 story points · Cincinnati Kroger banner cluster · 3 of 5 agents live in shadow-run. "
         "Exit gate is VP Retention + Food Safety Lead sign-off on shadow-run accuracy. If we can't pass the gate, we don't "
         "scale to W2. Risk-managed entry."),
        ("Q · Can we see this running on real Kroger data, not synthetic?",
         "Synthetic in W1 by design — it lets us de-risk the architecture without negotiating data-access scope upfront. "
         "In Wave 1, weeks 5–8, we begin a 30-day live shadow-run on Cincinnati cluster (read-only, no mutation). That's "
         "where we validate against your real data before any production action in W2."),
        ("Q · What happens if the FDA changes FSMA 204 requirements mid-engagement?",
         "Three things. (1) The FSMA 204 policy corpus is a versioned manifest — we update it in a PR, not a release. (2) "
         "BRZ.RC.LOT_TRACE.kde_schema_version field accommodates KDE format evolution. (3) The audit row preserves the "
         "policy version that fired at decision time, so historical decisions remain defensible against the regulation that "
         "applied to them. We're built for regulatory evolution, not against it."),
    ]
    for q, a in qa_pairs:
        addh(doc, q, level=2)
        addbody(doc, a, color=INK2, size=10.5)

    pb(doc)

    addh(doc, "7 · Demo failure-recovery playbook", level=1)
    addbody(doc, "What to do when something breaks in the live demo. The shorter the recovery, the more confident the audience.")
    addtable(doc, ["Failure", "Recovery"],
             [["Docker service crashes mid-demo",
               "Smile, say 'this is a great demo of our incident playbook' · 'docker compose restart [service]' · resume in 30 sec"],
              ["Adaptive Card doesn't deliver",
               "Pivot to backup screenshots · 'normally this is a Teams card · here's what Marisol sees on her device'"],
              ["Power BI dashboard slow to load",
               "Pivot to LEDGER explorer (faster surface) · explain that Power BI is one of multiple consumers"],
              ["Wi-Fi flakes / Azure OpenAI unreachable",
               "Restart with --offline flag · uses local Ollama · slightly slower but works"],
              ["Audience question outside scope",
               "Park it: 'great question · let me get you a precise answer in our follow-up tomorrow' · capture for action items"],
              ["Audience pushback on a specific number",
               "'Show me your assumption' · open the ROI Case sensitivity sliders · adjust live · re-emerge with their number"],
              ["Time pressure (running long)",
               "Skip Path B (NPV trending) — keep Path A (intervention loop) and Path C (audit trail) · those are the two "
               "most-asked stakeholder views"]],
             col_widths=[2.4, 4.0])

    addh(doc, "8 · Post-demo follow-up checklist (T+24 hrs)", level=1)
    addbullets(doc, [
        "Send personalized recap email · attach One-pager · ROI Case · relevant Personas",
        "Capture all parked questions · respond within 24 business hours",
        "Identify champions in the room · note for Stakeholder Map updates",
        "Identify skeptics · their objections inform Risk Register adds",
        "Schedule next-step working session within 5 business days",
        "Update engagement log · who attended · key reactions · decisions",
    ])

    addbody(doc, "Independence posture preserved.", italic=True, color=MUTE, size=9)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— End of Demo Script & Walkthrough Guide —"); setrun(r, italic=True, size=10, color=MUTE, font=SERIF)

    docsave(doc, RC03_T3 / "APEX-RC-E2E-03-Demo-Script-and-Walkthrough-Guide.docx")


# ============== EXECUTE ==============

print("Building Tier A artifacts...\n")
build_sad()
build_ops_playbook()
build_test_plan()
build_roadmap()
build_demo_script()
print("\nTier A complete.")
