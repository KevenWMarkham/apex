"""Append CFMP v0.2 section (incl. new Virtual Views + Pack catalog update)
into every relevant APEX book volume.

Strategy:
- Each HTML file gets backed up to `*-pre-cfmp-v0.2.html.bak` the first
  time we touch it; subsequent runs are idempotent (we strip our previous
  insertion before re-appending if you re-run).
- A common CFMP_HTML_SECTION block is composed per-volume (full version
  for the big books, short blurb for the Executive Summary, Pack-catalog
  entry for the Library).
- TRILOGY-MS-PLATFORM-VALIDATION.md gets a new CFMP row appended.
- APEX-Sellers-Guide-Runtime-Addendum.docx gets a new section added.

Idempotent: a sentinel comment marks our injected block so we can find +
remove it on re-run.

Run once after changes:
    python _update_book_volumes.py
"""

from __future__ import annotations

import re
from pathlib import Path
from datetime import date

BOOK_DIR = Path(r"C:\Stage\Clients\Industries\APEX\docs\book")

SENTINEL_START = "<!-- BEGIN CFMP-V0.2-ADDENDUM (auto-managed; do not edit by hand) -->"
SENTINEL_END   = "<!-- END CFMP-V0.2-ADDENDUM -->"

TODAY = date.today().isoformat()

# ─── New Virtual Views catalog the CFMP Pack contributes ──────────────────
# These are the VVs that ride on the new Interface #15 + customer persona.
NEW_VVS = [
    ("merml.product_with_planogram_location",
     "Per-SKU shelf location joined from MERML.Product × CFMP.StoreMap. "
     "Refreshed nightly from the planogram feed; drives the agent's "
     "route_to_product tool."),
    ("merml.osa_by_aisle",
     "On-shelf availability rolled up by aisle. Combines POS velocity, "
     "inventory position, and Vision AI Dev Kit shelf-void detections. "
     "Feeds OSA / shelf-gap-restock-dispatch scenarios."),
    ("cxml.customer_session_with_route",
     "Active customer phone session enriched with the most recent "
     "wayfinding route + current graph node. Backs the customer_phone "
     "Adaptive Card transport."),
    ("cxml.cart_dwell_event",
     "Cart-stationary >5min outside checkout zone event stream. "
     "Triggers cart-dwell-abandonment-rescue dispatches via Teams cards "
     "to the nearest associate."),
    ("cxml.dietary_profile",
     "Customer's active dietary preferences (loyalty-bound; consent-gated). "
     "Read by the proactive associate to flag conflicts and propose "
     "compliant alternatives."),
    ("cfmp.storemap_view",
     "Thin overlay binding the retailer store_id to the Azure Maps Creator "
     "Dataset ID + Stateset ID + nightly SKU↔unit_id table. Pack-level "
     "manifest captured once per store, refreshed on remodel."),
]


# Updated Pack catalog with CFMP slotted as the 7th pack.
PACK_CATALOG_NOTE = """\
The APEX Pack catalog now includes seven Industry Solution Packs.
CFMP is the first organized around a customer-moment spine rather than
an operator function — sibling to the operator-side Retail Merchandising
Pack v1 (§9.6 of APEX-Architecture-v5).
"""


def render_section(title: str, body_html: str) -> str:
    """Wrap a CFMP section with a distinct styled card so it stands out
    inside the Word-published HTML soup."""
    return f"""{SENTINEL_START}
<style>
  .cfmp-addendum {{
    margin: 2em auto; max-width: 980px;
    font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
    background: #f4f8ff; border: 2px solid #1c4e80; border-radius: 8px;
    padding: 1.5em 1.75em; color: #0d1117;
  }}
  .cfmp-addendum h2.cfmp {{
    color: #1c4e80; border-bottom: 2px solid #1c4e80; padding-bottom: 0.3em;
    margin-top: 0;
  }}
  .cfmp-addendum h3.cfmp {{ color: #1c4e80; margin-top: 1.2em; }}
  .cfmp-addendum .cfmp-tag {{
    display: inline-block; background: #1c4e80; color: white;
    padding: 0.15em 0.55em; border-radius: 4px; font-size: 0.75em;
    font-weight: 700; letter-spacing: 0.04em; margin-right: 0.4em;
  }}
  .cfmp-addendum .cfmp-live {{ background: #0a6e3a; }}
  .cfmp-addendum .cfmp-planned {{ background: #6e560a; }}
  .cfmp-addendum table.cfmp {{
    border-collapse: collapse; margin: 0.8em 0; width: 100%;
    background: white; border: 1px solid #c0d0e0;
  }}
  .cfmp-addendum table.cfmp th, .cfmp-addendum table.cfmp td {{
    border: 1px solid #c0d0e0; padding: 0.4em 0.7em; text-align: left;
    vertical-align: top; font-size: 0.92em;
  }}
  .cfmp-addendum table.cfmp th {{ background: #e2ecf7; }}
  .cfmp-addendum code {{
    background: #e8eff7; padding: 0.05em 0.3em; border-radius: 3px;
    font-size: 0.92em;
  }}
  .cfmp-addendum pre {{
    background: #0d1117; color: #cdd9e5; padding: 0.7em 0.9em;
    border-radius: 6px; overflow-x: auto; font-size: 0.86em;
  }}
  .cfmp-addendum ul.cfmp {{ padding-left: 1.4em; }}
  .cfmp-addendum ul.cfmp li {{ margin-bottom: 0.25em; }}
</style>
<section class="cfmp-addendum">
  <div>
    <span class="cfmp-tag">APEX-M · CFMP v0.2</span>
    <span class="cfmp-tag cfmp-live">PHASE 1 LIVE</span>
    <span style="color:#7d8590;font-size:0.85em">Addendum inserted {TODAY}</span>
  </div>
  <h2 class="cfmp">{title}</h2>
  {body_html}
</section>
{SENTINEL_END}
"""


def new_vv_table() -> str:
    rows = "\n".join(
        f"      <tr><td><code>{vv}</code></td><td>{desc}</td></tr>"
        for vv, desc in NEW_VVS
    )
    return f"""\
    <h3 class="cfmp">New Virtual Views shipped with CFMP v0.2</h3>
    <p>CFMP contributes <strong>6 new Virtual Views</strong> spanning MERML
    (merchandise), CXML (customer experience), and the new
    <code>cfmp</code> namespace. All are declarative manifests; the runtime
    federates them through the standard apex-orchestrator data path.</p>
    <table class="cfmp">
      <thead><tr><th>Virtual View ID</th><th>Purpose</th></tr></thead>
      <tbody>
{rows}
      </tbody>
    </table>
"""


def pack_catalog_table() -> str:
    return f"""\
    <h3 class="cfmp">Updated Pack catalog (7 packs)</h3>
    <p>{PACK_CATALOG_NOTE}</p>
    <table class="cfmp">
      <thead>
        <tr><th>#</th><th>Pack</th><th>Buyer</th><th>Schemas</th><th>Status</th></tr>
      </thead>
      <tbody>
        <tr><td>1</td><td>Manufacturing v1</td><td>Plant Manager, Ops Dir</td><td>MFGML</td><td>GA</td></tr>
        <tr><td>2</td><td>Finance / Controllership</td><td>Controller, CFO</td><td>FINML</td><td>v1 Q4 FY27</td></tr>
        <tr><td>3</td><td>Risk &amp; Compliance</td><td>SOX PMO Lead, CRO</td><td>RISKML</td><td>v1 Q4 FY27</td></tr>
        <tr><td>4</td><td>Retail Merchandising v1</td><td>Chief Merchant, DMM, Planning Lead</td><td>MERML</td><td>v1 Q2 FY27</td></tr>
        <tr><td>5</td><td>Customer Experience v1</td><td>Care Ops Lead</td><td>CXML</td><td>v1 Q4 FY27</td></tr>
        <tr><td>6</td><td>ESG</td><td>Sustainability Officer</td><td>ESGML</td><td>v1 Q3 FY28</td></tr>
        <tr><td><strong>7</strong></td><td><strong>Customer Focused Merchandise (CFMP) v0.2</strong></td><td>CMO, CX-VP, Loyalty Director</td><td>CXML + MERML + new <code>cfmp</code> namespace</td><td><strong>Phase 1 LIVE</strong> on APEX-M dev substrate</td></tr>
      </tbody>
    </table>
"""


def core_extensions_block() -> str:
    return """\
    <h3 class="cfmp">Core extensions CFMP introduces</h3>
    <ul class="cfmp">
      <li><strong>New persona type — <code>customer</code></strong> · identity bound to loyalty ID;
        <code>consent_gate: required</code>; channel <code>customer_phone</code>. Reusable across all
        future customer-facing Packs (Banking-Customer, Telco-Customer, Hospitality-Guest).</li>
      <li><strong>New HITL transport — <code>customer_phone</code></strong> · extends interface #9.
        APEX-M binding: Azure Notification Hubs signed-link push + MCP callback. Same model on G/A.</li>
      <li><strong>Proposed Interface #15 — Maps &amp; Wayfinding</strong> · APEX-M: Azure Maps
        Creator (Indoor Maps + Wayfinding REST + Web SDK). APEX-G: Google Maps Geospatial Creator.
        APEX-A: Amazon Location Service indoor maps. Industry Packs consume the abstract interface;
        the cloud-profile YAML resolves the implementation.</li>
    </ul>
"""


def phase_status_block() -> str:
    return """\
    <h3 class="cfmp">Phase 1 status (live on APEX-M dev substrate today)</h3>
    <ul class="cfmp">
      <li><span class="cfmp-tag cfmp-live">LIVE</span> Vendored apex_core + apex_audit + apex_purview
        in the orchestrator container; signed 14-field LedgerRow per <code>/agent/ask</code>;
        Atlas-shaped Purview lineage edges buffered offline.</li>
      <li><span class="cfmp-tag cfmp-live">LIVE</span> Microsoft Agent Framework 1.6.0 (GA) with
        provider toggle (Azure OpenAI default → Anthropic optional); default model
        <code>gpt-5-mini</code>; five agent tools including <code>route_to_product</code> riding
        Interface #15.</li>
      <li><span class="cfmp-tag cfmp-live">LIVE</span> Azure Speech (en-US-AvaMultilingualNeural)
        STT + TTS via <code>/agent/tts</code>; HITL consent gate at $50 cart-add.</li>
      <li><span class="cfmp-tag cfmp-planned">PHASE 2</span> Provision Fabric capacity (F2 trial);
        land 800-product catalog in Bronze → Silver (canonical MERML) → Gold (Direct Lake);
        AuditStore swaps in-memory for OneLake WORM Bronze.</li>
      <li><span class="cfmp-tag cfmp-planned">PHASE 3</span> Provision Purview; install APEX
        classifications via CI; flush lineage buffer to live Atlas REST; drop AOAI key in favor of
        Container App managed identity.</li>
      <li><span class="cfmp-tag cfmp-planned">PHASE 4</span> Per-tenant Azure Maps Creator resource;
        upload retailer Drawing Package; embed Web SDK indoor view in the portal.</li>
    </ul>
"""


def commercial_block() -> str:
    return """\
    <h3 class="cfmp">Commercial envelope (sub-tier ladder)</h3>
    <table class="cfmp">
      <thead><tr><th>Sub-tier</th><th>Scope</th><th>Price band</th><th>Duration</th><th>Funding</th></tr></thead>
      <tbody>
        <tr><td><strong>Pack Lite</strong></td><td>3 scenarios + wayfinding · 1 store</td>
            <td>$150K–$250K</td><td>4–6 weeks</td><td>BVA + DCIF</td></tr>
        <tr><td><strong>Pack Standard</strong></td><td>10 scenarios · 5 stores</td>
            <td>$500K–$1.5M</td><td>12–16 weeks</td><td>DCIF + Client + ISV burndown</td></tr>
        <tr><td><strong>Pack Enterprise</strong></td><td>All 18 scenarios · full chain · Operate-ready</td>
            <td>$1.5M–$3.5M</td><td>6–9 months</td><td>Client direct + T&amp;M</td></tr>
      </tbody>
    </table>
"""


def artifacts_block(book_name: str) -> str:
    return f"""\
    <h3 class="cfmp">Companion artifacts</h3>
    <ul class="cfmp">
      <li><code>docs/packs/CFMP-v0.2.md</code> — Pack design document (12 sections).</li>
      <li><code>docs/packs/CFMP-Scenario-Chains-v0.2.xlsx</code> — 18-scenario workbook
        (5 sheets: Summary · Scenario Library · Featured Chains · Scenario→KPI Chain · Pack Sub-Tiers).</li>
      <li><code>docs/packs/APEX-Architecture-v5.1-with-CFMP-chapter.docx</code> — parent architecture
        doc with new Chapter 26 (CFMP v0.2) inserted before Appendix A.</li>
      <li><strong>Live demo</strong> — <a href="https://ca-visionkit-portal.gentlestone-9b49b099.eastus2.azurecontainerapps.io/architecture"
        style="color:#1c4e80;">/architecture page</a> with planned-vs-deployed status, dependency tree,
        per-resource detail modals, and APEX accelerator value section.</li>
    </ul>
    <p style="color:#7d8590;font-size:0.85em;margin-top:1em">
      <em>Inserted into <code>{book_name}</code> automatically by
      <code>docs/packs/_update_book_volumes.py</code> · {TODAY}.
      Originals preserved as <code>*-pre-cfmp-v0.2.html.bak</code>.</em>
    </p>
"""


# ─── Per-book section templates ───────────────────────────────────────────

def section_full(book_name: str) -> str:
    """Long form for big books (Sellers Guide, Services Guide, Implementation,
    Deployment, Library, Sellers Handbook, M-Merch arch brief)."""
    body = (
        "<p><strong>This addendum extends the book with the seventh Industry "
        "Solution Pack — Customer Focused Merchandise Pack (CFMP) v0.2 — and "
        "documents the six new Virtual Views, three new APEX Core extensions, "
        "and Phase-1 live status of the working demo on APEX-M.</strong></p>"
        + pack_catalog_table()
        + new_vv_table()
        + core_extensions_block()
        + phase_status_block()
        + commercial_block()
        + artifacts_block(book_name)
    )
    return render_section(
        "Addendum · Customer Focused Merchandise Pack (CFMP) v0.2",
        body,
    )


def section_short(book_name: str) -> str:
    """Short form for the Executive Summary."""
    body = (
        "<p>APEX-M now ships <strong>seven Industry Solution Packs</strong> "
        "with the addition of <strong>Customer Focused Merchandise (CFMP) "
        "v0.2</strong> — the first Pack organized around a customer-moment "
        "spine (choose / select / buy / services) rather than an operator "
        "function. Sibling to the operator-side Retail Merchandising Pack.</p>"
        "<p><strong>Phase 1 live on APEX-M dev substrate.</strong> "
        "18 scenarios mapped to the customer journey, six new Virtual Views, "
        "two new Core extensions (<code>customer</code> persona type + "
        "<code>customer_phone</code> HITL transport), one proposed new Core "
        "interface (#15 Maps &amp; Wayfinding) with Azure Maps Creator as the "
        "APEX-M binding. Pack Lite at $150K–$250K · 4–6 weeks.</p>"
        + artifacts_block(book_name)
    )
    return render_section(
        "Addendum · CFMP v0.2 — 7th Industry Solution Pack",
        body,
    )


def section_snowflake(book_name: str) -> str:
    """Snowflake doc cross-reference."""
    body = (
        "<p>This Snowflake integration brief describes the data-plane "
        "alternative to the Microsoft Fabric path that the CFMP Pack uses on "
        "APEX-M. <strong>CFMP itself runs on APEX-M today</strong> with the "
        "Microsoft Fabric path (Bronze/Silver/Gold on OneLake). When a "
        "retailer is Snowflake-primary on the data side but APEX-M-primary on "
        "the agent runtime side, the federation pattern documented in this "
        "brief allows the same CFMP scenarios to land on Snowflake-backed "
        "Silver/Gold while keeping the LEDGER + lineage in the APEX-M "
        "primary. Interface #15 (Maps &amp; Wayfinding via Azure Maps "
        "Creator) and the <code>customer</code> persona type are unchanged.</p>"
        + artifacts_block(book_name)
    )
    return render_section(
        "Addendum · CFMP v0.2 + Snowflake data-plane",
        body,
    )


# ─── Per-volume routing table ─────────────────────────────────────────────

VOLUMES = [
    ("APEX-M-Merch-Agentic-Client-Architecture.html",            section_full),
    ("APEX-Snowflake-Integration-Architecture.html",             section_snowflake),
    ("Professional-APEX-M-Deployment-Guide.html",                section_full),
    ("Professional-APEX-M-Executive-Summary.html",               section_short),
    ("Professional-APEX-M-Implementation-Guide-(Vuori-Example).html", section_full),
    ("Professional-APEX-M-Library.html",                         section_full),
    ("Professional-APEX-M-Sellers-Guide.html",                   section_full),
    ("Professional-APEX-M-Services-Guide.html",                  section_full),
    ("Sellers-Handbook-Agentic-AI-on-Microsoft.html",            section_full),
]


# ─── HTML mutation ───────────────────────────────────────────────────────

def update_html(path: Path, section_fn) -> tuple[str, int]:
    """Update one HTML volume; return (status, new size).
    Backs up to *.pre-cfmp-v0.2.html.bak (idempotent first-time)."""
    raw = path.read_text(encoding="utf-8", errors="replace")
    backup = path.with_suffix(".pre-cfmp-v0.2.html.bak")
    if not backup.exists():
        backup.write_text(raw, encoding="utf-8")

    # Strip any existing CFMP addendum (idempotent re-runs)
    pattern = re.compile(
        re.escape(SENTINEL_START) + r".*?" + re.escape(SENTINEL_END) + r"\s*",
        re.DOTALL,
    )
    if pattern.search(raw):
        raw = pattern.sub("", raw)
        status = "re-inserted"
    else:
        status = "first-insert"

    section_html = section_fn(path.name)

    # Insert before </body> (case-insensitive)
    m = re.search(r"</body>", raw, re.IGNORECASE)
    if not m:
        # Should not happen given inventory; append at end as fallback.
        raw = raw + "\n" + section_html + "\n"
    else:
        idx = m.start()
        raw = raw[:idx] + section_html + "\n" + raw[idx:]

    path.write_text(raw, encoding="utf-8")
    return status, len(raw)


# ─── docx + md updates ────────────────────────────────────────────────────

def update_runtime_addendum_docx() -> None:
    """Append a CFMP section to the Sellers Guide Runtime Addendum docx."""
    docx_path = BOOK_DIR / "APEX-Sellers-Guide-Runtime-Addendum.docx"
    if not docx_path.exists():
        print(f"  (skip) {docx_path.name} not found")
        return
    from docx import Document
    doc = Document(str(docx_path))
    backup = docx_path.with_suffix(".pre-cfmp-v0.2.docx.bak")
    if not backup.exists():
        # python-docx doesn't preserve everything on save, so make a binary backup
        backup.write_bytes(docx_path.read_bytes())

    # If a "CFMP v0.2 Addendum" heading already exists, skip (idempotent)
    for p in doc.paragraphs:
        if "CFMP v0.2 Addendum" in p.text:
            print(f"  {docx_path.name}: already has CFMP addendum")
            return

    # Available styles vary across docs (some lack Heading 1 / List Bullet
    # entirely). Build a defensive style picker so we never crash.
    available = {s.name for s in doc.styles}
    def pick(*names: str) -> str | None:
        for n in names:
            if n in available:
                return n
        return None

    h1_style     = pick("Heading 1", "heading 1", "Title", "Subtitle")
    h2_style     = pick("Heading 2", "heading 2", "Heading 1", "Title")
    body_style   = pick("Normal", "Body Text")
    bullet_style = pick("List Bullet", "List Paragraph", "Normal", "Body Text")

    def add_h(text: str, level: int) -> None:
        style = h1_style if level == 1 else h2_style
        p = doc.add_paragraph(text)
        if style:
            try:
                p.style = doc.styles[style]
            except KeyError:
                pass
        # Force bold if we couldn't get a proper heading style
        if not style or "Heading" not in style:
            for r in p.runs:
                r.bold = True
                r.font.size = None  # let default apply

    def add_p(text: str, bullet: bool = False) -> None:
        p = doc.add_paragraph(text)
        style = bullet_style if bullet else body_style
        if style:
            try:
                p.style = doc.styles[style]
            except KeyError:
                pass

    doc.add_page_break()
    add_h("CFMP v0.2 Addendum — 7th Industry Solution Pack", 1)
    add_p(
        f"Inserted {TODAY}. APEX-M now ships seven Industry Solution Packs "
        "with Customer Focused Merchandise (CFMP) v0.2 — the first Pack "
        "organized around a customer-moment spine. Phase 1 live on APEX-M "
        "dev substrate today."
    )
    add_h("New Virtual Views", 2)
    for vv, desc in NEW_VVS:
        add_p(f"{vv} — {desc}", bullet=True)
    add_h("Core extensions", 2)
    for line in [
        "New persona type `customer` — loyalty-ID-bound, consent-gated, channel = customer_phone. Reusable across all customer-facing Packs.",
        "New HITL transport `customer_phone` extending interface #9. APEX-M binding: Azure Notification Hubs.",
        "Proposed Interface #15 — Maps & Wayfinding. APEX-M binding: Azure Maps Creator (Indoor Maps + Wayfinding REST + Web SDK).",
    ]:
        add_p(line, bullet=True)
    add_h("Phase 1 live status", 2)
    for line in [
        "Vendored apex_core + apex_audit + apex_purview in the orchestrator container.",
        "Signed 14-field LedgerRow per /agent/ask with HMAC-SHA256 chain and three-version stamps.",
        "Atlas-shaped Purview lineage edges (mcp→agent, agent→orchestration, orchestration→audit) buffered offline.",
        "Microsoft Agent Framework 1.6.0 (GA) with provider toggle (Azure OpenAI / Anthropic).",
        "Azure Speech (en-US-AvaMultilingualNeural) STT + TTS via /agent/tts.",
        "HITL consent gate at $50 cart-add.",
    ]:
        add_p(line, bullet=True)
    add_h("Companion artifacts", 2)
    for line in [
        "docs/packs/CFMP-v0.2.md — Pack design (12 sections)",
        "docs/packs/CFMP-Scenario-Chains-v0.2.xlsx — 18-scenario workbook",
        "docs/packs/APEX-Architecture-v5.1-with-CFMP-chapter.docx — parent arch doc + Chapter 26",
        "Live demo: https://ca-visionkit-portal.gentlestone-9b49b099.eastus2.azurecontainerapps.io/architecture",
    ]:
        add_p(line, bullet=True)

    doc.save(str(docx_path))
    print(f"  {docx_path.name}: updated (size {docx_path.stat().st_size:,} bytes)")


def update_trilogy_md() -> None:
    md = BOOK_DIR / "TRILOGY-MS-PLATFORM-VALIDATION.md"
    if not md.exists():
        print(f"  (skip) {md.name} not found")
        return
    raw = md.read_text(encoding="utf-8")
    backup = md.with_suffix(".pre-cfmp-v0.2.md.bak")
    if not backup.exists():
        backup.write_text(raw, encoding="utf-8")

    sentinel = "<!-- CFMP-V0.2-ADDENDUM -->"
    if sentinel in raw:
        # Strip prior addendum (idempotent)
        raw = re.sub(
            re.escape(sentinel) + r".*?" + re.escape("<!-- /CFMP-V0.2-ADDENDUM -->"),
            "",
            raw,
            flags=re.DOTALL,
        )

    addendum = f"""\

{sentinel}

## CFMP v0.2 — APEX-M validation update ({TODAY})

The APEX-M platform-validation matrix now includes a seventh Industry Pack: **Customer Focused Merchandise (CFMP) v0.2**, the first Pack organized around a customer-moment spine.

### What's validated as of Phase 1 live

- **Microsoft Agent Framework 1.6.0 (GA)** running gpt-5-mini against Azure OpenAI (Responses API, `api_version=v1`). Provider toggle to Anthropic verified in code.
- **apex_audit / apex_purview** vendored packages stamp a 14-field signed LedgerRow per `/agent/ask` and emit Atlas-shaped lineage edges (mcp→agent, agent→orchestration, orchestration→audit).
- **Azure Speech** (`en-US-AvaMultilingualNeural`) STT + TTS via `/agent/tts`.
- **Azure Container Apps** consumption-plan deployment (portal + orchestrator + ACR + cae-visionkit environment) in `rg-iot-visionkit` / `Global_RnD_Agentic_MERCH`.
- **PostgreSQL + pgvector** backing the 800-product MERML-aligned catalog with IVFFlat index and `+0.15` priced-row bias.
- **HITL consent gate** firing on cart-add ≥ $50 (configurable threshold).

### Proposed Interface #15 — Maps & Wayfinding

APEX-M binding: **Azure Maps Creator** (Indoor Maps + Wayfinding REST + Web SDK). Local `storemap.yaml` fallback active in `orchestrator/azure_maps.py` today; production activation on `DEMO_AZURE_MAPS_DATASET_ID` + key.

### Companion artifacts

- `docs/packs/CFMP-v0.2.md` — Pack design document
- `docs/packs/CFMP-Scenario-Chains-v0.2.xlsx` — 18-scenario workbook
- `docs/packs/APEX-Architecture-v5.1-with-CFMP-chapter.docx` — parent arch doc with new Chapter 26
- Live demo: <https://ca-visionkit-portal.gentlestone-9b49b099.eastus2.azurecontainerapps.io/architecture>

<!-- /CFMP-V0.2-ADDENDUM -->
"""
    md.write_text(raw.rstrip() + addendum, encoding="utf-8")
    print(f"  {md.name}: updated")


# ─── main ────────────────────────────────────────────────────────────────

def main() -> None:
    print(f"=== Updating book volumes in {BOOK_DIR} ===\n")
    for name, fn in VOLUMES:
        path = BOOK_DIR / name
        if not path.exists():
            print(f"  (skip) {name} — not found")
            continue
        before_size = path.stat().st_size
        status, after_size = update_html(path, fn)
        print(f"  {name}: {status} ({before_size:,} → {after_size:,} bytes)")
    print()
    update_runtime_addendum_docx()
    update_trilogy_md()
    print("\nDone. Originals saved as *.pre-cfmp-v0.2.html.bak / .docx.bak / .md.bak.")


if __name__ == "__main__":
    main()
