"""Dump the structure of the existing APEX RC E2E walkthrough so we can model the new one on it."""
from __future__ import annotations
import io, sys
from docx import Document
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

DOC = r"C:\Stage\Clients\Industries\Consumer\Retail\Kroger\02_projects\FY27_Pipeline\assortment-pricing-agentic\deliverables\Services\RC-E2E-03_Assortment-and-Pricing\Tier0-Foundation\APEX-RC-E2E-03-Walkthrough.docx"

doc = Document(DOC)
print(f"Paragraphs: {len(doc.paragraphs)}")
print(f"Tables: {len(doc.tables)}")
print(f"Sections: {len(doc.sections)}")
print()

# Outline (style + first chars)
print("=" * 100)
print("OUTLINE (style + 140 chars)")
print("=" * 100)
for i, p in enumerate(doc.paragraphs):
    style = p.style.name if p.style else "?"
    txt = p.text.strip()
    if not txt:
        continue
    print(f"[{i:4}] {style[:22]:<22} | {txt[:140]}")

print()
print("=" * 100)
print("TABLES")
print("=" * 100)
for ti, t in enumerate(doc.tables):
    print(f"\n--- Table {ti}: {len(t.rows)} rows x {len(t.columns)} cols ---")
    for ri, row in enumerate(t.rows[:6]):
        cells = [c.text.strip().replace("\n", " // ")[:60] for c in row.cells]
        print(f"  R{ri}: {cells}")
    if len(t.rows) > 6:
        print(f"  ... and {len(t.rows) - 6} more rows")
