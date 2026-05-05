"""Build the APEX · RC-E2E-03 service walkthrough (Kroger reference client).

Mirrors the Nike R2R / Disney TMT-MED-01 walkthrough pattern:
  · 11 sections + 2 appendices
  · ~30 tables, no figures
  · Reference-client framing — Kroger anchors RC-E2E-03 the same way Nike
    anchors RC-E2E-06 and Disney+ anchors TMT-MED-01

Service:        RC-E2E-03 — Assortment & Pricing  (highest-value RC service)
Secondary:      RC-E2E-09 — Product Tracking      (FSMA 204 / lot-provenance)
Reference:      The Kroger Co. — 2,800+ stores · ~$150B revenue · 60M households
APEX scenario:  RC-SUPCHN-01 — Cold-chain excursion · marquee chain
Output:         Kroger 02_projects/FY27_Pipeline/assortment-pricing-agentic/deliverables/
"""

from __future__ import annotations

import io
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

OUT_DIR = Path(r"C:\Stage\Clients\Industries\Consumer\Retail\Kroger\02_projects\FY27_Pipeline\assortment-pricing-agentic\deliverables")
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT = OUT_DIR / "APEX-RC-E2E-03-Walkthrough.docx"

# ---- styling helpers ----------------------------------------------------

INK   = "0B0B0F"
INK2  = "4A4A55"
MUTE  = "8a8a90"
GOLD  = "B8892E"
TEAL  = "2A7F73"
VIO   = "6648B0"
BAND  = "F3F3F5"
HEAD  = "0B0B0F"
LINE  = "D4D4D7"


def set_cell_bg(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex.lstrip("#"))
    tcPr.append(shd)


def set_cell_border(cell, color_hex=LINE):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), color_hex.lstrip("#"))
        borders.append(b)
    tcPr.append(borders)


def set_run(run, *, bold=False, italic=False, size=None, color=None, font="Arial"):
    if bold: run.bold = True
    if italic: run.italic = True
    if size is not None: run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor.from_string(color.lstrip("#"))
    if font: run.font.name = font


def add_body(doc, text, *, size=11, italic=False, color=INK, space_after=6, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None: p.alignment = align
    r = p.add_run(text)
    set_run(r, size=size, color=color, italic=italic)
    return p


def add_h(doc, text, *, level=1):
    sizes = {1: 18, 2: 14, 3: 12.5}
    p = doc.add_paragraph()
    p.style = doc.styles[f"Heading {level}"]
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(4)
    for r in list(p.runs):
        r.text = ""
    r = p.add_run(text)
    set_run(r, bold=True, size=sizes.get(level, 12), color=INK)
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), GOLD)
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p


def add_bullets(doc, items, *, size=11):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(it)
        set_run(r, size=size, color=INK)


def add_table(doc, header, rows, *, col_widths=None,
              header_fill=HEAD, header_color="FFFFFF", zebra=True):
    n = len(header)
    tbl = doc.add_table(rows=len(rows) + 1, cols=n)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(header):
        c = tbl.rows[0].cells[j]
        set_cell_bg(c, header_fill)
        set_cell_border(c)
        c.paragraphs[0].text = ""
        r = c.paragraphs[0].add_run(h)
        set_run(r, bold=True, size=10, color=header_color)
        if col_widths: c.width = Inches(col_widths[j])
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            c = tbl.rows[i].cells[j]
            set_cell_border(c)
            if zebra and i % 2 == 0:
                set_cell_bg(c, BAND)
            c.paragraphs[0].text = ""
            r = c.paragraphs[0].add_run(str(v))
            set_run(r, size=10, color=INK)
            if col_widths: c.width = Inches(col_widths[j])
    doc.add_paragraph()
    return tbl


def add_callout(doc, kind, body, *, accent=GOLD, fill="F6F1E4"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = tbl.rows[0].cells[0]
    set_cell_bg(c, fill)
    c.paragraphs[0].text = ""
    p_tag = c.paragraphs[0]
    r_tag = p_tag.add_run(kind.upper())
    set_run(r_tag, bold=True, size=9, color=accent)
    p = c.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    r = p.add_run(body)
    set_run(r, size=10.5, color=INK)
    doc.add_paragraph()


def add_kv(doc, pairs, *, lw=1.6, rw=4.9):
    tbl = doc.add_table(rows=len(pairs), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(pairs):
        ck, cv = tbl.rows[i].cells
        ck.width = Inches(lw)
        cv.width = Inches(rw)
        set_cell_bg(ck, BAND)
        set_cell_border(ck); set_cell_border(cv)
        ck.paragraphs[0].text = ""
        cv.paragraphs[0].text = ""
        rk = ck.paragraphs[0].add_run(k)
        set_run(rk, bold=True, size=10, color=INK)
        rv = cv.paragraphs[0].add_run(v)
        set_run(rv, size=10, color=INK)
    doc.add_paragraph()


def add_eyebrow(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text.upper())
    set_run(r, bold=True, size=9, color=MUTE)
    rPr = r._element.get_or_add_rPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:val"), "80")
    rPr.append(spacing)


def pagebreak(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, size=9.5, color=INK, font="Consolas")


# ============================================================================
# BUILD
# ============================================================================

def build():
    doc = Document()
    for section in doc.sections:
        section.page_width  = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = section.bottom_margin = Inches(0.9)
        section.left_margin = section.right_margin = Inches(1.0)
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # ============ COVER ============
    add_eyebrow(doc, "APEX · RC-E2E-03 · REFERENCE WALKTHROUGH · KROGER")
    p = doc.add_paragraph()
    r = p.add_run("APEX · RC-E2E-03")
    set_run(r, bold=True, size=26, color=INK)
    p1b = doc.add_paragraph()
    r1b = p1b.add_run("Assortment & Pricing")
    set_run(r1b, bold=True, size=20, color=INK)
    p2 = doc.add_paragraph()
    r2 = p2.add_run("Cold-chain markdown · perishables integrity · pricing decisioning — Kroger as reference client")
    set_run(r2, size=13, color=INK2, italic=True)

    doc.add_paragraph()
    add_body(doc,
             "A walkthrough of the APEX RC-E2E-03 Assortment & Pricing service across the APEX "
             "value-delivery chain — Scenario → Solution → Use Case → Service → Persona → KPI — "
             "anchored against The Kroger Co. as the Wave-1 reference client. The same five-agent "
             "service pattern documented here is the highest-leverage RC engagement for any "
             "grocery, mass, or perishables-heavy retailer at scale, and ports across to Albertsons, "
             "Ahold Delhaize, Walmart, Costco, and Target with only the manifest layer changing.",
             italic=True, color=INK2)

    doc.add_paragraph()
    add_kv(doc, [
        ("Document type",       "APEX Service Walkthrough · Service-to-KPI"),
        ("Service",             "RC-E2E-03 · Assortment & Pricing  (highest-value RC service)"),
        ("Secondary service",   "RC-E2E-09 · Product Tracking  (FSMA 204 / lot-provenance)"),
        ("APEX scenario",       "RC-SUPCHN-01 · Cold-chain excursion (store cooler) — marquee chain"),
        ("Practice",            "RC — Retail & Consumer · Grocery sub-Practice"),
        ("Reference client",    "The Kroger Co. — 2,800+ stores · ~$150B revenue · 60M Plus households"),
        ("Audience",            "CFO · COO · Chief Merchant · Food Safety Lead · CIO · External audit"),
        ("Chain links",         "Scenario · Solution · Use Case · Service · Persona · KPI"),
        ("Waves covered",       "Approach · W1 Foundation · W2 Pilot · W3 Scale & Fuse"),
        ("Companion artifacts", "RC-E2E-03-Sprint-Plan.xlsx · RC-E2E-03-Resource-Plan.xlsx · RC-E2E-03-ROI-Case.html"),
        ("Version / date",      "v1.0 · 27 Apr 2026"),
    ])

    pagebreak(doc)

    # ============ CONTENTS ============
    add_h(doc, "Contents", level=1)
    add_body(doc,
             "This document walks the RC-E2E-03 Assortment & Pricing service through every link of "
             "the APEX value-delivery chain. Sections 1–2 set context. Sections 3–8 are the chain "
             "itself: one section per chain link. Section 9 projects the chain across the three "
             "delivery waves; Section 10 summarizes the commercial envelope; Section 11 describes "
             "the reference Wave 1 implementation. Appendices A and B carry the glossary and the "
             "artifact cross-reference.",
             italic=True, color=INK2)

    add_table(doc,
              ["#", "Section", "What it covers"],
              [
                  ["1",  "The Reference Excursion at Kroger",
                   "The dairy-case event Kroger handles every day across 2,800+ stores and the "
                   "gap between detection and decision."],
                  ["2",  "Why Assortment & Pricing, Why Kroger, Why Now",
                   "The five forces converging on perishables and pricing; why Kroger is the RC "
                   "reference client; why FY27 is the build window."],
                  ["3",  "Chain Link 1 · Scenario",
                   "Kroger's perishables and pricing operating reality today; current-state "
                   "response; friction quantified; scope boundaries."],
                  ["4",  "Chain Link 2 · Solution",
                   "RC-E2E-03 in one sentence; five-agent orchestration; commercial envelope "
                   "tiering; three design commitments."],
                  ["5",  "Chain Link 3 · Use Cases",
                   "Six agent-level use cases composing the assortment-and-pricing scenario "
                   "family; mapping to the sprint plan; incremental delivery path."],
                  ["6",  "Chain Link 4 · Services",
                   "RC-E2E-03 Assortment & Pricing + RC-E2E-09 Product Tracking; service "
                   "composition; tenant manifest; candidate derived services."],
                  ["7",  "Chain Link 5 · Personas",
                   "Store Operations Lead as primary; five secondary personas (Food Safety Lead, "
                   "Pull-Team Lead, Pricing Manager, Inventory Control, FDA Auditor) with plane "
                   "bindings."],
                  ["8",  "Chain Link 6 · KPIs",
                   "Per-event KPIs, derived KPIs for Kroger commercial cases, audit-row "
                   "attribution."],
                  ["9",  "The Three Waves",
                   "W1 Foundation (single-banner pilot), W2 Pilot (250-store anchor cluster), "
                   "W3 Scale & Fuse (enterprise 2,800+)."],
                  ["10", "Commercial Envelope",
                   "Approach / W1 / W2 / W3 cost envelopes and durable-value reference."],
                  ["11", "Reference Wave 1 Implementation",
                   "What an RC-E2E-03 W1 prototype looks like at Kroger; demo paths; technology "
                   "stack; intentional scope boundaries."],
                  ["A",  "Glossary",
                   "Selected terms used in this document with APEX- or grocery-specific meaning."],
                  ["B",  "Artifact Cross-Reference",
                   "How this walkthrough fits with the sprint plan, resource plan, and ROI case."],
              ],
              col_widths=[0.4, 2.6, 3.5])

    pagebreak(doc)

    # ============ SECTION 1 · Reference Event ============
    add_h(doc, "Section 1 · The Reference Excursion at Kroger", level=1)
    add_body(doc,
             "At 6:08 AM on a Thursday in April 2027, a refrigeration sensor at Kroger Store #428 "
             "in Cincinnati reads 48.2°F. Dairy Case A3 — the floor-length open case carrying "
             "milk, yogurt, butter, and a daily-rotated selection of cheeses — has breached its "
             "41°F setpoint. Over the next 52 minutes, the case will drift between 46°F and 49°F "
             "before recovery. By 7:00 AM, when the store's Operations Lead arrives for opening, "
             "412 dairy units have been outside compliant temperature for an hour. The store opens "
             "in 30 minutes. Three decisions need to be made about those 412 units before the "
             "first customer touches a milk jug.")
    add_body(doc,
             "This is the scenario. Everything downstream — the RC-E2E-03 agent architecture, the "
             "sprint use cases, the productized services, the persona cast, the KPI rollups — "
             "exists to answer three questions about the 412 units and the 52 minutes between "
             "6:08 AM and 7:00 AM:")
    add_bullets(doc, [
        "Which units are safe to sell, which should be marked down, and which must be destroyed?",
        "Who at Kroger has the authority to make that call, and how do they make it in 90 seconds rather than the 12 minutes it takes today?",
        "What evidence survives the decision — auditable, FDA FSMA-grade, attributable to the specific 6:08 AM event at the specific Dairy Case A3 at the specific Kroger Store #428?",
    ])

    add_h(doc, "Today, without RC-E2E-03", level=3)
    add_body(doc,
             "Kroger's ~28,000-person store-operations workforce handles ~85 such events per week "
             "per banner across the enterprise. Today's response is a hand-off chain made of phone "
             "calls, text messages, printed temperature logs, and in-person walk-throughs: the "
             "Store Operations Lead discovers the excursion, walks the case (4–8 minutes), calls "
             "the refrigeration vendor (2–20 minutes hold time), pulls the overnight log, calls "
             "the Food Safety Compliance Lead for a judgment call (typical voicemail + callback "
             "delay), routes the disposition to the Pull Team for execution, and finally — typically "
             "within 24 hours — enters the event into a regulatory log for FSMA audit-trail "
             "purposes. The decision itself happens in roughly 12 minutes; the information assembly "
             "around it consumes the 52-minute pre-opening window.")

    add_h(doc, "Tomorrow, with RC-E2E-03", level=3)
    add_body(doc,
             "Five agents — Assess, Classify, Quantify, Approve, Act — run in parallel where the "
             "work allows, sequentially where it requires. Material disposition decisions (markdown "
             "above tier threshold, destroy-all events) route through the Store Operations Lead via "
             "Microsoft Teams Adaptive Card on her device. Every decision emits a 14-field audit "
             "row to LEDGER bound to FSMA 204 lot-provenance. The 12-minute decision compresses to "
             "90 seconds; the 5.2 hours per shift of manager attention reclaimed translates "
             "directly into store-floor presence. Audit evidence is generated, not reconstructed, "
             "at every regulatory touchpoint.")

    add_callout(doc, "Independence note",
                "RC-E2E-03's agent manifests, policy manifests, and LEDGER are Kroger artifacts; "
                "Deloitte's role is advisory and implementation enablement. The SOW is structured "
                "so that the engagement meets Deloitte independence requirements regardless of "
                "Kroger's external-audit firm relationship. The underlying APEX pattern is "
                "independence-safe by design; the specific engagement structure is designed in "
                "concert with Deloitte's Independence office.")

    # ============ SECTION 2 · Why ============
    add_h(doc, "Section 2 · Why Assortment & Pricing, Why Kroger, Why Now", level=1)

    add_h(doc, "2.1 · The five forces converge on the dairy case", level=2)
    add_body(doc,
             "APEX exists because five forces converged across FY25–FY26. For grocery assortment "
             "and pricing specifically, all five land on the perishables fixture: the "
             "decision-velocity gap that leaves operational signals as aged work; the agentic-AI "
             "inflection that crossed enterprise viability for store-by-store decisioning; the "
             "FDA FSMA 204 final-rule implementation deadline that demands lot-traceability "
             "evidence at the unit level; the food-safety reputational risk that compounds in the "
             "Walmart-Amazon era of consumer transparency; and the margin compression on shelf-"
             "stable categories that makes perishables-shrink reduction the single largest "
             "free-cash-flow lever available to a grocer of Kroger's scale.")

    add_h(doc, "2.2 · Why Kroger is the APEX RC reference client", level=2)
    add_body(doc,
             "Of the seven Practices APEX addresses, RC carries the reference assortment-and-"
             "pricing scenario because grocery has all the hardness of high-volume perishables "
             "operations plus distinctive operating realities that stress the agent architecture "
             "in ways that mass-merchant and apparel models do not. Kroger specifically anchors "
             "RC because Kroger combines four characteristics that no other US grocer matches at "
             "comparable scale:")
    add_bullets(doc, [
        "Multi-banner footprint — Kroger, Ralphs, Fred Meyer, King Soopers, Smith's, Fry's, Harris Teeter, Mariano's, QFC, Pick 'n Save, Roundy's, and others — each with distinct local pricing zones and assortment patterns",
        "Sophisticated pricing DNA — 84.51° subsidiary, a decade of dunnhumby relationship history, customer-data-driven pricing already in the muscle memory of the merchant organization",
        "Fresh-food intensity — ~30% of revenue is fresh (dairy, deli, meat, produce, bakery); ~$45B exposure means even a 1% shrink reduction is ~$450M of free cash",
        "FSMA 204 readiness mandate — final-rule effective 2026 demands lot-level traceability across high-risk foods, making Product Tracking (RC-E2E-09) operationally non-optional",
        "Loyalty footprint — 60M+ Kroger Plus households generates the demand-elasticity signal density that pricing decisions actually need",
    ])
    add_body(doc,
             "If RC-E2E-03's orchestration pattern handles Kroger's grocery assortment-and-pricing "
             "operations, it ports across the Deloitte DMTSP retail portfolio to Albertsons-Safeway, "
             "Ahold Delhaize (Stop & Shop, Food Lion, Giant), Walmart Supercenter grocery, Costco "
             "Wholesale, and Target. The architectural substrate is constant; manifest layer and "
             "persona binding shift per client.")

    add_h(doc, "2.3 · Why this document exists", level=2)
    add_body(doc,
             "Three conversations at Kroger need the same material with different emphasis. The "
             "CFO and COO need to see a credible shrink-and-margin story with ROI that survives a "
             "scrub against the existing 84.51°-driven pricing stack. The Chief Merchant and Food "
             "Safety Compliance Lead need to see an operating-model story that fits how the "
             "merchant and store-ops organizations actually work today. The CIO and Chief Data "
             "Officer need to see a Microsoft-native architecture that fits Kroger's existing data "
             "and identity estate without inventing net-new platform categories. This document is "
             "the executive walkthrough that anchors all three.")

    pagebreak(doc)

    # ============ SECTION 3 · SCENARIO ============
    add_h(doc, "Section 3 · Chain Link 1 · Scenario", level=1)
    add_body(doc,
             "The Scenario link of the APEX value-delivery chain is the customer's operational "
             "reality stated without abstraction. It is not yet a solution. It is the thing that "
             "has to be true before a solution can be designed against it.")

    add_h(doc, "3.1 · Kroger's situation", level=2)
    add_body(doc,
             "Kroger operates approximately 2,800 supermarkets across 35 states under multiple "
             "banners, supported by 33 manufacturing plants, 44 distribution centers, and the "
             "84.51° data subsidiary. Each store carries approximately 1.8 refrigerated fixtures "
             "in the dairy category alone, with additional fixture counts in deli, produce, "
             "frozen, and prepared-food departments. On an average operating day across the "
             "enterprise, approximately 85 cold-chain excursion events per banner are formally "
             "logged. These are events for which a temperature breach has produced a pending "
             "disposition decision — sell-through, mark down, or destroy.")
    add_body(doc,
             "The ambient conditions of each excursion matter. A dairy case that briefly exceeds "
             "threshold during a morning restocking with doors open is a different event from the "
             "same case drifting steadily above threshold for four hours overnight because a "
             "compressor failed. The current-state response treats both events similarly; the "
             "target-state response distinguishes them at decision time.")

    add_h(doc, "3.2 · The current-state response", level=2)
    add_body(doc,
             "Today the response across all 2,800 Kroger stores is a sequential hand-off chain "
             "made of phone calls, text messages, printed logs, and in-person walk-throughs. The "
             "typical sequence over a 52-minute pre-opening window:")
    add_bullets(doc, [
        "T+0: Store Operations Lead discovers the excursion (vendor-portal alert, sensor reading, or visual inspection during morning readiness).",
        "T+5: Lead walks the affected case and makes a visual assessment of the inventory (4–8 minutes depending on case size and SKU count).",
        "T+10: Lead contacts the refrigeration vendor's service line. Phone hold times vary 2–20 minutes; service ticket may or may not open.",
        "T+15: Lead pulls the overnight temperature log (printed nightly or accessible via a vendor portal with a shared-credential login).",
        "T+25: Lead contacts the Food Safety Compliance Lead for a judgment call. Phone or voicemail; callback delay is variable.",
        "T+35: Disposition is agreed; Lead directs the Pull Team to physically execute (markdown rack, destroy stack, sell-through).",
        "T+45: Inventory Control updates the store's inventory position manually or by overnight batch from shrink entries.",
        "T+24h: Food Safety Compliance Lead enters the event into a regulatory log for FSMA audit-trail purposes — the primary artifact available to FDA auditors.",
    ])

    add_h(doc, "3.3 · The friction, quantified", level=2)
    add_body(doc, "Measured over a representative 90-day period at a reference store cluster, the "
                  "current-state response produces the following friction profile:")
    add_table(doc,
              ["Dimension", "Typical value", "Implication"],
              [
                  ["Time from excursion detection to manager decision",
                   "17 min median", "Store-opening readiness compressed; first-customer-impact risk"],
                  ["Manager attention consumed per event (info gathering)",
                   "~9 min / event", "5.2 hrs/shift across all events; not store-floor presence"],
                  ["FSMA-204 audit-row completeness",
                   "~62% on 24-hour batch entry", "Regulator-grade evidence reconstructed manually"],
                  ["Markdown vs. destroy decision accuracy (vs. policy)",
                   "~71% conformance", "29% of decisions either over-destroy or sell-through-at-risk"],
                  ["Per-event shrink (destroy disposition)",
                   "$1,313 median", "$1.1B annualized at 85 events/wk × 2,800 stores"],
                  ["Vendor-ticket opening rate (post-event)",
                   "~46% of qualifying events", "Refrigeration root-cause cycle time elongated"],
                  ["Inventory-position update lag",
                   "Overnight batch median", "Shrink reporting trails operational reality"],
                  ["FDA FSMA 204 evidence-chain breaks",
                   "~7% of events", "Regulatory-finding exposure compounding into 2026 effective date"],
              ],
              col_widths=[2.4, 1.9, 2.2])

    add_h(doc, "3.4 · Scope of the Kroger scenario", level=2)
    add_body(doc, "The Kroger RC-E2E-03 scenario specifically includes:")
    add_bullets(doc, [
        "Refrigeration telemetry ingest from in-store sensors (OPC-UA over current EMS infrastructure)",
        "Per-event scoring against FSMA 204 policy corpus and per-SKU markdown elasticity",
        "Per-unit disposition recommendation (sell-through, markdown, destroy) with confidence and time-to-action",
        "Markdown-decisioning quantification using MERML.Price + MERML.Markdown + SCML.Lot",
        "HITL approval routing for above-threshold dispositions via Microsoft Teams Adaptive Card",
        "Downstream action orchestration: POS markdown upload, refrigeration-vendor ticket creation, FSMA regulatory entry",
        "Outcome attribution at trace-id granularity feeding shrink-NPV measurement and policy refinement",
        "FSMA 204 evidence packet generation — quarterly, on-demand for FDA inspection windows",
    ])
    add_body(doc, "It specifically excludes:")
    add_bullets(doc, [
        "Refrigeration-equipment maintenance scheduling (separate scenario — RC-ASSET-* asset-maintenance Practice)",
        "Customer loyalty engagement post-event (separate scenario under RC-E2E-04 Customer Lifecycle)",
        "Strategic merchandising decisions on category mix (separate scenario under RC-E2E-08 Customer Acquisition)",
        "Distribution-center cold-chain integrity (adjacent scenario; this engagement covers store-level only)",
    ])

    pagebreak(doc)

    # ============ SECTION 4 · SOLUTION ============
    add_h(doc, "Section 4 · Chain Link 2 · Solution", level=1)
    add_body(doc,
             "The Solution link is how APEX answers the Scenario. It is an architectural "
             "description — what runs, where it runs, what contracts bind the components, how "
             "humans stay in charge of the decisions that matter.")

    add_h(doc, "4.1 · RC-E2E-03 in one sentence", level=2)
    add_callout(doc, "Solution",
                "A five-agent assortment-and-pricing orchestration — Assess, Classify, Quantify, "
                "Approve, Act — composes with a HITL-gated commercial envelope, runs on "
                "Microsoft-native infrastructure (Fabric + Foundry + Azure OpenAI + Teams + Entra) "
                "and consumes CAF (Converge Agentic Framework) primitives via A2A, fuses MERML "
                "pricing with SCML lot-provenance and FSMA 204 policy, and writes every decision "
                "to a 14-field audit row so shrink-NPV and FSMA evidence are generated, not "
                "reconstructed.",
                accent=TEAL, fill="E4F1EE")

    add_h(doc, "4.2 · The five agents", level=2)
    add_table(doc,
              ["Agent", "Responsibility", "Input", "Output", "HITL?"],
              [
                  ["Assess",
                   "Characterize the excursion (peak temp, duration, avg, door-open events)",
                   "Refrigeration telemetry · SCML.Fixture state",
                   "Excursion profile + severity classification",
                   "No"],
                  ["Classify",
                   "Per-SKU disposition (sell-through / markdown / destroy) per FSMA 204 policy",
                   "Excursion profile + SCML.Lot + policy corpus",
                   "Per-SKU disposition table with confidence",
                   "If confidence < 0.85, escalate"],
                  ["Quantify",
                   "Compute dollar impact of each disposition path; propose markdown vector",
                   "Disposition table + MERML.Price + MERML.Markdown + MERML.Elasticity",
                   "Three disposition scenarios with $ impact",
                   "Rare (deterministic given pricing)"],
                  ["Approve",
                   "Route material decisions to named persona via Adaptive Card",
                   "Quantification + persona policy + commercial envelope",
                   "Approved disposition with HITL signature",
                   "Always for material entries"],
                  ["Act",
                   "Orchestrate downstream actions (POS markdown, vendor ticket, FSMA log)",
                   "Approved disposition + integration manifests",
                   "Three parallel downstream effects + close-out audit row",
                   "No (post-approval execution only)"],
              ],
              col_widths=[1.0, 1.7, 1.6, 1.5, 1.2])

    add_h(doc, "4.3 · Commercial envelope — Kroger tier thresholds", level=2)
    add_body(doc,
             "The commercial envelope is the policy layer that decides 'do this automatically vs. "
             "require human approval.' For Kroger's RC-E2E-03, the envelope is tiered by "
             "store-cluster materiality × disposition-dollar-amount × confidence:")
    add_table(doc,
              ["Cluster tier", "Disposition $", "Confidence", "Action"],
              [
                  ["Tier 1 (flagship banners — Kroger, Fred Meyer)", "Any",          "Any",     "Always HITL — Store Ops Lead"],
                  ["Tier 2 (standard)",                              "> $2,500 / event", "Any",  "HITL — Store Ops Lead"],
                  ["Tier 2 (standard)",                              "$500–$2,500",  "≥ 0.95",  "Auto-execute; notify Store Ops Lead"],
                  ["Tier 2 (standard)",                              "$500–$2,500",  "< 0.95",  "HITL — Store Ops Lead"],
                  ["Tier 3 (smaller-format)",                        "≤ $500",        "≥ 0.90",  "Auto-execute; log only"],
                  ["Any tier",                                        "Destroy-all event", "Any", "Always HITL — Food Safety Lead"],
                  ["Any tier",                                        "FSMA-evidence-gap event", "Any", "Always HITL — Food Safety Lead"],
              ],
              col_widths=[2.4, 1.6, 1.0, 1.6])

    add_h(doc, "4.4 · Three design commitments", level=2)

    add_h(doc, "4.4.1 · Manifest-driven, not model-driven", level=3)
    add_body(doc,
             "Every choice that could change between Kroger banners or close periods — which LLM "
             "is used for disposition narrative, which tools each agent can call, which policy "
             "corpus applies to FSMA 204 thresholds, which Store Operations Lead has approval "
             "authority for which fixture cluster — is declared in a Git-versioned manifest. "
             "Updating any of these is a manifest PR with review and merge discipline, not an "
             "ops-ticket. FSMA-policy versioning is part of the manifest history; auditors and "
             "FDA inspectors see exactly what the system was configured to do at the time of any "
             "specific disposition decision.")

    add_h(doc, "4.4.2 · Audit row, not log line", level=3)
    add_body(doc,
             "Every decision — every assessment, every classification, every quantification, every "
             "approval, every downstream action — emits a 14-field structured row carrying "
             "trace-id, three version stamps (manifest, policy, prompt), content-addressed input "
             "and output hashes, reasoning-trace reference, HITL status and approver, "
             "downstream-effect cross-reference, and KPI attribution. This is the artifact that "
             "turns 'Kroger food safety' from a binder of paper into an auditable substrate.")

    add_h(doc, "4.4.3 · HITL is first-class, not bolted on", level=3)
    add_body(doc,
             "The HITL gate between Quantify and Act is not a retro-fitted human check. It is "
             "declared in the orchestration manifest as a first-class primitive. Its gate kind "
             "(ZERO_TOUCH, ACK_ONLY, HITL, or ESCALATION) is deterministically mapped from the "
             "commercial envelope. The Store Operations Lead sees only what she must see; "
             "everything else dispatches automatically with audit-row attribution. Material "
             "events arrive as Microsoft Teams Adaptive Cards on her device with one-click "
             "approve / modify / reject.")

    pagebreak(doc)

    # ============ SECTION 5 · USE CASES ============
    add_h(doc, "Section 5 · Chain Link 3 · Use Cases", level=1)
    add_body(doc,
             "The Use Case link is the set of discrete operational patterns the Solution "
             "activates. One Solution typically enables multiple Use Cases — each Use Case is a "
             "bounded, testable, individually-measurable pattern that can be delivered as a piece "
             "of a larger engagement or as a standalone increment.")
    add_body(doc,
             "RC-E2E-03's six primary use cases (at the agent orchestration level) compose with "
             "Kroger's 24 detailed sprint-level use cases documented in RC-E2E-03-Sprint-Plan.xlsx. "
             "The six agent-level use cases below are individually deliverable, individually "
             "measurable, and individually commercially-envelopable.")

    for title, payload in [
        ("5.1 · Use Case · Excursion Triage",
         [("Trigger", "Data Activator reflex detects refrigerated fixture breaching setpoint > 15 min"),
          ("Agents",  "Assess"),
          ("Tools",   "OPC-UA telemetry ingest · SCML.Fixture schema validator"),
          ("HITL",    "None — deterministic profile computation"),
          ("Envelope","Included in W1 foundation")]),
        ("5.2 · Use Case · FSMA 204 Policy Classification",
         [("Trigger", "Excursion profile emitted by Triage"),
          ("Agents",  "Classify"),
          ("Tools",   "SCML.Lot retrieval · FSMA 204 policy corpus · LLM (CAF cap:disposition-reasoning)"),
          ("HITL",    "Confidence < 0.85 escalates to Food Safety Compliance Lead"),
          ("Envelope","~25% of W2 scope")]),
        ("5.3 · Use Case · Markdown Decisioning",
         [("Trigger", "Classification emitted by Policy-Classify"),
          ("Agents",  "Quantify"),
          ("Tools",   "MERML.Price · MERML.Markdown · MERML.Elasticity · SCML.Lot expiry"),
          ("HITL",    "Rare — deterministic given pricing/elasticity inputs"),
          ("Envelope","~20% of W2 scope (highest revenue surface)")]),
        ("5.4 · Use Case · Manager-Approval HITL",
         [("Trigger", "Quantification emitted by Markdown Decisioning"),
          ("Agents",  "None (approvals-mcp orchestrates the human)"),
          ("Tools",   "Microsoft Teams Adaptive Card · Entra identity binding · LEDGER write"),
          ("HITL",    "Always for material events; always for destroy-all"),
          ("Envelope","~15% of W2 scope")]),
        ("5.5 · Use Case · Parallel Downstream Action",
         [("Trigger", "HITL approval recorded"),
          ("Agents",  "Act"),
          ("Tools",   "POS markdown upload · refrigeration-vendor portal · FSMA regulatory log"),
          ("HITL",    "None — post-approval execution only"),
          ("Envelope","~25% of W2 scope")]),
        ("5.6 · Use Case · Outcome Attribution",
         [("Trigger", "Act completion"),
          ("Agents",  "None (orchestrator emits closeout audit row)"),
          ("Tools",   "LEDGER · 14-field audit-row writer"),
          ("HITL",    "None — write-only; consumed by Food Safety Lead and FDA Auditor"),
          ("Envelope","~15% of W2 scope; foundation for FSMA evidence packet")]),
    ]:
        add_h(doc, title, level=3)
        add_kv(doc, payload, lw=1.2, rw=5.0)

    add_h(doc, "5.7 · Mapping to the Kroger sprint plan's 24 use cases", level=2)
    add_body(doc,
             "The six agent-level use cases above are each realized by multiple sprint-level use "
             "cases in RC-E2E-03-Sprint-Plan.xlsx. The sprint plan refines the agent-level surface "
             "with specific user stories, acceptance criteria, and Wave assignment.")
    add_table(doc,
              ["Agent-level use case", "Sprint-level use cases (RC-E2E-03-Sprint-Plan.xlsx)"],
              [
                  ["5.1 Excursion Triage",                  "UC-01 · UC-02 · UC-03"],
                  ["5.2 FSMA 204 Policy Classification",    "UC-04 · UC-05 · UC-06 · UC-07"],
                  ["5.3 Markdown Decisioning",              "UC-08 · UC-09 · UC-10 · UC-11 · UC-12"],
                  ["5.4 Manager-Approval HITL",             "UC-13 · UC-14 · UC-15"],
                  ["5.5 Parallel Downstream Action",        "UC-16 · UC-17 · UC-18 · UC-19 · UC-20 · UC-21"],
                  ["5.6 Outcome Attribution",               "UC-22 · UC-23 · UC-24"],
              ],
              col_widths=[2.6, 3.7])

    add_h(doc, "5.8 · Incremental delivery path", level=2)
    add_body(doc,
             "Individually, each use case is a deliverable. Together they compose the full "
             "scenario. The composition is deliberate — Kroger can start with the first three "
             "agent-level use cases (Triage + Classify + Quantify) as a 'Cold-Chain Insight' "
             "package — read-only intelligence, no automated action, smaller envelope. Adding "
             "Manager-Approval HITL converts that into a managed-decisioning service. Adding "
             "Parallel Downstream Action and Outcome Attribution closes the loop end-to-end and "
             "lights up the FSMA evidence packet.")

    pagebreak(doc)

    # ============ SECTION 6 · SERVICES ============
    add_h(doc, "Section 6 · Chain Link 4 · Services", level=1)
    add_body(doc,
             "The Service link is what gets productized. A Service is a named, subscribable APEX "
             "capability with a defined persona, SLO, commercial terms, and an anchor in the APEX "
             "Service Catalog.")

    add_h(doc, "6.1 · Anchor services", level=2)

    add_h(doc, "RC-E2E-03 · Assortment & Pricing", level=3)
    add_body(doc,
             "The foundational pricing-and-disposition decisioning service. Owns the five-agent "
             "orchestration from Assess through Act. Primary persona: Store Operations Lead. SLO: "
             "decision-loop time 12 min → 90 sec at W2 steady state; ≥ 95% conformance to FSMA 204 "
             "policy. Commercial envelope: tiered by store-cluster size, banner mix, and "
             "perishables-category depth.")

    add_h(doc, "RC-E2E-09 · Product Tracking", level=3)
    add_body(doc,
             "The lot-provenance and FSMA 204 compliance service that sits underneath the "
             "Classify and Act agents. Primary persona: Food Safety Compliance Lead + FDA Auditor "
             "(consumer of evidence). SLO: lot-traceability query response under 100ms; quarterly "
             "FSMA evidence packet within 72 hours of period end. Commercial envelope: tiered by "
             "high-risk-food category breadth (the FSMA 204 final rule covers a defined Food "
             "Traceability List), supplier integration depth, and inspection-window cadence.")

    add_h(doc, "6.2 · Service composition", level=2)
    add_table(doc,
              ["Commitment", "RC-E2E-03 Assortment & Pricing",
                            "RC-E2E-09 Product Tracking"],
              [
                  ["Canonical schemas",
                   "MERML.Price · MERML.Markdown · MERML.Elasticity · SCML.Fixture",
                   "SCML.Lot · SCML.Inventory · SCML.Receipt"],
                  ["Agents deployed",
                   "Assess · Classify · Quantify · Approve · Act",
                   "Lot-Trace · Provenance-Reconcile · Evidence-Compose"],
                  ["MCP tool surfaces",
                   "POS markdown upload · refrigeration-vendor ticket · OPC-UA ingest",
                   "Lot-traceability query · FSMA 204 evidence export · supplier-integration"],
                  ["HITL policies",
                   "tier × disposition $ × confidence",
                   "exceptions on lot-trace gaps · destroy-all events"],
                  ["Power BI surfaces",
                   "Shrink dashboard · markdown-NPV trending · per-store cluster drilldown",
                   "FSMA evidence dashboard · lot-trace coverage · audit-trail completeness"],
                  ["Teams surfaces",
                   "Above-threshold disposition Adaptive Card",
                   "FSMA-gap notification card"],
                  ["SLO",
                   "Decision-loop time · markdown-NPV · policy conformance",
                   "Lot-query latency · evidence-packet assembly · audit-row coverage"],
                  ["Primary persona",
                   "Store Operations Lead",
                   "Food Safety Compliance Lead"],
              ],
              col_widths=[1.6, 2.5, 2.5])

    add_h(doc, "6.3 · Kroger tenant manifest subscription", level=2)
    add_body(doc, "The Kroger tenant manifest subscribes to the two services by declaration:")
    add_code(doc,
             '{ "tenant_id": "kr", "subscriptions": [\n'
             '  { "service_id": "RC-E2E-03", "service_version": "v1.2.0" },\n'
             '  { "service_id": "RC-E2E-09", "service_version": "v1.0.0" }\n'
             ']}')
    add_body(doc,
             "Version pinning is deliberate. When the APEX Service Catalog releases v1.3.0 of "
             "RC-E2E-03, Kroger decides when to upgrade. PATCH bumps are automatic; MINOR and "
             "MAJOR bumps require an explicit accept-upgrade HITL gate. This is how Kroger stays "
             "on stable versions while the catalog continues to evolve.")

    add_h(doc, "6.4 · Candidate derived services (post-W3)", level=2)
    add_body(doc,
             "RC-E2E-03 and RC-E2E-09 anchor the cold-chain scenario today. Three adjacent service "
             "candidates share the same architectural substrate and become fusion candidates at "
             "the end of Wave 3:")
    add_bullets(doc, [
        "RC-E2E-11 · Perishables Integrity Response — extends Excursion Triage with adjacent perishables scenarios (produce-ripeness, frozen-thaw-refreeze, prepared-foods-holding-time)",
        "RC-E2E-12 · Regulator-Grade Evidence Service — horizontal service offering pre-built FDA FSMA 204 auditor workbench, state-health-department complaint response, class-action discovery support",
        "RC-E2E-13 · Manager Time-Return Service — productizes the +5.2 hrs/shift reclamation across multiple RC scenarios (returns-fraud triage, OSA task generation, shift-handover automation)",
    ])

    pagebreak(doc)

    # ============ SECTION 7 · PERSONAS ============
    add_h(doc, "Section 7 · Chain Link 5 · Personas", level=1)
    add_body(doc,
             "The Persona link is the human-role cast that makes the scenario operational. "
             "Persona depth is what separates an architecturally correct solution from one that "
             "lands on the floor.")

    add_h(doc, "7.1 · Primary persona", level=2)
    add_callout(doc, "PRIMARY · Store Operations Lead",
                "Owns the per-store pre-opening readiness, the disposition policy, and the HITL "
                "approval path for every above-threshold event. Approves via Microsoft Teams "
                "Adaptive Card on phone or laptop; her approvals, modifications, and rejections "
                "write to the audit row with her named identity resolved through Entra. She "
                "reports to the District Manager and interfaces directly with the Food Safety "
                "Compliance Lead and refrigeration vendor. Her measure of success: time-to-"
                "decision compression, store-floor presence, and audit-row coverage.",
                accent="E59A3E")
    add_body(doc,
             "She is named, not generic. 'Store Operations Lead' is a title; the Store Ops Lead "
             "at Store #428 is a person with named identity, named device, named approval "
             "authority. The hitl-authority policy registered in LEDGER binds to her named "
             "identity (resolved through Entra group-lookup at runtime). The Adaptive Card fires "
             "to her device. Her approval is what unblocks the Act agent.")

    add_h(doc, "7.2 · Secondary personas", level=2)
    add_body(doc,
             "Five secondary personas each carry a distinct role in Kroger's store-operations "
             "model. Roles are deliberate — RC-E2E-03 does not invent roles, it maps to roles "
             "that already exist in Kroger's store-operations handbook.")
    add_table(doc,
              ["Role", "Function in the Kroger RC-E2E-03 scenario"],
              [
                  ["Pull-Team Lead",
                   "Task executor — receives the prioritized disposition list from the Act agent "
                   "and physically executes markdown / destroy / sell-through across the affected "
                   "fixture."],
                  ["Food Safety Compliance Lead",
                   "Escalation path for CRITICAL-severity events; final authority on destroy-all "
                   "and FSMA-evidence-gap events; consumes the LEDGER audit-row stream for "
                   "regulatory inquiry."],
                  ["Pricing Manager (banner level)",
                   "Owns the markdown-policy manifest and per-SKU elasticity tuning; reviews "
                   "weekly markdown-NPV trending; tunes tier thresholds quarterly."],
                  ["Inventory Control Specialist",
                   "Owns the inventory-position update path post-disposition; first-line response "
                   "for SCML.Inventory shape exceptions; reconciles physical-vs-system shrink."],
                  ["FDA Auditor / State Inspector",
                   "Consumes the LEDGER audit-row stream for FSMA 204 inspection; interfaces with "
                   "the Food Safety Compliance Lead on regulatory inquiry; signs off on quarterly "
                   "evidence packet shape."],
              ],
              col_widths=[2.0, 4.3])

    add_h(doc, "7.3 · Why persona depth matters", level=2)
    add_body(doc,
             "A seller who names only the primary persona says 'we built something for the Store "
             "Operations Lead.' A seller who names the full cast shows Kroger a complete "
             "operating model: who executes, who escalates, who approves above threshold, who "
             "handles the vendor touchpoint, who consumes the audit trail. That completeness is "
             "what turns a demo into an engagement.")
    add_body(doc,
             "Persona depth also matters for change management. When the engagement rolls out, "
             "Kroger's Field Enablement team has to train all six named roles — not just the "
             "Store Operations Lead. The six-role inventory informs training plan scope, "
             "org-readiness assessment, and the operational handoff checklist in the W2 → W3 "
             "gate review.")

    add_h(doc, "7.4 · Persona-to-plane mapping", level=2)
    add_table(doc,
              ["Persona", "Plane(s) touched", "Surface"],
              [
                  ["Store Operations Lead (primary)",  "Experience · Decision",
                   "Microsoft Teams Adaptive Card (mobile + desktop)"],
                  ["Pull-Team Lead",                    "Experience · Integration",
                   "Store-handheld task list · physical execution"],
                  ["Food Safety Compliance Lead",       "Experience · Decision · Ledger",
                   "Adaptive Card (CRITICAL events) · LEDGER query for regulatory inquiry"],
                  ["Pricing Manager",                   "Experience · Runtime",
                   "Markdown-policy manifest console · BI markdown-NPV dashboard"],
                  ["Inventory Control Specialist",      "Integration · Runtime",
                   "SCML.Inventory exception console · physical-vs-system reconciliation"],
                  ["FDA Auditor / State Inspector",     "Ledger",
                   "Quarterly FSMA evidence-packet export · audit-row query"],
                  ["District Manager (executive overlay)", "Ledger · Experience",
                   "Power BI shrink + decision-loop scorecard rolled up by store cluster"],
              ],
              col_widths=[2.1, 1.8, 2.4])

    pagebreak(doc)

    # ============ SECTION 8 · KPIs ============
    add_h(doc, "Section 8 · Chain Link 6 · KPIs", level=1)
    add_body(doc,
             "The KPI link is where the chain closes. APEX's signature is that every KPI "
             "attributes back to the specific decision audit row that produced it — not just to "
             "the aggregate. The distinction matters because it is what makes the KPI durable "
             "under audit and actionable under operating review.")

    add_h(doc, "8.1 · Per-event KPIs", level=2)
    add_body(doc,
             "A closed cold-chain event emits a closeout audit row carrying the following "
             "measured KPI fields. Each is bound to the specific trace-id of that event. A Power "
             "BI dashboard rolls them up across events; drill-through on any aggregated cell "
             "returns the individual audit rows.")
    add_table(doc,
              ["KPI", "Measurement", "Target (W2 Pilot)", "Target (W3 Scale)"],
              [
                  ["Time-to-decision",                "Sec from event to disposition",     "≤ 4 min", "≤ 90 sec"],
                  ["Manager attention reclaimed",     "Hours per shift per store",         "+3.5 hrs", "+5.2 hrs"],
                  ["Per-event shrink avoided",        "$ vs. destroy-all baseline",        "$640 / event", "$1,313 / event"],
                  ["Policy-conformance rate",         "% of dispositions matching policy", "≥ 88%", "≥ 95%"],
                  ["FSMA-204 audit-row completeness", "% with all 14 fields",              "≥ 96%", "≥ 99.5%"],
                  ["Auto-execute rate",               "% of events without HITL",          "55%", "75%"],
                  ["Vendor-ticket rate",              "% qualifying events with ticket",   "≥ 78%", "≥ 92%"],
                  ["Sell-through-at-risk rate",       "% sell-through later flagged",      "≤ 1.5%", "≤ 0.5%"],
              ],
              col_widths=[2.1, 2.0, 1.2, 1.2])

    add_h(doc, "8.2 · Derived KPIs for Kroger commercial cases", level=2)
    add_body(doc,
             "Beyond the per-event operational metrics, four derived KPIs anchor Kroger-specific "
             "commercial conversations. Each is computed from the per-event metrics aggregated "
             "over time and attributes back to specific events.")

    add_h(doc, "8.2.1 · Network-Wide Shrink Avoided", level=3)
    add_body(doc,
             "The headline financial metric. Computed as SUM(shrink_avoided_per_event) across all "
             "cold-chain closeout audit rows over a trailing 52 weeks. At 85 events/week × 2,800 "
             "stores × $1,313 avg avoided per event, the annualized number is approximately "
             "$200M. That is the single number CFO and COO will gate the engagement on. "
             "Sensitivity analysis on event-volume and per-event-shrink assumptions sits in "
             "RC-E2E-03-ROI-Case.html.")

    add_h(doc, "8.2.2 · Manager Capacity Returned", level=3)
    add_body(doc,
             "The operational metric that matters to the COO and the District Manager "
             "organization. Computed as 5.2 hrs per shift × 1.8 shifts per day × 2,800 stores = "
             "~26,200 hours per day of managerial attention reclaimed at W3 run-rate. Most of "
             "that converts to store-floor presence rather than dollars in the P&L; some "
             "monetizes through retention and reduced premium-pay.")

    add_h(doc, "8.2.3 · FSMA 204 Audit-Trail Completeness", level=3)
    add_body(doc,
             "The regulatory metric. Computed as the percentage of cold-chain events whose audit "
             "rows carry all 14 required fields, every version stamp, and a verifiable signature. "
             "Target ≥ 99.5% at W3. Matters because it is the direct answer to the FDA FSMA 204 "
             "auditor question and the audit-committee food-safety-controls question.")

    add_h(doc, "8.2.4 · Decision-Loop Compression", level=3)
    add_body(doc,
             "The metric that matters to store-operations leadership. Today's decision-loop time "
             "is 17 minutes median; W2 target is ≤ 4 min; W3 target is ≤ 90 sec. The value of "
             "compression isn't just operational — it's that decisions that today happen after "
             "store-opening now happen before. A dairy case that today drives a destroy-all "
             "because the manager couldn't get a Food Safety call back in time becomes a "
             "marked-down case with the right disposition.")

    add_h(doc, "8.3 · KPI attribution in the audit row", level=2)
    add_body(doc,
             "The closeout audit row — emitted once per resolved cold-chain event — carries the "
             "structure below. Every KPI attributes back to specific rows by trace-id.")
    add_code(doc,
             '{\n'
             '  "trace_id": "trc_2027-04-22_0608_kr_428_a3",\n'
             '  "scenario_id": "RC-SUPCHN-01-cold-chain-excursion",\n'
             '  "service_id": "RC-E2E-03",\n'
             '  "tenant_id": "kr",\n'
             '  "banner": "Kroger",\n'
             '  "event_kind": "rc_e2e_03_closeout",\n'
             '  "event_ts": "2027-04-22T07:00:18Z",\n'
             '  "manifest_version": "rc-cold-chain-orchestration-v1.2.0",\n'
             '  "policy_version": "fsma-204-cold-chain-corpus-v1.0.0",\n'
             '  "kpis": {\n'
             '    "time_to_decision_sec": 92,\n'
             '    "shrink_avoided_usd": 1287.40,\n'
             '    "policy_conformance": true,\n'
             '    "fsma_audit_row_complete": true,\n'
             '    "manager_time_reclaimed_min": 11.5\n'
             '  }\n'
             '}')

    pagebreak(doc)

    # ============ SECTION 9 · WAVES ============
    add_h(doc, "Section 9 · The Three Waves", level=1)
    add_body(doc,
             "The same chain — Scenario → Solution → Use Case → Service → Persona → KPI — "
             "projects across three delivery Waves. W1 Foundation is prerequisite work that "
             "enables the scenario. W2 Pilot is the scenario running live at proof-point scale. "
             "W3 Scale & Fuse is the scenario at enterprise run-rate, fused with adjacent "
             "scenarios.")

    add_h(doc, "9.1 · Wave 1 · Foundation — Single-banner anchor", level=2)
    add_body(doc,
             "W1 is the one-time investment per tenant. It produces the rails that make the "
             "scenario possible and make every subsequent scenario cheap. For Kroger, W1 scope is "
             "one banner cluster (Cincinnati-area Kroger stores), three of the five agents "
             "(Assess, Classify, Quantify), and the full APEX platform substrate via CAF. No "
             "automated downstream action in W1 — read-only intelligence only.")
    add_table(doc,
              ["W1 element", "What it looks like for Kroger"],
              [
                  ["Scope",            "One Cincinnati-area Kroger banner cluster · ~50 stores · dairy and deli only"],
                  ["Agents live",      "3 of 5 (Assess, Classify, Quantify)"],
                  ["Use cases",        "5.1 + 5.2 + 5.3 (read-only, no automated action)"],
                  ["Personas onboarded","Store Ops Lead + Food Safety Compliance Lead + Pricing Manager"],
                  ["Schemas",          "SCML.Fixture, SCML.Lot (W1 scope) · MERML.Price, MERML.Markdown"],
                  ["Audit row",        "Composed end-to-end on shadow data; replay against historical 4 quarters"],
                  ["Adaptive Card",    "Store Ops Lead's device tested; Entra binding live (notification only)"],
                  ["LEDGER plane",     "Stood up; pilot rows ingested; query path verified"],
                  ["Duration",         "5–8 months"],
                  ["Effort",           "~50 story points"],
                  ["Exit gate",        "Store Ops Lead + Food Safety Lead sign-off on shadow-run accuracy"],
              ],
              col_widths=[1.7, 4.6])

    add_h(doc, "9.2 · Wave 2 · Pilot — 250-store anchor cluster + FSMA evidence", level=2)
    add_body(doc,
             "W2 is where the scenario becomes real. A 250-store anchor cluster (typically the "
             "North America NA-Central or NA-South cluster) runs the live scenario end-to-end for "
             "90 consecutive operating days plus the first quarterly FSMA evidence packet "
             "delivered to FDA via Kroger's Food Safety organization. The full five-agent chain "
             "fires for every cold-chain event.")
    add_table(doc,
              ["W2 element", "What it looks like for Kroger"],
              [
                  ["Scope",            "250 stores · multi-banner · dairy + deli + produce + frozen"],
                  ["Agents live",      "All 5 (full orchestration with HITL gating)"],
                  ["Use cases",        "5.1 through 5.6 fully exercised"],
                  ["Personas onboarded","All 6 named roles operating live"],
                  ["Schemas",          "All MERML + SCML schemas live; FSMA 204 policy corpus v1.0"],
                  ["HITL volume",      "Tier-tuned to ~5–10% of events at steady state"],
                  ["Power BI",         "Shrink dashboard · markdown-NPV trending · FSMA evidence completeness"],
                  ["FSMA packet",      "First quarterly packet delivered to Kroger Food Safety + walked with FDA-relations team"],
                  ["Duration",         "8–12 months"],
                  ["Effort",           "~45 story points initial + ongoing enablement"],
                  ["Exit gate",        "Steering Committee approval to scale to enterprise"],
              ],
              col_widths=[1.7, 4.6])

    add_h(doc, "9.3 · Wave 3 · Scale & Fuse — 2,800+ stores · multi-banner · enterprise", level=2)
    add_body(doc,
             "W3 is where the scenario meets the Kroger enterprise and where it stops running "
             "alone. Scope expands from 250 stores to all 2,800+; banner coverage expands from one "
             "anchor banner to the full multi-banner footprint (Kroger, Ralphs, Fred Meyer, King "
             "Soopers, Smith's, Fry's, Harris Teeter, Mariano's, QFC, Pick 'n Save, Roundy's, "
             "and others); fixture-class coverage expands from dairy/deli/produce to the full "
             "refrigerated-and-frozen footprint. And — most importantly — RC-E2E-03 fuses with "
             "adjacent scenarios (Perishables Integrity Response, Regulator-Grade Evidence "
             "Service, Manager Time-Return Service) that share the architectural substrate.")
    add_table(doc,
              ["W3 element", "What it looks like for Kroger"],
              [
                  ["Scope",            "All 2,800+ stores · multi-banner · full refrigerated + frozen footprint"],
                  ["Agents live",      "All 5 with banner-aware policy manifests"],
                  ["Use cases",        "Full 24 sprint-level use cases plus fusion hooks"],
                  ["Schemas",          "All MERML + SCML shapes; cross-banner aggregation views"],
                  ["Cross-scenario fusion","RC-E2E-11 Perishables Integrity + RC-E2E-12 Regulator Evidence light up"],
                  ["Personas",         "Store-ops trained at scale; FDA-relations consuming evidence directly"],
                  ["KPI rollup",       "Enterprise-wide ~$200M annualized shrink-avoided at run-rate"],
                  ["Decision time",    "≤ 90 sec across all banners and fixture classes"],
                  ["Duration",         "12–18 months core; ongoing fusion"],
                  ["Effort",           "~70 story points initial + ongoing scenario-fusion"],
                  ["Exit gate",        "Steady-state operating review; expansion into adjacent scenarios"],
              ],
              col_widths=[1.7, 4.6])

    add_h(doc, "9.4 · Cross-wave load-bearing connections", level=2)
    add_body(doc,
             "Three structural dependencies bind the waves into one progression. They are not "
             "architectural coincidences — they are the load-bearing connections that make the "
             "wave model coherent.")
    add_callout(doc, "W1 → W2",
                "The HITL Adaptive Card prototyped in W1 (Store Ops Lead's device with Entra "
                "identity resolution, notification-only) powers the W2 HITL gate across all five "
                "agents in approval mode. What is prototyped in W1 is consumed unchanged in W2.",
                accent=TEAL, fill="E4F1EE")
    add_callout(doc, "W1 → W2",
                "The LEDGER plane established in W1 accumulates the W2 audit rows. The FSMA "
                "evidence packet format is designed in W1, exercised with shadow audit rows, and "
                "scales to thousands of rows in W2.")
    add_callout(doc, "W2 → W3",
                "The W2 pilot KPIs unlock W3 scaling. The Steering Committee decision to commit "
                "to all-banner enterprise rollout is specifically gated on W2 meeting "
                "decision-loop-time and FSMA-evidence-coverage targets.",
                accent=TEAL, fill="E4F1EE")

    pagebreak(doc)

    # ============ SECTION 10 · COMMERCIAL ENVELOPE ============
    add_h(doc, "Section 10 · Commercial Envelope", level=1)
    add_body(doc,
             "The Kroger engagement encompasses four commercial tiers corresponding to the "
             "Approach stage (Discovery and SOW build-out), Wave 1 Foundation, Wave 2 Pilot, and "
             "Wave 3 Scale & Fuse. Each carries a distinct cost-envelope range, duration, and "
             "deliverable signature.")
    add_table(doc,
              ["Tier", "Duration", "Cost envelope", "Effort", "Primary deliverable"],
              [
                  ["Approach",         "3–5 weeks",   "$200K–$500K",   "22 SP",   "Discovery report · banner scoping · SOW"],
                  ["Wave 1 Foundation","5–8 months",  "$1.2M–$3.5M",   "50 SP",   "Single-banner shadow-run · audit-row end-to-end"],
                  ["Wave 2 Pilot",     "8–12 months", "$5M–$12M",      "45 SP",   "250-store live · 90-day operating run · FSMA packet"],
                  ["Wave 3 Scale&Fuse","12–18 months","$18M–$45M",     "70+ SP",  "Enterprise 2,800+ stores · multi-banner · fusion"],
              ],
              col_widths=[1.4, 1.2, 1.3, 0.8, 2.5])

    add_h(doc, "10.1 · What the envelope buys", level=2)
    add_body(doc,
             "The commercial envelope is not a project-labor-hours calculation. It reflects the "
             "total value-delivery bundle: Deloitte-delivered solutioning and Practice expertise; "
             "Microsoft platform licensing and consumption (Fabric F-SKU, Foundry, Azure OpenAI, "
             "Teams, Entra, Power BI); CAF-substrate consumption and any Converge-tier costs that "
             "ride with it; ISV co-sell where applicable; client-side enablement and change "
             "management; FDA / state-inspection coordination.")

    add_h(doc, "10.2 · Envelope vs. durable value at Kroger scale", level=2)
    add_body(doc,
             "A 2,800-store W3 enterprise footprint produces approximately $200M annualized "
             "shrink-avoided value at run-rate, plus operational capacity returned to managers "
             "(~26K hours/day across the enterprise), plus regulatory-audit cost reduction (FSMA "
             "evidence collation hours), plus the strategic optionality that comes with auditable "
             "perishables operations. The $18M–$45M W3 envelope spans a rapid-payback range "
             "(under 6 months at run-rate at the low end) to a strategic-investment range (12–18 "
             "months at the high end, reflecting greater scenario-fusion scope in the upper tier).")

    add_h(doc, "10.3 · Independence posture in Kroger commercial conversations", level=2)
    add_callout(doc, "Independence",
                "The seller and delivery team use Deloitte-appropriate language throughout: "
                "'Deloitte's Microsoft practice,' 'DMTSP,' 'Microsoft platform capabilities,' "
                "'Microsoft-native deployment.' The terms 'partner,' 'alliance,' and 'strategic "
                "alliance' do not appear in RC-E2E-03 commercial materials. The engagement is "
                "structured so that manifests, policies, and LEDGER are Kroger artifacts; "
                "Deloitte's role is advisory and implementation enablement. Before any contract "
                "drafting, confirm Kroger's external-audit firm relationship and route the SOW "
                "through Deloitte's Independence office.")

    pagebreak(doc)

    # ============ SECTION 11 · REFERENCE W1 IMPLEMENTATION ============
    add_h(doc, "Section 11 · Reference Wave 1 Implementation", level=1)
    add_body(doc,
             "The RC-E2E-03 Wave 1 prototype is the concrete proof point that anchors this "
             "walkthrough. It will be a runnable web application that executes the three "
             "foundation use cases against synthetic Kroger cold-chain data with full audit-row "
             "capture.")

    add_h(doc, "11.1 · What the W1 prototype demonstrates", level=2)
    add_table(doc,
              ["Capability", "What it demonstrates"],
              [
                  ["Synthetic Kroger fixture telemetry",
                   "50K synthetic cold-chain events across 12 fixture classes with realistic excursion patterns"],
                  ["Assess agent (live)",
                   "Excursion-profile computation; severity classification; SCML.Fixture state ingestion"],
                  ["Classify agent (live)",
                   "Per-SKU disposition with confidence; FSMA 204 policy corpus retrieval"],
                  ["Quantify agent (live)",
                   "Three disposition scenarios with $ impact (sell-through / markdown / destroy)"],
                  ["Adaptive Card flow",
                   "Mock Store Ops Lead notification with Entra identity resolution (read-only mode)"],
                  ["LEDGER query surface",
                   "Drill-through from Power BI dashboard tile to specific trace_id"],
                  ["Audit-row export",
                   "Quarterly FSMA evidence packet ZIP with manifests, policies, and audit-row extract"],
                  ["FSMA 204 policy versioning",
                   "Policy manifest with version stamps; PR-merge discipline"],
                  ["Manifest versioning",
                   "Git-backed manifests for agents, policy, schema; PR-merge discipline"],
                  ["Shadow-run replay",
                   "Replay against 4 historical quarters of synthetic excursion data"],
                  ["Demo console",
                   "Three canonical demo paths · 3–5 minutes each · narrated for Kroger stakeholders"],
                  ["Self-host install",
                   "Local Docker + Ollama (offline-capable for sensitive demo environments)"],
              ],
              col_widths=[2.2, 4.1])

    add_h(doc, "11.2 · Three canonical demo paths", level=2)
    add_body(doc, "Three demo paths planned for Kroger stakeholder reviews, each ~3–5 minutes:")
    add_body(doc,
             "Demo Path A · The excursion loop. Pick a synthetic fixture from the dairy cluster. "
             "Watch Assess ingest the 52-minute temperature curve; watch Classify produce the "
             "per-SKU disposition table with FSMA 204 policy citations; watch Quantify produce "
             "the three-scenario $ impact analysis; watch the Adaptive Card fire to the Store Ops "
             "Lead's device with the proposed disposition; click through Approve / Modify / "
             "Reject; watch the audit row write.")
    add_body(doc,
             "Demo Path B · The shrink-NPV trending. Open the Power BI shrink dashboard. Show "
             "the per-banner shrink-avoided trending over 12 weeks of shadow-run data. Drill "
             "through into a single high-impact cell — get the contributing audit rows. Drill "
             "into one row — see the full reasoning trace, the version stamps, the HITL "
             "approver's identity.")
    add_body(doc,
             "Demo Path C · The FSMA evidence trail. Pick a synthetic excursion event. Open the "
             "audit-row stream for that trace_id. Show the full chain of custody from temperature "
             "breach through classification, through quantification, through HITL approval, "
             "through downstream action, through FSMA log entry. Show the manifest version and "
             "policy version stamps. Show the FSMA 204 policy citations captured at decision time.")

    add_h(doc, "11.3 · Prototype technology stack", level=2)
    add_table(doc,
              ["Layer", "Stack"],
              [
                  ["Frontend",            "React + Vite · Tailwind · Power BI Embedded"],
                  ["Backend",             "Node.js (Fastify) · pg-pool"],
                  ["Persistence",         "PostgreSQL (LEDGER + state)"],
                  ["Identity",            "Entra ID (mock binding for prototype)"],
                  ["Inference",           "Local Ollama (Llama-class model) for offline demos · Azure OpenAI for online demos"],
                  ["Orchestration",       "Lightweight DAG runner + manifest loader (CAF-shaped contract)"],
                  ["Observability",       "OpenTelemetry traces + LangFuse-equivalent local span store"],
                  ["Schema validation",   "Zod (MERML.* + SCML.* shapes)"],
                  ["Demo packaging",      "Docker Compose · 'docker compose up' single-command boot"],
                  ["Source control",      "Git repo with manifest history; PR-template for policy edits"],
                  ["Synthetic data",      "50K synthetic excursion events · 4 historical quarters · 12 fixture classes"],
              ],
              col_widths=[1.6, 4.7])

    add_h(doc, "11.4 · Mapping prototype → APEX architecture", level=2)
    add_body(doc,
             "The prototype is deliberately a subset of the full APEX architecture — enough to "
             "make the design concrete for Kroger stakeholders, not so much that it requires a "
             "full Azure/Fabric/Foundry deployment to demo:")
    add_table(doc,
              ["Prototype element", "APEX architecture element", "Kroger production surface"],
              [
                  ["Local Postgres",            "LEDGER plane (audit row store)",
                   "Azure Fabric Lakehouse + Postgres index in production"],
                  ["Mock POS / vendor-portal APIs", "Integration plane",
                   "Kroger's actual POS markdown stack and refrigeration-vendor portals"],
                  ["Mock Adaptive Card device", "Experience plane",
                   "Store Ops Lead's actual Microsoft Teams environment"],
                  ["Local Ollama model",        "Decision plane (LLM)",
                   "Azure OpenAI via AEF LLM Gateway (CAF substrate)"],
                  ["Lightweight DAG runner",    "Runtime plane",
                   "CAF Dynamic Orchestrator + A2A Swarm in production"],
                  ["Synthetic excursion data",  "MERML + SCML shapes",
                   "Kroger's actual refrigeration-telemetry stream + lot-provenance feed"],
                  ["Mock FSMA policy corpus",   "Policy/Decision-evidence manifest",
                   "Kroger's actual Food Safety policy corpus + FDA inspection workflow"],
              ],
              col_widths=[1.9, 2.0, 2.4])

    add_h(doc, "11.5 · Prototype limitations and next steps", level=2)
    add_body(doc, "The prototype is deliberately scoped for Wave 1 demonstration. It does not:")
    add_bullets(doc, [
        "Connect to real Kroger refrigeration-telemetry stream — synthetic data only; real integration is W2 scope",
        "Run on Azure — everything is local (Docker + Ollama on a laptop); Azure migration is W2",
        "Activate the Approve or Act agents — those land in W2 (W1 is read-only intelligence)",
        "Implement real Entra authentication — hardcoded Store Ops Lead persona for demo; Entra integration is W2",
        "Produce a true FDA-ready FSMA evidence packet — audit timeline view demonstrates the substrate; packet assembly is W2",
        "Handle multi-banner orchestration — single-banner scope only",
    ])
    add_body(doc,
             "These are not gaps in the architecture — they are intentional scope boundaries. "
             "The W1 prototype demonstrates the agentic flow end-to-end with enough fidelity to "
             "make the Wave 1 exit review credible. Wave 2 closes each gap with the appropriate "
             "Kroger production system.")

    pagebreak(doc)

    # ============ APPENDIX A · GLOSSARY ============
    add_h(doc, "Appendix A · Glossary", level=1)
    add_body(doc, "Selected terms used in this document with APEX- or grocery-specific meaning.")
    add_table(doc,
              ["Term", "Definition"],
              [
                  ["A2A Swarm",
                   "Cross-app agent collaboration primitive in CAF; the bridge by which APEX scenarios dispatch into CAF runtime."],
                  ["Adaptive Card",
                   "Microsoft Teams rich-card format used by RC-E2E-03 for HITL approval surfaces."],
                  ["Agent",
                   "A bounded unit of intent + tool-use that emits typed output and writes an audit row per decision."],
                  ["APEX",
                   "Agentic Platform for Enterprise eXecution — Deloitte's reference architecture for industry scenarios on Microsoft-native deployment."],
                  ["Audit row",
                   "14-field structured record emitted by every agent decision; primary substrate of FSMA evidence and KPI attribution."],
                  ["Banner",
                   "A retail brand identity within Kroger (Kroger, Ralphs, Fred Meyer, etc.). RC-E2E-03 manifests are banner-aware."],
                  ["CAF",
                   "Converge Agentic Framework — production platform substrate that RC-E2E-03 consumes via A2A."],
                  ["Cold-chain excursion",
                   "An event in which a refrigerated fixture breaches its FSMA-governed temperature setpoint for sufficient duration to require disposition."],
                  ["Disposition",
                   "The outcome decision for an excursion-affected unit: sell-through, markdown, or destroy."],
                  ["Entra",
                   "Microsoft identity platform; resolves named-persona HITL bindings at runtime."],
                  ["Fabric",
                   "Microsoft unified analytics platform; hosts Bronze/Silver/Gold data assets used by RC-E2E-03 in production."],
                  ["Foundry",
                   "Microsoft Azure AI Foundry; hosts model endpoints, tool registrations, and prompt management."],
                  ["FSMA 204",
                   "FDA Food Safety Modernization Act §204(d) final rule — lot-traceability requirement for high-risk foods, effective 2026."],
                  ["HITL",
                   "Human-in-the-loop gate; first-class orchestration primitive declared in manifest."],
                  ["LEDGER",
                   "The APEX audit-row persistence + retrieval plane; substrate of FSMA evidence and KPI attribution."],
                  ["Manifest",
                   "Git-versioned declarative artifact describing agents, policies, schemas, and tenant subscriptions."],
                  ["Markdown",
                   "Disposition path: lower price + accelerated sell-through. Quantified by Quantify agent against MERML.Markdown."],
                  ["MERML",
                   "APEX canonical schema family for Merchandising objects (Price, Markdown, Elasticity, etc.)."],
                  ["Persona",
                   "A named human role with declared authority, plane bindings, and surfaces. Not generic; named identity."],
                  ["Pull Team",
                   "Store-floor team that executes physical disposition (markdown rack, destroy stack, sell-through restock)."],
                  ["RC-E2E-03",
                   "APEX Service Catalog ID for Assortment & Pricing. The highest-value RC service for grocery clients at scale."],
                  ["RC-E2E-09",
                   "APEX Service Catalog ID for Product Tracking (FSMA 204 / lot-provenance). Secondary service in this engagement."],
                  ["SCML",
                   "APEX canonical schema family for Supply Chain & Lot objects (Lot, Inventory, Fixture, Receipt, etc.)."],
                  ["Shrink",
                   "Inventory loss from spoilage, theft, damage, or markdown. RC-E2E-03's headline financial KPI is shrink-avoided."],
                  ["Trace-id",
                   "Immutable correlation identifier that binds every audit row from one decision event."],
                  ["Wave",
                   "Delivery phase: W1 Foundation, W2 Pilot, W3 Scale & Fuse. Anchors commercial envelope and exit gates."],
                  ["84.51°",
                   "Kroger's data-and-analytics subsidiary; Kroger's existing pricing and customer-data muscle. RC-E2E-03 complements rather than replaces."],
              ],
              col_widths=[1.7, 4.6])

    pagebreak(doc)

    # ============ APPENDIX B · ARTIFACT CROSS-REFERENCE ============
    add_h(doc, "Appendix B · Artifact Cross-Reference", level=1)
    add_body(doc,
             "This document is one of several APEX RC-E2E-03 artifacts covering Kroger's "
             "assortment-and-pricing automation program. The cross-reference below shows which "
             "artifact serves which audience.")
    add_table(doc,
              ["Artifact", "Primary audience", "Purpose"],
              [
                  ["APEX-RC-E2E-03-Walkthrough.docx (this)",
                   "CFO · COO · Chief Merchant · Food Safety Lead · CIO · External Audit",
                   "Executive walkthrough of the value-delivery chain"],
                  ["RC-E2E-03-Sprint-Plan.xlsx",
                   "Engagement Manager · Lead Engineer · Scrum lead",
                   "24 sprint use cases with story points, dependencies, exit criteria"],
                  ["RC-E2E-03-Resource-Plan.xlsx",
                   "DMTSP PMO · Engagement Manager",
                   "Resource loading by phase, role, sprint"],
                  ["RC-E2E-03-ROI-Case.html",
                   "CFO · COO · Steering Committee",
                   "Shrink-avoided uplift case with sensitivity analysis"],
                  ["RC-E2E-03 service spec (APEX Service Catalog)",
                   "APEX Practice RC lead · Engagement Architect",
                   "Service composition contract: schemas, agents, SLOs, persona"],
                  ["APEX-Stacked-Architecture-Narrated.html",
                   "All audiences",
                   "Practice reference — six sub-Practices, planes, archetypes"],
                  ["APEX-CAF-Integration-Architecture.docx",
                   "AI & Eng leadership · CAF platform lead · APEX RC Practice lead",
                   "Substrate integration: how RC-E2E-03 consumes CAF via A2A"],
                  ["MERML + SCML schema spec",
                   "Engagement Architect · Data engineer",
                   "Price, Markdown, Elasticity, Lot, Fixture canonical shapes"],
                  ["APEX_Design.md",
                   "Build team · Practice Lead",
                   "Reference architecture spec for agents, LEDGER, audit row"],
                  ["Professional-APEX-Sellers-Guide.html",
                   "Seller · Account Team",
                   "Runtime plane, audit-row extension, LEDGER feedback loop"],
                  ["Kroger 01_account materials",
                   "Account Team · Pursuit Lead",
                   "Account context, Kroger organizational priors, prior engagements"],
                  ["W1 prototype repo (planned)",
                   "Lead Engineer · Kroger technical reviewers",
                   "Runnable Wave 1 demo · synthetic data · three demo paths"],
                  ["FSMA 204 evidence shape addendum (planned)",
                   "Kroger Food Safety Office · FDA-relations · External Audit",
                   "Decision-evidence shape and quarterly packet format for FDA inspection"],
              ],
              col_widths=[2.3, 2.0, 2.0])

    add_body(doc,
             "Independence posture preserved across every APEX RC-E2E-03 artifact: 'Deloitte's "
             "Microsoft practice,' 'DMTSP,' 'Microsoft platform capabilities,' 'Microsoft-native "
             "deployment.' The terms 'partner,' 'alliance,' and 'strategic alliance' do not "
             "appear in any artifact.",
             italic=True, color=MUTE, size=9)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— End of walkthrough —")
    set_run(r, italic=True, size=10, color=MUTE)

    # save
    target = OUT
    try:
        doc.save(str(target))
    except PermissionError:
        n = 2
        while True:
            alt = target.with_name(target.stem + f"-v{n}" + target.suffix)
            if not alt.exists():
                target = alt
                break
            n += 1
        doc.save(str(target))
        print(f"\n[notice] {OUT.name} was locked — saved as {target.name} instead.")
    print(f"Wrote {target}")


if __name__ == "__main__":
    build()
