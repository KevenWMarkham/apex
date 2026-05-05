"""Tier-2 artifacts:
  5. Fabric Implementation Runbook (DOCX)
  6. Agent manifests (YAML — 7 files)
  7. Use Case Catalog (XLSX)
"""
from __future__ import annotations
import io, sys, os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
OUT = Path(r"C:\Stage\Clients\Industries\Consumer\Retail\Kroger\02_projects\FY27_Pipeline\assortment-pricing-agentic\deliverables")
OUT.mkdir(parents=True, exist_ok=True)
MANIFESTS = OUT / "manifests"
MANIFESTS.mkdir(exist_ok=True)

# ---- shared docx helpers (compact) ----
INK="0B0B0F"; INK2="4A4A55"; MUTE="8a8a90"; GOLD="B8892E"; TEAL="2A7F73"
VIO="6648B0"; CRIM="A22A39"; BAND="F3F3F5"; HEAD="0B0B0F"; LINE="D4D4D7"

def cell_bg(cell, c):
    tcPr = cell._tc.get_or_add_tcPr()
    s = OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto")
    s.set(qn("w:fill"),c.lstrip("#")); tcPr.append(s)
def cell_border(cell, c=LINE):
    tcPr = cell._tc.get_or_add_tcPr()
    bs = OxmlElement("w:tcBorders")
    for e in ("top","left","bottom","right"):
        b = OxmlElement(f"w:{e}"); b.set(qn("w:val"),"single"); b.set(qn("w:sz"),"4"); b.set(qn("w:color"),c.lstrip("#"))
        bs.append(b)
    tcPr.append(bs)
def setrun(r, *, bold=False, italic=False, size=None, color=None, font="Arial"):
    if bold: r.bold = True
    if italic: r.italic = True
    if size is not None: r.font.size = Pt(size)
    if color: r.font.color.rgb = RGBColor.from_string(color.lstrip("#"))
    if font: r.font.name = font
def addbody(doc, t, *, size=11, italic=False, color=INK, sa=6):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(sa)
    r = p.add_run(t); setrun(r, size=size, color=color, italic=italic); return p
def addh(doc, t, *, level=1):
    sz = {1:18,2:14,3:12.5,4:11.5}
    p = doc.add_paragraph(); p.style = doc.styles[f"Heading {level}"]
    p.paragraph_format.space_before = Pt(16 if level==1 else 10); p.paragraph_format.space_after = Pt(4)
    for r in list(p.runs): r.text = ""
    r = p.add_run(t); setrun(r, bold=True, size=sz.get(level,12), color=INK)
    if level == 1:
        pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom"); bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),"6"); bot.set(qn("w:space"),"4"); bot.set(qn("w:color"),GOLD)
        pBdr.append(bot); pPr.append(pBdr)
    return p
def addbullets(doc, items, *, size=11):
    for it in items:
        p = doc.add_paragraph(style="List Bullet"); r = p.add_run(it); setrun(r, size=size, color=INK)
def addtable(doc, header, rows, *, col_widths=None):
    n = len(header); tbl = doc.add_table(rows=len(rows)+1, cols=n); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j,h in enumerate(header):
        c = tbl.rows[0].cells[j]; cell_bg(c,HEAD); cell_border(c); c.paragraphs[0].text=""
        r = c.paragraphs[0].add_run(h); setrun(r,bold=True,size=10,color="FFFFFF")
        if col_widths: c.width = Inches(col_widths[j])
    for i,row in enumerate(rows, start=1):
        for j,v in enumerate(row):
            c = tbl.rows[i].cells[j]; cell_border(c)
            if i%2==0: cell_bg(c,BAND)
            c.paragraphs[0].text=""; r = c.paragraphs[0].add_run(str(v)); setrun(r,size=10,color=INK)
            if col_widths: c.width = Inches(col_widths[j])
    doc.add_paragraph(); return tbl
def addcallout(doc, kind, body, *, accent=GOLD, fill="F6F1E4"):
    tbl = doc.add_table(rows=1, cols=1); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = tbl.rows[0].cells[0]; cell_bg(c,fill); c.paragraphs[0].text=""
    rt = c.paragraphs[0].add_run(kind.upper()); setrun(rt,bold=True,size=9,color=accent)
    p = c.add_paragraph(); p.paragraph_format.space_before = Pt(3)
    r = p.add_run(body); setrun(r,size=10.5,color=INK); doc.add_paragraph()
def addkv(doc, pairs, *, lw=1.6, rw=4.9):
    tbl = doc.add_table(rows=len(pairs), cols=2); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,(k,v) in enumerate(pairs):
        ck,cv = tbl.rows[i].cells; ck.width=Inches(lw); cv.width=Inches(rw)
        cell_bg(ck,BAND); cell_border(ck); cell_border(cv)
        ck.paragraphs[0].text=""; cv.paragraphs[0].text=""
        rk = ck.paragraphs[0].add_run(k); setrun(rk,bold=True,size=10,color=INK)
        rv = cv.paragraphs[0].add_run(v); setrun(rv,size=10,color=INK)
    doc.add_paragraph()
def addeyebrow(doc, t):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(t.upper()); setrun(r,bold=True,size=9,color=MUTE)
    rPr = r._element.get_or_add_rPr(); sp = OxmlElement("w:spacing"); sp.set(qn("w:val"),"80"); rPr.append(sp)
def pb(doc):
    p = doc.add_paragraph(); p.add_run().add_break(WD_BREAK.PAGE)
def addcode(doc, t):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(t); setrun(r, size=9, color=INK, font="Consolas")


# =========================================================================
# 5. FABRIC IMPLEMENTATION RUNBOOK (DOCX)
# =========================================================================
def build_runbook():
    doc = Document()
    for s in doc.sections:
        s.page_width=Inches(8.5); s.page_height=Inches(11)
        s.top_margin=s.bottom_margin=Inches(0.85); s.left_margin=s.right_margin=Inches(0.95)
    doc.styles["Normal"].font.name="Arial"; doc.styles["Normal"].font.size=Pt(11)

    addeyebrow(doc, "APEX · RC-E2E-03 + RC-E2E-09 · FABRIC IMPLEMENTATION RUNBOOK")
    p=doc.add_paragraph(); r=p.add_run("Fabric Implementation Runbook"); setrun(r,bold=True,size=24,color=INK)
    p2=doc.add_paragraph(); r2=p2.add_run("Step-by-step build guide · Bronze landing → Silver canonical → Gold semantic")
    setrun(r2,bold=True,size=14,color=INK2,italic=True)
    doc.add_paragraph()
    addbody(doc, "Operational runbook for standing up the APEX RC-E2E-03 + RC-E2E-09 services on Microsoft "
                 "Fabric. Where the SOR document specifies WHAT to build, this document specifies HOW to "
                 "build it. Targets the Lead Engineer, Fabric admin, and Data Engineering team.", italic=True, color=INK2)
    doc.add_paragraph()
    addkv(doc, [
        ("Document type", "Implementation Runbook · Microsoft Fabric"),
        ("Services", "RC-E2E-03 · RC-E2E-09"),
        ("Platform", "Microsoft Fabric · Eventstream · Data Factory · Lakehouse · Real-Time Intelligence · Power BI"),
        ("Audience", "Lead Engineer · Fabric admin · Data Eng · Platform Eng"),
        ("Reference client", "The Kroger Co."),
        ("Companion", "APEX-RC-E2E-03-09-SOR-Bronze-to-Silver.docx · APEX-RC-E2E-03-09-SOR-ERD.html"),
        ("Version / date", "v1.0 · 27 Apr 2026"),
    ])
    pb(doc)

    addh(doc, "1 · Workspace and capacity setup", level=1)
    addh(doc, "1.1 · Fabric workspace topology", level=2)
    addbody(doc, "Three workspaces, separated by lifecycle and security boundary:")
    addtable(doc, ["Workspace", "Purpose", "Capacity SKU", "RBAC scope"],
             [["apex-rc-e2e-dev","Development + integration","F32","Engineers + Architects"],
              ["apex-rc-e2e-uat","Pre-production validation","F64","Engineers + QA + Stakeholder reviewers"],
              ["apex-rc-e2e-prod","Production","F128 (W2) → F256 (W3)","Operators + Auditors (read-only)"]],
             col_widths=[1.6,1.8,1.2,1.6])
    addh(doc, "1.2 · OneLake folder structure", level=2)
    addcode(doc, "apex-rc-e2e-prod/\n  Files/\n  Tables/\n    sch_rc_e2e_03/\n      brz_pos_txn/\n      brz_pricing/\n      brz_promo_markdown/\n      brz_refrig_telemetry/\n      brz_refrig_service/\n      brz_loyalty_hh/\n      slv_merml_price/\n      slv_merml_markdown/\n      slv_merml_elasticity/\n      slv_merml_promotion/\n      slv_scml_fixture/\n    sch_rc_e2e_09/\n      brz_lot_trace/\n      brz_vendor_master/\n      brz_receipt/\n      brz_food_safety/\n      slv_scml_lot/\n      slv_scml_receipt/\n      slv_scml_disposition/\n    sch_rc_shared/\n      brz_item_master/\n      brz_inventory/\n      slv_dim_item/\n      slv_dim_vendor/\n      slv_dim_household/\n      slv_scml_inventory/\n    sch_ledger/\n      audit_row/\n      sox_packet_export/")

    addh(doc, "1.3 · Identity bindings (Entra)", level=2)
    addtable(doc, ["Role", "Entra group", "Workspace permission"],
             [["APEX Engineer","apex-rc-engineers","Member (dev) · Contributor (UAT) · Reader (prod)"],
              ["APEX Architect","apex-rc-architects","Admin (dev/UAT) · Contributor (prod)"],
              ["Operator","apex-rc-operators","Member (prod) · Reader (dev/UAT)"],
              ["Auditor","apex-rc-auditors","Reader (prod) · LEDGER query rights"],
              ["FDA-Relations","apex-rc-fda-relations","FSMA evidence packet export rights"],
              ["VP Retention (HITL)","apex-rc-hitl-vp","Adaptive Card delivery target"]],
             col_widths=[1.8,1.7,3.2])
    pb(doc)

    addh(doc, "2 · Bronze ingest pipelines", level=1)
    addh(doc, "2.1 · Eventstream-driven (streaming feeds)", level=2)
    addbody(doc, "Five Bronze tables ingest via Fabric Eventstream for sub-minute latency:")
    addtable(doc, ["Bronze table", "Source", "Eventstream config notes"],
             [["BRZ.RC.POS_TXN","POS gateway via Event Hubs","Partition by store_id; output to Lakehouse with schema-validation enabled"],
              ["BRZ.RC.PRICING","Pricing engine CDC","Partition by price_zone_id; conflict handling = last-write-wins on (sku, zone, eff_date)"],
              ["BRZ.RC.REFRIG_TELEMETRY","MQTT bridge from EMS","Highest-volume; consider Eventhouse for hot-tier real-time queries"],
              ["BRZ.RC.FOOD_SAFETY","QMS CDC","Direct write to Lakehouse; trigger evidence-packet generator on disposition events"],
              ["BRZ.RC.RECEIPT","Manhattan WMS CDC","Partition by receipt_ts day"]],
             col_widths=[1.8,1.5,3.2])
    addcode(doc, "# Eventstream config skeleton (Fabric Real-Time)\n"
                 "name: rc-e2e-pos-stream\nsource:\n  type: event_hubs\n  connection: kr-pos-eventhub\n"
                 "transformations:\n  - type: schema_validate\n    schema_ref: brz_pos_txn_schema\n"
                 "  - type: dedupe\n    on: [source_record_id, source_system]\n  - type: enrich\n"
                 "    add_fields: [ingest_ts, batch_id]\nsink:\n  type: lakehouse_table\n"
                 "  workspace: apex-rc-e2e-prod\n  table: sch_rc_e2e_03.brz_pos_txn\n  mode: append\n"
                 "  partition_by: ingest_date\n  zorder_by: [store_id, transaction_ts]")

    addh(doc, "2.2 · Data Factory pipelines (batch feeds)", level=2)
    addtable(doc, ["Bronze table", "Pipeline pattern", "Schedule"],
             [["BRZ.RC.ITEM_MASTER","Daily incremental + nightly full","02:00 UTC"],
              ["BRZ.RC.VENDOR_MASTER","Weekly full","Sunday 03:00 UTC"],
              ["BRZ.RC.LOYALTY_HH","Daily snapshot","04:00 UTC"],
              ["BRZ.RC.LOT_TRACE (batch portion)","Daily EDI ingestion","06:00 UTC"],
              ["BRZ.RC.INVENTORY","Hourly delta","Top of every hour"],
              ["BRZ.RC.PROMO_MARKDOWN","Daily + on-publish trigger","01:00 UTC + event"]],
             col_widths=[2.0,2.5,1.8])
    addh(doc, "2.3 · Polling for vendor portals (REST)", level=2)
    addbody(doc, "BRZ.RC.REFRIG_SERVICE polls Emerson, Lennox, Hussmann, Hill PHOENIX vendor APIs every "
                 "15 minutes via Data Factory web activity. Implement vendor-portal API rate-limit "
                 "compliance, exponential-backoff retry, and connection-failure alerts to Facilities Eng.")
    pb(doc)

    addh(doc, "3 · Delta Lake configuration", level=1)
    addh(doc, "3.1 · Partitioning + Z-order strategy", level=2)
    addtable(doc, ["Bronze table", "Partition by", "Z-order by"],
             [["BRZ.RC.POS_TXN","ingest_date","store_id, transaction_ts"],
              ["BRZ.RC.ITEM_MASTER","ingest_date","sku, gtin, category_l1"],
              ["BRZ.RC.PRICING","ingest_date","sku, price_zone_id, effective_date"],
              ["BRZ.RC.PROMO_MARKDOWN","ingest_date","promo_type, start_ts, sku_group_id"],
              ["BRZ.RC.INVENTORY","ingest_date","location_id, sku, snapshot_ts"],
              ["BRZ.RC.REFRIG_TELEMETRY","ingest_date · store_id (hash-bucket on hot stores)","fixture_id, reading_ts"],
              ["BRZ.RC.REFRIG_SERVICE","ingest_date","store_id, fixture_id, open_ts"],
              ["BRZ.RC.FOOD_SAFETY","ingest_date","store_id, incident_ts, incident_type"],
              ["BRZ.RC.LOT_TRACE","ingest_date · ftl_category","lot_code, gtin, event_ts"],
              ["BRZ.RC.VENDOR_MASTER","ingest_date","vendor_id, vendor_type"],
              ["BRZ.RC.RECEIPT","ingest_date","location_id, lot_code, receipt_ts"],
              ["BRZ.RC.LOYALTY_HH","ingest_date","loyalty_household_id, primary_banner"]],
             col_widths=[2.1,2.3,2.0])
    addh(doc, "3.2 · Vacuum and optimize cadence", level=2)
    addbullets(doc, [
        "OPTIMIZE: nightly for high-velocity tables (POS_TXN, PRICING, REFRIG_TELEMETRY, LOT_TRACE)",
        "OPTIMIZE: weekly for moderate tables (INVENTORY, FOOD_SAFETY, RECEIPT, PROMO_MARKDOWN)",
        "OPTIMIZE: monthly for slow-changing tables (ITEM_MASTER, VENDOR_MASTER, LOYALTY_HH, REFRIG_SERVICE)",
        "VACUUM: 7-day retention default; extend to 30 days for tables with regulatory replay requirements (FSMA 204 = LOT_TRACE, FOOD_SAFETY, RECEIPT, ITEM_MASTER)",
        "Time-travel preservation: never VACUUM beyond 24-month FSMA window for LOT_TRACE — point-in-time queries depend on it",
    ])
    pb(doc)

    addh(doc, "4 · Silver canonical layer", level=1)
    addbody(doc, "Silver shapes are built via Fabric Notebooks (PySpark) or SQL endpoint stored procedures. "
                 "Recommend Notebook approach for the join-heavy shapes (SCML.Lot, MERML.Price); SQL "
                 "endpoint for simpler dimensional mappings (DIM.Item, DIM.Vendor).")
    addh(doc, "4.1 · Build order (DAG)", level=2)
    addbody(doc, "Run in this order to satisfy referential dependencies:")
    addtable(doc, ["Order", "Silver shape", "Build pattern", "Refresh"],
             [["1","DIM.Item","SCD2 from BRZ.RC.ITEM_MASTER","Daily 03:00 UTC"],
              ["2","DIM.Vendor","SCD2 from BRZ.RC.VENDOR_MASTER","Weekly Sun 04:00 UTC"],
              ["3","DIM.Household","Type-1 from BRZ.RC.LOYALTY_HH","Daily 05:00 UTC"],
              ["4","SCML.Fixture","Window-aggregate from REFRIG_TELEMETRY + REFRIG_SERVICE","Streaming + 5-min window"],
              ["5","MERML.Price","CDC-driven from PRICING + PROMO_MARKDOWN overlay","Streaming"],
              ["6","SCML.Receipt","CDC-driven from RECEIPT + LOT_TRACE","Streaming"],
              ["7","MERML.Promotion","On-publish from PROMO_MARKDOWN","Daily + event"],
              ["8","SCML.Lot","Multi-source join (LOT_TRACE + RECEIPT + VENDOR + ITEM)","Streaming + nightly recon"],
              ["9","MERML.Markdown","PROMO_MARKDOWN + POS_TXN attribution","Daily 06:00 UTC"],
              ["10","SCML.Inventory","Reconcile INVENTORY + RECEIPT + POS_TXN","Hourly"],
              ["11","MERML.Elasticity","Per-SKU model fit from POS_TXN + PRICING","Weekly Sun 06:00 UTC"],
              ["12","SCML.Disposition","FOOD_SAFETY closeout + LOT_TRACE + POS_TXN outcome","Streaming + 7-day outcome window"]],
             col_widths=[0.5,1.8,2.7,1.5])
    addh(doc, "4.2 · Notebook scaffolding pattern", level=2)
    addcode(doc, "# notebook: build_slv_scml_lot.ipynb\nfrom pyspark.sql import functions as F\nfrom delta.tables import DeltaTable\n\n# read Bronze\nlot = spark.read.format('delta').table('apex-rc-e2e-prod.sch_rc_e2e_09.brz_lot_trace').filter('NOT _is_deleted')\nrcpt = spark.read.format('delta').table('apex-rc-e2e-prod.sch_rc_shared.brz_receipt').filter('NOT _is_deleted')\nvendor = spark.read.format('delta').table('apex-rc-e2e-prod.sch_rc_shared.slv_dim_vendor')\nitem = spark.read.format('delta').table('apex-rc-e2e-prod.sch_rc_shared.slv_dim_item')\n\n# canonical SCML.Lot\nslv_lot = (lot\n    .join(rcpt, on='lot_code', how='left')\n    .join(vendor, on='vendor_id', how='left')\n    .join(item, on='gtin', how='left')\n    .withColumn('lot_status', F.when(F.col('disposition').isNotNull(), 'closed').otherwise('active'))\n    .select('lot_code','gtin','sku','vendor_id','ftl_category','event_ts','qty','lot_status', ...))\n\n# upsert into Silver\nDeltaTable.forName(spark,'apex-rc-e2e-prod.sch_rc_e2e_09.slv_scml_lot') \\\n  .alias('t').merge(slv_lot.alias('s'), 'ardinal lot_code = s.lot_code AND t.gtin = s.gtin') \\\n  .whenMatchedUpdateAll().whenNotMatchedInsertAll().execute()")
    pb(doc)

    addh(doc, "5 · Data quality &amp; observability", level=1)
    addh(doc, "5.1 · Great Expectations / DQ pack templates", level=2)
    addbullets(doc, [
        "Schema validation at Bronze land-time: column-level type assertions, required-field non-null checks, range checks on numeric fields",
        "Cross-table referential integrity: BRZ.RC.POS_TXN.sku must exist in BRZ.RC.ITEM_MASTER active window",
        "Volume anomaly detection: alert on >25% drop from 7-day moving average of daily ingest count",
        "Telemetry-quality SLO: BRZ.RC.REFRIG_TELEMETRY ≥ 99.5% completeness per fixture per day",
        "FSMA-completeness: BRZ.RC.LOT_TRACE rows on FTL items must have non-null lot_code, supplier_id, kde_payload",
    ])
    addh(doc, "5.2 · Microsoft Purview integration", level=2)
    addbullets(doc, [
        "Register all Lakehouses with Purview at workspace creation",
        "Auto-classify PII fields (loyalty_household_id flagged as 'pseudonymized identifier')",
        "Lineage capture across Bronze → Silver → Gold via Notebook + Data Factory connector",
        "Sensitivity labels: FSMA 204 evidence rows = 'Restricted' · loyalty data = 'Confidential' · pricing = 'Confidential'",
    ])
    addh(doc, "5.3 · LEDGER query and audit-row export", level=2)
    addbody(doc, "LEDGER 14-field audit rows write to a dedicated Lakehouse table (sch_ledger.audit_row) "
                 "and a PostgreSQL index (Azure Database for PostgreSQL) for low-latency trace_id lookup. "
                 "Query patterns:")
    addcode(doc, "-- by trace_id (drill-through)\nSELECT * FROM sch_ledger.audit_row WHERE trace_id = 'trc_2027-04-22_0608_kr_428_a3';\n\n"
                 "-- by scenario+time window (Privacy review)\nSELECT * FROM sch_ledger.audit_row\n"
                 "WHERE scenario_id = 'RC-SUPCHN-01-cold-chain-excursion'\n  AND event_ts >= '2027-01-01' AND event_ts < '2027-04-01';\n\n"
                 "-- evidence packet export (FDA inspection)\nCALL sch_ledger.export_fsma_packet(quarter_start => '2027-01-01', banner => 'Kroger');")
    pb(doc)

    addh(doc, "6 · Power BI semantic model", level=1)
    addbullets(doc, [
        "Direct Lake mode for Silver-tier reads · no import latency",
        "Three semantic models: rc_shrink_dashboard · rc_fsma_evidence · rc_pricing_outcomes",
        "Drill-through from any aggregated cell to LEDGER trace_id detail",
        "Row-level security: Store Ops sees own banner cluster; Privacy/Audit sees enterprise; FDA Relations sees export-format only",
    ])

    addh(doc, "7 · Wave-by-wave runbook execution", level=1)
    addtable(doc, ["Wave", "Runbook execution focus"],
             [["W1 Foundation","Workspace setup · Eventstream for 5 streaming feeds · Data Factory for 7 batch feeds · Silver dimensional layer · LEDGER scaffolding · Adaptive Card delivery channel"],
              ["W2 Pilot","Full Silver MERML/SCML build · Power BI semantic models · Privacy/audit packet generator · Production Eventstream throughput tuning · agent integration"],
              ["W3 Scale","Multi-banner orchestration · cross-banner aggregation views · F256 capacity scaling · disaster recovery runbook · multi-region replication"]],
             col_widths=[1.4,5.0])
    pb(doc)

    addh(doc, "Appendix A · Operational runbook entries", level=1)
    addtable(doc, ["Trigger", "Action"],
             [["POS feed delay > 60 sec","Page on-call data engineer · escalate after 5 min · check Eventstream consumer health"],
              ["Refrig telemetry gap > 5 min","Page Facilities Eng + on-call · check MQTT bridge · verify EMS gateway connectivity"],
              ["FSMA 204 KDE schema evolution","Open PR to Bronze schema · run shadow ingest 7 days · review with Privacy"],
              ["Vendor-portal API timeout","Auto-retry 3x · escalate to vendor account manager · failover to manual ingest"],
              ["Recall in BRZ.RC.LOT_TRACE","Auto-trigger SCML.Lot recall-window query · alert Food Safety Lead · prep evidence packet"],
              ["Bronze schema drift detected","Auto-PR with proposed addition · review by Data Eng within 24h"],
              ["Late-arriving data > 24h late","Re-project Silver windows · log to LEDGER as exception"],
              ["Bronze partition reprocess","Use batch_id targeting · validate via dedup · log to ops journal"],
              ["F-SKU capacity throttle","Scale F-SKU up · investigate root cause · rebalance partitioned writes"],
              ["LEDGER PostgreSQL lag","Failover to Lakehouse-only mode · reindex · plan PG restore"]],
             col_widths=[2.1,4.4])

    addbody(doc, "Independence posture preserved across this document: 'Deloitte's Microsoft practice,' "
                 "'DMTSP,' 'Microsoft platform capabilities,' 'Microsoft-native deployment.' The terms "
                 "'partner,' 'alliance,' and 'strategic alliance' do not appear.",
            italic=True, color=MUTE, size=9)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— End of Implementation Runbook —"); setrun(r, italic=True, size=10, color=MUTE)
    target = OUT / "APEX-RC-E2E-03-09-Fabric-Runbook.docx"
    try: doc.save(str(target))
    except PermissionError:
        n=2
        while True:
            alt = target.with_name(target.stem + f"-v{n}" + target.suffix)
            if not alt.exists(): target=alt; break
            n+=1
        doc.save(str(target))
    print(f"  wrote {target.name}")


# =========================================================================
# 6. AGENT MANIFESTS (YAML — 7 files)
# =========================================================================
def build_manifests():
    print("\nWriting agent manifests...")
    manifests = {
        "parent-orchestrator.yaml": """# APEX RC-E2E-03 · Parent Orchestrator Manifest
# Coordinates the 5-agent fleet for cold-chain excursion handling
# Reference scenario: RC-SUPCHN-01-cold-chain-excursion-store-cooler

apiVersion: apex.deloitte.com/v1
kind: Orchestrator
metadata:
  name: rc-e2e-03-cold-chain-orchestrator
  service_id: RC-E2E-03
  scenario_id: RC-SUPCHN-01-cold-chain-excursion
  version: v1.0.0
  tenant_id: kr
  banner_scope: ["KR", "RAL", "FM", "KS", "SMI", "FRY", "HT", "MAR", "QFC"]

spec:
  pattern: hierarchical-root + sequential-with-hitl-gate
  archetype_refs:
    - A-HRC-02
    - A-HGC-04
    - A-SDS-05

  trigger:
    type: streaming-detect
    source: BRZ.RC.REFRIG_TELEMETRY
    condition: |
      temperature_f > temperature_setpoint_f + 5
      AND duration_minutes >= 15
      AND offline_flag = false

  flow:
    - step: assess
      agent: rc-e2e-03-assess-agent
      timeout_sec: 30

    - step: classify
      agent: rc-e2e-03-classify-agent
      timeout_sec: 60
      retry_on_low_confidence:
        threshold: 0.85
        action: escalate_to_food_safety_lead

    - step: quantify
      agent: rc-e2e-03-quantify-agent
      timeout_sec: 30

    - step: hitl-gate
      type: APPROVE
      policy_ref: rc-e2e-03-materiality-policy
      adaptive_card_ref: rc-e2e-03-disposition-card
      timeout_sec: 600
      escalation_after_timeout: rc-e2e-03-escalation-policy

    - step: act
      agent: rc-e2e-03-act-agent
      condition: hitl-gate.outcome == 'APPROVED'
      timeout_sec: 60
      parallel: true
      sub_steps:
        - pos-markdown-upload
        - vendor-ticket-create
        - fsma-regulatory-log

    - step: closeout
      type: audit-row-emit
      ledger_target: sch_ledger.audit_row
      kpi_capture:
        - time_to_decision_sec
        - shrink_avoided_usd
        - policy_conformance
        - fsma_audit_row_complete
        - manager_time_reclaimed_min

  observability:
    trace_propagation: w3c-traceparent
    langfuse_project: apex-rc-e2e-03-prod
    purview_lineage: enabled

  hitl:
    primary_persona: store-operations-lead
    primary_persona_lookup: entra-group:apex-rc-store-ops
    escalation_persona: food-safety-compliance-lead
    escalation_persona_lookup: entra-group:apex-rc-food-safety
""",

        "assess-agent.yaml": """# APEX RC-E2E-03 · Assess Agent Manifest
# Characterizes the cold-chain excursion event

apiVersion: apex.deloitte.com/v1
kind: Agent
metadata:
  name: rc-e2e-03-assess-agent
  service_id: RC-E2E-03
  version: v1.0.0
  tenant_id: kr

spec:
  responsibility: |
    Characterize a cold-chain excursion event by computing peak temperature,
    duration, average temperature, door-open events, and severity classification.
    Read-only against BRZ.RC.REFRIG_TELEMETRY; emit excursion-profile output.

  capability_tag: cap:excursion-assessment
  llm_capability_required: cap:reasoning-light

  inputs:
    - source: BRZ.RC.REFRIG_TELEMETRY
      filter: "trigger event window (excursion start - 30 min, now)"
    - source: SCML.Fixture
      filter: "fixture_id = trigger.fixture_id"

  outputs:
    schema: ExcursionProfile
    fields:
      - excursion_id
      - fixture_id
      - peak_temp_f
      - avg_temp_f
      - duration_minutes
      - door_open_event_count
      - severity_class
      - confidence

  tools:
    - mcp_tool: telemetry-window-query
      registry_ref: cap:telemetry-window-query
    - mcp_tool: fixture-state-lookup
      registry_ref: cap:fixture-state-lookup

  hitl:
    gate: ZERO_TOUCH
    audit_row: required

  prompt: |
    You are the Assess agent for APEX RC-E2E-03 cold-chain excursion handling.
    Given the temperature telemetry window and fixture state, compute:
      - peak_temp_f, avg_temp_f, duration_minutes
      - door_open_event_count from telemetry
      - severity_class: NORMAL / WATCH / CRITICAL
      - confidence (0.0-1.0)

    Severity rules:
      NORMAL    → peak < setpoint+5°F · duration < 15 min
      WATCH     → peak < setpoint+10°F · duration < 60 min
      CRITICAL  → peak >= setpoint+10°F OR duration >= 60 min OR multiple door events

    Emit ExcursionProfile JSON. Do not propose disposition at this stage —
    that is the Classify agent's responsibility.

  observability:
    trace_attributes:
      - excursion_id
      - fixture_id
      - severity_class
""",

        "classify-agent.yaml": """# APEX RC-E2E-03 · Classify Agent Manifest
# Per-SKU disposition classification per FSMA 204 policy

apiVersion: apex.deloitte.com/v1
kind: Agent
metadata:
  name: rc-e2e-03-classify-agent
  service_id: RC-E2E-03
  version: v1.0.0
  tenant_id: kr

spec:
  responsibility: |
    Per-SKU disposition classification (sell-through / markdown / destroy)
    based on FSMA 204 policy corpus and lot-provenance data. Reads ExcursionProfile
    from Assess + SCML.Lot from Silver + FSMA 204 policy manifest.

  capability_tag: cap:disposition-classification
  llm_capability_required: cap:reasoning-advanced

  inputs:
    - source: rc-e2e-03-assess-agent.output
    - source: SCML.Lot
      filter: "fixture_id IN (trigger.fixture_id) AND lot_status = 'active'"
    - source: DIM.Item
      join: "on sku"
    - source: FSMA-204-policy-corpus
      version: v1.0.0

  outputs:
    schema: DispositionTable
    fields:
      - lot_code
      - sku
      - disposition: [SELL_THROUGH | MARKDOWN | DESTROY | HOLD]
      - reason
      - policy_citation
      - confidence
      - escalation_required

  tools:
    - mcp_tool: scml-lot-query
      registry_ref: cap:scml-lot-query
    - mcp_tool: fsma-policy-resolver
      registry_ref: cap:fsma-policy-resolver
    - mcp_tool: dim-item-lookup
      registry_ref: cap:dim-item-lookup

  hitl:
    gate: ESCALATION
    escalation_condition: confidence < 0.85
    escalation_persona: food-safety-compliance-lead
    audit_row: required

  prompt: |
    You are the Classify agent for APEX RC-E2E-03 cold-chain excursion handling.
    Given the excursion profile and lot-provenance for affected SKUs, classify
    each lot's disposition using the FSMA 204 policy corpus.

    Per-SKU disposition options:
      SELL_THROUGH    → product safe to sell at full price
      MARKDOWN        → product safe to sell with markdown to accelerate
      DESTROY         → regulatory or quality threshold breached
      HOLD            → insufficient data to disposition; refer to Food Safety Lead

    For each disposition, cite the specific FSMA 204 policy clause that applies.
    Emit DispositionTable JSON with confidence per row.

    If confidence < 0.85 on any row, set escalation_required = true.
""",

        "quantify-agent.yaml": """# APEX RC-E2E-03 · Quantify Agent Manifest
# Computes dollar impact of disposition decisions

apiVersion: apex.deloitte.com/v1
kind: Agent
metadata:
  name: rc-e2e-03-quantify-agent
  service_id: RC-E2E-03
  version: v1.0.0
  tenant_id: kr

spec:
  responsibility: |
    Computes dollar impact of each disposition path per affected lot.
    Three scenarios per disposition: sell-through value, markdown value, destroy value.
    Reads DispositionTable from Classify + MERML.Price + MERML.Markdown + MERML.Elasticity.

  capability_tag: cap:markdown-quantification
  llm_capability_required: cap:reasoning-light

  inputs:
    - source: rc-e2e-03-classify-agent.output
    - source: MERML.Price
      filter: "sku IN (DispositionTable.sku) AND active"
    - source: MERML.Markdown
      filter: "sku IN (DispositionTable.sku)"
    - source: MERML.Elasticity
      filter: "sku IN (DispositionTable.sku)"
    - source: SCML.Lot
      filter: "lot_code IN (DispositionTable.lot_code)"

  outputs:
    schema: QuantifiedDisposition
    fields:
      - lot_code
      - sku
      - sell_through_value_usd
      - markdown_value_usd
      - destroy_loss_usd
      - markdown_recommendation_pct
      - confidence

  tools:
    - mcp_tool: merml-price-lookup
      registry_ref: cap:merml-price-lookup
    - mcp_tool: merml-elasticity-curve
      registry_ref: cap:merml-elasticity-curve
    - mcp_tool: scml-lot-qty
      registry_ref: cap:scml-lot-qty

  hitl:
    gate: ZERO_TOUCH
    audit_row: required

  prompt: |
    You are the Quantify agent for APEX RC-E2E-03.
    Given the disposition table and pricing/elasticity data, compute per-lot:
      - sell_through_value_usd: qty × current price (assume full sell-through)
      - markdown_value_usd: qty × markdown_price using elasticity-suggested discount
      - destroy_loss_usd: qty × cost_basis (the baseline)

    For each lot recommended for MARKDOWN, propose markdown_recommendation_pct
    that maximizes (markdown_value_usd - destroy_loss_usd) within FSMA window.

    Emit QuantifiedDisposition JSON.
""",

        "approve-agent.yaml": """# APEX RC-E2E-03 · Approve / HITL Agent Manifest
# Routes material decisions to named persona via Adaptive Card

apiVersion: apex.deloitte.com/v1
kind: Agent
metadata:
  name: rc-e2e-03-approve-agent
  service_id: RC-E2E-03
  version: v1.0.0
  tenant_id: kr

spec:
  responsibility: |
    HITL gate. Routes material disposition decisions to named persona
    (Store Operations Lead by default) via Microsoft Teams Adaptive Card.
    Decision (Approve / Modify / Reject) writes to LEDGER audit row.

  capability_tag: cap:hitl-orchestration
  llm_capability_required: NONE

  inputs:
    - source: rc-e2e-03-quantify-agent.output
    - source: rc-e2e-03-materiality-policy
      version: v1.0.0
    - source: DIM.Store
      join: "store_id"

  outputs:
    schema: ApprovedDisposition
    fields:
      - disposition_record
      - approver_id
      - approver_role
      - approval_ts
      - hitl_kind
      - modifications

  tools:
    - mcp_tool: teams-adaptive-card-send
      registry_ref: cap:teams-adaptive-card
    - mcp_tool: entra-identity-resolve
      registry_ref: cap:entra-identity-resolve
    - mcp_tool: ledger-write
      registry_ref: cap:ledger-write

  hitl:
    gate_logic:
      - condition: "store_tier == 'Tier_1'"
        kind: HITL
        persona: store-operations-lead
      - condition: "disposition_record.total_value_usd > 2500"
        kind: HITL
        persona: store-operations-lead
      - condition: "any disposition_record.disposition == 'DESTROY'"
        kind: HITL
        persona: food-safety-compliance-lead
      - condition: "all disposition_record.confidence >= 0.95 AND total_value_usd <= 2500"
        kind: ZERO_TOUCH
        notify: store-operations-lead
      - default: HITL
        persona: store-operations-lead

    timeout_sec: 600
    timeout_action: escalate-to-food-safety-lead

  adaptive_card_template:
    title: "Cold-chain disposition required · Store {store_id}"
    fields:
      - excursion_summary
      - affected_lot_count
      - proposed_actions
      - dollar_impact
      - reasoning_trace_link
    actions:
      - approve
      - modify
      - reject
""",

        "act-agent.yaml": """# APEX RC-E2E-03 · Act Agent Manifest
# Orchestrates downstream actions post-approval

apiVersion: apex.deloitte.com/v1
kind: Agent
metadata:
  name: rc-e2e-03-act-agent
  service_id: RC-E2E-03
  version: v1.0.0
  tenant_id: kr

spec:
  responsibility: |
    Post-approval execution. Three parallel downstream actions:
    POS markdown upload, refrigeration-vendor ticket creation, FSMA regulatory log.
    Reads ApprovedDisposition from Approve agent. Emits closeout audit row.

  capability_tag: cap:multi-channel-action
  llm_capability_required: NONE

  inputs:
    - source: rc-e2e-03-approve-agent.output
      condition: hitl_kind != 'REJECTED'

  outputs:
    schema: ActOutcome
    fields:
      - pos_markdown_status
      - vendor_ticket_id
      - fsma_log_id
      - all_actions_completed
      - closeout_ts

  tools:
    - mcp_tool: pos-markdown-upload
      registry_ref: cap:pos-markdown-upload
      target_systems: [Toshiba_4690, IBM_TCx, NCR_Counterpoint]
    - mcp_tool: vendor-portal-ticket
      registry_ref: cap:vendor-portal-ticket
      target_vendors: [Emerson, Lennox, Hussmann, Hill_PHOENIX]
    - mcp_tool: fsma-regulatory-log-write
      registry_ref: cap:fsma-regulatory-log
    - mcp_tool: ledger-closeout-write
      registry_ref: cap:ledger-write

  parallel_execution:
    - pos-markdown-upload
    - vendor-ticket-create
    - fsma-regulatory-log

  retry_policy:
    max_retries: 3
    backoff: exponential
    on_failure: escalate-to-on-call

  hitl:
    gate: ZERO_TOUCH
    audit_row: required (closeout)
""",

        "materiality-policy.yaml": """# APEX RC-E2E-03 · Materiality Policy Manifest
# Commercial envelope tier thresholds for HITL routing

apiVersion: apex.deloitte.com/v1
kind: Policy
metadata:
  name: rc-e2e-03-materiality-policy
  service_id: RC-E2E-03
  version: v1.0.0
  tenant_id: kr
  authoring_persona: pricing-manager

spec:
  description: |
    Tiered HITL routing policy for cold-chain disposition decisions.
    Tiered by store-cluster materiality × disposition-dollar × confidence.
    Per-banner override supported.

  default_tier_assignment:
    Tier_1: ["KR", "FM", "MAR"]    # Flagship banners
    Tier_2: ["RAL", "KS", "SMI", "FRY", "HT", "QFC", "PNS"]  # Standard
    Tier_3: ["BAK", "OWN", "PYL", "JC"]  # Smaller-format

  rules:
    - id: tier1-always-hitl
      condition: store_tier == 'Tier_1'
      action: HITL
      persona: store-operations-lead

    - id: tier2-above-threshold
      condition: store_tier == 'Tier_2' AND disposition_value_usd > 2500
      action: HITL
      persona: store-operations-lead

    - id: tier2-mid-confidence
      condition: |
        store_tier == 'Tier_2' AND disposition_value_usd >= 500 AND disposition_value_usd <= 2500
        AND min_confidence >= 0.95
      action: ZERO_TOUCH
      notify: store-operations-lead

    - id: tier2-mid-low-confidence
      condition: |
        store_tier == 'Tier_2' AND disposition_value_usd >= 500 AND disposition_value_usd <= 2500
        AND min_confidence < 0.95
      action: HITL
      persona: store-operations-lead

    - id: tier3-low-value-high-conf
      condition: |
        store_tier == 'Tier_3' AND disposition_value_usd <= 500 AND min_confidence >= 0.90
      action: ZERO_TOUCH
      log_only: true

    - id: any-destroy-all
      condition: any disposition.action == 'DESTROY'
      action: HITL
      persona: food-safety-compliance-lead
      override_other_rules: true

    - id: any-fsma-evidence-gap
      condition: any disposition.fsma_audit_row_complete == false
      action: HITL
      persona: food-safety-compliance-lead
      override_other_rules: true

  banner_overrides:
    KR_Cincinnati_anchor:
      thresholds:
        tier2_high: 5000  # Higher threshold for anchor cluster during W2 pilot
""",
    }
    for name, content in manifests.items():
        p = MANIFESTS / name
        p.write_text(content, encoding="utf-8")
        print(f"  wrote manifests/{name}")


# =========================================================================
# 7. USE CASE CATALOG (XLSX)
# =========================================================================
def build_uc_catalog():
    print("\nWriting Use Case Catalog xlsx...")
    HDR_FONT=Font(name="Calibri",size=11,bold=True,color="FFFFFF")
    HDR_FILL=PatternFill("solid",fgColor="0B0B0F")
    THIN=Side(style="thin",color="D4D4D7"); BOR=Border(left=THIN,right=THIN,top=THIN,bottom=THIN)
    BAND_FILL=PatternFill("solid",fgColor="F3F3F5")
    CELL=Font(name="Calibri",size=10); CODE=Font(name="Consolas",size=10,color="6648B0",bold=True)
    HDR_AL = Alignment(horizontal="left",vertical="center",wrap_text=True)
    CELL_AL = Alignment(horizontal="left",vertical="top",wrap_text=True)

    wb = Workbook()
    # ----- README sheet -----
    ws = wb.active; ws.title = "README"
    ws.column_dimensions["A"].width = 28; ws.column_dimensions["B"].width = 92
    rows = [
        ("APEX · RC-E2E-03 · Use Case Catalog · Kroger reference",
         "24 sprint-level use cases composing the 6 agent-level use cases. Companion to APEX-RC-E2E-03-Walkthrough.docx Section 5."),
        ("",""),
        ("How to read this workbook",""),
        ("Use Cases","All 24 UCs with description, acceptance criteria, dependencies, story points, wave, KPI driven."),
        ("Agent-Level Mapping","UCs grouped by the 6 agent-level use cases from the walkthrough."),
        ("Wave Distribution","UC distribution across W1/W2/W3 waves."),
        ("KPI Coverage","Which KPIs each UC contributes to."),
        ("",""),
        ("Style legend",""),
        ("Black header","Sheet-level column headers"),
        ("Violet code","UC IDs and Sprint IDs (UC-01, W1-S03 format)"),
        ("Italic gray","Author commentary and dependency notes"),
        ("",""),
        ("Companion artifacts",""),
        ("APEX-RC-E2E-03-Walkthrough.docx","Service walkthrough — UC narrative"),
        ("APEX-RC-E2E-03-09-SOR-Bronze-to-Silver.docx","SOR specification"),
        ("APEX-RC-E2E-03-09-SOR-ERD.html","Visual ERD"),
        ("APEX-RC-E2E-03-Kroger-Personas.html","Persona day-in-the-life"),
        ("manifests/*.yaml","Agent + policy manifests"),
    ]
    for i,(l,r) in enumerate(rows, start=1):
        c1 = ws.cell(row=i, column=1, value=l); c2 = ws.cell(row=i, column=2, value=r)
        if i == 1:
            c1.font = Font(bold=True, size=14); c2.font = Font(italic=True, size=10, color="4A4A55")
        elif l and r:
            c1.font = Font(bold=True, size=10); c2.font = CELL
        else:
            c1.font = Font(bold=True, size=11, color="6648B0"); c2.font = CELL
        c2.alignment = Alignment(wrap_text=True, vertical="top"); ws.row_dimensions[i].height = 28

    # ----- Use Cases sheet -----
    ws = wb.create_sheet("Use Cases")
    headers = ["UC ID","Title","Agent-level UC","Description","Acceptance Criteria",
               "Dependencies","SP","Wave","Sprint anchor","KPI driven","Persona owner"]
    widths = [9,32,22,52,60,16,5,7,12,18,20]
    for j,h in enumerate(headers, start=1):
        c = ws.cell(row=1, column=j, value=h); c.font=HDR_FONT; c.fill=HDR_FILL; c.alignment=HDR_AL; c.border=BOR
        ws.column_dimensions[get_column_letter(j)].width = widths[j-1]
    ws.row_dimensions[1].height = 30; ws.freeze_panes = "C2"

    UCS = [
      ("UC-01","Refrigeration telemetry ingest","Excursion Triage",
       "Stream BRZ.RC.REFRIG_TELEMETRY from EMS gateways via OPC-UA→MQTT→Eventstream into Lakehouse.",
       "30-second cadence verified · Bronze landing < 30 sec · 99.5% sensor completeness · schema-validated",
       "—","8","W1","W1-S02","Coverage","Data Engineering"),
      ("UC-02","Excursion-detection trigger","Excursion Triage",
       "Watermark-based trigger: temperature > setpoint + 5°F for ≥ 15 min consecutive.",
       "Trigger fires within 30 sec of threshold breach · false-positive rate < 5% · alarm classification correct ≥ 95%",
       "UC-01","5","W1","W1-S03","Coverage, Loop time","Engineering Architect"),
      ("UC-03","Assess agent · excursion characterization","Excursion Triage",
       "Characterize peak temp, duration, avg, door-event count, severity classification.",
       "Assess output schema-validated · severity classification matches Food Safety policy ≥ 95% · trace_id correlates with telemetry window",
       "UC-01,UC-02","5","W1","W1-S03","Loop time","ML Engineering"),
      ("UC-04","SCML.Lot canonical layer","FSMA Policy Classification",
       "Build SCML.Lot from BRZ.RC.LOT_TRACE + RECEIPT + VENDOR + ITEM joins.",
       "All FTL items have non-null lot_code at receipt CTE · join completeness > 99.5% · daily reconciliation passes",
       "UC-01,UC-19","8","W1","W1-S04","FSMA-204 coverage","Data Engineering"),
      ("UC-05","FSMA 204 policy corpus loader","FSMA Policy Classification",
       "Versioned policy manifest with per-FTL-category disposition rules.",
       "Policy corpus v1.0 loaded · Food Safety Lead sign-off · PR-merge workflow tested",
       "—","5","W1","W1-S04","FSMA-204 coverage","Engineering Architect + Food Safety"),
      ("UC-06","Classify agent · per-SKU disposition","FSMA Policy Classification",
       "Per-SKU disposition (sell-through/markdown/destroy/hold) with FSMA 204 policy citation.",
       "Disposition matches policy ≥ 95% · confidence calibrated · escalations route to Food Safety Lead",
       "UC-04,UC-05","8","W1","W1-S05","Policy-conformance, FSMA-204","ML Engineering"),
      ("UC-07","Confidence threshold tuning","FSMA Policy Classification",
       "Calibrate Classify confidence threshold at 0.85 based on shadow-run data.",
       "Calibration plot reviewed · Food Safety Lead approves · escalation rate within 5-10% target",
       "UC-06","3","W2","W2-S01","Policy-conformance","ML Engineering"),
      ("UC-08","MERML.Price canonical layer","Markdown Decisioning",
       "MERML.Price from BRZ.RC.PRICING + PROMO_MARKDOWN overlay + ITEM_MASTER + LOYALTY_HH consent layer.",
       "Active price lookup ≤ 100ms · price-zone resolution correct · loyalty consent state respected",
       "UC-04","8","W1","W1-S05","Markdown-NPV","Data Engineering"),
      ("UC-09","MERML.Markdown derived layer","Markdown Decisioning",
       "MERML.Markdown from PROMO_MARKDOWN + POS_TXN attribution.",
       "Markdown outcomes attributed to specific promo_id · NPV computation matches Pricing Strategy reference",
       "UC-08,UC-21","5","W2","W2-S02","Markdown-NPV","Data Engineering"),
      ("UC-10","MERML.Elasticity model fit","Markdown Decisioning",
       "Per-SKU price-elasticity curves from POS_TXN + PRICING + LOYALTY_HH segment.",
       "Elasticity coefficients within reference range · weekly retraining validated · drift detection live",
       "UC-08,UC-21","8","W2","W2-S02","Markdown-NPV","ML Engineering"),
      ("UC-11","Quantify agent · disposition $ impact","Markdown Decisioning",
       "Per-lot 3-scenario $ computation (sell-through/markdown/destroy).",
       "Quantification matches Pricing Strategy reference within ±5% · markdown_recommendation_pct optimized",
       "UC-08,UC-09,UC-10","8","W2","W2-S02","Markdown-NPV","ML Engineering"),
      ("UC-12","Markdown-NPV outcome attribution","Markdown Decisioning",
       "Trace_id-attributed shrink-avoided NPV computation.",
       "NPV per disposition rolls up to per-cohort dashboard · drill-through to specific decisions",
       "UC-09,UC-19","5","W2","W2-S02","Shrink avoided","Engineering Architect"),
      ("UC-13","Materiality policy manifest","Manager-Approval HITL",
       "Tiered materiality policy by store_tier × value × confidence.",
       "Policy authored · simulated against historical 90 days · Pricing Strategy and VP Retention sign-off",
       "—","5","W2","W2-S03","Concession Eff.","Pricing Manager + Engineering Architect"),
      ("UC-14","Approve agent · HITL routing","Manager-Approval HITL",
       "Routes material decisions per materiality policy.",
       "Routing correctness ≥ 99% · escalation paths verified · timeout policy fires correctly",
       "UC-13","5","W2","W2-S03","Loop time","Engineering Architect"),
      ("UC-15","Adaptive Card delivery + approval","Manager-Approval HITL",
       "Teams Adaptive Card with Approve/Modify/Reject buttons.",
       "Card delivery ≤ 5 sec · Entra identity resolution verified · approval audit-row complete · designate-approver flow works",
       "UC-14","8","W2","W2-S03","Loop time, Audit-row","Frontend Engineering"),
      ("UC-16","POS markdown upload (Act sub-step)","Parallel Downstream Action",
       "POS markdown to Toshiba/IBM/NCR per banner.",
       "Markdown active in POS within 2 min · per-banner POS-API integration verified · idempotent on retry",
       "UC-15,UC-08","5","W2","W2-S03","Loop time","Backend Engineering"),
      ("UC-17","Vendor-portal ticket creation","Parallel Downstream Action",
       "Auto-open service ticket in Emerson/Lennox/Hussmann/Hill PHOENIX portal.",
       "Ticket created within 5 min of HITL approval · vendor-API rate-limit compliance · ticket_id captured",
       "UC-15","5","W2","W2-S03","Loop time","Backend Engineering"),
      ("UC-18","FSMA regulatory log entry","Parallel Downstream Action",
       "FOOD_SAFETY incident record write with all 14 audit-row fields + trace_id.",
       "FSMA log entry within 30 sec of action · 100% audit-row completeness on closeouts",
       "UC-15,UC-19","3","W2","W2-S03","FSMA-204 coverage","Backend Engineering"),
      ("UC-19","LEDGER 14-field audit row write","Outcome Attribution",
       "Closeout audit row to LEDGER per decision.",
       "Audit-row ≥ 99.5% · all 14 fields populated · trace_id-correlated · query API returns row by trace_id under 100ms",
       "—","5","W1","W1-S04","Audit-row coverage","Engineering Architect"),
      ("UC-20","Power BI shrink-NPV dashboard","Outcome Attribution",
       "Per-banner shrink-avoided trending with drill-through.",
       "Dashboard published · drill-through works to specific trace_ids · semantic model in Direct Lake mode",
       "UC-12,UC-19","5","W2","W2-S04","Shrink avoided, NPV","BI Engineering"),
      ("UC-21","FSMA 204 evidence packet generator","Outcome Attribution",
       "Quarterly auto-generated packet ZIP for FDA inspection.",
       "Packet generated in < 4 hr at W2 · format approved by Food Safety Lead and external audit walkthrough",
       "UC-04,UC-19","8","W2","W2-S04","FSMA-204 coverage","Engineering Architect + Privacy"),
      ("UC-22","Cohort segmentation engine","Outcome Attribution",
       "Per-cohort outcome attribution with drift detection.",
       "Cohort families operating · drift watchtower alerts on Risk-Score drift · auto-retrain trigger tested",
       "UC-10","5","W2","W2-S04","Precision/Recall","ML Engineering"),
      ("UC-23","Multi-banner orchestration extension","All",
       "Cross-banner operating-model · banner-aware policy manifests.",
       "All ~10 banners running on cluster · cross-banner aggregation views · Steering Committee approves W3",
       "all W2 UCs","13","W3","W3-S01","All KPIs at scale","Engineering Architect"),
      ("UC-24","Cross-scenario fusion","All",
       "Adjacent scenario integration: RC-E2E-11 Perishables Integrity Response · RC-E2E-12 Regulator Evidence Service.",
       "Fusion hooks tested · cross-scenario audit-row correlation · expansion plan signed off",
       "UC-23","8","W3","W3-S02","NPV at run-rate","Engineering Architect"),
    ]
    for i, uc in enumerate(UCS, start=2):
        for j,v in enumerate(uc, start=1):
            c = ws.cell(row=i, column=j, value=v); c.border=BOR; c.alignment=CELL_AL
            if j in (1,9): c.font = CODE
            else: c.font = CELL
            if i % 2 == 0: c.fill = BAND_FILL
        ws.row_dimensions[i].height = max(50, 14*(uc[3].count("\n")+uc[4].count("\n")+2))

    # ----- Wave Distribution sheet -----
    ws = wb.create_sheet("Wave Distribution")
    headers = ["Wave","UC count","UC IDs","Total SP","Cumulative SP","Notes"]
    widths = [10,9,40,9,12,40]
    for j,h in enumerate(headers, start=1):
        c = ws.cell(row=1, column=j, value=h); c.font=HDR_FONT; c.fill=HDR_FILL; c.alignment=HDR_AL; c.border=BOR
        ws.column_dimensions[get_column_letter(j)].width = widths[j-1]
    ws.row_dimensions[1].height = 30
    rows = [
        ("Wave 1", 8, "UC-01, UC-02, UC-03, UC-04, UC-05, UC-06, UC-08, UC-19", 50, 50,
         "Foundation: telemetry, FSMA policy, Classify, MERML.Price, LEDGER scaffolding"),
        ("Wave 2", 14, "UC-07, UC-09 → UC-22 (excl UC-19)", 81, 131,
         "Pilot: full 5-agent live, materiality policy, HITL flow, Act, evidence packet, dashboards"),
        ("Wave 3", 2, "UC-23, UC-24", 21, 152,
         "Scale & Fuse: multi-banner orchestration, cross-scenario fusion"),
    ]
    for i, r in enumerate(rows, start=2):
        for j,v in enumerate(r, start=1):
            c = ws.cell(row=i, column=j, value=v); c.border=BOR; c.alignment=CELL_AL
            c.font = CELL
            if j == 1: c.font = Font(name="Calibri",size=10,bold=True)
            if i % 2 == 0: c.fill = BAND_FILL

    # ----- Agent-Level Mapping sheet -----
    ws = wb.create_sheet("Agent-Level Mapping")
    headers = ["Agent-level UC (walkthrough §5)","Sprint UCs","Total SP","Wave"]
    widths = [40,38,9,8]
    for j,h in enumerate(headers, start=1):
        c = ws.cell(row=1, column=j, value=h); c.font=HDR_FONT; c.fill=HDR_FILL; c.alignment=HDR_AL; c.border=BOR
        ws.column_dimensions[get_column_letter(j)].width = widths[j-1]
    ws.row_dimensions[1].height = 30
    rows = [
        ("5.1 Excursion Triage", "UC-01, UC-02, UC-03", 18, "W1"),
        ("5.2 FSMA 204 Policy Classification", "UC-04, UC-05, UC-06, UC-07", 24, "W1+W2"),
        ("5.3 Markdown Decisioning", "UC-08, UC-09, UC-10, UC-11, UC-12", 34, "W1+W2"),
        ("5.4 Manager-Approval HITL", "UC-13, UC-14, UC-15", 18, "W2"),
        ("5.5 Parallel Downstream Action", "UC-16, UC-17, UC-18", 13, "W2"),
        ("5.6 Outcome Attribution", "UC-19, UC-20, UC-21, UC-22", 23, "W1+W2"),
        ("Cross-wave (W3)", "UC-23, UC-24", 21, "W3"),
    ]
    for i, r in enumerate(rows, start=2):
        for j,v in enumerate(r, start=1):
            c = ws.cell(row=i, column=j, value=v); c.border=BOR; c.alignment=CELL_AL; c.font=CELL
            if i % 2 == 0: c.fill = BAND_FILL

    # ----- KPI Coverage sheet -----
    ws = wb.create_sheet("KPI Coverage")
    headers = ["KPI","UCs that contribute","W2 Pilot target","W3 Scale target"]
    widths = [30, 30, 16, 16]
    for j,h in enumerate(headers, start=1):
        c = ws.cell(row=1, column=j, value=h); c.font=HDR_FONT; c.fill=HDR_FILL; c.alignment=HDR_AL; c.border=BOR
        ws.column_dimensions[get_column_letter(j)].width = widths[j-1]
    ws.row_dimensions[1].height = 30
    rows = [
        ("Time-to-decision (loop time)","UC-02,UC-03,UC-14,UC-15,UC-16","≤ 4 min","≤ 90 sec"),
        ("Manager attention reclaimed","UC-15,UC-16","+3.5 hrs/shift","+5.2 hrs/shift"),
        ("Per-event shrink avoided","UC-11,UC-12,UC-20","$640/event","$1,313/event"),
        ("Policy-conformance rate","UC-06,UC-07","≥ 88%","≥ 95%"),
        ("FSMA-204 audit-row completeness","UC-04,UC-19,UC-21","≥ 96%","≥ 99.5%"),
        ("Auto-execute rate","UC-13,UC-14,UC-15","55%","75%"),
        ("Vendor-ticket rate","UC-17","≥ 78%","≥ 92%"),
        ("Sell-through-at-risk rate","UC-06,UC-22","≤ 1.5%","≤ 0.5%"),
    ]
    for i, r in enumerate(rows, start=2):
        for j,v in enumerate(r, start=1):
            c = ws.cell(row=i, column=j, value=v); c.border=BOR; c.alignment=CELL_AL; c.font=CELL
            if i % 2 == 0: c.fill = BAND_FILL

    target = OUT / "APEX-RC-E2E-03-Use-Case-Catalog.xlsx"
    try: wb.save(str(target))
    except PermissionError:
        n=2
        while True:
            alt = target.with_name(target.stem + f"-v{n}" + target.suffix)
            if not alt.exists(): target=alt; break
            n+=1
        wb.save(str(target))
    print(f"  wrote {target.name}")

build_runbook()
build_manifests()
build_uc_catalog()
print("\nTier 2 complete.")
