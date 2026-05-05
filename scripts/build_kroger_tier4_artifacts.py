"""Tier-4 artifacts:
  11. Cross-grocer comparison (XLSX)
  12. AI/ML Model Spec (DOCX)
  13. Privacy & Data Governance Spec (DOCX)
"""
from __future__ import annotations
import io, sys
from pathlib import Path
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
OUT = Path(r"C:\Stage\Clients\Industries\Consumer\Retail\Kroger\02_projects\FY27_Pipeline\assortment-pricing-agentic\deliverables\Tier4-Strategic")
OUT.mkdir(parents=True, exist_ok=True)

INK="0B0B0F"; INK2="4A4A55"; MUTE="8a8a90"; GOLD="B8892E"; TEAL="2A7F73"
VIO="6648B0"; CRIM="A22A39"; BAND="F3F3F5"; HEAD="0B0B0F"; LINE="D4D4D7"

# ---- DOCX helpers ----
def cell_bg(cell, c):
    tcPr = cell._tc.get_or_add_tcPr()
    s = OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto")
    s.set(qn("w:fill"),c.lstrip("#")); tcPr.append(s)
def cell_border(cell, c=LINE):
    tcPr = cell._tc.get_or_add_tcPr()
    bs = OxmlElement("w:tcBorders")
    for e in ("top","left","bottom","right"):
        b = OxmlElement(f"w:{e}"); b.set(qn("w:val"),"single"); b.set(qn("w:sz"),"4"); b.set(qn("w:color"),c.lstrip("#"))
        bs.append(b)
    tcPr.append(bs)
def setrun(r, *, bold=False, italic=False, size=None, color=None, font="Arial"):
    if bold: r.bold = True
    if italic: r.italic = True
    if size is not None: r.font.size = Pt(size)
    if color: r.font.color.rgb = RGBColor.from_string(color.lstrip("#"))
    if font: r.font.name = font
def addbody(doc, t, *, size=11, italic=False, color=INK, sa=6):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(sa)
    r = p.add_run(t); setrun(r, size=size, color=color, italic=italic); return p
def addh(doc, t, *, level=1):
    sz = {1:18,2:14,3:12.5,4:11.5}
    p = doc.add_paragraph(); p.style = doc.styles[f"Heading {level}"]
    p.paragraph_format.space_before = Pt(16 if level==1 else 10); p.paragraph_format.space_after = Pt(4)
    for r in list(p.runs): r.text = ""
    r = p.add_run(t); setrun(r, bold=True, size=sz.get(level,12), color=INK)
    if level == 1:
        pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom"); bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),"6"); bot.set(qn("w:space"),"4"); bot.set(qn("w:color"),GOLD)
        pBdr.append(bot); pPr.append(pBdr)
    return p
def addbullets(doc, items, *, size=11):
    for it in items:
        p = doc.add_paragraph(style="List Bullet"); r = p.add_run(it); setrun(r, size=size, color=INK)
def addtable(doc, header, rows, *, col_widths=None):
    n = len(header); tbl = doc.add_table(rows=len(rows)+1, cols=n); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j,h in enumerate(header):
        c = tbl.rows[0].cells[j]; cell_bg(c,HEAD); cell_border(c); c.paragraphs[0].text=""
        r = c.paragraphs[0].add_run(h); setrun(r,bold=True,size=10,color="FFFFFF")
        if col_widths: c.width = Inches(col_widths[j])
    for i,row in enumerate(rows, start=1):
        for j,v in enumerate(row):
            c = tbl.rows[i].cells[j]; cell_border(c)
            if i%2==0: cell_bg(c,BAND)
            c.paragraphs[0].text=""; r = c.paragraphs[0].add_run(str(v)); setrun(r,size=10,color=INK)
            if col_widths: c.width = Inches(col_widths[j])
    doc.add_paragraph(); return tbl
def addcallout(doc, kind, body, *, accent=GOLD, fill="F6F1E4"):
    tbl = doc.add_table(rows=1, cols=1); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = tbl.rows[0].cells[0]; cell_bg(c,fill); c.paragraphs[0].text=""
    rt = c.paragraphs[0].add_run(kind.upper()); setrun(rt,bold=True,size=9,color=accent)
    p = c.add_paragraph(); p.paragraph_format.space_before = Pt(3)
    r = p.add_run(body); setrun(r,size=10.5,color=INK); doc.add_paragraph()
def addkv(doc, pairs, *, lw=1.6, rw=4.9):
    tbl = doc.add_table(rows=len(pairs), cols=2); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,(k,v) in enumerate(pairs):
        ck,cv = tbl.rows[i].cells; ck.width=Inches(lw); cv.width=Inches(rw)
        cell_bg(ck,BAND); cell_border(ck); cell_border(cv)
        ck.paragraphs[0].text=""; cv.paragraphs[0].text=""
        rk = ck.paragraphs[0].add_run(k); setrun(rk,bold=True,size=10,color=INK)
        rv = cv.paragraphs[0].add_run(v); setrun(rv,size=10,color=INK)
    doc.add_paragraph()
def addeyebrow(doc, t):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(t.upper()); setrun(r,bold=True,size=9,color=MUTE)
    rPr = r._element.get_or_add_rPr(); sp = OxmlElement("w:spacing"); sp.set(qn("w:val"),"80"); rPr.append(sp)
def pb(doc):
    p = doc.add_paragraph(); p.add_run().add_break(WD_BREAK.PAGE)


# =========================================================================
# 11. CROSS-GROCER COMPARISON (XLSX)
# =========================================================================
def build_grocer_comparison():
    print("Building Cross-Grocer Comparison xlsx...")
    HDR_FONT=Font(name="Calibri",size=11,bold=True,color="FFFFFF")
    HDR_FILL=PatternFill("solid",fgColor="0B0B0F")
    THIN=Side(style="thin",color="D4D4D7"); BOR=Border(left=THIN,right=THIN,top=THIN,bottom=THIN)
    BAND_FILL=PatternFill("solid",fgColor="F3F3F5")
    CELL=Font(name="Calibri",size=10); CODE=Font(name="Consolas",size=10,color="6648B0",bold=True)
    HDR_AL = Alignment(horizontal="left",vertical="center",wrap_text=True)
    CELL_AL = Alignment(horizontal="left",vertical="top",wrap_text=True)
    CTR = Alignment(horizontal="center",vertical="center",wrap_text=True)

    wb = Workbook()
    ws = wb.active; ws.title = "README"
    ws.column_dimensions["A"].width = 28; ws.column_dimensions["B"].width = 92
    rows = [
        ("APEX · Cross-Grocer Comparison · RC-E2E-03 + RC-E2E-09",
         "Comparative analysis of US-grocery account portfolio for the same Tier-1 service set. Helps Account Team triage next pursuits."),
        ("",""),
        ("How to read this workbook",""),
        ("Account Comparison","Side-by-side comparison across 9 dimensions for 7 grocer accounts."),
        ("Pursuit Prioritization","Ranked pursuit list with rationale and recommended sequencing."),
        ("RC-E2E-03 Lever Sizing","Per-account shrink-avoided sizing with assumptions."),
        ("RC-E2E-09 Readiness","FSMA 204 readiness scorecard per account."),
        ("Strategic Notes","Per-account engagement-strategy notes, including incumbent-platform constraints."),
        ("",""),
        ("Use this for",""),
        ("Pursuit prioritization","Which next 3-4 grocery pursuits after Kroger reference"),
        ("Resource allocation","Account Team capacity planning across portfolio"),
        ("Reusable IP",  "Which Kroger artifacts (manifests, schemas, demos) port without rework"),
        ("Independence routing","Per-account audit-firm relationship lookup"),
    ]
    for i,(l,r) in enumerate(rows, start=1):
        c1 = ws.cell(row=i, column=1, value=l); c2 = ws.cell(row=i, column=2, value=r)
        if i == 1: c1.font = Font(bold=True, size=14); c2.font = Font(italic=True, size=10, color="4A4A55")
        elif l and r: c1.font = Font(bold=True, size=10); c2.font = CELL
        else: c1.font = Font(bold=True, size=11, color="6648B0"); c2.font = CELL
        c2.alignment = Alignment(wrap_text=True, vertical="top"); ws.row_dimensions[i].height = 28

    # Account Comparison
    ws = wb.create_sheet("Account Comparison")
    headers = ["Dimension","Kroger","Walmart (Grocery)","Albertsons-Safeway","Costco Wholesale","Target (Grocery)","Ahold Delhaize US","Publix"]
    widths = [28] + [22]*7
    for j,h in enumerate(headers, start=1):
        c = ws.cell(row=1, column=j, value=h); c.font=HDR_FONT; c.fill=HDR_FILL; c.alignment=HDR_AL; c.border=BOR
        ws.column_dimensions[get_column_letter(j)].width = widths[j-1]
    ws.row_dimensions[1].height = 30; ws.freeze_panes = "B2"

    DIMS = [
      ("Annual revenue","~$150B","~$280B (US grocery share)","~$80B","~$255B (overall)","~$60B (food/grocery)","~$55B (US)","~$60B"),
      ("Store count","~2,800","~4,700 (US Supercenters)","~2,300","~600 (US warehouses)","~1,950","~2,050","~1,400"),
      ("Banner mix","Multi-banner: Kroger, Ralphs, Fred Meyer, KS, Smith's, FRY, HT, Mariano's, QFC","Single brand (Walmart Supercenter)","Multi-banner: Albertsons, Safeway, Vons, Tom Thumb, ACME, Shaw's","Single brand (Costco)","Single brand (Target)","Multi-banner: Stop&Shop, Food Lion, Giant, Hannaford","Single brand (Publix)"),
      ("Fresh-food revenue %","~30% (~$45B)","~25% (~$70B in grocery share)","~32% (~$26B)","~25% (~$64B)","~28% (~$17B)","~35% (~$19B)","~33% (~$20B)"),
      ("Cold-chain shrink exposure (est)","~$1.1B","~$1.7B","~$650M","~$1.6B","~$425M","~$475M","~$500M"),
      ("FSMA 204 readiness","Mid · multi-vendor lot-trace","Low-Mid · scale challenge","Mid","Low-Mid · supplier-driven","Low-Mid","Mid","Low"),
      ("Pricing platform incumbent","Revionics + 84.51°","In-house · Element","Eversight (Instacart)","In-house","In-house","Symphony / RetailX","In-house"),
      ("Loyalty platform","Kroger Plus + 84.51° (~60M HH)","Walmart+","Albertsons for U","No public loyalty","Target Circle","Each-banner specific","Publix Pharmacy + Club"),
      ("Lot-trace platform","Trace One + FoodLogiQ","In-house","FoodLogiQ","In-house + supplier","Mixed","Trace One","Mixed"),
      ("84.51°-equivalent data analytics","84.51° (in-house subsidiary)","Walmart Connect (separate)","No equivalent","No equivalent","Roundel","No equivalent","No equivalent"),
      ("Microsoft footprint","Substantial · Azure + Fabric strategic","Limited · multi-cloud","Mixed","Limited · multi-cloud","Substantial","Mixed","Limited"),
      ("Audit firm (verify currency)","PwC","Ernst & Young","Deloitte","KPMG","Ernst & Young","KPMG","Ernst & Young"),
      ("Independence-office routing","Required","Required if audit overlap","CRITICAL · Deloitte audit client","Required","Required","Required","Required"),
      ("Deloitte engagement history","Strong DMTSP relationship","Limited","CRITICAL audit relationship","Mixed","Strong","Mixed","Limited"),
      ("Pursuit priority ranking","#1 (anchor)","#3","#7 (audit constraint)","#4","#2","#5","#6"),
    ]
    for i, d in enumerate(DIMS, start=2):
        for j, v in enumerate(d, start=1):
            c = ws.cell(row=i, column=j, value=v); c.border=BOR; c.alignment=CELL_AL
            c.font = CELL
            if j == 1: c.font = Font(bold=True, size=10)
            if i % 2 == 0: c.fill = BAND_FILL
        ws.row_dimensions[i].height = 38

    # Pursuit Prioritization
    ws = wb.create_sheet("Pursuit Prioritization")
    headers = ["Rank","Account","Why this rank","Recommended Wave-1 anchor","Reusable from Kroger","Sequencing"]
    widths = [6, 22, 38, 26, 30, 22]
    for j,h in enumerate(headers, start=1):
        c = ws.cell(row=1, column=j, value=h); c.font=HDR_FONT; c.fill=HDR_FILL; c.alignment=HDR_AL; c.border=BOR
        ws.column_dimensions[get_column_letter(j)].width = widths[j-1]
    ws.row_dimensions[1].height = 30
    PURSUITS = [
      (1,"Kroger","Strongest DMTSP relationship · multi-banner footprint · 84.51° pricing DNA · willing reference client",
       "RC-E2E-03 Cincinnati banner · ~50 stores",
       "Reference deployment · all artifacts apply directly",
       "Active — Wave 1"),
      (2,"Target","Strong Microsoft footprint · grocery growth-segment · pricing/markdown decisioning is high-priority",
       "RC-E2E-03 in NA-anchor cluster · ~100 stores",
       "Manifests + schemas reusable · banner-mix simpler · personas adapt",
       "FY27 Q3 (post-Kroger W1 exit)"),
      (3,"Walmart (Grocery)","Largest US grocery scale · multi-cloud preference complicates · scale advantage if won",
       "RC-E2E-09 Product Tracking (FSMA 204) — most-defensible angle",
       "RC-E2E-09 manifests directly · MERML/SCML schemas align",
       "FY27 Q4 — independence + scale risk"),
      (4,"Costco","Single-brand · supplier-driven SKU model · cold-chain exposure on perishables",
       "RC-E2E-03 in tier-1 warehouses",
       "Adapt elasticity calibration to membership model · simpler banner",
       "FY28 H1"),
      (5,"Ahold Delhaize US","Multi-banner · European parent · digital-first orientation",
       "RC-E2E-03 + RC-E2E-09 on Stop & Shop banner",
       "Multi-banner pattern from Kroger ports well · personas adapt",
       "FY28 H1"),
      (6,"Publix","Limited Microsoft footprint · single-brand · strong food-safety culture (asset)",
       "RC-E2E-09 (FSMA 204) as wedge",
       "FSMA 204 manifests + evidence packet · partial reuse",
       "FY28 H2 — needs MSFT-platform investment"),
      (7,"Albertsons-Safeway","CRITICAL: Deloitte external audit relationship · MUST route through Independence office FIRST",
       "Independence-office determination required before any pursuit conversation",
       "Pattern reuse possible IF independence cleared",
       "On hold pending independence-office determination"),
    ]
    for i, p in enumerate(PURSUITS, start=2):
        for j, v in enumerate(p, start=1):
            c = ws.cell(row=i, column=j, value=v); c.border=BOR; c.alignment=CELL_AL; c.font=CELL
            if j == 1: c.alignment = CTR; c.font = Font(bold=True, size=12, color="6648B0")
            if j == 2: c.font = Font(bold=True, size=10)
            if i % 2 == 0: c.fill = BAND_FILL
        # Color rank
        if i == 2: ws.cell(row=i, column=1).fill = PatternFill("solid", fgColor="C9E5DC")
        elif i in (3,4): ws.cell(row=i, column=1).fill = PatternFill("solid", fgColor="EFF7F5")
        elif i in (5,6): ws.cell(row=i, column=1).fill = PatternFill("solid", fgColor="FBF1DC")
        elif i == 8: ws.cell(row=i, column=1).fill = PatternFill("solid", fgColor="F8D5DA")
        ws.row_dimensions[i].height = 60

    # RC-E2E-03 Lever Sizing
    ws = wb.create_sheet("E2E-03 Lever Sizing")
    headers = ["Account","Stores","Events/wk/banner","$ avoided/event","Reduction rate","Annualized W3 shrink avoided"]
    widths = [22, 10, 16, 16, 14, 26]
    for j,h in enumerate(headers, start=1):
        c = ws.cell(row=1, column=j, value=h); c.font=HDR_FONT; c.fill=HDR_FILL; c.alignment=HDR_AL; c.border=BOR
        ws.column_dimensions[get_column_letter(j)].width = widths[j-1]
    ws.row_dimensions[1].height = 30
    SIZING = [
      ("Kroger",      2800, 85, "$1,313", "18%", "~$200M"),
      ("Walmart (Grocery)",4700,90,"$1,450","16%","~$330M"),
      ("Albertsons-Safeway",2300,80,"$1,200","18%","~$140M"),
      ("Costco",     600,150,"$2,800","18%","~$95M"),
      ("Target",      1950,55,"$1,180","16%","~$70M"),
      ("Ahold Delhaize",2050,75,"$1,250","18%","~$110M"),
      ("Publix",     1400,80,"$1,300","18%","~$95M"),
    ]
    for i, s in enumerate(SIZING, start=2):
        for j, v in enumerate(s, start=1):
            c = ws.cell(row=i, column=j, value=v); c.border=BOR; c.alignment=CELL_AL; c.font=CELL
            if j == 1: c.font = Font(bold=True, size=10)
            if j in (2,3,4,5,6): c.alignment = CTR
            if j == 6: c.font = Font(bold=True, size=10, color=GOLD)
            if i % 2 == 0: c.fill = BAND_FILL

    # RC-E2E-09 FSMA 204 Readiness
    ws = wb.create_sheet("E2E-09 FSMA Readiness")
    headers = ["Account","Lot-trace platform","TLC capture","KDE coverage","Recall response","FDA-response time","Overall readiness"]
    widths = [22, 22, 14, 14, 16, 18, 18]
    for j,h in enumerate(headers, start=1):
        c = ws.cell(row=1, column=j, value=h); c.font=HDR_FONT; c.fill=HDR_FILL; c.alignment=HDR_AL; c.border=BOR
        ws.column_dimensions[get_column_letter(j)].width = widths[j-1]
    ws.row_dimensions[1].height = 30
    READINESS = [
      ("Kroger",         "Trace One + FoodLogiQ", "Mid", "Mid", "1-3 days", "1-3 weeks", "Mid"),
      ("Walmart (Grocery)","In-house",            "Low-Mid", "Low-Mid", "Mid",     "1 week",   "Low-Mid"),
      ("Albertsons-Safeway","FoodLogiQ",          "Mid", "Mid", "1-3 days", "1-2 weeks","Mid"),
      ("Costco",         "Supplier-driven",       "Low", "Low", "Days",    "2 weeks",  "Low-Mid"),
      ("Target",         "Mixed",                 "Low-Mid","Low","Mid",   "2 weeks",  "Low-Mid"),
      ("Ahold Delhaize", "Trace One",             "Mid", "Mid", "Days",    "1 week",   "Mid"),
      ("Publix",         "Mixed",                 "Low", "Low", "Days",    "2-3 weeks","Low"),
    ]
    for i, r in enumerate(READINESS, start=2):
        for j, v in enumerate(r, start=1):
            c = ws.cell(row=i, column=j, value=v); c.border=BOR; c.alignment=CELL_AL; c.font=CELL
            if j == 1: c.font = Font(bold=True, size=10)
            if j in (3,4,5,6,7): c.alignment = CTR
            if i % 2 == 0: c.fill = BAND_FILL
            # Color readiness
            if j == 7:
                if v == "Mid": c.fill = PatternFill("solid", fgColor="EFF7F5")
                elif v == "Low-Mid": c.fill = PatternFill("solid", fgColor="FBF1DC")
                elif v == "Low": c.fill = PatternFill("solid", fgColor="F8D5DA")

    # Strategic Notes
    ws = wb.create_sheet("Strategic Notes")
    headers = ["Account","Strategic notes"]
    widths = [22, 80]
    for j,h in enumerate(headers, start=1):
        c = ws.cell(row=1, column=j, value=h); c.font=HDR_FONT; c.fill=HDR_FILL; c.alignment=HDR_AL; c.border=BOR
        ws.column_dimensions[get_column_letter(j)].width = widths[j-1]
    ws.row_dimensions[1].height = 30
    NOTES = [
      ("Kroger","Reference engagement in flight. Cincinnati banner cluster. 84.51° relationship is unique advantage. Independence routing proceeded; PwC audit relationship documented. Multi-banner pattern is the key reusable IP."),
      ("Walmart (Grocery)","Largest US grocery scale with $280B+ exposure. Multi-cloud preference (AWS-leaning) means MSFT-native pitch needs platform-justification angle. Scale + EY audit creates sizing-vs-independence tradeoff. Lead with FSMA 204 angle (RC-E2E-09); E&Y not Deloitte audit so engagement structure is straightforward."),
      ("Albertsons-Safeway","CRITICAL: Deloitte external audit relationship. Independence-office determination required BEFORE any pursuit conversation. Engagement structure must preserve audit independence. Hold pursuit until independence office signs off."),
      ("Costco","Membership model changes elasticity calibration. Single-brand simplifies banner-mix vs Kroger. KPMG audit relationship straightforward. Supplier-driven SKU model is unique constraint. Wave-1 anchor: tier-1 warehouses with high fresh mix."),
      ("Target","Strong Microsoft footprint = lower platform-investment friction. Grocery is growth segment. Roundel ad-platform provides 84.51°-equivalent data backbone. NA-anchor cluster Wave-1. Likely the most-clean execution after Kroger."),
      ("Ahold Delhaize US","Multi-banner footprint similar to Kroger pattern; reuses Cincinnati methodology. European parent's data-residency requirements may add complexity. Stop & Shop is best candidate banner. KPMG audit straightforward."),
      ("Publix","Strong food-safety culture is asset for FSMA 204 angle (RC-E2E-09 wedge). Limited Microsoft footprint requires platform-investment justification. Single-brand simpler. Long sales cycle expected. FY28 H2 timing realistic."),
    ]
    for i, n in enumerate(NOTES, start=2):
        for j, v in enumerate(n, start=1):
            c = ws.cell(row=i, column=j, value=v); c.border=BOR; c.alignment=CELL_AL; c.font=CELL
            if j == 1: c.font = Font(bold=True, size=10)
            if i % 2 == 0: c.fill = BAND_FILL
        ws.row_dimensions[i].height = 60

    target = OUT / "APEX-RC-E2E-03-Cross-Grocer-Comparison.xlsx"
    try: wb.save(str(target))
    except PermissionError:
        n=2
        while True:
            alt = target.with_name(target.stem + f"-v{n}" + target.suffix)
            if not alt.exists(): target=alt; break
            n+=1
        wb.save(str(target))
    print(f"  wrote {target.name}")


# =========================================================================
# 12. AI/ML MODEL SPEC (DOCX)
# =========================================================================
def build_ml_model_spec():
    print("\nBuilding AI/ML Model Spec docx...")
    doc = Document()
    for s in doc.sections:
        s.page_width = Inches(8.5); s.page_height = Inches(11)
        s.top_margin = s.bottom_margin = Inches(0.85); s.left_margin = s.right_margin = Inches(0.95)
    doc.styles["Normal"].font.name = "Arial"; doc.styles["Normal"].font.size = Pt(11)

    addeyebrow(doc, "APEX · RC-E2E-03 + RC-E2E-09 · AI/ML MODEL SPECIFICATION")
    p = doc.add_paragraph(); r = p.add_run("AI/ML Model Specification"); setrun(r, bold=True, size=24, color=INK)
    p2 = doc.add_paragraph(); r2 = p2.add_run("Three production models: Risk-Score · Elasticity · Disposition Classification")
    setrun(r2, bold=True, size=14, color=INK2, italic=True)
    doc.add_paragraph()
    addbody(doc, "Specifies the three production ML models that anchor RC-E2E-03 + RC-E2E-09 agent decisioning. "
                 "For each: features, training data, evaluation metrics, drift detection, retraining cadence, "
                 "and bias-evaluation discipline. Audience: ML Engineering, Engagement Architect, Privacy Office, Internal Audit.",
            italic=True, color=INK2)
    doc.add_paragraph()
    addkv(doc, [
        ("Document type",    "AI/ML Model Specification"),
        ("Models specified", "(1) Risk-Score · (2) Elasticity · (3) Disposition Classification"),
        ("Services",         "RC-E2E-03 Assortment & Pricing · RC-E2E-09 Product Tracking"),
        ("Audience",         "ML Engineering · Engagement Architect · Privacy · Internal Audit · External Audit"),
        ("Reference client", "The Kroger Co."),
        ("Version / date",   "v1.0 · 27 Apr 2026"),
    ])
    pb(doc)

    # Section 1: Overview
    addh(doc, "1 · Overview · three models, three roles", level=1)
    addbody(doc, "RC-E2E-03 + RC-E2E-09 require three production ML models. Each has a distinct purpose, "
                 "training cadence, and accountability surface.")
    addtable(doc, ["Model","Purpose","Used by agent","Update cadence"],
             [["Risk-Score","Per-fixture excursion-risk score · drives trigger sensitivity tuning","Assess + Classify","Weekly"],
              ["Elasticity","Per-SKU price-elasticity curves · drives markdown-recommendation_pct","Quantify","Weekly"],
              ["Disposition Classification","Per-lot disposition (sell-through/markdown/destroy/hold)","Classify","Bi-weekly"]],
             col_widths=[2.0,2.4,1.6,1.2])

    addcallout(doc, "Authoritative source",
               "All three models live in the CAF substrate via AEF LLM Gateway + ML model serving. "
               "APEX manifests reference them by capability tag (cap:risk-score, cap:elasticity-curve, cap:disposition-classification) "
               "rather than by specific model URL. This is what makes the engagement portable across reference clients.")

    pb(doc)

    # Section 2: Risk-Score Model
    addh(doc, "2 · Model 1 · Risk-Score", level=1)
    addh(doc, "2.1 · Purpose &amp; semantics", level=2)
    addbody(doc, "Per-fixture excursion-risk score (0.0-1.0) representing the probability of a "
                 "compliance-breaching excursion within a forward-looking window. Used by Assess agent "
                 "to prioritize triage urgency, and by Classify agent as a confidence calibrator.")
    addh(doc, "2.2 · Features", level=2)
    addtable(doc, ["Feature category","Specific features","Source"],
             [["Telemetry recent","Last 24h temperature mean/std/min/max · door-open count · setpoint distance",     "BRZ.RC.REFRIG_TELEMETRY"],
              ["Telemetry historical","30-day moving avg · 90-day excursion frequency · seasonal-adjusted",           "BRZ.RC.REFRIG_TELEMETRY"],
              ["Fixture state",     "Fixture age · recent service-ticket count · compressor-runtime trend",          "BRZ.RC.REFRIG_SERVICE + DIM.Fixture"],
              ["Banner / store",    "Store traffic volume · SKU density per fixture · banner-tier",                  "DIM.Store"],
              ["Item mix",          "FTL-flagged item count in fixture · category mix (dairy/deli/produce)",         "BRZ.RC.ITEM_MASTER"],
              ["Time / seasonality","Day of week · hour of day · season · weather context (where available)",        "Derived + external"]],
             col_widths=[1.7,3.0,2.0])

    addh(doc, "2.3 · Training data", level=2)
    addbullets(doc, [
        "Historical excursion events from past 4 quarters of BRZ.RC.REFRIG_TELEMETRY + BRZ.RC.FOOD_SAFETY",
        "Labeled outcomes from FOOD_SAFETY.severity (NORMAL / WATCH / CRITICAL) and disposition (SELL_THROUGH / MARKDOWN / DESTROY)",
        "Roughly 4M training examples per banner-quarter at full Kroger enterprise scale",
        "Class imbalance: ~92% NORMAL · ~6% WATCH · ~2% CRITICAL — handled via stratified sampling + focal loss",
    ])
    addh(doc, "2.4 · Model architecture", level=2)
    addbody(doc, "Gradient-boosted ensemble (XGBoost or LightGBM) for primary scoring; LSTM for temporal "
                 "feature extraction over the 24h window. Architecture choices are CAF-substrate decisions; "
                 "APEX manifest contracts on capability not architecture.")
    addh(doc, "2.5 · Evaluation metrics", level=2)
    addtable(doc, ["Metric","W2 target","W3 target"],
             [["Precision @ risk-threshold (0.85)","≥ 0.78","≥ 0.85"],
              ["Recall @ risk-threshold (0.85)","≥ 0.72","≥ 0.82"],
              ["AUC-ROC","≥ 0.91","≥ 0.95"],
              ["Brier score (calibration)","≤ 0.08","≤ 0.05"],
              ["Per-banner performance gap","< 8 pts","< 3 pts"],
              ["Per-fixture-class consistency","Within ±5%","Within ±2%"]],
             col_widths=[3.4,1.4,1.4])

    addh(doc, "2.6 · Drift detection &amp; retraining", level=2)
    addbullets(doc, [
        "Cohort Drift Watchtower (UC-22) monitors weekly performance per cohort family",
        "Population Stability Index (PSI) computed on feature distribution vs. baseline · alert at PSI > 0.1",
        "Performance-drift trigger: precision drop > 5pp from 7-day moving average",
        "Auto-retrain triggered with HITL approval gate (Audience Lead) before deployment",
        "Champion-challenger A/B: new model shadows live for 14 days before promotion",
    ])

    addh(doc, "2.7 · Bias evaluation", level=2)
    addbullets(doc, [
        "Per-banner performance breakdown reviewed weekly — no banner more than 8pp behind enterprise mean (W2) / 3pp (W3)",
        "Per-fixture-class breakdown — no class more than 5% performance gap",
        "Demographic-segment review: per-store-tier and per-region — no systematic underperformance",
        "Fairness audit at quarterly model deployment · Privacy Office sign-off",
    ])

    pb(doc)

    # Section 3: Elasticity Model
    addh(doc, "3 · Model 2 · Elasticity", level=1)
    addh(doc, "3.1 · Purpose &amp; semantics", level=2)
    addbody(doc, "Per-SKU price-elasticity curves used by Quantify agent to compute optimal markdown_recommendation_pct. "
                 "Curve form: log-log demand response with promo-period adjustment. Output: elasticity coefficient + "
                 "expected unit-volume response per markdown percentage.")
    addh(doc, "3.2 · Features", level=2)
    addtable(doc, ["Feature category","Specific features","Source"],
             [["Historical sales","52-week unit volume per SKU · price-point histogram · promo-period flag",        "BRZ.RC.POS_TXN + BRZ.RC.PRICING"],
              ["Promotional context","Past TPR/markdown response · ad-block exposure · loyalty-tier overlay",       "BRZ.RC.PROMO_MARKDOWN + DIM.Household"],
              ["Cross-elasticity",  "Substitutable-SKU response · category-cannibalization signals",                "BRZ.RC.POS_TXN basket-level"],
              ["Seasonality / shelf-life","Time-to-expiration · seasonal demand · perishable-shelf-life pressure",  "BRZ.RC.LOT_TRACE + BRZ.RC.ITEM_MASTER"],
              ["Loyalty cohort",    "Per-cohort price-sensitivity index from 84.51° segment data",                  "BRZ.RC.LOYALTY_HH"],
              ["Banner / region",   "Banner-specific elasticity baseline · regional-economic factor",               "DIM.Store"]],
             col_widths=[1.8,3.0,1.9])

    addh(doc, "3.3 · Training data", level=2)
    addbullets(doc, [
        "Past 52-week BRZ.RC.POS_TXN price × demand observations per SKU",
        "Per-SKU elasticity computed via Bayesian hierarchical model with category-level priors",
        "Approximately 1.2M SKU-week observations across the active Kroger SKU universe",
        "Cold-start handling for new SKUs: borrow priors from category + brand level",
    ])
    addh(doc, "3.4 · Model architecture", level=2)
    addbody(doc, "Bayesian hierarchical regression with category-level pooling for cold-start. Per-SKU posterior "
                 "elasticity coefficients with credible intervals exposed to Quantify agent. Curve refit weekly per SKU; "
                 "category-level pooling refit monthly.")

    addh(doc, "3.5 · Evaluation metrics", level=2)
    addtable(doc, ["Metric","W2 target","W3 target"],
             [["Mean absolute % error on 1-week-ahead unit forecast","≤ 14%","≤ 9%"],
              ["Per-category R² fit","≥ 0.72","≥ 0.85"],
              ["95% credible interval coverage","≥ 92%","≥ 95%"],
              ["Markdown-recommendation lift over flat baseline","+8% NPV","+15% NPV"],
              ["Per-banner consistency","Within ±10%","Within ±5%"]],
             col_widths=[3.4,1.4,1.4])

    addh(doc, "3.6 · Retraining cadence", level=2)
    addbullets(doc, [
        "Per-SKU posterior refit weekly (Sun 06:00 UTC)",
        "Category-level pooling refit monthly · validates category boundaries",
        "Promo-response calibration per ad-block cycle (every 2 weeks)",
        "Cold-start refit within 24 hours of new-SKU activation",
    ])

    pb(doc)

    # Section 4: Disposition Classification Model
    addh(doc, "4 · Model 3 · Disposition Classification", level=1)
    addh(doc, "4.1 · Purpose &amp; semantics", level=2)
    addbody(doc, "Per-lot 4-class classification (SELL_THROUGH / MARKDOWN / DESTROY / HOLD) consumed by Classify agent. "
                 "The most consequential model — drives ~$200M annualized value lever. Confidence calibration is critical "
                 "because confidence < 0.85 escalates to Food Safety Lead.")
    addh(doc, "4.2 · Features", level=2)
    addtable(doc, ["Feature category","Specific features","Source"],
             [["Excursion profile","Peak temp · duration · avg temp · door-open count · severity classification",     "Assess agent output"],
              ["Lot context",     "Lot age · time-to-expiration · supplier · FTL-category · prior-handling history",   "BRZ.RC.LOT_TRACE + BRZ.RC.RECEIPT"],
              ["FSMA 204 policy", "Applicable policy clause · severity threshold · category-specific rules",           "FSMA-204-policy-corpus"],
              ["Item attributes", "Temperature class · shelf-life-days · hazmat flag · private-label flag",            "BRZ.RC.ITEM_MASTER"],
              ["Historical outcomes","Past disposition for similar excursion + lot characteristics",                    "BRZ.RC.FOOD_SAFETY history"],
              ["Confidence calibration","Per-class probability + uncertainty estimate · escalation flag",              "Model output (calibrated)"]],
             col_widths=[1.8,3.0,1.9])

    addh(doc, "4.3 · Training data", level=2)
    addbullets(doc, [
        "Past 8 quarters of BRZ.RC.FOOD_SAFETY closeouts paired with excursion profiles",
        "Approximately 280K labeled training examples (95% of which are MARKDOWN or SELL_THROUGH)",
        "DESTROY class is rare (~3%) and highly consequential — handled via class-weighted loss + focal loss",
        "Augmented with policy-corpus synthetic examples for edge cases (FTL-category-specific severity thresholds)",
        "FSMA 204 policy clause embedding learned from corpus — grounds the policy_citation output",
    ])

    addh(doc, "4.4 · Model architecture", level=2)
    addbody(doc, "Multi-task transformer encoder (per-feature-category embedding) with calibrated classification head. "
                 "Calibration via temperature scaling + Platt scaling for confidence-threshold accuracy. "
                 "Policy_citation generated via retrieval-augmented head over the policy corpus.")

    addh(doc, "4.5 · Evaluation metrics", level=2)
    addtable(doc, ["Metric","W2 target","W3 target"],
             [["Multi-class accuracy (4 classes)","≥ 0.88","≥ 0.95"],
              ["Per-class precision (DESTROY)","≥ 0.92","≥ 0.97"],
              ["Per-class recall (DESTROY)","≥ 0.85","≥ 0.92"],
              ["Confidence calibration (ECE)","≤ 0.04","≤ 0.02"],
              ["Policy-citation accuracy","≥ 0.91","≥ 0.96"],
              ["Sell-through-at-risk rate (false-positive on SELL_THROUGH)","≤ 1.5%","≤ 0.5%"]],
             col_widths=[3.4,1.4,1.4])

    addh(doc, "4.6 · Confidence threshold tuning (UC-07)", level=2)
    addbody(doc, "The 0.85 confidence threshold for HITL escalation is calibrated based on the cost of "
                 "false-positive (un-needed escalation) vs. false-negative (auto-disposition with low confidence). "
                 "Calibration plot reviewed by Food Safety Lead; rebalanced quarterly. Below 0.85, escalation routes "
                 "to Food Safety Compliance Lead automatically.")

    addh(doc, "4.7 · Drift detection &amp; bias", level=2)
    addbullets(doc, [
        "Per-FTL-category accuracy tracked weekly · soft-cheese, leafy-greens, mollusks segregated for separate scrutiny",
        "Per-banner performance breakdown — no banner more than 5pp behind enterprise mean",
        "Per-supplier accuracy breakdown — flag suppliers with systematically misclassified lots for vendor review",
        "Quarterly bias-evaluation suite · Privacy Office + Food Safety Lead sign-off before production deployment",
        "Champion-challenger A/B: new model shadows live for 30 days · Food Safety Lead approves promotion",
    ])

    pb(doc)

    # Section 5: Cross-cutting concerns
    addh(doc, "5 · Cross-cutting model governance", level=1)
    addh(doc, "5.1 · Model registry &amp; lineage", level=2)
    addbullets(doc, [
        "All three models registered in Microsoft Purview with version history",
        "Lineage from training data (BRZ.RC.* tables) through model artifact to inference endpoint",
        "Per-prediction audit trail — every Classify agent inference logs (model_version, feature_hash, confidence, output)",
        "LangFuse tracing on every inference call — joinable to APEX trace_id",
    ])
    addh(doc, "5.2 · Privacy &amp; consent", level=2)
    addbullets(doc, [
        "Training data uses hashed loyalty_household_id only · no clear-text customer PII in any model feature",
        "Consent state from BRZ.RC.LOYALTY_HH respected at inference: WITHDRAWN consent excluded from elasticity feature aggregation",
        "Per-decision evidence-of-access logged in Purview · queryable by Privacy Office on inquiry",
        "Model retraining excludes WITHDRAWN-consent households from training data",
    ])

    addh(doc, "5.3 · Independence &amp; auditability", level=2)
    addcallout(doc, "Model artifacts are Kroger artifacts",
               "Trained model weights, training data lineage, and policy-corpus embeddings are Kroger artifacts. "
               "Deloitte's role is advisory and implementation enablement. Model governance (training cadence, "
               "deployment approval, drift triage) flows through Kroger ML Engineering with Audience Lead and "
               "Food Safety Lead as approval persona. This preserves audit independence and aligns with the "
               "Independence-office routing of the engagement.")

    addh(doc, "5.4 · Operational SLOs", level=2)
    addtable(doc, ["Model","Inference latency (p99)","Throughput (peak)","Cold-start fallback"],
             [["Risk-Score","≤ 50ms","~30K predictions/min","Heuristic-based scoring for new fixtures"],
              ["Elasticity","≤ 100ms","~5K predictions/min","Category-level prior posteriors"],
              ["Disposition Classification","≤ 200ms","~50K predictions/min","Policy-corpus rule-based fallback"]],
             col_widths=[2.6,1.8,1.8,2.0])

    pb(doc)

    addh(doc, "Appendix A · Cross-references", level=1)
    addtable(doc, ["Artifact","Purpose"],
             [["APEX-RC-E2E-03-Walkthrough.docx","Service walkthrough with agent definitions"],
              ["APEX-RC-E2E-03-09-SOR-Bronze-to-Silver.docx","Source-of-record schemas for training data"],
              ["APEX-RC-E2E-03-Use-Case-Catalog.xlsx","UC-22 Cohort Drift Watchtower"],
              ["APEX-FSMA-204-Compliance-Checklist.docx","Policy-corpus reference"],
              ["manifests/classify-agent.yaml","Capability tag references"],
              ["APEX-RC-E2E-03-Kroger-Risk-Register.xlsx","Risks R-06, R-07 — model drift + bias"]],
             col_widths=[3.4,3.0])

    addbody(doc, "Independence posture preserved: 'Deloitte's Microsoft practice,' 'DMTSP,' "
                 "'Microsoft platform capabilities,' 'Microsoft-native deployment.' The terms 'partner,' "
                 "'alliance,' and 'strategic alliance' do not appear.",
            italic=True, color=MUTE, size=9)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— End of AI/ML Model Specification —"); setrun(r, italic=True, size=10, color=MUTE)
    target = OUT / "APEX-RC-E2E-03-09-AI-ML-Model-Spec.docx"
    try: doc.save(str(target))
    except PermissionError:
        n=2
        while True:
            alt = target.with_name(target.stem + f"-v{n}" + target.suffix)
            if not alt.exists(): target = alt; break
            n+=1
        doc.save(str(target))
    print(f"  wrote {target.name}")


# =========================================================================
# 13. PRIVACY & DATA GOVERNANCE SPEC (DOCX)
# =========================================================================
def build_privacy_spec():
    print("\nBuilding Privacy & Data Governance Spec docx...")
    doc = Document()
    for s in doc.sections:
        s.page_width = Inches(8.5); s.page_height = Inches(11)
        s.top_margin = s.bottom_margin = Inches(0.85); s.left_margin = s.right_margin = Inches(0.95)
    doc.styles["Normal"].font.name = "Arial"; doc.styles["Normal"].font.size = Pt(11)

    addeyebrow(doc, "APEX · RC-E2E-03 + RC-E2E-09 · PRIVACY &amp; DATA GOVERNANCE")
    p = doc.add_paragraph(); r = p.add_run("Privacy &amp; Data Governance Specification"); setrun(r, bold=True, size=24, color=INK)
    p2 = doc.add_paragraph(); r2 = p2.add_run("Consent state machine · data residency · RBAC · Purview · retention · FSMA 204 governance overlay")
    setrun(r2, bold=True, size=14, color=INK2, italic=True)
    doc.add_paragraph()
    addbody(doc, "Privacy and data-governance specification for RC-E2E-03 + RC-E2E-09 deployment at Kroger. "
                 "Goes deeper than SOR §6.3. Targets the Privacy Office, CISO, Internal Audit, and external-audit "
                 "liaisons. Establishes the consent state machine, RBAC discipline, Purview tagging strategy, "
                 "retention policy, and FSMA 204 governance overlay.",
            italic=True, color=INK2)
    doc.add_paragraph()
    addkv(doc, [
        ("Document type",    "Privacy &amp; Data Governance Specification"),
        ("Services",         "RC-E2E-03 · RC-E2E-09"),
        ("Reference client", "The Kroger Co."),
        ("Audience",         "Privacy Office · CISO · Internal Audit · External Audit · Data Engineering"),
        ("Regulatory frame", "FSMA 204 · GDPR (where applicable) · CCPA · state-level privacy laws · Kroger Privacy Policy"),
        ("Companion",        "APEX-RC-E2E-03-09-SOR-Bronze-to-Silver.docx · APEX-RC-E2E-03-09-AI-ML-Model-Spec.docx"),
        ("Version / date",   "v1.0 · 27 Apr 2026"),
    ])
    pb(doc)

    addh(doc, "1 · Consent state machine", level=1)
    addh(doc, "1.1 · Consent states", level=2)
    addbody(doc, "Every loyalty household has a consent state captured in BRZ.RC.LOYALTY_HH.consent_state. "
                 "The state machine drives downstream data-use eligibility.")
    addtable(doc, ["State","Meaning","Personalization eligible?","Data retention policy"],
             [["FULL","Full consent · all marketing + analytics use cases","Yes, full","Standard 24-month retention"],
              ["RESTRICTED","Limited consent · analytics only · no targeted marketing","Limited (analytics only)","Standard retention; redacted use"],
              ["WITHDRAWN","Consent withdrawn · cease-personalization request","No","30-day deletion of personalization-relevant data; transactional record retained per FSMA"],
              ["CHILD","Account is minor's profile · COPPA-compliant","No personalization","Strict retention; no profile-building"],
              ["UNKNOWN","No consent state captured (legacy account)","No","Treated as RESTRICTED until consent established"]],
             col_widths=[1.3,2.2,1.8,2.0])

    addh(doc, "1.2 · State transitions", level=2)
    addbullets(doc, [
        "FULL → RESTRICTED → WITHDRAWN: triggered by household opt-out via Plus app, web, or call-center",
        "WITHDRAWN → FULL: requires explicit re-opt-in; cannot auto-revert",
        "UNKNOWN → RESTRICTED: default at Plus account creation if consent not captured",
        "Any → CHILD: triggered by COPPA-validation flag · cannot revert without identity verification",
    ])

    addh(doc, "1.3 · State respect at decision time", level=2)
    addbody(doc, "Every Quantify agent inference checks consent_state before incorporating loyalty-derived features. "
                 "Per-decision evidence-of-access (which household_id was accessed at what time) logs to Purview.")
    addcallout(doc, "Consent at decision time, not training time",
               "Training data uses 90-day-prior consent state · inference uses real-time consent state. "
               "If a household withdraws consent on Day N, Day-N+1 elasticity calibration uses 90-day-prior data "
               "(legal under most regimes), but Day-N+1 personalization decisions exclude that household.")

    pb(doc)

    addh(doc, "2 · Data residency &amp; sovereignty", level=1)
    addbody(doc, "All RC-E2E-03 + RC-E2E-09 data resides in US-East and US-West Azure regions only. No cross-border "
                 "data movement except for backup/DR which stays within US sovereign cloud boundaries.")
    addtable(doc, ["Data category","Residency","Backup region","Notes"],
             [["BRZ.RC.LOYALTY_HH","US-East primary","US-West","Hashed identifiers only · no clear-text PII"],
              ["BRZ.RC.POS_TXN","US-East primary","US-West","Transaction-level; loyalty-linked rows carry hashed hh_id only"],
              ["BRZ.RC.LOT_TRACE","US-East primary","US-West","Supplier-confidential · access-controlled in Silver"],
              ["BRZ.RC.FOOD_SAFETY","US-East primary","US-West","FSMA 204 evidence · FDA-accessible · long-retention"],
              ["BRZ.RC.REFRIG_TELEMETRY","US-East primary","US-West","Operational data · 90-day hot tier"],
              ["LEDGER audit-row","US-East primary","US-West","Quarterly export to FDA Relations packet"],
              ["Model artifacts","US-East","US-West","Versioned in Purview · Kroger-owned"]],
             col_widths=[2.0,1.3,1.3,2.5])

    pb(doc)

    addh(doc, "3 · RBAC discipline", level=1)
    addh(doc, "3.1 · Role-to-data-zone mapping", level=2)
    addtable(doc, ["Role","Bronze access","Silver access","Gold (Power BI)","LEDGER"],
             [["APEX Engineer (dev)","Read · all tables","Read/write","Read","Read"],
              ["APEX Engineer (prod)","Read · all tables","Read · write via PR","Read","Read"],
              ["APEX Architect","Read · all tables","Read · write","Read","Read"],
              ["Operator","No","Read · own banner","Read · own banner","No"],
              ["Auditor / Privacy","No","Read · all (read-only)","Read · all","Read · all"],
              ["FDA-Relations","No","Read · LOT_TRACE + FOOD_SAFETY only","Read · FSMA dashboard","Read · LOT_TRACE + FOOD_SAFETY only"],
              ["Store Ops Lead","No","No","Read · own store cluster","No"],
              ["Category Manager","No","No","Read · own category","No"],
              ["External Auditor","No","No","Read · packet-format export only","Read · packet-export only"]],
             col_widths=[1.7,1.6,1.6,1.5,1.4])

    addh(doc, "3.2 · Privileged access management (PAM)", level=2)
    addbullets(doc, [
        "Production-write access requires PAM activation · time-bounded · approval workflow",
        "Manifest-merge requires 2-reviewer approval · committed via Git workflow",
        "Policy-manifest changes require Pricing Manager + Engineering Architect approval",
        "Model deployment requires ML Engineering + Audience Lead approval",
        "All PAM activations logged to Purview · Privacy Office reviews quarterly",
    ])

    pb(doc)

    addh(doc, "4 · Microsoft Purview tagging strategy", level=1)
    addh(doc, "4.1 · Sensitivity labels", level=2)
    addtable(doc, ["Sensitivity label","Applied to","Rationale"],
             [["Public","DIM.Calendar · DIM.Banner static refs","Non-sensitive reference data"],
              ["Internal","BRZ.RC.PROMO_MARKDOWN · MERML.Promotion","Commercial-confidential pricing"],
              ["Confidential","BRZ.RC.PRICING · BRZ.RC.POS_TXN · MERML.* · DIM.Item","Commercial-sensitive · competitor-readable"],
              ["Restricted","BRZ.RC.LOYALTY_HH · DIM.Household · personalization-relevant aggregates","Privacy-relevant · consent-gated"],
              ["FSMA-204","BRZ.RC.LOT_TRACE · BRZ.RC.FOOD_SAFETY · LEDGER FSMA rows","Regulator-accessible · 24-month retention minimum · supplier-confidential"],
              ["Audit-Critical","LEDGER audit-row · evidence-packet artifacts","Auditor-required · immutable · long-retention"]],
             col_widths=[1.5,3.0,2.0])

    addh(doc, "4.2 · Auto-classification rules", level=2)
    addbullets(doc, [
        "Hashed household_id auto-detected and labeled Restricted",
        "Pricing fields auto-detected and labeled Confidential",
        "Lot-trace fields auto-detected and labeled FSMA-204",
        "Audit-row fields auto-detected and labeled Audit-Critical",
        "Purview rules versioned · changes require Privacy Office approval",
    ])

    pb(doc)

    addh(doc, "5 · Retention policy", level=1)
    addtable(doc, ["Data category","Hot tier","Warm tier","Cold archive","Notes"],
             [["BRZ.RC.POS_TXN","30 days","13 months","7 years","SOX/audit retention drives cold-archive"],
              ["BRZ.RC.PRICING","90 days","Indefinite (SCD2)","—","Price-history queries common"],
              ["BRZ.RC.PROMO_MARKDOWN","90 days","13 months","7 years","Audit trail for promo decisions"],
              ["BRZ.RC.INVENTORY","30 days","13 months","—","Reconciliation needs"],
              ["BRZ.RC.REFRIG_TELEMETRY","7 days","90 days","Aggregate-only after 90d","Massive volume · aggregate retention"],
              ["BRZ.RC.REFRIG_SERVICE","90 days","5 years","—","Equipment-history data"],
              ["BRZ.RC.FOOD_SAFETY","180 days","Indefinite","Indefinite","FDA evidence retention · no destroy"],
              ["BRZ.RC.LOT_TRACE","180 days","Indefinite","Indefinite (FSMA 24-month minimum)","FSMA 204 retention drives indefinite"],
              ["BRZ.RC.VENDOR_MASTER","180 days","Indefinite","—","Vendor-history needs"],
              ["BRZ.RC.RECEIPT","30 days","5 years","—","Audit + recall lookback"],
              ["BRZ.RC.LOYALTY_HH","30 days","Indefinite (privacy-aware)","—","WITHDRAWN consent triggers redaction"],
              ["LEDGER audit-row","Indefinite","Indefinite","Indefinite","Cannot be deleted · audit pillar"]],
             col_widths=[1.9,1.0,1.4,1.4,1.6])

    addh(doc, "5.1 · WITHDRAWN-consent redaction", level=2)
    addbody(doc, "When a household transitions to WITHDRAWN consent state, the following data lifecycle events trigger:")
    addbullets(doc, [
        "Day 0: consent_state updated to WITHDRAWN in BRZ.RC.LOYALTY_HH · personalization decisions exclude household effective immediately",
        "Day 7: profile-building features redacted from Silver elasticity training data",
        "Day 30: BRZ.RC.LOYALTY_HH row marked _is_deleted = true · DIM.Household removes profile attributes",
        "Day 90: any model-feature aggregations excluding the household are recomputed and republished",
        "Transactional records (BRZ.RC.POS_TXN) retained per FSMA + financial-audit obligations · only loyalty linkage removed",
        "All redaction events logged to Purview · queryable by Privacy Office",
    ])

    pb(doc)

    addh(doc, "6 · FSMA 204 governance overlay", level=1)
    addbody(doc, "FSMA 204 lot-traceability data is supplier-confidential. Access controls in Silver and Gold limit "
                 "visibility to Food Safety / Compliance / Vendor-Mgmt roles · external auditors see export-format only.")
    addtable(doc, ["Constraint","Implementation"],
             [["FSMA records retained 24 months minimum","Indefinite cold archive in Delta Lake · time-travel preserved"],
              ["FDA-accessible within 24 hours of request","Auto-generated evidence packet · ≤ 4 hr at W2 · ≤ 60 min at W3"],
              ["Records protected against alteration","Bronze + Silver immutability · soft-delete only · Purview audit on access"],
              ["Supplier-confidential lot data","Vendor-Mgmt role-only access in Silver · external-audit sees export only"],
              ["Recall response · lot-pull within 2 hours","Auto-shelf-pull via Act agent · SCML.Lot recall-window query"],
              ["TLC namespace integrity","TLC-namespacing convention enforced · per-banner hash buckets"]],
             col_widths=[3.4,3.0])

    pb(doc)

    addh(doc, "7 · Audit log review cadence", level=1)
    addtable(doc, ["Cadence","Reviewer","What's reviewed"],
             [["Daily","Operations","Eventstream lag · DQ pack failures · PAM activations"],
              ["Weekly","Audience Lead + Pricing Manager","Cohort drift · model performance · concession utilization"],
              ["Bi-weekly","Engineering Architect","Bronze schema drift · Silver join correctness · LEDGER coverage"],
              ["Monthly","Privacy Office","Consent state changes · WITHDRAWN redaction events · Purview access logs"],
              ["Monthly","CISO + Engineering Architect","RBAC review · PAM activations · landing-zone hardening"],
              ["Quarterly","Internal Audit + Food Safety Lead","FSMA 204 evidence packet · auditor-walkthrough preparation"],
              ["Quarterly","Steering Committee","Risk register · KPI trajectory · Wave-progression gates"],
              ["Annual","External Audit","SOC-2-adjacent evidence · independence-office review"]],
             col_widths=[1.4,2.4,2.6])

    pb(doc)

    addh(doc, "8 · Privacy and audit incident response", level=1)
    addtable(doc, ["Trigger","Response"],
             [["Suspected data breach (any source)","CISO + Privacy Office activation · 24-hr containment · root-cause analysis · Audit Committee notification"],
              ["Consent-state desync detected","Personalization for affected households suspended · Privacy Office investigation · 30-day remediation SLA"],
              ["FDA inspection request","Food Safety Lead + FDA Relations + Engineering Architect · evidence packet generated within 4 hours · auditor walkthrough scheduled"],
              ["Recall event","Food Safety Lead activation · auto-shelf-pull · evidence chain validated · External communications coordinated"],
              ["Model bias detection","ML Engineering + Privacy Office · model rolled back · investigation · re-deployment with sign-off"],
              ["LEDGER write failure","Engineering Architect on-call · failover to backup write path · root-cause within 24 hr"]],
             col_widths=[2.8,3.6])

    pb(doc)

    addh(doc, "Appendix A · Cross-references", level=1)
    addtable(doc, ["Artifact","Purpose"],
             [["APEX-RC-E2E-03-09-SOR-Bronze-to-Silver.docx","Source-of-record specification"],
              ["APEX-RC-E2E-03-09-AI-ML-Model-Spec.docx","Model bias + privacy governance"],
              ["APEX-FSMA-204-Compliance-Checklist.docx","FSMA 204 final-rule mapping"],
              ["APEX-RC-E2E-03-Kroger-Risk-Register.xlsx","Risks R-07, R-18 — bias and consent desync"],
              ["APEX-RC-E2E-03-09-Fabric-Runbook.docx","Implementation runbook §6.3"]],
             col_widths=[3.4,3.0])

    addbody(doc, "Independence posture preserved across this document: 'Deloitte's Microsoft practice,' "
                 "'DMTSP,' 'Microsoft platform capabilities,' 'Microsoft-native deployment.' The terms "
                 "'partner,' 'alliance,' and 'strategic alliance' do not appear.",
            italic=True, color=MUTE, size=9)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— End of Privacy &amp; Data Governance Specification —"); setrun(r, italic=True, size=10, color=MUTE)
    target = OUT / "APEX-RC-E2E-03-09-Privacy-Data-Governance-Spec.docx"
    try: doc.save(str(target))
    except PermissionError:
        n=2
        while True:
            alt = target.with_name(target.stem + f"-v{n}" + target.suffix)
            if not alt.exists(): target = alt; break
            n+=1
        doc.save(str(target))
    print(f"  wrote {target.name}")


print("Writing Tier-4 artifacts...")
build_grocer_comparison()
build_ml_model_spec()
build_privacy_spec()
print("\nTier 4 complete.")
