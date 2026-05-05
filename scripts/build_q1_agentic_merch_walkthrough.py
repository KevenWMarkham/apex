"""Build the APEX Agentic-Merch Q1 walkthrough docx.

Modeled on:
    APEX-RC-E2E-03-Walkthrough.docx (Kroger cold-chain reference)

Sourced from:
    mvp-scope-q1.html              -- Q1 scope document
    mvp-scope-q1-journey-map.html  -- Persona journey map
    mvp-scope-q1-usecases.xlsx     -- 20 Q1 use cases
    APEX-RC-Scenario-Chains.xlsx   -- RC scenario library (anchor scenarios)

Output:
    C:\\Users\\kmarkham\\OneDrive - Deloitte (O365D)\\Documents\\APEX-Agentic-Merch-Q1-Walkthrough.docx
"""
from __future__ import annotations
import io, sys, os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

OUT = r"C:\Users\kmarkham\OneDrive - Deloitte (O365D)\Documents\APEX-Agentic-Merch-Q1-Walkthrough.docx"

INK_DARK = RGBColor(0x0D, 0x28, 0x18)
INK_GREEN = RGBColor(0x86, 0xBC, 0x25)
INK_GREY = RGBColor(0x60, 0x60, 0x60)
INK_AMBER = RGBColor(0xF0, 0xA8, 0x39)
TABLE_HEAD_FILL = "0D2818"
TABLE_BAND_FILL = "F4F7EF"
CALLOUT_FILL = "FDF6E3"

doc = Document()

# ---------- styling helpers ----------
def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)

def shade_table_header(table):
    for cell in table.rows[0].cells:
        set_cell_bg(cell, TABLE_HEAD_FILL)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.bold = True
                r.font.size = Pt(10)

def shade_alternating(table, start_row=1):
    for i, row in enumerate(table.rows[start_row:], start=start_row):
        if i % 2 == 0:
            for cell in row.cells:
                set_cell_bg(cell, TABLE_BAND_FILL)

def style_table(table, head=True, alt=True):
    table.style = "Table Grid"
    if head: shade_table_header(table)
    if alt: shade_alternating(table)
    # set fonts
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            for p in cell.paragraphs:
                for r in p.runs:
                    if not r.font.size:
                        r.font.size = Pt(10)
                    r.font.name = "Calibri"

def add_h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = INK_DARK
    return p

def add_h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = INK_DARK
    return p

def add_h3(text):
    p = doc.add_heading(text, level=3)
    for r in p.runs:
        r.font.color.rgb = INK_DARK
    return p

def add_para(text, bold=False, italic=False, size=11, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return p

def add_bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.size = Pt(11)
    return p

def add_callout(text, fill_hex=CALLOUT_FILL, label=None):
    """Single-cell shaded box used for INDEPENDENCE / SOLUTION / SHIFT notes."""
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    set_cell_bg(cell, fill_hex)
    p = cell.paragraphs[0]
    if label:
        rl = p.add_run(label + " // ")
        rl.bold = True
        rl.font.size = Pt(10)
        rl.font.color.rgb = INK_DARK
    rb = p.add_run(text)
    rb.font.size = Pt(10)
    rb.font.color.rgb = INK_DARK

def add_table(headers, rows, widths_in=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        c.text = ""
        p = c.paragraphs[0]
        r = p.add_run(h)
        r.font.size = Pt(10)
        r.font.bold = True
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            c = t.rows[i].cells[j]
            c.text = ""
            for line_idx, line in enumerate(str(val).split("\n")):
                if line_idx == 0:
                    p = c.paragraphs[0]
                else:
                    p = c.add_paragraph()
                r = p.add_run(line)
                r.font.size = Pt(10)
    if widths_in:
        for j, w in enumerate(widths_in):
            for row in t.rows:
                row.cells[j].width = Inches(w)
    style_table(t)
    return t

# =====================================================================
# DOCUMENT
# =====================================================================

# ---------- TITLE / COVER ----------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = title.add_run("APEX · AGENTIC MERCH · Q1 REFERENCE WALKTHROUGH · APPAREL RETAILER")
r.font.size = Pt(11)
r.font.color.rgb = INK_GREEN
r.font.bold = True

big = doc.add_paragraph()
r = big.add_run("APEX · Agentic Merch")
r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = INK_DARK

sub = doc.add_paragraph()
r = sub.add_run("Q1 MVP — Chief Merchant WBR")
r.font.size = Pt(20); r.font.color.rgb = INK_GREY

dek = doc.add_paragraph()
r = dek.add_run("Friday-night detection · Saturday overnight chain · Monday 8 AM brief — Apparel Retailer as reference client")
r.font.italic = True; r.font.size = Pt(11); r.font.color.rgb = INK_GREY

doc.add_paragraph()
add_para("A walkthrough of the APEX Agentic-Merch Q1 service across the APEX value-delivery chain — Scenario → Solution → Use Case → Service → Persona → KPI — anchored on an Apparel Retailer Week-16 performance miss and the overnight agent chain that turns the exception into a Monday-morning Done / Your Call / Watch brief.", italic=True, size=10, color=INK_GREY)

# Title metadata table
add_table(
    headers=["Field", "Value"],
    rows=[
        ["Document type",     "APEX Service Walkthrough · Service-to-KPI"],
        ["Anchor service",    "RC-E2E-03 · Assortment & Pricing  (markdown decisioning surface)"],
        ["Co-anchor service", "RC-E2E-02 · Inventory & Replenishment  (allocation + new-PO surface)"],
        ["Backstop service",  "RC-E2E-01 · Demand Sensing & Planning  (consumer-intent confirmation)"],
        ["APEX scenario",     "WBR Exception Triage — composite of rc-dynamic-markdown-optimization-rc + rc-auto-replenishment-tuning + rc-baseline-demand-forecasting"],
        ["Practice",          "RC — Retail & Consumer · Apparel sub-Practice"],
        ["Reference client",  "Apparel Retailer (illustrative — figures for alignment purposes only)"],
        ["Demo event",        "Apparel · Midwest · Week 16 · −12% miss · supply disruption root cause"],
        ["Cadence",           "Weekly Business Review · Friday detect → Saturday execute → Monday brief"],
        ["Q1 agents",         "Agents 1 · 4 · 8 · 13 · 13.3 (5 of 6 — Agents 6 & 7 deferred to Sprint 2)"],
        ["Authors",           "Joe G., Aditi G."],
        ["For",               "Apurva L., Saurabh V."],
    ],
    widths_in=[2.0, 4.5],
)

# Independence callout
doc.add_paragraph()
add_callout(
    "Agentic-Merch Q1 manifests, decision-rights tables, and policy thresholds use Deloitte-approved Microsoft platform language throughout. "
    "Independence rules require: 'Deloitte's Microsoft practice,' 'DMTSP,' 'Microsoft platform capabilities.' "
    "Cost-offset partnerships (BVA, ECIF workarounds) are positioned to clients via approved language only.",
    label="INDEPENDENCE NOTE",
)

doc.add_page_break()

# ---------- CONTENTS ----------
add_h1("Contents")
add_para(
    "This document walks the Agentic-Merch Q1 service through every link of the APEX value-delivery chain. "
    "Sections 1–2 set context (the reference moment, why now). Sections 3–8 traverse the six chain links — "
    "Scenario, Solution, Use Cases, Services, Personas, KPIs. Sections 9–11 stage delivery — three waves, commercial envelope, "
    "and the Wave-1 reference implementation."
)
add_table(
    headers=["#", "Section", "What it covers"],
    rows=[
        ["1",  "The Reference Exception at the Apparel Retailer", "The Friday 9:43 PM Apparel miss the retailer detects every week"],
        ["2",  "Why Agentic Merch, Why Apparel Retail, Why Now", "Five forces converging on weekly-cycle merchandising; why apparel is the right reference"],
        ["3",  "Chain Link 1 · Scenario", "Today's manual WBR vs. the agentic overnight chain"],
        ["4",  "Chain Link 2 · Solution", "5-of-6 agent orchestration · WBR cadence · commercial envelope"],
        ["5",  "Chain Link 3 · Use Cases", "Six agent-level use cases composing 20 sprint-level UCs"],
        ["6",  "Chain Link 4 · Services", "RC-E2E-03 anchor + RC-E2E-02 / RC-E2E-01 composition"],
        ["7",  "Chain Link 5 · Personas", "9 personas — 4 human gates, 5 agent personas"],
        ["8",  "Chain Link 6 · KPIs", "Per-event KPIs and four derived commercial KPIs"],
        ["9",  "The Three Waves", "W1 Foundation · W2 Pilot · W3 Scale & Fuse"],
        ["10", "Commercial Envelope", "Tier durations, cost ranges, effort points"],
        ["11", "Reference Wave-1 Implementation", "What the W1 prototype demonstrates — three demo paths"],
        ["A",  "Appendix · Glossary",                  "Selected terms with APEX- and merchandising-specific meaning"],
        ["B",  "Appendix · Artifact Cross-Reference",  "How this document fits the broader Q1 artifact set"],
    ],
    widths_in=[0.4, 2.6, 4.5],
)

doc.add_page_break()

# ---------- SECTION 1 ----------
add_h1("Section 1 · The Reference Exception at the Apparel Retailer")
add_para(
    "At 9:43 PM on a Friday in Week 16 of FY27, the Apparel Retailer's Performance Analytics agent finishes the weekly close. "
    "Apparel — the retailer's largest category — has missed plan by 12%. It is the second consecutive week-over-week miss. "
    "Twenty-four SKUs are understocked across eight Midwest stores. The Chief Merchant is offline for the weekend."
)
add_para(
    "This is the scenario. Everything downstream — the Q1 agent chain, the sprint use cases, the productized services, the persona cast, "
    "the KPI definitions — exists to answer three questions about this single recurring event:"
)
add_bullet("Is the miss a demand problem (markdown) or a supply problem (replenish)?")
add_bullet("Which actions are safe to execute overnight, and which require the Chief Merchant on Monday morning?")
add_bullet("What evidence — auditable, attributable to the specific Friday-night ledger row — survives the decisions taken during the eight hours the merchant was asleep?")

add_h3("Today, without Agentic Merch")
add_para(
    "Apparel-retail merchandising organizations typically run a Monday morning WBR with category leads spending 4–6 hours over the weekend "
    "or first thing Monday assembling exception lists, pulling ad-hoc reports, calling Planning to confirm OTB headroom, and individually triaging 6–10 "
    "category misses. The Chief Merchant arrives Monday to a 60-page deck, no actions taken yet, and the operating week half-lost before the day starts."
)
add_h3("Tomorrow, with Agentic Merch")
add_para(
    "Five agents work the chain in deterministic sequence — The Analyst (detect) → The Demand Checker (diagnose) → The Finance Lead (guardrail) → "
    "The Operations Lead (act) → The Briefer (synthesize). Material decisions emit a structured 14-field LEDGER row before they execute. "
    "By Monday 8 AM, the Chief Merchant opens a Microsoft Teams Adaptive Card with three buckets — Done / Your Call / Watch — "
    "and a single decision is staged with two pre-modelled paths."
)

# ---------- SECTION 2 ----------
add_h1("Section 2 · Why Agentic Merch, Why Apparel Retail, Why Now")

add_h2("2.1 · The five forces converge on the apparel WBR")
add_para(
    "APEX exists because five forces converged across FY25–FY26. For apparel-retail merchandising specifically, all five land on the Monday WBR: "
    "(1) cycle-time pressure on weekly merchandising decisions; (2) demand-signal proliferation outpacing manual analytics; "
    "(3) margin compression making over-markdown the dominant P&L leak; (4) FY27 enterprise AI mandates requiring auditable agent chains; "
    "(5) Microsoft-platform alignment with Adaptive-Card-based HITL on Teams."
)

add_h2("2.2 · Why an apparel retailer is the right Q1 reference")
add_para(
    "Apparel retail carries the right mix of complexity and controllability for the Q1 reference. It has weekly cycle structure (WBR cadence), "
    "discrete actionable categories (apparel, accessories, footwear), tractable inventory units (SKUs in DCs), demand-elasticity history (markdown depth × demand response), "
    "and approval-routing maturity (Planning Lead / VP Supply Chain / Category DMM / Chief Merchant). It is also a legitimate pricing surface "
    "where wrong markdowns measurably erode margin — making the markdown-block decision a credible headline for the Q1 demo."
)
add_bullet("Weekly-cycle merchandising — Friday close → Monday brief is the canonical apparel WBR cadence")
add_bullet("Mature OTB discipline — apparel retailers carry sophisticated Open-To-Buy management, making Agent 4 guardrails a natural fit")
add_bullet("Markdown elasticity history — prior season's price-vs-demand history informs intent-vs-supply diagnosis (Agent 1)")
add_bullet("Multi-tier approval routing — Planning Lead ($ ≤50K) → VP SC ($50–100K) → Chief Merchant ($ >100K) is the standard apparel governance ladder")
add_bullet("Microsoft Teams maturity — apparel HQs run on Teams, making Adaptive Card HITL the dominant approval surface")

add_para(
    "If Agentic-Merch Q1's orchestration pattern handles an apparel retailer's WBR, it ports across the Deloitte DMTSP retail portfolio "
    "with manifest-only configuration changes — same decision-rights model, same MERML / SCML / CXML schemas, same Adaptive Card surface, same LEDGER audit shape."
)

add_h2("2.3 · Why this document exists")
add_para(
    "Three conversations at the Apparel Retailer need the same material with different emphasis. The Chief Merchant and CFO need to see a credible "
    "margin-protection-and-cycle-compression story. The CIO needs to see manifest-driven, audit-grade architecture — agents are not magic; they are "
    "deterministic state machines with LLM augmentation. The Steering Committee needs a wave-by-wave delivery cadence with auditable W1 deliverables. "
    "This walkthrough provides all three views without compromising on the technical substance under each."
)

doc.add_page_break()

# ---------- SECTION 3 ----------
add_h1("Section 3 · Chain Link 1 · Scenario")
add_para(
    "The Scenario link of the APEX value-delivery chain is the customer's operational reality stated without abstraction. It is not yet a solution, "
    "not yet an agent architecture, not yet a service. It is what the Apparel Retailer experiences every week."
)

add_h2("3.1 · The Apparel Retailer's situation")
add_para(
    "The reference Apparel Retailer operates an apparel banner with multi-region distribution centers, ~800 stores across North America, "
    "10–15 active categories per banner, and a Chief Merchant accountable for ~$3B annual category revenue. WBR cadence is canonical: "
    "Friday-close performance → weekend exception synthesis → Monday morning steering. Each WBR surfaces 6–10 material category misses "
    "ranked by severity-times-consecutive-miss-count. The current weekly approval load on the Chief Merchant is ~12 hours of reading and "
    "decisioning across the WBR, mid-week pulse review, and end-of-week sign-offs."
)

add_h2("3.2 · The current-state response")
add_para(
    "Today the response across the merchandising organization is a sequential hand-off chain made of spreadsheets, emails, Teams chats, "
    "and informal phone calls. Walked end-to-end on a representative Friday-to-Monday cycle, it looks like:"
)
add_bullet("T+0 (Fri 9:43 PM): Performance close lands. No human reads it until Saturday morning.")
add_bullet("T+12h (Sat 9 AM): Category Director scans top-line numbers, identifies obvious misses, cc's Planning.")
add_bullet("T+18h (Sat 3 PM): Planning Lead pulls OTB position from the planning system, sends back to Category Director.")
add_bullet("T+24h (Sun 9 AM): Category Director assembles slide deck — typically 8–12 slides per material miss, no consistent template.")
add_bullet("T+36h (Sun 9 PM): Chief Merchant reads brief in inbox, formulates questions for Monday WBR.")
add_bullet("T+44h (Mon 8 AM): WBR convenes — 90 minutes of discussion, no decisions captured in a system, follow-up actions in spreadsheets.")
add_bullet("T+50h (Mon 2 PM): Allocation team begins manual rebalance work; new POs route through email approval chains.")
add_bullet("T+72h (Tue noon): First inventory movement reaches Midwest stores — three-day lag from Friday's signal.")

add_h2("3.3 · The friction, quantified")
add_table(
    headers=["Dimension", "Typical value", "Implication"],
    rows=[
        ["Time from Friday-close miss to first inventory movement", "~72 hours median", "Three-day lag costs sell-through during peak weekend traffic"],
        ["Chief Merchant attention consumed per WBR (preparation + meeting)", "~6 hours/week", "$300K+/year of merchant capacity locked in synthesis, not strategy"],
        ["LEDGER-grade audit-row completeness on weekly decisions", "~28% of decisions captured systematically", "Regulatory and ESG-audit traceability reconstructed from email"],
        ["Markdown-vs-replen decision accuracy (vs. demand-elasticity policy)", "~64% conformance", "36% of weekly markdowns either over-cut margin or miss the supply-only fix"],
        ["Aggregate weekly action spend visibility", "Tracked in a planning spreadsheet, refreshed Wed", "$200K–$500K guardrail breached frequently and discovered post-fact"],
        ["Approval-routing self-healing on absence (PTO, OOO)", "Manual reroute via email", "1–3 day approval stalls when default approver is out"],
    ],
    widths_in=[2.6, 1.7, 3.2],
)

add_h2("3.4 · Scope of the Q1 Apparel-Retailer scenario")
add_para("The Q1 Agentic-Merch scenario specifically includes:")
add_bullet("Friday performance close ingest from MERML.Plan + MERML.Actuals (Agent 13)")
add_bullet("Per-category, per-region exception ranking with severity + consecutive-miss-count rule")
add_bullet("Consumer demand-intent confirmation via CXML.Signal (Agent 1) — 0–1.0 score with baseline comparison")
add_bullet("Markdown-block decision when intent ≥ baseline AND root cause = supply disruption")
add_bullet("Four-guardrail validation via MERML.OTB + MERML.Margin + MERML.Markdown + MERML.AggregateSpend (Agent 4)")
add_bullet("Inventory rebalance execution within DC-existing-stock (Agent 8 · ZT)")
add_bullet("New-PO staging routed by $-tier to Planning Lead / VP SC / Chief Merchant via Microsoft Teams Adaptive Card")
add_bullet("Done / Your Call / Watch synthesis into the Monday morning brief (Agent 13.3 ADAPT)")
add_bullet("LEDGER 14-field audit row on every decision · WORM enforcement · hash-chain validation")

add_para("It specifically excludes:")
add_bullet("Pricing optimization for full-margin price changes (Agent 6 — deferred to Sprint 2)")
add_bullet("Promotion-effectiveness optimization across competing price events (Agent 7 — deferred to Sprint 2)")
add_bullet("Net-new-vendor onboarding (separate scenario — RC-E2E-09 Product Tracking)")
add_bullet("Strategic merchandising decisions on category mix (separate scenario — RC-E2E-08 Customer Acquisition)")
add_bullet("Distribution-center labor planning (adjacent — RC-E2E-02 Inventory & Replenishment expansion)")

doc.add_page_break()

# ---------- SECTION 4 ----------
add_h1("Section 4 · Chain Link 2 · Solution")
add_para(
    "The Solution link is how APEX answers the Scenario. It is an architectural description — what runs, where it runs, what contracts bind the components, "
    "what guarantees the audit trail, what makes the whole substrate cheap to extend to the next scenario."
)

add_h2("4.1 · Agentic Merch Q1 in one sentence")
add_callout(
    "Five agents — The Analyst, The Demand Checker, The Finance Lead, The Operations Lead, The Briefer — run a deterministic Friday-night-to-Monday-morning chain "
    "that detects performance exceptions, diagnoses demand vs. supply root cause, validates four financial guardrails, executes safe inventory rebalance "
    "autonomously, stages new-PO commitments for human approval, and delivers a Done / Your Call / Watch brief into Microsoft Teams — every step emitting "
    "a 14-field LEDGER WORM audit row.",
    label="SOLUTION",
)

add_h2("4.2 · The five Q1 agents")
add_table(
    headers=["Agent (Q1 Persona)", "Responsibility", "Input", "Output", "HITL?"],
    rows=[
        [
            "Agent 13 · The Analyst\n(Performance Analytics)",
            "Detect performance exception · rank by severity + consecutive-miss · route context to demand checker and guardrail governor",
            "MERML.Plan · MERML.Actuals · MERML.Inventory · DC health",
            "Ranked exception list · root-cause hypothesis · routing manifest",
            "ACK only — observe-and-route",
        ],
        [
            "Agent 1 · The Demand Checker\n(Consumer Signals)",
            "Confirm consumer demand intent · diagnose miss as demand vs. supply",
            "CXML.Signal (search/social/traffic/conversion) · prior-week baseline",
            "Intent score 0–1.0 · diagnosis label · confidence",
            "ACK only — observe-and-route",
        ],
        [
            "Agent 4 · The Finance Lead\n(MFP Guardrail Governor)",
            "Validate four guardrails — OTB buffer · margin impact · markdown depth · aggregate weekly spend",
            "Proposed action package · MERML.OTB · MERML.Margin · MERML.WeeklyEnvelope",
            "Pass / HITL / Block per guardrail · combined verdict",
            "HITL when guardrail thresholds breached (5–15% OTB, 50–150 bps margin, 10–20% markdown, $200K–$500K aggregate)",
        ],
        [
            "Agent 8 · The Operations Lead\n(Allocation & Replenishment)",
            "Execute inventory rebalance within DC-existing-stock · stage new POs for human approval · activate secondary suppliers",
            "Approved action package · SCML.Inventory · SCML.Supplier · TMS connectors",
            "Transfer orders · staged PO records · supplier activations",
            "Always for new POs · ZT for rebalance · ZT for pre-approved suppliers",
        ],
        [
            "Agent 13.3 · The Briefer\n(ADAPT — Brief Synthesis)",
            "Synthesize Done / Your Call / Watch buckets · render Adaptive Cards · deliver Monday brief in Teams",
            "Week's LEDGER rows · pending HITL queue · Teams identity bindings",
            "Brief markdown · Adaptive Cards · response webhooks",
            "ACK for Done · HITL for Your Call · ACK for Watch",
        ],
    ],
    widths_in=[1.7, 2.2, 1.5, 1.5, 1.4],
)

add_h2("4.3 · Mapping to the APEX 6-agent reference architecture")
add_para(
    "Q1 is a 5-of-6 deployment of the APEX agent reference architecture. The mapping is deliberate — Q1 takes the agent shape that "
    "already exists in APEX and projects it onto the WBR cadence. Each Q1 agent corresponds to an APEX agent role, with naming chosen "
    "for merchandising readability (the Chief Merchant says 'The Analyst,' not 'Assess agent')."
)
add_table(
    headers=["APEX agent role", "Q1 persona name", "Q1 agent ID", "What changes vs. canonical APEX"],
    rows=[
        ["Assess",          "The Analyst",        "Agent 13",   "Performance-miss focus rather than asset-state focus"],
        ["Classify",        "The Demand Checker", "Agent 1",    "Demand-intent classification rather than disposition classification"],
        ["Quantify",        "The Finance Lead",   "Agent 4",    "Four-guardrail composite rather than single $-impact"],
        ["Approve",         "(human only)",       "—",          "Decision Rights routing — Planning Lead / VP SC / DMM / Chief Merchant"],
        ["Act",             "The Operations Lead","Agent 8",    "Two-mode action — ZT rebalance vs. HITL new-PO"],
        ["Evidence-Write",  "The Briefer",        "Agent 13.3", "Done / Your Call / Watch synthesis with Adaptive Card delivery"],
    ],
    widths_in=[1.5, 1.7, 1.2, 3.0],
)

add_h2("4.4 · Commercial envelope — Q1 decision-rights thresholds")
add_para(
    "The commercial envelope is the policy layer that decides 'do this autonomously vs. require human approval.' For Q1, the envelope is "
    "expressed as the Decision Rights table — proposed thresholds based on apparel-retail industry norms, locked per client during W1 Foundation. "
    "Amber values are proposed defaults; the Chief Merchant confirms or overrides during onboarding."
)
add_table(
    headers=["Action type", "Default approver", "Escalation trigger", "Escalation to", "Window"],
    rows=[
        ["New PO ≤ $50K",                                "Planning Lead",        "No response",  "VP Supply Chain",       "24h"],
        ["New PO $50K–$100K",                            "VP Supply Chain",      "No response",  "Chief Merchant",        "48h"],
        ["New PO > $100K",                               "Chief Merchant (block + escalate)", "—",  "Block — resurfaces next week", "—"],
        ["Markdown depth 10–20%",                        "Category DMM",         "No response",  "Chief Merchant",        "24h"],
        ["Markdown depth > 20%",                         "Chief Merchant (block + escalate)", "—",  "Block — resurfaces next week", "—"],
        ["Aggregate weekly spend $200K–$500K",           "Chief Merchant",       "No response",  "Block — resurfaces",    "48h"],
        ["Inventory rebalance (existing DC stock)",      "Auto (Zero Touch)",    "—",            "—",                     "—"],
        ["Secondary supplier (pre-approved list)",       "Auto (Zero Touch)",    "—",            "—",                     "—"],
        ["Secondary supplier (new / unapproved)",        "Chief Merchant",       "No response",  "Block",                 "48h"],
    ],
    widths_in=[2.3, 1.7, 1.4, 1.7, 0.7],
)

add_h2("4.5 · Three design commitments")

add_h3("4.5.1 · Manifest-driven, not model-driven")
add_para(
    "Every choice that could change between client banners or close periods — which LLM is used for brief synthesis, which tools each agent calls, "
    "which thresholds gate which decisions, which approvers receive which Adaptive Cards — is declared in a tenant-scoped manifest. "
    "Agents read against the manifest at runtime; no routing logic is embedded in the agent. New approver? Update manifest. Threshold change? "
    "Update manifest. Same agents continue running; same LEDGER schema continues being emitted."
)

add_h3("4.5.2 · Audit row, not log line")
add_para(
    "Every decision — every detection, every diagnosis, every guardrail evaluation, every approval, every downstream action, every brief tile — "
    "emits a 14-field structured LEDGER record on a write-once-read-many store. Hash-chain links each row to its predecessor in the same scenario. "
    "Every brief item deep-links to the LEDGER row that produced it. The audit row is the contract between the agent chain and the regulatory + "
    "ESG audit functions; it is not an instrumentation afterthought."
)

add_h3("4.5.3 · HITL is first-class, not bolted on")
add_para(
    "The HITL gates between Quantify and Act are declared in the orchestration manifest as first-class primitives, not retrofitted human checks. "
    "Adaptive Cards are versioned schema artifacts; response webhooks write back to LEDGER and notify the originating agent. "
    "Timeout escalation rules are declared, not coded. The merchant's Teams identity is the primary binding — there is no separate approval app to install."
)

doc.add_page_break()

# ---------- SECTION 5 ----------
add_h1("Section 5 · Chain Link 3 · Use Cases")
add_para(
    "The Use Case link is the set of discrete operational patterns the Solution activates. One Solution typically enables multiple Use Cases — "
    "each a defensible incremental delivery surface. Q1's six agent-level use cases compose from 20 sprint-level use cases documented in "
    "mvp-scope-q1-usecases.xlsx. The sprint plan refines acceptance criteria, story points, and exit conditions for each."
)

add_h3("5.1 · Use Case · Exception Detection & Diagnosis")
add_table(
    headers=["Trigger", "Friday performance close lands · Agent 13 picks up MERML.Plan vs. MERML.Actuals delta"],
    rows=[
        ["Agents",   "13 · The Analyst, then 1 · The Demand Checker"],
        ["Tools",    "MERML.Plan retrieval · MERML.Actuals · severity-ranking model · CXML.Signal aggregation · LLM (root-cause narrative)"],
        ["HITL",     "None — observe-and-route only · markdown auto-blocked when intent healthy AND root cause = supply"],
        ["Envelope", "~15% of W1 scope · highest detection-quality value"],
        ["Sprint UCs", "UC-01 · UC-02 · UC-03"],
    ],
    widths_in=[1.5, 5.0],
)

add_h3("5.2 · Use Case · Four-Guardrail Validation")
add_table(
    headers=["Trigger", "Proposed action package emitted by Agent 8 — runs four guardrails in parallel"],
    rows=[
        ["Agents",   "4 · The Finance Lead"],
        ["Tools",    "MERML.OTB · MERML.Margin · MERML.Markdown · MERML.WeeklyEnvelope · LEDGER writer"],
        ["HITL",     "Per guardrail — 5–15% OTB → CM, 50–150 bps margin → CM, 10–20% markdown → DMM, $200K–$500K aggregate → CM"],
        ["Envelope", "~20% of W1 scope · highest financial-control value"],
        ["Sprint UCs", "UC-04 · UC-05 · UC-06 · UC-07"],
    ],
    widths_in=[1.5, 5.0],
)

add_h3("5.3 · Use Case · Inventory Action Execution")
add_table(
    headers=["Trigger", "Agent 4 returns Pass on all four guardrails for a proposed action"],
    rows=[
        ["Agents",   "8 · The Operations Lead"],
        ["Tools",    "SCML.Inventory · SCML.Supplier · TMS connector · POS connector · LEDGER writer"],
        ["HITL",     "ZT for rebalance within DC-existing-stock · HITL for new POs (per $ tier) · HITL for new suppliers"],
        ["Envelope", "~25% of W1 scope · highest cycle-compression value"],
        ["Sprint UCs", "UC-08 · UC-09 · UC-10"],
    ],
    widths_in=[1.5, 5.0],
)

add_h3("5.4 · Use Case · Brief Synthesis & Delivery")
add_table(
    headers=["Trigger", "Sunday-evening synthesis run picks up week's LEDGER rows"],
    rows=[
        ["Agents",   "13.3 · The Briefer (ADAPT)"],
        ["Tools",    "LEDGER reader · Adaptive Card renderer · Microsoft Teams Bot Framework · Entra identity resolution"],
        ["HITL",     "ACK for Done bucket · HITL for Your Call bucket · ACK for Watch bucket · 24h/48h response windows"],
        ["Envelope", "~20% of W1 scope · highest UX-quality value"],
        ["Sprint UCs", "UC-11 · UC-12 · UC-13 · UC-20"],
    ],
    widths_in=[1.5, 5.0],
)

add_h3("5.5 · Use Case · Human Approval Routing")
add_table(
    headers=["Trigger", "Agent 13.3 emits a Your Call item with approver routing"],
    rows=[
        ["Agents",   "None (approvals-mcp orchestrates the human)"],
        ["Tools",    "Microsoft Teams Adaptive Card · Entra identity · LEDGER writer · response-window timer"],
        ["HITL",     "Always for material entries · escalation chain self-heals on timeout (24h/48h)"],
        ["Envelope", "~10% of W1 scope · highest persona-experience value"],
        ["Sprint UCs", "UC-14 · UC-15 · UC-16 · UC-17 · UC-18"],
    ],
    widths_in=[1.5, 5.0],
)

add_h3("5.6 · Use Case · Audit-Row Persistence (LEDGER WORM)")
add_table(
    headers=["Trigger", "Any agent emits a decision (action, block, route, ack)"],
    rows=[
        ["Agents",   "Cross-cutting — every Q1 agent writes here"],
        ["Tools",    "LEDGER · 14-field schema validator · WORM enforcement · hash-chain link"],
        ["HITL",     "None — write-only · consumed by audit, regulatory, ESG, and brief deep-links"],
        ["Envelope", "~10% of W1 scope · foundation for evidence packets"],
        ["Sprint UCs", "UC-19"],
    ],
    widths_in=[1.5, 5.0],
)

add_h2("5.7 · Mapping to the 20 sprint-level Q1 use cases")
add_para(
    "The six agent-level use cases above are each realized by multiple sprint-level use cases in mvp-scope-q1-usecases.xlsx. "
    "The sprint workbook carries detailed acceptance criteria, story points, dependencies, and Q1-status per UC."
)
add_table(
    headers=["Agent-level use case", "Sprint-level use cases (mvp-scope-q1-usecases.xlsx)"],
    rows=[
        ["5.1 Exception Detection & Diagnosis", "UC-01 · UC-02 · UC-03"],
        ["5.2 Four-Guardrail Validation",        "UC-04 · UC-05 · UC-06 · UC-07"],
        ["5.3 Inventory Action Execution",       "UC-08 · UC-09 · UC-10"],
        ["5.4 Brief Synthesis & Delivery",       "UC-11 · UC-12 · UC-13 · UC-20"],
        ["5.5 Human Approval Routing",           "UC-14 · UC-15 · UC-16 · UC-17 · UC-18"],
        ["5.6 Audit-Row Persistence (LEDGER)",   "UC-19"],
    ],
    widths_in=[2.7, 4.0],
)

add_h2("5.8 · Incremental delivery path")
add_para(
    "Individually, each use case is a deliverable. Together they compose the full WBR scenario. The composition is deliberate — the Apparel "
    "Retailer can start with detection-only (5.1), add guardrails (5.2), then execution (5.3), then brief delivery (5.4) — without rearchitecting. "
    "Each incremental UC unlocks an incremental cycle-compression and audit-quality benefit."
)

doc.add_page_break()

# ---------- SECTION 6 ----------
add_h1("Section 6 · Chain Link 4 · Services")
add_para(
    "The Service link is what gets productized. A Service is a named, subscribable APEX capability with a defined persona, SLO, commercial terms, "
    "and version contract. The Apparel Retailer subscribes to Services; the Apparel Retailer does not subscribe to agents directly."
)

add_h2("6.1 · Anchor service")
add_h3("RC-E2E-03 · Assortment & Pricing")
add_para(
    "The foundational pricing-and-disposition decisioning service. Anchors the markdown-block decision (Agent 13 → Agent 1 → Agent 4 → DMM bypass) "
    "and the WBR brief surface (Agent 13.3 ADAPT). Primary persona: Chief Merchant. "
    "Secondary personas: Category DMM, Planning Lead, VP Supply Chain. Sprint-2 expansion adds Agent 6 (Pricing) and Agent 7 (Promotion Optimization) "
    "for the full pricing-and-promotion surface."
)

add_h2("6.2 · Co-anchor service")
add_h3("RC-E2E-02 · Inventory & Replenishment")
add_para(
    "The inventory-action service. Owns Agent 8 — rebalance + new-PO + secondary-supplier activation. Primary persona: VP Supply Chain. "
    "Co-binds with RC-E2E-03 because the WBR exception triage spans both services — a markdown-block decision (RC-E2E-03) typically pairs "
    "with a rebalance + new-PO action (RC-E2E-02). Q1 is the first scenario where Deloitte demonstrates RC-E2E-03 + RC-E2E-02 fused at the orchestration layer."
)

add_h2("6.3 · Backstop service")
add_h3("RC-E2E-01 · Demand Sensing & Planning")
add_para(
    "The demand-intelligence service. Owns Agent 1 — consumer demand-intent confirmation. Primary persona: Demand Planner. "
    "Q1 uses RC-E2E-01 as a context provider — Agent 1's intent score is the diagnostic that flips a markdown into a rebalance. "
    "RC-E2E-01's broader baseline-forecasting and new-product-demand-sensing capabilities are out of Q1 scope."
)

add_h2("6.4 · Service composition contract")
add_table(
    headers=["Commitment", "RC-E2E-03 Assortment & Pricing", "RC-E2E-02 Inventory & Replenishment", "RC-E2E-01 Demand Sensing & Planning"],
    rows=[
        ["Canonical schemas",
            "MERML.Plan · MERML.Actuals · MERML.OTB · MERML.Margin · MERML.Markdown · MERML.WeeklyEnvelope",
            "SCML.Inventory · SCML.Supplier · SCML.Transfer · SCML.PO",
            "CXML.Signal · CXML.Intent · CXML.Baseline"],
        ["Q1 agents deployed",
            "Agent 13 · Agent 4 · Agent 13.3",
            "Agent 8",
            "Agent 1"],
        ["MCP tool surfaces",
            "Performance close ingest · OTB query · margin model · markdown elasticity · Adaptive Card writer · LEDGER",
            "DC-inventory query · TMS transfer · supplier activation · PO writer · LEDGER",
            "Search/social/traffic ingest · intent score · baseline retrieval · LEDGER"],
        ["HITL policies",
            "Markdown 10–20% → DMM · aggregate $200K–$500K → CM",
            "PO ≤$50K → Planning Lead · PO $50K–$100K → VP SC · PO >$100K → Block + CM",
            "ACK only — no autonomous actions"],
        ["Power BI surfaces",
            "WBR brief dashboard · Done/Your Call/Watch trending · markdown-NPV trending",
            "Allocation completeness · in-stock rate · PO velocity",
            "Demand-intent trending · category-by-region intent heatmap"],
        ["Sprint-2 expansion",
            "Agent 6 (Pricing) · Agent 7 (Promotion Optimization)",
            "Vendor-managed inventory · ladder replenishment for promos",
            "Baseline forecasting · new-product demand sensing"],
    ],
    widths_in=[1.4, 2.0, 1.7, 1.4],
)

add_h2("6.5 · Apparel-Retailer tenant manifest subscription")
add_para("The Apparel-Retailer tenant manifest subscribes to the three services by declaration:")
add_para(
    '{ "tenant_id": "appretl", "subscriptions": [\n'
    '  { "service_id": "RC-E2E-03", "service_version": "v1.0.0-q1" },\n'
    '  { "service_id": "RC-E2E-02", "service_version": "v1.0.0-q1" },\n'
    '  { "service_id": "RC-E2E-01", "service_version": "v1.0.0-q1" }\n'
    ']}',
    italic=True, size=10, color=INK_GREY,
)
add_para(
    "Version pinning is deliberate. When the APEX Service Catalog releases v1.1.0 of RC-E2E-03 (adding Agent 6 Pricing in Sprint 2), "
    "the Apparel Retailer decides when to upgrade. PATCH bumps are auto-applied. MINOR/MAJOR bumps require explicit subscription change."
)

doc.add_page_break()

# ---------- SECTION 7 ----------
add_h1("Section 7 · Chain Link 5 · Personas")
add_para(
    "The Persona link is the human-and-agent role cast that makes the scenario operational. Persona depth is what separates an architecturally "
    "correct solution from one that lands. Q1 names nine personas — four humans, five agents — and each has a defined surface, decision right, "
    "and audit footprint. Generic role names ('the merchant') are not used; every persona is named with the specific responsibility carried in the WBR cycle."
)

add_h2("7.1 · Primary persona — Chief Merchant")
add_callout(
    "PRIMARY · CHIEF MERCHANT // Owns the Weekly Business Review. Reads the Monday 8 AM brief in Microsoft Teams. Approves Your Call items "
    "(aggregate weekly spend $200K–$500K · markdown >20% escalation · PO >$100K escalation). Sees Done items in confirmation-only tier. "
    "Watches Watch items for next-week escalation. Persona-level binding: persona_level = 'chief_merchant'.",
    label=None, fill_hex="EEF5DD",
)

add_h2("7.2 · Secondary human personas")
add_para("Three secondary human personas each carry a distinct gate in the WBR decision-rights ladder. Roles are deliberate — Q1 does not invent roles; it captures the apparel-retail organizational structure as it already exists.")
add_table(
    headers=["Role", "Function in the Q1 scenario"],
    rows=[
        ["Planning Lead",
            "First-line PO approver. Owns POs ≤ $50K. 24h response window. Escalation target: VP Supply Chain. Bypassed in Q1 demo because the action was a DC-existing-stock rebalance (no new PO at this tier)."],
        ["VP Supply Chain",
            "Mid-tier PO approver. Owns POs $50K–$100K. 48h response window. Escalation target: Chief Merchant. Receives the demo $62K expedited PO via Adaptive Card on Monday morning."],
        ["Category DMM",
            "Markdown approver. Owns markdown depth 10–20%. 24h response window. Escalation target: Chief Merchant. Bypassed in Q1 demo because the markdown was auto-blocked (intent healthy → supply diagnosis)."],
    ],
    widths_in=[1.7, 5.0],
)

add_h2("7.3 · Five agent personas")
add_para(
    "Q1 names its agents. The Chief Merchant reads Done / Your Call / Watch tiles attributed to 'The Analyst noticed,' 'The Demand Checker confirmed,' "
    "'The Finance Lead approved,' 'The Operations Lead acted,' 'The Briefer summarised.' Naming is not stylistic — it is auditable: every LEDGER row "
    "carries the agent_id and persona_name fields, and the brief deep-links each tile to the row that produced it."
)
add_table(
    headers=["Agent persona", "Agent ID", "MCP", "Q1 chapter", "Surface in the brief"],
    rows=[
        ["The Analyst",        "Agent 13",   "merml-mcp", "Detect (Fri 9:43 PM)",   "Done — exception #1 ranked"],
        ["The Demand Checker", "Agent 1",    "cxml-mcp",  "Diagnose (Sat 12:30 AM)", "Done — markdown blocked rationale"],
        ["The Finance Lead",   "Agent 4",    "merml-mcp", "Guardrail (Sat 2:15 AM)", "Done — $190K headroom confirmed"],
        ["The Operations Lead","Agent 8",    "scml-mcp",  "Act (Sat 3:45 AM)",       "Done — 24 SKUs · 8 stores · Your Call — $62K PO"],
        ["The Briefer",        "Agent 13.3", "merml-mcp (sub-agent of 13)", "Brief (Mon 8:00 AM)", "Owns the entire brief surface"],
    ],
    widths_in=[1.5, 1.0, 1.2, 1.7, 1.6],
)

add_h2("7.4 · Why persona depth matters")
add_para(
    "A seller who names only the Chief Merchant says 'we built something for the Chief Merchant.' A seller who names the full nine-persona cast "
    "shows competence — Q1 understands that approval routing self-heals on Planning Lead's PTO, that DMM bypass on healthy demand prevents margin erosion, "
    "and that audit traceability requires the agent personas to be named, not anonymous. The nine personas are also the change-management surface — "
    "Field Enablement trains all nine roles in W2 onboarding, not just the Chief Merchant."
)

add_h2("7.5 · Persona-to-plane mapping")
add_table(
    headers=["Persona", "APEX plane(s) touched", "Surface"],
    rows=[
        ["Chief Merchant (primary)",       "Experience · Decision",            "Microsoft Teams Adaptive Card · Power BI WBR dashboard"],
        ["Planning Lead",                  "Experience · Decision",            "Adaptive Card · email fallback"],
        ["VP Supply Chain",                "Experience · Decision",            "Adaptive Card · email fallback"],
        ["Category DMM",                   "Experience · Decision",            "Adaptive Card · pricing-system writeback console"],
        ["The Analyst (Agent 13)",         "Decision · Runtime · Ledger",      "merml-mcp · LEDGER · brief author"],
        ["The Demand Checker (Agent 1)",   "Decision · Runtime · Ledger",      "cxml-mcp · LEDGER"],
        ["The Finance Lead (Agent 4)",     "Decision · Runtime · Ledger",      "merml-mcp · LEDGER"],
        ["The Operations Lead (Agent 8)",  "Decision · Runtime · Integration · Ledger", "scml-mcp · TMS connector · LEDGER"],
        ["The Briefer (Agent 13.3)",       "Experience · Decision · Runtime · Ledger",  "Teams Bot Framework · Adaptive Card renderer · LEDGER deep-links"],
    ],
    widths_in=[1.9, 2.3, 2.5],
)

doc.add_page_break()

# ---------- SECTION 8 ----------
add_h1("Section 8 · Chain Link 6 · KPIs")
add_para(
    "The KPI link is where the chain closes. APEX's signature is that every KPI attributes back to the specific decision audit row that produced it. "
    "Q1 KPIs split into two layers — per-WBR-cycle operational metrics and four derived commercial KPIs that anchor Apparel-Retailer steering conversations."
)

add_h2("8.1 · Per-WBR-cycle KPIs")
add_para(
    "A closed WBR cycle (Friday close → Monday brief → Wednesday EOD response window expiry) emits a closeout audit row carrying the following measured "
    "KPI fields. Each is bound to the specific trace_id of the underlying exception."
)
add_table(
    headers=["KPI", "Measurement", "Target (W2 Pilot)", "Target (W3 Scale)"],
    rows=[
        ["Time-to-decision",                              "Friday close → first inventory action (median)",        "≤ 6 hours",     "≤ 4 hours"],
        ["Chief-Merchant attention reclaimed",            "Hours/week reduced from current ~6h WBR prep",          "+3.0 hrs",      "+4.5 hrs"],
        ["Margin protected per week",                     "$ over auto-blocked-markdown baseline",                  "$80K / week",   "$140K / week"],
        ["Decision-conformance rate",                     "% of decisions matching policy",                          "≥ 90%",         "≥ 96%"],
        ["LEDGER audit-row completeness",                 "% with all 14 fields populated",                         "≥ 97%",         "≥ 99.5%"],
        ["Aggregate-spend guardrail effectiveness",       "% of weeks with cumulative ≤ $500K",                     "≥ 95%",         "100%"],
        ["Adaptive-Card response latency",                "Median minutes from delivery to response",                "≤ 90 min",      "≤ 30 min"],
        ["Approval-routing self-heal",                    "% of timeouts auto-escalated correctly",                  "100%",          "100%"],
    ],
    widths_in=[2.4, 2.4, 1.0, 1.0],
)

add_h2("8.2 · Derived commercial KPIs")
add_para(
    "Beyond the per-cycle operational metrics, four derived KPIs anchor Apparel-Retailer steering conversations. Each is computed from the per-cycle "
    "audit-row stream — none are estimates."
)

add_h3("8.2.1 · Network-Wide Margin Protected")
add_para(
    "The headline financial metric. Computed as SUM(margin_protected_per_week) across all WBR-cycle closeout audit rows over a trailing 52 weeks, "
    "with the destroy-margin baseline being the markdown $-impact that would have executed in the absence of the demand-intent check. "
    "$140K/week × 52 weeks = ~$7.3M annualized at W3 run-rate for a single category × region."
)

add_h3("8.2.2 · Chief-Merchant Capacity Returned")
add_para(
    "The operational metric that matters to the Chief Merchant and the COO. Computed as 4.5 hrs/week reclaimed × 52 weeks = ~234 hours/year per merchant. "
    "Across a typical 8–12 senior merchant headcount, this translates to ~2,300 hours/year of reclaimed senior-merchant capacity — equivalent to one "
    "FTE of strategic merchandising attention reinjected into the operating model."
)

add_h3("8.2.3 · WBR Audit-Trail Completeness")
add_para(
    "The audit metric. Computed as the percentage of WBR-cycle decisions whose audit rows carry all 14 required fields, every version stamp, and a "
    "valid hash chain back to the upstream agent decision. W2 target ≥ 97%; W3 target ≥ 99.5% — the residual 0.5% being acknowledged data-quality gaps "
    "in source systems (resolved as Sprint-2 data-contract work)."
)

add_h3("8.2.4 · Decision-Loop Compression")
add_para(
    "The cycle-time metric. Today's Friday-close-to-first-inventory-action loop is 72 hours median; W2 target is ≤ 6 hours; W3 target is ≤ 4 hours. "
    "The compression unlocks weekend sell-through capture — exception responses arrive at stores before peak Saturday and Sunday shopping rather than "
    "three days late."
)

add_h2("8.3 · KPI attribution in the audit row")
add_para(
    "The closeout audit row — emitted once per resolved WBR exception — carries the structure below. Every KPI attributes back to specific rows."
)
add_para(
    "{\n"
    '  "trace_id": "trc_2027-04-22_2143_appretl_apparel_midwest_wk16",\n'
    '  "scenario_id": "rc-wbr-exception-triage-apparel",\n'
    '  "service_id": "RC-E2E-03",\n'
    '  "tenant_id": "appretl",\n'
    '  "wbr_week": 16,\n'
    '  "category": "Apparel",\n'
    '  "region": "Midwest",\n'
    '  "miss_pct": -0.12,\n'
    '  "consecutive_miss_count": 2,\n'
    '  "intent_score": 0.72,\n'
    '  "diagnosis": "supply",\n'
    '  "guardrail_otb_buffer_post": 0.21,\n'
    '  "guardrail_margin_bps": 18,\n'
    '  "actions_taken": ["rebalance_24sku_8stores"],\n'
    '  "actions_staged": ["po_62k_apparel_8sku"],\n'
    '  "approver_route": "vp_supply_chain",\n'
    '  "margin_protected_usd": 142000,\n'
    '  "decision_loop_hours": 5.7,\n'
    '  "ledger_row_count": 11,\n'
    '  "hash_chain_valid": true\n'
    "}",
    italic=True, size=9, color=INK_GREY,
)

doc.add_page_break()

# ---------- SECTION 9 ----------
add_h1("Section 9 · The Three Waves")
add_para(
    "The same chain — Scenario → Solution → Use Case → Service → Persona → KPI — projects across three delivery Waves. "
    "W1 Foundation is prerequisite; W2 Pilot is where the scenario becomes real; W3 Scale & Fuse is where the scenario meets the enterprise."
)

add_h2("9.1 · Wave 1 · Foundation — single-banner shadow run")
add_para(
    "W1 is the one-time investment per tenant. It produces the rails that make the scenario possible and make every subsequent scenario cheap. "
    "For the Apparel Retailer, W1 establishes the LEDGER plane, the MERML / SCML / CXML canonical schemas, the Microsoft Teams Adaptive Card "
    "binding for the Chief Merchant, and three of five agents (13, 1, 4) running in shadow mode against synthetic Apparel Week-16 fixtures."
)
add_table(
    headers=["W1 element", "What it looks like for the Apparel Retailer"],
    rows=[
        ["Scope",            "One banner · Apparel category · Midwest region cluster · ~50 stores"],
        ["Agents live",      "3 of 5 (Agents 13, 1, 4) — read-only, no automated action"],
        ["Use cases",        "5.1 + 5.2 + 5.6 (detection, guardrail, audit-row)"],
        ["Personas onboarded", "Chief Merchant + Category DMM + Planning Lead (read-only brief access)"],
        ["Schemas",          "MERML.Plan · MERML.Actuals · MERML.OTB · MERML.Margin · MERML.Markdown · CXML.Signal (W1 scope)"],
        ["MCPs live",        "merml-mcp · cxml-mcp (subset) — scml-mcp scaffolded, not yet connected to TMS"],
        ["Microsoft surfaces","Adaptive Card preview (no live send) · Power BI dashboard · Entra identity binding scaffolded"],
        ["LEDGER",           "Live · WORM enforcement on · 14-field schema · hash chain validated end-to-end"],
        ["HITL flow",        "Read-only — Chief Merchant sees brief preview, no Approve/Hold buttons"],
        ["Synthetic fixtures","50K synthetic Apparel exceptions across 12 weeks of synthetic actuals × 24 categories × 4 regions"],
        ["KPIs measured",    "Detection accuracy · diagnosis accuracy · LEDGER completeness · brief-rendering quality"],
        ["Exit criteria",    "≥ 95% detection on synthetic fixtures · ≥ 90% diagnosis on synthetic fixtures · 100% LEDGER hash-chain valid · CM-readable brief preview"],
    ],
    widths_in=[2.0, 4.5],
)

add_h2("9.2 · Wave 2 · Pilot — 250-store cluster live + first-decision Adaptive Card")
add_para(
    "W2 is where the scenario becomes real. A 250-store anchor cluster (typically the apparel banner's North-Central or Mid-Atlantic cluster) "
    "runs the live overnight chain. All five agents activate. Adaptive Cards land in the Chief Merchant's Teams. The first auto-blocked markdown "
    "and the first auto-rebalance happen for real, with audit rows joining the LEDGER."
)
add_table(
    headers=["W2 element", "What it looks like for the Apparel Retailer"],
    rows=[
        ["Scope",            "250 stores · multi-region · 8 categories · live WBR cadence"],
        ["Agents live",      "All 5 (full orchestration with HITL gating)"],
        ["Use cases",        "5.1 through 5.6 fully exercised"],
        ["Personas onboarded","All 9 named personas operating live"],
        ["Schemas",          "All MERML + SCML + CXML schemas live"],
        ["MCPs live",        "merml-mcp + cxml-mcp + scml-mcp · TMS connector · POS connector live"],
        ["Microsoft surfaces","Adaptive Cards live · Power BI WBR dashboard · Entra approver-binding live · approval webhook signing"],
        ["LEDGER",           "Live · WORM enforcement · 14-field schema · per-tenant query API for audit"],
        ["HITL flow",        "Live — Approve/Hold buttons write to LEDGER and notify Agent 8"],
        ["Real telemetry",   "Live MERML.Plan + MERML.Actuals + CXML.Signal feeds · synthetic supplemented during cold-start"],
        ["KPIs measured",    "All per-WBR-cycle KPIs · margin protected · cycle compression · audit completeness · response latency"],
        ["Exit criteria",    "12 consecutive WBR cycles green on conformance · ≥ 90% policy-conformance rate · ≥ $80K/week margin protected"],
    ],
    widths_in=[2.0, 4.5],
)

add_h2("9.3 · Wave 3 · Scale & Fuse — enterprise · multi-banner · cross-scenario")
add_para(
    "W3 is where the scenario meets the apparel-retail enterprise and where it stops running alone. Scope expands from 250 stores to all 800+; "
    "additional banners onboard via manifest-only configuration; Sprint-2 agents (6 Pricing, 7 Promotion Optimization) join the Q1 chain; "
    "fusion hooks into adjacent RC-E2E services (RC-E2E-08 customer acquisition, RC-E2E-04 customer lifecycle) activate."
)
add_table(
    headers=["W3 element", "What it looks like for the Apparel Retailer"],
    rows=[
        ["Scope",                "All 800+ stores · all banners · all categories · live WBR cadence"],
        ["Agents live",          "Q1 5 agents + Sprint-2 Agents 6 (Pricing) + 7 (Promotion Optimization) — 7 of 8"],
        ["Use cases",            "Full 20 sprint-level use cases plus Sprint-2 expansion"],
        ["Schemas",              "All MERML + SCML + CXML shapes; cross-banner aggregation views"],
        ["Cross-scenario fusion","Fuse with RC-E2E-08 Customer Acquisition · RC-E2E-04 Customer Lifecycle · RC-E2E-09 Product Tracking"],
        ["Microsoft surfaces",   "Adaptive Cards across full approver cast · 4 Power BI surfaces · Sentinel monitoring · Defender for Cloud audit"],
        ["LEDGER",               "Cross-tenant query API · regulator export profile · ESG reporting integration"],
        ["KPIs measured",        "All W2 KPIs at enterprise scale + cross-scenario fusion KPIs"],
        ["Exit criteria",        "52 consecutive WBR cycles green · ≥ 96% policy-conformance · ≥ $140K/week margin protected · ≥ 99.5% LEDGER completeness"],
    ],
    widths_in=[2.0, 4.5],
)

add_h2("9.4 · Cross-wave load-bearing connections")
add_para("Three structural dependencies bind the waves into one progression. They are not architectural coincidences — they are the load-bearing connections that make the next wave cheap given the previous one.")
add_callout(
    "W1 → W2 // The Adaptive Card prototyped in W1 (Chief Merchant preview, no live send) becomes the live Approve/Hold surface in W2. "
    "Same schema, same Entra binding, same LEDGER write-back contract. W1 invests in the substrate; W2 turns on the action.",
    label=None,
)
add_callout(
    "W1 → W2 // The LEDGER plane established in W1 accumulates the audit-row corpus that W2 uses to demonstrate KPI attribution. "
    "12 weeks of W1 shadow rows become the baseline against which W2's first 90 days are measured.",
    label=None,
)
add_callout(
    "W2 → W3 // The W2 pilot KPIs unlock W3 scaling. The Steering Committee accepts the W2 audit-row corpus as evidence; the W3 enterprise rollout "
    "is a configuration-only expansion (more stores, more banners, more categories — same agents, same schemas, same LEDGER).",
    label=None,
)

doc.add_page_break()

# ---------- SECTION 10 ----------
add_h1("Section 10 · Commercial Envelope")
add_para(
    "The Apparel-Retailer Q1 engagement encompasses four commercial tiers corresponding to the Approach stage (Discovery and SOW build-out), "
    "Wave 1 Foundation (single-banner shadow run), Wave 2 Pilot (250-store live + first-decision Adaptive Card), and Wave 3 Scale & Fuse (enterprise + multi-banner)."
)
add_table(
    headers=["Tier", "Duration", "Cost envelope", "Effort", "Primary deliverable"],
    rows=[
        ["Approach",            "3–5 weeks",      "$200K–$400K",      "20 SP",   "Discovery report · banner scoping · WBR-decision-rights confirmation · SOW"],
        ["Wave 1 Foundation",   "4–6 months",     "$1.0M–$2.5M",      "40 SP",   "Single-banner shadow-run · audit-row end-to-end · brief preview · Adaptive-Card scaffolding"],
        ["Wave 2 Pilot",        "6–9 months",     "$3.5M–$8M",        "40 SP",   "250-store live · 12-week WBR-cycle proof · first auto-blocked markdown · first auto-rebalance"],
        ["Wave 3 Scale & Fuse", "9–15 months",    "$10M–$28M",        "55+ SP",  "Enterprise 800+ stores · multi-banner · Sprint-2 agents 6 + 7 live · cross-scenario fusion"],
    ],
    widths_in=[1.6, 1.0, 1.4, 0.6, 2.4],
)

add_h2("10.1 · What the envelope buys")
add_para(
    "The commercial envelope is not a project-labor-hours calculation. It reflects the total value-delivery bundle: Deloitte-delivered solutioning, "
    "Microsoft platform commitments, third-party ISV teaming where applicable, manifest-driven configuration, change-management for nine personas, "
    "and the LEDGER substrate that makes every subsequent scenario cheap."
)

add_h2("10.2 · Envelope vs. durable value at Apparel-Retailer scale")
add_para(
    "An 800-store W3 enterprise footprint produces approximately $7.3M annualized margin-protected value at run-rate per category-region pair, "
    "plus 2,300 reclaimed senior-merchant hours, plus a regulator-grade audit-row corpus that satisfies ESG-reporting and FY27 board-level AI-governance "
    "requirements. The W3 cost envelope ($10M–$28M) recovers in 6–18 months on margin-protection alone, before counting capacity or audit value."
)

add_h2("10.3 · Independence posture in commercial conversations")
add_callout(
    "The seller and delivery team use Deloitte-approved language throughout the Apparel-Retailer commercial conversation. "
    "Cost-offset partnerships with Microsoft (BVA-led entry, ECIF workarounds via ISV teaming, MACC-burndown structures) are positioned to the "
    "client only via approved language: 'Deloitte's Microsoft practice,' 'DMTSP,' 'Microsoft platform capabilities,' 'co-investment.' "
    "Deloitte does not receive ECIF directly from Microsoft.",
    label="INDEPENDENCE",
)

doc.add_page_break()

# ---------- SECTION 11 ----------
add_h1("Section 11 · Reference Wave-1 Implementation")
add_para(
    "The Agentic-Merch Q1 Wave-1 prototype is the concrete proof point that anchors this walkthrough. It is a runnable web application that executes "
    "the Friday-night-to-Monday-morning chain end-to-end against synthetic Apparel Week-16 fixtures, with three demo paths designed for stakeholder reviews."
)

add_h2("11.1 · What the W1 prototype demonstrates")
add_table(
    headers=["Capability", "What it demonstrates"],
    rows=[
        ["Synthetic Apparel performance fixtures", "50K synthetic exceptions across 12 categories × 4 regions × 12 weeks of Apparel-Retailer-pattern actuals"],
        ["The Analyst (Agent 13, live)",          "Severity ranking · consecutive-miss detection · root-cause hypothesis routing"],
        ["The Demand Checker (Agent 1, live)",    "Intent score 0–1.0 · baseline comparison · demand-vs-supply diagnosis · markdown auto-block"],
        ["The Finance Lead (Agent 4, live)",      "Four-guardrail evaluation · OTB buffer · margin impact · markdown depth · aggregate spend"],
        ["The Briefer (Agent 13.3, preview)",     "Done / Your Call / Watch synthesis · brief preview rendering · LEDGER deep-link generation"],
        ["LEDGER plane",                          "14-field WORM audit row · hash-chain validation · brief deep-link API"],
        ["Adaptive-Card schema",                  "Versioned card template · Entra identity binding · response-webhook scaffolding (no live send)"],
        ["Power BI WBR dashboard",                "Done/Your Call/Watch trending · margin-protected trending · LEDGER coverage trending"],
        ["Manifest-driven configuration",         "Tenant manifest · approver chain · threshold table · synthesizable to live W2 manifests"],
    ],
    widths_in=[2.4, 4.2],
)

add_h2("11.2 · Three canonical demo paths")
add_para("Three demo paths planned for Apparel-Retailer stakeholder reviews, each ~3–5 minutes:")
add_para(
    "Demo Path A · The exception loop. Pick a synthetic Apparel Week-16 fixture. Watch The Analyst rank it #1 of 7 exceptions; "
    "watch The Demand Checker confirm intent 0.72 and diagnose 'supply'; watch The Finance Lead pass all four guardrails with $190K headroom; "
    "watch The Operations Lead stage the rebalance + $62K PO; watch The Briefer render the Done / Your Call / Watch tiles. ~4 minutes."
)
add_para(
    "Demo Path B · The margin-protected trending. Open the Power BI WBR dashboard. Show 12 weeks of synthetic margin-protected trending. "
    "Show the auto-blocked-markdown event distribution. Show the Adaptive-Card response-latency histogram. ~3 minutes."
)
add_para(
    "Demo Path C · The audit trail. Pick a synthetic exception trace_id. Open the LEDGER row stream for that trace. Walk the 11 rows from "
    "detection through brief delivery. Show the hash-chain validation. Show the regulator-export profile. ~4 minutes."
)

add_h2("11.3 · Prototype technology stack")
add_table(
    headers=["Layer", "Stack"],
    rows=[
        ["Frontend",     "React + Vite · Tailwind · Power BI Embedded"],
        ["Backend",      "Node.js (Fastify) · pg-pool"],
        ["Persistence",  "PostgreSQL (LEDGER + state)"],
        ["Identity",     "Entra ID (mock binding for prototype)"],
        ["Inference",    "Local Ollama (Llama-class model) for offline demos · Azure OpenAI in W2"],
        ["Orchestration","Lightweight DAG runner (W1) → CAF Dynamic Orchestrator + A2A Swarm (W2+)"],
        ["Adaptive Cards","Microsoft Bot Framework SDK · scaffolded send (no live Teams in W1) · live in W2"],
        ["Observability","OpenTelemetry · Power BI · LEDGER query API"],
        ["Hosting",      "Local Docker for W1 demos · Azure App Service + Azure Container Apps in W2+"],
    ],
    widths_in=[1.6, 5.0],
)

add_h2("11.4 · Mapping prototype → APEX architecture")
add_para(
    "The prototype is deliberately a subset of the full APEX architecture — enough to make the design concrete for Apparel-Retailer stakeholders, "
    "not so much that it duplicates W2's investment. Each prototype element maps cleanly to its production counterpart."
)
add_table(
    headers=["Prototype element", "APEX architecture element", "Apparel-Retailer production surface"],
    rows=[
        ["Local Postgres",                  "LEDGER plane (audit row store)",       "Azure Fabric Lakehouse + Postgres index in production"],
        ["Mock TMS / POS APIs",             "Integration plane",                    "The Apparel Retailer's actual TMS and POS markdown stack"],
        ["Mock Adaptive Card device",       "Experience plane",                     "Chief Merchant's actual Microsoft Teams environment"],
        ["Local Ollama model",              "Decision plane (LLM)",                 "Azure OpenAI via AEF LLM Gateway (CAF substrate)"],
        ["Lightweight DAG runner",          "Runtime plane",                        "CAF Dynamic Orchestrator + A2A Swarm in production"],
        ["Local synthetic Apparel fixtures","SOR + Real-Time Hub + Bronze · Silver","MERML.Plan + MERML.Actuals + CXML.Signal feeds in production"],
        ["Manifest in source repo",         "Governance Plane · tenant manifest",   "Purview-cataloged manifest + Sentinel-monitored deployment"],
    ],
    widths_in=[2.0, 2.4, 2.2],
)

add_h2("11.5 · Prototype limitations and next steps")
add_para("The prototype is deliberately scoped for Wave 1 demonstration. It does not:")
add_bullet("Connect to real Apparel-Retailer plan / actuals streams — synthetic data only; real integration is W2 scope")
add_bullet("Run on Azure — everything is local (Docker + Ollama on a laptop); Azure migration is W2")
add_bullet("Activate The Operations Lead live — staging-only in W1; live TMS / POS writes are W2")
add_bullet("Implement live Entra authentication — hardcoded Chief Merchant persona for demo; Entra integration is W2")
add_bullet("Live-send Adaptive Cards in Teams — preview only; live send is W2")
add_bullet("Handle multi-banner orchestration — single banner / single category / single region only")

add_para(
    "These are not gaps in the architecture — they are intentional scope boundaries. The W1 prototype demonstrates the agentic flow end-to-end "
    "with auditable rigor; W2 turns on the live integrations."
)

doc.add_page_break()

# ---------- APPENDIX A · GLOSSARY ----------
add_h1("Appendix A · Glossary")
add_para("Selected terms used in this document with APEX- and merchandising-specific meaning.")
glossary = [
    ("ADAPT",            "Action Router / Brief Synthesis sub-agent (Agent 13.3) — Q1's brief-synthesis primitive · sub-agent of Agent 13"),
    ("Adaptive Card",    "Microsoft Teams rich-card format used by Q1 for HITL approval surfaces; versioned schema; webhook write-back to LEDGER"),
    ("Agent",            "A bounded unit of intent + tool-use that emits typed output and writes a LEDGER audit row"),
    ("Apparel Retailer", "The illustrative Q1 reference client — generic apparel banner, ~800 stores, weekly WBR cadence"),
    ("APEX",             "Agentic Platform for Enterprise eXecution — Deloitte's reference 5-plane / 6-agent runtime architecture"),
    ("Audit row",        "14-field structured record emitted by every agent decision; WORM-stored in LEDGER with hash-chain link to predecessor"),
    ("Briefer",          "Q1 persona name for Agent 13.3 — synthesizes Done / Your Call / Watch tiers and renders Adaptive Cards"),
    ("Chief Merchant",   "Q1 primary persona — owns the WBR, approves Your Call items, sees Done items, watches Watch items"),
    ("CXML",             "Customer Experience Model Language — canonical schema family for consumer signals, intent, baselines"),
    ("Decision Rights",  "Q1's threshold table — which agent action requires which human approver under which condition"),
    ("DMM",              "Divisional Merchandise Manager — secondary human persona, default approver for markdown depth 10–20%"),
    ("HITL",             "Human-In-The-Loop — explicit decision gate where the agent stages an action awaiting a named human's response"),
    ("LEDGER",           "Q1's WORM audit-row store — 14-field schema, hash-chain links, regulator-export profile, brief deep-links"),
    ("MERML",            "Merchandising Model Language — canonical schema family for plan, actuals, OTB, margin, markdown, weekly envelope"),
    ("OTB",              "Open-To-Buy — apparel-retail term for remaining weekly purchase budget after committed and proposed actions"),
    ("Persona level",    "LEDGER field that binds a decision to its persona owner (e.g., persona_level = 'chief_merchant')"),
    ("SCML",             "Supply Chain Model Language — canonical schema family for inventory, supplier, transfer, PO"),
    ("WBR",              "Weekly Business Review — apparel-retail merchandising cadence, Q1's anchoring rhythm"),
    ("WORM",             "Write-Once-Read-Many — LEDGER persistence guarantee preventing audit-row tampering"),
    ("Your Call",        "Brief tier for HITL-pending items requiring Chief Merchant approval before execution"),
    ("Zero Touch (ZT)",  "Decision-rights gate where the agent executes autonomously — e.g., DC-existing-stock rebalance"),
]
add_table(
    headers=["Term", "Definition"],
    rows=glossary,
    widths_in=[1.6, 5.0],
)

# ---------- APPENDIX B · ARTIFACT CROSS-REFERENCE ----------
add_h1("Appendix B · Artifact Cross-Reference")
add_para(
    "This document is one of several APEX Agentic-Merch Q1 artifacts covering the Apparel-Retailer engagement. "
    "The cross-reference below shows how each artifact composes the full Q1 deliverable bundle."
)
add_table(
    headers=["Artifact", "Primary audience", "Purpose"],
    rows=[
        ["APEX-Agentic-Merch-Q1-Walkthrough.docx (this)",
            "Chief Merchant · CFO · CIO · Steering Committee · DMTSP",
            "Executive walkthrough of the value-delivery chain"],
        ["mvp-scope-q1.html",
            "Apurva L. · Saurabh V. · architects",
            "Q1 scope authoritative source — 6 tabs covering scope, walk-through, contracts, sequence, coverage, decision rights"],
        ["mvp-scope-q1-journey-map.html",
            "Chief Merchant · all 9 personas",
            "Persona journey map — 9 personas × 5-chapter timeline · cross-persona swimlane"],
        ["mvp-scope-q1-usecases.xlsx",
            "Engagement Manager · Lead Engineer · Scrum lead",
            "20 sprint use cases with acceptance criteria, persona matrix, decision rights, WBR timeline"],
        ["APEX-RC-Scenario-Chains.xlsx",
            "RC Practice lead · Engagement Architect",
            "RC scenario library — anchor scenarios for Q1 (markdown optimization, replenishment tuning, demand forecasting)"],
        ["RC-E2E-03 service spec (APEX Service Catalog)",
            "APEX Practice RC lead · Engagement Architect",
            "Service composition contract: schemas, agents, SLOs, persona contract"],
        ["RC-E2E-02 service spec (APEX Service Catalog)",
            "APEX Practice RC lead · Engagement Architect",
            "Co-anchor service spec for inventory & replenishment surface"],
        ["Q1-Sprint-Plan.xlsx (forthcoming)",
            "Engagement Manager · Lead Engineer · Scrum lead",
            "20 sprint use cases with story points, dependencies, exit criteria, sprint assignment"],
        ["Q1-Resource-Plan.xlsx (forthcoming)",
            "DMTSP PMO · Engagement Manager",
            "Resource loading by phase, role, sprint"],
        ["Q1-ROI-Case.html (forthcoming)",
            "CFO · COO · Steering Committee",
            "Margin-protected uplift case with sensitivity analysis"],
        ["Q1-Tenant-Manifest.json (forthcoming)",
            "Lead Engineer · Platform Engineer",
            "Apparel-Retailer-specific manifest: approver chain, thresholds, persona bindings, schema versions"],
        ["LEDGER-Schema.json (forthcoming)",
            "Lead Engineer · Audit · Compliance",
            "14-field audit-row contract · WORM enforcement specification · regulator-export profile"],
    ],
    widths_in=[2.4, 1.8, 2.4],
)

doc.add_paragraph()
add_para(
    "Independence posture preserved across every APEX Agentic-Merch Q1 artifact: 'Deloitte's Microsoft practice,' 'DMTSP,' 'Microsoft platform capabilities,' "
    "'co-investment.' Deloitte does not receive ECIF directly from Microsoft.",
    italic=True, size=10, color=INK_GREY,
)

doc.add_paragraph()
end = doc.add_paragraph()
r = end.add_run("— End of walkthrough —")
r.italic = True; r.font.size = Pt(10); r.font.color.rgb = INK_GREY
end.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ---------- save ----------
os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print(f"Wrote: {OUT}")
print(f"Paragraphs: {len(doc.paragraphs)}")
print(f"Tables: {len(doc.tables)}")
