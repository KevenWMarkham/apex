"""Apply Google fonts to existing DOCX deliverables.

Matches the HTML deliverables' typography:
  - Fraunces        → Heading 1/2/3/4 (serif display)
  - IBM Plex Sans   → all body text (sans-serif)
  - JetBrains Mono  → code/monospace runs

Walks paragraphs + tables (recursive) and updates every run's font name
to the appropriate face. Also updates each Heading style and the Normal style.

Note: For Word to render these fonts directly, install them on the system from
Google Fonts. Otherwise Word substitutes its closest fallback (typically Calibri
for sans, Cambria for serif, Consolas for mono) — visual feel is preserved.
"""
from __future__ import annotations
import io, sys
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

# ---- font policy ----
SERIF = "Fraunces"        # Headings
SANS  = "IBM Plex Sans"   # Body
MONO  = "JetBrains Mono"  # Code (was Consolas)

# Old fonts to detect for code-run rebinding
MONO_OLD = ("Consolas", "Courier New", "Courier")

# Heading styles that should use the serif face
HEADING_STYLES = ("Heading 1", "Heading 2", "Heading 3", "Heading 4",
                  "Heading 5", "Heading 6", "Title", "Subtitle")


def set_run_font(run, font_name: str):
    """Set font name with full asian-script + complex-script attributes."""
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), font_name)
    # Also clear the python-docx-managed name so it doesn't conflict
    run.font.name = font_name


def style_set_font(style, font_name: str):
    """Set the font on a paragraph style."""
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), font_name)
    style.font.name = font_name


def target_for_run(run, paragraph) -> str:
    """Decide which font face this run should use."""
    style_name = paragraph.style.name if paragraph.style else ""
    # Headings
    if any(style_name.startswith(h) for h in HEADING_STYLES):
        return SERIF
    # Existing-mono runs (code samples, JSON blocks, IDs)
    current_font = run.font.name or ""
    if current_font in MONO_OLD or current_font == MONO:
        return MONO
    # Default: body sans
    return SANS


def update_paragraphs(paragraphs):
    for p in paragraphs:
        for run in p.runs:
            target = target_for_run(run, p)
            set_run_font(run, target)


def update_doc(path: Path) -> tuple[int, int]:
    """Update fonts in a DOCX file. Returns (paragraphs_touched, runs_touched)."""
    doc = Document(path)

    # 1. Default style
    try:
        style_set_font(doc.styles["Normal"], SANS)
    except KeyError:
        pass

    # 2. Heading styles
    for h_style in HEADING_STYLES:
        try:
            style_set_font(doc.styles[h_style], SERIF)
        except KeyError:
            pass

    # 3. Body paragraphs
    pcount = 0; rcount = 0
    update_paragraphs(doc.paragraphs)
    pcount += len(doc.paragraphs)
    for p in doc.paragraphs:
        rcount += len(p.runs)

    # 4. Table cells (recursive — handles nested tables)
    def visit_table(tbl):
        nonlocal pcount, rcount
        for row in tbl.rows:
            for cell in row.cells:
                update_paragraphs(cell.paragraphs)
                pcount += len(cell.paragraphs)
                for p in cell.paragraphs:
                    rcount += len(p.runs)
                # nested tables
                for nested in cell.tables:
                    visit_table(nested)
    for tbl in doc.tables:
        visit_table(tbl)

    # 5. Headers / footers
    for section in doc.sections:
        for hf in (section.header, section.footer,
                   section.first_page_header, section.first_page_footer,
                   section.even_page_header, section.even_page_footer):
            try:
                update_paragraphs(hf.paragraphs)
                for tbl in hf.tables:
                    visit_table(tbl)
            except Exception:
                pass

    # 6. Save (in place)
    try:
        doc.save(str(path))
    except PermissionError:
        # Try a v2 sibling if locked
        alt = path.with_name(path.stem + "-v2" + path.suffix)
        doc.save(str(alt))
        print(f"  [WARN] {path.name} was locked; saved as {alt.name}")
        return pcount, rcount

    return pcount, rcount


# -------------------------------------------------------------------------
# DRIVER
# -------------------------------------------------------------------------

ROOTS = [
    Path(r"C:\Stage\Clients\Industries\Consumer\Retail\Kroger\02_projects\FY27_Pipeline\assortment-pricing-agentic\deliverables"),
    Path(r"C:\Stage\Clients\Industries\TMT\Global_Media_Entertainment\Walt_Disney\02_projects\FY27_Pipeline\subscriber-lifecycle-agentic\deliverables"),
]

print("Applying Google fonts to DOCX deliverables...")
print(f"  Serif (headings):  {SERIF}")
print(f"  Sans (body):       {SANS}")
print(f"  Mono (code):       {MONO}")
print()

total_files = 0
total_paragraphs = 0
total_runs = 0
errors = []

for root in ROOTS:
    if not root.exists():
        print(f"  [SKIP] {root} not found")
        continue
    docx_files = list(root.rglob("*.docx"))
    docx_files = [f for f in docx_files if not f.name.startswith("~$")]
    if not docx_files:
        continue
    print(f"[{root.parts[-3]}/.../{root.parts[-1]}]")
    for docx in sorted(docx_files):
        try:
            p, r = update_doc(docx)
            total_files += 1
            total_paragraphs += p
            total_runs += r
            rel = docx.relative_to(root)
            print(f"  ✓ {rel}  · {p} paragraphs · {r} runs")
        except Exception as e:
            errors.append((docx, str(e)))
            print(f"  ✗ {docx.name}  · ERROR: {e}")
    print()

print(f"Done. {total_files} file(s) updated · {total_paragraphs} paragraphs · {total_runs} runs.")
if errors:
    print(f"\nErrors ({len(errors)}):")
    for f, e in errors:
        print(f"  {f.name}: {e}")
