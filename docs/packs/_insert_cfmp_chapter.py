"""Insert a new Chapter 26 — CFMP v0.2 — into APEX-Architecture-v5.docx.

Reads the parent architecture docx, locates the "Appendix A — Reference
Manifests" heading, and inserts the new chapter immediately before it.
Saves as APEX-Architecture-v5.1-with-CFMP-chapter.docx in this folder.

Run once; the output docx is the durable artifact. The parent docx in
~/Downloads is never modified.

Style matching: uses the same Heading 1 / Heading 2 / Normal / List Bullet
styles already in the parent doc.
"""

from __future__ import annotations

from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

SRC = Path(r"C:\Users\kmarkham\Downloads\APEX-Architecture-v5.docx")
OUT = Path(r"C:\Stage\Clients\Industries\APEX\docs\packs\APEX-Architecture-v5.1-with-CFMP-chapter.docx")

# ---------------------------------------------------------------------------
# Chapter content
# ---------------------------------------------------------------------------

CHAPTER_H1 = "Chapter 26 — Customer Focused Merchandise Pack (CFMP) v0.2"

CHAPTER_INTRO = (
    "This chapter documents the Customer Focused Merchandise Pack (CFMP) — "
    "the seventh Industry Solution Pack in the APEX catalog and the first "
    "organized around a customer-moment spine rather than an operator function. "
    "CFMP is a sibling to (not a replacement for) the Retail Merchandising Pack "
    "documented in §9.6: that pack serves the Chief Merchant / DMM / Planning Lead "
    "buyers via MERML; CFMP serves the CMO / CX-VP buyer via a CXML-led journey "
    "spine, with MERML and SCML as supporting data planes. "
    "This chapter records CFMP's position in the family, the Core-interface "
    "extensions it requires (a new persona type, a new HITL surface, and a "
    "proposed Interface #15), and the Phase-1 live status of the working demo "
    "running on APEX-M."
)

SECTIONS: list[tuple[str, list[str | list[str]]]] = [
    (
        "26.1  Position in the Pack Catalog",
        [
            "CFMP is the customer-side counterpart to the Retail Merchandising Pack v1 documented in §9.6. Both packs share the RC practice bundle (apex-rc) and the MERML + CXML schemas, but the buyer, persona, HITL surface, and KPI roll-up differ:",
            ("bullets", [
                "Buyer — Retail Merchandising serves the Chief Merchant / DMM / Planning Lead via internal data. CFMP serves the CMO / CX-VP / Loyalty Director via customer-facing flows.",
                "Primary persona — Retail Merchandising binds operator roles to Entra group OIDs. CFMP introduces a NEW persona type (`customer`) bound to loyalty ID + consent gate (see §26.4).",
                "HITL surface — Retail Merchandising uses Teams Adaptive Cards. CFMP adds the `customer_phone` transport (see §26.5).",
                "KPI roll-up — trip conversion, basket size, in-store NPS, loyalty retention, cart-abandon recovery, BOPIS attach. These are the metrics CMOs are graded on.",
            ]),
            "Pack ID: cfmp. Status: v0.2 draft (May 2026). Cloud profiles supported: APEX-M primary, APEX-G / APEX-A via the shared 10-asset bundle and the proposed Interface #15.",
        ],
    ),
    (
        "26.2  The Customer-Journey Spine",
        [
            "CFMP organizes all 18 scenarios in its v0.2 manifest around four customer moments. Every Adaptive Card, Virtual View, and agent tool maps to exactly one phase:",
            ("bullets", [
                "CHOOSE — \"What should I buy?\" Recipes, pairings, dietary filters, personalized offers, in-store ad attribution.",
                "SELECT — \"Where is it / is it on the shelf?\" Indoor wayfinding (see §26.3), OSA / shelf-gap detection, planogram compliance.",
                "BUY — \"How do I pay and leave?\" Scan-and-go, self-checkout assist, queue prediction, BOPIS pickup, returns trust.",
                "SERVICES — \"What's next from you?\" Loyalty winback (featured chain rc-loyalty-churn-prediction-winback), complaint triage, review summarization, sentiment, NPS.",
            ]),
            "Sub-tier sizing follows the standard APEX additive rule (Lite ⊂ Standard ⊂ Enterprise): Pack Lite ships 3 scenarios at $150–250K, Standard ships 10 at $500K–$1.5M, Enterprise ships all 18 at $1.5M–$3.5M. Scenario detail is in the companion workbook CFMP-Scenario-Chains-v0.2.xlsx (Featured Chains sheet matches the 14-column schema from APEX-Scenario-Chains.xlsx).",
        ],
    ),
    (
        "26.3  Proposed Interface #15 — Maps & Wayfinding",
        [
            "Architecture v5 §6 enumerates 14 Cloud Profile interfaces, with §21 (Live Translation Layer) adding #14. CFMP proposes #15 — Maps & Wayfinding — to support the cfmp-wayfinding-walk-to-product scenario and the eight downstream capabilities that ride on the same indoor-map dataset.",
            "The abstract Core contract is: an indoor-map dataset (geometry + units + levels), a routing service that returns turn-by-turn directions plus a GeoJSON polyline, and a live state set for dynamic occupancy or availability.",
            "Cloud-profile bindings:",
            ("bullets", [
                "APEX-M — Azure Maps Creator. Drawing Package upload → Conversion → Dataset + Tileset + Stateset. Wayfinding REST: POST /wayfinding/route. Web SDK indoor view renders the Tileset + route polyline.",
                "APEX-G — Google Maps Geospatial Creator (equivalent indoor dataset + routing surface).",
                "APEX-A — Amazon Location Service indoor maps + routing.",
            ]),
            "CFMP runtime: orchestrator/azure_maps.py wraps the APEX-M binding. When the per-tenant Creator dataset is configured (env vars DEMO_AZURE_MAPS_DATASET_ID + DEMO_AZURE_MAPS_KEY), the agent's route_to_product tool calls the Wayfinding REST API. When unconfigured, the orchestrator falls back to a local storemap.yaml graph for dev (the response carries a `source: local_storemap_yaml` field so the demo narrative is honest about which path executed).",
            "Strategic note: once a retailer's CAD floor plans are captured into the Creator dataset, eight other CFMP capabilities ride on the same data — \"find an associate,\" \"skip the queue,\" \"get a tour,\" BOPIS counter routing, spill-hazard distance, aisle-engagement attribution, shelf-gap dispatch, and accessibility routing. The dataset is the moat.",
        ],
    ),
    (
        "26.4  New Persona Type — `customer`",
        [
            "Existing APEX personas are employees with Entra group OID bindings. The customer has none of that. CFMP extends personas.yaml with a new persona type:",
            ("code", "- id: customer\n  identity_binding:\n    source: loyalty_id\n    scope: opted_in_only\n  consent_gate: required\n  channel: customer_phone\n"),
            "Identity is the retailer's loyalty ID (anonymous-customer flow is a Wave-1 design item — see §26.8). The `consent_gate: required` declaration means every customer-facing action stamps a consent hash into the LedgerRow (see §11.2 audit-row schema). Auditors can replay any customer-facing decision and verify the consent was in effect at the time.",
            "This persona type is reusable across future customer-facing Packs (Banking-Customer, Telco-Customer, Hospitality-Guest) and should be promoted from Pack-specific to Foundation-level when CFMP exits Pack Lite into Standard.",
        ],
    ),
    (
        "26.5  New HITL Surface — `customer_phone`",
        [
            "Architecture v5 §6 interface #9 (HITL surface) binds to Teams Adaptive Cards by default on APEX-M. The customer doesn't have Teams. CFMP adds the `customer_phone` transport binding:",
            ("bullets", [
                "APEX-M — Azure Notification Hubs signed-link push → phone app → MCP callback.",
                "APEX-G — Firebase Cloud Messaging.",
                "APEX-A — Amazon Pinpoint / SNS.",
            ]),
            "The runtime contract is identical to the Teams binding: an Adaptive Card is built, the persona's channel is resolved, the card is pushed, an approval/decline/timeout returns via the MCP callback, and the LedgerRow's hitl_status moves from PENDING through APPROVED / REJECTED / OVERRIDDEN per §11.2.",
            "CFMP today: cart-add subtotals crossing the consent threshold (default $50, configurable via DEMO_CART_HITL_THRESHOLD) mark the LedgerRow `hitl_status=pending` with rationale stamped. The orchestrator pauses cleanly in apex_orchestrator.SequentialRunner once Phase 1.5 wraps in (see §26.6 status).",
        ],
    ),
    (
        "26.6  Phase 1 Live Status",
        [
            "The Phase 1 framework wedge is live on APEX-M in the dev (Lab) substrate today. The demo is reachable at https://ca-visionkit-portal.gentlestone-9b49b099.eastus2.azurecontainerapps.io and surfaces the live framework state via the APEX panel + /architecture page.",
            "What's live:",
            ("bullets", [
                "Vendored apex_core + apex_audit + apex_purview packages inside the orchestrator container. The orchestrator stamps a sealed 14-field LedgerRow per /agent/ask, with HMAC-SHA256 chain, three-version stamps (manifest/policy/prompt/model), and content-addressed ToolCallRecord entries.",
                "Atlas-shaped Purview lineage edges (apex_orchestration_to_audit, apex_agent_to_orchestration, apex_mcp_to_agent per distinct tool) buffered offline; Phase 3 swaps the buffer for live /atlas/v2/entity/bulk uploads.",
                "HITL consent gate live at $50 cart-add (see §26.5).",
                "Microsoft Agent Framework 1.6.0 (GA) with provider toggle (Azure OpenAI default → Anthropic optional). Default model gpt-5-mini.",
                "Azure Speech (en-US-AvaMultilingualNeural) for STT (mic) and TTS (chip/typed replies via /agent/tts).",
                "Five agent tools: search_products (pgvector), get_product, recipe_for_items, suggest_pairings, route_to_product (Interface #15 binding).",
            ]),
            "What's planned (Phase 1.5 → 2 → 3):",
            ("bullets", [
                "Phase 1.5 — wrap the agent in apex_orchestrator.SequentialRunner with the canonical 4×4 gate matrix; replace untyped Product dict with apex_merml entities.",
                "Phase 2 — provision Fabric capacity (F2 trial). Bronze (raw) → Silver (canonical MERML) → Gold (Direct Lake) catalog. AuditStore swaps from in-memory to OneLake WORM Bronze. fabric-mcp + merml-mcp become the data-plane MCPs.",
                "Phase 3 — provision Purview account; install APEX classifications via CI; flush lineage buffer to live Atlas REST; drop AOAI subscription key in favor of Container App managed identity.",
                "Phase 4 — provision per-tenant Azure Maps Creator resource; upload retailer Drawing Package; embed Web SDK indoor view in the portal.",
            ]),
        ],
    ),
    (
        "26.7  10-Asset Pack Bundle (CFMP specifics)",
        [
            "CFMP follows the standard 10-asset bundle structure documented in §14.1. The CFMP-specific contents:",
            ("bullets", [
                "VV manifests — merml.product_with_planogram_location, merml.osa_by_aisle, cxml.customer_session_with_route, cxml.cart_dwell_event, cxml.dietary_profile, cfmp.storemap_view (new, references the Azure Maps Creator Dataset).",
                "Scenario manifests — 18 YAMLs spanning the four journey phases (3 Lite / 10 Standard / 18 Enterprise per the additive sub-tier rule).",
                "Source adapters — POS + Loyalty CRM + Planogram + Beacon vendor SDK (Estimote / Kontakt) + Vision AI Dev Kit MQTT + phone-app event stream.",
                "Adaptive Cards — operator cards for Teams (associate dispatch, restock alerts) AND customer-facing cards for the new customer_phone transport (route, recipe suggestion, consent prompt, loyalty milestone).",
                "Persona map — `customer` (new type, see §26.4) plus store_associate, assistant_manager, store_manager, merch_director, loyalty_director.",
                "Demo data — synthetic 800-product catalog with planogram, 30 beacons, 50 simulated shopper sessions; runs on a laptop (Lite substrate).",
                "BVA worksheet — cfmp-roi.xlsx aggregates trip-conversion, basket-size, NPS, churn, BOPIS-attach uplift.",
                "Sample SOW — Lite / Standard / Enterprise templates at $150–250K / $500K–$1.5M / $1.5–3.5M.",
                "Acceptance tests — 80+ tests including the Pack-specific wayfinding shortest-path correctness, beacon-loss camera-recovery, and customer-consent gate enforcement.",
                "Runbook + training — beacon field install + storemap capture procedure + store-team training deck + customer-consent FAQ.",
            ]),
        ],
    ),
    (
        "26.8  ARB Asks (decisions required before Pack Lite ships)",
        [
            "Three Foundation-level enhancements need ARB sign-off before CFMP can move from dev (Lab) to stage:",
            ("bullets", [
                "Promote `customer` persona type (§26.4) from Pack-specific to Foundation. It is reusable across all future customer-facing Packs.",
                "Add `customer_phone` HITL transport (§26.5) to interface #9 implementations on all three cloud profiles (APEX-M / G / A).",
                "Ratify Interface #15 — Maps & Wayfinding (§26.3) as a Core interface, with the proposed bindings for APEX-M / G / A.",
            ]),
            "Three Wave-1 BVA-workshop decisions for the first CFMP engagement:",
            ("bullets", [
                "Localization primary — BLE beacons (Estimote / Kontakt) vs. Wi-Fi RTT (802.11mc). Pick one per Pack v1; the adapter abstraction supports either.",
                "Phone-app deployment model — SDK into the retailer's existing app (production path) vs. CFMP-branded Deloitte demo app (closeable-pilot path).",
                "Anonymous-customer flow design — consent-on-first-tap UX and identity-binding fallback when loyalty ID is unavailable.",
            ]),
            "Recommendation: CFMP is sold separately from Retail Merchandising Pack. Different buyer (CMO / CX-VP), different envelope-curve, different KPI dashboard. Cross-pack Fuse opportunities to Retail Merchandising (planogram refresh, OTB simulation outputs) and to ESG (sustainability-attribute filter on dietary prefs) are tracked as Wave-2 expansions.",
        ],
    ),
    (
        "26.9  Related Artifacts",
        [
            ("bullets", [
                "docs/packs/CFMP-v0.2.md — full Pack design document (12 sections).",
                "docs/packs/CFMP-Scenario-Chains-v0.2.xlsx — 18-scenario workbook (Summary · Scenario Library · Featured Chains · Scenario→KPI Chain · Pack Sub-Tiers).",
                "Live demo — https://ca-visionkit-portal.gentlestone-9b49b099.eastus2.azurecontainerapps.io with /architecture and /api/apex/* endpoints surfacing real-time framework state.",
                "orchestrator/apex_integration.py — 341 LOC integration glue wrapping apex_audit and apex_purview at /agent/ask time.",
                "orchestrator/azure_maps.py — Interface #15 client (Azure Maps Creator Wayfinding REST + local storemap.yaml fallback).",
            ]),
        ],
    ),
]


def find_appendix_a_paragraph(doc: Document):
    """Locate the first paragraph that is the 'Appendix A — Reference Manifests' heading."""
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading") and "Appendix A" in p.text and "Reference Manifests" in p.text:
            return p
    return None


def insert_paragraph_before(target_paragraph, doc: Document, text: str, style: str | None = None):
    """Insert a new paragraph with `text` immediately before `target_paragraph`.
    Returns the new paragraph."""
    new_p = deepcopy(target_paragraph._element)
    # Clear runs from the cloned element
    for r in new_p.findall(qn("w:r")):
        new_p.remove(r)
    for hyper in new_p.findall(qn("w:hyperlink")):
        new_p.remove(hyper)
    target_paragraph._element.addprevious(new_p)

    # Convert into a python-docx paragraph object
    from docx.text.paragraph import Paragraph
    new_para = Paragraph(new_p, target_paragraph._parent)
    if style:
        try:
            new_para.style = doc.styles[style]
        except KeyError:
            pass
    new_para.add_run(text)
    return new_para


def add_block(doc: Document, target, content_block):
    """Add a content block (string OR ('bullets', [...]) OR ('code', '...'))
    just before `target`."""
    if isinstance(content_block, str):
        insert_paragraph_before(target, doc, content_block, style="Normal")
        return
    kind, payload = content_block
    if kind == "bullets":
        for item in payload:
            insert_paragraph_before(target, doc, item, style="List Bullet")
    elif kind == "code":
        # No "Code" style in this doc; use Normal but with a leading indent marker.
        for line in payload.rstrip().split("\n"):
            insert_paragraph_before(target, doc, "    " + line, style="Normal")


def main() -> None:
    print(f"Opening {SRC} …")
    doc = Document(str(SRC))

    target = find_appendix_a_paragraph(doc)
    if target is None:
        # Fallback: append at end
        print("WARN: Appendix A not found; appending at end of document")
        doc.add_paragraph()  # spacer
        doc.add_paragraph(CHAPTER_H1, style="Heading 1")
        doc.add_paragraph(CHAPTER_INTRO, style="Normal")
        for title, blocks in SECTIONS:
            doc.add_paragraph(title, style="Heading 2")
            for b in blocks:
                if isinstance(b, str):
                    doc.add_paragraph(b, style="Normal")
                elif b[0] == "bullets":
                    for item in b[1]:
                        doc.add_paragraph(item, style="List Bullet")
                elif b[0] == "code":
                    for line in b[1].rstrip().split("\n"):
                        doc.add_paragraph("    " + line, style="Normal")
    else:
        print(f"Inserting Chapter 26 before: {target.text[:60]}")
        insert_paragraph_before(target, doc, CHAPTER_H1, style="Heading 1")
        insert_paragraph_before(target, doc, CHAPTER_INTRO, style="Normal")
        for title, blocks in SECTIONS:
            insert_paragraph_before(target, doc, title, style="Heading 2")
            for b in blocks:
                add_block(doc, target, b)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Wrote {OUT}")
    print(f"Size: {OUT.stat().st_size:,} bytes")


if __name__ == "__main__":
    main()
