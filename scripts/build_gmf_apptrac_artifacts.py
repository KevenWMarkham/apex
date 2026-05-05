"""APEX documentation set for AI Loan Origination & Underwriting at GMF.

Service: AXLE-FIN-01 · Auto Loan Origination & Underwriting Automation
Microsoft stack: Doc Intelligence + Foundry agents + Fabric + Teams + Entra + Purview
Reference client: General Motors Financial (GMF) · AppTrac platform integration

Four deliverables:
  1. Walkthrough (DOCX)         · Tier0-Foundation
  2. GMF One-pager (HTML)       · Tier1-Executive
  3. Six-Agents Architecture (DOCX) · Tier2-Build
  4. AppTrac SOR Integration (DOCX) · Tier0-Foundation
"""
from __future__ import annotations
import io, sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
ROOT = Path(r"C:\Stage\Clients\Industries\Automotive\General_Motors\02_projects\FY27_Pipeline\apptrac-loan-origination-agentic\deliverables")
T0 = ROOT / "Tier0-Foundation"
T1 = ROOT / "Tier1-Executive"
T2 = ROOT / "Tier2-Build"
for d in (T0, T1, T2): d.mkdir(parents=True, exist_ok=True)

INK="0B0B0F"; INK2="4A4A55"; MUTE="8a8a90"; GOLD="B8892E"; TEAL="2A7F73"
VIO="6648B0"; CRIM="A22A39"; BAND="F3F3F5"; HEAD="0B0B0F"; LINE="D4D4D7"
SERIF="Fraunces"; SANS="IBM Plex Sans"; MONO="JetBrains Mono"

def cell_bg(cell, c):
    tcPr = cell._tc.get_or_add_tcPr()
    s = OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto")
    s.set(qn("w:fill"),c.lstrip("#")); tcPr.append(s)
def cell_border(cell, c=LINE):
    tcPr = cell._tc.get_or_add_tcPr()
    bs = OxmlElement("w:tcBorders")
    for e in ("top","left","bottom","right"):
        b = OxmlElement(f"w:{e}"); b.set(qn("w:val"),"single"); b.set(qn("w:sz"),"4"); b.set(qn("w:color"),c.lstrip("#"))
        bs.append(b)
    tcPr.append(bs)
def setrun(r, *, bold=False, italic=False, size=None, color=None, font=SANS):
    if bold: r.bold = True
    if italic: r.italic = True
    if size is not None: r.font.size = Pt(size)
    if color: r.font.color.rgb = RGBColor.from_string(color.lstrip("#"))
    if font:
        r.font.name = font
        rPr = r._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts"); rPr.insert(0, rFonts)
        for attr in ("w:ascii","w:hAnsi","w:eastAsia","w:cs"):
            rFonts.set(qn(attr), font)
def addbody(doc, t, *, size=11, italic=False, color=INK, sa=6, font=SANS):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(sa)
    r = p.add_run(t); setrun(r, size=size, color=color, italic=italic, font=font); return p
def addh(doc, t, *, level=1):
    sz = {1:18,2:14,3:12.5,4:11.5}
    p = doc.add_paragraph(); p.style = doc.styles[f"Heading {level}"]
    p.paragraph_format.space_before = Pt(16 if level==1 else 10); p.paragraph_format.space_after = Pt(4)
    for r in list(p.runs): r.text = ""
    r = p.add_run(t); setrun(r, bold=True, size=sz.get(level,12), color=INK, font=SERIF)
    if level == 1:
        pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom"); bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),"6"); bot.set(qn("w:space"),"4"); bot.set(qn("w:color"),GOLD)
        pBdr.append(bot); pPr.append(pBdr)
    return p
def addbullets(doc, items, *, size=11):
    for it in items:
        p = doc.add_paragraph(style="List Bullet"); r = p.add_run(it); setrun(r, size=size, color=INK, font=SANS)
def addtable(doc, header, rows, *, col_widths=None):
    n = len(header); tbl = doc.add_table(rows=len(rows)+1, cols=n); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j,h in enumerate(header):
        c = tbl.rows[0].cells[j]; cell_bg(c,HEAD); cell_border(c); c.paragraphs[0].text=""
        r = c.paragraphs[0].add_run(h); setrun(r,bold=True,size=10,color="FFFFFF", font=SANS)
        if col_widths: c.width = Inches(col_widths[j])
    for i,row in enumerate(rows, start=1):
        for j,v in enumerate(row):
            c = tbl.rows[i].cells[j]; cell_border(c)
            if i%2==0: cell_bg(c,BAND)
            c.paragraphs[0].text=""; r = c.paragraphs[0].add_run(str(v)); setrun(r,size=10,color=INK, font=SANS)
            if col_widths: c.width = Inches(col_widths[j])
    doc.add_paragraph(); return tbl
def addcallout(doc, kind, body, *, accent=GOLD, fill="F6F1E4"):
    tbl = doc.add_table(rows=1, cols=1); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = tbl.rows[0].cells[0]; cell_bg(c,fill); c.paragraphs[0].text=""
    rt = c.paragraphs[0].add_run(kind.upper()); setrun(rt,bold=True,size=9,color=accent, font=MONO)
    p = c.add_paragraph(); p.paragraph_format.space_before = Pt(3)
    r = p.add_run(body); setrun(r,size=10.5,color=INK, font=SANS); doc.add_paragraph()
def addkv(doc, pairs, *, lw=1.6, rw=4.9):
    tbl = doc.add_table(rows=len(pairs), cols=2); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,(k,v) in enumerate(pairs):
        ck,cv = tbl.rows[i].cells; ck.width=Inches(lw); cv.width=Inches(rw)
        cell_bg(ck,BAND); cell_border(ck); cell_border(cv)
        ck.paragraphs[0].text=""; cv.paragraphs[0].text=""
        rk = ck.paragraphs[0].add_run(k); setrun(rk,bold=True,size=10,color=INK, font=SANS)
        rv = cv.paragraphs[0].add_run(v); setrun(rv,size=10,color=INK, font=SANS)
    doc.add_paragraph()
def addeyebrow(doc, t):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(t.upper()); setrun(r,bold=True,size=9,color=MUTE, font=MONO)
    rPr = r._element.get_or_add_rPr(); sp = OxmlElement("w:spacing"); sp.set(qn("w:val"),"80"); rPr.append(sp)
def pb(doc):
    p = doc.add_paragraph(); p.add_run().add_break(WD_BREAK.PAGE)
def addcode(doc, t):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(t); setrun(r, size=9.5, color=INK, font=MONO)

def docinit():
    doc = Document()
    for s in doc.sections:
        s.page_width = Inches(8.5); s.page_height = Inches(11)
        s.top_margin = s.bottom_margin = Inches(0.85); s.left_margin = s.right_margin = Inches(0.95)
    doc.styles["Normal"].font.name = SANS
    doc.styles["Normal"].font.size = Pt(11)
    return doc

def docsave(doc, path):
    target = path
    try: doc.save(str(target))
    except PermissionError:
        n=2
        while True:
            alt = target.with_name(target.stem + f"-v{n}" + target.suffix)
            if not alt.exists(): target = alt; break
            n+=1
        doc.save(str(target))
    print(f"  wrote {target.relative_to(ROOT)}")


# =========================================================================
# 1. WALKTHROUGH
# =========================================================================
def build_walkthrough():
    print("Building Walkthrough...")
    doc = docinit()
    addeyebrow(doc, "APEX · AXLE-FIN-01 · GMF · APPTRAC · WALKTHROUGH")
    p = doc.add_paragraph(); r = p.add_run("APEX · AXLE-FIN-01"); setrun(r, bold=True, size=26, color=INK, font=SERIF)
    p1b = doc.add_paragraph(); r1b = p1b.add_run("AI Loan Origination & Underwriting"); setrun(r1b, bold=True, size=20, color=INK, font=SERIF)
    p2 = doc.add_paragraph(); r2 = p2.add_run("Doc Intelligence + Azure AI Foundry agents · AppTrac at GM Financial as reference build")
    setrun(r2, size=13, color=INK2, italic=True, font=SERIF)
    doc.add_paragraph()
    addbody(doc, "A walkthrough of the APEX AXLE-FIN-01 service across the value-delivery chain — Scenario → "
                 "Solution → Use Case → Service → Persona → KPI — anchored against GM Financial (GMF) as the "
                 "Wave-1 reference client. The same six-agent service pattern documented here ports across "
                 "captive auto finance (Ford Credit · Chrysler Capital · Toyota Financial · Honda Financial), "
                 "consumer lending platforms (auto · marine · powersports), and equipment finance with only "
                 "the manifest layer changing.",
            italic=True, color=INK2)
    doc.add_paragraph()
    addkv(doc, [
        ("Document type",     "APEX Service Walkthrough · Service-to-KPI"),
        ("Service",           "AXLE-FIN-01 · Auto Loan Origination & Underwriting Automation"),
        ("Microsoft stack",   "Azure AI Document Intelligence · Azure AI Foundry · Microsoft Fabric · Teams · Entra · Purview · Azure OpenAI"),
        ("Reference client",  "General Motors Financial (GMF) · AppTrac integration"),
        ("Practice",          "AXLE — Industrial & Manufacturing (automotive · captive finance sub-Practice)"),
        ("Audience",          "GMF CFO · Chief Credit Officer · CIO · Compliance · Underwriting leadership · DMTSP Account Team"),
        ("Chain links",       "Scenario · Solution · Use Case · Service · Persona · KPI"),
        ("Waves covered",     "Approach · W1 Foundation · W2 Pilot · W3 Scale & Fuse"),
        ("Companion",         "AXLE-FIN-01-GMF-OnePager.html · Six-Agents-Architecture.docx · AppTrac-SOR-Integration.docx"),
        ("Version / date",    "v1.0 · 27 Apr 2026"),
    ])
    pb(doc)

    # Section 1
    addh(doc, "Section 1 · The Reference Application at GMF", level=1)
    addbody(doc,
            "At 9:14 AM on a Tuesday in March 2027, a Chevy Silverado leaves a Houston-area Chevrolet "
            "dealership lot. The dealer's F&I manager submitted the loan application 11 minutes earlier "
            "through AppTrac, GMF's loan origination platform. In those 11 minutes: identity documents "
            "extracted, paystubs validated, vehicle title verified against the vehicle of sale, debt-to-"
            "income computed against the credit bureau pull, custom credit policy applied for the dealer's "
            "tier and the buyer's FICO band, decision composed (Approved · Tier B · 6.99% APR · 72 months · "
            "stipulations: proof of insurance), dealer notification, adverse-action evidence (none required), "
            "ECOA-compliant audit trail captured. Buyer signs, drives off. No human underwriter touched the "
            "application — yet every decision is fully attributable, reproducible, and ECOA-defensible.")
    addbody(doc,
            "This is the scenario. Everything downstream — the AXLE-FIN-01 agent architecture, the six "
            "use cases, the productized service, the persona cast, the KPI rollups — exists to answer "
            "three questions about this Tuesday morning origination at GMF:")
    addbullets(doc, [
        "Which of the ~3M annual GMF loan applications can be auto-decisioned in minutes vs. routed to a human underwriter?",
        "Who at GMF has the authority to approve material edge cases (subprime overrides · high-LTV · compensating factors), and how do they do it without leaving Microsoft Teams?",
        "What evidence survives the decision — auditable, ECOA/TILA/FCRA-grade, attributable to the specific application at the specific dealer at the specific date and time?",
    ])

    addh(doc, "Today, without AXLE-FIN-01", level=3)
    addbody(doc,
            "GMF's current underwriting flow on AppTrac mixes auto-decisioning rules with manual review. "
            "Approximately 30-40% of applications auto-decision through static rules; the remaining 60-70% "
            "queue for an underwriter. Manual review takes 4-24 hours depending on tier and queue depth. "
            "Documents arrive as scanned PDFs from dealers; underwriters open each, eyeball the data, "
            "transcribe key fields into the underwriting system, run additional checks, decide, document. "
            "Adverse-action rationales (where required by ECOA) are composed manually, often from "
            "templates. Quality varies by underwriter and queue pressure. Audit reconstruction at "
            "compliance-review time pulls from email threads, screenshots, and underwriter notes.")
    addh(doc, "Tomorrow, with AXLE-FIN-01", level=3)
    addbody(doc,
            "Six agents — Ingest, Validate, Underwrite, Decide, Approve, Commit — run in sequence. Azure "
            "AI Document Intelligence extracts every field from every document at submission time. Foundry "
            "agents apply credit policy, compute DTI/PTI/LTV, detect fraud signals, propose decisions with "
            "confidence scores. Material edge cases route through HITL via Teams Adaptive Card on the "
            "underwriter's device. Every decision emits a 14-field audit row to LEDGER bound to ECOA, "
            "TILA, and FCRA evidence requirements. Decision time compresses from 4-24 hours to ~10 minutes; "
            "auto-decision rate moves from 35% to 75% over Wave 2; manual review hours drop ~60%.")
    addcallout(doc, "Independence note",
               "AXLE-FIN-01 agent manifests, credit policy manifests, and LEDGER are GMF artifacts; "
               "Deloitte's role is advisory and implementation enablement. Before any contract, the SOW "
               "is routed through Deloitte's Independence office; if Deloitte audits GM (parent) or GMF, "
               "engagement structure is designed accordingly. The underlying APEX pattern is "
               "independence-safe by construction.")
    pb(doc)

    # Section 2
    addh(doc, "Section 2 · Why AI Underwriting, Why GMF, Why Now", level=1)
    addh(doc, "2.1 · The forces converging on captive auto finance", level=2)
    addbody(doc,
            "Five forces converged across FY25-FY26. For captive auto finance specifically, all five land "
            "on the AppTrac queue: (1) the decision-velocity gap — dealer F&I closes deals in minutes, "
            "manual underwriting takes hours; (2) the agentic-AI inflection that crossed enterprise "
            "viability for document-heavy decisioning; (3) the Doc Intelligence maturity inflection — "
            "Azure's document-AI capabilities now reliably extract from messy real-world dealer documents; "
            "(4) the ECOA/CFPB scrutiny on adverse-action accuracy and disparate-impact testing; "
            "(5) the captive-finance margin pressure as automotive OEMs (GM included) lean harder on "
            "captive lending as a differentiator.")

    addh(doc, "2.2 · Why GMF is the AXLE captive-finance reference client", level=2)
    addbody(doc,
            "Of the major captive auto-finance organizations (GMF · Ford Credit · Chrysler Capital · "
            "Toyota Financial · Honda Financial · Hyundai Capital), GMF anchors the AXLE-FIN-01 reference "
            "build for four reasons:")
    addbullets(doc, [
        "Scale — ~3M originations/year · ~$120B portfolio · ~4,500 GM dealers · enough volume to drive a real automation curve",
        "Mixed-tier book — prime + near-prime + subprime in one pipeline · the hardest underwriting surface · everything else is easier",
        "AppTrac as a known integration target — central origination platform with documented APIs · the integration cost is bounded",
        "Strong DMTSP relationship — established Microsoft footprint at GM and GMF · Azure AI Foundry strategic alignment · Azure OpenAI in-place",
    ])
    addbody(doc,
            "If AXLE-FIN-01's pattern handles GMF's mixed-tier captive-finance underwriting, it ports "
            "across the Deloitte DMTSP captive-finance portfolio with manifest-layer adaptation: Ford "
            "Credit (similar scale, different policy corpus), Chrysler Capital (Stellantis), Toyota "
            "Financial (different tier mix). Beyond auto: powersports lending (Polaris, Harley-Davidson "
            "captives), marine, RV, equipment finance — all share the architectural substrate.")

    addh(doc, "2.3 · Why this document exists", level=2)
    addbody(doc,
            "Three conversations at GMF need the same material with different emphasis. The CFO and Chief "
            "Credit Officer need a credible decision-velocity-and-margin story with ROI that survives a "
            "rigorous scrub against the existing AppTrac+rules infrastructure. The Compliance Officer and "
            "ECOA program lead need an evidence story that's auditor-grade by construction and CFPB-"
            "defensible. The CIO and CDO need a Microsoft-native architecture that fits GMF's existing "
            "Azure estate without inventing net-new platform categories. This document anchors all three.")
    pb(doc)

    # Section 3
    addh(doc, "Section 3 · Chain Link 1 · Scenario", level=1)
    addh(doc, "3.1 · GMF's situation", level=2)
    addbody(doc,
            "GMF originates approximately 3M loans and leases per year across roughly 4,500 GM dealers in "
            "the US. Application flow: dealer F&I manager captures buyer info + chosen vehicle + intended "
            "structure into the dealer's DMS (CDK Global, Reynolds & Reynolds, Dealertrack, RouteOne) → "
            "DMS submits via AppTrac integration → AppTrac normalizes, pulls credit bureau, queues for "
            "underwriting decision → decision returns to dealer typically within 4-24 hours.")
    addbody(doc,
            "Each application carries 8-15 documents on average: state-issued ID, two recent paystubs (or "
            "tax returns for self-employed), proof of residence, vehicle title or buyer's order, insurance "
            "binder, sometimes a co-applicant's documents, sometimes additional verification artifacts "
            "(WW-2 history, bank statements). All scanned PDFs from dealer F&I scanners. Quality variance "
            "is significant.")

    addh(doc, "3.2 · The current-state response", level=2)
    addbody(doc, "Today's flow:")
    addbullets(doc, [
        "T+0: dealer submits application via DMS → AppTrac integration",
        "T+5min: AppTrac runs static auto-decision rules · ~30-40% auto-decisioned (mostly clean prime applications)",
        "T+5min – T+24hr: remaining ~60-70% queue for underwriter · depth varies by hour and tier",
        "T+queue: underwriter opens application · reviews documents one by one · transcribes income and DTI fields manually",
        "T+30-90min: underwriter applies credit policy · runs additional verification · composes decision",
        "T+queue+effort: decision returns to AppTrac · dealer notification · funding pipeline if approved",
        "T+24h: ECOA adverse-action letter generated (where applicable) · template-driven · variable quality",
        "T+30 days: portfolio-level ECOA disparate-impact testing · retrospective · sampled",
    ])

    addh(doc, "3.3 · The friction, quantified", level=2)
    addbody(doc, "Measured over four representative quarters at GMF, current-state friction:")
    addtable(doc, ["Dimension","Typical value","Implication"],
             [["Application volume",                          "~3M / yr · ~250K / mo · ~12K / day",         "Sustained capacity strain on underwriting org"],
              ["Auto-decision rate today",                    "~30-40%",                                    "60-70% require manual underwriter touch"],
              ["Manual decision time",                        "4-24 hr median · longer for subprime",      "Dealer F&I waits · deal close-rate impact"],
              ["Underwriter capacity",                        "~600 underwriters · ~30 apps/day each",      "Linear cost scaling · weekend backlogs"],
              ["Document extraction time per app",            "~12-18 min manual",                          "~1.2-1.8M underwriter-hours/yr on doc transcription alone"],
              ["Document quality issues",                     "~15% of apps need re-pull from dealer",     "Cycle-time hit · dealer relationship friction"],
              ["ECOA adverse-action rationale variance",      "Template-driven · variable quality",         "CFPB exposure · disparate-impact risk"],
              ["Audit reconstruction (post-hoc)",             "~3-5 days for compliance review sample",     "Resource-intensive · evidence-quality risk"]],
             col_widths=[2.4,1.9,2.2])

    addh(doc, "3.4 · Scope of the GMF scenario", level=2)
    addbody(doc, "AXLE-FIN-01 specifically includes:")
    addbullets(doc, [
        "Document ingestion from AppTrac and dealer DMS attachments · all loan-application document types",
        "Doc Intelligence extraction · ID · paystubs · W-2s · tax returns · bank statements · titles · insurance binders · custom dealer-specific deal jackets",
        "Credit-bureau orchestration (existing GMF integration to Equifax/Experian/TransUnion · MCP-wrapped)",
        "Underwriting decisioning · DTI/PTI/LTV/FICO/custom credit policy per dealer and tier",
        "Decision composition with reasoning · stipulations (proof of insurance, GAP, etc.) · pricing tier",
        "HITL routing for material edge cases · subprime overrides · compensating-factor reviews",
        "Adverse-action narrative generation (ECOA-compliant) · approved/conditional/declined decisions",
        "Decision write-back to AppTrac · dealer notification · audit-row to LEDGER",
        "Portfolio-level disparate-impact testing surface (continuous · not retrospective)",
    ])
    addbody(doc, "It specifically excludes:")
    addbullets(doc, [
        "Loan servicing post-funding (separate scenario · could be AXLE-FIN-02)",
        "Collections / loss mitigation (separate scenario · AXLE-FIN-03 candidate)",
        "Lease residual valuation (adjacent · could fuse at W3)",
        "Dealer floorplan financing (separate captive product · AXLE-FIN-04 candidate)",
        "Vehicle-level wholesale pricing for trade-in (uses MERML.Elasticity but is a separate decisioning surface)",
    ])
    pb(doc)

    # Section 4
    addh(doc, "Section 4 · Chain Link 2 · Solution", level=1)
    addh(doc, "4.1 · AXLE-FIN-01 in one sentence", level=2)
    addcallout(doc, "Solution",
               "A six-agent loan-origination orchestration — Ingest, Validate, Underwrite, Decide, Approve, "
               "Commit — composed with a HITL-gated commercial envelope, runs on Microsoft-native "
               "infrastructure (Azure AI Document Intelligence + Azure AI Foundry + Fabric + Teams + "
               "Entra + Purview), integrates AppTrac as the system of record, and writes every decision "
               "to a 14-field audit row so ECOA evidence is generated, not collated.",
               accent=TEAL, fill="E4F1EE")

    addh(doc, "4.2 · The six agents", level=2)
    addtable(doc, ["Agent","Responsibility","Microsoft stack","HITL?"],
             [["Ingest",    "Pull docs from AppTrac · run Doc Intelligence extraction · normalize to FINML.Application",
                            "Doc Intelligence (prebuilt + custom models)",
                            "No"],
              ["Validate",  "Schema completeness · authenticity · fraud signals · cross-doc consistency",
                            "Foundry agent · Doc Intelligence confidence scores",
                            "If fraud-signal triggered"],
              ["Underwrite","Apply credit policy · DTI/PTI/LTV/FICO · stipulations · pricing tier",
                            "Foundry agent · Azure OpenAI for narrative · cap:credit-policy-resolver",
                            "If policy edge case"],
              ["Decide",    "Compose decision · approve/conditional/declined · adverse-action narrative",
                            "Foundry agent · Azure OpenAI for ECOA-compliant narrative",
                            "If subprime override or low confidence"],
              ["Approve",   "Route material decisions to Underwriter via Teams Adaptive Card · audit-row HITL log",
                            "Microsoft Teams · Entra identity",
                            "Always for material entries"],
              ["Commit",    "Write decision to AppTrac · dealer notification · adverse-action letter · audit-row closeout",
                            "AppTrac API · email/dealer-portal · LEDGER",
                            "No"]],
             col_widths=[1.0,2.1,2.0,1.4])

    addh(doc, "4.3 · Commercial envelope · GMF tier thresholds", level=2)
    addbody(doc, "The commercial envelope is the policy layer that decides 'auto-decision vs. require human "
                 "underwriter.' For AXLE-FIN-01 at GMF, tiered by buyer-tier × confidence × deal characteristics:")
    addtable(doc, ["Tier","Loan amount","Confidence","HITL?"],
             [["Tier 1 (FICO 740+ · DTI < 35%)",   "Any",                "Any",        "Auto · ZERO_TOUCH · log only"],
              ["Tier 2-3 standard (FICO 660-739)",  "≤ $40K · DTI < 45%", "≥ 0.92",     "Auto · ACK_ONLY"],
              ["Tier 2-3 standard",                  "≤ $40K · DTI < 45%", "< 0.92",     "HITL · Underwriter"],
              ["Tier 2-3 standard",                  "> $40K OR DTI ≥ 45%","Any",        "HITL · Underwriter"],
              ["Subprime (FICO < 660)",              "Any",                "Any",        "HITL · Senior Underwriter"],
              ["Any tier · adverse-action proposed", "Any",                "Any",        "HITL · ECOA review · adverse-action narrative QC"],
              ["Any tier · fraud signal detected",   "Any",                "Any",        "HITL · Fraud Specialist · escalation path"]],
             col_widths=[2.4,1.6,1.2,2.0])

    addh(doc, "4.4 · Three design commitments", level=2)
    addh(doc, "4.4.1 · Manifest-driven, not model-driven", level=3)
    addbody(doc, "Every choice that could change between dealer cohorts or tier policies — which LLM is "
                 "used for adverse-action narrative, which custom Doc Intelligence model applies for which "
                 "dealer DMS, which credit policy version applies, which underwriter has approval authority "
                 "for which tier — is declared in a Git-versioned manifest. ECOA policy versioning is part of "
                 "the manifest history; auditors and CFPB inspectors see exactly what the system was "
                 "configured to do at the time of any specific decision.")
    addh(doc, "4.4.2 · Audit row, not log line", level=3)
    addbody(doc, "Every decision — every document extraction, every UW computation, every HITL approval, "
                 "every committed loan — emits a 14-field structured row carrying trace_id, three version "
                 "stamps (manifest, policy, prompt), content-addressed input/output hashes, reasoning-trace "
                 "reference, HITL status and approver, and the specific ECOA reason codes that fired. This is "
                 "the substrate for ECOA disparate-impact testing and CFPB audit response.")
    addh(doc, "4.4.3 · HITL is first-class, not bolted on", level=3)
    addbody(doc, "The HITL gate between Decide and Commit is declared in the orchestration manifest as a "
                 "first-class primitive. Material decisions arrive as Microsoft Teams Adaptive Cards on the "
                 "underwriter's device with one-click Approve / Modify / Decline. The underwriter never logs "
                 "into another tool. The audit row captures her named identity, decision, modifications, and "
                 "reasoning.")
    pb(doc)

    # Section 5
    addh(doc, "Section 5 · Chain Link 3 · Use Cases", level=1)
    addbody(doc,
            "AXLE-FIN-01's six primary agent-level use cases compose the loan-origination scenario. Each is "
            "individually deliverable, individually measurable, and individually commercially-envelopable.")

    for title, payload in [
        ("5.1 · Use Case · Document Ingestion + Extraction",
         [("Trigger", "Application submitted via AppTrac · documents attached"),
          ("Agents",  "Ingest"),
          ("Tools",   "Doc Intelligence prebuilt models (ID · paystub · W-2 · 1099 · bank statement) + custom dealer DMS deal-jacket models"),
          ("HITL",    "None — deterministic extraction with confidence scores"),
          ("Envelope","Included in W1 foundation")]),
        ("5.2 · Use Case · Validation & Fraud Signals",
         [("Trigger", "Ingest output produced"),
          ("Agents",  "Validate"),
          ("Tools",   "Cross-doc consistency · authenticity heuristics · ID-vs-paystub matching · synthetic-identity detection"),
          ("HITL",    "Fraud signal → Fraud Specialist · routes via Teams Adaptive Card"),
          ("Envelope","~10% of W2 scope")]),
        ("5.3 · Use Case · Underwriting Decisioning",
         [("Trigger", "Validation passed"),
          ("Agents",  "Underwrite"),
          ("Tools",   "Credit-policy resolver (versioned per dealer + tier) · DTI/PTI/LTV/FICO compute · MERML.VehicleValue lookup · MERML.PricingTier"),
          ("HITL",    "Confidence < 0.92 OR subprime → escalation"),
          ("Envelope","~25% of W2 scope")]),
        ("5.4 · Use Case · Decision Composition & Adverse-Action Narrative",
         [("Trigger", "Underwriting computed"),
          ("Agents",  "Decide"),
          ("Tools",   "Azure OpenAI (ECOA-compliant narrative) · stipulation library · pricing-tier resolver"),
          ("HITL",    "ECOA review for adverse-action narratives at W2 (relax at W3)"),
          ("Envelope","~20% of W2 scope")]),
        ("5.5 · Use Case · Underwriter Approval (HITL)",
         [("Trigger", "Material decision per commercial envelope"),
          ("Agents",  "None (approvals MCP orchestrates the human)"),
          ("Tools",   "Microsoft Teams Adaptive Card · Entra identity binding · cap:ledger-write"),
          ("HITL",    "Always for material entries"),
          ("Envelope","~25% of W2 scope")]),
        ("5.6 · Use Case · Commit + Outcome Attribution",
         [("Trigger", "Approval recorded · or auto-decision per envelope"),
          ("Agents",  "Commit"),
          ("Tools",   "AppTrac decision-write-back · dealer-portal notification · ECOA adverse-action letter generator · LEDGER closeout"),
          ("HITL",    "None — write-only · consumed by Compliance + Risk"),
          ("Envelope","~20% of W2 scope")]),
    ]:
        addh(doc, title, level=3)
        addkv(doc, payload, lw=1.2, rw=5.0)

    pb(doc)

    # Section 6
    addh(doc, "Section 6 · Chain Link 4 · Services", level=1)
    addh(doc, "6.1 · Anchor service", level=2)
    addh(doc, "AXLE-FIN-01 · Auto Loan Origination & Underwriting Automation", level=3)
    addbody(doc, "The foundational underwriting-automation service. Owns the six-agent orchestration from "
                 "Ingest through Commit. Primary persona: GMF Underwriter. SLO: decision-loop time 4-24 hr → "
                 "10 min (W2) → 4 min (W3). Auto-decision rate 35% → 75% (W3). Commercial envelope tiered by "
                 "application volume × tier mix × HITL policy depth.")

    addh(doc, "6.2 · Adjacent / candidate services", level=2)
    addbullets(doc, [
        "AXLE-FIN-02 · Loan Servicing & Hardship Decisioning — extends the agent pattern to in-life servicing decisions (deferrals · modifications · payoff quotes)",
        "AXLE-FIN-03 · Collections & Loss Mitigation — late-stage decisioning · skip-trace orchestration · charge-off analytics",
        "AXLE-FIN-04 · Dealer Floorplan Risk — wholesale lending to dealers · inventory aging · concentration risk",
        "AXLE-FIN-05 · Lease Residual & End-of-Term Valuation — post-funding decisioning surface · MERML.VehicleValue fusion",
    ])

    addh(doc, "6.3 · GMF tenant manifest subscription", level=2)
    addbody(doc, "The GMF tenant manifest subscribes to the service:")
    addcode(doc, '{ "tenant_id": "gmf", "subscriptions": [\n'
                 '  { "service_id": "AXLE-FIN-01", "service_version": "v1.0.0" }\n'
                 ']}')
    pb(doc)

    # Section 7 - Personas
    addh(doc, "Section 7 · Chain Link 5 · Personas", level=1)
    addh(doc, "7.1 · Primary persona", level=2)
    addcallout(doc, "PRIMARY · Underwriter · Tier 2-3 originations",
               "Owns the HITL approval path for material decisions. Approves via Teams Adaptive Card; "
               "decisions write to LEDGER with named Entra identity. Reports to Senior Credit Officer. "
               "Reviews ~30-50 HITL decisions/day at W3 steady-state (vs. ~30-50 full applications/day "
               "today). Measure of success: HITL queue time, decision quality, ECOA compliance.",
               accent="E59A3E")

    addh(doc, "7.2 · Secondary personas", level=2)
    addtable(doc, ["Role","Function in AXLE-FIN-01"],
             [["Senior Credit Officer (SCO)",            "Escalation path for subprime overrides · adverse-action quality control · policy-tuning input"],
              ["Compliance Officer / ECOA Program Lead", "Consumes audit-row stream for ECOA disparate-impact testing · CFPB exam response"],
              ["Fraud Specialist",                       "Receives Adaptive Card for fraud-signal triggered cases · synthetic-identity investigation"],
              ["Funding Specialist",                     "Post-approval pipeline · stipulation tracking · dealer follow-through"],
              ["Dealer F&I Manager",                     "External persona · consumes decision via dealer DMS · time-to-decision is their primary KPI"],
              ["Risk Officer (portfolio)",               "Power BI dashboards · portfolio-level KPIs · early-warning indicators · LEDGER drill-through"],
              ["CIO · Application Operations",           "Owns AppTrac integration · system-of-record stability"],
              ["CDO · Data Governance",                  "Owns Doc Intelligence model lifecycle · Purview classification · PII handling"]],
             col_widths=[2.4,3.9])

    addh(doc, "7.3 · Persona-to-plane mapping", level=2)
    addtable(doc, ["Persona","Plane(s) touched","Surface"],
             [["Underwriter (primary)",   "Experience · Decision",        "Microsoft Teams Adaptive Card (mobile + desktop)"],
              ["Senior Credit Officer",    "Experience · Decision · Ledger","Adaptive Card escalation · LEDGER query for policy review"],
              ["Compliance Officer",       "Ledger",                       "Power BI compliance dashboard · audit-row export · CFPB-format export"],
              ["Fraud Specialist",         "Experience · Decision",        "Adaptive Card on fraud-signal events"],
              ["Funding Specialist",       "Experience · Integration",     "AppTrac post-approval pipeline view"],
              ["Dealer F&I Manager",       "Integration",                  "Dealer DMS notification (CDK · Reynolds · Dealertrack · RouteOne)"],
              ["Risk Officer",             "Ledger · Experience",          "Power BI portfolio scorecard · per-cohort drill-through"]],
             col_widths=[2.0,1.8,2.5])
    pb(doc)

    # Section 8 - KPIs
    addh(doc, "Section 8 · Chain Link 6 · KPIs", level=1)
    addh(doc, "8.1 · Per-cycle KPIs", level=2)
    addtable(doc, ["KPI","Measurement","Today","Target W2","Target W3"],
             [["Decision time",                       "Submit-to-decision",                 "4-24 hr",     "≤ 30 min",      "≤ 10 min"],
              ["Auto-decision rate",                  "% no human touch",                   "30-40%",      "55%",            "75%"],
              ["Manual review hours / day",           "Underwriter time on doc transcription","~20K hrs",   "-40%",           "-60%"],
              ["Decision-quality (override rate)",   "% Decide outputs SCO-overridden",     "—",           "≤ 5%",           "≤ 2%"],
              ["ECOA adverse-action accuracy",       "% narratives passing QC",             "Variable",    "≥ 98%",          "≥ 99.5%"],
              ["Audit-row coverage",                  "% material decisions attributed",    "—",           "≥ 99%",          "≥ 99.5%"],
              ["Funding rate (approval → booked)",    "% approvals actually booked",        "Baseline",    "+1pt",           "+2-3pt"],
              ["Dealer F&I NPS (decision speed)",     "Per-dealer survey",                  "Baseline",    "+10pt",          "+15-20pt"]],
             col_widths=[1.9,1.8,1.0,1.1,1.1])

    addh(doc, "8.2 · Derived KPIs for GMF commercial cases", level=2)
    addh(doc, "8.2.1 · Capacity Reclaimed", level=3)
    addbody(doc, "~600 underwriters × ~5 hrs/day on doc transcription × 250 working days × $85/hr loaded = "
                 "~$64M/yr in capacity available for redeployment. At W3 steady-state, ~60% of that "
                 "(~$38M/yr) converts directly to capacity for portfolio growth, complex cases, dealer "
                 "relationship management, and exception handling.")
    addh(doc, "8.2.2 · Dealer Close-Rate Lift", level=3)
    addbody(doc, "~3M applications × ~88% close-rate baseline × ~+2-3pt close-rate lift from faster "
                 "decisioning × ~$1,200 avg captive-finance margin per loan = ~$70-105M/yr revenue uplift "
                 "potential. Sensitivity is principally on the close-rate lift estimate; bench testing in "
                 "W1 shadow-run validates against historical lost-deal data.")
    addh(doc, "8.2.3 · ECOA Compliance Posture", level=3)
    addbody(doc, "Continuous disparate-impact testing replaces quarterly retrospective sampling. CFPB "
                 "exam-response time drops from 4-6 weeks to under 1 week. Adverse-action narrative quality "
                 "improves measurably; CFPB-finding risk reduces. Hard to monetize directly but informs a "
                 "$XX-million risk-adjustment in the business case (D&O · regulatory-fine avoidance · "
                 "audit-fee reduction).")
    addh(doc, "8.2.4 · Audit-Trail Completeness", level=3)
    addbody(doc, "100% of material decisions audit-row attributed at W3. Direct answer to CFPB exam "
                 "questions and Internal Audit / SOX-adjacent control validation. Eliminates the 4-6 week "
                 "retrospective audit reconstruction current state.")

    addh(doc, "8.3 · KPI attribution in the audit row", level=2)
    addcode(doc, '{\n'
                 '  "trace_id": "trc_2027-03-09_0914_gmf_dealer-12834_app-7298321",\n'
                 '  "scenario_id": "AXLE-FIN-01-loan-origination",\n'
                 '  "service_id": "AXLE-FIN-01",\n'
                 '  "tenant_id": "gmf",\n'
                 '  "event_kind": "axle_fin_01_closeout",\n'
                 '  "event_ts": "2027-03-09T13:25:18Z",\n'
                 '  "manifest_version": "axle-fin-01-orchestration-v1.0.0",\n'
                 '  "policy_version": "gmf-credit-policy-v3.4.0",\n'
                 '  "kpis": {\n'
                 '    "decision_time_min": 11.4,\n'
                 '    "auto_decisioned": true,\n'
                 '    "buyer_tier": "B",\n'
                 '    "fico_at_decision": 712,\n'
                 '    "decision_outcome": "APPROVED",\n'
                 '    "ecoa_adverse_action": false,\n'
                 '    "audit_row_complete": true\n'
                 '  }\n'
                 "}")
    pb(doc)

    # Section 9 - Three Waves
    addh(doc, "Section 9 · The Three Waves", level=1)
    addh(doc, "9.1 · Wave 1 · Foundation — Single-tier shadow-run", level=2)
    addtable(doc, ["W1 element","What it looks like for GMF"],
             [["Scope",            "Tier-1 prime applications only · single dealer cluster (Houston metro · ~120 dealers) · shadow-run only"],
              ["Agents live",      "3 of 6 (Ingest · Validate · Underwrite) · Decide/Approve/Commit shadow"],
              ["Use cases",        "5.1 + 5.2 + partial 5.3 · read-only · no automated commit"],
              ["Personas onboarded","UW + SCO + ECOA Program Lead + 1 dealer pilot group"],
              ["Schemas",          "FINML.Application · FINML.Document · DIM.Dealer · DIM.Tier"],
              ["Doc Intelligence", "Prebuilt models (ID/paystub/W-2/bank statement) live · 1 custom model in training (Houston dealer DMS deal jacket)"],
              ["Audit row",        "Composed end-to-end on shadow data; replay against historical 4 quarters"],
              ["Adaptive Card",    "Underwriter pilot · 5 named identities · Entra binding live"],
              ["LEDGER plane",     "Stood up · pilot rows ingested · query path verified"],
              ["Duration",         "5-8 months"],
              ["Effort",           "~50 story points"],
              ["Exit gate",        "SCO + ECOA Program Lead sign-off on shadow-run accuracy + 30-day live shadow comparison"]],
             col_widths=[1.7,4.6])

    addh(doc, "9.2 · Wave 2 · Pilot — Houston cluster live + Tier 2 expansion", level=2)
    addtable(doc, ["W2 element","What it looks like for GMF"],
             [["Scope",            "Houston cluster live · Tier-1 + Tier-2 prime · ~250K applications over 90 days"],
              ["Agents live",      "All 6 (full orchestration with HITL routing per envelope)"],
              ["Use cases",        "5.1 through 5.6 fully exercised"],
              ["Doc Intelligence", "All prebuilt models live · 4 custom dealer-DMS models live (top 4 dealer DMS platforms)"],
              ["Schemas",          "Full FINML.* live · DIM.Dealer + DIM.Bureau + DIM.Vehicle"],
              ["HITL volume",      "Tier-tuned to ~5-10% of decisions at steady state · Underwriter daily volume ~30-50 HITL cases"],
              ["Power BI",         "Decision-time dashboard · auto-decision rate trending · ECOA disparate-impact continuous testing"],
              ["ECOA evidence",    "First quarterly compliance packet exercised with Compliance Officer + Internal Audit"],
              ["Duration",         "8-12 months"],
              ["Effort",           "~45 story points + ongoing enablement"],
              ["Exit gate",        "Steering Committee approval to scale · ECOA continuous-monitoring signoff · CFPB exam-readiness review"]],
             col_widths=[1.7,4.6])

    addh(doc, "9.3 · Wave 3 · Scale & Fuse — National + subprime + lease + adjacent-product fusion", level=2)
    addtable(doc, ["W3 element","What it looks like for GMF"],
             [["Scope",            "All ~4,500 GM dealers · all tiers including subprime · leases included · adjacent products"],
              ["Agents live",      "All 6 with banner + tier + product-aware policy manifests"],
              ["Doc Intelligence", "Full custom-model coverage across major dealer DMS platforms · co-applicant + business-purpose docs"],
              ["Cross-product fusion","AXLE-FIN-02 (servicing) · AXLE-FIN-04 (dealer floorplan) · MERML.VehicleValue fusion"],
              ["Personas",         "Underwriting org trained at scale · CFPB consuming evidence directly · dealer NPS measurement live"],
              ["KPI rollup",       "Enterprise-wide ~$70-105M annualized close-rate uplift · ~$38M+ capacity reclaimed at run-rate"],
              ["Decision time",    "≤ 4 min across all tiers and products"],
              ["Duration",         "12-18 months core · ongoing fusion"],
              ["Effort",           "~70 story points + ongoing scenario-fusion"],
              ["Exit gate",        "Steady-state operating review · expansion into adjacent products"]],
             col_widths=[1.7,4.6])
    pb(doc)

    # Section 10
    addh(doc, "Section 10 · Commercial Envelope", level=1)
    addtable(doc, ["Tier","Duration","Cost envelope","Effort","Primary deliverable"],
             [["Approach",          "3-5 wks",   "$200K-$500K",   "22 SP",  "Discovery · dealer scoping · SOW · Independence routing"],
              ["Wave 1 Foundation", "5-8 mos",   "$1.2M-$3.5M",   "50 SP",  "Houston shadow-run · 3-agent · audit-row e2e · 1 custom DocInt model"],
              ["Wave 2 Pilot",      "8-12 mos",  "$5M-$12M",      "45 SP",  "Houston live · 6 agents · 90-day operating run · ECOA quarterly packet"],
              ["Wave 3 Scale&Fuse", "12-18 mos", "$18M-$45M",     "70+ SP", "All ~4,500 dealers · subprime · leases · adjacent-product fusion"]],
             col_widths=[1.4,1.2,1.3,0.8,2.5])

    addh(doc, "10.1 · Envelope vs. durable value at GMF scale", level=2)
    addbody(doc, "A national W3 enterprise footprint at GMF produces ~$38M/yr in capacity reclaimed, plus "
                 "~$70-105M/yr in close-rate-driven revenue uplift, plus regulatory-risk-adjustment value "
                 "(harder to monetize but routinely appears in CFPB / ECOA program risk modeling). The "
                 "$18M-$45M W3 envelope spans rapid-payback (under 12 months at the low end) to strategic-"
                 "investment (24-36 months at the high end with greater fusion scope).")

    addh(doc, "10.2 · Independence posture in GMF commercial conversations", level=2)
    addcallout(doc, "Independence",
               "Before contracting, confirm Deloitte's external-audit relationship with GM (parent) and GMF. "
               "If audit-relationship-constrained, structure SOW so manifests, policies, and LEDGER are GMF "
               "artifacts; Deloitte's role advisory and implementation enablement. Route the SOW through "
               "Deloitte's Independence office. The underlying APEX pattern is independence-safe by design.")
    pb(doc)

    # Section 11
    addh(doc, "Section 11 · Reference Wave 1 Implementation", level=1)
    addh(doc, "11.1 · What the W1 prototype demonstrates", level=2)
    addtable(doc, ["Capability","What it demonstrates"],
             [["Synthetic GMF application population", "100K synthetic applications across 8 dealer cohorts · 4 historical quarters of realistic patterns"],
              ["Ingest agent (live)",                  "Doc Intelligence extraction · prebuilt + 1 custom dealer-DMS model · ~99% schema fidelity"],
              ["Validate agent (live)",                "Cross-doc consistency · synthetic-identity heuristics · authenticity scoring"],
              ["Underwrite agent (live)",              "GMF credit policy v3.4 · DTI/PTI/LTV compute · MERML.VehicleValue lookup"],
              ["Adaptive Card flow",                   "Mock Underwriter approval · Entra identity resolution"],
              ["LEDGER query surface",                 "Drill-through from dashboard to specific trace_id"],
              ["ECOA evidence packet",                 "Quarterly evidence packet ZIP · disparate-impact testing surface"],
              ["Manifest versioning",                  "Git-backed manifests · PR-merge discipline"],
              ["Shadow-run replay",                    "4 historical quarters of synthetic application data"],
              ["Demo console",                         "3 canonical demo paths · 3-5 min each · narrated for GMF stakeholders"],
              ["Self-host install",                    "Local Docker + Ollama (offline-capable for sensitive demo environments)"]],
             col_widths=[2.2,4.1])

    addh(doc, "11.2 · Three canonical demo paths", level=2)
    addbody(doc, "Demo Path A · The application loop. Pick a synthetic Houston-area Tier-2 application. "
                 "Watch Ingest extract documents in 30 seconds; watch Validate run cross-doc consistency; "
                 "watch Underwrite apply credit policy with reasoning trace; watch Decide compose decision "
                 "with stipulations; watch Adaptive Card fire to the Underwriter's device with full context; "
                 "click Approve; watch the audit row write.")
    addbody(doc, "Demo Path B · The decision-time trending. Open the Power BI decision-time dashboard. Show "
                 "the per-tier decision-time distribution over 12 weeks of shadow-run data. Drill through "
                 "into a single high-latency cell — get the contributing applications. Drill into one — see "
                 "the full reasoning trace, the version stamps, the HITL approver's identity.")
    addbody(doc, "Demo Path C · The ECOA evidence trail. Pick a synthetic adverse-action application. Open "
                 "the audit-row stream for that trace_id. Show the full chain of custody from document "
                 "submission through extraction, validation, underwriting, decision, narrative, approval, "
                 "and adverse-action letter. Show the manifest version and policy version stamps. Show the "
                 "ECOA reason codes captured at decision time.")

    addh(doc, "11.3 · Prototype technology stack", level=2)
    addtable(doc, ["Layer","Stack"],
             [["Frontend",            "React + Vite · Tailwind · Power BI Embedded"],
              ["Backend",             "Node.js (Fastify) · pg-pool"],
              ["Persistence",         "PostgreSQL (LEDGER + state)"],
              ["Identity",            "Microsoft Entra ID (mock binding for prototype)"],
              ["Doc extraction",      "Azure AI Document Intelligence · prebuilt + 1 custom model"],
              ["Inference",           "Local Ollama (offline) · Azure OpenAI for online demos · Foundry agent runtime"],
              ["Orchestration",       "Lightweight DAG runner + manifest loader (CAF-shaped contract)"],
              ["Observability",       "OpenTelemetry traces · LangFuse-equivalent local span store"],
              ["Schema validation",   "Zod (FINML.* shapes)"],
              ["Demo packaging",      "Docker Compose · 'docker compose up' single-command boot"],
              ["Synthetic data",      "100K synthetic applications · 4 historical quarters · 8 dealer cohorts"]],
             col_widths=[1.6,4.7])

    addh(doc, "11.4 · Prototype limitations and next steps", level=2)
    addbullets(doc, [
        "Connect to real AppTrac · synthetic data only · real integration is W2 scope",
        "Run on Azure · everything is local (Docker + Ollama on a laptop) · Azure migration is W2",
        "Custom DocInt models for all major DMS platforms · 1 in W1 · 4 in W2 · all in W3",
        "Real Entra authentication · hardcoded persona for demo · Entra integration is W2",
        "True CFPB-ready evidence packet · audit timeline view demonstrates substrate · packet assembly is W2",
    ])
    pb(doc)

    addh(doc, "Appendix A · Glossary", level=1)
    addtable(doc, ["Term","Definition"],
             [["AppTrac",         "GMF's loan origination tracking platform · system of record for application + decision data"],
              ["DTI",             "Debt-to-Income ratio · core underwriting metric"],
              ["PTI",             "Payment-to-Income ratio · auto-finance specific metric"],
              ["LTV",             "Loan-to-Value ratio · vehicle valuation comparison"],
              ["FICO",            "Credit-bureau score · primary tier-determination input"],
              ["ECOA",            "Equal Credit Opportunity Act · regulates underwriting + adverse-action"],
              ["TILA",            "Truth in Lending Act · disclosure requirements"],
              ["FCRA",            "Fair Credit Reporting Act · governs credit-bureau usage"],
              ["GLBA",            "Gramm-Leach-Bliley Act · privacy of consumer financial info"],
              ["Adverse Action",  "ECOA-required notice when application denied or modified · narrative + reason codes"],
              ["F&I",             "Finance & Insurance · dealer department that originates loan applications"],
              ["DMS",             "Dealer Management System · CDK Global · Reynolds & Reynolds · Dealertrack · RouteOne"],
              ["Doc Intelligence","Azure AI Document Intelligence (formerly Form Recognizer)"],
              ["Foundry",         "Azure AI Foundry · agent runtime · capability registry · model serving"],
              ["FINML",           "APEX canonical schema family for Financial Services objects (Application · Document · Decision · etc.)"],
              ["LEDGER",          "APEX 14-field audit-row store · primary substrate for ECOA evidence and KPI attribution"],
              ["Stipulation",     "Condition attached to loan approval (proof of insurance · GAP · etc.) · cleared before funding"]],
             col_widths=[1.7,4.6])

    addbody(doc, "Independence posture preserved.", italic=True, color=MUTE, size=9)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— End of walkthrough —"); setrun(r, italic=True, size=10, color=MUTE, font=SERIF)
    docsave(doc, T0 / "APEX-AXLE-FIN-01-Walkthrough.docx")


# =========================================================================
# 2. ONE-PAGER (HTML)
# =========================================================================

ONE_PAGER = """<!DOCTYPE html>
<html lang="en"><head>
<meta charset="UTF-8" /><meta name="viewport" content="width=device-width, initial-scale=1.0" />
<title>APEX · AXLE-FIN-01 · GMF · One-Pager</title>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
<link href="https://fonts.googleapis.com/css2?family=Fraunces:ital,opsz,wght@0,9..144,400;0,9..144,500;0,9..144,600&family=IBM+Plex+Sans:wght@0,300;400;500;600;700&family=JetBrains+Mono:wght@400;500;600&display=swap" rel="stylesheet">
<style>
:root{--bg-0:#07090F;--bg-1:#0C1119;--bg-2:#141B27;--bg-3:#1C2433;--line:rgba(255,255,255,.10);--line-strong:rgba(255,255,255,.20);--ink-0:#F4EFE5;--ink-1:#D9D2C3;--ink-2:#9A9383;--ink-3:#6B6757;--gold:#E3B657;--teal:#3DD9C4;--crimson:#D94B5A;--violet:#9B7BE0;--sky:#6FB6E5;--amber:#E8A33D;--ember:#E87B3C}
body.light{--bg-0:#F5F1E8;--bg-1:#EEE7D7;--bg-2:#E5DCC7;--bg-3:#D9CFB6;--line:rgba(20,14,4,.10);--line-strong:rgba(20,14,4,.22);--ink-0:#1A1508;--ink-1:#3D3420;--ink-2:#6B5F45;--ink-3:#8E8367;--gold:#8E6A1E;--teal:#147A6C;--crimson:#A22A39;--violet:#6548B0;--sky:#3E7FAE;--amber:#B8791F;--ember:#B0551B}
*{box-sizing:border-box}html,body{margin:0;padding:0}
body{background:var(--bg-0);color:var(--ink-0);font-family:"IBM Plex Sans",sans-serif;font-size:14.5px;line-height:1.55;min-height:100vh;transition:background .5s,color .5s}
body::before{content:"";position:fixed;inset:0;background:radial-gradient(1200px 600px at 10% -10%,rgba(232,163,61,.08),transparent 60%),radial-gradient(900px 500px at 95% 110%,rgba(155,123,224,.06),transparent 60%);pointer-events:none;z-index:0}
.theme-toggle{position:fixed;top:24px;right:24px;z-index:100;display:flex;align-items:center;gap:8px;background:var(--bg-1);border:1px solid var(--line-strong);border-radius:999px;padding:8px 14px;cursor:pointer;font-family:"JetBrains Mono",monospace;font-size:11px;letter-spacing:.12em;text-transform:uppercase;color:var(--ink-1)}
.page{position:relative;z-index:1;max-width:1200px;margin:0 auto;padding:56px 48px 40px}
.eyebrow{display:inline-block;font-family:"JetBrains Mono",monospace;font-size:11px;letter-spacing:.18em;text-transform:uppercase;color:var(--gold);border:1px solid var(--line-strong);padding:6px 14px;border-radius:999px;margin-bottom:22px}
h1{font-family:"Fraunces",serif;font-weight:500;font-size:clamp(36px,5vw,56px);line-height:1.04;letter-spacing:-0.02em;margin:0 0 14px}
h1 em{font-style:italic;font-weight:400;color:var(--amber)}
.subhead{font-size:17px;color:var(--ink-1);max-width:880px;margin:0 0 28px}

.headline-band{display:grid;grid-template-columns:2fr 1fr 1fr;gap:18px;margin:24px 0 32px}
@media (max-width:800px){.headline-band{grid-template-columns:1fr}}
.headline-card{background:var(--bg-1);border:1px solid var(--line);border-left:4px solid var(--accent);border-radius:12px;padding:22px 24px}
.headline-card .label{font-family:"JetBrains Mono",monospace;font-size:10.5px;letter-spacing:.14em;text-transform:uppercase;color:var(--ink-3);margin-bottom:8px}
.headline-card .value{font-family:"Fraunces",serif;font-size:36px;font-weight:500;line-height:1;color:var(--accent);margin-bottom:6px}
.headline-card .value.big{font-size:54px}
.headline-card .sub{font-size:13px;color:var(--ink-2)}

.three-col{display:grid;grid-template-columns:repeat(3,1fr);gap:18px;margin-bottom:32px}
@media (max-width:900px){.three-col{grid-template-columns:1fr}}
.col-card{background:var(--bg-1);border:1px solid var(--line);border-top:3px solid var(--accent);border-radius:12px;padding:22px}
.col-card h3{font-family:"Fraunces",serif;font-size:20px;font-weight:500;margin:0 0 4px;color:var(--ink-0)}
.col-card .col-eye{font-family:"JetBrains Mono",monospace;font-size:10.5px;letter-spacing:.14em;text-transform:uppercase;color:var(--accent);margin-bottom:12px}
.col-card p{color:var(--ink-1);font-size:13.5px;margin:0 0 10px}
.col-card ul{margin:8px 0 0;padding-left:18px;color:var(--ink-1);font-size:13px}
.col-card ul li{margin-bottom:4px}

.agent-row{display:grid;grid-template-columns:repeat(6,1fr);gap:10px;margin-top:24px}
@media (max-width:1000px){.agent-row{grid-template-columns:1fr 1fr}}
.agent-pill{background:var(--bg-1);border:1px solid var(--line);border-top:3px solid var(--accent);border-radius:8px;padding:14px;text-align:center}
.agent-pill .num{font-family:"JetBrains Mono",monospace;font-size:10.5px;color:var(--accent);margin-bottom:6px}
.agent-pill h4{font-family:"Fraunces",serif;font-size:15px;font-weight:500;margin:0 0 4px}
.agent-pill p{font-size:11.5px;color:var(--ink-2);margin:0}

.wave-strip{background:var(--bg-1);border:1px solid var(--line);border-radius:12px;padding:22px 24px;margin-top:24px}
.wave-strip h3{font-family:"JetBrains Mono",monospace;font-size:11px;letter-spacing:.15em;text-transform:uppercase;color:var(--ink-3);margin:0 0 14px}
.wave-row{display:grid;grid-template-columns:repeat(4,1fr);gap:12px}
@media (max-width:900px){.wave-row{grid-template-columns:1fr 1fr}}
.wave-item{border-left:3px solid var(--w-color);padding:10px 14px;background:var(--bg-2);border-radius:4px}
.wave-item .w-label{font-family:"JetBrains Mono",monospace;font-size:10px;letter-spacing:.14em;text-transform:uppercase;color:var(--w-color);margin-bottom:4px}
.wave-item .w-name{font-family:"Fraunces",serif;font-size:15px;font-weight:500;margin-bottom:4px}
.wave-item .w-detail{font-size:12px;color:var(--ink-2);line-height:1.45}
.wave-item .w-env{margin-top:8px;font-family:"JetBrains Mono",monospace;font-size:11px;color:var(--gold)}

.footer-band{background:var(--bg-1);border:1px solid var(--line);border-radius:12px;padding:22px 28px;display:grid;grid-template-columns:2fr 1fr;gap:18px;align-items:center;margin-top:24px}
@media (max-width:800px){.footer-band{grid-template-columns:1fr}}
.footer-band .left p{margin:0 0 4px;color:var(--ink-1)}
.footer-band .left .signoff{font-family:"Fraunces",serif;font-size:18px;color:var(--ink-0);margin-top:8px}
.footer-band .right{font-family:"JetBrains Mono",monospace;font-size:11px;color:var(--ink-2);text-align:right}
.footer-band .right b{color:var(--ink-0);display:block;margin-bottom:4px}
</style></head>
<body>
<div class="theme-toggle" onclick="toggleTheme()"><span id="themeIcon">☀</span><span id="themeLabel">Light</span></div>

<div class="page">
  <div class="eyebrow">APEX · AXLE-FIN-01 · GMF · APPTRAC · ONE-PAGER</div>
  <h1>AI loan origination at <em>3M-app/year</em> scale.</h1>
  <p class="subhead">APEX AXLE-FIN-01 — Auto Loan Origination & Underwriting Automation — anchored on GM Financial as the Wave-1 reference client. Doc Intelligence + Azure AI Foundry agents on top of AppTrac. Six agents collapse 4-24 hour underwriting decisions to ~10 minutes; auto-decision rate moves from 35% to 75%; ECOA evidence is generated, not collated.</p>

  <div class="headline-band">
    <div class="headline-card" style="--accent:var(--gold)">
      <div class="label">W3 annualized durable value · GMF</div>
      <div class="value big">~$110M</div>
      <div class="sub">~$70-105M close-rate uplift + ~$38M capacity reclaimed at run-rate</div>
    </div>
    <div class="headline-card" style="--accent:var(--teal)">
      <div class="label">Decision time</div>
      <div class="value">4-24 hr → 10 min</div>
      <div class="sub">≤ 4 min at W3 · across all tiers</div>
    </div>
    <div class="headline-card" style="--accent:var(--violet)">
      <div class="label">Auto-decision rate</div>
      <div class="value">35% → 75%</div>
      <div class="sub">Audit-row attributed · ECOA-defensible</div>
    </div>
  </div>

  <div class="three-col">
    <div class="col-card" style="--accent:var(--crimson)">
      <div class="col-eye">Why now</div>
      <h3>Three pressures converging</h3>
      <ul>
        <li><b>Doc Intelligence maturity</b> — Azure now reliably extracts from messy dealer documents</li>
        <li><b>CFPB scrutiny</b> on adverse-action accuracy + disparate-impact testing</li>
        <li><b>Captive-finance margin pressure</b> — speed is competitive differentiation in F&I</li>
      </ul>
    </div>
    <div class="col-card" style="--accent:var(--amber)">
      <div class="col-eye">What we deliver</div>
      <h3>Six-agent orchestration</h3>
      <ul>
        <li><b>Ingest</b> — Doc Intelligence extraction</li>
        <li><b>Validate</b> — fraud + consistency</li>
        <li><b>Underwrite</b> — credit policy + bureau</li>
        <li><b>Decide</b> — composition + adverse-action narrative</li>
        <li><b>Approve</b> — HITL via Teams Adaptive Card</li>
        <li><b>Commit</b> — AppTrac write-back + LEDGER</li>
      </ul>
    </div>
    <div class="col-card" style="--accent:var(--teal)">
      <div class="col-eye">How it runs</div>
      <h3>Microsoft-native</h3>
      <ul>
        <li><b>Azure AI Document Intelligence</b> — prebuilt + custom dealer DMS models</li>
        <li><b>Azure AI Foundry</b> — agent runtime · capability tags</li>
        <li><b>Microsoft Fabric</b> — FINML Bronze + Silver + LEDGER</li>
        <li><b>Microsoft Teams + Entra</b> — HITL Adaptive Card · named-identity binding</li>
        <li><b>Microsoft Purview</b> — PII classification · ECOA evidence lineage</li>
      </ul>
    </div>
  </div>

  <div class="agent-row">
    <div class="agent-pill" style="--accent:var(--sky)"><div class="num">01</div><h4>Ingest</h4><p>Doc Intelligence extraction</p></div>
    <div class="agent-pill" style="--accent:var(--violet)"><div class="num">02</div><h4>Validate</h4><p>Cross-doc + fraud signals</p></div>
    <div class="agent-pill" style="--accent:var(--teal)"><div class="num">03</div><h4>Underwrite</h4><p>DTI · PTI · LTV · FICO</p></div>
    <div class="agent-pill" style="--accent:var(--gold)"><div class="num">04</div><h4>Decide</h4><p>Decision + ECOA narrative</p></div>
    <div class="agent-pill" style="--accent:var(--ember)"><div class="num">05</div><h4>Approve</h4><p>HITL via Teams</p></div>
    <div class="agent-pill" style="--accent:var(--crimson)"><div class="num">06</div><h4>Commit</h4><p>AppTrac + LEDGER</p></div>
  </div>

  <div class="wave-strip">
    <h3>Three-wave path · Approach → W1 Foundation → W2 Pilot → W3 Scale & Fuse</h3>
    <div class="wave-row">
      <div class="wave-item" style="--w-color:var(--ink-2)"><div class="w-label">Approach</div><div class="w-name">Discovery + SOW</div><div class="w-detail">3-5 wks · dealer scoping · Independence routing</div><div class="w-env">$200K-$500K</div></div>
      <div class="wave-item" style="--w-color:var(--sky)"><div class="w-label">Wave 1</div><div class="w-name">Foundation</div><div class="w-detail">5-8 mo · Houston cluster ~120 dealers · Tier-1 prime · 3 agents · shadow-run</div><div class="w-env">$1.2M-$3.5M</div></div>
      <div class="wave-item" style="--w-color:var(--teal)"><div class="w-label">Wave 2</div><div class="w-name">Pilot</div><div class="w-detail">8-12 mo · Houston live · Tier-1+2 · 6 agents · 90-day run · ECOA packet</div><div class="w-env">$5M-$12M</div></div>
      <div class="wave-item" style="--w-color:var(--gold)"><div class="w-label">Wave 3</div><div class="w-name">Scale & Fuse</div><div class="w-detail">12-18 mo · ~4,500 dealers · all tiers · subprime · leases · adjacent products</div><div class="w-env">$18M-$45M</div></div>
    </div>
  </div>

  <div class="footer-band">
    <div class="left">
      <p style="font-family:'Fraunces',serif;font-size:16px;color:var(--ink-0)">Companion artifacts:</p>
      <p style="font-family:'JetBrains Mono',monospace;font-size:12px;color:var(--ink-2)">Walkthrough · Six-Agents Architecture · AppTrac SOR Integration</p>
      <p class="signoff">— Deloitte Microsoft Technology Solutions & Practices</p>
    </div>
    <div class="right">
      <b>Independence-safe by design</b>
      All AXLE-FIN-01 commercial materials use "Deloitte's Microsoft practice," "DMTSP," "Microsoft platform capabilities," "Microsoft-native deployment." The terms "partner," "alliance," and "strategic alliance" do not appear.
      <br><br>v1.0 · 27 Apr 2026
    </div>
  </div>
</div>

<script>
function toggleTheme(){const b=document.body,i=document.getElementById('themeIcon'),l=document.getElementById('themeLabel');b.classList.toggle('light');if(b.classList.contains('light')){i.textContent='☾';l.textContent='Dark';localStorage.setItem('apex-theme','light')}else{i.textContent='☀';l.textContent='Light';localStorage.setItem('apex-theme','dark')}}
(function(){const s=localStorage.getItem('apex-theme');if(s==='light'){document.body.classList.add('light');document.getElementById('themeIcon').textContent='☾';document.getElementById('themeLabel').textContent='Dark'}})();
</script></body></html>"""


def build_one_pager():
    print("\nWriting GMF One-pager...")
    target = T1 / "APEX-AXLE-FIN-01-GMF-OnePager.html"
    target.write_text(ONE_PAGER, encoding="utf-8")
    print(f"  wrote {target.relative_to(ROOT)}")


# =========================================================================
# 3. SIX-AGENTS ARCHITECTURE (DOCX) - DocInt + Foundry deep dive
# =========================================================================

def build_six_agents():
    print("\nBuilding Six-Agents Architecture...")
    doc = docinit()
    addeyebrow(doc, "APEX · AXLE-FIN-01 · SIX AGENTS · DOC INTELLIGENCE + FOUNDRY")
    p = doc.add_paragraph(); r = p.add_run("Six Agents · Doc Intelligence + Foundry"); setrun(r, bold=True, size=24, color=INK, font=SERIF)
    p2 = doc.add_paragraph(); r2 = p2.add_run("AXLE-FIN-01 architecture · how Document Intelligence and Azure AI Foundry compose into the GMF underwriting flow")
    setrun(r2, bold=True, size=14, color=INK2, italic=True, font=SERIF)
    doc.add_paragraph()
    addbody(doc, "Architectural deep dive on the six agents that compose AXLE-FIN-01 — Ingest, Validate, "
                 "Underwrite, Decide, Approve, Commit — and how they integrate Azure AI Document Intelligence "
                 "(for document extraction) with Azure AI Foundry (for agent runtime, model serving, and "
                 "capability registry). Per-agent specifications, Doc Intelligence model strategy "
                 "(prebuilt + custom), Foundry capability tags, and the AppTrac integration surface.",
            italic=True, color=INK2)
    doc.add_paragraph()
    addkv(doc, [
        ("Document type",     "Architectural Deep Dive · 6 Agents + DocInt + Foundry"),
        ("Service",           "AXLE-FIN-01 Auto Loan Origination & Underwriting"),
        ("Microsoft stack",   "Azure AI Document Intelligence · Azure AI Foundry · Azure OpenAI · Microsoft Teams · Entra · Purview"),
        ("Reference client",  "GM Financial (GMF) · AppTrac integration"),
        ("Audience",          "Lead Engineer · Solution Architect · ML Engineering · GMF IT · CAF MCP team"),
        ("Companion",         "AXLE-FIN-01-Walkthrough.docx · AppTrac-SOR-Integration.docx · GMF-OnePager.html"),
        ("Version / date",    "v1.0 · 27 Apr 2026"),
    ])
    pb(doc)

    # 1. Overview
    addh(doc, "1 · The six agents at a glance", level=1)
    addtable(doc, ["#","Agent","Responsibility","HITL gate","Audit row?"],
             [["1","Ingest",     "Pull docs from AppTrac · Doc Intelligence extraction · normalize to FINML.Application", "ZERO_TOUCH",                       "Yes (capture)"],
              ["2","Validate",   "Cross-doc consistency · fraud signals · authenticity heuristics",                       "ESCALATION (fraud signals)",       "Yes (capture)"],
              ["3","Underwrite", "DTI/PTI/LTV/FICO compute · custom credit policy · pricing tier",                       "ESCALATION (low confidence/policy edge)", "Yes (capture)"],
              ["4","Decide",     "Decision composition · approve/conditional/declined · ECOA narrative",                  "ESCALATION (subprime override · adverse-action)", "Yes (capture)"],
              ["5","Approve",    "HITL gate · Adaptive Card to Underwriter · audit-row HITL log",                         "ALWAYS for material entries",      "Yes (HITL log)"],
              ["6","Commit",     "AppTrac decision-write · dealer notification · adverse-action letter · LEDGER closeout","ZERO_TOUCH",                       "Yes (closeout)"]],
             col_widths=[0.4,1.2,2.7,1.7,0.9])

    # 2. Doc Intelligence strategy
    addh(doc, "2 · Azure AI Document Intelligence strategy", level=1)
    addbody(doc, "Doc Intelligence is the spine of the Ingest agent. Three model categories combine to cover "
                 "the full document landscape at GMF.")

    addh(doc, "2.1 · Prebuilt Doc Intelligence models (immediate use)", level=2)
    addtable(doc, ["Doc type","Prebuilt model","Coverage at GMF"],
             [["State-issued ID (driver license)","prebuilt-idDocument",      "~95% of US states · co-applicant ID where applicable"],
              ["Paystub",                          "prebuilt-payStub.us",      "Most US employer formats · self-employed handled separately"],
              ["W-2 wage statement",               "prebuilt-tax.us.w2",       "Annual filing for prior 2 years · employer-issued"],
              ["1099 forms",                       "prebuilt-tax.us.1099*",    "Self-employed or contract income"],
              ["Tax return forms (1040)",          "prebuilt-tax.us.1040",     "Self-employed income verification"],
              ["Bank statement",                   "prebuilt-bankStatement.us","Recent 2-3 months · cross-verification of income + assets"],
              ["Vehicle title",                    "prebuilt-document (general)+layout", "Title transfer state-by-state · combined with custom fine-tuning"],
              ["Insurance binder",                 "prebuilt-document + layout","Policy verification · stipulation tracking"]],
             col_widths=[2.0,2.0,2.6])

    addh(doc, "2.2 · Custom Doc Intelligence models (per-DMS deal jacket)", level=2)
    addbody(doc, "GMF's dealer network spans four major DMS platforms (CDK Global · Reynolds & Reynolds · "
                 "Dealertrack · RouteOne). Each DMS produces a different deal-jacket format with specific "
                 "field positions, terminology, and embedded forms. Custom models are trained per DMS to "
                 "reach extraction parity with prebuilt-document accuracy.")
    addtable(doc, ["DMS platform","Deal jacket format","Custom model strategy","Wave"],
             [["CDK Global",         "DealerSocket DealCenter format",   "Custom Doc Intelligence model · 250+ training samples", "W1"],
              ["Reynolds & Reynolds","ERA-IGNITE format",                "Custom model · 250+ training samples",                  "W2-S01"],
              ["Dealertrack",        "Cox Automotive credit application","Custom model · 250+ training samples",                  "W2-S01"],
              ["RouteOne",           "RouteOne credit application",      "Custom model · 250+ training samples",                  "W2-S01"]],
             col_widths=[1.8,2.4,2.0,0.8])

    addh(doc, "2.3 · Custom-model lifecycle", level=2)
    addbullets(doc, [
        "Initial training: 250-500 labeled samples per DMS · 2-week training cycle in Foundry · pilot test with shadow extraction",
        "Production deployment: champion-challenger A/B for 14 days · promote to production on accuracy gate (≥ 98% field-level)",
        "Continuous improvement: per-week sample of low-confidence extractions reviewed and added to training set",
        "Drift detection: per-week extraction-confidence distribution monitored · alerts on > 5pp drop",
        "Versioning: semver per custom model · pinned in agent manifest · MAJOR rev requires manifest version-up",
    ])

    addcallout(doc, "Independence-safe model artifact ownership",
               "Custom Doc Intelligence models trained for GMF are GMF artifacts. The training data — "
               "labeled deal-jacket samples from GMF dealers — is GMF's. The model weights are stored in "
               "GMF's Foundry workspace. Deloitte's role is advisory + implementation enablement on the "
               "training pipeline; the resulting model and its lifecycle are GMF-owned.",
               accent=GOLD)

    pb(doc)

    # 3. Foundry agent runtime
    addh(doc, "3 · Azure AI Foundry agent runtime", level=1)
    addbody(doc, "Foundry hosts the agent runtime, the model serving for ML capabilities, and the capability "
                 "registry that makes manifests portable across reference clients. Three Foundry-specific "
                 "patterns:")

    addh(doc, "3.1 · Foundry agent definition (per agent)", level=2)
    addbody(doc, "Each of the six AXLE-FIN-01 agents is registered as a Foundry agent. Sample registration "
                 "snippet for the Ingest agent:")
    addcode(doc, '{\n'
                 '  "agent_id": "axle-fin-01-ingest-agent",\n'
                 '  "version": "1.0.0",\n'
                 '  "foundry_workspace": "gmf-axle-fin-prod",\n'
                 '  "model": {\n'
                 '    "type": "Azure OpenAI",\n'
                 '    "deployment": "gpt-4o-2024-11-20",\n'
                 '    "via_aef_gateway": true\n'
                 '  },\n'
                 '  "tools": [\n'
                 '    "cap:doc-intelligence-extract",\n'
                 '    "cap:apptrac-fetch-application",\n'
                 '    "cap:apptrac-fetch-documents",\n'
                 '    "cap:finml-application-validate",\n'
                 '    "cap:ledger-write"\n'
                 '  ],\n'
                 '  "system_prompt_ref": "prompts/ingest-agent-v1.0.0.md",\n'
                 '  "evaluation_pack": "axle-fin-01-ingest-eval-v1.0",\n'
                 '  "observability": {\n'
                 '    "trace": true,\n'
                 '    "audit_row": true,\n'
                 '    "purview_lineage": true\n'
                 '  }\n'
                 "}")

    addh(doc, "3.2 · Foundry capability registry", level=2)
    addbody(doc, "Sixteen MCP tool capabilities registered for AXLE-FIN-01:")
    addtable(doc, ["Capability","Backing","Used by"],
             [["cap:doc-intelligence-extract",       "Azure AI Document Intelligence (prebuilt + custom models)", "Ingest"],
              ["cap:apptrac-fetch-application",      "AppTrac REST API",                                          "Ingest"],
              ["cap:apptrac-fetch-documents",        "AppTrac document store",                                    "Ingest"],
              ["cap:finml-application-validate",     "FINML.Application schema validator",                        "Ingest · Validate"],
              ["cap:fraud-signal-detect",            "Custom Foundry agent · synthetic-identity heuristics",      "Validate"],
              ["cap:bureau-pull",                    "Equifax/Experian/TransUnion (existing GMF integration)",   "Underwrite"],
              ["cap:credit-policy-resolver",         "Versioned policy corpus · per dealer/tier",                  "Underwrite"],
              ["cap:vehicle-value-lookup",           "MERML.VehicleValue (KBB/Black Book/Manheim feeds)",         "Underwrite"],
              ["cap:dti-pti-ltv-compute",           "Deterministic computation",                                  "Underwrite"],
              ["cap:adverse-action-narrative",       "Azure OpenAI · ECOA-compliant narrative",                   "Decide"],
              ["cap:teams-adaptive-card",            "Microsoft Teams (CAF-shared)",                              "Approve"],
              ["cap:entra-identity-resolve",         "Microsoft Entra (CAF-shared)",                              "Approve"],
              ["cap:apptrac-decision-write",         "AppTrac REST API",                                          "Commit"],
              ["cap:dealer-notification",            "Dealer DMS callback · email · portal",                      "Commit"],
              ["cap:adverse-action-letter-generator","Letter composition · USPS-deliverable format",              "Commit"],
              ["cap:ledger-write",                   "LEDGER (CAF-shared)",                                       "All 6 agents"]],
             col_widths=[2.4,2.4,1.6])

    addh(doc, "3.3 · Foundry evaluation packs", level=2)
    addbody(doc, "Each agent has a Foundry evaluation pack — automated test cases that exercise the agent "
                 "against ground-truth labels. Evaluation packs run continuously; regressions block "
                 "production deployment.")
    addtable(doc, ["Agent","Eval pack","Pass criteria"],
             [["Ingest",    "100 labeled GMF historical applications",  "≥ 99% field-level extraction accuracy on prebuilt; ≥ 98% on custom"],
              ["Validate",  "100 labeled fraud + 100 clean applications", "≥ 95% true positive · ≤ 2% false positive on fraud signals"],
              ["Underwrite","500 labeled GMF historical decisions",     "≥ 92% match to historical decisions on hold-out set"],
              ["Decide",    "200 ECOA-narrative ground-truth pairs",     "≥ 99% pass on ECOA-compliance lint"],
              ["Approve",   "(orchestration only · routing tests)",     "Routing correctness ≥ 99% on materiality gate"],
              ["Commit",    "Idempotency tests · partial-failure tests","Compensating-action correctness ≥ 100%"]],
             col_widths=[1.2,3.0,2.4])

    pb(doc)

    # 4. Per-agent details (condensed)
    addh(doc, "4 · Per-agent details (condensed)", level=1)
    addbody(doc, "Per-agent specifications. Full system prompts, tool registrations, and HITL gating live in "
                 "the Foundry agent registry; this section summarizes the contract.")

    addh(doc, "4.1 · Ingest", level=2)
    addbody(doc, "Pulls application + documents from AppTrac, runs Document Intelligence extraction, "
                 "normalizes to FINML.Application + FINML.Document records. Emits per-document confidence "
                 "scores. Watermark trigger: AppTrac webhook on application submission.")
    addh(doc, "4.2 · Validate", level=2)
    addbody(doc, "Cross-document consistency (ID name vs. paystub name vs. application name), authenticity "
                 "heuristics (paystub format consistency, employer ID validation), synthetic-identity "
                 "detection (SSN-vs-name-vs-DOB triangulation against bureau pull). Routes fraud signals "
                 "to Fraud Specialist via Adaptive Card.")
    addh(doc, "4.3 · Underwrite", level=2)
    addbody(doc, "Bureau pull orchestration (Equifax/Experian/TransUnion via existing GMF integration). "
                 "DTI/PTI/LTV computation. Custom credit policy applied per dealer + tier. Vehicle value "
                 "lookup via MERML.VehicleValue (KBB/Black Book/Manheim feeds). Outputs underwriting "
                 "decision recommendation with confidence score and policy citations.")
    addh(doc, "4.4 · Decide", level=2)
    addbody(doc, "Composes the decision: approve / conditional / declined, with stipulations (proof of "
                 "insurance, GAP, etc.) and pricing tier. For adverse-action decisions, generates an "
                 "ECOA-compliant narrative via Azure OpenAI with strict prompt scaffolding and "
                 "post-generation lint (no protected-class language, only ECOA reason codes from approved "
                 "list, plain-English explanation per 12 CFR Part 1002 Appendix C).")
    addh(doc, "4.5 · Approve (HITL)", level=2)
    addbody(doc, "Routes material decisions per the commercial envelope (subprime · low-confidence · "
                 "policy edge cases · adverse-action narratives) to the Underwriter via Microsoft Teams "
                 "Adaptive Card. Card carries: application summary, all extracted document data, decision "
                 "recommendation with reasoning, ECOA narrative for review, three buttons (Approve / Modify / "
                 "Decline). Modifications write back per-field deltas.")
    addh(doc, "4.6 · Commit", level=2)
    addbody(doc, "Post-approval execution. Three sub-actions in parallel: AppTrac decision-write (REST API), "
                 "dealer notification (DMS callback or email), adverse-action letter generation (where "
                 "applicable, USPS-deliverable format). Writes 14-field closeout audit row to LEDGER "
                 "with all upstream version stamps and ECOA reason codes.")

    pb(doc)

    # 5. ECOA + compliance
    addh(doc, "5 · ECOA + regulatory compliance posture", level=1)
    addh(doc, "5.1 · ECOA (Equal Credit Opportunity Act) hooks", level=2)
    addbullets(doc, [
        "Adverse-action narrative must use only ECOA reason codes from approved list (12 CFR Part 1002 Appendix C)",
        "Decide agent prompts include strict scaffolding to prevent protected-class language",
        "Post-generation lint catches any reason-code drift before commit",
        "Continuous disparate-impact testing surface in LEDGER · per-cohort decision-distribution analysis",
        "Per-decision audit-row captures: ECOA reason codes (if applicable), narrative version, model version, policy version",
    ])
    addh(doc, "5.2 · TILA (Truth in Lending Act) integration", level=2)
    addbody(doc, "TILA disclosures (APR · finance charge · payment schedule · total of payments) are "
                 "computed by Underwrite agent and written to AppTrac at decision time. Audit row captures "
                 "TILA-disclosure-version stamp.")
    addh(doc, "5.3 · FCRA (Fair Credit Reporting Act) compliance", level=2)
    addbody(doc, "Bureau pull authorization captured at AppTrac submission time. Pull is recorded in the "
                 "audit row with timestamp, bureau, and authorization reference. Permissible-purpose code "
                 "stamped on every pull.")
    addh(doc, "5.4 · GLBA + privacy", level=2)
    addbody(doc, "Heavy PII in this scenario (SSN, DOB, income, employer). Microsoft Purview classifies "
                 "every field at Bronze landing; tokenization for downstream serving where appropriate; "
                 "consent state captured at application submission; access logs to Purview for every PII "
                 "touch.")

    pb(doc)

    # 6. AppTrac integration
    addh(doc, "6 · AppTrac integration touchpoints", level=1)
    addbody(doc, "Detailed integration specs in the AppTrac SOR Integration document. Summary here.")
    addtable(doc, ["AppTrac surface","Integration pattern","Used by"],
             [["Application submission webhook",      "AppTrac → Eventstream",                "Trigger Ingest agent"],
              ["Application detail REST API",         "Pull on agent invocation",             "Ingest"],
              ["Document store REST API",             "Pull document binaries on Ingest",     "Ingest"],
              ["Bureau pull orchestration",            "Existing AppTrac integration",          "Underwrite (via cap:bureau-pull)"],
              ["Decision write-back REST API",         "Commit-time write",                    "Commit"],
              ["Stipulation tracking REST API",        "Commit-time write + post-funding",     "Commit"],
              ["Dealer notification callback",         "AppTrac → DMS",                        "Commit (delegate)"]],
             col_widths=[2.4,2.4,1.6])

    addbody(doc, "Independence posture preserved.", italic=True, color=MUTE, size=9)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— End of Six-Agents Architecture —"); setrun(r, italic=True, size=10, color=MUTE, font=SERIF)
    docsave(doc, T2 / "APEX-AXLE-FIN-01-Six-Agents-Architecture.docx")


# =========================================================================
# 4. APPTRAC SOR INTEGRATION (DOCX)
# =========================================================================

def build_apptrac_sor():
    print("\nBuilding AppTrac SOR Integration...")
    doc = docinit()
    addeyebrow(doc, "APEX · AXLE-FIN-01 · APPTRAC · SOR INTEGRATION")
    p = doc.add_paragraph(); r = p.add_run("AppTrac SOR Integration"); setrun(r, bold=True, size=24, color=INK, font=SERIF)
    p2 = doc.add_paragraph(); r2 = p2.add_run("Bronze landing · FINML canonical · AppTrac as the system of record at GMF")
    setrun(r2, bold=True, size=14, color=INK2, italic=True, font=SERIF)
    doc.add_paragraph()
    addbody(doc, "Integration specification for AppTrac as the system of record for AXLE-FIN-01 at GMF. "
                 "Defines the Bronze tables that capture AppTrac state, the Silver canonical FINML schema, "
                 "the touchpoints between APEX agents and AppTrac (read + write), the integration patterns "
                 "(webhook + polling + REST), and the failure-handling story.",
            italic=True, color=INK2)
    doc.add_paragraph()
    addkv(doc, [
        ("Document type",     "SOR Integration Specification · Bronze + Silver mapping"),
        ("Service",           "AXLE-FIN-01 Auto Loan Origination & Underwriting"),
        ("Reference client",  "GM Financial · AppTrac platform"),
        ("Audience",          "Lead Engineer · Solution Architect · GMF Application Operations · Data Engineering"),
        ("Companion",         "AXLE-FIN-01-Walkthrough.docx · Six-Agents-Architecture.docx · GMF-OnePager.html"),
        ("Version / date",    "v1.0 · 27 Apr 2026"),
    ])
    pb(doc)

    # 1. AppTrac as the SOR
    addh(doc, "1 · AppTrac as the system of record", level=1)
    addbody(doc, "AppTrac is GMF's loan origination platform — it is, and remains, the system of record for "
                 "loan applications, decisions, stipulations, and funding state. AXLE-FIN-01 does not "
                 "replace AppTrac. AXLE-FIN-01 reads from AppTrac (documents, application data, bureau "
                 "results), reasons over the data via the six-agent orchestration, and writes back the "
                 "decision + stipulations + adverse-action artifacts to AppTrac at commit time.")
    addcallout(doc, "Architectural commitment",
               "AppTrac remains the SOR. APEX wraps it with intelligence; APEX does not replace it. This "
               "preserves GMF's existing dealer integrations (DMS callbacks, funding pipelines, servicing "
               "handoffs) and keeps the regulatory artifact trail rooted in GMF's authoritative platform.")

    addh(doc, "2 · Bronze landing tables (10)", level=1)
    addbody(doc, "Ten Bronze tables capture AppTrac state for AXLE-FIN-01. Each follows the APEX Bronze "
                 "convention: append-only, ingest-metadata mandatory (source_system · source_record_id · "
                 "ingest_ts · batch_id · _is_deleted), partitioned by ingest_date, Z-ordered on join keys.")
    addtable(doc, ["Bronze table","What it captures","Source pattern","Refresh"],
             [["BRZ.GMF.APPLICATION",       "Application header · buyer info · vehicle · structure",       "AppTrac webhook + REST",     "Streaming"],
              ["BRZ.GMF.DOCUMENT",          "Document references + Doc Intelligence extraction outputs",   "AppTrac doc store + DocInt", "Streaming"],
              ["BRZ.GMF.BUREAU_PULL",       "Credit bureau pulls · Equifax/Experian/TransUnion",          "Bureau APIs via AppTrac",     "Streaming"],
              ["BRZ.GMF.DEALER",            "Dealer master · DMS · tier · scorecard",                     "AppTrac dealer master",       "Daily snapshot"],
              ["BRZ.GMF.UNDERWRITER",       "Underwriter directory · roles · authority levels",           "AppTrac UW master",           "Weekly snapshot"],
              ["BRZ.GMF.CREDIT_POLICY",     "Versioned credit policy corpus · per dealer/tier",            "Policy-mgmt system",          "On-publish"],
              ["BRZ.GMF.VEHICLE_VALUATION", "MERML.VehicleValue feeds · KBB · Black Book · Manheim",      "External feeds",              "Daily"],
              ["BRZ.GMF.DECISION",          "Historical underwriting decisions · pre-AXLE + post-AXLE",   "AppTrac decision history",    "Streaming + nightly"],
              ["BRZ.GMF.STIPULATION",       "Decision stipulations + clearance state · pre/post-funding", "AppTrac stipulation tracking","Streaming"],
              ["BRZ.GMF.FUNDING",           "Post-decision funding state · booked/cancelled/expired",     "AppTrac funding pipeline",   "Daily"]],
             col_widths=[2.0,2.6,1.6,1.0])

    addh(doc, "2.1 · Volume estimates", level=2)
    addtable(doc, ["Table","Daily volume","Notes"],
             [["BRZ.GMF.APPLICATION",       "~12K rows/day",    "~3M/yr ÷ 250 working days"],
              ["BRZ.GMF.DOCUMENT",          "~120K rows/day",   "~10 docs/app on average"],
              ["BRZ.GMF.BUREAU_PULL",       "~30K rows/day",    "~2-3 bureau pulls per app at submission"],
              ["BRZ.GMF.DECISION",          "~12K rows/day",    "1 per app at decision time"],
              ["BRZ.GMF.STIPULATION",       "~30K rows/day",    "Avg 2-3 stipulations per approved app"],
              ["BRZ.GMF.FUNDING",           "~10K rows/day",    "~85% funding rate of approved apps"]],
             col_widths=[2.4,1.6,3.0])

    pb(doc)

    addh(doc, "3 · Silver canonical layer (FINML)", level=1)
    addbody(doc, "Silver canonical layer for financial-services scenarios. Five FINML shapes that AXLE-FIN-01 "
                 "agents read against.")
    addtable(doc, ["Silver shape","Composition","Primary key","Update pattern"],
             [["FINML.Application",   "BRZ.GMF.APPLICATION + DIM.Dealer + DIM.Vehicle + BRZ.GMF.BUREAU_PULL", "(application_id)",         "Streaming · per-application state"],
              ["FINML.Document",      "BRZ.GMF.DOCUMENT + Doc Intelligence extraction structured fields",     "(document_id)",            "Streaming · append on extraction"],
              ["FINML.Decision",      "BRZ.GMF.DECISION + audit-row trace_id correlation",                   "(decision_id)",            "Streaming · per-decision"],
              ["FINML.PolicyApplication","BRZ.GMF.CREDIT_POLICY (current effective) + decision-time application","(decision_id, policy_clause)","Append per decision"],
              ["FINML.Outcome",        "BRZ.GMF.FUNDING + BRZ.GMF.STIPULATION + post-funding signals",       "(application_id, outcome_event)","Append on outcome events"]],
             col_widths=[1.7,2.7,1.3,1.5])

    addh(doc, "3.1 · DIM dimensions", level=2)
    addtable(doc, ["Dimension","Composition","SCD"],
             [["DIM.Dealer",      "BRZ.GMF.DEALER · history preserved (DMS migrations, scorecard changes)", "Type 2"],
              ["DIM.Underwriter", "BRZ.GMF.UNDERWRITER · role and authority history",                       "Type 2"],
              ["DIM.Vehicle",     "Vehicle attributes · year/make/model/trim/VIN-decoded",                  "Type 1"],
              ["DIM.Bureau",      "Static reference · 3 bureaus + score products",                         "Type 0"]],
             col_widths=[1.6,4.0,1.0])

    pb(doc)

    addh(doc, "4 · Integration touchpoints (read)", level=1)
    addtable(doc, ["Touchpoint","Trigger","API/pattern","Used by agent"],
             [["Application submission",   "AppTrac webhook",           "Eventstream consumer",  "Ingest (orchestrator)"],
              ["Application detail",       "Agent invocation",          "REST GET /applications/:id",   "Ingest"],
              ["Document binaries",         "Agent invocation",          "REST GET /applications/:id/documents", "Ingest"],
              ["Bureau pull results",       "AppTrac → Bronze CDC",      "AppTrac CDC stream",   "Underwrite"],
              ["Dealer master",             "Daily snapshot",            "REST batch export",     "Validate · Underwrite · Approve"],
              ["Credit policy current",     "On-publish",                "Policy webhook + REST", "Underwrite"]],
             col_widths=[2.0,1.8,2.2,1.6])

    addh(doc, "5 · Integration touchpoints (write)", level=1)
    addtable(doc, ["Touchpoint","Trigger","API/pattern","Idempotency","Used by agent"],
             [["Decision write",        "Commit-time",   "REST POST /applications/:id/decision",     "decision_id",       "Commit"],
              ["Stipulation creation",   "Commit-time",   "REST POST /applications/:id/stipulations","stipulation_id × app","Commit"],
              ["Adverse-action narrative","Commit-time",   "REST POST /applications/:id/aa-letter",  "decision_id",       "Commit"],
              ["Dealer notification",     "Commit-time",   "AppTrac → DMS callback (delegated)",     "decision_id",       "Commit (indirect)"],
              ["Audit-row reference",     "Commit-time",   "REST POST /applications/:id/audit-ref",  "trace_id",          "Commit"]],
             col_widths=[2.0,1.4,2.4,1.4,1.0])

    addh(doc, "5.1 · Compensating actions on partial failure", level=2)
    addbullets(doc, [
        "AppTrac decision-write failure: retry 3× with exponential backoff · circuit-break after 5 consecutive failures · escalate to AppTrac on-call",
        "If decision-write succeeds but stipulation-creation fails: queue locally · retry · do not roll back decision",
        "If decision-write succeeds but adverse-action letter fails: alert ECOA Program Lead · letter regenerated and sent within 30 days regulatory deadline",
        "If audit-row reference write fails: write succeeds locally to LEDGER first · AppTrac reference is reconciliation metadata · backfill on AppTrac recovery",
    ])

    pb(doc)

    addh(doc, "6 · Failure handling and observability", level=1)
    addh(doc, "6.1 · AppTrac availability dependency", level=2)
    addbody(doc, "AXLE-FIN-01's dependency on AppTrac is total — AppTrac down means the service is down. "
                 "Mitigation:")
    addbullets(doc, [
        "AppTrac SLA visibility · GMF Application Operations is the SLA owner",
        "Webhook delivery retry · AppTrac retries failed webhook deliveries automatically",
        "Read failures: agent operations queue locally up to 4 hours · resume on AppTrac recovery",
        "Write failures: per-write idempotency key allows safe retry on AppTrac recovery",
        "Cross-system reconciliation: nightly job validates LEDGER decisions against AppTrac state · alerts on drift",
    ])

    addh(doc, "6.2 · Observability", level=2)
    addbullets(doc, [
        "OpenTelemetry traces propagate from AppTrac webhook through agent chain to AppTrac decision-write",
        "Microsoft Purview lineage: AppTrac → Bronze → Silver → agent → AppTrac decision-write",
        "Per-AppTrac-touchpoint latency monitored · alerts on > 2× baseline",
        "Per-decision reconciliation: trace_id present in AppTrac audit-ref + LEDGER · daily cross-check",
    ])

    pb(doc)

    addh(doc, "7 · PII handling and classification", level=1)
    addbody(doc, "Heavy PII in this scenario (SSN, DOB, income, employer name, account numbers). Microsoft "
                 "Purview classifies every field at Bronze landing per the Privacy & Data Governance Spec.")
    addtable(doc, ["Field category","Sensitivity","Handling"],
             [["SSN",                          "Restricted · PII · ECOA-relevant",  "Tokenized at ingest · clear-text only at agent decision time"],
              ["DOB",                          "Restricted · PII",                  "Stored encrypted at rest · decrypted at agent decision time"],
              ["Driver license number",        "Restricted · PII",                  "Tokenized at ingest"],
              ["Income figures + employer",    "Confidential · ECOA-relevant",      "Encrypted at rest · ECOA reason-code-mapped"],
              ["Bureau score + tradelines",    "Confidential · FCRA-protected",     "Encrypted at rest · FCRA permissible-purpose stamp"],
              ["Bank account numbers",          "Restricted · PII · GLBA-protected","Tokenized at ingest · clear-text only at funding pipeline"],
              ["Vehicle VIN",                   "Internal",                          "Standard handling"],
              ["Decision audit row",            "Audit-Critical",                    "Immutable · indefinite retention · ECOA-defensible"]],
             col_widths=[2.4,2.0,2.2])

    addbody(doc, "Independence posture preserved.", italic=True, color=MUTE, size=9)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— End of AppTrac SOR Integration —"); setrun(r, italic=True, size=10, color=MUTE, font=SERIF)
    docsave(doc, T0 / "APEX-AXLE-FIN-01-AppTrac-SOR-Integration.docx")


# =========================================================================
# EXECUTE
# =========================================================================

print(f"Output root: {ROOT}\n")
build_walkthrough()
build_one_pager()
build_six_agents()
build_apptrac_sor()
print("\nGMF AppTrac documentation set complete.")
